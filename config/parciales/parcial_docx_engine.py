# -*- coding: utf-8 -*-
"""Motor DOCX para parciales UNIAJC 2026-2 (estudiante + solución)."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xA0, 0x20, 0x30)
FONT = "Calibri"


def _shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_run(run, *, size=11, bold=False, color=GRIS, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_para(
    doc,
    text: str,
    *,
    size=11,
    bold=False,
    color=GRIS,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_after=6,
    space_before=0,
    shade: str | None = None,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if shade:
        _shade(p, shade)
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, color=color)
    return p


def add_rich(doc, parts, *, size=11, space_after=6, space_before=0):
    """parts: list of (text, bold, color?)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for item in parts:
        if len(item) == 2:
            text, bold = item
            color = GRIS
        else:
            text, bold, color = item
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold, color=color)
    return p


def banda(doc, text: str):
    p = add_para(
        doc,
        text,
        size=13,
        bold=True,
        color=BLANCO,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=10,
        space_before=12,
        shade="095292",
    )
    return p


def h2(doc, text: str):
    return add_para(doc, text, size=12, bold=True, color=AZUL, space_before=12, space_after=6)


def h3(doc, text: str):
    return add_para(doc, text, size=11, bold=True, color=CIAN, space_before=8, space_after=4)


def body(doc, text: str, *, space_after=4):
    return add_para(doc, text, size=11, color=GRIS, space_after=space_after)


def mono(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _shade(p, "F2F2F2")
    r = p.add_run(text)
    _set_run(r, size=9.5, bold=False, color=GRIS, name="Consolas")
    return p


def bullets(doc, items: list[str], *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        _set_run(r, size=size, color=GRIS)


def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)


def portada(doc, meta: dict[str, Any], *, es_solucion: bool):
    add_para(
        doc,
        "Institución Universitaria Antonio José Camacho — UNIAJC",
        size=10,
        bold=True,
        color=AZUL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "Facultad de Ingenierías · Programa de Ingeniería de Sistemas",
        size=10,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    titulo = meta["titulo_parcial"]
    if es_solucion:
        titulo = f"{titulo} — SOLUCIÓN / CLAVE"
    banda(doc, f"  {titulo}")
    add_para(doc, meta["asignatura"], size=14, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)
    add_para(
        doc,
        f"Código {meta['codigo']}  ·  Grupo {meta['grupo']}  ·  Periodo {meta['periodo']}",
        size=11,
        bold=True,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        f"Parcial {meta['n']} — Corte {meta['corte']}  ·  Valor en el corte: {meta['valor_corte']}",
        size=11,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        # meta['clase'] = numero de SESION del calendario 2026-2 (13 sesiones).
        # Las «Clase N» del material (1..15) se listan en meta['temas'].
        # La modalidad NO es parametro porque no varia: `regla_parciales` de
        # config/calendario/semestre_2026_2.json dice «Modalidad VIRTUAL en los 4 cursos:
        # los parciales se presentan VIRTUAL SINCRONO, no en el campus». Decia
        # «Presencial (síncrono)» en los 12 parciales y sus 12 soluciones, que es el
        # error factual que puede hacer que un estudiante viaje al campus el dia del
        # examen.
        f"Fecha: {meta['fecha']}  ·  Sesión {meta['clase']} de 13  ·  "
        "Modalidad: virtual síncrona por Google Meet",
        size=11,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        f"Horario del bloque: {meta['horario']}  ·  Tiempo sugerido: {meta['tiempo']}",
        size=11,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Docente: Julian Andres Castaño Espinosa  ·  julianacastano@profesores.uniajc.edu.co",
        size=10,
        color=GRIS,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    if es_solucion:
        add_para(
            doc,
            "DOCUMENTO DOCENTE — No distribuir a estudiantes. Incluye clave y rúbrica breve.",
            size=10,
            bold=True,
            color=ROJO,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            shade="FBE4E4",
            space_after=10,
        )
    else:
        add_para(
            doc,
            "Versión estudiante — No incluir material de apoyo no autorizado.",
            size=10,
            bold=True,
            color=AZUL,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            shade="E8F4FA",
            space_after=10,
        )

    # Datos estudiante
    if not es_solucion:
        h2(doc, "Identificación del estudiante")
        for label in [
            "Nombre completo: _______________________________________________",
            "Documento / código: ______________________    Firma: ______________________",
        ]:
            body(doc, label, space_after=6)

    h2(doc, "Temas evaluados en este parcial (Plan de curso 2026-2)")
    body(doc, meta["cobertura"])
    body(doc, "Lista explícita de temas evaluados (solo estos; no se preguntan temas de otros cortes):", space_after=4)
    bullets(doc, meta["temas"])

    h2(doc, "Instrucciones generales")
    # Las dos primeras vinetas asumian examen en papel en un salon: «Escriba con letra
    # clara» y «No se permite el uso de dispositivos electronicos». El parcial se
    # presenta por Meet y se responde en este mismo archivo, o sea EN un dispositivo
    # electronico: la regla anterior prohibia el unico medio de entrega que hay.
    bullets(
        doc,
        [
            "Lea con atención cada sección antes de responder. Responda en este mismo "
            "documento, sobre las líneas marcadas, y entréguelo por el medio que el "
            "docente indique al abrir la sesión.",
            "El parcial se califica sobre 100 puntos. Nota sobre 5.0 = (puntos obtenidos / 100) × 5.0.",
            "Es individual: no se permite consultar con otras personas ni compartir "
            "respuestas. El material de apoyo autorizado es el que el docente indique al "
            "abrir la sesión.",
            "Justifique las respuestas de desarrollo. En ejercicios de código o SQL, la claridad y la corrección cuentan.",
            "Este parcial evalúa únicamente los temas del corte indicado. El desglose del Acuerdo (30%/30%/40%) se aplica a nivel de corte, no dentro de este documento.",
            f"Duración sugerida: {meta['tiempo']} dentro del bloque de 120 minutos.",
        ],
    )

    h2(doc, "Distribución de puntaje")
    for s in meta["secciones_resumen"]:
        body(doc, f"• {s}")
    body(doc, "Total: 100 puntos  →  Nota = puntos/20 sobre 5.0", space_after=10)


def por_que_opciones(doc, item: dict[str, Any]):
    """Justificacion de TODAS las opciones, no solo de la correcta (solo en la solucion).

    Lo exige el checklist del repo para preguntas cerradas. Sin esto, el docente puede
    senalar la «b» pero no responder por que la «c» no era, que es justo lo que el
    estudiante pregunta cuando reclama la nota.
    """
    pq = item.get("por_que")
    if not pq:
        return
    h3(doc, "Por qué cada opción")
    for etiqueta, razon in pq.items():
        add_rich(doc, [(f"    {etiqueta}  ", True), (razon, False)], size=10, space_after=2)


def codigo_solucion(doc, item: dict[str, Any]):
    """El codigo/SQL resuelto, docente-only.

    Va aparte de `codigo`, que se imprime tambien en la version del estudiante: si la
    respuesta se pusiera ahi, el parcial traeria la solucion impresa.
    """
    cod = item.get("solucion_codigo")
    if not cod:
        return
    h3(doc, "Respuesta que corre tal cual")
    for line in cod.splitlines() or [""]:
        mono(doc, line if line else " ")


def errores_frecuentes(doc, item: dict[str, Any]):
    """Que llega mal y que hacer con ello: decisiones de calificacion tomadas antes."""
    errs = item.get("errores")
    if not errs:
        return
    h3(doc, "Errores frecuentes y qué hacer")
    for e in errs:
        body(doc, f"• {e}", space_after=2)


def render_item(doc, item: dict[str, Any], *, es_solucion: bool):
    tipo = item["tipo"]
    if tipo == "enunciado":
        body(doc, item["texto"], space_after=6)
        return
    if tipo == "mcq":
        body(doc, f"{item['id']}. ({item['pts']} pts) {item['pregunta']}", space_after=2)
        for opt in item["opciones"]:
            body(doc, f"    {opt}", space_after=1)
        if es_solucion:
            add_para(doc, f"    → Clave: {item['clave']}", size=10, bold=True, color=AZUL, space_after=2)
            por_que_opciones(doc, item)
            errores_frecuentes(doc, item)
            if item.get("nota"):
                body(doc, f"    Nota docente: {item['nota']}", space_after=6)
        else:
            body(doc, "", space_after=4)
        return
    if tipo == "vf":
        body(doc, f"{item['id']}. ({item['pts']} pts) {item['enunciado']}", space_after=1)
        if es_solucion:
            add_para(
                doc,
                f"    → Clave: {item['clave']} — {item.get('justificacion', '')}",
                size=10,
                bold=True,
                color=AZUL,
                space_after=4,
            )
            errores_frecuentes(doc, item)
        else:
            body(doc, "    Respuesta (V / F): ________    Justificación breve: _______________________________", space_after=6)
        return
    if tipo == "match":
        body(doc, f"{item['id']}. ({item['pts']} pts) {item['instruccion']}", space_after=2)
        body(doc, "Columna A:", space_after=1)
        for a in item["col_a"]:
            body(doc, f"    {a}", space_after=1)
        body(doc, "Columna B:", space_after=1)
        for b in item["col_b"]:
            body(doc, f"    {b}", space_after=1)
        if es_solucion:
            add_para(doc, f"    → Emparejamiento: {item['clave']}", size=10, bold=True, color=AZUL, space_after=4)
            por_que_opciones(doc, item)
            errores_frecuentes(doc, item)
        else:
            # El marcador decia «(ej. 1-c, 2-a…)», y «1-c» es la primera pareja de la clave
            # real del emparejamiento en el Parcial 2 de Programacion II y en el Parcial 1
            # de Seminario: el ejemplo del formato regalaba una de las cuatro parejas. Se
            # deja solo el formato, que no puede colisionar con ninguna clave.
            body(doc, "    Respuestas (formato 1-x, 2-x, 3-x, 4-x): _____________________________________", space_after=6)
        return
    if tipo == "desarrollo":
        body(doc, f"{item['id']}. ({item['pts']} pts) {item['enunciado']}", space_after=2)
        if item.get("subitems"):
            for s in item["subitems"]:
                body(doc, f"    {s}", space_after=1)
        if item.get("codigo"):
            for line in item["codigo"].splitlines() or [""]:
                mono(doc, line if line else " ")
        if es_solucion:
            h3(doc, "Respuesta esperada / rúbrica")
            for line in item["solucion"]:
                body(doc, f"• {line}", space_after=2)
            codigo_solucion(doc, item)
            errores_frecuentes(doc, item)
        else:
            for _ in range(item.get("lineas", 4)):
                body(doc, "_" * 78, space_after=2)
        return
    if tipo == "practica":
        body(doc, f"{item['id']}. ({item['pts']} pts) {item['enunciado']}", space_after=2)
        if item.get("contexto"):
            add_para(doc, item["contexto"], size=10, color=GRIS, shade="FFF8D6", space_after=4)
        if item.get("requerimientos"):
            bullets(doc, item["requerimientos"])
        if item.get("codigo"):
            for line in item["codigo"].splitlines() or [""]:
                mono(doc, line if line else " ")
        if es_solucion:
            h3(doc, "Solución / rúbrica breve")
            for line in item["solucion"]:
                body(doc, f"• {line}", space_after=2)
            codigo_solucion(doc, item)
            errores_frecuentes(doc, item)
        else:
            for _ in range(item.get("lineas", 6)):
                body(doc, "_" * 78, space_after=2)
        return


def build_docx(meta: dict[str, Any], secciones: list[dict[str, Any]], out: Path, *, es_solucion: bool):
    doc = Document()
    set_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = GRIS

    portada(doc, meta, es_solucion=es_solucion)

    for sec in secciones:
        banda(doc, f"  {sec['titulo']}  ({sec['pts']} puntos)")
        if sec.get("intro"):
            body(doc, sec["intro"], space_after=6)
        for item in sec["items"]:
            render_item(doc, item, es_solucion=es_solucion)

    if not es_solucion:
        h2(doc, "Espacio de borrador (opcional)")
        for _ in range(5):
            body(doc, "_" * 78, space_after=4)

    add_para(
        doc,
        "Fin del parcial — UNIAJC 2026-2 · Julian Andres Castaño · julianacastano@profesores.uniajc.edu.co",
        size=9,
        color=CIAN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
