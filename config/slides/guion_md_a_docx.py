# -*- coding: utf-8 -*-
"""
Convierte un GUIÓN DOCENTE en Markdown a un archivo .docx bien formateado:
- Comandos de consola: `inline` en monoespaciado y bloques ``` ``` como caja
  sombreada monoespaciada (aspecto de terminal).
- Tablas Markdown -> tablas de Word (encabezado gris, texto blanco).
- Encabezados con jerarquía; **negrita** e `inline code` respetados.
- Líneas de 📸 Pantallazo -> caja de "marcador de captura"; si hay una captura
  real que coincide, se EMBEBE la imagen. Las capturas viven POR CURSO en
  '<curso>/Kit docente/Clase N/Capturas/' (misma carpeta del guion + /Capturas).
  Si la imagen NO existe, en vez de una caja vacía se imprime el paso a paso para
  producirla: la del token '[[captura: x.png | receta: ...]]' si viene escrita, o
  una genérica con el nombre y la ruta exactos con que hay que guardarla.

Uso:  python guion_md_a_docx.py "ruta/al/guion.md"  [--out "ruta.docx"]
"""
import sys, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1A, 0x1C, 0x1D)
GRAY = RGBColor(0x4A, 0x4A, 0x49)
ORANGE = RGBColor(0xC2, 0x3E, 0x12)   # naranja un poco más oscuro para texto sobre blanco
CODE_BG = "F2F2F3"
HEAD_BG = "3A3A3C"
ALT_BG = "ECECEE"
MONO = "Consolas"

# Las capturas viven junto al guion en Kit docente: <curso>/Kit docente/Clase N/Capturas/.
# convert() fija CAPTURAS a esa carpeta según la ruta del guion.
CAPTURAS = ""

def _shade(el, fill):
    pPr = el.get_or_add_pPr() if el.tag.endswith('}p') else el.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def shade_par(p, fill): _shade(p._p, fill)
def shade_cell(c, fill):
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _runs_codigo(p, text, size, color, bold):
    """Emite los `código` de un tramo que ya sabe si va en negrita o no."""
    for part in re.split(r'(`[^`]+`)', text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith('`') and part.endswith('`'):
            r.text = part[1:-1]; r.font.name = MONO; r.font.size = Pt(size - 0.5)
            r.font.color.rgb = RGBColor(0x11, 0x11, 0x11); r.bold = bold
        else:
            r.text = part; r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color


def add_inline(p, text, base_size=11, base_color=GRAY, base_bold=False):
    """Parseo inline: **negrita**, @@negrita@@ (convención usada en datos de taller) y `código`.

    La negrita se resuelve primero y el `código` se busca DENTRO de ella, porque van
    anidados: con una sola alternancia, el `INSERT` de «**5 pts — el `INSERT` del caso
    valido**» quedaba dentro del tramo de negrita y salia con los acentos graves
    impresos. Pasaba en cada linea del desglose de puntos de las soluciones docentes.
    """
    for part in re.split(r'(\*\*.*?\*\*|@@.*?@@)', text):
        if not part:
            continue
        if (part.startswith('**') and part.endswith('**')) or (part.startswith('@@') and part.endswith('@@')):
            _runs_codigo(p, part[2:-2], base_size, base_color, True)
        else:
            _runs_codigo(p, part, base_size, base_color, base_bold)

def find_capture(caption):
    c = caption.lower()
    base = CAPTURAS
    if not base or not os.path.isdir(base):
        return None
    def pick(sub):
        for f in sorted(os.listdir(base)):
            if sub in f.lower() and f.lower().endswith(".png"):
                return os.path.join(base, f)
        return None
    if "subnet" in c or "subred" in c:
        return pick("subnetting")
    if "jslinux" in c or "webminal" in c or ("terminal" in c and "linux" in c) or "distrosea" in c or "copy.sh" in c:
        return pick("jslinux")
    # Acceso/login a examlab: estricto (evita falsos positivos como "tecnologías de acceso").
    if re.search(r'(inicia\w*\s+sesi|inicio de sesi|iniciar sesi|\blogin\b|autentic|'
                 r'(acces\w*|acceder|entrar|ingres\w*|iniciar)\s+(a\s+)?examlab|pantalla de acceso)', c):
        return pick("login-form") or pick("login")
    return None

def _capture_token(text):
    """Extrae un token explícito de la línea de pantallazo.

    Sintaxis aceptada:
        [[captura: archivo.png]]
        [[captura: archivo.png | receta: 1) abre X  2) ejecuta Y  3) recorta ...]]

    La parte `receta:` es opcional y se usa cuando el PNG todavía no existe: en vez
    de dejar una caja vacía que nadie llena, el .docx imprime el paso a paso para
    producirla. Devuelve (texto_sin_token, nombre_archivo|None, receta|None)."""
    m = re.search(r'\[\[\s*captura:\s*([^\]]+?)\s*\]\]', text, re.I)
    if not m:
        return text, None, None
    cuerpo = m.group(1).strip()
    receta = None
    if "|" in cuerpo:
        fn, resto = cuerpo.split("|", 1)
        fn = fn.strip()
        resto = resto.strip()
        if re.match(r'(?i)receta\s*:', resto):
            receta = re.sub(r'(?i)^receta\s*:\s*', '', resto).strip() or None
    else:
        fn = cuerpo
    if not fn.lower().endswith(".png"):
        fn += ".png"
    limpio = re.sub(r'\[\[\s*captura:[^\]]+\]\]', '', text, flags=re.I).strip()
    return limpio, fn, receta


def _ruta_capturas_legible(nombre=None):
    """`Kit docente/Clase N/Capturas/archivo.png` — ruta que el docente reconoce.

    CAPTURAS es una ruta absoluta de esta máquina; en el documento sirve más la
    ruta relativa al curso, que es como está escrita en el resto del material."""
    partes = [x for x in os.path.normpath(CAPTURAS or "").split(os.sep) if x]
    cola = "/".join(partes[-3:]) if len(partes) >= 3 else (CAPTURAS or "Capturas")
    return cola + "/" + nombre if nombre else cola + "/"


def _receta_generica(caption, explicit):
    """Paso a paso por defecto para producir una captura que todavía no existe.

    Existe porque el material referencia imágenes que nadie alcanzó a tomar y la
    caja «inserta aquí la captura» no le decía al docente ni qué abrir ni con qué
    nombre guardar; el resultado era que la caja seguía vacía el día de la clase."""
    destino = _ruta_capturas_legible(explicit or "<nombre-corto-con-guiones>.png")
    return (
        "1) Abre la herramienta del bloque y reproduce exactamente los pasos "
        "descritos arriba, con el mismo dominio del ejemplo.  "
        "2) Captura solo la ventana útil (no el escritorio completo) en el momento "
        "en que se ve el resultado: " + (caption or "el resultado del paso") + ".  "
        "3) Recorta a ~1200 px de ancho para que se lea dentro del documento.  "
        "4) Guárdala como " + destino + ".  "
        "5) Vuelve a generar el guion: la imagen queda embebida aquí sola."
    )


def screenshot_box(doc, caption, explicit=None, receta=None):
    # Caja marcador + imagen real si coincide (token explícito o por palabra clave).
    p = doc.add_paragraph(); shade_par(p, "FFF4EC")
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    # El rotulo es un solo run en negrita: un `docker ps` saldria con los acentos
    # graves impresos, asi que aqui el span va a «». Se convierte solo para pintar;
    # find_capture() sigue viendo el caption original, que es con el que casa el PNG.
    rotulo = re.sub(r'`([^`\n]+)`', r'«\1»', caption)
    r = p.add_run("📸  " + rotulo); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = ORANGE
    img = None
    if explicit:
        cand = os.path.join(CAPTURAS, explicit)
        if os.path.isfile(cand):
            img = cand
    if img is None:
        img = find_capture(caption)
    if img:
        try:
            doc.add_picture(img, width=Inches(5.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    else:
        # Recurso que no se pudo generar: en vez de una caja vacía, el paso a paso
        # para producirlo. Regla transversal del material 2026-2.
        ph = doc.add_paragraph(); shade_par(ph, CODE_BG)
        ph.paragraph_format.space_after = Pt(1)
        rr = ph.add_run("Esta imagen todavía no existe. Cómo producirla:")
        rr.bold = True; rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = GRAY
        pr = doc.add_paragraph(); shade_par(pr, CODE_BG)
        pr.paragraph_format.left_indent = Inches(0.15)
        r2 = pr.add_run(receta or _receta_generica(caption, explicit))
        r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRAY

def add_code_block(doc, lines):
    p = doc.add_paragraph(); shade_par(p, CODE_BG)
    p.paragraph_format.left_indent = Inches(0.15); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    for i, ln in enumerate(lines):
        if i: p.add_run().add_break()
        r = p.add_run(ln); r.font.name = MONO; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

def add_heading(doc, level, text):
    sizes = {1: 20, 2: 16, 3: 13.5, 4: 12, 5: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 7); p.paragraph_format.space_after = Pt(3)
    add_inline(p, text, base_size=sizes.get(level, 12), base_color=DARK, base_bold=True)
    for r in p.runs: r.bold = True

def is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:|-]+\|?\s*$', line)) and '-' in line

def parse_table_row(line):
    line = line.strip().strip('|')
    # Se parte en los | que NO vengan escapados. Una celda puede llevar un pipe
    # literal —«docker images \| grep bibliolite»— y sin esto se volvia dos celdas:
    # corria el resto de la fila y partia en dos el span de codigo, que salia con
    # los acentos graves impresos a cada lado.
    return [c.strip().replace('\\|', '|') for c in re.split(r'(?<!\\)\|', line)]

def convert(md_path, out_path):
    global CAPTURAS
    # Capturas: misma carpeta del guion + /Capturas (p. ej. Kit docente/Clase N/Capturas)
    CAPTURAS = os.path.join(os.path.dirname(os.path.abspath(md_path)), "Capturas")
    with open(md_path, encoding="utf-8") as f:
        raw = f.read().split("\n")
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(11)
    title_done = False
    i = 0
    while i < len(raw):
        line = raw[i]
        s = line.strip()
        # Bloque de código ```
        if s.startswith("```"):
            block = []; i += 1
            while i < len(raw) and not raw[i].strip().startswith("```"):
                block.append(raw[i]); i += 1
            add_code_block(doc, block); i += 1; continue
        # Tabla
        if s.startswith("|") and i + 1 < len(raw) and is_table_sep(raw[i + 1]):
            header = parse_table_row(raw[i]); i += 2
            rows = []
            while i < len(raw) and raw[i].strip().startswith("|"):
                rows.append(parse_table_row(raw[i])); i += 1
            ncol = len(header)
            t = doc.add_table(rows=1 + len(rows), cols=ncol); t.style = 'Table Grid'
            for j, h in enumerate(header[:ncol]):
                c = t.cell(0, j); c.paragraphs[0].text = ""; shade_cell(c, HEAD_BG)
                add_inline(c.paragraphs[0], h, base_size=10.5, base_color=RGBColor(0xFF, 0xFF, 0xFF), base_bold=True)
                for r in c.paragraphs[0].runs: r.bold = True
            for ri, row in enumerate(rows):
                for j in range(ncol):
                    c = t.cell(ri + 1, j); c.paragraphs[0].text = ""
                    if ri % 2 == 1: shade_cell(c, ALT_BG)
                    add_inline(c.paragraphs[0], row[j] if j < len(row) else "", base_size=10)
            doc.add_paragraph(); continue
        # Vacío
        if not s:
            i += 1; continue
        # Regla horizontal
        if re.match(r'^(---+|\*\*\*+|___+)$', s):
            i += 1; continue
        # Encabezado #
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            hashes, text = m.group(1), m.group(2)
            lvl = len(hashes)
            if not title_done:
                p = doc.add_paragraph(); add_inline(p, text, base_size=22, base_color=DARK, base_bold=True)
                for r in p.runs: r.bold = True
                title_done = True
            else:
                add_heading(doc, max(2, lvl - 1), text)
            i += 1; continue
        # Blockquote
        if s.startswith(">"):
            content = s.lstrip(">").strip()
            if "📸" in content or "🖼" in content or re.search(r'\[\[\s*captura:', content, re.I):
                content, explicit, receta = _capture_token(content)
                cap = re.sub(r'^\s*(📸|🖼️?)\s*', '', content).replace("**", "").replace("Pantallazo —", "").replace("Imagen —", "").strip()
                cap = re.sub(r'^\W+', '', cap)
                screenshot_box(doc, cap or "Imagen", explicit, receta)
            else:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
                shade_par(p, "F7F7F8")
                add_inline(p, content, base_size=10.5, base_color=GRAY);
                for r in p.runs: r.italic = True
            i += 1; continue
        # 📸 / 🖼️ sin blockquote (o cualquier línea con token [[captura:]])
        if s.startswith("📸") or s.startswith("**📸") or "📸 **Pantallazo" in s or s.startswith("🖼") or re.search(r'\[\[\s*captura:', s, re.I):
            s2, explicit, receta = _capture_token(s)
            cap = re.sub(r'^\W*(📸|🖼️?)\W*', '', s2).replace("**", "").replace("Pantallazo —", "").replace("Imagen —", "").strip()
            screenshot_box(doc, cap or "Imagen", explicit, receta); i += 1; continue
        # Lista
        mli = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
        if mli:
            indent = len(mli.group(1)); ordered = bool(re.match(r'\d+\.', mli.group(2)))
            style = 'List Number' if ordered else 'List Bullet'
            if indent >= 2 and not ordered: style = 'List Bullet 2'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25 + (0.25 if indent >= 2 else 0))
            add_inline(p, mli.group(3), base_size=11)
            i += 1; continue
        # Párrafo normal
        p = doc.add_paragraph(); add_inline(p, s, base_size=11)
        i += 1
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    src = sys.argv[1]
    out = None
    if "--out" in sys.argv:
        out = sys.argv[sys.argv.index("--out") + 1]
    else:
        out = os.path.splitext(src)[0] + ".docx"
    convert(src, out)
    print("OK", out)
