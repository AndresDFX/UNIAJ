# -*- coding: utf-8 -*-
"""Genera Proyecto Integrador 2026-2 (BD II + Arquitectura).

Estudiante (compartible):
  <Curso>/Clases/Proyecto Integrador/Enunciado Proyecto Integrador - <Curso> - 2026-2.docx

Docente (privado):
  <Curso>/Kit docente/Proyecto Integrador/Guia Docente PI - <Curso> - 2026-2.docx
  <Curso>/Kit docente/Proyecto Integrador/Guia Docente PI - <Curso> - 2026-2.md
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xA0, 0x20, 0x30)
FONT = "Calibri"

DOCENTE = "Julian Andres Castaño Espinosa"
CORREO = "julianacastano@profesores.uniajc.edu.co"


def _shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _run(run, *, size=11, bold=False, color=GRIS, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _run_inline(p, text, *, size=11, bold=False, color=GRIS):
    """Soporta el marcador @@negrita@@ (convencion del workspace): parte el texto
    en runs y fuerza negrita solo en los tramos marcados."""
    for part in re.split(r'(@@.*?@@)', text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith('@@') and part.endswith('@@'):
            r.text = part[2:-2]
            _run(r, size=size, bold=True, color=color)
        else:
            r.text = part
            _run(r, size=size, bold=bold, color=color)


def para(doc, text, *, size=11, bold=False, color=GRIS, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=6, space_before=0, shade=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if shade:
        _shade(p, shade)
    _run_inline(p, text, size=size, bold=bold, color=color)
    return p


def banda(doc, text):
    return para(doc, f"  {text}", size=13, bold=True, color=BLANCO, shade="095292",
                space_before=10, space_after=8)


def h2(doc, text):
    return para(doc, text, size=12, bold=True, color=AZUL, space_before=12, space_after=6)


def h3(doc, text):
    return para(doc, text, size=11, bold=True, color=CIAN, space_before=8, space_after=4)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run_inline(p, it, size=11, color=GRIS)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        _run(r, size=10, bold=True, color=BLANCO)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "095292")
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            _run(r, size=10, color=GRIS)
            if ri % 2 == 1:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F2F2F2")
                cell._tc.get_or_add_tcPr().append(shd)
    para(doc, "", space_after=4)


def margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)


def portada_estudiante(doc, meta):
    para(doc, "Institución Universitaria Antonio José Camacho — UNIAJC",
         size=10, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Facultad de Ingenierías · Programa de Ingeniería de Sistemas",
         size=10, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    banda(doc, meta["titulo"])
    para(doc, meta["asignatura"], size=14, bold=True, color=AZUL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)
    para(doc, f"Código {meta['codigo']}  ·  Grupo {meta['grupo']}  ·  Periodo {meta['periodo']}",
         size=11, bold=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"Horario: {meta['horario']}  ·  Modalidad: Virtual (sesiones por Meet)",
         size=11, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"Docente: {DOCENTE}  ·  {CORREO}",
         size=10, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "Versión estudiante — Entrega en ExamLab según hitos del plan.",
         size=10, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER,
         shade="E8F4FA", space_after=10)


def portada_docente(doc, meta):
    para(doc, "Institución Universitaria Antonio José Camacho — UNIAJC",
         size=10, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    banda(doc, meta["titulo_docente"])
    para(doc, meta["asignatura"], size=13, bold=True, color=AZUL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)
    para(doc, f"{meta['codigo']} · Grupo {meta['grupo']} · {meta['periodo']}",
         size=11, bold=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "DOCUMENTO DOCENTE — No distribuir a estudiantes. Incluye rúbrica e hitos.",
         size=10, bold=True, color=ROJO, align=WD_ALIGN_PARAGRAPH.CENTER,
         shade="FBE4E4", space_after=10)


# ---------------------------------------------------------------------------
# BD II — VetCare DB
# ---------------------------------------------------------------------------

BD2_META = {
    "titulo": "Proyecto Integrador 2026-2 — VetCare DB",
    "titulo_docente": "Guía docente · Proyecto Integrador — Bases de Datos II",
    "asignatura": "Bases de Datos II",
    "codigo": "FI303215",
    "grupo": "641A-2",
    "periodo": "2026-2",
    "horario": "Lunes 18:00–20:00 (120 min)",
    "dominio": "VetCare DB",
}


def enunciado_bd2(doc):
    portada_estudiante(doc, BD2_META)

    h2(doc, "1. Propósito")
    para(doc,
         "Diseñar, administrar y optimizar una base de datos relacional avanzada para "
         "VetCare DB (clínica veterinaria), integrando seguridad/respaldo, objetos "
         "programables (procedimientos, funciones y disparadores), optimización e "
         "integración conceptual con una aplicación externa. Hilo conductor ABPr del curso.")
    para(doc,
         "Peso: 20% del Corte 3 (Acuerdo pedagógico). El Parcial 3 (15%) y la asistencia "
         "(5%) se evalúan por separado.")

    h2(doc, "2. Dominio — VetCare DB")
    para(doc,
         "Una clínica veterinaria necesita gestionar mascotas, dueños, veterinarios, "
         "citas, historial clínico, insumos/medicamentos y facturación básica. Usted "
         "modela y opera la capa de datos avanzada (no se exige una app de escritorio).")
    bullets(doc, [
        "Entidades mínimas sugeridas: Dueño, Mascota, Veterinario, Cita, Consulta/Historial, "
        "Insumo/Medicamento, DetalleFactura (ajuste el modelo con justificación).",
        "Reglas de negocio: no citar mascotas inactivas; stock de insumos no negativo; "
        "auditoría de cambios sensibles (precios, cancelaciones).",
        "Puede ampliar el dominio con 1–2 módulos propios si mantiene coherencia.",
    ])

    h2(doc, "3. Entregables obligatorios")
    bullets(doc, [
        "Modelo de datos: diagrama ER (draw.io / diagrams.net o Excalidraw) + script DDL.",
        "Administración y seguridad: roles/usuarios (o plan de privilegios), política de "
        "respaldo/recuperación documentada.",
        "Objetos programables: ≥2 procedimientos, ≥1 función, ≥2 disparadores con casos de prueba.",
        "Optimización: ≥2 consultas «antes/después» con justificación de índices o reescritura.",
        "Integración: propuesta de cómo una app (API/cliente) consumiría la BD "
        "(contrato de operaciones + ejemplo SQL o pseudocódigo de llamada a procedimientos).",
        "Informe breve (PDF o DOCX) + carpeta de scripts SQL + diagrama exportado (PNG/SVG).",
        "Sustentación / presentación (5–8 min) en el cierre del PI.",
    ])

    h2(doc, "4. Hitos y fechas orientativas (Plan 2026-2)")
    table(doc,
          ["Hito", "Sesión", "Fecha", "Qué entregar / hacer"],
          [
              ["Arranque (opcional)", "Sesiones 1–5", "ago–sep",
               "Definir dominio, borrador ER y alcance"],
              ["Construcción", "Sesiones 6–9", "sep–oct",
               "DDL, seguridad, procs/triggers, tuning"],
              ["Avance PI + integración", "Sesión 10 (doble · Clases 11+12)", "26/10/2026",
               "Demo parcial + checklist · contrato app ↔ BD y pruebas "
               "(cierre: domingo 01/11 23:59)"],
              ["Casos reales", "Sesión 11 (autónoma · Clase 13)", "02/11/2026",
               "Aplicar el análisis de casos al PI (cierre: domingo 08/11 23:59)"],
              ["Parcial 3 · paquete final", "Sesión 12 (Clase 14)", "09/11/2026",
               "Día de parcial = solo evaluación (presencial). El paquete final "
               "cierra el domingo 15/11 23:59; el ensayo y la preparación fueron "
               "en la Sesión 10 (Clase 12)"],
              ["Sustentación / cierre", "Sesión 13 (Clase 15)", "16/11/2026",
               "Defensa en vivo del PI + cierre del curso"],
          ])
    para(doc,
         "Nota: la sesión 13 (16/11) cae en festivo, pero por decisión docente es la "
         "sesión de sustentaciones: no es autónoma y no es parcial. El Parcial 3 es el "
         "09/11 (sesión 12 · material Clase 14). El paquete final se sube a ExamLab "
         "antes del 16/11; ese día solo se defiende.")

    h2(doc, "5. Herramientas (gratis + navegador)")
    bullets(doc, [
        "SQL: DB Fiddle, OneCompiler SQL, SQLTest.online o RunSQL.",
        "Procedimientos / PL-SQL: Oracle Live SQL (cuenta free, sin tarjeta).",
        "Diagramas: draw.io / diagrams.net · Excalidraw.",
        "Entregas: Google Docs/Drive o Word Online → subir a ExamLab.",
    ])
    para(doc,
         "No se exige instalar Oracle/MySQL/PostgreSQL/SQL Server ni Docker en el PC. "
         "No se pide cloud IaaS con tarjeta.",
         shade="FFF8D6", space_after=8)

    h2(doc, "6. Criterios de evaluación (rúbrica resumida — 100 pts → 20% del corte)")
    table(doc,
          ["Criterio", "Pts", "Evidencia"],
          [
              ["Modelo + DDL coherente", "20", "ER + scripts ejecutables"],
              ["Seguridad y respaldo", "15", "Roles/privilegios + plan backup"],
              ["Procs / funciones / triggers", "25", "Código + pruebas"],
              ["Optimización", "15", "Antes/después + justificación"],
              ["Integración app ↔ BD", "10", "Contrato + ejemplos"],
              ["Informe + sustentación", "15", "Claridad, demo, respuestas"],
          ])

    h2(doc, "7. Modalidad de trabajo")
    bullets(doc, [
        "Individual por defecto: cada estudiante desarrolla y entrega su propio proyecto.",
        "Opcional: el docente puede autorizar equipos de 2 o 3 integrantes.",
        "Debes poder explicar cualquier parte en la sustentación (si hay equipo, cualquier integrante).",
        "La entrega en ExamLab es siempre individual, incluso si el artefacto se trabajó en equipo.",
    ])

    h2(doc, "8. Qué NO es este proyecto")
    bullets(doc, [
        "No sustituye el Parcial 3 (evaluación síncrona presencial del 09/11).",
        "No es una app GUI completa: el foco es la capa de datos avanzada.",
        "No se copia un enunciado de otra institución: el dominio es VetCare DB UNIAJC.",
    ])


def guia_bd2(doc):
    portada_docente(doc, BD2_META)
    h2(doc, "Contexto")
    para(doc,
         "Patrón heredado de Prog. II / Seminario: PI continuo ABPr, 20% Corte 3. "
         "Enunciados viejos viven como .gdoc en Clase 1; aquí se materializa en "
         "Clases/Proyecto Integrador/ (estudiante) + Kit docente/Proyecto Integrador/.")
    h2(doc, "Alineación RAA")
    bullets(doc, [
        "RAA1 Seguridad y respaldo → roles, privilegios, plan de backup/restore.",
        "RAA2 Procedimientos y disparadores → procs, funciones, triggers con pruebas.",
        "RAA3 Optimización → índices, reescritura, evidencia antes/después.",
    ])
    h2(doc, "Hitos docentes (sesiones del calendario · 13 sesiones)")
    table(doc,
          ["Sesión", "Fecha", "Material", "Rol docente"],
          [
              ["10 (doble)", "26/10", "Clases 11+12",
               "Revisión de avance con checklist viva + empujar integración app↔BD "
               "y casos de prueba en el mismo bloque"],
              ["11 (autónoma)", "02/11", "Clase 13",
               "Publicar el caso real y revisar la entrega asíncrona (dom 08/11)"],
              ["12", "09/11", "Clase 14",
               "Parcial 3 presencial + ensayo de presentación PI"],
              ["13", "16/11", "Clase 15",
               "Sesión síncrona de sustentaciones: turnos, preguntas y cierre"],
          ])
    h2(doc, "Rúbrica detallada (calificar sobre 100 → escala a 20% del corte)")
    bullets(doc, [
        "Modelo/DDL (20): normalización razonable, FK, constraints, nombres claros.",
        "Seguridad/respaldo (15): al menos 2 roles con privilegios distintos; RPO/RTO "
        "en lenguaje simple; procedimiento de restore descrito.",
        "Programables (25): triggers de auditoría o integridad; procs de negocio "
        "(agendar cita, registrar consulta, descontar stock).",
        "Optimización (15): EXPLAIN o narrativa equivalente en playground; índices justificados.",
        "Integración (10): operaciones CRUD vía procs; no hace falta desplegar API real.",
        "Informe+sustentación (15): 5–8 min; preguntas cruzadas al azar sobre cualquier parte del trabajo.",
    ])
    h2(doc, "Evidencias a pedir")
    bullets(doc, [
        "Enlace o ZIP: scripts .sql, diagrama, informe.",
        "Capturas de ejecución en Live SQL / DB Fiddle.",
        "Reparto del guion por bloques temáticos (y, si hubo equipo autorizado, quién presentó qué).",
    ])
    h2(doc, "Post-clase Padlet")
    para(doc, "Si usó Padlet en Clase 1: ⋯ → Clear posts → código → Delete (reutilizar cupo gratis).")


# ---------------------------------------------------------------------------
# Arquitectura — CloudLite App
# ---------------------------------------------------------------------------

ARQ_META = {
    "titulo": "Proyecto Integrador 2026-2 — CloudLite App",
    "titulo_docente": "Guía docente · Proyecto Integrador — Arquitectura de Sistemas",
    "asignatura": "Arquitectura de Sistemas Computacionales",
    "codigo": "FI303380",
    "grupo": "6303C",
    "periodo": "2026-2",
    "horario": "Lunes 10:00–12:00 (120 min)",
    "dominio": "CloudLite App",
}


def enunciado_arq(doc):
    portada_estudiante(doc, ARQ_META)

    h2(doc, "1. Propósito")
    para(doc,
         "Diseñar y simular el despliegue de una arquitectura cloud para CloudLite App "
         "(aplicación web/API de un dominio realista a su elección: citas, academia, "
         "inventario liviano, etc.). Integra diagramas, contenedores en lab de navegador "
         "y CI/CD conceptual con GitHub Actions — sin cloud de pago ni tarjeta.")
    para(doc,
         "Peso: 20% del Corte 3 (Acuerdo pedagógico). El Parcial 3 (15%) y la asistencia "
         "(5%) se evalúan por separado.")

    h2(doc, "2. Alcance de CloudLite App")
    bullets(doc, [
        "Elegir un dominio concreto y justificar 3–5 capacidades funcionales.",
        "Decidir modelo de servicio dominante (IaaS / PaaS / SaaS) y por qué.",
        "Arquitectura lógica + de despliegue (capas, componentes, red, almacenamiento).",
        "Contenedorización de al menos un servicio (demo en LabEx Docker Playground o Killercoda).",
        "Pipeline CI/CD conceptual (GitHub Actions): build + test + artefactos (sin runner de pago).",
        "Seguridad, monitoreo, costos/sostenibilidad y escalabilidad (al menos un escenario).",
    ])

    h2(doc, "3. Entregables obligatorios")
    bullets(doc, [
        "Documento de arquitectura (PDF/DOCX): contexto, ADRs breves, riesgos, costos estimados "
        "en lenguaje cualitativo (bajo/medio — sin factura real).",
        "Diagramas (draw.io / diagrams.net): al menos (a) componentes/C4-lite, (b) despliegue, "
        "(c) flujo CI/CD.",
        "Lab de contenedores: Dockerfile o compose mínimo + captura/enlace de sesión "
        "LabEx Docker Playground o Killercoda.",
        "Repo o carpeta con workflow GitHub Actions (.yml) que compile/pruebe un stub "
        "(aunque el «deploy» sea simulado).",
        "Presentación de sustentación (5–8 min) con demo de diagrama + lab.",
    ])

    h2(doc, "4. Hitos y fechas orientativas (Plan 2026-2)")
    table(doc,
          ["Hito", "Sesión", "Fecha", "Qué entregar / hacer"],
          [
              ["Fundamentos", "Sesiones 1–5", "ago–sep",
               "Dominio, IaaS/PaaS/SaaS, boceto C4"],
              ["Profundización", "Sesiones 6–9", "sep–oct",
               "Seguridad, redes, monitoreo, CI/CD"],
              ["Avance PI + rendimiento", "Sesión 10 (doble · Clases 11+12)", "26/10/2026",
               "Diagramas v1 + checklist · escenario de prueba y métricas objetivo "
               "(cierre: domingo 01/11 23:59)"],
              ["Escalabilidad", "Sesión 11 (autónoma · Clase 13)", "02/11/2026",
               "Sección de autoescalado del informe (cierre: domingo 08/11 23:59)"],
              ["Parcial 3 · paquete final", "Sesión 12 (Clase 14)", "09/11/2026",
               "Día de parcial = solo evaluación (presencial). El paquete final "
               "cierra el domingo 15/11 23:59; el ensayo y la preparación fueron "
               "en la Sesión 10 (Clase 12)"],
              ["Sustentación / cierre", "Sesión 13 (Clase 15)", "16/11/2026",
               "Defensa en vivo del PI CloudLite + cierre del curso"],
          ])
    para(doc,
         "Nota: la sesión 13 (16/11) cae en festivo, pero por decisión docente es la "
         "sesión de sustentaciones: no es autónoma y no es parcial. El Parcial 3 es el "
         "09/11 (sesión 12 · material Clase 14).")

    h2(doc, "5. Herramientas (gratis + navegador)")
    bullets(doc, [
        "Diagramas: draw.io / diagrams.net · Excalidraw.",
        "Contenedores: LabEx Docker Playground · Killercoda (sin Docker Desktop obligatorio).",
        "CI/CD: GitHub Actions (cuenta free) — pipelines simples.",
        "Entregas: Google Docs/Drive o Word Online → subir a ExamLab.",
    ])
    para(doc,
         "Prohibido como requisito: AWS/GCP/Oracle Cloud/Azure Free Tier con tarjeta; "
         "instalar VirtualBox/VMware/Docker Desktop/WSL; software de modelado de pago.",
         shade="FBE4E4", space_after=8)

    h2(doc, "6. Criterios de evaluación (rúbrica resumida — 100 pts → 20% del corte)")
    table(doc,
          ["Criterio", "Pts", "Evidencia"],
          [
              ["Dominio + decisión IaaS/PaaS/SaaS", "15", "Justificación clara"],
              ["Diagramas de arquitectura", "25", "Componentes + despliegue"],
              ["Contenedores (lab navegador)", "20", "Dockerfile/compose + captura"],
              ["CI/CD conceptual", "15", "Workflow .yml + explicación"],
              ["Seguridad / costos / escalabilidad", "10", "Sección en informe"],
              ["Informe + sustentación", "15", "Claridad, demo, respuestas"],
          ])

    h2(doc, "7. Modalidad de trabajo")
    bullets(doc, [
        "Individual por defecto: cada estudiante desarrolla y entrega su propio proyecto.",
        "Opcional: el docente puede autorizar equipos de 2 o 3 integrantes.",
        "Debes poder explicar los diagramas y el workflow CI/CD (si hay equipo, cualquier integrante).",
        "La entrega en ExamLab es siempre individual, incluso si el artefacto se trabajó en equipo.",
    ])

    h2(doc, "8. Qué NO es este proyecto")
    bullets(doc, [
        "No sustituye el Parcial 3 (09/11, presencial).",
        "No exige cuenta cloud de pago ni gastos del estudiante.",
        "No se pide producción real en Internet: el foco es diseño + simulación en labs gratis.",
    ])


def guia_arq(doc):
    portada_docente(doc, ARQ_META)
    h2(doc, "Contexto")
    para(doc,
         "Mismo patrón de peso (20% Corte 3) que Prog. II / Seminario. Enfoque microcurrículo: "
         "arquitecturas cloud, virtualización/contenedores, seguridad, rendimiento y sostenibilidad.")
    h2(doc, "Alineación RAA")
    bullets(doc, [
        "RAA1 IaaS/PaaS/SaaS → decisión de modelo de servicio en el informe.",
        "RAA2 Virtualización y distribuidos → contenedores + diagrama de despliegue.",
        "RAA3 Seguridad, rendimiento y sostenibilidad → secciones explícitas + escenario de escala.",
    ])
    h2(doc, "Hitos docentes (sesiones del calendario · 13 sesiones)")
    table(doc,
          ["Sesión", "Fecha", "Material", "Rol docente"],
          [
              ["10 (doble)", "26/10", "Clases 11+12",
               "Revisar diagramas v1 y bloquear dominios demasiado amplios + "
               "métricas/pruebas de rendimiento en el mismo bloque"],
              ["11 (autónoma)", "02/11", "Clase 13",
               "Publicar el ejercicio de autoescalado y revisar la entrega (dom 08/11)"],
              ["12", "09/11", "Clase 14", "Parcial 3 + ensayo de pitch (5–8 min)"],
              ["13", "16/11", "Clase 15",
               "Sesión síncrona de sustentaciones: turnos, preguntas y cierre"],
          ])
    h2(doc, "Rúbrica detallada")
    bullets(doc, [
        "Dominio/servicio (15): problema real, límites claros, no «todo AWS».",
        "Diagramas (25): legibles, convenciones, separación de concerns.",
        "Contenedores (20): lab reproducible; si el lab caduca, capturas + Dockerfile bastan.",
        "CI/CD (15): stages claros; deploy puede ser «echo/simulate».",
        "Seguridad/costos/escala (10): amenazas básicas, estimación cualitativa, autoescalado conceptual.",
        "Sustentación (15): preguntas al azar sobre cualquier parte; si hubo equipo, penalizar que hable solo uno.",
    ])
    h2(doc, "Evidencias")
    bullets(doc, [
        "Repo GitHub o ZIP con .yml, Dockerfile, diagramas PNG, informe.",
        "Enlace Killercoda/LabEx o capturas con timestamp.",
    ])


# ---------------------------------------------------------------------------
# Programacion II y Seminario de Sistemas — VetCare (PI COMPARTIDO)
#
# Los dos cursos trabajan el MISMO producto: el sistema de la Clinica Veterinaria
# "Huellitas". Prog II construye el software; Seminario disena los planos. Por eso
# el enunciado contempla tres casos segun la matricula del estudiante (A/B/C):
# un estudiante puede estar en ambos cursos, o en uno solo, y no puede quedar
# bloqueado por lo que se hace en el otro.
#
# Contenido heredado de los enunciados originales del docente (que vivian sueltos
# como .gdoc: el de Prog II en Kit docente y el de Seminario dentro de Clase 1).
# ---------------------------------------------------------------------------

CASOS_MATRICULA = [
    "@@Caso A — cursa AMBAS materias (modo full stack):@@ trabaja sobre un unico "
    "VetCare. En Programacion II se califica la calidad del codigo Java; en Seminario, la "
    "calidad de los diagramas y la documentacion de ESE mismo sistema.",
    "@@Caso B — solo Programacion II (modo developer):@@ el docente le entrega los diagramas "
    "UML ya hechos. Su mision es traducir esos planos a codigo Java funcional.",
    "@@Caso C — solo Seminario (modo arquitecto/QA):@@ se encarga del diseno y el plan de "
    "pruebas. Para validar su diseno entrega un prototipo navegable (mockup) o un prototipo "
    "minimo en consola; no se le exige software con interfaz grafica.",
]

CONTEXTO_VETCARE = (
    "La Clinica Veterinaria «Huellitas» atiende un alto volumen de pacientes y hoy lleva "
    "toda su gestion en carpetas de papel. La administracion reporta tres problemas: se "
    "extravian fichas de pacientes, buscar un historial en el archivo fisico genera filas "
    "en la sala de espera, y no hay metricas (no saben cuantas especies atienden al mes). "
    "Ustedes fueron contratados para resolverlo con un sistema llamado @@VetCare@@."
)

PROG2_META = {
    "titulo": "Proyecto Integrador 2026-2 — VetCare (aplicacion Java)",
    "titulo_docente": "Guia docente · Proyecto Integrador — Programacion II",
    "asignatura": "Programacion II",
    "codigo": "FI303204",
    "grupo": "341C",
    "periodo": "2026-2",
    "horario": "Miercoles 18:00–20:00 (120 min)",
    "dominio": "VetCare",
}

SEMIN_META = {
    "titulo": "Proyecto Integrador 2026-2 — VetCare (arquitectura y diseno)",
    "titulo_docente": "Guia docente · Proyecto Integrador — Seminario de Sistemas",
    "asignatura": "Seminario de Sistemas",
    "codigo": "FI303301",
    "grupo": "341C",
    "periodo": "2026-2",
    "horario": "Jueves 18:00–20:00 (120 min)",
    "dominio": "VetCare",
}


def enunciado_prog2(doc):
    portada_estudiante(doc, PROG2_META)

    h2(doc, "1. Proposito")
    para(doc,
         "Construir en Java la aplicacion de escritorio VetCare para la clinica veterinaria "
         "«Huellitas», aplicando los pilares de la POO, estructuras de datos, interfaz grafica, "
         "manejo de errores y persistencia. Es el hilo conductor ABPr del curso: cada clase "
         "aporta una pieza del mismo producto, no ejercicios sueltos.")
    para(doc,
         "Peso: 20% del Corte 3 (Acuerdo pedagogico). El Parcial 3 (15%) y la asistencia (5%) "
         "se evaluan por separado.")

    h2(doc, "2. Contexto del cliente")
    para(doc, CONTEXTO_VETCARE)
    h3(doc, "El sistema debe permitir, como minimo")
    bullets(doc, [
        "Registrar duenos (clientes) y sus mascotas (pacientes).",
        "Agendar citas medicas.",
        "Guardar un historial clinico basico por mascota.",
        "Buscar rapidamente el expediente de un animal por su identificador.",
        "Conservar la informacion aunque se cierre el programa.",
    ])

    h2(doc, "3. Requisitos tecnicos obligatorios")
    bullets(doc, [
        "@@Pilares de POO:@@ clases del dominio (Mascota, Dueno, Cita), objetos, "
        "encapsulamiento y herencia usados con criterio, no decorativos.",
        "@@Estructuras de datos:@@ colecciones de Java (List / Map) para gestionar los "
        "registros en memoria.",
        "@@Interfaz grafica:@@ ventanas funcionales con Swing o JavaFX.",
        "@@Manejo de errores:@@ bloques try-catch que eviten que el programa se cierre "
        "(ej.: el usuario escribe texto donde va la edad de la mascota).",
        "@@Persistencia:@@ guardar y leer desde archivos .txt o .csv para no perder datos "
        "al cerrar la aplicacion.",
    ])

    h2(doc, "4. Como se conecta con Seminario de Sistemas")
    para(doc,
         "VetCare es el MISMO producto en las dos asignaturas: aqui se construye el software, "
         "en Seminario se disenan los planos (requisitos, UML, pruebas, manuales). Su "
         "situacion depende de su matricula:")
    bullets(doc, CASOS_MATRICULA)

    h2(doc, "5. Hitos de entrega")
    table(doc,
          ["Sesion (material)", "Fecha", "Que se entrega"],
          [
              ["S10 (doble · Clases 11+12)", "28/10",
               "Avance de codigo + revision cruzada e integracion de modulos: "
               "la aplicacion corre de punta a punta"],
              ["S11 (Clase 13)", "04/11", "Manejo de excepciones y validaciones incorporado"],
              ["S12 (Clase 14)", "11/11", "Version final + ensayo de la presentacion"],
              ["S13 (Clase 15)", "18/11", "Sustentacion final (mismo dia del Parcial 3)"],
          ])

    h2(doc, "6. Entregables finales")
    bullets(doc, [
        "Codigo fuente del proyecto (ZIP o enlace al repositorio).",
        "Archivo de datos de ejemplo (.txt o .csv) con registros de prueba.",
        "Manual de usuario breve (1-2 paginas) con capturas de la interfaz.",
        "Sustentacion de 5-8 minutos (si hay equipo autorizado, deben participar todos los integrantes).",
    ])

    h2(doc, "7. Criterios de evaluacion (100 puntos)")
    table(doc,
          ["Criterio", "Pts", "Que se mira"],
          [
              ["POO", "20", "Clases del dominio bien modeladas; encapsulamiento y herencia con proposito"],
              ["Estructuras de datos", "15", "Coleccion adecuada al caso y justificada"],
              ["Interfaz grafica", "20", "Ventanas funcionales; el usuario puede completar las tareas"],
              ["Manejo de errores", "15", "La aplicacion no se cae ante entradas invalidas"],
              ["Persistencia", "15", "Los datos sobreviven al cierre del programa"],
              ["Sustentacion", "15", "Se pregunta al azar sobre cualquier parte; con equipo, todos explican"],
          ])

    h2(doc, "8. Entrega")
    para(doc, "Entrega en @@ExamLab@@ segun los hitos del plan. Envio individual.")


def guia_prog2(doc):
    portada_docente(doc, PROG2_META)
    h2(doc, "Contexto")
    para(doc,
         "PI compartido con Seminario de Sistemas: mismo producto VetCare, distinto entregable. "
         "El enunciado original vivia suelto como .gdoc en Kit docente; ahora el documento del "
         "estudiante vive en Clases/Proyecto Integrador/ y esta guia en Kit docente/.")
    h2(doc, "Alineacion RAA")
    bullets(doc, [
        "RAA1 Estructuras de datos → colecciones para registros de VetCare.",
        "RAA2 Eventos y componentes graficos → interfaz Swing/JavaFX funcional.",
        "RAA3 Patrones de diseno → organizacion del codigo, separacion de responsabilidades.",
    ])
    h2(doc, "Los tres casos de matricula (lo que mas dudas genera)")
    bullets(doc, [
        "Caso A (ambas materias): NO calificar dos veces lo mismo. Aqui se califica el codigo; "
        "los diagramas se califican en Seminario.",
        "Caso B (solo Prog II): usted les entrega los diagramas UML ya hechos. Tenga listo un "
        "juego de diagramas base de VetCare para repartir en Clase 11.",
        "Caso C (solo Seminario): no aparece en este curso, pero si un estudiante pregunta, "
        "sepa que alla entrega mockup/prototipo, no codigo con interfaz.",
    ])
    h2(doc, "Hitos docentes")
    table(doc,
          ["Sesion (material)", "Fecha", "Rol docente"],
          [
              ["S10 (doble · Clases 11+12)", "28/10",
               "Revision cruzada (cada estudiante lee el codigo de otro) y empujar la "
               "integracion; detectar trabajos con modulos sueltos"],
              ["S11 (Clase 13)", "04/11", "Verificar que las validaciones existan de verdad, no en el papel"],
              ["S12 (Clase 14)", "11/11", "Ensayo cronometrado de la sustentacion"],
              ["S13 (Clase 15)", "18/11", "Parcial 3 + sustentacion final"],
          ])
    h2(doc, "Errores frecuentes a vigilar")
    bullets(doc, [
        "Herencia forzada solo para «cumplir el requisito» (una clase Animal que nadie usa).",
        "GUI que se ve bien pero no guarda nada: pedir siempre la prueba de cerrar y reabrir.",
        "try-catch vacio que se traga el error sin avisar al usuario.",
        "Si hubo equipo, que uno solo haya programado todo: por eso las preguntas al azar en la sustentacion.",
    ])
    h2(doc, "Evidencias a pedir")
    bullets(doc, [
        "ZIP o repositorio con el codigo fuente compilable.",
        "Archivo .csv/.txt con datos de prueba.",
        "Captura de la aplicacion corriendo + manual breve.",
    ])


def enunciado_seminario(doc):
    portada_estudiante(doc, SEMIN_META)

    h2(doc, "1. Proposito")
    para(doc,
         "Disenar la arquitectura, los requisitos y la interfaz del sistema VetCare para la "
         "clinica veterinaria «Huellitas». Su rol en esta asignatura es de @@analista funcional "
         "y arquitecto de software@@: aqui no se construye la casa, se dibujan los planos para "
         "que cualquier equipo pueda construirla.")
    para(doc,
         "Peso: 20% del Corte 3 (Acuerdo pedagogico). El Parcial 3 (15%) y la asistencia (5%) "
         "se evaluan por separado.")

    h2(doc, "2. Contexto del cliente")
    para(doc, CONTEXTO_VETCARE)
    h3(doc, "Requerimientos que el veterinario jefe pidio en la primera entrevista")
    bullets(doc, [
        "«Necesito registrar el ID, el nombre y la especie del animal que llega.»",
        "«Necesito buscar rapido el expediente de un animal usando solo su ID.»",
        "«Quiero ver una lista de quienes estan en la sala de espera.»",
        "«Tiene que ser muy facil de usar: no somos expertos en computadoras.»",
        "«La informacion no puede borrarse si se va la luz o apagamos el computador.»",
    ])
    para(doc,
         "Esas cinco frases son material crudo de entrevista, no requisitos formales. "
         "Convertirlas en RF/RNF bien escritos es parte del trabajo.")

    h2(doc, "3. Fases de entrega")
    table(doc,
          ["Fase", "Que produce"],
          [
              ["1. Ingenieria de requisitos",
               "Requisitos funcionales (RF) y no funcionales (RNF) formales y trazables"],
              ["2. Modelado UML",
               "Casos de uso, diagrama de clases del dominio y un diagrama dinamico "
               "(actividad o secuencia)"],
              ["3. Diseno de interfaz (UX/UI)",
               "Wireframes y mockup navegable de las pantallas (Figma, Penpot o Balsamiq)"],
              ["4. Arquitectura de datos",
               "Diccionario de datos y formato de almacenamiento (estructura del CSV o "
               "modelo entidad-relacion basico)"],
          ])

    h2(doc, "4. Como se conecta con Programacion II")
    para(doc,
         "VetCare es el MISMO producto en las dos asignaturas: aqui se disenan los planos y "
         "en Programacion II se construye el software. Su situacion depende de su matricula:")
    bullets(doc, CASOS_MATRICULA)
    para(doc,
         "@@Independencia curricular:@@ si NO cursa Programacion II, su proyecto termina con el "
         "documento de diseno y el prototipo visual navegable. No necesita escribir ni compilar "
         "codigo: su nota depende de la calidad, coherencia y profesionalismo de sus planos.",
         shade="E8F4FA")

    h2(doc, "5. Hitos de entrega")
    table(doc,
          ["Sesion (material)", "Fecha", "Que se entrega"],
          [
              ["S10 (doble · Clases 11+12)", "29/10",
               "Avance: requisitos (RF/RNF) + casos de uso y diagramas UML avanzados "
               "(clases + dinamico)"],
              ["S11 (Clase 13)", "05/11", "Diseno de interfaz: wireframes y mockup navegable"],
              ["S12 (Clase 14)", "12/11", "Documento consolidado + ensayo de la sustentacion"],
              ["S13 (Clase 15)", "19/11", "Sustentacion final (mismo dia del Parcial 3)"],
          ])

    h2(doc, "6. Entregables finales")
    bullets(doc, [
        "Documento de diseno de arquitectura con RF/RNF, UML y diccionario de datos.",
        "Mockup navegable (enlace o PDF exportado).",
        "Plan de pruebas: casos de prueba documentados sobre los requisitos.",
        "Sustentacion de 5-8 minutos (si hay equipo autorizado, deben participar todos los integrantes).",
    ])

    h2(doc, "7. Criterios de evaluacion (100 puntos)")
    table(doc,
          ["Criterio", "Pts", "Que se mira"],
          [
              ["Requisitos", "20", "RF/RNF verificables y trazables al pedido del cliente"],
              ["Modelado UML", "25", "Diagramas coherentes entre si y con los requisitos"],
              ["Diseno de interfaz", "20", "Wireframes usables; el flujo se entiende sin explicacion"],
              ["Arquitectura de datos", "15", "Diccionario de datos completo y formato justificado"],
              ["Plan de pruebas", "10", "Casos de prueba que cubren los requisitos criticos"],
              ["Sustentacion", "10", "Se pregunta al azar sobre cualquier parte; con equipo, todos explican"],
          ])

    h2(doc, "8. Entrega")
    para(doc, "Entrega en @@ExamLab@@ segun los hitos del plan. Envio individual.")


def guia_seminario(doc):
    portada_docente(doc, SEMIN_META)
    h2(doc, "Contexto")
    para(doc,
         "PI compartido con Programacion II: mismo producto VetCare, distinto entregable. "
         "El enunciado original vivia suelto como .gdoc dentro de la carpeta de Clase 1; ahora "
         "el documento del estudiante vive en Clases/Proyecto Integrador/ y esta guia en Kit docente/.")
    h2(doc, "Alineacion RAA")
    bullets(doc, [
        "RAA1 Patrones y modularidad → organizacion del diseno, separacion de responsabilidades.",
        "RAA2 Documentacion y validacion → SDD, diccionario de datos, plan de pruebas.",
        "RAA3 Presentacion y sustentacion → defensa clara de las decisiones de diseno.",
    ])
    h2(doc, "Los tres casos de matricula (lo que mas dudas genera)")
    bullets(doc, [
        "Caso A (ambas materias): aqui se califican los planos; el codigo se califica en Prog II. "
        "No exigir codigo compilable en esta asignatura.",
        "Caso C (solo Seminario): cierra con SDD + mockup navegable. Es una ruta completa y "
        "valida, no una version reducida: dejarlo claro desde Clase 1 para que nadie se sienta "
        "en desventaja.",
        "Caso B (solo Prog II): no aparece aca; sus diagramas base los provee el docente en el "
        "otro curso. Mantener un juego de diagramas VetCare listo para repartir alla.",
    ])
    h2(doc, "Hitos docentes")
    table(doc,
          ["Sesion (material)", "Fecha", "Rol docente"],
          [
              ["S10 (doble · Clases 11+12)", "29/10",
               "Revisar que los RF sean verificables (no deseos vagos) y la coherencia "
               "entre casos de uso y diagrama de clases"],
              ["S11 (Clase 13)", "05/11", "Wireframes: exigir flujo completo de una tarea real"],
              ["S12 (Clase 14)", "12/11", "Consolidacion del documento + ensayo cronometrado"],
              ["S13 (Clase 15)", "19/11", "Parcial 3 + sustentacion final"],
          ])
    h2(doc, "Errores frecuentes a vigilar")
    bullets(doc, [
        "Requisitos no verificables («el sistema debe ser rapido») sin criterio medible.",
        "Diagrama de clases que no corresponde a los casos de uso entregados.",
        "Mockup bonito que no cubre el flujo critico (registrar → buscar → ver historial).",
        "Diccionario de datos sin tipos ni restricciones: queda inservible para quien programe.",
    ])
    h2(doc, "Evidencias a pedir")
    bullets(doc, [
        "Documento de diseno (PDF o DOCX) con RF/RNF, UML y diccionario de datos.",
        "Enlace al mockup navegable (Figma/Penpot) o PDF exportado.",
        "Plan de pruebas con casos trazados a requisitos.",
    ])


# ---------------------------------------------------------------------------
# Markdown kits (fuente editable docente)
# ---------------------------------------------------------------------------

MD_BD2 = """# Guía docente — Proyecto Integrador · Bases de Datos II · 2026-2

**Privado docente** · No compartir en `Clases/`.

- **Curso:** Bases de Datos II (FI303215 · 641A-2)
- **Peso:** 20% Corte 3 (Acuerdo)
- **Dominio estudiante:** VetCare DB
- **Enunciado:** `Clases/Proyecto Integrador/Enunciado Proyecto Integrador - Bases de Datos II - 2026-2.docx`

## Hitos (Plan 2026-2 · semestre acortado a 13 sesiones)

Las carpetas de material siguen numeradas 1–15; lo que cambia es en qué sesión se dicta cada una.
Toda entrega intermedia cierra el **domingo 23:59 siguiente** a la sesión correspondiente.

| Sesión | Fecha | Material | Foco | Cierre de entrega |
|---|---|---|---|---|
| S10 (doble) | 26/10/2026 | Clases 11 + 12 | Avance PI — checklist · Integración app ↔ BD · **preparación de la sustentación** | dom 01/11/2026 23:59 |
| S11 (autónoma) | 02/11/2026 | Clase 13 | Análisis de casos reales aplicado al PI | dom 08/11/2026 23:59 |
| S12 | 09/11/2026 | Clase 14 | **Parcial 3** (presencial · solo evaluación) | **paquete final:** dom 15/11/2026 23:59 |
| S13 | 16/11/2026 | Clase 15 | **Sustentación en vivo del PI** + cierre (sesión síncrona) | — (el paquete ya está entregado) |

> S13 cae en festivo (Independencia de Cartagena), pero por decisión docente es sesión de
> sustentaciones: no es autónoma y no es parcial. El paquete final se sube a ExamLab antes
> de la sesión; el 16/11 solo se defiende.

## Rúbrica (100 pts)

Modelo/DDL 20 · Seguridad/respaldo 15 · Programables 25 · Optimización 15 · Integración 10 · Informe+sustentación 15.

## Evidencias

Scripts SQL · ER · capturas Live SQL/DB Fiddle · informe · presentación 5–8 min.
"""

MD_ARQ = """# Guía docente — Proyecto Integrador · Arquitectura · 2026-2

**Privado docente** · No compartir en `Clases/`.

- **Curso:** Arquitectura de Sistemas Computacionales (FI303380 · 6303C)
- **Peso:** 20% Corte 3 (Acuerdo)
- **Dominio estudiante:** CloudLite App
- **Enunciado:** `Clases/Proyecto Integrador/Enunciado Proyecto Integrador - Arquitectura de Sistemas Computacionales - 2026-2.docx`

## Hitos (Plan 2026-2 · semestre acortado a 13 sesiones)

Las carpetas de material siguen numeradas 1–15; lo que cambia es en qué sesión se dicta cada una.
Toda entrega intermedia cierra el **domingo 23:59 siguiente** a la sesión correspondiente.

| Sesión | Fecha | Material | Foco | Cierre de entrega |
|---|---|---|---|---|
| S10 (doble) | 26/10/2026 | Clases 11 + 12 | Avance PI — diagramas v1 · Pruebas de rendimiento · **preparación de la sustentación** | dom 01/11/2026 23:59 |
| S11 (autónoma) | 02/11/2026 | Clase 13 | Escalabilidad automática aplicada al PI | dom 08/11/2026 23:59 |
| S12 | 09/11/2026 | Clase 14 | **Parcial 3** (presencial · solo evaluación) | **paquete final:** dom 15/11/2026 23:59 |
| S13 | 16/11/2026 | Clase 15 | **Sustentación en vivo del PI** + cierre (sesión síncrona) | — (el paquete ya está entregado) |

> S13 cae en festivo (Independencia de Cartagena), pero por decisión docente es sesión de
> sustentaciones: no es autónoma y no es parcial. El paquete final se sube a ExamLab antes
> de la sesión; el 16/11 solo se defiende.

## Rúbrica (100 pts)

Dominio/servicio 15 · Diagramas 25 · Contenedores 20 · CI/CD 15 · Seg/costos/escala 10 · Informe+sustentación 15.

## Herramientas

draw.io · LabEx Docker Playground / Killercoda · GitHub Actions · **sin** AWS/GCP/Oracle Cloud con tarjeta.
"""


def build_one(curso_dir: Path, meta: dict, build_est, build_doc, md_text: str, slug: str):
    est_dir = curso_dir / "Clases" / "Proyecto Integrador"
    kit_dir = curso_dir / "Kit docente" / "Proyecto Integrador"
    est_dir.mkdir(parents=True, exist_ok=True)
    kit_dir.mkdir(parents=True, exist_ok=True)

    est_path = est_dir / f"Enunciado Proyecto Integrador - {slug} - 2026-2.docx"
    doc_path = kit_dir / f"Guia Docente PI - {slug} - 2026-2.docx"
    md_path = kit_dir / f"Guia Docente PI - {slug} - 2026-2.md"

    d1 = Document()
    margins(d1)
    build_est(d1)
    d1.save(est_path)

    d2 = Document()
    margins(d2)
    build_doc(d2)
    d2.save(doc_path)

    md_path.write_text(md_text, encoding="utf-8")
    print("OK ->", est_path)
    print("OK ->", doc_path)
    print("OK ->", md_path)

MD_PROG2 = """# Guía docente — Proyecto Integrador · Programación II · 2026-2

**Privado docente** · No compartir en `Clases/`.

## Producto
**VetCare** — aplicación de escritorio en Java para la Clínica Veterinaria «Huellitas».
PI **compartido con Seminario de Sistemas**: mismo producto, distinto entregable
(aquí el software, allá los planos).

## Casos de matrícula
- **A** (ambas materias): aquí se califica el código; los diagramas se califican en Seminario.
- **B** (solo Prog II): el docente entrega los diagramas UML ya hechos; el estudiante programa.
- **C** (solo Seminario): no aplica en este curso.

## Hitos
Semestre acortado a **13 sesiones**; las carpetas de material siguen numeradas 1–15.
Toda entrega intermedia cierra el **domingo 23:59 siguiente** a la sesión.

| Sesión | Fecha | Material | Entrega | Cierre |
|---|---|---|---|---|
| S10 (doble) | 28/10 | Clases 11 + 12 | Avance de código + revisión cruzada e integración de módulos | dom 01/11/2026 23:59 |
| S11 | 04/11 | Clase 13 | Excepciones y validaciones | dom 08/11/2026 23:59 |
| S12 | 11/11 | Clase 14 | Versión final + ensayo | **paquete final:** dom 15/11/2026 23:59 |
| S13 | 18/11 | Clase 15 | Sustentación (mismo día del Parcial 3) | — |

## Rúbrica (100 pts → 20% del Corte 3)
POO 20 · Estructuras de datos 15 · GUI 20 · Manejo de errores 15 · Persistencia 15 · Sustentación 15

## Errores frecuentes
- Herencia forzada solo para cumplir el requisito.
- GUI que no persiste: pedir siempre la prueba de cerrar y reabrir.
- `try-catch` vacío que oculta el error.
- Si hubo equipo autorizado, que uno solo haya programado todo → preguntas al azar en la sustentación.

Entrega en **ExamLab**. Enunciado del estudiante en `Clases/Proyecto Integrador/`.
"""

MD_SEMIN = """# Guía docente — Proyecto Integrador · Seminario de Sistemas · 2026-2

**Privado docente** · No compartir en `Clases/`.

## Producto
**VetCare** — arquitectura y diseño del sistema para la Clínica Veterinaria «Huellitas».
PI **compartido con Programación II**: mismo producto, distinto entregable
(aquí los planos, allá el software).

## Casos de matrícula
- **A** (ambas materias): aquí se califican los planos; el código se califica en Prog II.
- **C** (solo Seminario): cierra con SDD + mockup navegable. Ruta completa y válida,
  **no** una versión reducida — decirlo desde Clase 1.
- **B** (solo Prog II): no aplica aquí.

## Fases
1. Ingeniería de requisitos (RF/RNF)
2. Modelado UML (casos de uso, clases, dinámico)
3. Diseño de interfaz (wireframes + mockup navegable)
4. Arquitectura de datos (diccionario de datos)

## Hitos
Semestre acortado a **13 sesiones**; las carpetas de material siguen numeradas 1–15.
Toda entrega intermedia cierra el **domingo 23:59 siguiente** a la sesión.

| Sesión | Fecha | Material | Entrega | Cierre |
|---|---|---|---|---|
| S10 (doble) | 29/10 | Clases 11 + 12 | Requisitos + casos de uso y UML avanzado | dom 01/11/2026 23:59 |
| S11 | 05/11 | Clase 13 | Wireframes y mockup | dom 08/11/2026 23:59 |
| S12 | 12/11 | Clase 14 | Documento consolidado + ensayo | **paquete final:** dom 15/11/2026 23:59 |
| S13 | 19/11 | Clase 15 | Sustentación (mismo día del Parcial 3) | — |

## Rúbrica (100 pts → 20% del Corte 3)
Requisitos 20 · UML 25 · Interfaz 20 · Arquitectura de datos 15 · Plan de pruebas 10 · Sustentación 10

## Errores frecuentes
- Requisitos no verificables («debe ser rápido») sin criterio medible.
- Diagrama de clases que no corresponde a los casos de uso.
- Mockup que no cubre el flujo crítico (registrar → buscar → ver historial).
- Diccionario de datos sin tipos ni restricciones.

Entrega en **ExamLab**. Enunciado del estudiante en `Clases/Proyecto Integrador/`.
"""


def build():
    build_one(
        ROOT / "Bases de Datos II",
        BD2_META,
        enunciado_bd2,
        guia_bd2,
        MD_BD2,
        "Bases de Datos II",
    )
    build_one(
        ROOT / "Arquitectura de Sistemas Computacionales",
        ARQ_META,
        enunciado_arq,
        guia_arq,
        MD_ARQ,
        "Arquitectura de Sistemas Computacionales",
    )
    build_one(
        ROOT / "Programacion II",
        PROG2_META,
        enunciado_prog2,
        guia_prog2,
        MD_PROG2,
        "Programacion II",
    )
    build_one(
        ROOT / "Seminario de Sistemas",
        SEMIN_META,
        enunciado_seminario,
        guia_seminario,
        MD_SEMIN,
        "Seminario de Sistemas",
    )


if __name__ == "__main__":
    build()
