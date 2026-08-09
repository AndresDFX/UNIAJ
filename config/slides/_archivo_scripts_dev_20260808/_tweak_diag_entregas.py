# -*- coding: utf-8 -*-
"""Ajuste ligero de wording en DIAGNOSTICO institucional BD II / Arq."""
from pathlib import Path
from docx import Document

ROOT = Path(r"G:\Mi unidad\Trabajos\Empleo\UNIAJ\Cursos")

def tweak(path, prev_snip, new_prev=None, hall_old=None, hall_new=None):
    doc = Document(path)
    # Table 1 = conceptos previos (first content table after header often index 1)
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if hall_old and hall_old in t:
                    # replace paragraph text carefully
                    for p in cell.paragraphs:
                        if hall_old in p.text:
                            # clear and rewrite first paragraph of cell if matches
                            pass
                    # simpler: set cell text if exact-ish
                    if cell.text.strip().startswith("[PENDIENTE]") or hall_old in cell.text:
                        # Keep structure: join with newlines if multi
                        cell.text = hall_new
                        changed += 1
                if new_prev and prev_snip in t and new_prev not in t:
                    # prepend note if missing
                    if "conocimientos previos" not in t.lower():
                        cell.text = "Prerrequisitos (conocimientos previos del curso anterior / fundamentos):\n" + t
                        changed += 1
    doc.save(path)
    print("OK", path.name, "cells", changed)

bd = ROOT / "Bases de Datos II" / "Entregas docente" / "2026-2" / "DIAGNOSTICO - Bases de Datos II - 2026-2.docx"
arq = ROOT / "Arquitectura de Sistemas Computacionales" / "Entregas docente" / "2026-2" / "DIAGNOSTICO - Arquitectura de Sistemas Computacionales - 2026-2.docx"

for path in (bd, arq):
    doc = Document(path)
    # find table with conceptos previos header
    for table in doc.tables:
        header = table.rows[0].cells[0].text.strip() if table.rows else ""
        if "conceptos previos" in header.lower() or "Liste los conceptos previos" in header:
            # body cell is usually row 1
            if len(table.rows) > 1:
                body = table.rows[1].cells[0]
                txt = body.text.strip()
                if not txt.lower().startswith("prerrequisitos"):
                    body.text = "Prerrequisitos (conocimientos previos — no el temario avanzado del semestre):\n" + txt
                    print("previos OK", path.name)
        if "hallazgos en el diagnostico" in header.lower() or "hallazgos en el diagn" in header.lower():
            if len(table.rows) > 1:
                body = table.rows[1].cells[0]
                txt = body.text
                if "conocimientos previos" not in txt.lower():
                    body.text = (
                        "[PENDIENTE] Aplicar en Clase 1 la prueba de conocimientos previos (Kit docente/Clase 1/Prueba Diagnostica…). "
                        "Registrar nivel general, fortalezas y debilidades tras calificar. "
                        "No inventar listado ni promedios hasta tener resultados reales."
                    )
                    print("hallazgos OK", path.name)
    doc.save(path)
    print("saved", path.name)