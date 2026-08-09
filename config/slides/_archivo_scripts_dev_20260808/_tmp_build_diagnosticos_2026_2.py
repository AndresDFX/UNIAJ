# -*- coding: utf-8 -*-
"""Genera diagnósticos BD II / Arquitectura + actualiza CSV Clase 1 / regenera PPTX."""
from __future__ import annotations

import csv
import os
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Twips

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

NAVY = RGBColor(0x09, 0x52, 0x92)
CIAN = RGBColor(0x26, 0x9C, 0xCB)
GRAY = RGBColor(0x2B, 0x2B, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DOCENTE = "Julian Andres Castaño Espinosa"
CORREO = "julianacastano@profesores.uniajc.edu.co"

# --- helpers docx ---

def set_run(run, *, size=11, bold=False, color=GRAY, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_p(doc, text, *, size=11, bold=False, color=GRAY, space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): hex_color,
        },
    )
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=GRAY, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)


def write_multiline_cell(cell, lines, *, size=10):
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run(r, size=size, color=GRAY)


def clear_merged_row_text(table, row_idx, text):
    """Set text on first cell of a (possibly merged) row."""
    cell = table.rows[row_idx].cells[0]
    # also clear duplicates from merge
    seen = set()
    for c in table.rows[row_idx].cells:
        if id(c) in seen:
            continue
        seen.add(id(c))
        if c is cell:
            continue
        c.text = ""
    write_multiline_cell(cell, text if isinstance(text, list) else [text])


def build_institucional_diagnostico(
    *,
    out_path: Path,
    asignatura: str,
    grupo: str,
    periodo: str,
    prerrequisitos: list[str],
    acciones_previstas: list[str],
):
    """Copia plantilla Prog II y rellena Parte 1; Partes 2-3 quedan pendientes."""
    template = (
        ROOT
        / "Programacion II"
        / "Entregas docente"
        / "DIAGNOSTICO Y SEGUMIENTO ACTUALIZADO.docx"
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, out_path)
    doc = Document(str(out_path))

    # Table 0: header parte 1
    t0 = doc.tables[0]
    # Profesor already set in template; ensure docente
    for c in t0.rows[0].cells[1:]:
        set_cell_text(c, DOCENTE, size=10)
    for c in t0.rows[1].cells[1:]:
        set_cell_text(c, asignatura, size=10)
    # Grupo / Periodo
    set_cell_text(t0.rows[2].cells[1], grupo, size=10)
    set_cell_text(t0.rows[2].cells[3], periodo, size=10)

    # Table 1: prerrequisitos
    clear_merged_row_text(doc.tables[1], 1, prerrequisitos)

    # Table 2: hallazgos — pendiente
    clear_merged_row_text(
        doc.tables[2],
        1,
        [
            "[PENDIENTE] Aplicar el instrumento de evaluación diagnóstica en Clase 1.",
            "Registrar aquí el nivel general, fortalezas y debilidades del grupo tras calificar la prueba.",
            "No inventar listado ni promedios hasta tener resultados reales.",
        ],
    )

    # Table 3: acciones
    clear_merged_row_text(doc.tables[3], 1, acciones_previstas + [
        "[PENDIENTE] Ajustar acciones según hallazgos reales del diagnóstico de Clase 1.",
    ])

    # Table 4: firma
    set_cell_text(doc.tables[4].rows[0].cells[1], DOCENTE, size=10)
    set_cell_text(doc.tables[4].rows[1].cells[1], "[PENDIENTE]", size=10)

    # Parte 2 header (table 5)
    t5 = doc.tables[5]
    for c in t5.rows[0].cells[1:]:
        set_cell_text(c, DOCENTE, size=9)
    for c in t5.rows[1].cells[1:]:
        set_cell_text(c, asignatura, size=9)
    # grupo / periodo cells — keep structure, fill known
    set_cell_text(t5.rows[2].cells[1], grupo, size=9)
    # Periodo cells typically at index 5+
    for ci, c in enumerate(t5.rows[2].cells):
        if ci >= 5:
            set_cell_text(c, periodo, size=9)
    # estudiantes matriculados etc
    # row 3 has counts — set pendiente
    for c in t5.rows[3].cells:
        txt = c.text.strip()
        if txt.isdigit() or txt == "" or (txt and txt[0].isdigit()):
            # leave labels, clear numbers
            pass
    # Safer: rewrite numeric-looking middle cells to [PENDIENTE]
    # Cells: label, value, label, value pattern — set value cells
    cells = t5.rows[3].cells
    # Known layout from Prog II: matriculados value around index 2, perdieron ~4, no presentaron ~8
    try:
        set_cell_text(cells[2], "[PENDIENTE]", size=9)
        set_cell_text(cells[4], "[PENDIENTE]", size=9)
        set_cell_text(cells[8], "[PENDIENTE]", size=9)
    except Exception:
        pass

    # Table 6: reportes parcial — clear filled narrative
    clear_merged_row_text(
        doc.tables[6],
        1,
        ["[PENDIENTE] Completar tras el Parcial 1 / cierre del Corte 1."],
    )
    clear_merged_row_text(
        doc.tables[6],
        3,
        ["[PENDIENTE] Identificar estudiantes con dificultades académicas o actitudinales (sin inventar nombres)."],
    )

    clear_merged_row_text(
        doc.tables[7],
        1,
        ["[PENDIENTE] Registrar nuevas acciones tras el análisis del primer parcial."],
    )
    set_cell_text(doc.tables[8].rows[0].cells[1], DOCENTE, size=10)
    set_cell_text(doc.tables[8].rows[1].cells[1], "[PENDIENTE]", size=10)

    # Parte 3
    t9 = doc.tables[9]
    for c in t9.rows[0].cells[1:]:
        set_cell_text(c, DOCENTE, size=10)
    for c in t9.rows[1].cells[1:]:
        set_cell_text(c, asignatura, size=10)
    set_cell_text(t9.rows[2].cells[1], grupo, size=10)
    set_cell_text(t9.rows[2].cells[3], periodo, size=10)

    set_cell_text(doc.tables[10].rows[0].cells[1], "[PENDIENTE]", size=10)
    set_cell_text(doc.tables[10].rows[1].cells[1], "[PENDIENTE]", size=10)
    set_cell_text(doc.tables[10].rows[2].cells[1], "[PENDIENTE]", size=10)

    for ti, msg in [
        (11, "[PENDIENTE] Completar al cierre del curso (post Parcial 2 / PI)."),
        (12, "[PENDIENTE] Reportar aspectos del microcurrículo inconclusos, si aplica."),
        (13, "[PENDIENTE] Registrar estrategias implementadas y lecciones aprendidas."),
        (14, "[PENDIENTE] Recomendaciones para el Programa / próxima oferta."),
    ]:
        clear_merged_row_text(doc.tables[ti], 1, [msg])

    set_cell_text(doc.tables[15].rows[0].cells[1], DOCENTE, size=10)
    set_cell_text(doc.tables[15].rows[1].cells[1], "[PENDIENTE]", size=10)

    doc.save(str(out_path))
    print("OK institucional ->", out_path)


def build_prueba_estudiante(
    *,
    out_path: Path,
    asignatura: str,
    codigo: str,
    grupo: str,
    periodo: str,
    horario: str,
    intro: str,
    secciones: list[dict],
):
    """Instrumento de aplicación en Clase 1 (Kit docente). No va a Clases/ hasta el día."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Cabecera marca
    add_p(
        doc,
        "Institución Universitaria Antonio José Camacho — UNIAJC",
        size=10,
        bold=True,
        color=NAVY,
        space_after=2,
    )
    add_p(doc, "Prueba diagnóstica de entrada (Clase 1)", size=16, bold=True, color=NAVY, space_after=4)
    add_p(doc, f"{asignatura} · {codigo}", size=12, bold=True, color=CIAN, space_after=2)
    add_p(
        doc,
        f"Grupo: {grupo}  ·  Periodo: {periodo}  ·  {horario}",
        size=10,
        bold=True,
        color=GRAY,
        space_after=2,
    )
    add_p(doc, f"Docente: {DOCENTE} · {CORREO}", size=9, color=GRAY, space_after=8)

    # Meta estudiante
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    meta = [
        ("Nombre completo:", "________________________________"),
        ("Código / documento:", "________________________________"),
        ("Fecha (Clase 1):", "________________  ·  Duración: ~25–30 min"),
    ]
    for i, (a, b) in enumerate(meta):
        set_cell_text(t.rows[i].cells[0], a, bold=True, size=10)
        shade_cell(t.rows[i].cells[0], "E8F4FA")
        set_cell_text(t.rows[i].cells[1], b, size=10)

    add_p(doc, "", space_after=4)
    add_p(doc, "Propósito", size=12, bold=True, color=NAVY, space_after=2)
    add_p(
        doc,
        intro
        + " No afecta la nota del corte; sirve para ajustar el ritmo de las primeras clases y el Proyecto Integrador.",
        size=10,
        space_after=6,
    )
    add_p(
        doc,
        "Indicaciones: responda con letra clara. En ítems de selección marque una sola opción. "
        "En ejercicios cortos escriba SQL/diagrama/texto según se pida. Herramientas de clase: gratis + navegador (sin tarjeta).",
        size=9,
        space_after=8,
    )

    qnum = 1
    for sec in secciones:
        add_p(doc, sec["titulo"], size=12, bold=True, color=NAVY, space_after=4)
        for item in sec["items"]:
            add_p(doc, f"{qnum}. {item['enunciado']}", size=10, bold=True, space_after=2)
            if item.get("opciones"):
                for op in item["opciones"]:
                    add_p(doc, f"   ○  {op}", size=10, space_after=1)
            if item.get("espacio"):
                for _ in range(item["espacio"]):
                    add_p(doc, "_" * 78, size=9, color=RGBColor(0x8F, 0x98, 0x9D), space_after=1)
            qnum += 1
        add_p(doc, "", space_after=4)

    add_p(doc, "Cierre reflexivo (2–3 líneas)", size=12, bold=True, color=NAVY, space_after=2)
    add_p(
        doc,
        f"{qnum}. ¿Qué tema de {asignatura} te genera más expectativa o temor, y por qué?",
        size=10,
        bold=True,
        space_after=2,
    )
    for _ in range(3):
        add_p(doc, "_" * 78, size=9, color=RGBColor(0x8F, 0x98, 0x9D), space_after=1)

    add_p(doc, "", space_after=8)
    add_p(
        doc,
        "Uso docente: aplicar en Clase 1 (tras Presentación del curso / Padlet). "
        "Archivo institucional de seguimiento: Entregas docente/2026-2/DIAGNOSTICO…. "
        "No publicar en carpeta Clases/ hasta el día de aplicación.",
        size=8,
        color=RGBColor(0x8F, 0x98, 0x9D),
        space_after=2,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print("OK prueba ->", out_path)


def write_csv_bom(path: Path, rows: list[dict], fieldnames: list[str]):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for r in rows:
            w.writerow(r)
    print("OK csv ->", path)


FIELDS = [
    "curso",
    "codigo_fi",
    "grupo",
    "clase_n",
    "fecha",
    "dia",
    "hora_inicio",
    "hora_fin",
    "tipo_clase",
    "es_parcial",
    "parcial_n",
    "sesion_etiqueta",
    "tema",
    "notas",
]


def load_csv(path: Path) -> list[dict]:
    with open(path, encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def normalize_notas(rows: list[dict], *, clase1_tema: str | None = None) -> list[dict]:
    out = []
    for r in rows:
        rr = dict(r)
        notas = (rr.get("notas") or "").strip()
        # normalize pendiente listado
        if "pendiente listado" in notas.lower() or "[pendiente listado]" in notas.lower():
            # keep other notes, ensure [PENDIENTE listado]
            parts = [p.strip() for p in notas.replace("[PENDIENTE listado]", "").replace("pendiente listado", "").split(";") if p.strip()]
            parts.append("[PENDIENTE listado]")
            rr["notas"] = "; ".join(parts)
        elif not notas:
            rr["notas"] = "[PENDIENTE listado]"
        else:
            if "[PENDIENTE listado]" not in notas and "pendiente listado" not in notas.lower():
                rr["notas"] = notas.rstrip("; ") + "; [PENDIENTE listado]"
            else:
                rr["notas"] = notas
        if clase1_tema and str(rr.get("clase_n")) == "1":
            rr["tema"] = clase1_tema
        out.append(rr)
    return out


def patch_text_file(path: Path, replacements: list[tuple[str, str]]):
    text = path.read_text(encoding="utf-8")
    orig = text
    for a, b in replacements:
        if a not in text:
            print(f"WARN missing in {path.name}: {a[:60]!r}")
        else:
            text = text.replace(a, b)
    if text != orig:
        path.write_text(text, encoding="utf-8")
        print("OK patch ->", path)
    else:
        print("SKIP (no change) ->", path)


def main():
    # ========== A) Diagnósticos ==========
    bd2_dir = ROOT / "Bases de Datos II"
    arq_dir = ROOT / "Arquitectura de Sistemas Computacionales"

    prerreq_bd2 = [
        "Modelo entidad-relación (entidades, atributos, relaciones, cardinalidad).",
        "Normalización básica (1FN–3FN) y claves (primaria, foránea, candidata).",
        "SQL fundamental: DDL/DML (CREATE/ALTER, INSERT/UPDATE/DELETE) y consultas SELECT con WHERE, JOIN, GROUP BY.",
        "Integridad referencial y restricciones (NOT NULL, UNIQUE, CHECK).",
        "Competencia: traducir un caso sencillo de negocio a un esquema relacional y consultas básicas en navegador (sin instalar SGBD local).",
    ]
    acciones_bd2 = [
        "Aplicar prueba diagnóstica en Clase 1 (tras Presentación del curso / Padlet) y registrar hallazgos en este formato.",
        "Reforzar SQL y modelo relacional en las primeras clases regulares con laboratorios en navegador (DB Fiddle / OneCompiler / Live SQL).",
        "Metodología ABPr con Proyecto Integrador continuo de gestión avanzada de BD (seguridad, procedimientos, tuning).",
        "Si el diagnóstico muestra debilidad en JOINs/normalización: micro-refuerzos de 10–15 min al inicio de Clases 2–4.",
    ]

    prerreq_arq = [
        "Fundamentos de sistemas computacionales: hardware, software, SO y redes básicas (IP, DNS, cliente-servidor).",
        "Noción introductoria de servicios en internet / nube (almacenamiento, apps web) sin exigir experiencia en un proveedor específico.",
        "Lectura de diagramas simples de componentes o despliegue (cajas y flechas).",
        "Competencia: explicar, a nivel conceptual, cómo una aplicación se separa en capas (presentación, lógica, datos) y por qué importa la escalabilidad.",
    ]
    acciones_arq = [
        "Aplicar prueba diagnóstica en Clase 1 (tras Presentación del curso / Padlet) y registrar hallazgos en este formato.",
        "Arrancar con vocabulario cloud (IaaS/PaaS/SaaS) usando analogías y diagramas en navegador (draw.io / Excalidraw).",
        "Metodología ABPr con Proyecto Integrador de arquitectura cloud simulada (sin AWS/GCP/Oracle Cloud con tarjeta).",
        "Si el diagnóstico muestra debilidad en redes/capas: micro-refuerzos conceptuales al inicio de Clases 2–4 + labs Killercoda / Play with Docker solo en browser.",
    ]

    build_institucional_diagnostico(
        out_path=bd2_dir / "Entregas docente" / "2026-2" / "DIAGNOSTICO - Bases de Datos II - 2026-2.docx",
        asignatura="Bases de Datos II",
        grupo="641A-2",
        periodo="2026-2",
        prerrequisitos=prerreq_bd2,
        acciones_previstas=acciones_bd2,
    )
    build_institucional_diagnostico(
        out_path=arq_dir / "Entregas docente" / "2026-2" / "DIAGNOSTICO - Arquitectura de Sistemas Computacionales - 2026-2.docx",
        asignatura="Arquitectura de Sistemas Computacionales",
        grupo="6303C",
        periodo="2026-2",
        prerrequisitos=prerreq_arq,
        acciones_previstas=acciones_arq,
    )

    # Instrumentos estudiante (Kit docente / Clase 1) — patrón Prog II
    build_prueba_estudiante(
        out_path=bd2_dir / "Kit docente" / "Clase 1" / "Prueba Diagnostica - Bases de Datos II.docx",
        asignatura="Bases de Datos II",
        codigo="FI303215",
        grupo="641A-2",
        periodo="2026-2",
        horario="Lunes 18:00–20:00",
        intro=(
            "Esta prueba explora prerrequisitos de Bases de Datos I (modelo relacional, SQL básico e integridad) "
            "antes de avanzar a administración, procedimientos, seguridad y optimización."
        ),
        secciones=[
            {
                "titulo": "A. Modelo relacional y diseño (BD I)",
                "items": [
                    {
                        "enunciado": "En un modelo ER, una relación N:M entre ESTUDIANTE y CURSO típicamente se implementa en el modelo relacional como:",
                        "opciones": [
                            "Una sola tabla con ambos identificadores sin claves foráneas",
                            "Una tabla intermedia (asociación) con FKs a ambas entidades",
                            "Dos tablas sin ninguna relación física",
                            "Solo con un atributo multivaluado en ESTUDIANTE",
                        ],
                    },
                    {
                        "enunciado": "Explique en 2–3 líneas la diferencia entre clave primaria y clave foránea.",
                        "espacio": 3,
                    },
                    {
                        "enunciado": "Una tabla está en 1FN si:",
                        "opciones": [
                            "No tiene claves foráneas",
                            "Todos los atributos son atómicos (sin grupos repetitivos)",
                            "Está indexada",
                            "Usa únicamente tipos numéricos",
                        ],
                    },
                ],
            },
            {
                "titulo": "B. SQL fundamental",
                "items": [
                    {
                        "enunciado": "Escriba un SELECT que liste nombre y correo de la tabla cliente donde ciudad = 'Cali' (sintaxis estándar).",
                        "espacio": 4,
                    },
                    {
                        "enunciado": "¿Qué hace un INNER JOIN entre pedido y cliente sobre cliente_id?",
                        "opciones": [
                            "Devuelve todos los clientes aunque no tengan pedidos",
                            "Devuelve solo filas con coincidencia en ambas tablas",
                            "Elimina pedidos huérfanos automáticamente",
                            "Crea un índice compuesto",
                        ],
                    },
                    {
                        "enunciado": "Indique si cada sentencia es DDL o DML: (a) CREATE TABLE  (b) UPDATE  (c) ALTER TABLE  (d) INSERT",
                        "espacio": 3,
                    },
                ],
            },
            {
                "titulo": "C. Integridad y administración intro",
                "items": [
                    {
                        "enunciado": "Si existe FK de detalle_pedido.producto_id → producto.id, ¿qué suele impedir un DELETE de un producto referenciado?",
                        "opciones": [
                            "Un trigger de auditoría",
                            "La integridad referencial (restricción de FK)",
                            "El comando COMMIT",
                            "La normalización 2FN",
                        ],
                    },
                    {
                        "enunciado": "Caso corto: una tienda necesita Producto(id, nombre, precio) y Venta(id, fecha, producto_id, cantidad). Dibuje o liste las tablas con PK/FK y escriba un JOIN que sume cantidad vendida por producto.",
                        "espacio": 6,
                    },
                ],
            },
        ],
    )

    build_prueba_estudiante(
        out_path=arq_dir / "Kit docente" / "Clase 1" / "Prueba Diagnostica - Arquitectura de Sistemas Computacionales.docx",
        asignatura="Arquitectura de Sistemas Computacionales",
        codigo="FI303380",
        grupo="6303C",
        periodo="2026-2",
        horario="Lunes 10:00–12:00",
        intro=(
            "Esta prueba explora fundamentos de sistemas (capas, cliente-servidor, redes básicas) "
            "y nociones introductorias de servicios en la nube, antes de IaaS/PaaS/SaaS, virtualización y arquitecturas distribuidas."
        ),
        secciones=[
            {
                "titulo": "A. Fundamentos de sistemas",
                "items": [
                    {
                        "enunciado": "En una arquitectura en capas (presentación / lógica / datos), ¿dónde ubica típicamente las reglas de negocio?",
                        "opciones": [
                            "Solo en el navegador del usuario",
                            "En la capa de lógica (servicios / backend)",
                            "Únicamente en el SGBD como índices",
                            "En el cableado de red",
                        ],
                    },
                    {
                        "enunciado": "Explique con un ejemplo cotidiano qué es un modelo cliente-servidor.",
                        "espacio": 3,
                    },
                    {
                        "enunciado": "Relacione: (1) DNS  (2) IP  (3) HTTP — con: resolución de nombres / dirección de host / protocolo de aplicación web.",
                        "espacio": 3,
                    },
                ],
            },
            {
                "titulo": "B. Nube e infraestructura intro",
                "items": [
                    {
                        "enunciado": "¿Cuál afirmación describe mejor un servicio tipo SaaS?",
                        "opciones": [
                            "Alquilo máquinas virtuales y administro el SO completo",
                            "Uso una aplicación lista (p. ej. correo web) sin gestionar servidores",
                            "Solo alquilo el cableado del datacenter",
                            "Es sinónimo exclusivo de virtualización de escritorio",
                        ],
                    },
                    {
                        "enunciado": "Diferencie en 2–3 líneas máquina virtual vs contenedor (idea general, sin comandos).",
                        "espacio": 3,
                    },
                    {
                        "enunciado": "Mencione una ventaja y un riesgo de desplegar una app en la nube (seguridad, costo o disponibilidad).",
                        "espacio": 3,
                    },
                ],
            },
            {
                "titulo": "C. Lectura de arquitectura",
                "items": [
                    {
                        "enunciado": "Un sistema web tiene: navegador → API → base de datos. Dibuje o liste los componentes y una flecha de dependencia. Señale un posible cuello de botella.",
                        "espacio": 5,
                    },
                    {
                        "enunciado": "¿Para qué sirve, en una frase, un balanceador de carga?",
                        "opciones": [
                            "Cifrar discos duros",
                            "Repartir tráfico entre varias instancias del servicio",
                            "Reemplazar por completo a la base de datos",
                            "Compilar el código fuente",
                        ],
                    },
                ],
            },
        ],
    )

    # ========== Actualizar planes / calendarios / CSV ==========
    tema_bd2 = "Presentación del curso · Diagnóstico · Revisión de Bases de Datos I"
    tema_arq = "Presentación del curso · Diagnóstico · Introducción a arquitecturas cloud"

    patch_text_file(
        bd2_dir / "Plan curso" / "2026-2" / "PLAN_DE_CURSO_2026-2.md",
        [
            (
                "| 1 | 10/08/2026 | Presencial | Presentación del curso · Revisión de Bases de Datos I |",
                f"| 1 | 10/08/2026 | Presencial | {tema_bd2} |",
            )
        ],
    )
    patch_text_file(
        arq_dir / "Plan curso" / "2026-2" / "PLAN_DE_CURSO_2026-2.md",
        [
            (
                "| 1 | 10/08/2026 | Presencial | Presentación del curso · Introducción a arquitecturas cloud |",
                f"| 1 | 10/08/2026 | Presencial | {tema_arq} |",
            )
        ],
    )

    for cal_path, nota in [
        (bd2_dir / "Plan curso" / "2026-2" / "CALENDARIO_2026-2.md", "Presentación + Diagnóstico + revisión BD I"),
        (arq_dir / "Plan curso" / "2026-2" / "CALENDARIO_2026-2.md", "Presentación + Diagnóstico + intro cloud"),
    ]:
        patch_text_file(
            cal_path,
            [
                ("| 1 | 10/08/2026 | Presencial | — |", f"| 1 | 10/08/2026 | Presencial | {nota} |"),
            ],
        )

    # CSVs BD II / Arq
    for curso_path, tema, cfg_name in [
        (bd2_dir / "Plan curso" / "2026-2" / "calendario_eventos_2026-2.csv", tema_bd2, "eventos_bases_datos_ii_2026-2.csv"),
        (arq_dir / "Plan curso" / "2026-2" / "calendario_eventos_2026-2.csv", tema_arq, "eventos_arquitectura_2026-2.csv"),
    ]:
        rows = normalize_notas(load_csv(curso_path), clase1_tema=tema)
        write_csv_bom(curso_path, rows, FIELDS)
        write_csv_bom(ROOT / ".config" / "calendario" / cfg_name, rows, FIELDS)

    # ========== B) CSV cursos viejos (Prog II / Seminario) ==========
    for curso_path, cfg_name in [
        (
            ROOT / "Programacion II" / "Plan curso" / "2026-1" / "calendario_eventos_2026-1.csv",
            "eventos_programacion_ii_2026-1.csv",
        ),
        (
            ROOT / "Seminario de Sistemas" / "Plan curso" / "2026-1" / "calendario_eventos_2026-1.csv",
            "eventos_seminario_2026-1.csv",
        ),
    ]:
        rows = normalize_notas(load_csv(curso_path))
        write_csv_bom(curso_path, rows, FIELDS)
        write_csv_bom(ROOT / ".config" / "calendario" / cfg_name, rows, FIELDS)

    # Refresh eventos_todos if exists — rebuild from four CSVs
    todos = []
    for p in [
        ROOT / ".config" / "calendario" / "eventos_programacion_ii_2026-1.csv",
        ROOT / ".config" / "calendario" / "eventos_seminario_2026-1.csv",
        ROOT / ".config" / "calendario" / "eventos_bases_datos_ii_2026-2.csv",
        ROOT / ".config" / "calendario" / "eventos_arquitectura_2026-2.csv",
    ]:
        todos.extend(load_csv(p))
    write_csv_bom(ROOT / ".config" / "calendario" / "eventos_todos_cursos_2026-2.csv", todos, FIELDS)

    # ========== Builds PPTX: patch Clase 1 tema ==========
    patch_text_file(
        ROOT / ".config" / "slides" / "build_uniajc_bd2_curso.py",
        [
            (
                '"tema": "Presentación del curso · Revisión de Bases de Datos I"',
                f'"tema": "{tema_bd2}"',
            ),
        ],
    )
    patch_text_file(
        ROOT / ".config" / "slides" / "build_uniajc_arq_curso.py",
        [
            (
                '"tema": "Presentación del curso · Introducción a arquitecturas cloud"',
                f'"tema": "{tema_arq}"',
            ),
        ],
    )

    # Reglas / agente / uniajc.json clase_1
    for path in [
        ROOT / ".cursor" / "rules" / "uniajc-docente.mdc",
        ROOT / ".claude" / "agents" / "disenador-curricular-uniajc.md",
        ROOT / ".cursor" / "agents" / "disenador-curricular-uniajc.md",
        ROOT / ".claude" / "agents" / "uniajc-dudas-material.md",
        ROOT / ".cursor" / "agents" / "uniajc-dudas-material.md",
    ]:
        if not path.exists():
            print("SKIP missing", path)
            continue
        patch_text_file(
            path,
            [
                (
                    "La **Clase 1** siempre incluye: (1) Presentación del curso (acuerdo, logística, Padlet, evaluación, cronograma) + (2) arranque temático de la primera unidad.",
                    "La **Clase 1** siempre incluye: (1) Presentación del curso (acuerdo, logística, Padlet, evaluación, cronograma) + (2) **Diagnóstico de entrada** + (3) arranque temático de la primera unidad.",
                ),
                (
                    "Guion/slides de Clase 1 = Presentación del Curso + primer bloque temático.",
                    "Guion/slides de Clase 1 = Presentación del Curso + Diagnóstico + primer bloque temático.",
                ),
                (
                    "Wording del plan/CONTENIDO: `Presentación del curso · [tema intro]`.",
                    "Wording del plan/CONTENIDO: `Presentación del curso · Diagnóstico · [tema intro]`.",
                ),
                (
                    "Siempre: Presentación del curso + arranque temático de la primera unidad. Wording: `Presentación del curso · [tema intro]`. Guion/slides Clase 1 = PPTX del curso + primer bloque temático.",
                    "Siempre: Presentación del curso + Diagnóstico + arranque temático de la primera unidad. Wording: `Presentación del curso · Diagnóstico · [tema intro]`. Guion/slides Clase 1 = PPTX del curso + Diagnóstico + primer bloque temático.",
                ),
            ],
        )

    # uniajc.json clase_1 block
    uj = ROOT / ".config" / "universidades" / "uniajc.json"
    patch_text_file(
        uj,
        [
            (
                '"_regla": "La Clase 1 siempre combina Presentación del curso + arranque temático de la primera unidad (no solo logística)."',
                '"_regla": "La Clase 1 siempre combina Presentación del curso + Diagnóstico de entrada + arranque temático de la primera unidad (no solo logística)."',
            ),
            (
                '"Presentación del curso (acuerdo, logística, Padlet, evaluación, cronograma)",\n        "Un poco del tema de la primera unidad"',
                '"Presentación del curso (acuerdo, logística, Padlet, evaluación, cronograma)",\n        "Diagnóstico de entrada (instrumento en Kit docente / registro en Entregas docente)",\n        "Un poco del tema de la primera unidad"',
            ),
            (
                '"guion_y_slides": "Guion/slides de Clase 1 = Presentación del Curso + primer bloque temático.",\n      "wording_plan": "Presentación del curso · [tema intro]"',
                '"guion_y_slides": "Guion/slides de Clase 1 = Presentación del Curso + Diagnóstico + primer bloque temático.",\n      "wording_plan": "Presentación del curso · Diagnóstico · [tema intro]"',
            ),
        ],
    )

    # README eventos
    patch_text_file(
        ROOT / ".config" / "calendario" / "README_eventos_csv.md",
        [
            (
                "Archivos `eventos_*_2026-2.csv` (copia en `.config/calendario/`) y `<Curso>/Plan curso/2026-1|2026-2/calendario_eventos_….csv`: 15 filas/clase por curso, UTF-8 con BOM.",
                "Archivos `eventos_*_2026-1.csv` / `eventos_*_2026-2.csv` (copia en `.config/calendario/`) y `<Curso>/Plan curso/2026-1|2026-2/calendario_eventos_….csv`: 15 filas/clase por curso, UTF-8 con BOM. Notas: `[PENDIENTE listado]` hasta tener nómina.",
            ),
        ],
    )

    print("DONE generation/patches")


if __name__ == "__main__":
    main()
