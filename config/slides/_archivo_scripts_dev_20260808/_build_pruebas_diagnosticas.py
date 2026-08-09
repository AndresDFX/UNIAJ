# -*- coding: utf-8 -*-
"""Regenera pruebas diagnosticas Kit Clase 1 (BD II + Arquitectura)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"G:\Mi unidad\Trabajos\Empleo\UNIAJ\Cursos")
BLUE = RGBColor(0x09, 0x52, 0x92)


def set_run(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_p(doc, text, bold=False, size=11, space_after=4, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    set_run(r, bold=bold, size=size, color=color)
    return p


def add_opt(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("\u25cb  " + text)
    set_run(r, size=10)


def add_lines(doc, n=3):
    for _ in range(n):
        add_p(doc, "_" * 78, size=10, space_after=2)


def header_common(doc, curso, codigo, grupo, periodo, horario):
    add_p(doc, "Instituci\u00f3n Universitaria Antonio Jos\u00e9 Camacho \u2014 UNIAJC", bold=True, size=12, color=BLUE)
    add_p(doc, "Diagn\u00f3stico de conocimientos previos \u2014 no califica corte; sirve para adecuar el ritmo.", bold=True, size=12)
    add_p(doc, f"Prueba diagn\u00f3stica de entrada (Clase 1) \u00b7 {curso} \u00b7 {codigo}", bold=True, size=11)
    add_p(doc, f"Grupo: {grupo}  \u00b7  Periodo: {periodo}  \u00b7  {horario}", size=10)
    add_p(doc, "Docente: Julian Andres Casta\u00f1o Espinosa \u00b7 julianacastano@profesores.uniajc.edu.co", size=10)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    cells = t.rows[0].cells
    cells[0].text = "Nombre completo: ________________________________"
    cells[1].text = "C\u00f3digo / documento: ________________________________"
    add_p(doc, "Fecha (Clase 1): ________________  \u00b7  Duraci\u00f3n: ~25\u201330 min  \u00b7  12\u201314 \u00edtems", size=10, space_after=8)


def build_bd2():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    header_common(doc, "Bases de Datos II", "FI303215", "641A-2", "2026-2", "Lunes 18:00\u201320:00")
    add_p(doc, "Prop\u00f3sito", bold=True, size=11, color=BLUE)
    add_p(
        doc,
        "Esta prueba mide saberes previos de Bases de Datos I (modelo relacional, SQL b\u00e1sico, integridad y normalizaci\u00f3n). "
        "No eval\u00faa a\u00fan administraci\u00f3n avanzada, tuning ni el temario nuevo del semestre. "
        "Diagn\u00f3stico de conocimientos previos \u2014 no califica corte; sirve para adecuar el ritmo.",
        size=10,
        space_after=6,
    )
    add_p(doc, "Indicaciones: una sola opci\u00f3n en selecci\u00f3n; en abiertos escriba con letra clara. Herramientas de clase: gratis + navegador.", size=10, space_after=8)

    add_p(doc, "A. Modelo relacional y dise\u00f1o (BD I)", bold=True, size=11, color=BLUE)
    add_p(doc, "1. En un modelo ER, una relaci\u00f3n N:M entre ESTUDIANTE y CURSO t\u00edpicamente se implementa como:", bold=False)
    for o in [
        "Una sola tabla con ambos identificadores sin claves for\u00e1neas",
        "Una tabla intermedia (asociaci\u00f3n) con FKs a ambas entidades",
        "Dos tablas sin ninguna relaci\u00f3n f\u00edsica",
        "Solo con un atributo multivaluado en ESTUDIANTE",
    ]:
        add_opt(doc, o)
    add_p(doc, "2. Explique en 2\u20133 l\u00edneas la diferencia entre clave primaria y clave for\u00e1nea.")
    add_lines(doc, 3)
    add_p(doc, "3. Una tabla est\u00e1 en 1FN si:")
    for o in [
        "No tiene claves for\u00e1neas",
        "Todos los atributos son at\u00f3micos (sin grupos repetitivos)",
        "Est\u00e1 indexada",
        "Usa \u00fanicamente tipos num\u00e9ricos",
    ]:
        add_opt(doc, o)
    add_p(doc, "4. Indique en una frase qu\u00e9 busca la 3FN (idea general, sin formalismo).")
    add_lines(doc, 2)

    add_p(doc, "B. SQL fundamental (BD I)", bold=True, size=11, color=BLUE)
    add_p(doc, "5. Escriba un SELECT que liste nombre y correo de la tabla cliente donde ciudad = 'Cali'.")
    add_lines(doc, 3)
    add_p(doc, "6. \u00bfQu\u00e9 hace un INNER JOIN entre pedido y cliente sobre cliente_id?")
    for o in [
        "Devuelve todos los clientes aunque no tengan pedidos",
        "Devuelve solo filas con coincidencia en ambas tablas",
        "Elimina pedidos hu\u00e9rfanos autom\u00e1ticamente",
        "Crea un \u00edndice compuesto",
    ]:
        add_opt(doc, o)
    add_p(doc, "7. Clasifique DDL o DML: (a) CREATE TABLE  (b) UPDATE  (c) ALTER TABLE  (d) INSERT")
    add_lines(doc, 2)
    add_p(doc, "8. Escriba un ejemplo de agregaci\u00f3n: contar pedidos por cliente (SELECT \u2026 GROUP BY).")
    add_lines(doc, 3)

    add_p(doc, "C. Integridad y restricciones (BD I)", bold=True, size=11, color=BLUE)
    add_p(doc, "9. Si existe FK detalle_pedido.producto_id \u2192 producto.id, \u00bfqu\u00e9 suele impedir un DELETE del producto referenciado?")
    for o in [
        "Un trigger de auditor\u00eda",
        "La integridad referencial (restricci\u00f3n de FK)",
        "El comando COMMIT",
        "La normalizaci\u00f3n 2FN",
    ]:
        add_opt(doc, o)
    add_p(doc, "10. Relacione: NOT NULL / UNIQUE / CHECK \u2014 con: valor obligatorio / sin duplicados / regla de dominio.")
    add_lines(doc, 2)
    add_p(doc, "11. \u00bfQu\u00e9 problema de integridad aparece si hay pedidos con cliente_id que no existe en cliente?")
    add_lines(doc, 2)

    add_p(doc, "D. Caso corto (BD I)", bold=True, size=11, color=BLUE)
    add_p(doc, "12. Una tienda necesita Producto(id, nombre, precio) y Venta(id, fecha, producto_id, cantidad). Liste PK/FK y un JOIN que sume cantidad por producto.")
    add_lines(doc, 5)
    add_p(doc, "13. En una frase: diferencia entre entidad d\u00e9bil y entidad fuerte (idea de BD I).")
    add_lines(doc, 2)

    add_p(doc, "Cierre (metacognici\u00f3n \u2014 no califica)", bold=True, size=11, color=BLUE)
    add_p(doc, "14. \u00bfQu\u00e9 tema de BD I te sientes m\u00e1s s\u00f3lido y cu\u00e1l necesitas reforzar? (2\u20133 l\u00edneas)")
    add_lines(doc, 3)

    add_p(
        doc,
        "Uso docente: aplicar en Clase 1 (tras Presentaci\u00f3n del curso / Padlet). "
        "Registro institucional: Entregas docente/2026-2/DIAGNOSTICO\u2026. No publicar en Clases/ hasta el d\u00eda de aplicaci\u00f3n.",
        size=9,
        space_after=2,
    )
    out = ROOT / "Bases de Datos II" / "Kit docente" / "Clase 1" / "Prueba Diagnostica - Bases de Datos II.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("OK", out)


def build_arq():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    header_common(doc, "Arquitectura de Sistemas Computacionales", "FI303380", "6303C", "2026-2", "Lunes 10:00\u201312:00")
    add_p(doc, "Prop\u00f3sito", bold=True, size=11, color=BLUE)
    add_p(
        doc,
        "Esta prueba mide saberes previos de fundamentos de sistemas, redes b\u00e1sicas, sistemas operativos y nociones introductorias de virtualizaci\u00f3n/servicios en internet. "
        "No eval\u00faa a\u00fan IaaS/PaaS/SaaS en profundidad, microservicios avanzados, CI/CD ni el temario nuevo del semestre. "
        "Diagn\u00f3stico de conocimientos previos \u2014 no califica corte; sirve para adecuar el ritmo.",
        size=10,
        space_after=6,
    )
    add_p(doc, "Indicaciones: una sola opci\u00f3n en selecci\u00f3n; en abiertos escriba texto o diagrama simple. Herramientas de clase: gratis + navegador.", size=10, space_after=8)

    add_p(doc, "A. Fundamentos de sistemas y capas", bold=True, size=11, color=BLUE)
    add_p(doc, "1. En una arquitectura en capas (presentaci\u00f3n / l\u00f3gica / datos), \u00bfd\u00f3nde ubica t\u00edpicamente las reglas de negocio?")
    for o in [
        "Solo en el navegador del usuario",
        "En la capa de l\u00f3gica (servicios / backend)",
        "\u00danicamente en el SGBD como \u00edndices",
        "En el cableado de red",
    ]:
        add_opt(doc, o)
    add_p(doc, "2. Explique con un ejemplo cotidiano qu\u00e9 es un modelo cliente-servidor.")
    add_lines(doc, 3)
    add_p(doc, "3. Mencione dos diferencias entre hardware y software en un sistema computacional.")
    add_lines(doc, 2)
    add_p(doc, "4. \u00bfPara qu\u00e9 sirve, en una frase, un sistema operativo?")
    add_lines(doc, 2)

    add_p(doc, "B. Redes b\u00e1sicas", bold=True, size=11, color=BLUE)
    add_p(doc, "5. Relacione: (1) DNS  (2) IP  (3) HTTP \u2014 con: resoluci\u00f3n de nombres / direcci\u00f3n de host / protocolo de aplicaci\u00f3n web.")
    add_lines(doc, 2)
    add_p(doc, "6. \u00bfQu\u00e9 describe mejor el puerto 443 en la web cotidiana?")
    for o in [
        "Puerto t\u00edpico de HTTPS (tr\u00e1fico web cifrado)",
        "Direcci\u00f3n MAC de la tarjeta de red",
        "Nombre del dominio principal",
        "Tama\u00f1o del disco duro",
    ]:
        add_opt(doc, o)
    add_p(doc, "7. En una frase: diferencia entre LAN e Internet (idea general).")
    add_lines(doc, 2)

    add_p(doc, "C. Virtualizaci\u00f3n e internet (intro)", bold=True, size=11, color=BLUE)
    add_p(doc, "8. \u00bfCu\u00e1l afirmaci\u00f3n describe mejor un servicio tipo SaaS a nivel introductorio?")
    for o in [
        "Alquilo m\u00e1quinas virtuales y administro el SO completo",
        "Uso una aplicaci\u00f3n lista (p. ej. correo web) sin gestionar servidores",
        "Solo alquilo el cableado del datacenter",
        "Es sin\u00f3nimo exclusivo de virtualizaci\u00f3n de escritorio",
    ]:
        add_opt(doc, o)
    add_p(doc, "9. Diferencie en 2\u20133 l\u00edneas m\u00e1quina virtual vs contenedor (idea general, sin comandos).")
    add_lines(doc, 3)
    add_p(doc, "10. Mencione una ventaja y un riesgo de desplegar una app en la nube (seguridad, costo o disponibilidad).")
    add_lines(doc, 2)

    add_p(doc, "D. Lectura de arquitectura simple", bold=True, size=11, color=BLUE)
    add_p(doc, "11. Un sistema web tiene: navegador \u2192 API \u2192 base de datos. Liste componentes y se\u00f1ale un posible cuello de botella.")
    add_lines(doc, 4)
    add_p(doc, "12. \u00bfPara qu\u00e9 sirve, en una frase, un balanceador de carga?")
    for o in [
        "Cifrar discos duros",
        "Repartir tr\u00e1fico entre varias instancias del servicio",
        "Reemplazar por completo a la base de datos",
        "Compilar el c\u00f3digo fuente",
    ]:
        add_opt(doc, o)
    add_p(doc, "13. Explique en una frase qu\u00e9 entiende por escalabilidad de un sistema.")
    add_lines(doc, 2)

    add_p(doc, "Cierre (metacognici\u00f3n \u2014 no califica)", bold=True, size=11, color=BLUE)
    add_p(doc, "14. \u00bfQu\u00e9 fundamento (redes, SO, capas, virtualizaci\u00f3n) te sientes m\u00e1s s\u00f3lido y cu\u00e1l necesitas reforzar?")
    add_lines(doc, 3)

    add_p(
        doc,
        "Uso docente: aplicar en Clase 1 (tras Presentaci\u00f3n del curso / Padlet). "
        "Registro institucional: Entregas docente/2026-2/DIAGNOSTICO\u2026. No publicar en Clases/ hasta el d\u00eda de aplicaci\u00f3n.",
        size=9,
        space_after=2,
    )
    out = ROOT / "Arquitectura de Sistemas Computacionales" / "Kit docente" / "Clase 1" / "Prueba Diagnostica - Arquitectura de Sistemas Computacionales.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("OK", out)


if __name__ == "__main__":
    build_bd2()
    build_arq()