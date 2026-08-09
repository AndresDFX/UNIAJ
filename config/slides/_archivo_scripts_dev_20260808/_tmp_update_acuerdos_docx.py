# -*- coding: utf-8 -*-
"""Update horario/modalidad in Acuerdos pedagogicos 2026-2 .docx."""
from pathlib import Path
from docx import Document

ROOT = Path(r"G:\Mi unidad\Trabajos\Empleo\UNIAJ\Cursos")

REPLACEMENTS = [
    (
        "Presencialidad asistida (Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)",
        "Virtual (clases y parciales síncronos · festivos autónomos)",
    ),
    (
        "Presencialidad asistida (Clase 1 presencial · resto virtual · parciales presencial)",
        "Virtual (clases y parciales síncronos)",
    ),
    ("modalidad del curso = Presencialidad asistida", "modalidad del curso = Virtual"),
    ("Modalidad: Presencialidad asistida", "Modalidad: Virtual"),
    ("Modalidad:Presencialidad asistida", "Modalidad: Virtual"),
    ("Presencialidad asistida (franja virtual síncrona)", "Virtual"),
    ("Presencialidad asistida (confirmar calendario virtual", "Virtual (confirmar calendario virtual"),
    ("Presencialidad asistida", "Virtual"),
    ("18:00–20:00", "20:00–22:00"),
    ("18:00 – 20:00", "20:00 – 22:00"),
    ("18:00-20:00", "20:00-22:00"),
]


def replace_in_paragraph(paragraph) -> bool:
    changed = False
    # Full-paragraph text first (runs may split)
    full = paragraph.text
    new_full = full
    for old, new in REPLACEMENTS:
        if old in new_full:
            new_full = new_full.replace(old, new)
    if new_full != full and paragraph.runs:
        # Put all text in first run; clear others to avoid duplication
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
        return True
    # Fallback: per-run
    for run in paragraph.runs:
        t = run.text
        n = t
        for old, new in REPLACEMENTS:
            if old in n:
                n = n.replace(old, new)
        if n != t:
            run.text = n
            changed = True
    return changed


def update_doc(path: Path, night: bool) -> None:
    doc = Document(str(path))
    reps = list(REPLACEMENTS)
    if not night:
        # Architecture: do not change 18->20 (none expected); keep 10:00
        reps = [r for r in reps if "18:00" not in r[0]]
    count = 0
    for p in doc.paragraphs:
        # temporarily swap global
        global REPLACEMENTS
        old_reps = REPLACEMENTS
        REPLACEMENTS = reps
        if replace_in_paragraph(p):
            count += 1
        REPLACEMENTS = old_reps
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    old_reps = REPLACEMENTS
                    REPLACEMENTS = reps
                    if replace_in_paragraph(p):
                        count += 1
                    REPLACEMENTS = old_reps
    # Also headers/footers
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                old_reps = REPLACEMENTS
                REPLACEMENTS = reps
                if replace_in_paragraph(p):
                    count += 1
                REPLACEMENTS = old_reps
    doc.save(str(path))
    print(f"OK {path.relative_to(ROOT)} paragraphs_touched~{count}")


docs = [
    (ROOT / "Programacion II" / "Entregas docente" / "2026-2", True),
    (ROOT / "Seminario de Sistemas" / "Entregas docente" / "2026-2", True),
    (ROOT / "Bases de Datos II" / "Entregas docente" / "2026-2", True),
    (ROOT / "Arquitectura de Sistemas Computacionales" / "Entregas docente" / "2026-2", False),
]
for folder, night in docs:
    for path in folder.glob("ACUERDO PEDAGOGICO*.docx"):
        update_doc(path, night=night)
