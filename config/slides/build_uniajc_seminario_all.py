# -*- coding: utf-8 -*-
"""Genera TODO el material de Seminario de Sistemas 2026-2 (PI VetCare, diseño).

Por que existe
--------------
El material del curso venia del periodo anterior y no era regenerable: cualquier cambio
de convencion habia que aplicarlo a mano en 14 carpetas, y los guiones traian la agenda
de un bloque que ya no existe. Este build lo deja alineado al plan 2026-2 y regenerable
desde una sola fuente, igual que Prog II, BD II y Arquitectura.

Diferencia clave con Programacion II: esta asignatura es de **analisis y diseño**, no de
programacion. El estudiante no entrega codigo: entrega requisitos, diagramas UML,
wireframes y documentos. Por eso donde Prog II genera `Codigo/*.java`, aqui se genera
`Plantillas/*.md` con el artefacto de diseño de la clase.

Salidas
-------
  Clases/Clase N - <slug>/Presentacion.pptx      (estudiante)
  Clases/Clase N - <slug>/Taller PI - Clase N - VetCare.docx
  Kit docente/Clase N/Guion Docente ….md|.docx   (regla de oro: teoria desarrollada)
  Kit docente/Clase N/Quiz ….docx + Quiz … CLAVE DOCENTE.docx
  Kit docente/Clase N/Solucion Taller ….md|.docx
  Kit docente/Clase N/Plantillas/<artefacto>.md
  Kit docente/Clase N/Guia aplicacion Parcial N ….md|.docx   (dias 5/10/15)

Los datos pedagogicos viven en `seminario_clases_data.py` (una sola fuente).
"""
from __future__ import annotations

import os
import re
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
SLIDES = Path(__file__).resolve().parent
sys.path.insert(0, str(SLIDES))

from uniajc_slides_engine import (  # noqa: E402
    block_timeline_slide,
    box_note_slide,
    checklist_slide,
    class_cover,
    closing_slide,
    content_slide,
    herramientas_slide,
    new_prs,
    pseudo_code_slide,
    steps_visual_slide,
)
from uniajc_quiz_helpers import clave_text, q_abierta, q_om, q_vf  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt as DocPt, RGBColor  # noqa: E402

from seminario_clases_data import CLASES  # noqa: E402

CURSO = ROOT / "Seminario de Sistemas"
CLASES_DIR = CURSO / "Clases"
KIT_DIR = CURSO / "Kit docente"

AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN_D = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xA0, 0x20, 0x30)
FONT = "Calibri"
EXAMLAB = "ExamLab (examlab.lovable.app/app)"

# Logo + para que sirve, por herramienta. El campo `herramienta` de cada clase es
# texto libre separado por "·", asi que se mapea por subcadena. Figma y Penpot van
# sin logo a proposito: no hay asset de marca y no se inventa uno.
HERRAMIENTAS = [
    ("draw.io", "drawio.png", "Diagramas UML y de contexto"),
    ("excalidraw", "excalidraw.png", "Bocetos rapidos a mano alzada"),
    ("google docs", "google_docs.png", "Documentos del paquete de diseño"),
    ("mermaid", "mermaid.png", "Diagramas como texto versionable"),
    ("figma", None, "Wireframes y prototipo navegable"),
    ("penpot", None, "Alternativa libre a Figma"),
]


def _herramientas_del_dia(texto):
    """Convierte 'draw.io · Mermaid Live Editor' en items para herramientas_slide."""
    t = (texto or "").lower()
    items = [{"name": n.title() if n.islower() and " " not in n else n,
              "logo": logo, "note": nota}
             for n, logo, nota in HERRAMIENTAS if n in t]
    # ExamLab siempre aparece: es el canal de entrega de todos los talleres
    items.append({"name": "ExamLab", "logo": "examlab.png",
                  "note": "Entrega del taller · domingo 23:59"})
    return items


PARCIALES = {
    5: ("Parcial 1", "Parcial 1 - Ciclos de vida y metodologias.docx"),
    10: ("Parcial 2", "Parcial 2 - Requerimientos UML y casos de uso.docx"),
    15: ("Parcial 3", "Parcial 3 - UML avanzado interfaces y proyecto.docx"),
}


# ----------------------------------------------------------------- helpers docx
def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def run(r, *, size=11, bold=False, color=GRIS):
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = DocPt(size)
    r.bold = bold
    r.font.color.rgb = color


def para(doc, text, *, size=11, bold=False, color=GRIS, space_after=6, shade_fill=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DocPt(space_after)
    if shade_fill:
        shade(p, shade_fill)
    r = p.add_run(text)
    run(r, size=size, bold=bold, color=color)
    return p


def banda(doc, text):
    return para(doc, "  " + text, size=13, bold=True, color=BLANCO,
                shade_fill="095292", space_after=8)


def add_inline_docx(p, text, *, size=11, color=GRIS):
    """Soporta @@negrita@@, misma convencion que el motor de slides."""
    for part in re.split(r"(@@.*?@@)", text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith("@@") and part.endswith("@@"):
            r.text = part[2:-2]
            run(r, size=size, bold=True, color=color)
        else:
            r.text = part
            run(r, size=size, color=color)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = DocPt(2)
        add_inline_docx(p, str(it))


def _resumen(bullets_, max_chars=115):
    """La teoria del guion es larga a proposito; la slide lleva solo la idea central."""
    out = []
    for b in bullets_:
        first = re.split(r"(?<=[a-záéíóúü0-9\)])\.\s", b, maxsplit=1)[0].strip()
        if len(first) > max_chars or len(first) < 12:
            base = b if len(first) < 12 else first
            first = base[:max_chars].rsplit(" ", 1)[0].rstrip(":,;") + "…"
        out.append(first)
    return out


def _quiz_items(c):
    out = []
    for q in c.get("quiz", []):
        t = q.get("tipo")
        if t == "om":
            out.append(q_om(q["q"], q.get("opciones", []), q.get("clave", "")))
        elif t == "vf":
            out.append(q_vf(q["q"], q.get("clave", "V")))
        else:
            out.append(q_abierta(q["q"], q.get("clave", "")))
    return out


# ----------------------------------------------------------------------- slides
def build_pptx(c):
    n = c["n"]
    if n in PARCIALES:
        prs = new_prs()
        class_cover(prs, PARCIALES[n][0], subtitulo="Solo evaluacion", clase_n=n, idx=1)
        content_slide(prs, "Indicaciones", [
            "Hoy es **solo Parcial** (presencial sincrono).",
            "No hay tema nuevo ni taller del PI en esta sesion.",
            "Duracion sugerida: **90–100 min** dentro del bloque de 120.",
            "La preparacion del PI continua en la siguiente clase regular.",
        ], idx=2)
        closing_slide(prs, f"{PARCIALES[n][0]} · Clase {n}",
                      ["Enfocados en la evaluacion del corte",
                       "El PI VetCare continua la proxima clase"],
                      accent="Solo evaluacion")
        out_dir = CLASES_DIR / f"Clase {n} - {PARCIALES[n][0]}"
        out_dir.mkdir(parents=True, exist_ok=True)
        prs.save(str(out_dir / "Presentacion.pptx"))
        print("PPTX", out_dir.name)
        return

    prs = new_prs()
    class_cover(prs, c["titulo"], subtitulo=c["subtitulo"], clase_n=n, idx=1)
    idx = 2
    content_slide(prs, "Encuadre de hoy · Objetivo del PI", [
        f"**Hoy avanzamos VetCare en:** {c['hito_pi']}",
        f"Herramienta: **{c['herramienta']}** · Bloque **120 min**",
        f"**Entregable de hoy:** {c['entregable']}",
        "La teoria esta al servicio del producto: al salir, VetCare avanzo.",
    ], idx=idx); idx += 1
    block_timeline_slide(prs, "Mapa del bloque de hoy (120 min)", [
        ("0-10", "Encuadre y repaso del avance"),
        ("10-40", "Teoria Core del tema de hoy"),
        ("40-60", "Demo en vivo sobre VetCare"),
        ("60-105", "Taller guiado = avance del PI"),
        ("105-120", "Criterios de exito y cierre"),
    ], idx=idx); idx += 1
    content_slide(prs, "Teoria Core", _resumen(c["teoria"]), idx=idx, size=15); idx += 1
    if c.get("codigo_slide_lineas"):
        pseudo_code_slide(prs, c.get("codigo_slide_titulo", "Codigo de hoy"),
                          c["codigo_slide_lineas"],
                          caption=c.get("codigo_slide_caption"), idx=idx); idx += 1
    content_slide(prs, "Demo del dia", [
        f"**Herramienta:** {c['herramienta']}",
        f"**Demo:** {c['demo']}",
        "Mismo dominio VetCare — no otro ejemplo. Aqui se diseña, no se programa.",
    ], idx=idx); idx += 1
    herramientas_slide(prs, _herramientas_del_dia(c["herramienta"]),
                       title="Herramientas de hoy",
                       sub="Gratis · funcionan en el navegador", idx=idx); idx += 1
    if c.get("contexto"):
        content_slide(prs, "Taller PI — por que importa", c["contexto"], idx=idx, size=16); idx += 1
    if c.get("escenario"):
        content_slide(prs, "Taller PI — punto de partida", c["escenario"], idx=idx, size=16); idx += 1
    steps_visual_slide(prs, "Taller PI — pasos guiados",
                       [(t, "") for t in c["taller"]], idx=idx); idx += 1
    if c.get("pistas"):
        checklist_slide(prs, "Antes de entregar — autochequeo", c["pistas"], idx=idx); idx += 1
    content_slide(prs, "Criterios de exito / entregable", [
        f"**Entregable:** {c['entregable']}",
        *[f"@@Exito:@@ {x}" for x in c.get("criterios", [])],
        f"@@Entrega en {EXAMLAB}@@ — domingo 23:59.",
    ], idx=idx); idx += 1
    box_note_slide(prs, "Para el PI esta semana", [
        ("info", f"Hito: {c['hito_pi']}"),
        ("aclaracion", "Enunciado completo: Clases/Proyecto Integrador/ (VetCare)."),
        ("advertencia", f"Entrega del taller en {EXAMLAB} · domingo 23:59."),
    ], idx=idx); idx += 1
    closing_slide(prs, f"Clase {n} · VetCare avanza", [
        c["hito_pi"],
        f"Entregable: {c['entregable']}",
        "Siguiente clase: continuamos el hilo del PI",
    ], accent="Teoria al servicio del proyecto")
    out_dir = CLASES_DIR / f"Clase {n} - {c['slug']}"
    out_dir.mkdir(parents=True, exist_ok=True)
    prs.save(str(out_dir / "Presentacion.pptx"))
    print("PPTX", out_dir.name)


# ---------------------------------------------------------------------- taller
def build_taller_docx(c):
    n = c["n"]
    if n in PARCIALES:
        return
    doc = Document()
    banda(doc, f"Taller PI · Clase {n} · Seminario de Sistemas")
    para(doc, c["titulo"], size=14, bold=True, color=AZUL)
    para(doc, "Hilo conductor: Proyecto Integrador VetCare — diseño, no programacion.",
         size=11, bold=True)
    para(doc, f"Herramienta: {c['herramienta']}")
    para(doc, f"Hoy avanzamos el PI en: {c['hito_pi']}", shade_fill="FFF8D6")
    para(doc, "1. Contexto / por que importa al PI", size=12, bold=True, color=AZUL)
    bullets(doc, c.get("contexto") or ["Trabaje sobre el VetCare de su equipo."])
    para(doc, "2. Punto de partida", size=12, bold=True, color=AZUL)
    bullets(doc, c.get("escenario") or ["Use el codigo que ya tiene del avance anterior."])
    para(doc, "3. Pasos guiados", size=12, bold=True, color=AZUL)
    bullets(doc, c["taller"])
    para(doc, "4. Entregable", size=12, bold=True, color=AZUL)
    para(doc, c["entregable"], shade_fill="E8F4FA")
    para(doc, "5. Criterios de exito", size=12, bold=True, color=AZUL)
    bullets(doc, c.get("criterios") or ["Avance real y verificable del VetCare del equipo."])
    para(doc, "6. Antes de entregar (autochequeo)", size=12, bold=True, color=AZUL)
    bullets(doc, [f"☐ {p}" for p in c.get("pistas", [])])
    para(doc, "7. Entrega", size=12, bold=True, color=AZUL)
    p = doc.add_paragraph()
    add_inline_docx(p, f"@@Sube tu taller en {EXAMLAB}@@ — domingo 23:59. Un envio por equipo.")
    out_dir = CLASES_DIR / f"Clase {n} - {c['slug']}"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_dir / f"Taller PI - Clase {n} - VetCare.docx"))


# -------------------------------------------------------------------- solucion
def build_solucion_docx(c):
    n = c["n"]
    if n in PARCIALES or not c.get("solucion_pasos"):
        return
    kit = KIT_DIR / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    stem = f"Solucion Taller Clase {n} - VetCare"
    lines = [f"# Solucion Taller · Clase {n} · {c['titulo']}", "",
             "> DOCUMENTO DOCENTE — PRIVADO. No publicar en Clases/.", "",
             "## Solucion paso a paso"]
    lines += [f"{i}. {s}" for i, s in enumerate(c["solucion_pasos"], 1)]
    lines += ["", "## Rubrica corta"] + [f"- [ ] {r}" for r in c.get("solucion_rubrica", [])]
    lines += ["", "## Errores frecuentes"] + [f"- {e}" for e in c.get("solucion_errores", [])]
    if c.get("artefacto_archivo"):
        lines += ["", f"Plantilla de apoyo: `Kit docente/Clase {n}/Plantillas/{c['artefacto_archivo']}`"]
    (kit / f"{stem}.md").write_text("\n".join(lines), encoding="utf-8")

    doc = Document()
    banda(doc, f"Solucion Taller · Clase {n} · VetCare")
    para(doc, "DOCUMENTO DOCENTE — PRIVADO (no va en Clases/)", bold=True,
         color=ROJO, shade_fill="FBE4E4")
    para(doc, c["titulo"], size=12, bold=True, color=AZUL)
    para(doc, "Solucion paso a paso", size=12, bold=True, color=AZUL)
    bullets(doc, c["solucion_pasos"])
    para(doc, "Rubrica corta", size=12, bold=True, color=AZUL)
    bullets(doc, ["[ ] " + r for r in c.get("solucion_rubrica", [])])
    para(doc, "Errores frecuentes", size=12, bold=True, color=AZUL)
    bullets(doc, c.get("solucion_errores", []))
    doc.save(str(kit / f"{stem}.docx"))


# ------------------------------------------------------------------------ quiz
def build_quiz(c):
    n = c["n"]
    items = _quiz_items(c)
    if n in PARCIALES or not items:
        return
    kit = KIT_DIR / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)

    doc = Document()
    banda(doc, f"Quiz · Clase {n} · {c['titulo']}")
    para(doc, "Version estudiante — SOLO enunciados. No proyectar la Clave docente.",
         size=10, bold=True, color=AZUL, shade_fill="E8F4FA")
    para(doc, "Individual · 8-10 min.", size=10)
    for i, it in enumerate(items, 1):
        para(doc, f"{i}. {it['q']}", size=11, bold=True)
        if it.get("tipo") == "om":
            for op in it.get("opciones", []):
                para(doc, "     " + str(op), size=10.5, space_after=2)
        elif it.get("tipo") == "vf":
            para(doc, "     ( ) Verdadero    ( ) Falso", size=10.5, space_after=2)
        else:
            para(doc, "     Respuesta: ______________________________", size=10.5, space_after=2)
    doc.save(str(kit / f"Quiz Clase {n} - VetCare.docx"))

    dk = Document()
    banda(dk, f"CLAVE DOCENTE · Quiz Clase {n}")
    para(dk, "DOCUMENTO DOCENTE — PRIVADO. No proyectar.", bold=True,
         color=ROJO, shade_fill="FBE4E4")
    for i, it in enumerate(items, 1):
        para(dk, clave_text(it, i), size=10, shade_fill="E8F4FA")
    dk.save(str(kit / f"Quiz Clase {n} - CLAVE DOCENTE.docx"))


# ---------------------------------------------------------------------- codigo
def build_codigo(c):
    """En esta asignatura el artefacto no es codigo: es una plantilla de diseño
    (requisitos, historia de usuario, especificacion de caso de uso, diccionario
    de datos) que el docente reparte y el estudiante llena."""
    n = c["n"]
    if n in PARCIALES or not c.get("artefacto_contenido"):
        return
    dest = KIT_DIR / f"Clase {n}" / "Plantillas"
    dest.mkdir(parents=True, exist_ok=True)
    nombre = c.get("artefacto_archivo") or f"Plantilla Clase {n}.md"
    (dest / nombre).write_text(c["artefacto_contenido"], encoding="utf-8")


# ---------------------------------------------------------------------- guiones
def build_guion_md(c):
    n = c["n"]
    kit = KIT_DIR / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    (kit / "Capturas").mkdir(exist_ok=True)

    if n in PARCIALES:
        titulo, archivo = PARCIALES[n]
        md = f"""# Guia de aplicacion · Clase {n} · {titulo} (solo evaluacion)

> Dia de **parcial = solo evaluacion**. No hay tema nuevo ni avance del PI en clase.
> Enunciado y solucion: `Parciales/{archivo}`

- **Curso:** Seminario de Sistemas (FI303301) · 120 min · **presencial sincrono**

## Checklist 120 min

| Min | Accion |
|---|---|
| 0-10 | Ingreso, asistencia, normas (sin material no autorizado). |
| 10-15 | Entregar enunciado. Aclarar tiempo y canal de dudas de forma. |
| 15-100 | Desarrollo del parcial (silencio de trabajo). |
| 100-110 | Aviso de 10 min; revision de integridad. |
| 110-120 | Recoleccion y cierre. |

## Notas
- No mezclar «tema + parcial» el mismo dia.
- La solucion es privada: archivo `* - SOLUCION.docx` en `Parciales/`.
- El PI VetCare continua en la siguiente clase regular.
"""
        path = kit / f"Guia aplicacion {titulo} - Clase {n}.md"
        path.write_text(md, encoding="utf-8")
        return path

    teoria = "\n\n".join(c["teoria"])
    pasos = "\n".join(f"{i}. {t}" for i, t in enumerate(c["taller"], 1))
    md = f"""# Guion docente · Clase {n} · {c['titulo']}

- **Curso:** Seminario de Sistemas (FI303301) · 120 min
- **Hilo:** Proyecto Integrador **VetCare** — planos del sistema de la clinica «Huellitas»
- **Hoy avanzamos el PI en:** {c['hito_pi']}
- **Entregable de hoy:** {c['entregable']}
- **Herramienta:** {c['herramienta']}
- **Slides:** `Clases/Clase {n} - {c['slug']}/Presentacion.pptx`

> Sin mapa del curso, sin bio del docente, sin fechas de periodo: eso vive en la Sesion 0.

## Fundamento teorico para el docente

{teoria}

**Demo que usted debe poder repetir:** {c['demo']}

## Plan minuto a minuto (120 min)

### 0-10 · Encuadre
**Decir:** «Hoy avanzamos VetCare en: {c['hito_pi'].rstrip('. ')}. La teoria es corta; el peso esta en
el taller del proyecto.»
Pasar asistencia. Recordar donde quedo el avance de la clase pasada.

### 10-40 · Teoria Core
Cubrir el fundamento de arriba apoyandose en la slide «Teoria Core» y en la de codigo
proyectable. Cada 8-10 min, amarrar al producto: «esto es lo que van a dejar hoy en VetCare».
Pregunta al aire (2 min): ¿donde encaja esto en el VetCare de su equipo?

### 40-60 · Demo en vivo
**Decir:** «Miren mi pantalla. Dominio VetCare — no otro ejemplo.»
Demo: {c['demo']}
Escribir el codigo en vivo (no copiar-pegar). Codigo de apoyo:
`Kit docente/Clase {n}/Plantillas/{c.get('artefacto_archivo', '(sin plantilla)')}`

### 60-105 · Taller guiado = avance del PI
**Decir:** «Equipos: abran su proyecto VetCare. Esto suma a la rubrica del PI.»
Actividades:
{pasos}
Circular por los equipos. Empujar evidencia funcionando, no perfeccionismo.
Entregable: {c['entregable']}

### 105-120 · Criterios de exito y cierre
Repasar el checklist de la slide de criterios.
Aplicar el quiz corto de `Kit docente/Clase {n}/Quiz Clase {n} - VetCare.docx`
(la clave va aparte y **no se proyecta**).
**Decir:** «Queda avanzado: {c['hito_pi']}. Entrega en ExamLab, domingo 23:59.»

## Solucion del taller (privada)
`Kit docente/Clase {n}/Solucion Taller Clase {n} - VetCare.docx` — no proyectar completa.
"""
    path = kit / f"Guion Docente Clase {n} - {c['slug']}.md"
    path.write_text(md, encoding="utf-8")
    return path


def convert_guion(md_path: Path):
    conv = SLIDES / "guion_md_a_docx.py"
    if conv.exists():
        subprocess.run([sys.executable, str(conv), str(md_path)], check=False)


def build_readme():
    KIT_DIR.mkdir(parents=True, exist_ok=True)
    (KIT_DIR / "README.md").write_text("""# Kit docente — Seminario de Sistemas (2026-2)

Material **privado** del docente. Los estudiantes solo ven `Clases/`.

## Enfoque
Todo el curso avanza el **Proyecto Integrador VetCare**: aqui se producen los PLANOS del
sistema de la Clinica Veterinaria «Huellitas» (requisitos, UML, interfaz), no el codigo.
El mismo producto se programa en Programacion II.

## Por clase
- `Guion Docente Clase N - ….md` + `.docx` (fundamento teorico + minuto a minuto)
- `Quiz Clase N - VetCare.docx` (sin claves) + `Quiz Clase N - CLAVE DOCENTE.docx`
- `Solucion Taller Clase N - VetCare.md|.docx` (privada)
- `Plantillas/` artefactos de diseño · `Capturas/` evidencias
- Dias 5 / 10 / 15: `Guia aplicacion Parcial N` (solo evaluacion)

## Regenerar
```bash
python config/slides/build_uniajc_seminario_all.py
```

## Proyecto Integrador
- Estudiante: `Clases/Proyecto Integrador/`
- Docente: `Kit docente/Proyecto Integrador/`
""", encoding="utf-8")


def main():
    KIT_DIR.mkdir(parents=True, exist_ok=True)
    CLASES_DIR.mkdir(parents=True, exist_ok=True)
    build_readme()
    for c in CLASES:
        print(f"=== Clase {c['n']} ===")
        build_pptx(c)
        build_taller_docx(c)
        build_solucion_docx(c)
        build_quiz(c)
        build_codigo(c)
        md = build_guion_md(c)
        if md:
            convert_guion(md)
    print("DONE Seminario de Sistemas")


if __name__ == "__main__":
    main()
