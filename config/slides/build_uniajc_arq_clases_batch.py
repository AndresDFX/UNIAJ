# -*- coding: utf-8 -*-
"""Genera material Clases 1–15 · Arquitectura (UNIAJC 2026-2) con enfoque PI CloudLite App.

Por clase regular/autónoma:
  - Clases/Clase N - <Tema>/Presentacion.pptx + Taller ….docx
  - Kit docente/Clase N/Guion….md (+ .docx vía guion_md_a_docx) + Quiz….docx + Capturas/
Días de parcial (5, 9, 14): solo guía breve en Kit (enunciados ya en Parciales/).

Enfoque: teoría breve al servicio del PI; talleres = avance de entregables CloudLite.
"""
from __future__ import annotations

import os
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import calendario_2026_2 as cal  # noqa: E402
from arq_fundamentos import FUNDAMENTOS  # noqa: E402
from arq_examlab_data import (  # noqa: E402
    ACTIVIDAD_CORTE1,
    ACTIVIDAD_CORTE2,
    ACTIVIDAD_CORTE3,
    EXAMLAB as TALLERES_EXAMLAB,
)
import arq_solucion_data as soluciones  # noqa: E402
import solucion_taller  # noqa: E402
import examlab_talleres  # noqa: E402
from uniajc_slides_engine import (  # noqa: E402
    before_after_slide,
    box_note_slide,
    closing_slide,
    content_slide,
    new_prs,
    herramientas_slide,
    pseudo_code_slide,
    steps_visual_slide,
    table_content,
)
from uniajc_slides_engine import (  # noqa: E402
    AMARILLO,
    CONTENT_W,
    CIAN,
    GRAY,
    MARGIN,
    MSO_ANCHOR,
    NAVY,
    PP_ALIGN,
    Pt as EPt,
    SW,
    SH,
    WHITE,
    _rich,
    _run,
    add_logo,
    bg_white,
    blank,
    diagram_boxes_slide,
    footer_num,
    rect,
    textbox,
)

ROOT = Path(__file__).resolve().parents[2]
CURSO = ROOT / "Arquitectura de Sistemas Computacionales"
SLIDES_DIR = Path(__file__).resolve().parent

# Marcador interno "[CAP: token] descripcion" en las bullets de slides_extra: le
# recuerda al docente que tome esa captura antes de clase (ver Kit docente/Capturas/
# README.txt). NUNCA debe llegar literal a la diapositiva del estudiante. Donde ya
# existe la captura real (o su ilustracion "salida esperada"), se inserta la imagen
# y se limpia el texto; donde todavia no existe, solo se limpia el texto.
_CAP_RE = re.compile(r"📸?\s*\[CAP:\s*([\w-]+)\]\s*")
CAPTURAS_REALES = {
    "docker-ps": CURSO / "Kit docente" / "Clase 3" / "Capturas" / "salida-docker-ps.png",
    "actions-yml": CURSO / "Kit docente" / "Clase 8" / "Capturas" / "salida-actions-run.png",
}


def _limpiar_y_capturar(bullets_):
    """Quita el marcador [CAP:] de las bullets y devuelve la ruta de imagen real
    a insertar en la diapositiva, si existe alguna para los tokens presentes."""
    limpias, imagen = [], None
    for b in bullets_:
        tokens = _CAP_RE.findall(b)
        b2 = _CAP_RE.sub("", b).strip()
        b2 = re.sub(r"\s*·\s*$", "", b2).strip()
        limpias.append(b2)
        for t in tokens:
            ruta = CAPTURAS_REALES.get(t)
            if ruta and ruta.exists():
                imagen = ruta
    return limpias, imagen


def _add_captura(slide, ruta_png):
    """Inserta la captura real en la esquina inferior derecha de una content_slide,
    sin tapar las vinetas (ancho fijo moderado, alto proporcional real)."""
    from PIL import Image as _PILImage
    iw, ih = _PILImage.open(str(ruta_png)).size
    w = 3.4
    h = w * ih / iw
    x = SW - MARGIN - w
    y = SH - MARGIN - h
    slide.shapes.add_picture(str(ruta_png), Inches(x), Inches(y), width=Inches(w), height=Inches(h))

AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN_D = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xA0, 0x20, 0x30)
FONT = "Calibri"
# Monoespaciada para las plantillas del entregable: sin ella, las columnas de una
# matriz separadas por «|» no quedan alineadas y el estudiante no ve la tabla.
MONO = "Consolas"

# ---------------------------------------------------------------------------
# Modalidad de trabajo — decision docente 2026-2
# ---------------------------------------------------------------------------
# El curso paso de «equipos de 2-3, individual si el docente lo autoriza» a lo
# contrario: INDIVIDUAL POR DEFECTO, y el docente puede autorizar equipos de 2 o 3.
# Cuando hay equipo autorizado el artefacto puede ser compartido, pero la entrega en
# ExamLab SIEMPRE es individual (cada estudiante responde con sus propias palabras) y
# cualquier integrante debe poder explicar cualquier parte en 60 segundos.
#
# Estos son los textos por DEFECTO de las 15 clases. Una clase puede sobreescribir
# cualquiera de estas claves en su dict de CLASSES (p. ej. la Clase 1, donde el trabajo
# es estrictamente individual porque cada estudiante define su propio dominio).
MODALIDAD_DEFAULTS = {
    # Slides
    "agenda_taller_nota": "avance individual",
    "equipo_note": "Individual por defecto · el docente puede autorizar equipos de 2–3; "
                   "la entrega en ExamLab siempre es individual.",
    # Taller del estudiante
    "explica_60s_note": "Puedes explicar tu decisión en 60 segundos "
                        "(si trabajas en equipo autorizado, cualquier integrante debe poder hacerlo).",
    "entrega_unidad_note": "Modalidad de trabajo: individual por defecto; el docente puede autorizar "
                           "equipos de 2 o 3 y en ese caso el artefacto puede ser compartido, pero el "
                           "envío en ExamLab es siempre individual (responde con tus propias palabras).",
    # Guion docente
    "arranque_cita": "¿En qué quedó tu CloudLite la clase pasada?",
    "arranque_nota": "pregunta de arranque (1 min) para detectar estudiantes rezagados antes de avanzar:",
    "voluntario_word": "estudiante",
    "taller_modalidad_word": "individual · equipos de 2–3 solo si tú los autorizaste",
    "retro_word": "estudiantes",
    "criterio_60s_note": "el estudiante explica su artefacto en 60 s",
    "criterio_oral_note": "Explicación oral de 60 s por estudiante (muestreo; si autorizaste equipos, "
                          "pregunta a cualquier integrante).",
    "entregan_word": "los estudiantes",
}


def mod(c: dict, key: str) -> str:
    """Texto de modalidad de trabajo: override de la clase o default individual."""
    return c.get(key, MODALIDAD_DEFAULTS[key])


# ---------------------------------------------------------------------------
# Tipo de bloque: SE DERIVA DEL CALENDARIO, no se escribe a mano por clase
# ---------------------------------------------------------------------------
# Historia del defecto que esto previene: el material se genero con el calendario
# viejo, donde el festivo del 17/08 caia en la Clase 2 y el del 16/11 cerraba el
# curso. Al acortarse el semestre (24/08-22/11, 13 sesiones para 15 temas) el
# 17/08 quedo FUERA del rango y la Sesion 13 (16/11) se destino a las
# sustentaciones del PI en vivo. Como el tipo estaba escrito a mano en cada dict,
# el material siguio anunciando «actividad autonoma» para dos clases que si
# tienen encuentro sincrono. Leyendolo del calendario en cada build, un cambio de
# fechas no puede volver a desincronizar el material.
CAL_KEY = "arquitectura"


def tipo_de_clase(n_clase: int) -> str:
    """'parcial' | 'autonoma' | 'sustentacion' | 'regular' para una Clase de material.

    'presencial' y 'virtual' colapsan a 'regular': la diferencia entre ellas es de
    modalidad (aula vs Meet) y no cambia la estructura del bloque, porque las dos
    son sincronas. Lo que si cambia la estructura es que no haya encuentro
    (autonoma), que el bloque sea solo evaluacion (parcial) o que se dedique a
    sustentar el PI (sustentacion).
    """
    s = cal.sesion_de_clase(CAL_KEY, n_clase)
    if s is None:
        raise SystemExit(
            "Clase %d no aparece en 'clases_material' de ninguna sesion de '%s' en %s. "
            "Corrija el calendario, no este build." % (n_clase, CAL_KEY, cal.JSON_PATH)
        )
    if s.get("parcial"):
        return "parcial"
    t = (s.get("tipo") or "").lower()
    return t if t in ("autonoma", "sustentacion") else "regular"


TIPO_LABEL = {
    "regular": "Clase regular (teoría + taller PI) · encuentro síncrono",
    "autonoma": "Actividad autónoma (festivo, sin encuentro síncrono)",
    "sustentacion": "Sustentación del Proyecto Integrador · **en vivo** (síncrona)",
    "parcial": "Solo evaluación (Parcial)",
}


# ---------------------------------------------------------------------------
# Catálogo de clases (Plan 2026-2) — PI-first
# ---------------------------------------------------------------------------
# Sin clave "tipo": se inyecta desde el calendario justo despues de esta lista.

CLASSES = [
    {
        "n": 1,
        "slug": "Introduccion a arquitecturas cloud",
        "tema": "Introducción a arquitecturas cloud",
        "sub": "Diagnóstico · CloudLite App · primer boceto",
        "pi_hoy": "Definir dominio CloudLite App + 3–5 capacidades + problema en 2–3 frases",
        "entregable": "Ficha PI de 5 bloques + C4 Context en Mermaid renderizado en ExamLab (boceto previo en Excalidraw/draw.io)",
        "herramienta": "Padlet · Excalidraw / draw.io",
        # Modalidad individual por defecto: desde 2026-2 los textos por defecto del
        # curso ya estan escritos en modo individual (ver MODALIDAD_DEFAULTS abajo),
        # asi que esta clase solo necesita el matiz propio de la Clase 1: aqui el
        # trabajo es estrictamente individual porque cada estudiante define SU dominio.
        "agenda_taller_nota": "avance individual",
        "equipo_note": "Actividad individual: cada estudiante define su propio dominio CloudLite.",
        "ficha_bloques_note": "Ficha de 5 bloques: DOMINIO · PROBLEMA · ACTORES (con sus sistemas externos) · CAPACIDADES · FUERA DE ALCANCE.",
        "explica_60s_note": "El estudiante puede explicar su decisión en 60 segundos.",
        "entrega_unidad_note": "Un envío por estudiante.",
        "entrega_oficial_nota": "La entrega oficial se hace respondiendo las preguntas abiertas del taller dentro de ExamLab (https://uniaj.examlab.workers.dev/). El documento/ficha en Word o Google Docs es opcional, solo para que el estudiante conserve sus respuestas; lo que califica es lo que quede escrito en las preguntas de ExamLab.",
        "separar_notas_docente": True,
        "arranque_cita": "¿En qué quedó tu CloudLite la clase pasada?",
        "arranque_nota": "pregunta de arranque (1 min) para detectar estudiantes rezagados antes de avanzar:",
        "voluntario_word": "estudiante",
        "taller_modalidad_word": "individual",
        "retro_word": "estudiantes",
        "criterio_60s_note": "el estudiante explica su artefacto en 60 s",
        "criterio_oral_note": "Explicación oral de 60 s por el estudiante (muestreo).",
        "entregan_word": "los estudiantes",
        "objetivos": [
            "Ubicar el curso como diseño de arquitecturas cloud al servicio del **PI CloudLite App**.",
            "Cerrar el dominio del PI en una ficha de cinco bloques, con capacidades de negocio y no piezas tecnicas.",
            "Dejar el **dominio y alcance** del PI escritos y compartibles.",
        ],
        "slides_extra": [
            ("Qué es arquitectura cloud (mapa mental)", [
                "@@Arquitectura@@ = componentes + relaciones + despliegue + calidad (seguridad, escala, costo).",
                "Cloud: recursos **bajo demanda**, multi-tenant, API/automatización.",
                "Para CloudLite App: no «comprar servidores»; **diseñar** capas y simular en labs gratis.",
                "Bloques tipicos: cliente → API → lógica → datos → observabilidad.",
            ]),
            ("CloudLite App — el hilo conductor", [
                "Aplicación web/API de un dominio realista (citas, academia, inventario liviano…).",
                "Entregables del semestre: diagramas + contenedor (lab) + CI/CD conceptual + informe.",
                "Hoy solo: **problema + capacidades + boceto de contexto**.",
                "Sin AWS/GCP/Oracle: draw.io, Killercoda, GitHub Actions.",
            ]),
            ("De dominio a arquitectura (mini-método)", [
                "1) Actor y problema. 2) Capacidades. 3) Contenedores lógicos. 4) Datos. 5) Riesgos.",
                "Ejemplo: *AgendaU* — estudiantes reservan tutoría; API + auth + agenda + notificaciones.",
                "Pregunta de diseño: ¿qué es **núcleo** vs satélite?",
                "Salida: diagrama C4 Context (sistema + actores externos).",
            ]),
        ],
        "taller_titulo": "Taller Clase 1 — Ficha y boceto CloudLite App",
        "taller_pasos": [
            "Elija dominio concreto (no «red social genérica») y escriba: problema (2–3 frases), 4 capacidades.",
            "Escriba actores, sistemas externos (2–3) y fuera de alcance.",
            "Boceto del **C4 Context** en Excalidraw o draw.io (CloudLite + actores + sistemas externos), y después convertirlo a **Mermaid** con ayuda de una IA para pegarlo renderizado en ExamLab.",
            "Revise que los nombres de actores y sistemas externos coincidan entre la ficha y el diagrama, y entregue en **ExamLab** las preguntas 1 a 4 de la actividad del Corte 1.",
        ],
        "quiz": [
            ("¿Qué hace que un dominio sea concreto y no genérico?",
             "Que nombre a quién sufre el problema con un rol y cómo se mide con una cifra; si el enunciado sirve para cualquier sistema, es genérico."),
            ("Nombre dos entregables del PI CloudLite App.",
             "Cualquiera de: diagramas C4/despliegue, lab contenedores, workflow Actions, informe, sustentación."),
            ("¿Por qué el curso prohíbe pedir AWS/GCP con tarjeta?",
             "Política gratis+navegador; el estudiante no asume costos ni tarjeta."),
        ],
    },
    {
        "n": 2,
        "slug": "Modelos de servicio IaaS PaaS SaaS",
        "tema": "Modelos de servicio: IaaS, PaaS, SaaS",
        "sub": "ADR del PI · elección del modelo de servicio",
        "pi_hoy": "Decidir modelo dominante (IaaS/PaaS/SaaS) para CloudLite + ADR breve",
        "entregable": "ADR-001: decisión de modelo de servicio + matriz de comparación aplicada al dominio",
        "herramienta": "Google Docs · draw.io (opcional)",
        "objetivos": [
            "Comparar IaaS, PaaS y SaaS con criterios de control, operación y velocidad.",
            "Elegir el **modelo dominante** de CloudLite con justificación.",
            "Documentar la decisión como **ADR** reutilizable en el informe PI.",
        ],
        "slides_extra": [
            ("IaaS · PaaS · SaaS (sin cloud de pago)", [
                "@@IaaS@@: usted administra SO/red/runtime; el proveedor da cómputo/red/disco.",
                "@@PaaS@@: usted despliega app; el proveedor gestiona runtime/escala básica.",
                "@@SaaS@@: consume el servicio listo (correo, CRM); poca personalización profunda.",
                "En este curso **simulamos** con labs/navegador; no abrimos cuentas IaaS con tarjeta.",
            ]),
            ("Cómo decidir para CloudLite", [
                "Pregunte: ¿necesito controlar red/SO o solo desplegar API+datos?",
                "Para un MVP académico suele ganar **PaaS conceptual** + contenedores (portable).",
                "Si el dominio exige mucho control de red → justifique IaaS *simulado* en diagrama.",
                "SaaS solo como **satélite** (auth, email, analytics) — no como toda la app.",
            ]),
            # Las 6 secciones son EXACTAMENTE las que califican las preguntas 6 y 7, y van
            # en la diapositiva porque el formato del ADR vale 18.75 de los 25 puntos de la
            # clase: titulo y estado no aparecian en ningun deck y se cobraban igual.
            ("Plantilla ADR-001", [
                "@@6 secciones rotuladas@@: Título · Estado · Contexto · Decisión · Alternativas descartadas · Consecuencias.",
                "@@Título@@: «ADR-001 Modelo de servicio dominante de CloudLite App».  @@Estado@@: Aceptado + fecha de hoy.",
                "@@Contexto@@ = las restricciones (quién lo sostiene, cuánto tiempo, qué presupuesto), no el resumen del tema.",
                "@@Decisión@@: 1 frase, 1 modelo dominante.  @@Alternativas descartadas@@: exactamente 2, con motivo del dominio.",
                "Máximo 1 página. Las 5 primeras secciones son la **pregunta 6**; @@Consecuencias@@ es la **pregunta 7**.",
            ]),
        ],
        "taller_titulo": "Taller Clase 2 — ADR modelo de servicio CloudLite",
        "taller_pasos": [
            "Retome su ficha y su C4 Context de la Clase 1: dominio, capacidades y actores (no cambie de dominio).",
            "Complete la matriz «Criterio | IaaS | PaaS | SaaS» con las 4 filas en orden: control, costo cualitativo, operación, time-to-demo (pregunta 5).",
            "Redacte las 5 primeras secciones del ADR-001: Título · Estado con fecha · Contexto · Decisión (1 frase, 1 modelo dominante) · Alternativas descartadas (exactamente 2) — pregunta 6.",
            "Escriba la sección 6 del mismo ADR, Consecuencias, en los 3 ejes (operación, costo, aprendizaje) con un + y un - en cada uno (pregunta 7).",
            "Actualice el informe PI (sección «Modelo de servicio») pegando el ADR completo, las 6 secciones.",
            "Entrega domingo 23:59 en **ExamLab** (Talleres) — mismo doc del PI o anexo.",
        ],
        "quiz": [
            ("Si solo despliega código y el proveedor gestiona el runtime, ¿qué modelo es?", "PaaS."),
            ("¿SaaS puede ser el modelo dominante de CloudLite App? ¿Cuándo sí/no?",
             "Rara vez como dominante (poca personalización); sí como servicios satélite."),
            ("¿Qué es un ADR?", "Architecture Decision Record: documenta una decisión y sus trade-offs."),
        ],
    },
    {
        "n": 3,
        "slug": "Virtualizacion y contenedores",
        "tema": "Virtualización y contenedores",
        "sub": "Lab Killercoda → stub CloudLite",
        "pi_hoy": "Contenerizar un stub del servicio principal de CloudLite",
        # Sin «compose opcional»: la actividad pide un servicio, un Dockerfile y un
        # `docker run`, y la solucion docente responde que compose «hoy suma ruido».
        # La bitacora se nombra porque son 6 de los 25 puntos de la clase.
        "entregable": "Dockerfile del stub + bitácora de 5 comandos con la salida real + captura del lab",
        "herramienta": "Killercoda · alterna si no carga: LabEx Docker Playground",
        "objetivos": [
            "Diferenciar VM vs contenedor y el rol de la imagen.",
            "Ejecutar un contenedor en lab de **navegador** (sin Docker Desktop obligatorio).",
            "Publicar el puerto y verificar el servicio con un **endpoint de salud** (ruta, código, cuerpo).",
            "Dejar evidencia PI: Dockerfile del stub CloudLite + bitácora + captura.",
        ],
        "slides_extra": [
            ("VM vs contenedor", [
                "VM: hipervisor + SO completo → aislamiento fuerte, más pesado.",
                "Contenedor: comparte kernel del host → portable y rápido para demos.",
                "Imagen = capas inmutables; contenedor = instancia en ejecución.",
                "CloudLite: contenerizamos al menos **un** servicio (API stub o front estático).",
            ]),
            ("Lab en navegador (pasos demo)", [
                "Abrir **Killercoda** (killercoda.com, escenario Ubuntu) con cuenta gratuita.",
                "`docker run` de un nginx/hello y luego **su** imagen stub.",
                "📸 [CAP: docker-ps] `docker ps` con el contenedor arriba.",
                "@@La sesión caduca a 1 h:@@ el Dockerfile se escribe en **tu** carpeta del PI y se "
                "pega en el lab, nunca al contrario. Guarda capturas antes de cerrar.",
                "Si Killercoda no carga: **LabEx Docker Playground** como alterna (misma hora de "
                "sesión, pero solo **3 al día** en el plan gratuito).",
            ]),
            ("Dockerfile mínimo para el stub", [
                "Siete instrucciones: **FROM · WORKDIR · COPY** de dependencias **· RUN · COPY** del código **· EXPOSE · CMD**.",
                "`FROM` con **etiqueta fija** (`node:20-alpine`, `python:3.12-slim`): con `latest`, la imagen de hoy no es la de mañana.",
                "Dependencias **antes** que el código: así cambiar una línea no reinstala nada (es la caché por capas).",
                "@@Nunca un secreto en la imagen:@@ las capas se acumulan y `docker history` las lee aunque el archivo se borre después.",
                "Si haces `COPY . .`, va con un **`.dockerignore`** al lado (`.env`, `node_modules`, `.git`) — y se menciona en la entrega.",
                "`EXPOSE` **documenta**, no publica. Publicar es el `-p` del `docker run`.",
            ]),
            # Diapositiva anadida por el criterio rector: la pregunta 10 vale 5 puntos por
            # el build con etiqueta, el mapeo de puertos explicado y el contrato de salud
            # (ruta, codigo y cuerpo), y ninguna diapositiva del deck proyectaba nada de
            # eso: el estudiante tenia que deducirlo del enunciado mientras se le calificaba.
            ("Construir, correr y verificar el contenedor", [
                "**Construir:** `docker build -t cloudlite-api:0.1.0 .` — nombre **y** etiqueta; sin etiqueta vale la mitad.",
                "**Correr:** `docker run -d -p 8081:8080 --name api cloudlite-api:0.1.0`.",
                "En `-p 8081:8080` la **izquierda es el anfitrión** (por donde entras tú) y la **derecha es el contenedor** (la del `EXPOSE`).",
                "@@Si los inviertes,@@ `docker ps` muestra el contenedor **Up** y la petición muere sin respuesta: el síntoma no dice la causa.",
                "**Verificar:** el contrato de salud son **tres** datos — **ruta** (`GET /health`), **código** (`200` vivo · `503` sin base) y **cuerpo** con su formato.",
                "Un `200` con cuerpo vacío no distingue «vivo» de «vivo pero roto»: el cuerpo lleva al menos un campo verificable.",
            ]),
        ],
        "taller_titulo": "Taller Clase 3 — Contenedor stub CloudLite",
        # Respaldo si algun dia falta TALLERES_EXAMLAB[3]["pasos"], que es lo que se
        # publica. Se corrigio porque seguia nombrando compose y un diagrama de
        # despliegue que la actividad no pide, y no nombraba la bitacora ni el
        # endpoint de salud, que son 11 de los 25 puntos de la clase.
        "taller_pasos": [
            "Definan qué servicio contenerizan hoy (API stub o front estático del dominio).",
            "Escriban el Dockerfile de siete instrucciones **y** su `.dockerignore` al lado.",
            "En Killercoda: construyan con nombre y etiqueta, corran publicando el puerto y verifiquen `GET /health` (si no carga, LabEx como alterna).",
            "Llenen la bitácora de 5 comandos pegando la salida real, con una fila de incidente.",
            "Capturen evidencia (PNG con prompt, `docker ps` y hora del sistema) o enlace de sesión + nota de caducidad.",
        ],
        "quiz": [
            ("¿Qué comparte un contenedor con el host que una VM típicamente no comparte?", "El kernel del SO."),
            ("En `docker run -p 8081:8080`, ¿qué lado es el anfitrión y qué lado el contenedor?",
             "8081 es el anfitrión, por donde se entra desde fuera; 8080 es el del contenedor, el mismo del EXPOSE. Invertirlos deja el contenedor «Up» y la petición sin respuesta."),
            ("¿Cuáles son los tres datos del contrato de un endpoint de salud?",
             "Ruta, código de estado y cuerpo con su formato. Un 200 con cuerpo vacío no distingue «vivo» de «vivo pero roto»."),
            ("¿Por qué no poner secretos en el Dockerfile?", "Quedan en capas/historial de la imagen: borrarlos en un paso posterior no los elimina, solo los oculta."),
        ],
    },
    {
        "n": 4,
        "slug": "Microservicios y arquitecturas distribuidas",
        "tema": "Microservicios · Arquitecturas distribuidas",
        # El nivel del modelo C4 que se entrega hoy es Containers, no Components: la
        # pregunta 13 exige que la PRIMERA linea del codigo sea `C4Container`. Decir
        # «Componentes» en el subtitulo, en el entregable y en el titulo del taller
        # mandaba al estudiante a un nivel distinto del que se le califica.
        "sub": "C4 Containers, contratos y riesgos",
        "pi_hoy": "Diagramar componentes/servicios de CloudLite y sus contratos",
        "entregable": "Diagrama C4 Container en Mermaid + tabla de 3 contratos + 3 riesgos de distribución",
        "herramienta": "draw.io o Excalidraw para bocetar · Mermaid dentro de ExamLab para entregar",
        "objetivos": [
            "Contrastar monolito vs microservicios con criterios de equipo y acoplamiento.",
            "Modelar CloudLite en **C4 Container** (Mermaid) con 2–5 cajas justificadas.",
            "Definir 3 contratos con verbo, ruta y **error de negocio**, y nombrar 3 **riesgos de distribución**.",
        ],
        "slides_extra": [
            ("Monolito vs microservicios (para el PI)", [
                "Monolito modular ≠ malo: a menudo mejor para equipos pequeños.",
                "Microservicio: frontera de **despliegue** e independencia de datos (ideal).",
                "Anti-patrón: 12 microservicios para 3 estudiantes = diagrama teatro.",
                "La decisión se escribe en @@una sola frase@@ y se elige UNA: «un poco de los dos» no es una decisión.",
                "Se sostiene con **dos** criterios: tamaño del equipo (con número y plazo) y acoplamiento (qué partes cambian juntas).",
                "Y con lo que se gana **y** lo que se pierde: sin la segunda mitad, no hubo decisión.",
            ]),
            # Titulo renombrado: la herramienta de entrega no es draw.io sino Mermaid
            # dentro de ExamLab, y el nivel es Containers. El marcador [CAP: drawio-c4]
            # apuntaba a una captura que no existe en CAPTURAS_REALES, asi que se
            # borraba en silencio y publicaba una vineta colgando.
            ("C4-lite: del Context a los Containers", [
                "Se abre la ÚNICA caja del Context de la Clase 1: los actores y los externos siguen ahí, **fuera** del sistema.",
                "@@Tres datos por caja, obligatorios:@@ nombre · tecnología · responsabilidad en una frase.",
                "Lo que guarda datos se marca como **almacén** (`ContainerDb`), no como una caja más.",
                "Los nombres son **idénticos** a los del C4 Context: si allí decía «Pasarela de pagos», aquí no puede decir «Pagos».",
                "Cada flecha lleva **protocolo Y formato** (`HTTPS/JSON`, `TCP/SQL`): una línea muda no cuenta.",
                "Entre 2 y 5 cajas. El número tiene que ser coherente con la decisión de la diapositiva anterior.",
            ]),
            # Diapositiva anadida por el criterio rector: la pregunta 14 vale 7 puntos
            # por una tabla de cuatro columnas con un 409 obligatorio, y el deck no
            # proyectaba ninguna tabla de contratos: solo la frase «listen 3 contratos».
            ("Los tres contratos de CloudLite: cuatro datos por fila", [], {
                "headers": ["Contrato", "Quién llama a quién", "Verbo y ruta", "Error de negocio"],
                "rows": [
                    ["Reservar turno", "App web → API de turnos", "POST /turnos",
                     "409 · la franja de las 10:00 ya la tomó otro cliente"],
                    ["Avisar al cliente", "API de turnos → Correo transaccional (SaaS)",
                     "POST /v1/mensajes", "422 · la dirección del cliente no existe"],
                    ["Guardar el turno", "API de turnos → Base de turnos",
                     "INSERT INTO turno (es SQL, no una ruta REST)",
                     "Restricción única (barbero, franja) rechaza el duplicado"],
                ],
                "note": "Tres pares de cajas DISTINTOS, y al menos un 409 de conflicto: sin él se descuenta. "
                        "Un 500 no es error de negocio, es una falla.",
                "col_w": [2.3, 3.1, 2.9, 3.83],
                "fs_body": 11,
            }),
            ("Distribuido implica fallos", [
                "Cada flecha del diagrama es una llamada de red: puede tardar, perderse o llegar dos veces.",
                "@@Riesgo 1 · qué se cae:@@ nombra **una** caja, y di qué deja de funcionar **y qué sigue funcionando**. «Se cae todo» vale la mitad.",
                "@@Riesgo 2 · cuántos saltos de red@@ tiene una operación de punta a punta: cuéntalos sobre tu propio diagrama.",
                "@@Riesgo 3 · un dato escrito en dos pasos:@@ nómbralo y di qué pasa si el segundo falla.",
                "Caja de herramientas para mitigarlos: **timeout · reintento con espera creciente · idempotencia · circuit breaker**.",
                "«Los microservicios son más complejos» no es un riesgo: no nombra caja, ni salto, ni dato.",
            ]),
        ],
        "taller_titulo": "Taller Clase 4 — C4 Containers de CloudLite",
        # Respaldo si algun dia falta TALLERES_EXAMLAB[4]["pasos"], que es lo que se
        # publica. Mandaba a exportar PNG y archivo .drawio, cuando la entrega es
        # codigo Mermaid pegado en ExamLab, y no nombraba el error de negocio ni los
        # tres riesgos, que son 10 de los 25 puntos de la clase.
        "taller_pasos": [
            "Decidan en una frase: monolito modular o microservicios, con dos criterios y lo que se gana y se pierde.",
            "Bocetan en draw.io o Excalidraw el C4 **Containers**: 2–5 cajas justificadas, almacenes marcados como tales.",
            "Pasen el boceto a **Mermaid** (`C4Container` en la primera línea) y péguenlo en la pregunta de diagrama de ExamLab.",
            "Listen 3 contratos con las cuatro columnas: contrato, quién llama a quién, verbo y ruta, y error de negocio.",
            "Nombren los 3 riesgos de distribución: qué caja se cae, cuántos saltos de red, y qué dato se escribe en dos pasos.",
        ],
        "quiz": [
            ("¿Cuándo preferiría un monolito modular en CloudLite?",
             "Equipo pequeño, dominio acotado, menos overhead operativo."),
            ("¿Qué debe etiquetar una flecha en C4?",
             # Sin `**`: el quiz y su clave se imprimen con para(), que no interpreta
             # negrita — los asteriscos salian a la vista en la CLAVE DOCENTE.
             "Protocolo Y formato (`HTTPS/JSON`, `TCP/SQL`), no solo «usa». Una flecha muda no cuenta."),
            ("¿Cómo se marca una base de datos en un C4 Container en Mermaid?",
             "Con `ContainerDb(...)`, no con `Container(...)`: es un almacén, no un servicio."),
            ("Cite un riesgo de sistemas distribuidos.",
             "Debe nombrar una caja, un salto de red o un dato escrito en dos pasos. «Es más complejo» no cuenta."),
        ],
    },
    {
        "n": 5,
        "slug": "Parcial 1",
        "tema": "Parcial 1",
        "sub": "Solo evaluación",
        "pi_hoy": "Sin avance dirigido de PI (día solo evaluación)",
        "entregable": "—",
        "herramienta": "—",
        "objetivos": ["Evaluar RAA1–intro RAA2 (cloud, modelos, virtualización, distribuidos)."],
        "slides_extra": [],
        "taller_titulo": "",
        "taller_pasos": [],
        "quiz": [],
    },
    {
        "n": 6,
        "slug": "Seguridad en la nube",
        "tema": "Seguridad en la nube",
        "sub": "Amenazas y controles → sección PI",
        "pi_hoy": "Modelo de amenazas mínimo + controles para CloudLite",
        "entregable": "Sección Seguridad PI: 5 amenazas STRIDE-lite + controles + secretos/CI",
        # Decia «Excalidraw · Google Docs» y omitia donde se califica, que es la seccion
        # del taller que el estudiante lee antes de empezar. Hoy no se dibuja nada nuevo:
        # se escribe una tabla y una politica, y se senala sobre el C4 que ya existe.
        "herramienta": "Google Docs para la tabla y la política · ExamLab para entregar",
        "objetivos": [
            "Aplicar un modelo de amenazas simple al dominio CloudLite.",
            "Mapear controles (authn/z, secretos, superficie de red) sin cloud de pago.",
            "Dejar la sección Seguridad del informe lista en borrador.",
        ],
        # Las diapositivas 7 y 8 se anadieron porque sus puntos se calificaban sin
        # haberse proyectado: el menor privilegio «aplicado a un componente concreto,
        # diciendo que deja de poder hacer» (1.25 pts de la pregunta 2) estaba en el
        # deck como las tres palabras «least privilege en roles», y de la politica de
        # secretos (7.5 pts) solo se proyectaba donde viven — quien rota, cada cuanto y
        # el procedimiento ante filtracion, 4.5 pts, no aparecian en ninguna
        # diapositiva. El fundamento del guion si los explicaba, que es exactamente el
        # defecto que el criterio rector del repo persigue: enseniado al docente, no
        # proyectado al estudiante. Las dos estan resueltas sobre el ejemplo de turnos
        # del C4 Containers de la Clase 4 —el unico diagrama que el estudiante ya vio
        # proyectado— y NO sobre BiblioLite, que es el dominio en el que esta resuelta
        # la solucion docente: el mecanismo se muestra, la respuesta no se regala.
        "slides_extra": [
            ("Amenazas que sí importan al PI", [
                "Credenciales en repo · APIs abiertas · datos PII sin cifrado en tránsito.",
                "STRIDE-lite: Spoofing, Tampering, Repudiation, Info disclosure, DoS, Elevation.",
                "Elijan 5 amenazas **del dominio**, no genéricas de Internet.",
                "La forma que se califica: **actor o dato concreto** + **el camino** por el que "
                "ocurre. Sin esas dos partes, la amenaza vale la mitad.",
                "«Fuga de información» es una categoría. «Un cliente lee el turno de otro porque "
                "`GET /turnos/17` no valida a quién pertenece» **sí** es una amenaza.",
            ]),
            ("Controles prácticos (gratis)", [
                "Identidad: autenticación con token · autorización por rol · **menor privilegio**.",
                "Red: publicar solo el punto de entrada; la base de datos sin acceso desde internet.",
                "Aplicación: validar en el **servidor** (el formulario no cuenta) y limitar la tasa.",
                # `_rich` no interpreta la cursiva de un asterisco y los imprimia crudos:
                # el deck decia «en los *secrets* del repositorio». Va con acentos graves,
                # que es lo que el motor convierte en comillas angulares.
                "Secretos: en los `secrets` del repositorio, inyectados como variable de entorno.",
                "Los dos que más puntos valen se abren en las siguientes dos diapositivas.",
            ]),
            ("Menor privilegio: qué deja de poder hacer", [
                "No es una definición, es una **resta**: se nombra el componente y lo que **deja "
                "de poder hacer** al aplicarlo. La pregunta 2 pide las dos mitades.",
                "Ejemplo sobre el C4 de la Clase 4: la `API de turnos` no entra a la "
                "`Base de turnos` como dueña de la base, sino con un rol propio.",
                "**Puede**: leer, insertar y actualizar sus tablas de turnos.",
                "**No puede**: borrar filas, cambiar la estructura, ni leer otro esquema.",
                "Por qué importa: si mañana hay una inyección de SQL, el atacante hereda **esos** "
                "permisos y no los del dueño. Hace daño, pero no borra la evidencia.",
            ]),
            ("Política de secretos: las cuatro preguntas", [], {
                "headers": ["Lo que se pregunta", "Respuesta concreta (así se califica)"],
                "rows": [
                    ["**1. Dónde viven**",
                     "En los `secrets` del repositorio y en las variables de entorno del "
                     "servicio. En local, un `.env` que está en `.gitignore` **y** en "
                     "`.dockerignore`; se versiona `.env.example`, con los nombres y **ningún** valor."],
                    ["**2. Quién los rota**",
                     "Un **responsable con rol**, escrito en el README: el dueño del repositorio. "
                     "«Se rotan automáticamente» no responde quién; alguien responde por que ocurra."],
                    ["**3. Cada cuánto**",
                     "Un número o un evento del calendario: **al cierre de cada corte**, y en la "
                     "entrega final. «Periódicamente» no es una frecuencia."],
                    ["**4. Qué está prohibido**",
                     "El `Dockerfile`, el `README`, el YAML en claro — y también imprimir el "
                     "secreto en el log del pipeline «para verificar que llegó»."],
                    ["**Si se filtra**",
                     "**Primero rotar**, después limpiar. Borrar el commit no arregla nada: el "
                     "historial ya salió del equipo y sigue en cada clon."],
                ],
                "col_w": [2.5, 9.833],
                "fs_body": 11,
                "note": "Un secreto en el Dockerfile queda en el historial de capas de la imagen: "
                        "«docker history» lo lee aunque una capa posterior borre el archivo.",
            }),
            ("Ejercicio guiado", [
                "Amenaza → control → dónde se ve, en el diagrama que ya tienen.",
                "Ej.: llaman la API sin autenticar → token verificado → **flecha** «App web → API».",
            ]),
            ("La tabla que se califica: una fila por amenaza", [], {
                # Encabezados LITERALMENTE los tres de la pregunta 2, de la solucion
                # docente y de la plantilla del taller. Traian una cuarta columna
                # «STRIDE» al frente, y como la diapositiva se titula «la tabla que se
                # califica», proyectaba una forma que no es la que se califica: la letra
                # de STRIDE se registra con la amenaza en la pregunta 1. Se conserva
                # como prefijo de la celda, que ensena lo mismo sin inventar columna.
                # La columna «Donde se ve» solo admite una caja o una flecha del C4
                # Containers o del Despliegue: en una version anterior de esta diapositiva
                # dos de las cuatro filas ponian «.dockerignore» y «contrato del
                # endpoint», que no son ninguna de las dos, y son los 2.5 pts que la
                # rubrica reparte justamente por esa columna.
                "headers": ["Amenaza", "Control", "Dónde se ve (caja o flecha)"],
                "rows": [
                    ["**S ·** Un cliente reserva **a nombre de otro**: `POST /turnos` toma el id "
                     "del cuerpo de la petición.",
                     "El id se toma del token verificado, no del cuerpo. Se prueba mandando un id "
                     "ajeno y esperando 403.",
                     "**Flecha** «App web → API de turnos»"],
                    ["**T ·** Un cliente **mueve la franja** de un turno ajeno porque la API no "
                     "revisa de quién es.",
                     "Validar el rol y la propiedad del turno antes de aceptar el cambio.",
                     "**Caja** «API de turnos»"],
                    ["**I ·** La **llave del servicio de avisos** queda dentro de la imagen del "
                     "contenedor.",
                     "La llave vive en los `secrets` y entra como variable de entorno en ejecución.",
                     "**Flecha** «API de turnos → Worker de avisos»"],
                    ["**D ·** Un script sin autenticar golpea `GET /turnos` mil veces por minuto "
                     "y agota las conexiones.",
                     "Límite de tasa por identidad o por IP en el punto de entrada.",
                     "**Caja** «App web» (punto de entrada)"],
                ],
                "col_w": [4.6, 3.9, 3.833],
                "fs_body": 10,
                "note": "Cuatro filas de ejemplo; el entregable pide cinco, sobre SU dominio. Una "
                        "lista genérica de buenas prácticas no es un modelo de amenazas, y dos "
                        "filas que dicen lo mismo con otras palabras cuentan como una.",
            }),
        ],
        "taller_titulo": "Taller Clase 6 — Seguridad CloudLite",
        "taller_pasos": [
            "Listen 5 amenazas STRIDE-lite aplicadas a su dominio.",
            "Para cada una: control + dónde se ve en el diagrama.",
            "Definan política de secretos del repo/Actions.",
            "Redacten sección Seguridad del informe PI (1–1.5 páginas).",
            "Entrega domingo 23:59 (avance PI).",
        ],
        "quiz": [
            ("¿Por qué no hardcodear API keys en el Dockerfile?", "Quedan en la imagen y el historial."),
            ("Nombre un control frente a spoofing de API.", "Autenticación con tokens/credenciales + HTTPS."),
            ("¿Qué es least privilege?", "Dar solo los permisos mínimos necesarios a cada rol/servicio."),
        ],
    },
    {
        "n": 7,
        "slug": "Redes y almacenamiento cloud",
        "tema": "Redes y almacenamiento cloud",
        "sub": "Diagrama de despliegue CloudLite",
        "pi_hoy": "Diagrama de despliegue: red, zonas, almacenamiento",
        # El entregable decia «draw.io» y la herramienta del dia tambien, pero lo que se
        # califica son 14 pts de **codigo Mermaid pegado en ExamLab**, con 2 pts que dependen
        # de que renderice sin error. draw.io y Excalidraw siguen sirviendo para el boceto —
        # es el paso 1 de la diapositiva «Del boceto a ExamLab» — pero no son la entrega.
        "entregable": "Diagrama Deployment en Mermaid dentro de ExamLab (3 zonas + puertos) + tipo de almacenamiento por componente",
        "herramienta": "ExamLab (Mermaid) · boceto en draw.io o Excalidraw",
        "objetivos": [
            "Modelar red lógica (cliente, edge, app, datos) sin VPC de pago.",
            "Elegir tipo de almacenamiento según el caso de uso CloudLite.",
            "Completar el diagrama de **despliegue** del PI.",
        ],
        "slides_extra": [
            ("Red lógica para el diagrama", [
                "Cliente → edge/balanceador → app → datos: **tres zonas**, no dos.",
                "**La base de datos va en la zona de datos**, nunca en la pública: es el error que la nota castiga.",
                "**Frontera de confianza:** la flecha donde termina lo que tú controlas (un SaaS externo).",
                "No inventen subnets AWS: usen zonas **Pública / Privada / Datos**.",
            ]),
            ("Almacenamiento", [
                "Tres tipos y sus nombres exactos: **Relacional** · **Bloque** · **Objeto**.",
                "Relacional: el dato se cruza con otro · Bloque: lo monta un solo proceso · Objeto: se recupera entero.",
                "Se justifica por la **característica del dato**, no por preferencia.",
                "Si tu dominio no maneja archivos, **declara que no necesitas objeto**: eso suma.",
            ]),
            ("Checklist del diagrama Deployment", [
                "Tres zonas rotuladas · cada componente en su zona · puerto de cada uno.",
                "Fronteras de confianza marcadas · **que renderice sin error** en ExamLab.",
                "Debe alinearse con el C4 Containers (mismos nombres).",
            ]),
        ],
        "taller_titulo": "Taller Clase 7 — Despliegue y storage CloudLite",
        "taller_pasos": [
            "Bocetén el Deployment en draw.io o Excalidraw con las **tres** zonas (pública / privada / datos).",
            "Tradúzcanlo a **Mermaid** y péguenlo en la pregunta 4 de ExamLab: se califica el diagrama **renderizado**.",
            "Etiqueten el puerto de cada componente y marquen las fronteras de confianza.",
            "Clasifiquen el almacenamiento de cada componente: Relacional / Bloque / Objeto, con la característica del dato.",
            "Llenen la tabla de correspondencia C4 Containers → Despliegue → Zona y listen los renombres.",
            "Entrega domingo 23:59.",
        ],
        "quiz": [
            ("¿Por qué separar zona pública y de datos en el diagrama?",
             "Reducir superficie de ataque y aclarar trust boundaries."),
            ("¿Cuándo usaría almacenamiento de objetos?", "Media/archivos no estructurados / backups."),
            ("¿El diagrama Deployment debe usar los mismos nombres que C4?", "Sí, para trazabilidad."),
        ],
    },
    {
        "n": 8,
        "slug": "Monitoreo optimizacion y CI-CD",
        "tema": "Monitoreo y optimización · CI/CD",
        "sub": "GitHub Actions + plan de observabilidad",
        "pi_hoy": "Workflow Actions (build/test/simulate) + métricas de monitoreo del PI",
        "entregable": ".github/workflows/ci.yml + sección Monitoreo/CI del informe",
        "herramienta": "GitHub Actions · Google Docs",
        "objetivos": [
            "Explicar pipeline CI vs CD y qué es realista sin cloud de pago.",
            "Crear un workflow Actions que construya/pruebe un stub.",
            "Definir 4–6 señales de monitoreo para CloudLite.",
        ],
        "slides_extra": [
            ("CI/CD sin tarjeta", [
                "CI: build + test en cada push. CD: deploy — aquí **simulado** (echo/artifact).",
                "GitHub Actions free: runners hosted; YAML en `.github/workflows/`.",
                "**«Ya tenemos CD» porque el YAML tiene un paso `deploy` resta puntos.** Aquí llega a «listo para desplegar».",
            ]),
            ("YAML mínimo", [
                "Tres bloques que se califican: **disparadores** (`on`) · **entorno** (`runs-on`) · **pasos**.",
                "Los pasos, en este orden: **construcción** → **prueba** → **despliegue simulado**.",
                "Secrets solo vía Settings; nunca en el YAML en claro (**cero en la pregunta si aparece uno**).",
                "**Un CI que solo imprime «OK» no es CI:** tienes que poder decir qué error lo pondría rojo.",
            ]),
            ("Monitoreo y optimización", [
                "Golden signals-lite: latencia, tráfico, errores, saturación.",
                "**Cada señal va con su umbral**: «medimos la latencia» no permite decidir nada.",
                "Al menos una debe ser un **registro** (log), no una métrica numérica.",
                "Optimización: caché conceptual, paginación, límites de rate (anotar en informe).",
            ]),
        ],
        "taller_titulo": "Taller Clase 8 — Actions + monitoreo CloudLite",
        "taller_pasos": [
            "Crea un repo free (o usa el que ya tengas del PI) con stub mínimo.",
            "Agreguen `.github/workflows/ci.yml` (build/test + deploy simulado).",
            "Listen 4–6 métricas/logs a observar en producción hipotética.",
            "Peguen captura del run verde (o YAML + explicación si Actions falla por cuota).",
            "Actualicen informe secciones CI/CD y Monitoreo.",
        ],
        "quiz": [
            ("¿Qué hace el stage CI que no hace necesariamente CD?", "Validar build/tests antes de desplegar."),
            ("¿Dónde van los secretos en Actions?", "En Settings → Secrets, no en el YAML."),
            ("Nombre dos golden signals.", "Latencia, tráfico, errores, saturación."),
        ],
    },
    {
        "n": 9,
        "slug": "Parcial 2",
        "tema": "Parcial 2",
        "sub": "Solo evaluación",
        "pi_hoy": "Sin avance dirigido de PI",
        "entregable": "—",
        "herramienta": "—",
        "objetivos": ["Evaluar seguridad, redes/storage, monitoreo y CI/CD conceptual."],
        "slides_extra": [],
        "taller_titulo": "",
        "taller_pasos": [],
        "quiz": [],
    },
    {
        "n": 10,
        "slug": "Costos y sostenibilidad cloud",
        "tema": "Costos y sostenibilidad cloud",
        "sub": "Actividad autónoma · sección PI",
        "pi_hoy": "Estimación cualitativa de costos + notas de sostenibilidad",
        "entregable": "Sección Costos/Sostenibilidad del informe (bajo/medio + drivers)",
        "herramienta": "Google Docs",
        "objetivos": [
            "Identificar drivers de costo (cómputo, datos, transferencia, idle).",
            "Proponer 3 apalancamientos de ahorro sin romper el diseño.",
            "Redactar sostenibilidad (apagado de labs, imágenes ligeras, sobredimensionamiento).",
        ],
        "slides_extra": [
            ("Costos sin factura real", [
                "Escala cualitativa: Bajo / Medio / Alto por componente.",
                "Drivers: siempre-on, egress, storage caliente, builds CI frecuentes.",
                "CloudLite: justifiquen por qué su diseño no es «siempre XL».",
            ]),
            ("Sostenibilidad", [
                "Menos capas innecesarias · imágenes slim · apagar labs · right-sizing.",
                "Relación con escalabilidad: escala cuando hay carga, no por vanidad.",
            ]),
        ],
        "taller_titulo": "Actividad autónoma Clase 10 — Costos CloudLite",
        "taller_pasos": [
            "Tabla componente → driver de costo → nivel (B/M/A) → apalancamiento.",
            "3 acciones de sostenibilidad aplicables al diseño.",
            "Integre en el informe PI (1 página).",
            "Entrega domingo 23:59.",
        ],
        "quiz": [
            ("Cite un driver típico de costo cloud.", "Cómputo idle, transferencia de datos, storage caliente…"),
            ("¿Qué es right-sizing?", "Ajustar capacidad al uso real, no sobredimensionar."),
            ("¿Por qué labs temporales ayudan a sostenibilidad académica?", "Evitan recursos siempre-on y costo/energía innecesaria."),
        ],
    },
    {
        "n": 11,
        "slug": "Avance del proyecto final",
        "tema": "Avance del proyecto final",
        "sub": "Checkpoint diagramas v1 CloudLite",
        "pi_hoy": "Integrar diagramas v1 + checklist de avance PI",
        "entregable": "Paquete v1: Context + Containers + Deployment + Dockerfile + Actions + informe 60%+",
        "herramienta": "draw.io · GitHub · Google Docs",
        "objetivos": [
            "Consolidar evidencias PI en un paquete revisable.",
            "Detectar huecos (nombres inconsistentes, servicios de más, sin seguridad).",
            "Salir con backlog claro hacia Clase 12/15.",
        ],
        "slides_extra": [
            ("Checklist de avance (obligatorio)", [
                "☐ Dominio + capacidades  ☐ ADR modelo  ☐ C4 Context/Containers",
                "☐ Deployment  ☐ Dockerfile/lab  ☐ Actions YAML  ☐ Seguridad  ☐ Costos",
                "Hoy el docente revisa en vivo; no es sustentación final.",
            ]),
            ("Errores frecuentes a corregir", [
                "Microservicios teatro · nombres distintos entre diagramas · secretos en imagen.",
                "CI sin tests · diagrama sin puertos · dominio infinito.",
            ]),
            ("Rúbrica (recordatorio)", [
                "Diagramas 25 · Contenedores 20 · CI/CD 15 · Dominio/servicio 15 · etc.",
                "Detalle en enunciado PI (no se repite evaluación global aquí).",
            ]),
        ],
        "taller_titulo": "Taller Clase 11 — Checkpoint PI CloudLite v1",
        "taller_pasos": [
            "Completen el checklist en el informe (sí/no + enlace evidencia).",
            "Unifiquen nombres entre C4 y Deployment.",
            "Empaqueten ZIP/repo: diagramas PNG, Dockerfile, YAML, informe.",
            "Feedback docente 1:1 corto (cola).",
            "Backlog escrito: 5 ítems para Clase 12.",
        ],
        "quiz": [
            ("¿Qué debe coincidir entre C4 y Deployment?", "Nombres de componentes/servicios."),
            ("¿El checkpoint de Clase 11 es el Parcial 3?", "No. Parcial 3 es evaluación aparte (Clase 14)."),
            ("Nombre un anti-patrón a corregir hoy.", "Microservicios excesivos / secretos en imagen / CI vacío."),
        ],
    },
    {
        "n": 12,
        "slug": "Pruebas de rendimiento y preparacion final",
        "tema": "Pruebas de rendimiento · Preparación de presentación final",
        "sub": "Métricas objetivo + ensayo de pitch PI",
        "pi_hoy": "Escenario de rendimiento + ensayo 5–8 min de sustentación",
        "entregable": "Sección Rendimiento + guion de pitch + paquete casi-final",
        "herramienta": "Google Docs · draw.io · (opcional) lab contenedor",
        "objetivos": [
            "Definir métricas/objetivos de rendimiento realistas para CloudLite.",
            "Diseñar un escenario de prueba (aunque sea cualitativo/simulado).",
            "Ensayar el pitch de sustentación (prep PI; Parcial 3 es otro día).",
        ],
        "slides_extra": [
            ("Rendimiento sin stress-tool de pago", [
                "Definan: RPS/usuarios concurrentes objetivo · p95 latencia · error rate.",
                "Escenario: pico de matrícula / hora pico de citas — narren carga.",
                "Bottlenecks probables: DB, auth, I/O de objetos.",
            ]),
            ("Preparación de presentación (5–8 min)", [
                "1 min problema · 2 min arquitectura · 1 min contenedor · 1 min CI · 1 min seguridad/costos · Q&A.",
                "Demo: diagrama + captura lab/Actions (no improvisar login cloud).",
                "Sustentas tú los 5 bloques; si hay equipo autorizado, hablan todos.",
            ]),
            ("Paquete de entrega", [
                "Informe + diagramas + Dockerfile + YAML + capturas.",
                "Fecha/canal: coordinación del periodo.",
            ]),
        ],
        "taller_titulo": "Taller Clase 12 — Rendimiento y ensayo CloudLite",
        "taller_pasos": [
            "Escriban escenario de carga + 3 métricas objetivo + bottleneck esperado.",
            "Ensaya el pitch 5–8 min (cronómetro) y da feedback cruzado a otro estudiante (o a otro equipo, si el docente los autorizó).",
            "Cierren backlog de Clase 11.",
            "Dejen paquete casi-final en Drive/repo.",
            "Entrega de avance domingo 23:59.",
        ],
        "quiz": [
            ("¿Qué es p95 de latencia?", "El 95% de las solicitudes están por debajo de ese tiempo."),
            ("¿La prep de pitch reemplaza el Parcial 3?", "No; Parcial 3 es evaluación síncrona por Meet del corte."),
            ("Cite un bottleneck típico.", "Base de datos, autenticación, almacenamiento de objetos…"),
        ],
    },
    {
        "n": 13,
        "slug": "Escalabilidad automatica",
        "tema": "Escalabilidad automática",
        "sub": "Actividad autónoma · escenario de escala PI",
        "pi_hoy": "Documentar política de autoescalado conceptual de CloudLite",
        "entregable": "Sección Escalabilidad: triggers, límites, qué escala y qué no",
        "herramienta": "Google Docs · draw.io (opcional nota en Deployment)",
        "objetivos": [
            "Distinguir escala vertical vs horizontal y cuándo aplicarlas.",
            "Definir triggers cualitativos (CPU, cola, RPS) sin cloud de pago.",
            "Actualizar el informe PI con la política de escala.",
        ],
        "slides_extra": [
            ("Escala para CloudLite", [
                "Horizontal: más réplicas del API. Vertical: más CPU/RAM a un nodo.",
                "Datos: escala distinta (read replicas conceptuales / partición — solo si aplica).",
                "Triggers: RPS, latencia p95, profundidad de cola, CPU.",
            ]),
            ("Límites y costos", [
                "max replicas · min replicas · cooldown — anótenlos aunque sean hipotéticos.",
                "Escalar mal = costo (enlace con Clase 10).",
            ]),
        ],
        "taller_titulo": "Actividad autónoma Clase 13 — Autoescalado CloudLite",
        "taller_pasos": [
            "Describan qué componente escala y por qué.",
            "Definan 2 triggers + min/max + qué NO se escala.",
            "Anoten impacto en costos/sostenibilidad.",
            "Opcional: marca «ASG/replicas» en diagrama Deployment.",
            "Entrega domingo 23:59 (sección Escalabilidad).",
        ],
        "quiz": [
            ("¿Horizontal vs vertical?", "Más instancias vs más recursos en la misma instancia."),
            ("¿Por qué poner max replicas?", "Controlar costo y fallos de cascada."),
            ("¿La base de datos escala igual que el API?", "No necesariamente; suele ser más delicada."),
        ],
    },
    {
        "n": 14,
        "slug": "Parcial 3",
        "tema": "Parcial 3",
        "sub": "Solo evaluación",
        "pi_hoy": "Sin trabajo dirigido nuevo de PI (evaluación del corte)",
        "entregable": "—",
        "herramienta": "—",
        "objetivos": ["Evaluar rendimiento, escalabilidad y cierre conceptual del diseño."],
        "slides_extra": [],
        "taller_titulo": "",
        "taller_pasos": [],
        "quiz": [],
    },
    {
        "n": 15,
        "slug": "Presentacion del proyecto y cierre",
        "tema": "Presentación del proyecto + cierre",
        "sub": "Sustentación en vivo del PI CloudLite · cierre del curso",
        "pi_hoy": "Sustentar en vivo el PI CloudLite App y entregar el paquete final",
        "entregable": "Paquete final en ExamLab (módulo Proyectos) + pitch de 5–8 min sustentado hoy en clase + Q&A",
        "herramienta": "Google Docs/Slides · diagramas · capturas lab",
        "objetivos": [
            "Sustentar **en vivo** CloudLite App con evidencias completas.",
            "Responder **en vivo** preguntas de arquitectura (ADRs, amenazas, escala).",
            "Cerrar el curso con reflexión de aprendizaje.",
        ],
        "slides_extra": [
            ("Cómo se ordena la sesión de hoy", [
                "Sustentación **en vivo**, en este bloque: no se reemplaza por video grabado.",
                "Turnos de **6 min de pitch + 2–4 min de Q&A**; el orden se sortea al empezar.",
                "Ten el paquete ya subido a ExamLab **antes** de tu turno (no se sube presentando).",
                "Mientras otros presentan, escuchas: el cierre del curso se hace con todo el grupo.",
            ]),
            ("Rúbrica de sustentación (recordatorio)", [
                "Claridad del problema · calidad de diagramas · demo lab/CI · respuestas.",
                "Cubres los 5 bloques del guion; en equipo autorizado hablan todos (penalización si solo uno presenta).",
            ]),
            ("Checklist final", [
                "Informe completo · 3 diagramas · Dockerfile+captura · Actions · pitch.",
                "Sin cuentas cloud de pago · sin instalar hipervisores como requisito.",
            ]),
            ("Cierre del curso", [
                "RAA1–3 aplicados al PI. Conserven el repo como portafolio.",
                "Gracias — arquitectura es trade-offs documentados, no logos de proveedores.",
            ]),
        ],
        "taller_titulo": "Guía de sustentación Clase 15 — PI CloudLite",
        "taller_pasos": [
            "Sube el paquete final a **ExamLab** (módulo Proyectos) **antes** de tu turno: informe + evidencias.",
            "Sustenta **en vivo** el pitch de 5–8 min con la lámina de arquitectura en pantalla.",
            "Responde el **Q&A en vivo** (3–4 preguntas del docente, dirigidas al azar).",
            "Entrega el Q&A escrito (3 preguntas duras que te harías + respuestas) como preparación del anterior.",
            "Reflexión final (½ página): qué trade-off fue el más difícil.",
        ],
        "quiz": [
            ("¿Qué evidencias mínimas debe mostrar la sustentación?",
             "Diagramas + lab contenedor o captura + workflow CI + justificación de decisiones."),
            ("¿El PI reemplaza el Parcial 3?", "No."),
            ("¿Se exige AWS Free Tier?", "No; está prohibido como requisito."),
        ],
    },
]

# Fuente de verdad del tipo de bloque: el calendario del periodo (ver tipo_de_clase).
for _c in CLASSES:
    _c["tipo"] = tipo_de_clase(_c["n"])


def _shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _run_d(run, *, size=11, bold=False, color=GRIS, font=None):
    f = font or FONT
    run.font.name = f
    run._element.rPr.rFonts.set(qn("w:eastAsia"), f)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _sin_code_spans(text) -> str:
    """`code span` de Markdown -> «code span». El .docx no es Markdown: los acentos
    graves salian impresos tal cual en lo que lee el estudiante. El fuente puede
    seguir escribiendose en Markdown, que es lo que necesita el guion .md."""
    return re.sub(r"`([^`\n]+)`", r"«\1»", str(text))


def para(doc, text, *, size=11, bold=False, color=GRIS, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=6, space_before=0, shade=None, font=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if shade:
        _shade(p, shade)
    r = p.add_run(_sin_code_spans(text))
    _run_d(r, size=size, bold=bold, color=color, font=font)
    return p


def banda(doc, text):
    return para(doc, f"  {text}", size=13, bold=True, color=BLANCO, shade="095292",
                space_before=10, space_after=8)


def h2(doc, text):
    return para(doc, text, size=12, bold=True, color=AZUL, space_before=12, space_after=6)


def add_inline_docx(p, text, *, size=11, color=GRIS):
    """Soporta @@negrita@@ dentro de un run de docx (mismo formato usado en las slides).

    Los `code spans` de Markdown pasan a «comillas angulares», igual que en las slides:
    el docx del estudiante los imprimia con los acentos graves a la vista.
    """
    text = _sin_code_spans(text)
    for part in re.split(r'(@@.*?@@)', text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith('@@') and part.endswith('@@'):
            r.text = part[2:-2]
            _run_d(r, size=size, bold=True, color=color)
        else:
            r.text = part
            _run_d(r, size=size, color=color)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_inline_docx(p, it)


def margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)


def cover_slide(prs, n: int, tema: str, sub: str, pi_hoy: str, *, tipo: str = "regular"):
    s = blank(prs)
    bg_white(s)
    rect(s, 0, 0, SW, 3.0, NAVY)
    rect(s, 0, 3.0, SW, 0.08, CIAN)
    add_logo(s, width=2.0, corner="left-top", mt=0.3, mr=0.5, variant="white")
    tn = textbox(s, SW - 2.2, 0.35, 1.8, 0.4)
    pn = tn.paragraphs[0]
    pn.alignment = PP_ALIGN.RIGHT
    _run(pn.add_run(), f"Clase {n}", 12, CIAN, bold=True)
    tf = textbox(s, MARGIN, 1.0, CONTENT_W, 1.5, anchor=MSO_ANCHOR.MIDDLE)
    p = tf.paragraphs[0]
    p.alignment = PP_ALIGN.CENTER
    _run(p.add_run(), tema, 28, WHITE, bold=True)
    ps = tf.add_paragraph()
    ps.alignment = PP_ALIGN.CENTER
    ps.space_before = EPt(8)
    _run(ps.add_run(), sub, 15, CIAN)
    tm = textbox(s, MARGIN, 3.5, CONTENT_W, 2.8)
    if tipo == "sustentacion":
        # El bloque no se reparte en teoría + taller: son turnos de sustentación.
        lineas_cover = [
            f"**Hoy cerramos el PI:** {pi_hoy}",
            "Bloque **120 min** · sesión **síncrona** de sustentaciones · turnos consecutivos.",
            "Paquete subido a ExamLab **antes** de tu turno · defensa **en vivo**, no video grabado.",
        ]
    elif tipo == "parcial":
        # Tampoco hay teoría ni taller: la portada anunciaba las dos y ademas rellenaba
        # «Hoy avanzamos el PI en:» con «Sin avance dirigido de PI», que se leia como una
        # frase sin sentido en la primera diapositiva que ve el grupo.
        lineas_cover = [
            "**Hoy es solo el parcial:** no hay tema nuevo ni taller del PI.",
            "Bloque **120 min** · sesión **virtual síncrona** por Google Meet.",
            "El enunciado se comparte al empezar · **no** se distribuye antes.",
        ]
    else:
        lineas_cover = [
            f"**Hoy avanzamos el PI en:** {pi_hoy}",
            "Bloque **120 min** · Teoría breve · Taller PI · cierre.",
            "Herramientas gratis + navegador · sin AWS/GCP/Oracle Cloud.",
        ]
    for i, ln in enumerate(lineas_cover):
        p = tm.paragraphs[0] if i == 0 else tm.add_paragraph()
        p.space_after = EPt(8)
        _rich(p, ln, 15, GRAY)
    footer_num(s, 1)
    return s


# Diagramas reales (cajas + flechas dibujadas con formas pptx, no placeholders de
# imagen) para las clases donde el concepto es inherentemente visual.
DIAGRAMAS = {
    1: {
        "titulo": "Ejemplo de diagrama C4 — nivel Context",
        "sub": "Reemplacen actor / sistema / externo por su propio dominio CloudLite",
        "boxes": [
            {"id": "actor", "label": "Estudiante\n(actor)", "x": 0.9, "y": 3.0, "w": 2.3, "h": 1.2, "color": AMARILLO, "text_color": NAVY},
            {"id": "sys", "label": "CloudLite App\n(el sistema = una sola caja)", "x": 5.3, "y": 2.6, "w": 3.1, "h": 1.8, "color": NAVY},
            {"id": "ext", "label": "Sistema externo\n(ej. correo, auth)", "x": 10.1, "y": 3.0, "w": 2.5, "h": 1.2, "color": CIAN},
        ],
        "arrows": [
            {"src": "actor", "dst": "sys", "label": "usa"},
            {"src": "sys", "dst": "ext", "label": "notifica / consume"},
        ],
        "note": "Nivel Context: el sistema es UNA caja (sin abrir por dentro). Actores a la izquierda, sistemas externos a la derecha. El interior se dibuja en Clase 4 (Containers).",
    },
    3: {
        "titulo": "Máquinas virtuales vs. contenedores",
        "sub": "Misma capacidad de cómputo, distinto nivel de aislamiento",
        "boxes": [
            {"id": "h_vm", "label": "Máquinas virtuales", "x": 0.9, "y": 1.55, "w": 4.3, "h": 0.5, "color": NAVY, "size": 12},
            {"id": "hw1", "label": "Hardware físico", "x": 0.9, "y": 2.25, "w": 4.3, "h": 0.65, "color": GRAY},
            {"id": "hyp", "label": "Hipervisor", "x": 0.9, "y": 2.95, "w": 4.3, "h": 0.65, "color": CIAN},
            {"id": "os1", "label": "SO invitado A\n(completo)", "x": 0.9, "y": 3.65, "w": 2.05, "h": 0.7, "color": CIAN, "size": 11},
            {"id": "os2", "label": "SO invitado B\n(completo)", "x": 3.15, "y": 3.65, "w": 2.05, "h": 0.7, "color": CIAN, "size": 11},
            {"id": "appA", "label": "App A", "x": 0.9, "y": 4.4, "w": 2.05, "h": 0.6, "color": AMARILLO, "text_color": NAVY},
            {"id": "appB", "label": "App B", "x": 3.15, "y": 4.4, "w": 2.05, "h": 0.6, "color": AMARILLO, "text_color": NAVY},
            {"id": "h_ct", "label": "Contenedores", "x": 8.0, "y": 1.55, "w": 4.3, "h": 0.5, "color": NAVY, "size": 12},
            {"id": "hw2", "label": "Hardware físico", "x": 8.0, "y": 2.25, "w": 4.3, "h": 0.65, "color": GRAY},
            {"id": "kernel", "label": "SO anfitrión + motor de contenedores (kernel compartido)", "x": 8.0, "y": 2.95, "w": 4.3, "h": 0.65, "color": CIAN, "size": 10.5},
            {"id": "c1", "label": "Contenedor A\n(App A)", "x": 8.0, "y": 3.65, "w": 2.05, "h": 0.7, "color": AMARILLO, "text_color": NAVY, "size": 11},
            {"id": "c2", "label": "Contenedor B\n(App B)", "x": 10.25, "y": 3.65, "w": 2.05, "h": 0.7, "color": AMARILLO, "text_color": NAVY, "size": 11},
        ],
        "note": "VM: cada instancia carga un SO completo (aislamiento fuerte, más pesado). Contenedor: comparte el kernel del anfitrión y solo empaqueta app + dependencias (arranca en segundos, pesa MB no GB).",
    },
    # Los rotulos dicen ahora el TIPO de elemento C4 de cada caja (Person, Container,
    # ContainerDb) porque la pregunta 13 califica por separado los 2 puntos de marcar
    # los almacenes como `ContainerDb`, y el ejemplo proyectado rotulaba la base de
    # datos como «contenedor», igual que la API. Y las flechas decian «HTTP/JSON» y
    # «SQL»: la rubrica exige protocolo Y formato en cada una, asi que el ejemplo
    # mostraba justo la etiqueta incompleta que se descuenta.
    # Faltaba la «App web». El guion narra sobre ESTA diapositiva «tres contenedores:
    # una aplicacion web, una API y una base de datos», y la cuarta caja —el worker—
    # como la que «aparece solo si hay una razon»; el molde Mermaid de la diapositiva
    # siguiente tambien declara `Container(spa, "App web", "React", ...)`. El dibujo
    # mostraba tres cajas distintas (API, base y worker) con el navegador degradado a
    # `Person`, asi que el docente decia una cosa y se proyectaba otra, y el
    # estudiante que copiaba el ejemplo modelaba su front como actor. Los tres
    # `Container` quedan en NAVY y solo el `ContainerDb` en CIAN: el color deja de
    # contradecir la nota, que es la que cobra 2 puntos en la pregunta 13.
    4: {
        "titulo": "Ejemplo de diagrama C4 — nivel Containers",
        "sub": "2–5 cajas justificadas, cada flecha con protocolo Y formato",
        "boxes": [
            {"id": "spa", "label": "App web\n(Container · React)", "x": 0.9, "y": 2.2, "w": 2.6, "h": 1.2, "color": NAVY, "size": 11},
            {"id": "api", "label": "API de turnos\n(Container · Node.js)", "x": 5.2, "y": 2.2, "w": 2.6, "h": 1.2, "color": NAVY, "size": 11},
            {"id": "db", "label": "Base de turnos\n(ContainerDb · PostgreSQL)", "x": 9.4, "y": 2.2, "w": 2.6, "h": 1.2, "color": CIAN, "size": 11},
            {"id": "cliente", "label": "Cliente\n(Person)", "x": 0.9, "y": 4.3, "w": 2.6, "h": 1.0, "color": AMARILLO, "text_color": NAVY},
            {"id": "notif", "label": "Worker de avisos\n(Container · cola)", "x": 5.2, "y": 4.3, "w": 2.6, "h": 1.0, "color": NAVY, "size": 11},
        ],
        "arrows": [
            {"src": "cliente", "dst": "spa", "label": "usa · HTTPS"},
            {"src": "spa", "dst": "api", "label": "HTTPS/JSON"},
            {"src": "api", "dst": "db", "label": "TCP/SQL"},
            {"src": "api", "dst": "notif", "label": "evento/cola (AMQP)"},
        ],
        "note": "Lo que guarda datos va como ContainerDb, no como un Container más. La cuarta caja (el worker) existe porque el correo tarda: sin esa razón, no va. Y estos nombres deben reaparecer igual en el Deployment (Clase 7).",
    },
    7: {
        "titulo": "Ejemplo de diagrama de despliegue (Deployment)",
        "sub": "Zonas de confianza: qué es público, qué es privado, dónde viven los datos",
        "boxes": [
            {"id": "publica", "label": "Zona pública\n(balanceador / edge)", "x": 0.9, "y": 2.4, "w": 3.6, "h": 1.1, "color": AMARILLO, "text_color": NAVY},
            {"id": "privada", "label": "Zona privada\n(app CloudLite)", "x": 5.2, "y": 2.4, "w": 3.6, "h": 1.1, "color": NAVY},
            {"id": "datos", "label": "Zona de datos\n(base de datos)", "x": 9.4, "y": 2.4, "w": 3.6, "h": 1.1, "color": CIAN},
        ],
        "arrows": [
            # Los puertos son los de CloudLite y no otros: el 443 entra al edge desde
            # internet, pero del edge a la API se habla por el `EXPOSE 8080` de la Clase 3.
            # Decir «solo puerto 443» aqui contradecia el molde de Mermaid, la solucion
            # docente y el Dockerfile del Corte 1, y los puertos valen 2 de los 14 pts.
            {"src": "publica", "dst": "privada", "label": "solo el 8080"},
            {"src": "privada", "dst": "datos", "label": "solo el 5432"},
        ],
        "note": "La base de datos NUNCA vive en la zona pública. Los nombres deben coincidir con los del C4 Containers (Clase 4) — mismo sistema, otro ángulo.",
    },
}


# Codigo/artefacto PROYECTABLE por clase: lo minimo que el estudiante debe ver
# en pantalla mientras se explica (el archivo completo va al repo del PI).
CODIGO_SLIDE = {
    # Los rotulos son literalmente los de las preguntas 6 y 7. Antes faltaban Titulo y
    # Estado, que la pregunta 6 cobra, y sobraba «Riesgos», que nadie califica.
    2: ("ADR-001 — las 6 secciones caben en una pagina", [
        "## 1. Titulo     ADR-001 Modelo de servicio dominante de CloudLite App",
        "## 2. Estado     Aceptado — <fecha de hoy>",
        "",
        "## 3. Contexto                    <- restricciones, no resumen del tema",
        "Lo sostengo yo solo, 12 semanas, sin presupuesto cloud ni tarjeta; demo viva el dia de la sustentacion.",
        "",
        "## 4. Decision                    <- UNA frase, UN modelo dominante",
        "La aplicacion de CloudLite se despliega sobre PaaS.",
        "",
        "## 5. Alternativas descartadas    <- exactamente 2, con motivo del dominio",
        "- IaaS: control total, pero yo opero SO y parches -> tiempo que no tengo.",
        "- SaaS como nucleo: no queda nada que arquitecturar; solo satelite (auth/email).",
        "",
        "## 6. Consecuencias               <- pregunta 7: 3 ejes, cada uno con + y -",
        "Operacion + / -        Costo + / -        Aprendizaje + / -",
    ], "Sin Titulo y Estado con fecha no hay ADR que citar en la Clase 15: son 1.5 puntos de la pregunta 6."),
    # El `.dockerignore` va en la MISMA diapositiva y no en una nota al pie: la rubrica
    # de la pregunta 8 anula los 5 puntos del Dockerfile si se hace `COPY . .` «sin
    # .dockerignore ni mencionarlo», y esta diapositiva proyecta justo ese `COPY . .`.
    # Sin el archivo de al lado, el deck estaba ensenando el error que corta la nota.
    3: ("Dockerfile minimo del stub CloudLite", [
        "FROM node:20-alpine          # base ligera y etiqueta FIJA: nunca latest",
        "WORKDIR /app",
        "COPY package*.json ./        # dependencias PRIMERO: cambian poco",
        "RUN npm ci --omit=dev        # se instalan DENTRO de la imagen",
        "COPY . .                     # el codigo despues: cambia en cada commit",
        "EXPOSE 8080                  # documenta el puerto; NO lo publica",
        'CMD ["node", "server.js"]    # UN proceso principal por contenedor',
        "",
        "# .dockerignore  <- archivo aparte, en la misma carpeta que el Dockerfile.",
        "# Sin el, el COPY . . de arriba se lleva el .env a las capas de la imagen.",
        "node_modules",
        ".env",
        ".env.*",
        ".git",
    ], "EXPOSE, CMD y el `-p` del run llevan el MISMO puerto. Y el `COPY . .` solo es seguro con `.dockerignore` al lado."),
    # Diapositiva anadida por el criterio rector: la pregunta 13 vale 11 puntos y se
    # califica sobre codigo Mermaid —primera linea exactamente `C4Container`, almacenes
    # como `ContainerDb`, cada `Rel` con protocolo Y formato—, y el deck no proyectaba
    # una sola linea de esa sintaxis. El estudiante la deducia del enunciado mientras
    # se le calificaba, o se la inventaba pidiendosela a una IA sin saber que revisar.
    # El molde tiene que ser EL MISMO sistema del ejemplo de la diapositiva anterior.
    # Antes le faltaba el worker y la API llamaba al correo directo, que es justo el
    # diseno que el guion acaba de descartar por heredar 300-2000 ms de latencia: dos
    # respuestas distintas a la misma necesidad, en el mismo dominio, una diapositiva
    # despues. Se agrega el `Container(worker...)` y el correo pasa a colgar de el. Las
    # dos lineas en blanco se van para no pasar de 15 renglones, que es lo que cabe.
    4: ("C4Container en Mermaid: el molde que ExamLab renderiza", [
        "C4Container",
        "title Diagrama de contenedores - CloudLite Turnos",
        'Person(cliente, "Cliente de la barberia", "Reserva y consulta sus turnos")',
        'System_Boundary(cloudlite, "CloudLite App") {',
        '  Container(spa, "App web", "React", "Muestra franjas libres y crea la reserva")',
        '  Container(api, "API de turnos", "Node.js", "Valida la franja y registra el turno")',
        '  ContainerDb(db, "Base de turnos", "PostgreSQL", "Turnos, clientes y horarios")',
        '  Container(worker, "Worker de avisos", "cola", "Envia el aviso, con reintentos")',
        "}",
        'System_Ext(correo, "Correo transaccional", "Entrega el correo al cliente")',
        'Rel(cliente, spa, "Reserva un turno", "HTTPS")',
        'Rel(spa, api, "POST /turnos", "HTTPS/JSON")',
        'Rel(api, db, "INSERT / SELECT de turnos", "TCP/SQL")',
        'Rel(api, worker, "Publica aviso-de-turno", "evento/cola (AMQP)")',
        'Rel(worker, correo, "Envia la confirmacion", "API REST sobre HTTPS")',
    ], "La base va como `ContainerDb` y cada `Rel` lleva protocolo Y formato. Los externos, FUERA del `System_Boundary`."),
    # Antes esta diapositiva era la tabla amenaza -> control -> evidencia, que ahora
    # vive en `slides_extra` como tabla de verdad: alli caben los encabezados exactos
    # de la pregunta 2 y las cuatro columnas no tienen que caber en 85 caracteres de
    # ancho fijo. Este hueco lo ocupa el artefacto de la demo del guion (paso 4 de
    # DEMO_ARQ[6]), que no estaba proyectado en ninguna parte: es la razon tecnica de
    # por que un secreto en el Dockerfile no se arregla borrandolo, y es la unica de
    # las cuatro respuestas de la politica que se puede DEMOSTRAR en clase.
    # La salida de «docker history» iba en columnas alineadas con espacios, y
    # `pseudo_code_slide` no fija fuente monoespaciada: el motor la pinta con la
    # tipografia proporcional del tema y las columnas se desmoronaban en pantalla.
    # Se usa el «--format» real del comando, que emite campos separados y se lee igual
    # en cualquier fuente. De paso corrige el orden: «docker history» lista la capa mas
    # reciente primero, asi que el RUN va ARRIBA del ENV, y es justo lo que hay que
    # ver — la capa que borra el archivo esta despues y no elimina la de antes.
    # La imagen es `cloudlite-api:0.1.0`, la que el estudiante construyo en la Clase 3
    # y la que muestra la captura de la demo de hoy: decia `turnos-api:1.0`, que no se
    # construye en ninguna clase, y en la misma sesion el docente proyectaba la
    # diapositiva y la captura con dos nombres distintos para el mismo comando.
    6: ("El secreto en la imagen: por qué borrarlo no sirve", [
        "# Dockerfile — el anti-patron",
        'ENV CORREO_API_KEY="sk_live_9f3a...c21"   # <- queda en la capa',
        "RUN rm -f /app/.env                       # <- NO la borra: la tapa",
        "",
        "$ docker history --format '{{.ID}} · {{.CreatedBy}} · {{.Size}}' cloudlite-api:0.1.0",
        "a91d0c33 · RUN rm -f /app/.env · 12kB",
        "b7c1e2f4 · ENV CORREO_API_KEY=sk_live_9f3a...c21 · 0B",
        "",
        "La capa de abajo sigue ahí: el «rm» de la de arriba no la elimina.",
        "Cualquiera que tenga la imagen lee la llave con este mismo comando.",
        "Lo mismo con Git: el commit borrado sigue en el historial y en cada clon.",
    ], "Por eso el primer paso ante una filtración es **rotar** la credencial, no borrar el commit."),
    # La pregunta 4 cobra 14 pts sobre codigo Mermaid pegado en ExamLab —3 zonas rotuladas,
    # cada componente en su zona, el puerto de cada uno, las fronteras de confianza y 2 pts
    # de «que renderice sin error»— y el molde no estaba proyectado en ninguna diapositiva:
    # el estudiante veia el diagrama dibujado (DIAGRAMAS[7]) pero nunca el codigo que se
    # entrega. Es el mismo molde que la Clase 4 ya usa para el C4Container.
    7: ("El Despliegue en Mermaid: el molde que ExamLab renderiza", [
        "flowchart LR",
        # La `App web` va en la zona publica y esta en el molde a proposito: su ubicacion es
        # parte de los 4 pts de «cada componente en su zona» y es la duda que el grupo
        # pregunta todos los semestres. El `Cliente / navegador` NO entra aqui —cabe en 15
        # lineas o cabe el externo, no las dos— y el guion lo resuelve: es el actor y va
        # fuera de las tres zonas, como en la tabla de la pregunta 6.
        '  subgraph publica["Zona publica - internet"]',
        '    web["App web<br/>React estatico - 443"]',
        '    edge["Edge / balanceador<br/>443 HTTPS"]',
        "  end",
        '  subgraph privada["Zona privada - solo desde el edge"]',
        '    api["API CloudLite<br/>8080 HTTP"]',
        "  end",
        # El aviso va DENTRO del rotulo de la zona, no como comentario `%%` al final de la
        # linea: Mermaid solo acepta comentarios en linea propia y un `%%` pegado al nodo
        # puede tumbar el renderizado, que son 2 de los 14 pts.
        '  subgraph datos["Zona de datos - sin internet: la BD va AQUI"]',
        '    db[("Base de datos<br/>5432 TCP")]',
        "  end",
        '  web -->|"HTTPS 443"| edge',
        '  edge -->|"HTTP 8080"| api',
        '  api -->|"TCP 5432"| db',
        '  api -->|"HTTPS 443 - frontera de confianza"| pagos["Pasarela de pagos externa"]',
    ], "Una `subgraph` por zona, el puerto en cada caja y la base con `[( )]`. La flecha al externo ES la frontera de confianza. Nombres: los mismos del C4 Containers de la Clase 4."),
    8: (".github/workflows/ci.yml — CI real, no un echo", [
        "name: CI",
        "on: [push, pull_request]          # 1. disparadores",
        "jobs:",
        "  build:",
        "    runs-on: ubuntu-latest        # 2. entorno de ejecucion",
        "    steps:",
        "      - uses: actions/checkout@v4",
        "      - uses: actions/setup-node@v4",
        "        with: { node-version: '20' }",
        "      - name: Construir",
        "        run: npm ci && docker build -t cloudlite-api:0.1.0 .",
        "      - name: Probar",
        "        run: npm test              # <- si esto no puede fallar, no es CI",
        "      - name: Despliegue SIMULADO (no despliega a ningun servidor)",
        "        run: echo \"Artefacto cloudlite-api:0.1.0 listo para desplegar\"",
    ], "Los tres pasos, en orden, y el ultimo rotulado como simulado. Secretos con ${{ secrets.NOMBRE }}, NUNCA en claro dentro del YAML."),
    13: ("Politica de autoescalado (tabla, no prosa)", [
        "Componente   | Tipo        | Trigger up      | Min | Max | Cooldown",
        "-------------|-------------|-----------------|-----|-----|---------",
        "API          | Horizontal  | CPU > 70% / 5m  |  2  |  6  | 5 min",
        "Worker cola  | Horizontal  | cola > 100 msg  |  1  |  4  | 3 min",
        "BD primaria  | NO escala   | -               |  1  |  1  | -",
    ], "El 'max' es lo que evita que la factura escale sola. La BD primaria casi nunca escala horizontal."),
}

ANTES_DESPUES_ARQ = {
    3: {
        "titulo": "Maquina virtual vs contenedor — que cambia de verdad",
        "b_t": "Maquina virtual",
        "b": ["Cada instancia trae un **SO completo** (kernel propio)",
              "Aislamiento fuerte a nivel de hardware virtual",
              "Arranque en **minutos**, tamaño en **GB**",
              "Util cuando se necesitan SO distintos"],
        "a_t": "Contenedor",
        "a": ["**Comparte el kernel** del anfitrion",
              "Empaqueta solo app + dependencias",
              "Arranque en **segundos**, tamaño en **MB**",
              "Util para desplegar el mismo servicio muchas veces"],
    },
    4: {
        "titulo": "Microservicios de verdad vs microservicios teatro",
        "b_t": "Teatro (lo que NO queremos)",
        "b": ["12 servicios para un equipo de 3",
              "Se separan por capricho tecnico, no por negocio",
              "Flechas sin protocolo ni verbo",
              "Nadie sabe explicar por que ese corte"],
        "a_t": "Frontera justificada",
        "a": ["2-5 contenedores logicos, cada uno con su razon",
              "Frontera = responsabilidad de negocio",
              "Cada flecha: protocolo + verbo (HTTP/JSON, 'reserva')",
              "Se explica en 60 s sin leer notas"],
    },
    12: {
        "titulo": "«Que sea rapido» no es un requisito",
        "b_t": "Enunciado vago",
        "b": ["«La app tiene que ser rapida»",
              "«Aguanta harta gente»",
              "«Si se pone lenta, le subimos recursos»",
              "No hay con que comparar despues"],
        "a_t": "Enunciado medible",
        "a": ["**p95 < 300 ms** en el endpoint de reserva",
              "Escenario: **150 RPS** en el pico de inicio de semestre",
              "Bottleneck sospechado: **consulta a la BD**",
              "Se puede verificar: pasa o no pasa"],
    },
}


#: Clase -> actividad de corte a la que aporta preguntas. Arquitectura no tiene una
#: actividad por clase: cuatro clases comparten UNA sola, con numeracion continua.
ACTIVIDAD_DE_CLASE = {
    n: act
    for act in (ACTIVIDAD_CORTE1, ACTIVIDAD_CORTE2, ACTIVIDAD_CORTE3)
    for n in act["clases"]
}


def _cierre_de(c):
    """Plazo real de entrega del taller de esta clase, o None si es el del domingo.

    La seccion «Entrega» del taller decia «domingo 23:59 (regla del Acuerdo)» en las
    quince clases. Es cierto solo para una actividad semanal, y en Arquitectura no hay
    ninguna: las Clases 1-4, 6-8-10 y 11-13-15 aportan preguntas a UNA actividad por
    corte que se entrega completa al cierre de ese corte. El estudiante de la Clase 6
    leia dos plazos contradictorios en el mismo documento —el domingo en la seccion 10
    y el cierre del Corte 2 en el Paso 4 y en la seccion 11— y el que aparecia primero
    era el equivocado.
    """
    act = ACTIVIDAD_DE_CLASE.get(c["n"])
    return act["cierre"] if act else None


#: Lo del dia de parcial que NO esta en el instrumento. Todo lo que si esta —archivo,
#: temas evaluados, secciones con sus puntos, peso en el corte, fecha— se lee de
#: `contenido_parciales_2026_2`, que es la fuente del .docx que abre el estudiante.
#: Copiarlo aqui era el defecto: la nota docente traia su propio `mapping` de archivos y
#: el guion decia «selección múltiple, emparejamiento, desarrollo y caso de diseño» para
#: los tres parciales, cuando el P2 tiene «Verdadero / Falso» y secciones C y D con otro
#: nombre. Un docente que leyera el guion del P2 anunciaba secciones que no existen.
#:   corte    · numero del corte que cierra; indexa ARQ_P1/P2/P3
#:   ultima   · ultima clase dictada antes del parcial, la que el grupo tiene fresca
#:   prep_pi  · clase donde se preparo el pitch del PI, o None si en este corte no aplica
#:   dudas_no · tres dudas REALES de contenido de este corte, para que el ejemplo de «esto
#:              no se responde» sea del parcial que se esta aplicando y no del primero
PARCIALES_ARQ = {
    5: {"corte": 1, "ultima": 4, "prep_pi": None,
        "dudas_no": ["¿PaaS incluye el sistema operativo?",
                     "¿los contenedores usan hipervisor?",
                     "¿esta opción es la correcta?"]},
    9: {"corte": 2, "ultima": 10, "prep_pi": None,
        "dudas_no": ["¿de la seguridad se encarga el proveedor o el cliente?",
                     "¿un bucket es almacenamiento de bloques o de objetos?",
                     "¿esta opción es la correcta?"]},
    14: {"corte": 3, "ultima": 13, "prep_pi": 12,
         "dudas_no": ["¿escalar horizontal es lo mismo que subirle la RAM?",
                      "¿la prueba de carga mide latencia o throughput?",
                      "¿esta opción es la correcta?"]},
}


def _cierre_parcial_pi(n, *, hablada=False):
    """Que pasa con el PI despues de este parcial.

    El deck y el guion decian los dos «el PI continúa en la siguiente clase», que es
    falso en el Parcial 3: lo que sigue es la sustentacion, y es justo el aviso que ese
    dia hay que dar. Se calcula de `CLASSES` para no depender de un numero escrito a
    mano que se rompa si el curso cambia de estructura.
    """
    siguiente = next((x["n"] for x in CLASSES if x["n"] > n), None)
    if siguiente and siguiente == CLASSES[-1]["n"]:
        return (f"En la Clase {siguiente} es la **sustentación del Proyecto Integrador**: "
                "lo que presentan es el CloudLite que ya construyeron, no algo nuevo."
                if hablada else
                f"Clase {siguiente}: sustentación del PI CloudLite, no hay tema nuevo")
    return ("El PI CloudLite continúa en la siguiente clase; hoy no hay tarea nueva."
            if hablada else "El PI CloudLite continúa en la siguiente clase")


def _parcial_meta(corte):
    """Portada del instrumento del corte `corte`, leida de la fuente del .docx.

    Se lee y no se copia por la misma razon que la clave de una pregunta cerrada se lee
    del banco: si manana cambia una seccion del parcial, el guion tiene que cambiar con
    ella o queda anunciando algo que el estudiante no va a encontrar.
    """
    _dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "parciales")
    if _dir not in sys.path:
        sys.path.insert(0, _dir)
    import contenido_parciales_2026_2 as cp  # noqa: PLC0415
    return getattr(cp, f"ARQ_P{corte}")["meta"]


def _pasos(c):
    """Pasos del taller de esta clase.

    Los pasos originales eran one-liners vagos ("creen diagrama Containers") sin
    cantidad ni criterio de verificacion, asi que el estudiante no sabia cuando
    habia terminado. La version reescrita vive en
    `arq_examlab_data.EXAMLAB[n]["pasos"]`; si falta, se cae a los originales
    para no romper el build.
    """
    t = TALLERES_EXAMLAB.get(c["n"]) or {}
    return t.get("pasos") or c["taller_pasos"]


# Herramientas por clase. El orden importa: primero la herramienta con la que se
# trabaja, y al final siempre Mermaid (cuando el taller pide un diagrama) y ExamLab,
# que es donde se entrega todo. Los nombres deben coincidir con el campo
# `herramienta` del dict de la clase; los logos viven en assets/herramientas/ y se
# normalizan con `python config/slides/normalizar_iconos.py`.
HERRAMIENTAS_DIA = {
    1: [{"name": "Padlet", "logo": "padlet.png", "note": "Rompehielos"},
        {"name": "Excalidraw", "logo": "excalidraw.png", "note": "Boceto rapido"},
        {"name": "draw.io", "logo": "drawio.png", "note": "C4 Context"}],
    2: [{"name": "Google Docs", "logo": "google_docs.png", "note": "ADR-001"},
        {"name": "draw.io", "logo": "drawio.png", "note": "Matriz de modelos"}],
    3: [{"name": "Killercoda", "logo": "killercoda.png", "note": "Lab del dia · sesion 1 h"},
        {"name": "LabEx Docker Playground", "logo": "labex.png", "note": "Alterna · 3 al dia"},
        {"name": "Google Docs", "logo": "google_docs.png", "note": "Informe PI"}],
    4: [{"name": "draw.io", "logo": "drawio.png", "note": "Boceto del C4"},
        {"name": "Excalidraw", "logo": "excalidraw.png", "note": "Boceto rapido"}],
    6: [{"name": "Google Docs", "logo": "google_docs.png", "note": "Tabla STRIDE y politica"},
        {"name": "Excalidraw", "logo": "excalidraw.png", "note": "Marcar controles en el C4 · opcional"}],
    7: [{"name": "draw.io", "logo": "drawio.png", "note": "Deployment"},
        {"name": "Google Docs", "logo": "google_docs.png", "note": "Informe PI"}],
    8: [{"name": "GitHub Actions", "logo": "github_actions.png", "note": "CI del PI"},
        {"name": "Google Docs", "logo": "google_docs.png", "note": "Informe PI"}],
    10: [{"name": "Google Docs", "logo": "google_docs.png", "note": "Tabla de costos"}],
    11: [{"name": "draw.io", "logo": "drawio.png", "note": "Auditoria del paquete"},
         {"name": "GitHub Actions", "logo": "github_actions.png", "note": "Evidencia CI"},
         {"name": "Google Docs", "logo": "google_docs.png", "note": "Informe PI"}],
    12: [{"name": "Google Docs", "logo": "google_docs.png", "note": "Objetivo de rendimiento"},
         {"name": "draw.io", "logo": "drawio.png", "note": "Bottleneck"}],
    13: [{"name": "Google Docs", "logo": "google_docs.png", "note": "Politica de escalado"},
         {"name": "draw.io", "logo": "drawio.png", "note": "Nota en Deployment"}],
    15: [{"name": "Google Docs", "logo": "google_docs.png", "note": "Pitch y portafolio"},
         {"name": "draw.io", "logo": "drawio.png", "note": "Diagramas del paquete"}],
}


def _herramientas_de(c: dict) -> list:
    """Herramientas de la clase + Mermaid (si hay diagrama) + ExamLab.

    Mermaid y ExamLab no se escriben clase por clase porque son transversales: si
    el taller tiene pregunta de diagrama, el estudiante entrega codigo Mermaid, y
    la entrega siempre ocurre en ExamLab. Anadirlos aqui evita que una clase se
    quede sin nombrarlos.
    """
    base = list(HERRAMIENTAS_DIA.get(c["n"], []))
    if not base:
        return []
    if _tiene_diagrama(c["n"]):
        base.append({"name": "Mermaid", "logo": "mermaid.png", "note": "Codigo del diagrama"})
    base.append({"name": "ExamLab", "logo": "examlab.png", "note": "Donde se entrega"})
    return base


# Titulo de la diapositiva del flujo de diagramacion (Excalidraw -> IA -> Mermaid
# -> ExamLab). El CONTENIDO de los 4 pasos es compartido (examlab_talleres), este
# es solo el rotulo con el que aparece en el deck del curso.
# Laboratorio de contenedores del curso. Se define aqui una sola vez para que el
# orden primario/alterna y sus limites no queden repetidos a mano por clase.
# Los dos cumplen la politica (gratis, navegador, sin tarjeta) y los dos caducan
# la sesion a 1 h; el desempate es el tope diario, ver LAB_LIMITES.
LAB_LIMITES = ("Cuenta gratuita, sin tarjeta · la sesión caduca a 1 h · "
               "un escenario a la vez")

FLUJO_SLIDE_TITULO = "Del boceto a ExamLab (diagrama)"


def _tiene_diagrama(n: int) -> bool:
    """True si el taller de ExamLab de esta clase tiene pregunta tipo `diagrama`.

    Decide si la clase necesita la diapositiva del flujo: 12 de las 15 clases de
    Arquitectura piden un diagrama, y hasta ahora ninguna explicaba que la respuesta
    es codigo Mermaid y no una imagen exportada.
    """
    taller = TALLERES_EXAMLAB.get(n) or {}
    return any(p.get("tipo") == "diagrama" for p in taller.get("preguntas", []))


def _slide_map(c: dict) -> list:
    """Titulos de las diapositivas de esta clase, EN ORDEN.

    Refleja los mismos condicionales que `build_pptx`. El guion docente numera sus
    referencias con esta lista para que se pueda leer con la presentacion
    proyectada, y `build_pptx` verifica al final que las dos coincidan: si alguien
    agrega una diapositiva y olvida el mapa, el build falla en vez de publicar un
    guion que apunta a la diapositiva equivocada.
    """
    n = c["n"]
    if c["tipo"] == "parcial":
        return [f"Portada · Clase {n} · {c['tema']}",
                "Que se evalua hoy",
                "Como se responde y como se entrega",
                f"Parcial · Clase {n}"]
    m = [f"Portada · Clase {n} · {c['tema']}",
         "Agenda de hoy (120 min)",
         "Objetivos de la clase",
         "PI CloudLite — entregable de hoy"]
    m += [x[0] for x in c.get("slides_extra", [])]
    dg = DIAGRAMAS.get(n)
    if dg:
        m.append(dg["titulo"])
    ad = ANTES_DESPUES_ARQ.get(n)
    if ad:
        m.append(ad["titulo"])
    cs = CODIGO_SLIDE.get(n)
    if cs:
        m.append(cs[0])
    if _herramientas_de(c):
        m.append("Herramientas de hoy")
    if _tiene_diagrama(n):
        m.append(FLUJO_SLIDE_TITULO)
    m.append("Sustentación (paso a paso)" if c["tipo"] == "sustentacion"
             else "Taller PI (paso a paso)")
    m.append("Para continuar (PI)")
    m.append(f"Clase {n} · cierre del PI CloudLite" if c["tipo"] == "sustentacion"
             else f"Clase {n} · PI en movimiento")
    return m


def _plano(s: str) -> str:
    """Minusculas y sin tildes, para comparar fragmentos con titulos de diapositiva.

    Los titulos del deck llevan tildes («Que es arquitectura cloud», «Como decidir»,
    «Rubrica de sustentacion») pero los modulos de datos se escriben SIN tildes por
    convencion. Sin plegar los acentos, el token «{{slide:Como decidir}}» no
    encontraba «Cómo decidir para CloudLite» y habia que buscar a mano un trozo del
    titulo que no llevara ninguna: eso es lo que hacia los fragmentos fragiles.
    """
    return "".join(
        ch for ch in unicodedata.normalize("NFD", (s or "").lower())
        if unicodedata.category(ch) != "Mn"
    )


def _slide_no(mapa, *fragmentos):
    """Numero (1-based) de la primera diapositiva cuyo titulo contiene el fragmento.

    La portada se excluye a proposito: su titulo es «Portada · Clase N · <tema>», y
    el tema repite justo las palabras que uno querria usar como fragmento. Sin
    excluirla, «{{slide:Almacenamiento}}» de la Clase 7 apuntaba a la portada
    («Redes y almacenamiento cloud») en vez de a la diapositiva 6, y lo mismo pasaba
    con «Monitoreo y» en la 8, «Sostenibilidad» en la 10 y «Preparacion de
    presentacion» en la 12. Ningun bloque de teoria se ancla nunca a la portada, asi
    que descartarla no quita nada y elimina la clase entera de falsos positivos.
    """
    for frag in fragmentos:
        f = _plano(frag)
        for i, t in enumerate(mapa, 1):
            if _plano(t).startswith("portada"):
                continue
            if f in _plano(t):
                return i
    return None


def _slide_tag(mapa, *fragmentos) -> str:
    """`[Slide 8]` listo para pegar en el guion, o cadena vacia si no aplica."""
    i = _slide_no(mapa, *fragmentos)
    return f"[Slide {i}] " if i else ""


# --- Referencias a diapositivas dentro del fundamento teorico ---------------
# En la prosa se escribe «{{slide:Teoria Core}}» y aqui se convierte en
# «diapositiva 4» usando el mapa real del deck. Asi el numero no se escribe a mano
# en ningun sitio y no puede quedar corrido.
_SLIDE_TOKEN = re.compile(r"\{\{\s*slide:\s*([^}]+?)\s*\}\}")


def _resolver_slides(texto, mapa, n_clase):
    def _rep(m):
        frag = m.group(1)
        i = _slide_no(mapa, frag)
        if i is None:
            raise SystemExit(
                f"Clase {n_clase}: el fundamento referencia la diapositiva "
                f"«{frag}», que ya no existe en el deck. Corrige el texto o "
                "_slide_map()."
            )
        return f"diapositiva {i}"
    return _SLIDE_TOKEN.sub(_rep, texto)


def _verificar_mapa(c: dict, prs) -> None:
    """Aborta si `_slide_map` y el deck real dejaron de coincidir."""
    esperado, real = len(_slide_map(c)), len(prs.slides)
    if esperado != real:
        raise SystemExit(
            f"Clase {c['n']}: el deck tiene {real} diapositivas y _slide_map() "
            f"declara {esperado}. Actualiza _slide_map() en "
            "build_uniajc_arq_clases_batch.py para que el guion siga apuntando a "
            "las diapositivas correctas."
        )


def build_pptx(c: dict) -> Path:
    n = c["n"]
    folder = CURSO / "Clases" / f"Clase {n} - {c['slug']}"
    folder.mkdir(parents=True, exist_ok=True)
    out = folder / "Presentacion.pptx"

    if c["tipo"] == "parcial":
        # Este deck es lo UNICO que el estudiante recibe el dia del parcial, y decia
        # tres cosas y ninguna util: la portada anunciaba «Teoría breve · Taller PI»
        # —que hoy no existen— y «Hoy avanzamos el PI en: Sin avance dirigido de PI»;
        # las indicaciones afirmaban que «la prep del PI / pitch quedó en la clase
        # regular anterior», falso en los Cortes 1 y 2; y en ninguna parte estaba el
        # alcance, el reparto de puntos, el canal ni la hora de cierre. Todo lo que hay
        # que anunciar al minuto 0 esta ahora proyectado, y sale de la portada del
        # propio instrumento.
        p = PARCIALES_ARQ[n]
        m = _parcial_meta(p["corte"])
        prs = new_prs()
        cover_slide(prs, n, c["tema"], "Solo evaluación · sin tema ni taller dirigido",
                    c["pi_hoy"], tipo="parcial")
        content_slide(prs, "Qué se evalúa hoy",
                      [f"**{s.split(' — ')[0]}** — {s.split(' — ')[1]}"
                       for s in m["secciones_resumen"]]
                      + [f"Total **100 puntos** · nota = puntos / 20 · "
                         f"este parcial pesa {m['valor_corte']}."],
                      sub="**Solo** "
                          + " · ".join(t.split(" · ")[0] for t in m["temas"])
                          + " — fuera de esa lista no hay nada",
                      idx=2)
        content_slide(prs, "Cómo se responde y cómo se entrega", [
            f"Tiempo previsto **{m['tiempo']}** dentro del bloque de 120.",
            "El envío **cierra en el minuto 110**: lo que llegue después no se recibe.",
            "Canal de entrega: el que se anuncia ahora. Confirmo cada recibido por el chat.",
            "Pregunta de **forma** sí (cuántas líneas, si pide tabla). De **contenido** no.",
            "Si se te cae el internet: sigue respondiendo y avisa por correo al volver.",
        ], idx=3)
        closing_slide(prs, f"Parcial · Clase {n}", [
            "Hoy solo se evalúa el corte",
            _cierre_parcial_pi(n),
        ], accent="Solo evaluación")
        _verificar_mapa(c, prs)
        prs.save(str(out))
        print("OK pptx parcial ->", out)
        return out

    prs = new_prs()
    cover_slide(prs, n, c["tema"], c["sub"], c["pi_hoy"], tipo=c["tipo"])
    idx = 2
    if c["tipo"] == "sustentacion":
        # El bloque no se reparte en teoría + taller: son turnos de sustentación
        # en vivo. Proyectar la agenda genérica aquí haría creer que todavía queda
        # tiempo de trabajo en clase, y el estudiante llegaría a subir archivos.
        agenda = [
            "**0–10** Encuadre + sorteo del orden de turnos.",
            "**10–110** Sustentaciones: **6 min de pitch + 2–4 min de Q&A** por turno.",
            "**110–120** Cierre del curso.",
            "El paquete debe estar **subido a ExamLab antes** de tu turno.",
            "Sesión **síncrona**: la defensa no se reemplaza por video grabado.",
        ]
    else:
        agenda = [
            "**0–10** Encuadre: hoy avanzamos el PI en… + entregable concreto.",
            "**10–40** Teoría Core breve (solo lo necesario para el taller PI).",
            f"**40–100** Taller guiado PI (demo en vivo + {mod(c, 'agenda_taller_nota')}).",
            "**100–115** Revisión de evidencias del PI.",
            "**115–120** Cierre: criterio de éxito + plazo domingo 23:59.",
        ]
    content_slide(prs, "Agenda de hoy (120 min)", agenda, idx=idx)
    idx += 1
    content_slide(prs, "Objetivos de la clase", c["objetivos"], idx=idx)
    idx += 1
    entregable_bullets = [
        f"@@Entregable:@@ {c['entregable']}",
        f"Herramienta: **{c['herramienta']}**",
        ("Hoy no se construye: se **defiende** lo que ya está en el paquete del PI."
         if c["tipo"] == "sustentacion"
         else "Todo lo que construyan hoy entra al **informe/repo del PI** (no es lab suelto)."),
        mod(c, "equipo_note"),
    ]
    if c.get("ficha_bloques_note"):
        entregable_bullets[-1:-1] = [c["ficha_bloques_note"]]
    content_slide(prs, "PI CloudLite — entregable de hoy", entregable_bullets, idx=idx)
    idx += 1
    for extra in c["slides_extra"]:
        # Una entrada es (titulo, vinetas) o (titulo, vinetas, tabla). La tercera
        # forma existe porque hay conceptos que son una comparacion y salen mejor
        # como tabla; renderizarla aqui, y no en un diccionario aparte, la mantiene
        # en el ORDEN de la teoria y dentro de `conceptos` (que es de donde el
        # guion saca la lista de temas y el reparto de minutos).
        title, bullets_ = extra[0], extra[1]
        tabla = extra[2] if len(extra) > 2 else None
        if tabla:
            table_content(prs, title, tabla["headers"], tabla["rows"],
                          note=tabla.get("note"), col_w=tabla.get("col_w"),
                          fs_body=tabla.get("fs_body", 11), idx=idx)
        else:
            bullets_limpias, imagen = _limpiar_y_capturar(bullets_)
            slide = content_slide(prs, title, bullets_limpias, idx=idx)
            if imagen:
                _add_captura(slide, imagen)
        idx += 1
    dg = DIAGRAMAS.get(n)
    if dg:
        diagram_boxes_slide(
            prs, dg["titulo"], dg["boxes"], arrows=dg.get("arrows"),
            sub=dg.get("sub"), note=dg.get("note"), idx=idx,
        )
        idx += 1
    ad = ANTES_DESPUES_ARQ.get(n)
    if ad:
        before_after_slide(prs, ad["titulo"], ad["b_t"], ad["b"], ad["a_t"], ad["a"], idx=idx)
        idx += 1
    cs = CODIGO_SLIDE.get(n)
    if cs:
        pseudo_code_slide(prs, cs[0], cs[1], caption=cs[2], idx=idx)
        idx += 1
    tools = _herramientas_de(c)
    if tools:
        herramientas_slide(prs, tools, title="Herramientas de hoy",
                           sub="Gratis · navegador o free tier · sin cuenta de pago",
                           idx=idx)
        idx += 1
    # Del boceto al codigo: solo donde el taller pide un diagrama. El estudiante
    # disena en Excalidraw/draw.io y entrega Mermaid; sin esta diapositiva llegaba
    # con un PNG a una caja de texto.
    if _tiene_diagrama(n):
        dialectos = examlab_talleres._dialectos_del_taller(TALLERES_EXAMLAB[n])
        steps_visual_slide(
            prs, FLUJO_SLIDE_TITULO,
            examlab_talleres.flujo_diagrama_pasos(
                dialectos[0] if len(dialectos) == 1 else "el tipo que pide el enunciado"),
            sub="El diagrama se entrega como código Mermaid dentro de ExamLab, no como imagen",
            idx=idx)
        idx += 1
    pasos_titulo = ("Sustentación (paso a paso)" if c["tipo"] == "sustentacion"
                    else "Taller PI (paso a paso)")
    content_slide(prs, pasos_titulo, [f"**{i+1}.** {p}" for i, p in enumerate(_pasos(c))], idx=idx)
    idx += 1
    if c["tipo"] == "sustentacion":
        box_note_slide(prs, "Para continuar (PI)", [
            ("info", f"Entregable: {c['entregable']}"),
            ("aclaracion", "El paquete se sube a ExamLab (https://uniaj.examlab.workers.dev/ · módulo Proyectos) **antes** del bloque de sustentaciones."),
            ("advertencia", "La sustentación es **en vivo** y con Q&A: no se acepta video grabado en su lugar."),
        ], idx=idx)
    else:
        box_note_slide(prs, "Para continuar (PI)", [
            ("info", f"Entregable: {c['entregable']}"),
            ("aclaracion", "Subir evidencias al paquete CloudLite (Drive/repo) y a ExamLab (https://uniaj.examlab.workers.dev/) domingo 23:59."),
            ("advertencia", "Sin cloud de pago ni instalaciones obligatorias de hipervisores/Docker Desktop."),
        ], idx=idx)
    idx += 1
    if c["tipo"] == "sustentacion":
        closing_slide(
            prs,
            f"Clase {n} · cierre del PI CloudLite",
            [
                c["pi_hoy"],
                f"Evidencia: {c['entregable']}",
                "Conserva el repo (informe, diagramas, Dockerfile, ci.yml) como portafolio",
            ],
            accent="Arquitectura = decisiones documentadas con sus consecuencias",
        )
    else:
        closing_slide(
            prs,
            f"Clase {n} · PI en movimiento",
            [
                c["pi_hoy"],
                f"Evidencia: {c['entregable']}",
                "Siguiente paso = siguiente hito del PI CloudLite",
            ],
            accent="Teoría al servicio del proyecto",
        )
    _verificar_mapa(c, prs)
    prs.save(str(out))
    print("OK pptx ->", out)
    return out


# Plantilla en blanco del entregable, dentro del taller del estudiante.
#
# Va SOLO donde la FORMA de la respuesta se califica. En la Clase 2 son 18.75 de los 25
# puntos: la matriz de la pregunta 5 (encabezados y filas exactos, celdas de maximo 2
# lineas) y las secciones rotuladas del ADR de las preguntas 6 y 7. Antes el taller las
# enumeraba en una linea de prosa; el estudiante inventaba la estructura y perdia puntos
# por la forma, que es justo lo que nunca se le habia dado. Los rotulos son LITERALMENTE
# los de TALLERES_EXAMLAB[n] y los de la solucion docente: si alli cambia un nombre,
# cambia aqui. No se le inventa formulario a una pregunta sin forma fija (SQL, diagrama,
# seleccion multiple, prosa libre): por eso el dict solo tiene las clases que lo piden.
#
# Ancho maximo de linea: 96 caracteres (Consolas 9 pt en la caja del docx). Mas largo
# envuelve y descuadra las columnas de la matriz.
PLANTILLA_TALLER = {
    # Clase 3: la bitacora de la pregunta 11 se califica POR SU FORMA (tres columnas
    # exactas, cinco filas en un orden dado, mas la descripcion de la captura y la fila
    # de incidente) y el contrato de salud de la pregunta 10 se descuenta si falta
    # cualquiera de sus tres datos. Las dos son estructura calificada, asi que la
    # estructura se entrega. El Dockerfile y las explicaciones de la pregunta 9 no
    # llevan plantilla: no tienen forma fija.
    3: [
        "A) CONTRATO DEL ENDPOINT DE SALUD   (pregunta 10)",
        "   Los tres datos se califican por separado: si falta uno, se descuenta.",
        "",
        "   Ruta                 : GET /______________",
        "   Codigo de estado     : ______  cuando ____________________________________",
        "                          ______  cuando ____________________________________",
        "   Cuerpo y su formato  : ______________  (di el formato: JSON, texto plano...)",
        "",
        "      [ pega aqui el cuerpo exacto que devuelve tu endpoint ]",
        "",
        "   Un 200 con el cuerpo vacio no distingue «vivo» de «vivo pero roto»: el cuerpo",
        "   lleva al menos un campo que se pueda verificar.",
        "",
        "B) BITACORA DEL LABORATORIO   (pregunta 11)",
        "   Cinco filas, en ESTE orden, una por comando. En la tercera columna va el",
        "   FRAGMENTO TEXTUAL de la salida real (numero de capas, identificador corto,",
        "   el 200): pegado, no parafraseado. «Salio bien» no suma nada en esa fila.",
        "",
        "   | Comando                          | Que esperaba          | Que salio realmente |",
        "   |----------------------------------|-----------------------|---------------------|",
        "   | 1. docker build -t ______:___ .  |                       |                     |",
        "   | 2. docker images | grep ________ |                       |                     |",
        "   | 3. docker run -d -p ___:___ ...  |                       |                     |",
        "   | 4. docker ps                     |                       |                     |",
        "   | 5. curl -i http://localhost:___/ |                       |                     |",
        "",
        "C) DESCRIPCION DE LA CAPTURA   (pregunta 11 · vale 0.5 pts cada elemento)",
        "   Los tres se tienen que ver EN LA MISMA imagen, sin recortar:",
        "",
        "   [ ] Prompt del laboratorio         : ______________________________________",
        "   [ ] Salida de docker ps            : ______________________________________",
        "   [ ] Hora del sistema               : ______________________________________",
        "   Truco: ejecuta `date` en la linea inmediatamente anterior al `docker ps`.",
        "",
        "D) FILA DE INCIDENTE   (pregunta 11)",
        "   No se acepta en blanco. Si nada fallo, va el que estuvo a punto de fallar.",
        "",
        "   Comando que fallo   : ______________________________________________________",
        "   Mensaje textual     : ______________________________________________________",
        "   Causa               : ______________________________________________________",
        "   Como lo resolvi     : ______________________________________________________",
        "",
        "E) VERIFICACION DE COHERENCIA   (1 pt de la pregunta 11, y se pierde solo)",
        "   Los tres numeros y nombres tienen que ser LOS MISMOS en las tres preguntas.",
        "   Escribelos una vez aqui y copialos, no los vuelvas a teclear:",
        "",
        "   Nombre de la imagen : ________________   Etiqueta : ____________",
        "   Puerto del contenedor (EXPOSE = lado derecho del -p) : ____________",
        "   Puerto del anfitrion (lado izquierdo del -p)         : ____________",
        "",
        "   [ ] El puerto del contenedor es el mismo en el EXPOSE (p8), en el run (p10)",
        "       y en la bitacora (p11).",
        "   [ ] El nombre y la etiqueta de la imagen son los mismos en el build, en el run",
        "       y en la fila 2 de la bitacora.",
    ],
    # Clase 4: la pregunta 12 se califica por sus tres partes (decision de UNA frase, dos
    # criterios, lo que se gana y lo que se pierde), la 14 por una tabla de cuatro columnas
    # con nombres exactos, y la 15 por tres riesgos que tienen que responder tres cosas
    # distintas. Las tres son estructura calificada. La pregunta 13 NO lleva plantilla: es
    # un diagrama, y el molde de su codigo se proyecta en la diapositiva de Mermaid.
    4: [
        "A) LA DECISION DE ARQUITECTURA   (pregunta 12)",
        "   Se elige UNA. «Un poco de los dos» vale cero: es lo que dice quien no decidio.",
        "",
        "   Decision, en UNA frase : ______________________________________________________",
        "                            (monolito modular  |  microservicios  ->  tacha una)",
        "",
        "   Criterio 1 · tamano del equipo, CON numero y plazo:",
        "     Somos ______ persona(s) y tenemos ______ semanas. Por eso ____________________",
        "     ______________________________________________________________________________",
        "",
        "   Criterio 2 · acoplamiento: di QUE partes cambian juntas:",
        "     Cuando cambia ______________________, tambien hay que cambiar _________________",
        "     ______________________________________________________________________________",
        "",
        "   Que se GANA con esta decision  : ______________________________________________",
        "   Que se PIERDE con esta decision: ______________________________________________",
        "   (sin la segunda linea no hubo decision, hubo justificacion escrita despues)",
        "",
        "B) LOS TRES CONTRATOS   (pregunta 14)",
        "   Cuatro columnas, tres filas. Los nombres de las cajas son los EXACTOS de tu",
        "   diagrama de la pregunta 13, letra por letra. Los tres contratos NO pueden ser",
        "   entre el mismo par de cajas. Al menos uno lleva un 409 de conflicto: si ninguno",
        "   lo lleva, se descuenta. Un 500 no es error de negocio, es una falla.",
        "",
        "   | Contrato        | Quien llama a quien       | Verbo y ruta      | Error de negocio |",
        "   |-----------------|---------------------------|-------------------|------------------|",
        "   |                 |                    ->     |                   |                  |",
        "   |                 |                           |                   |                  |",
        "   |                 |                    ->     |                   |                  |",
        "   |                 |                           |                   |                  |",
        "   |                 |                    ->     |                   |                  |",
        "   |                 |                           |                   |                  |",
        "",
        "   Si el contrato es con la base de datos, en «Verbo y ruta» va la sentencia (INSERT,",
        "   SELECT...), no una ruta REST; y si es un evento, va el nombre del evento.",
        "",
        "C) LOS TRES RIESGOS DE DISTRIBUCION   (pregunta 15)",
        "   Cada uno responde algo DISTINTO. «Los microservicios son mas complejos» vale cero:",
        "   no nombra caja, ni salto, ni dato.",
        "",
        "   Riesgo 1 · que se cae",
        "     Caja que falla (UNA, con su nombre): _______________________________________",
        "     Deja de funcionar                  : _______________________________________",
        "     SIGUE funcionando                  : _______________________________________",
        "     («se cae todo» vale la mitad: la otra mitad es decir que sobrevive)",
        "",
        "   Riesgo 2 · cuantos saltos de red",
        "     Operacion de punta a punta que elegi: _______________________________________",
        "     Saltos, contados sobre mi diagrama  : ______   Cuales: ______________________",
        "     ______________________________________________________________________________",
        "",
        "   Riesgo 3 · un dato escrito en dos pasos",
        "     Dato                        : _______________________________________________",
        "     Paso 1 lo escribe en        : _______________________________________________",
        "     Paso 2 lo escribe en        : _______________________________________________",
        "     Si el paso 2 falla, queda   : _______________________________________________",
        "",
        "D) VERIFICACION DE COHERENCIA   (se pierde sola, y cuesta 3 pts de la p13)",
        "   [ ] El numero de cajas del diagrama (p13) es coherente con la decision de la p12:",
        "       si dije monolito modular, no hay cinco servicios en el dibujo.",
        "   [ ] Los nombres de las cajas de la p14 son identicos a los de la p13.",
        "   [ ] Los nombres de la p13 son identicos a los del C4 Context de la pregunta 3.",
        "   [ ] Lo que guarda datos esta como ContainerDb, no como Container.",
        "   [ ] Ninguna flecha quedo sin protocolo Y formato.",
    ],
    2: [
        "A) MATRIZ IaaS / PaaS / SaaS DE TU DOMINIO   (pregunta 5)",
        "   Encabezados y filas exactos, en este orden. Ninguna celda queda vacia y ninguna pasa",
        "   de 2 lineas. Cada celda habla de TU dominio y de TUS capacidades, no de teoria general:",
        "   «mas control» no dice nada; «puedo instalar la libreria que necesita el prestamo» si.",
        "",
        "   | Criterio          | IaaS                 | PaaS                 | SaaS                 |",
        "   |-------------------|----------------------|----------------------|----------------------|",
        "   | Control           |                      |                      |                      |",
        "   |                   |                      |                      |                      |",
        "   | Costo cualitativo |                      |                      |                      |",
        "   |                   |                      |                      |                      |",
        "   | Operacion         |                      |                      |                      |",
        "   |                   |                      |                      |                      |",
        "   | Time-to-demo      |                      |                      |                      |",
        "   |                   |                      |                      |                      |",
        "",
        "   Fila Operacion: di QUIEN opera el sistema operativo y el runtime en cada modelo. Es la",
        "   que mas se equivoca: la responsabilidad no desaparece al subir de nivel, se reparte, y",
        "   en los tres modelos tu sigues respondiendo por tu aplicacion, tus permisos y tus datos.",
        "",
        "B) ADR-001  —  6 secciones rotuladas   (1 a 5 = pregunta 6 · 6 = pregunta 7)",
        "   Copia los rotulos tal cual, no agregues secciones y no pases de 1 pagina en total.",
        "",
        "   1. Titulo",
        "      ADR-001 Modelo de servicio dominante de CloudLite App",
        "",
        "   2. Estado",
        "      Aceptado - [fecha de hoy: __ / __ / 2026]",
        "",
        "   3. Contexto        <- RESTRICCIONES reales, no un resumen del tema (2 o 3 lineas)",
        "      Dominio: [el mismo de la ficha y del C4 Context de la Clase 1]",
        "      Quien sostiene el proyecto: [ ... ]      En cuanto tiempo: [ ... semanas]",
        "      Presupuesto y medios de pago: [ ... ]",
        "      Prueba: del contexto se deduce por que descartaste las dos alternativas. Si no se",
        "      deduce, todavia no es contexto.",
        "",
        "   4. Decision        <- UNA sola frase, UN solo modelo dominante (dos modelos = 0 pts aqui)",
        "      [ ... ]",
        "",
        "   5. Alternativas descartadas        <- exactamente 2, ni 1 ni 3",
        "      - [modelo]: [motivo atado a tu dominio: mas caro o mas complejo PARA QUE de tu sistema]",
        "      - [modelo]: [ ... ]",
        "      (si consumes SaaS satelite para identidad o correo, aclaralo aqui: no rompe la regla)",
        "",
        "   6. Consecuencias   (pregunta 7)   <- los tres ejes, cada uno con una + y una -",
        "      Operacion     + [que empiezas a hacer]",
        "                    - [ ... ]",
        "      Costo         + [que se abarata]",
        "                    - [que se encarece]",
        "      Aprendizaje   + [ ... ]",
        "                    - [ ... ]",
        "      Al menos UNA de las tres negativas habla de amarre al proveedor o de perdida de",
        "      control. Marca cual: [ Operacion / Costo / Aprendizaje ]",
    ],
    # Clase 6: no tenia plantilla, y de los 25 puntos de la clase 18.75 se califican POR
    # SU FORMA. La pregunta 2 pide literalmente una tabla de tres columnas —y reparte
    # 2.5 pts solo por la tercera—, la 3 pide cuatro respuestas rotuladas mas el
    # procedimiento ante filtracion, y la 1 califica cada amenaza por sus dos partes
    # (actor o dato, y camino). Las tres tienen forma fija, asi que la forma se entrega.
    # Los rotulos son los de TALLERES_EXAMLAB[6] y los de la solucion docente.
    6: [
        "A) LAS CINCO AMENAZAS   (pregunta 1 · 1.75 pts cada una)",
        "   Cada amenaza necesita las DOS partes, o vale la mitad: el actor o el dato concreto",
        "   de TU dominio, y el camino por el que ocurre. No hace falta una por cada letra.",
        "",
        "   1. STRIDE: ___  Quien o que dato: ______________________________________________",
        "      Camino por el que ocurre: ______________________________________________________",
        "   2. STRIDE: ___  Quien o que dato: ______________________________________________",
        "      Camino por el que ocurre: ______________________________________________________",
        "   3. STRIDE: ___  Quien o que dato: ______________________________________________",
        "      Camino por el que ocurre: ______________________________________________________",
        "   4. STRIDE: ___  Quien o que dato: ______________________________________________",
        "      Camino por el que ocurre: ______________________________________________________",
        "   5. STRIDE: ___  Quien o que dato: ______________________________________________",
        "      Camino por el que ocurre: ______________________________________________________",
        "",
        "B) TABLA AMENAZA - CONTROL - DONDE SE VE   (pregunta 2 · 5 pts + 2.5 pts)",
        "   Los encabezados son estos tres y no se cambian. La tercera columna vale 0.5 pts",
        "   por fila y solo admite una CAJA o una FLECHA de tu C4 Containers o de tu",
        "   Despliegue, escrita con el mismo nombre que tiene alli.",
        "",
        "   | Amenaza (n de A) | Control concreto y verificable | Donde se ve (caja o flecha) |",
        "   |------------------|--------------------------------|-----------------------------|",
        "   | 1.               |                                |                             |",
        "   | 2.               |                                |                             |",
        "   | 3.               |                                |                             |",
        "   | 4.               |                                |                             |",
        "   | 5.               |                                |                             |",
        "",
        "   Prueba de «verificable»: escribe en una linea que se hace para comprobar que el",
        "   control esta puesto. Si no puedes, todavia es una intencion.",
        "",
        "C) MENOR PRIVILEGIO   (pregunta 2 · 1.25 pts · las dos mitades pesan igual)",
        "   Definirlo sin aplicarlo vale la mitad. Aplicarlo sin decir que deja de poder",
        "   hacer, la otra mitad.",
        "",
        "   Componente de MI sistema : ______________________________________________________",
        "   Con que permisos entra   : ______________________________________________________",
        "   Que DEJA DE PODER hacer  : ______________________________________________________",
        "                              ______________________________________________________",
        "",
        "D) POLITICA DE SECRETOS DEL REPOSITORIO Y DE LA CI   (pregunta 3)",
        "   Las cuatro primeras valen 1.5 pts cada una y se califican una por una. La quinta",
        "   linea vale otros 1.5 y es la que mas se falla.",
        "",
        "   1. Donde viven        : ___________________________________________________________",
        "      Mis secretos son   : ________________, ________________, ________________",
        "   2. Quien los rota     : ________________________  (un rol responsable, no «solas»)",
        "   3. Cada cuanto        : ________________________  (un numero o un evento fijo)",
        "   4. Que esta prohibido, marca todo lo que aplique:",
        "      [ ] en el Dockerfile    [ ] en el README    [ ] en el YAML en claro",
        "      [ ] impreso en el log del pipeline          [ ] otro: _________________________",
        "   5. Si un secreto se filtra, el PRIMER paso es: ____________________________________",
        "      y despues: _____________________________________________________________________",
        "",
        "E) VERIFICACION DE COHERENCIA   (no da puntos; los quita si falla)",
        "   [ ] La tabla de B tiene UNA fila por cada amenaza de A, con el mismo numero.",
        "   [ ] Cada nombre de la tercera columna existe TAL CUAL en tu C4 Containers o en tu",
        "       Despliegue. Si no lo encuentras, al diagrama le falta esa pieza.",
        "   [ ] Ningun control cae en la aplicacion web: ocultar un boton no es un control,",
        "       porque la peticion se puede enviar sin pasar por la interfaz.",
        "   [ ] Ninguna de las cinco amenazas es otra de la lista con otras palabras.",
    ],
    # Las preguntas 5 y 6 son 11 de los 25 pts de hoy y las dos se califican por la FORMA
    # de la respuesta: una tabla de tres columnas cada una, con encabezados fijados en el
    # enunciado. La pregunta 4 NO lleva plantilla: es codigo Mermaid, y el molde va
    # proyectado en la diapositiva, no como formulario.
    7: [
        "A) ALMACENAMIENTO DE CADA COMPONENTE   (pregunta 5 · 5.5 pts)",
        "   Los encabezados son estos tres y no se cambian. La segunda columna solo admite una",
        "   de TRES palabras: Relacional, Bloque u Objeto. La tercera vale 2.5 pts del total y",
        "   tiene que nombrar la CARACTERISTICA DEL DATO, no una preferencia: «se cruza con",
        "   otros datos», «lo monta un solo proceso», «se recupera entero». «Es mas rapido» o",
        "   «es lo normal» no son caracteristicas del dato y no suman.",
        "",
        "   | Componente | Tipo       | Que caracteristica del dato lo exige |",
        "   |------------|------------|--------------------------------------|",
        "   |            |            |                                      |",
        "   |            |            |                                      |",
        "   |            |            |                                      |",
        "   |            |            |                                      |",
        "",
        "   Si tu dominio NO maneja archivos, imagenes ni adjuntos, escribe esta linea y suma",
        "   completo. Agregar un almacen de objetos «porque suena a cloud» descuenta.",
        "   Mi dominio no necesita almacenamiento de objetos porque: ___________________________",
        "   __________________________________________________________________________________",
        "",
        "B) CORRESPONDENCIA C4 CONTAINERS -> DESPLIEGUE   (pregunta 6 · 5.5 pts)",
        "   UNA fila por componente. Se descuenta si falta un componente que SI aparece en",
        "   alguno de los dos diagramas, asi que llena esta tabla con los dos abiertos al lado.",
        "   La tercera columna solo admite: Publica, Privada o Datos.",
        "",
        "   | Componente en el C4 Containers | Componente en el Despliegue | Zona |",
        "   |--------------------------------|-----------------------------|------|",
        "   |                                |                             |      |",
        "   |                                |                             |      |",
        "   |                                |                             |      |",
        "   |                                |                             |      |",
        "   |                                |                             |      |",
        "",
        "   Por que los nombres tienen que coincidir (2 pts · en terminos de que son el MISMO",
        "   sistema visto desde otro angulo, no dos sistemas):",
        "   __________________________________________________________________________________",
        "   __________________________________________________________________________________",
        "",
        "   Renombres que apliqué (1 pt · si no hubo ninguno, escribelo con esas palabras):",
        "   Antes: ______________________  Ahora: ______________________",
        "   Antes: ______________________  Ahora: ______________________",
        "   Diagrama que actualicé para que queden iguales: ___________________________________",
        "",
        "C) VERIFICACION ANTES DE ENVIAR   (no da puntos; los quita si falla)",
        "   [ ] Mi diagrama de la pregunta 4 RENDERIZO sin error dentro de ExamLab (2 pts).",
        "   [ ] Las TRES zonas estan rotuladas: publica, privada y de datos.",
        "   [ ] La base de datos NO esta en la zona publica. Si lo esta, pierdo 4 pts completos.",
        "   [ ] Cada componente lleva su puerto.",
        "   [ ] Marque la frontera de confianza: la flecha donde termina lo que yo controlo.",
        "   [ ] No aparece ningun nombre de subred, de zona de disponibilidad ni de servicio de",
        "       un proveedor concreto (nada de VPC, ni de nombres de marca).",
        "   [ ] Los nombres de las cajas son los MISMOS del C4 Containers del Corte 1.",
    ],
    # Las preguntas 8 y 10 son 11 de los 25 pts y las dos se califican por la forma: tres
    # campos nombrados en la 8, y una tabla de tres columnas con umbral obligatorio en la
    # 10 («una senal sin umbral no suma»). Las preguntas 7 y 9 no llevan plantilla: la 7 es
    # el YAML —codigo, y el molde va proyectado— y la 9 es prosa sin estructura calificada.
    8: [
        "A) QUE HACE DE VERDAD TU CONSTRUCCION Y TU PRUEBA   (pregunta 8 · 5 pts)",
        "   Los tres campos se califican por separado y el tercero vale casi la mitad. Es sobre",
        "   TU ci.yml, no sobre CI en general.",
        "",
        "   1. Que se compila o se instala (1.5 pts) : _________________________________________",
        "      ______________________________________________________________________________",
        "   2. Que se ejecuta en la prueba, y que comprueba exactamente (1.5 pts):",
        "      ______________________________________________________________________________",
        "      ______________________________________________________________________________",
        "   3. Con que condicion el pipeline debe FALLAR (2 pts). Hazte la prueba mental: que",
        "      error tendrias que introducir tu para que el check salga rojo. Si no encuentras",
        "      ninguno, tu pipeline todavia no valida nada y este punto vale CERO.",
        "      El check sale rojo cuando: _____________________________________________________",
        "      ______________________________________________________________________________",
        "",
        "B) SENALES DE MONITOREO DE MI DOMINIO   (pregunta 10 · 6 pts)",
        "   Entre 4 y 6 filas. Las cuatro primeras valen 1 pt cada una; la 5 y la 6 suman hasta",
        "   1 pt mas entre las dos. La tercera columna NO es opcional: una senal sin umbral no",
        "   suma, aunque este bien elegida. Y la segunda columna tiene que hablar de una",
        "   operacion de TU dominio, no de «el sistema».",
        "",
        "   | Senal | Que se mide en MI dominio | Umbral u objetivo |",
        "   |-------|---------------------------|-------------------|",
        "   |       |                           |                   |",
        "   |       |                           |                   |",
        "   |       |                           |                   |",
        "   |       |                           |                   |",
        "   |       |                           |                   |",
        "   |       |                           |                   |",
        "",
        "   Cual de mis filas es un REGISTRO y no una metrica numerica (1 pt · algo que se",
        "   escribe para poder reconstruir que paso despues): fila numero ______",
        "",
        "C) VERIFICACION ANTES DE ENVIAR   (no da puntos; los quita si falla)",
        "   [ ] Mi ci.yml tiene los tres bloques: on, runs-on y steps.",
        "   [ ] Los pasos estan en orden: construccion, prueba, despliegue simulado.",
        "   [ ] El paso de despliegue dice SIMULADO en su nombre y no promete un servidor real.",
        "   [ ] NINGUN secreto escrito en claro en el YAML. Si hay uno, la pregunta 7 vale cero.",
        "   [ ] La imagen y el puerto son los del Dockerfile del Corte 1, no otros.",
        "   [ ] En la pregunta 9 digo que llego hasta «listo para desplegar». Afirmar que ya",
        "       tengo CD descuenta la mitad de esa pregunta.",
        "   [ ] Cada senal de la tabla B tiene umbral. Cuento las filas sin umbral: ______",
    ],
}


# Bloque ampliado por clase (contexto/por-que-importa, escenario, pistas) — fusiona
# el contenido de los dos talleres duplicados que existian antes por clase, quedandose
# con lo mejor de cada uno en un solo documento generado por el pipeline.
TALLER_BLOQUE = {
    1: {
        "contexto": [
            "@@Por qué importa al PI:@@ sin dominio concreto no hay arquitectura que defender.",
            "La ficha + C4 Context son la semilla del informe y de todos los diagramas del semestre.",
            "Si el problema es vago (app de la universidad), el resto del PI se vuelve teatro.",
        ],
        "escenario": [
            "Actividad individual. Elegir un dominio concreto.",
            "Sugeridos: AgendaU · BiblioLite · InventarioLab · TurnosClinica · EventosCampus.",
            "Plantilla ficha (5 bloques): DOMINIO · PROBLEMA · ACTORES (con sus sistemas externos) · CAPACIDADES · FUERA DE ALCANCE.",
            "Diagrama: boceto visual en @@Excalidraw o draw.io@@ → conversión a @@Mermaid (C4Context)@@ con ayuda de una IA → pegar y @@renderizar en ExamLab@@.",
        ],
        "pistas": [
            "¿Quién sufre el problema y como lo miden?",
            "¿La caja grande es el sistema CloudLite (no un módulo interno)?",
            "¿Las flechas tienen verbo (reservar, notificar, autenticar)?",
            "¿Los 2-3 sistemas externos coinciden con los System_Ext del diagrama C4?",
            "¿El diagrama quedó pegado como Mermaid y renderizado en ExamLab (no solo como imagen)?",
            "¿Fuera de alcance está escrito (que NO haran hoy)?",
        ],
    },
    2: {
        "contexto": [
            "@@Por qué importa al PI:@@ el modelo IaaS/PaaS/SaaS define quién opera SO, runtime y costos.",
            "Sin ADR, el PI no puede justificar trade-offs en sustentación ni en parciales.",
        ],
        "escenario": [
            "Partir de la ficha/C4 de Clase 1 (mismo dominio, no se cambia).",
            "MVP académico típico: PaaS conceptual + SaaS satélite (auth/email).",
            "El ADR-001 son @@6 secciones rotuladas@@: Título · Estado · Contexto · Decisión · Alternativas descartadas · Consecuencias. Las 5 primeras van en la pregunta 6; Consecuencias es la pregunta 7.",
        ],
        # El checklist cubre las TRES preguntas de la clase (5, 6 y 7), no solo el ADR:
        # la matriz vale 6.25 puntos y no tenia ni una linea aqui.
        "pistas": [
            "¿La matriz tiene los 4 encabezados exactos (Criterio · IaaS · PaaS · SaaS) y las 4 filas en el orden pedido?",
            "¿Cada celda de la matriz nombra una capacidad de TU dominio y no cabe en más de 2 líneas?",
            "¿La fila de operación dice quién opera SO y runtime en cada modelo, sin afirmar que dejas de responder por tu app?",
            "¿El ADR trae el Título con el número y el Estado con la fecha de hoy?",
            "¿El Contexto son restricciones tuyas (quién lo sostiene, cuánto tiempo, qué presupuesto) y no un resumen del tema?",
            "¿La decisión nombra un modelo dominante en UNA frase (no un poco de todo)?",
            "¿Hay exactamente dos alternativas descartadas, cada una con el motivo atado a tu dominio?",
            "¿Las consecuencias cubren operación, costo y aprendizaje, con al menos un + y un - en cada eje?",
            "¿Alguna consecuencia negativa habla de amarre al proveedor o de pérdida de control?",
        ],
    },
    3: {
        "contexto": [
            "@@Por qué importa al PI:@@ CloudLite debe mostrar al menos un servicio contenerizado con evidencia.",
            "El contenedor es el puente entre el diagrama C4 y el despliegue realista (sin cloud de pago).",
            "Lab en navegador Killercoda: sin Docker Desktop obligatorio.",
        ],
        "escenario": [
            "Elegir el servicio principal del C4 (API o web).",
            "Abrir Killercoda (killercoda.com, escenario Ubuntu). " + LAB_LIMITES + ".",
            "Prohibido: copiar .env / API keys a la imagen.",
        ],
        # Una pista por pregunta, como minimo: antes eran 4 y cubrian la 8 y la 11,
        # asi que la 9 (capas y kernel compartido) y la 10 (etiqueta, lados del -p y
        # los tres datos del contrato de salud) llegaban sin checklist.
        "pistas": [
            "¿El Dockerfile está en tu carpeta del PI y no solo dentro del lab?",
            "¿Están las @@siete@@ instrucciones, y el `COPY` de dependencias antes del `COPY` del código? (p8)",
            "¿La imagen base tiene @@etiqueta fija@@ —no `latest`— y hay un `.dockerignore` al lado del `COPY . .`? (p8)",
            "¿Secretos fuera de la imagen?",
            "¿Nombraste @@dos@@ instrucciones de TU archivo que crean capa, y la diferencia con la VM en términos de @@kernel compartido@@? (p9)",
            "¿El `build` lleva nombre @@y@@ etiqueta, y dijiste qué lado del `-p` es el anfitrión y qué pasa si se invierten? (p10)",
            "¿El contrato de salud tiene los @@tres@@ datos: ruta, código de estado y cuerpo con su formato? (p10)",
            "¿La bitácora trae las @@5 filas en orden@@ con la salida pegada textualmente, y una fila de incidente? (p11)",
            "¿Hay evidencia con timestamp (captura o enlace)?",
            "¿La captura muestra los tres elementos en la MISMA imagen: prompt, `docker ps` y hora del sistema? (p11)",
            "¿El nombre de imagen, la etiqueta y el puerto son los mismos en la p8, la p10 y la p11?",
        ],
    },
    4: {
        "contexto": [
            "@@Por qué importa al PI:@@ el C4 Containers es el mapa lógico que luego alinea Deployment y CI.",
            "Anti-patrón: 12 microservicios para 3 estudiantes = teatro, no arquitectura.",
            "Regla CloudLite: 2-5 cajas justificadas + contratos etiquetados.",
        ],
        # «escenario» y «pistas» solo se renderizan en el .docx del estudiante, y ahi
        # add_inline_docx() entiende @@negrita@@ — no la negrita de Markdown. Con `**`
        # los asteriscos salian impresos en la ficha PI.
        "escenario": [
            "Partir del C4 Context de la pregunta 3 (mismos nombres de sistema, actores y externos).",
            "Bocetar en draw.io o Excalidraw la vista @@Containers@@ (no solo Context)…",
            "…y entregar el diagrama como código @@Mermaid@@ dentro de ExamLab, que es donde se renderiza y se califica.",
        ],
        # Una pista por pregunta, como minimo: antes eran 3 y ninguna cubria la 12
        # (la decision y sus dos criterios) ni la 15 (los tres riesgos), que juntas
        # son 7 de los 25 puntos de la clase.
        "pistas": [
            "¿La decisión de la p12 elige @@una@@ —monolito modular o microservicios— en una sola frase?",
            "¿Los dos criterios traen @@número y plazo@@ del equipo, y dicen qué partes cambian juntas?",
            "¿Dijiste qué se gana @@y@@ qué se pierde? (sin la segunda mitad no hay decisión)",
            "¿Hay 2-5 cajas (no 1 monolito innominado ni 12 microservicios), coherentes con la p12?",
            "¿La primera línea del código del diagrama es exactamente `C4Container`?",
            "¿Cada caja tiene sus @@tres@@ datos: nombre, tecnología y responsabilidad?",
            "¿Lo que guarda datos está como `ContainerDb` y los externos @@fuera@@ del `System_Boundary`?",
            "¿Cada flecha tiene protocolo @@y@@ formato (`HTTPS/JSON`, `TCP/SQL`)? ¿Ninguna quedó muda?",
            "¿Los nombres coinciden con el C4 Context de la p3, y coincidirán luego con Deployment?",
            "¿Los 3 contratos van entre pares de cajas @@distintos@@?",
            "¿Cada contrato tiene error de negocio, y al menos uno es un @@409@@ de conflicto? (un 500 no cuenta)",
            "¿Los 3 riesgos responden cosas distintas: qué caja se cae (y qué sigue vivo), cuántos saltos de red, y qué dato se escribe en dos pasos?",
        ],
    },
    6: {
        "contexto": [
            "@@Por qué importa al PI:@@ seguridad = amenazas del dominio + controles visibles.",
            "Si la API key está en el Dockerfile, ya filtraron el secreto.",
            "STRIDE-lite: 5 amenazas concretas, no lista genérica de internet.",
        ],
        "escenario": [
            "Actividad individual. Se parte del @@C4 Containers@@ del Corte 1: hoy no se dibuja nada nuevo, se señala sobre lo que ya existe.",
            "Amenazas típicas del curso, como referencia de la @@forma@@ y no para copiarlas: secretos en la imagen, API sin autenticación, registros que guardan tokens, datos personales sin TLS.",
            "Son tres respuestas escritas dentro de ExamLab: las cinco amenazas, la tabla de tres columnas y la política de secretos. La @@plantilla@@ de la sección siguiente trae la estructura exacta que se califica.",
        ],
        # Eran 3 pistas para 3 preguntas de 25 puntos, y dejaban sin cubrir la mitad de
        # la rubrica: la frecuencia de rotacion, el responsable, el procedimiento ante
        # filtracion y la forma que se exige a cada amenaza. El checklist tiene que
        # cubrir TODAS las preguntas, no las primeras.
        "pistas": [
            "¿Cada una de las cinco amenazas nombra el actor o el dato concreto de @@tu@@ dominio @@y@@ el camino por el que ocurre?",
            "¿Ninguna es una frase de manual que sirva igual para cualquier sistema? (esa vale la mitad)",
            "¿Ninguna repite otra con otras palabras? ¿Listaste amenazas y no controles? («falta HTTPS» es la pregunta 2)",
            "¿Cada control es @@verificable@@: puedes decir en una línea qué se hace para comprobar que está puesto?",
            "¿La tercera columna nombra una @@caja o una flecha@@ que existe tal cual en tu C4 Containers o en tu Despliegue? (son 2.5 pts)",
            "¿Ningún control cae en la aplicación web? Ocultar un botón no es un control: la petición se envía sin abrir la interfaz.",
            "¿Aparece el @@menor privilegio@@ sobre un componente concreto, diciendo qué @@deja de poder hacer@@? (las dos mitades pesan igual)",
            "¿Dónde viven los secretos, y qué archivo se versiona @@sin@@ valores reales?",
            "¿«Quién los rota» es un rol responsable? «Se rotan automáticamente» no responde quién.",
            "¿«Cada cuánto» es un número o un evento del calendario? «Periódicamente» no es una frecuencia.",
            "¿Lo prohibido incluye el `Dockerfile`, el `README` @@y@@ el YAML en claro?",
            "¿El primer paso ante una filtración es @@rotar@@ la credencial, y no borrar el commit?",
        ],
    },
    7: {
        "contexto": [
            "@@Por qué importa al PI:@@ sin zonas, el Deployment no demuestra fronteras de confianza.",
            "Son @@tres@@ zonas: pública, privada y de datos. Si la BD queda en la pública, el diagrama ya falló.",
            "Nombres del Deployment deben = nombres del C4 Containers.",
        ],
        "escenario": [
            "Cliente -> edge -> app -> datos.",
            "Sin inventar subnets AWS; trust boundaries sí.",
            "El diagrama se entrega como @@código Mermaid pegado en ExamLab@@, no como imagen.",
        ],
        # El checklist cubre las TRES preguntas de hoy, criterio por criterio calificado:
        # antes solo tenía 3 líneas y dejaba fuera puertos, fronteras, render y renombres.
        "pistas": [
            "P4 · ¿Están las TRES zonas rotuladas: pública, privada y de datos?",
            "P4 · ¿La base de datos quedó en la zona de datos y no en la pública?",
            "P4 · ¿Cada componente lleva su puerto?",
            "P4 · ¿Marcaste la frontera de confianza (la flecha hacia lo que no controlas)?",
            "P4 · ¿Pegaste el Mermaid y lo VISTE renderizado antes de enviar?",
            "P4 · ¿Quitaste todo nombre de subred o de servicio de un proveedor concreto?",
            "P5 · ¿Cada fila usa una de las tres palabras: Relacional, Bloque u Objeto?",
            "P5 · ¿Cada justificación nombra la característica del dato, no una preferencia?",
            "P5 · Si no manejas archivos, ¿declaraste que no necesitas objeto?",
            "P6 · ¿Hay una fila por cada componente de los DOS diagramas, con su zona?",
            "P6 · ¿Listaste los renombres, o dijiste explícitamente que no hubo ninguno?",
        ],
    },
    8: {
        "contexto": [
            "@@Por qué importa al PI:@@ CI que solo hace echo success no es CI.",
            "GitHub Actions free = evidencia de build/test sin tarjeta.",
            "Observabilidad: 4-6 métricas atadas al dominio (golden signals-lite).",
        ],
        "escenario": [
            "Repo free + stub de Clase 3 (o mínimo).",
            "Secrets en Settings; nunca en el YAML en claro.",
            "Misma imagen y mismo puerto del Dockerfile del Corte 1: @@cloudlite-api:0.1.0@@ y 8080.",
        ],
        # Cubre las CUATRO preguntas de hoy: antes dejaba fuera el orden de los pasos, el
        # rótulo de simulado, la condición de fallo y la señal de tipo registro.
        "pistas": [
            "P7 · ¿El YAML tiene los tres bloques: disparadores (on), entorno (runs-on) y pasos?",
            "P7 · ¿Los pasos están en orden: construcción, prueba y despliegue simulado?",
            "P7 · ¿El último paso dice SIMULADO en su propio nombre?",
            "P7 · ¿La imagen y el puerto son los del Dockerfile del Corte 1?",
            "P7 · ¿NINGÚN secreto escrito en claro dentro del YAML?",
            "P8 · ¿Puedes nombrar un error concreto que pondría este check en rojo?",
            "P9 · ¿Dices que tu pipeline llega hasta «listo para desplegar»?",
            "P10 · ¿Las 4-6 señales se refieren a operaciones de TU dominio?",
            "P10 · ¿CADA señal tiene umbral u objetivo?",
            "P10 · ¿Al menos una señal es un registro y no una métrica numérica?",
        ],
    },
    10: {
        "contexto": [
            "@@Por qué importa al PI:@@ lo más caro suele ser lo que dejan encendido sin usar.",
            "Costo cualitativo B/M/A es aceptable: no inventar precios USD de cloud de pago.",
        ],
        "escenario": [
            "Componentes: API, DB, object, CI, edge.",
            "Drivers: idle, egress, storage, minutos CI.",
        ],
        "pistas": [
            "¿Cada fila tiene driver (no solo \"caro\")?",
            "¿Las 3 acciones son verificables en el PI?",
            "¿Evitaron inventar facturas de AWS/GCP?",
        ],
    },
    11: {
        "contexto": [
            "@@Por qué importa al PI:@@ checkpoint v1 = integrar evidencias; no es sustentación ni Parcial 3.",
            "Si C4 y Deployment no comparten nombres, el PI está roto.",
        ],
        "escenario": [
            "Revisar: dominio, ADR, C4, Deployment, Dockerfile, Actions, Seguridad, Costos.",
            "Demo corta por estudiante (o por equipo, si el docente los autorizó) si el tiempo alcanza.",
        ],
        "pistas": [
            "¿Cada \"sí\" tiene enlace o ruta de archivo?",
            "¿Hay backlog priorizado (no lista infinita)?",
            "¿Anti-patrones (teatro microservicios / secretos / CI vacío) marcados?",
        ],
    },
    12: {
        "contexto": [
            "@@Por qué importa al PI:@@ sin métrica objetivo, \"rápido\" es opinión — no arquitectura.",
            "Ensayo de pitch 5-8 min reduce riesgo en la sustentación (Clase 15).",
        ],
        "escenario": [
            "Pico del dominio (ej. inicio de semestre / día de citas).",
            "Demo permitida: diagrama + captura lab/Actions — no K8s de pago.",
        ],
        "pistas": [
            "¿p95 / error rate / RPS tienen número u orden de magnitud?",
            "¿Bottleneck nombrado (DB/auth/storage)?",
            "¿Guion de pitch con tiempos por sección?",
        ],
    },
    13: {
        "contexto": [
            "@@Por qué importa al PI:@@ escalar la API no escala sola la base de datos.",
            "Política de autoescalado conceptual = evidencia de diseño (sin cloud de pago).",
        ],
        "escenario": [
            "Plantilla: componente, tipo, trigger up/down, min/max, cooldown.",
            "Impacto en costo cualitativo (enlace a Clase 10).",
        ],
        "pistas": [
            "¿Triggers medibles (RPS, p95, cola)?",
            "¿Hay techo (max) para no escalar infinito?",
            "¿DB/sesión sticky justificados como no-escalables?",
        ],
    },
    15: {
        "contexto": [
            "@@Por qué importa al PI:@@ sustentación = evidencias + decisiones, no tour de logos.",
            "PI 20% Corte 3 no sustituye el Parcial 3 (ya ocurrió en Clase 14).",
        ],
        "escenario": [
            "Orden del pitch: problema -> arquitectura -> contenedor/CI -> seguridad/costos/escala -> Q&A.",
        ],
        "pistas": [
            "¿Evidencias mínimas presentes (diagramas + lab + CI + decisiones)?",
            "¿Tiempos del pitch ensayados?",
            "¿PI y Parcial 3 no se confunden en el discurso?",
        ],
    },
}


def build_taller_docx(c: dict) -> Path | None:
    if c["tipo"] == "parcial":
        return None
    n = c["n"]
    tb = TALLER_BLOQUE.get(n, {})
    folder = CURSO / "Clases" / f"Clase {n} - {c['slug']}"
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{c['taller_titulo']}.docx"
    doc = Document()
    margins(doc)
    banda(doc, c["taller_titulo"])
    para(doc, "Arquitectura de Sistemas Computacionales · CloudLite App (PI)",
         size=11, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Documento estudiante — avance del Proyecto Integrador",
         size=10, color=CIAN_D, align=WD_ALIGN_PARAGRAPH.CENTER, shade="E8F4FA")
    # Las secciones se numeran corridas: cuatro de ellas son condicionales y con los
    # numeros escritos a mano un taller sin «contexto» empezaba en «2.» y otro sin
    # «escenario» saltaba del 3 al 5.
    _sec = 1
    if tb.get("contexto"):
        h2(doc, f"{_sec}. Contexto / por qué importa al PI")
        bullets(doc, tb["contexto"])
        _sec += 1
    h2(doc, f"{_sec}. Hoy avanzamos el PI en…")
    para(doc, c["pi_hoy"])
    _sec += 1
    h2(doc, f"{_sec}. Entregable concreto")
    para(doc, c["entregable"], shade="E8F4FA")
    _sec += 1
    if tb.get("escenario"):
        h2(doc, f"{_sec}. Escenario / datos de partida")
        bullets(doc, tb["escenario"])
        _sec += 1
    h2(doc, f"{_sec}. Herramientas (gratis + navegador)")
    para(doc, c["herramienta"])
    para(doc, "Prohibido como requisito: AWS/GCP/Oracle/Azure con tarjeta; VirtualBox/VMware/Docker Desktop obligatorio.",
         shade="FBE4E4")
    _sec += 1
    h2(doc, f"{_sec}. Pasos guiados")
    bullets(doc, _pasos(c))
    _sec += 1
    if PLANTILLA_TALLER.get(n):
        # La plantilla va DENTRO del taller a proposito: si el formato del entregable se
        # califica, el formato se entrega, no se adivina.
        h2(doc, f"{_sec}. Plantilla del entregable (copia esto y llénalo)")
        para(doc, "Esta es exactamente la estructura que se califica. Copia el bloque tal cual "
                  "en tu documento, llénalo, y pega cada parte en la pregunta de ExamLab que "
                  "corresponda. Los nombres no se cambian.", size=10)
        for linea in PLANTILLA_TALLER[n]:
            para(doc, linea or " ", size=9, shade="F2F2F3", space_after=0, font=MONO)
        para(doc, " ", size=6, space_after=0)
        _sec += 1
    h2(doc, f"{_sec}. Criterio de éxito")
    bullets(doc, [
        "El artefacto queda en el paquete PI (informe y/o repo) con nombres consistentes.",
        mod(c, "explica_60s_note"),
        "Evidencia adjunta (PNG, enlace lab, YAML, etc.).",
    ])
    _sec += 1
    if tb.get("pistas"):
        h2(doc, f"{_sec}. Pistas (checklist vacío — sin solución)")
        bullets(doc, [f"☐ {p}" for p in tb["pistas"]])
        _sec += 1
    _dudas = soluciones.DUDAS_ESTUDIANTE.get(n)
    if _dudas:
        # Van en el documento del estudiante a proposito: son las mismas preguntas
        # que aparecen cada semestre en la hora de taller.
        h2(doc, f"{_sec}. Dudas frecuentes (lee esto antes de preguntar)")
        for preg, resp in _dudas:
            para(doc, preg, bold=True, size=10.5, space_after=2)
            para(doc, resp, size=10.5)
        _sec += 1
    h2(doc, f"{_sec}. Entrega")
    _sec += 1
    if c["tipo"] == "sustentacion":
        # No es un taller con plazo del domingo: la sesión es la sustentación en vivo,
        # así que el paquete tiene que estar arriba ANTES del bloque.
        para(doc, "El paquete final se sube a ExamLab (https://uniaj.examlab.workers.dev/ · módulo Proyectos) "
                  "ANTES del bloque de sustentaciones: quien llega a subir archivos consume su propio "
                  "turno. La sustentación es en vivo (5–8 min de pitch + Q&A) en la sesión de clase; no "
                  "se reemplaza por un video grabado. " + mod(c, "entrega_unidad_note"))
    else:
        _cierre = _cierre_de(c)
        if _cierre:
            _act = ACTIVIDAD_DE_CLASE[c["n"]]
            # «compartida con» excluye la clase de hoy: con las cuatro dentro, el taller
            # de la Clase 6 decia «compartida con las Clases 6, 7, 8 y 10».
            _hermanas = [str(x) for x in _act["clases"] if x != c["n"]]
            _otras = (" y ".join([", ".join(_hermanas[:-1]), _hermanas[-1]])
                      if len(_hermanas) > 1 else _hermanas[0])
            para(doc, "Entrega en ExamLab (https://uniaj.examlab.workers.dev/ · módulo Talleres). "
                      f"Las preguntas de hoy son parte de UNA sola actividad, compartida con las "
                      f"Clases {_otras}: se guarda el avance de hoy y la actividad se entrega "
                      f"completa el {_cierre}. No hay entrega este domingo. "
                      + mod(c, "entrega_unidad_note"))
        else:
            para(doc, "Entrega en ExamLab (https://uniaj.examlab.workers.dev/ · módulo Talleres) · domingo 23:59 (regla del Acuerdo). "
                      + mod(c, "entrega_unidad_note"))
    if c.get("entrega_oficial_nota"):
        para(doc, c["entrega_oficial_nota"], shade="E8F4FA")
    # 10. Que encuentra en la plataforma. Antes el taller decia «suba el PNG a ExamLab»
    # sin explicar en que forma se responde; ademas pedia exportar de draw.io cuando la
    # plataforma dibuja Mermaid nativo (incluido C4). Esto lo hace explicito.
    _taller_el = TALLERES_EXAMLAB.get(c["n"])
    if _taller_el:
        examlab_talleres.render_estudiante(
            doc, _taller_el, para=para, bullets=bullets,
            add_inline=add_inline_docx, color_titulo=AZUL,
            titulo=f"{_sec}. Qué vas a resolver en ExamLab",
        )
    doc.save(str(path))
    print("OK taller ->", path)
    return path


def build_quiz_docx(c: dict) -> Path | None:
    """Genera DOS archivos: version proyectable (sin respuestas) + clave docente
    (solo respuestas). Antes se generaba un unico docx con la respuesta debajo de
    cada pregunta, lo que lo volvia imposible de proyectar sin regalar la clave."""
    if c["tipo"] == "parcial" or not c.get("quiz"):
        return None
    n = c["n"]
    kit = CURSO / "Kit docente" / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)

    # Version estudiante: solo enunciados, proyectable en clase.
    path = kit / f"Quiz Clase {n} - {c['slug']}.docx"
    doc = Document()
    margins(doc)
    banda(doc, f"Quiz Clase {n} — {c['tema']}")
    para(doc, "Version estudiante — SOLO enunciados. No proyectar la Clave docente.",
         size=10, bold=True, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, shade="E8F4FA")
    para(doc, "5–8 min · individual · al servicio de verificar el avance PI de hoy.")
    for i, (q, _a) in enumerate(c["quiz"], 1):
        h2(doc, f"Pregunta {i}")
        para(doc, q)
    doc.save(str(path))
    print("OK quiz ->", path)

    # Clave docente: enunciado + respuesta, documento privado.
    clave_path = kit / f"Quiz Clase {n} - CLAVE DOCENTE.docx"
    doc_k = Document()
    margins(doc_k)
    banda(doc_k, f"Quiz Clase {n} — CLAVE DOCENTE")
    para(doc_k, "DOCUMENTO DOCENTE — PRIVADO. No proyectar en Presentacion.pptx.",
         size=10, bold=True, color=ROJO, align=WD_ALIGN_PARAGRAPH.CENTER, shade="FBE4E4")
    for i, (q, a) in enumerate(c["quiz"], 1):
        h2(doc_k, f"Pregunta {i}")
        para(doc_k, q)
        para(doc_k, f"Respuesta: {a}", shade="E8F4FA")
    doc_k.save(str(clave_path))
    print("OK clave ->", clave_path)
    return path


# Imagenes de SALIDA ESPERADA por clase (generadas por config/slides/mockups.py).
# El guion las embebe con [[captura: archivo.png]]; si falta el archivo,
# guion_md_a_docx.py deja la caja "inserta aqui la captura" sin romper el build.
CAPTURAS_CLASE = {
    1: [("C4 Context de la demo en vivo: asi debe quedar el tablero al terminar",
         "demo-clase01.png")],
    3: [("Build y run del stub en el lab del navegador (lo que debe verse en pantalla)",
         "salida-docker-build-run.png"),
        ("Evidencia del entregable: el contenedor corriendo (`docker ps`)",
         "salida-docker-ps.png")],
    6: [("Por que un secreto NUNCA va dentro de la imagen (demo de 1 minuto)",
         "salida-secreto-en-imagen.png")],
    8: [("Run verde del workflow: build + test reales, no un `echo ok`",
         "salida-actions-run.png")],
}


def _capturas_md(n: int) -> str:
    items = CAPTURAS_CLASE.get(n)
    if not items:
        return ""
    return "".join(f"📸 {cap} [[captura: {fn}]]\n" for cap, fn in items)


# ---------------------------------------------------------------------------
# Guion docente: material que el docente necesita ADEMAS del fundamento teorico.
# Antes el plan minuto a minuto era una plantilla generica ("recorre las slides
# de conceptos") que no le decia al docente que hacer ni que decir. Estos tres
# diccionarios son lo que convierte el guion en algo dictable sin preparacion:
#   DEMO_ARQ      -> la demo concreta, paso a paso, que el docente debe repetir
#   ERRORES_ARQ   -> lo que el estudiante hace mal y como corregirlo en el momento
#   PREGUNTAS_ARQ -> preguntas de comprobacion oral (no del quiz) para el cierre
# ---------------------------------------------------------------------------

DEMO_ARQ = {
    1: ("Dibujar en vivo el C4 Context de un CloudLite de ejemplo", [
        "Abra draw.io en blanco y dibuje UNA caja al centro rotulada «CloudLite App».",
        "Agregue 2 monigotes a la izquierda (Usuario final, Administrador) con flechas rotuladas «consulta», «administra».",
        "Agregue 1 caja gris a la derecha rotulada «Pasarela de pagos (externo)» y una flecha «cobra».",
        "Diga en voz alta: «no dibuje que hay ADENTRO de la caja; eso es Clase 4».",
    ]),
    # Los 6 rotulos de la demo son los que califican las preguntas 6 y 7. Antes eran
    # otros cuatro (Contexto, Opciones, Decision, Consecuencias) y el estudiante copiaba
    # del tablero una estructura que la plataforma penalizaba.
    2: ("Llenar un ADR-001 delante del grupo, con sus 6 secciones rotuladas", [
        "Abra un Google Doc y escriba los 6 encabezados en orden: 1. Titulo · 2. Estado · 3. Contexto · 4. Decision · 5. Alternativas descartadas · 6. Consecuencias.",
        "Titulo: «ADR-001 Modelo de servicio dominante de CloudLite App». Estado: «Aceptado» y la fecha de hoy. Diga en voz alta: «estos dos rotulos valen 1.5 puntos y son los que se citan en la sustentacion».",
        "Contexto: «lo desarrolla una persona en doce semanas, sin presupuesto ni tarjeta, y tiene que estar en linea el dia de la sustentacion». Subraye que son RESTRICCIONES: «existen tres modelos y hay que elegir uno» no es contexto, es el apunte de clase.",
        "Decision, en una sola frase: «la aplicacion de CloudLite se despliega sobre PaaS». Tache en vivo un segundo modelo si alguien lo propone: «esta seccion vale cero si nombra dos».",
        "Alternativas descartadas, exactamente dos: IaaS, porque habria que operar el sistema operativo sin tiempo para ello; SaaS como nucleo, porque no quedaria arquitectura que disenar. Aclare aqui —y no en la decision— que identidad y correo siguen siendo SaaS satelite.",
        "Consecuencias: escriba UN eje (operacion) con su + y su -, y deje los otros dos al grupo. Diga: «un ADR de una pagina que se entiende vale mas que 5 paginas que nadie lee».",
    ]),
    # La demo es la bitacora de la pregunta 11 ejecutada delante del grupo, en el mismo
    # orden y con los mismos cinco comandos. Antes construia un nginx sin etiqueta con
    # `-p 80:80`: el estudiante copiaba del tablero un build que pierde 1.5 puntos (sin
    # etiqueta) y un mapeo con los dos lados iguales, que es justo el que no permite
    # distinguir anfitrion de contenedor cuando la pregunta 10 lo pide explicado.
    3: ("Construir, correr y verificar el stub en Killercoda — los 5 comandos de la bitacora", [
        "Abra killercoda.com, inicie sesion con la cuenta gratuita y lance un escenario Ubuntu (advierta en voz alta: la sesion caduca a 1 h, guarden capturas antes de cerrarla).",
        "Escriba el Dockerfile del stub en vivo, en el mismo orden de la diapositiva «Dockerfile minimo del stub CloudLite»: FROM node:20-alpine, WORKDIR, COPY package*.json, RUN npm ci --omit=dev, COPY . ., EXPOSE 8080, CMD. Y cree al lado un `.dockerignore` con `.env` y `node_modules` — diga: «sin este archivo, el COPY . . se lleva el .env a la imagen y son 5 puntos».",
        "Comando 1 — `docker build -t cloudlite-api:0.1.0 .` Senale la etiqueta `0.1.0`: «sin ella la imagen queda como latest y la de hoy no es la de manana». Senale en el log que `COPY package*.json` corre ANTES que `COPY . .`.",
        "Comando 2 — `docker images | grep cloudlite-api` y lea en voz alta el TAG y el SIZE: «esto es lo que va en la fila 2 de la bitacora, pegado, no descrito».",
        "Comando 3 — `docker run -d -p 8081:8080 --name api cloudlite-api:0.1.0`. Escriba en el tablero «8081 = anfitrion, por donde entro yo» y «8080 = contenedor, el del EXPOSE», y aclare por que los puse DISTINTOS: para que se vea cual es cual.",
        "Comando 4 — `docker ps`: senale IMAGE, STATUS y la columna PORTS con `0.0.0.0:8081->8080/tcp`. Ejecute `date` justo antes: «la hora del sistema en la misma captura vale 0.5 puntos».",
        "Comando 5 — `curl -i http://localhost:8081/health` y lea los TRES datos del contrato: la ruta, el `HTTP/1.1 200 OK` y el cuerpo JSON con su campo verificable.",
        "Error a proposito, 60 segundos: pare el contenedor y relancelo con los puertos invertidos (`-p 8080:8081`). `docker ps` sigue diciendo Up y el `curl` se queda colgado: «el sintoma no dice la causa; por eso la pregunta 10 pide explicar que pasa si los inviertes».",
        "Si Killercoda no carga, la alterna es LabEx Docker Playground (ojo: solo 3 sesiones al dia en el plan gratuito); si falla la red, proyecte las capturas de `Kit docente/Clase 3/Capturas/`.",
    ]),
    # La demo termina DENTRO de ExamLab, con el Mermaid pegado y renderizado, porque
    # es ahi donde se califica la pregunta 13 y es el paso que el estudiante no adivina.
    # Antes cerraba en el tablero de draw.io y nadie veia nunca la sintaxis que la
    # plataforma exige (`C4Container` en la primera linea, `ContainerDb` para la base).
    4: ("Convertir el Context de la Clase 1 en Containers, y dejarlo renderizado en ExamLab", [
        "Abra el diagrama C4 Context de la demo de Clase 1 y haga zoom a la caja «CloudLite App». Diga: «hoy no dibujamos otro sistema, abrimos este».",
        "Reemplace esa caja por 3 cajas internas: «App web», «API de turnos» y «Base de turnos». Escriba en cada una sus TRES datos: nombre, tecnologia y responsabilidad en una frase.",
        "Senale la base de datos y diga: «esta no es un Container mas, es un ALMACEN; en el codigo va como ContainerDb y son 2 puntos». Deje el cliente y el correo FUERA del recuadro del sistema.",
        "Rotule CADA flecha con protocolo Y formato: «HTTPS/JSON», «TCP/SQL». Borre a proposito una etiqueta y pregunte que se pierde: sin ella nadie puede decir por donde se rompe.",
        "Proponga una cuarta caja, el worker de avisos, y pida la razon de negocio. Si nadie la da, borrela en vivo: «eso es microservicios teatro». Si alguien la da (el correo tarda y puede fallar), quedese con ella y anote la razon al lado.",
        "Verifique nombre por nombre contra el C4 Context de la Clase 1: si alli decia «Pasarela de pagos», aqui no puede decir «Pagos». Son 2 puntos de la pregunta 13.",
        "Cierre en ExamLab: pegue el codigo Mermaid de la diapositiva del molde, cambie los nombres por los del ejemplo del tablero y proyecte el resultado RENDERIZADO. Diga: «si no renderiza, no hay diagrama; se revisa antes de enviar».",
    ]),
    6: ("De amenaza STRIDE a control verificable, en vivo", [
        # El paso 1 decia «alguien cambia el precio de un item»: ningun dominio del curso
        # vende nada, y «alguien» es justo el sujeto vago que la pregunta 1 penaliza. Y el
        # paso 3 pedia «en que archivo o diagrama se ve» — un archivo NO suma en la
        # pregunta 2, que reparte 2.5 pts por senalar una caja o una flecha. El docente
        # estaba modelando en vivo la respuesta que su propia rubrica descuenta.
        "Escriba en el tablero, con las dos partes que exige la rubrica: «Tampering: un cliente mueve la franja de un turno ajeno porque la API no revisa de quien es el turno».",
        "Pregunte al grupo cual seria el control; guie hasta «validar el rol y la propiedad del turno antes de aceptar el cambio».",
        "Agregue la tercera columna preguntando «sobre que CAJA o sobre que FLECHA del C4 Containers cae ese control». Aqui la respuesta es la caja «API de turnos». Un nombre de archivo no vale: si no se puede senalar en el diagrama, el control todavia es una intencion.",
        "Repita con una segunda fila cuyo control caiga en una FLECHA, para que se vea que las dos formas cuentan: «un cliente reserva a nombre de otro» -> «el id se toma del token» -> flecha «App web -> API de turnos».",
        "Demo de 1 minuto del anti-patron, con la diapositiva del historial de capas proyectada: un Dockerfile con la llave en texto plano, el `docker history` que la lee, y el `rm` posterior que no la borra sino que la tapa.",
    ]),
    # La demo dibujaba DOS zonas y dejaba la base de datos en la privada. La pregunta 4
    # califica TRES zonas y da cero en los 4 pts de ubicacion si la base de datos no esta
    # en la zona de datos: el docente estaba proyectando el diagrama que la rubrica castiga.
    7: ("Del boceto de tres zonas al Mermaid que se califica", [
        "En draw.io o Excalidraw dibuje TRES rectangulos, rotulados «Zona publica», «Zona privada» y «Zona de datos».",
        "Reparta las cajas de CloudLite: `Edge / balanceador` y `App web` en la publica, `API CloudLite` en la privada, `Base de datos` en la de datos — nunca en la publica. El `Cliente / navegador` va FUERA de las tres zonas: es el actor, no algo que usted despliegue, y esa es una de las dos filas sin par de la pregunta 6.",
        "Etiquete cada flecha con su puerto (443 al edge, 8080 a la API, 5432 a la base de datos) y saque una flecha aparte a la `Pasarela de pagos` externa: ahi esta la frontera de confianza, y son 2 de los 14 pts.",
        "Pregunte: «si un atacante llega desde internet, con que se topa primero?» — eso es superficie de exposicion.",
        "Traduzca ese boceto a Mermaid (el codigo de referencia esta abajo), peguelo en la pregunta 4 de ExamLab y proyectelo RENDERIZADO: 2 de los 14 pts son que renderice sin error.",
        "Verifique en voz alta que los nombres de los servicios son LOS MISMOS del C4 Containers de la Clase 4.",
    ]),
    # La demo montaba «3 steps: checkout, setup y una prueba»: le faltaban los dos pasos que
    # la pregunta 7 califica (la construccion y el despliegue simulado rotulado) y no hacia
    # la prueba mental de la condicion de fallo, que son otros 2 pts en la pregunta 8. Los
    # nombres de los pasos son los de la diapositiva del ci.yml, para que el codigo
    # proyectado y la captura del kit digan lo mismo.
    8: ("Un workflow de GitHub Actions que corra de verdad, con los tres pasos calificados", [
        "Cree `.github/workflows/ci.yml` copiando la diapositiva del ci.yml: `on: [push, pull_request]`, `runs-on: ubuntu-latest` y los pasos en ORDEN — Construir, Probar, Despliegue SIMULADO.",
        "Senale los tres bloques mientras los escribe: «disparadores, entorno y pasos: son 2, 1.5 y 4 puntos de la pregunta 7».",
        "En el paso Construir use la MISMA imagen de la Clase 3: `npm ci && docker build -t cloudlite-api:0.1.0 .` — la coherencia con el Dockerfile del Corte 1 vale 1 pt.",
        "Haga commit y push, abra la pestana Actions y espere el check verde; abra el log del paso Probar: «esto es evidencia, no una diapositiva que dice que tenemos CI».",
        "Rompa el pipeline a proposito, 60 segundos: cambie la asercion de la prueba (o borre `server.js`), haga push y muestre el check ROJO. Diga: «esta es la respuesta de la pregunta 8: con que condicion falla. Si no pueden romperlo, no estan validando nada».",
        "Vuelva a dejarlo verde y lea en voz alta el nombre del ultimo paso: «Despliegue SIMULADO (no despliega a ningun servidor)». Aclare la frontera: el pipeline llega hasta «listo para desplegar», y decirlo asi SUMA en la pregunta 9 — afirmar que ya hay CD resta la mitad.",
        "Cierre en Settings > Secrets and variables > Actions: «los secretos viven aqui y se referencian por nombre. Un secreto escrito en claro dentro del YAML es cero en toda la pregunta 7».",
    ]),
    10: ("Tabla de costo cualitativo en 5 minutos", [
        "Dibuje 3 columnas: Componente | Costo (Bajo/Medio/Alto) | Driver del costo.",
        "Llene 3 filas de CloudLite: base de datos gestionada (Alto, computo+almacenamiento constante 24/7), API en contenedor (Medio, numero de instancias), object storage de imagenes (Bajo, volumen de datos).",
        "Pregunte cual bajaria primero si el presupuesto se corta a la mitad, y exija que justifiquen con el driver, no con intuicion.",
    ]),
    11: ("Auditar en vivo el paquete de un voluntario", [
        "Pida a un estudiante voluntario (o a un equipo, si autorizo equipos) que proyecte su C4 Containers y su diagrama de despliegue lado a lado.",
        "Compare nombre por nombre: todo servicio del Containers debe existir en el despliegue y viceversa.",
        "Senale en voz alta el primer gap concreto que encuentre y escribalo como accion con responsable y fecha.",
        "Modele el tono: el hallazgo es sobre el artefacto, nunca sobre la persona.",
    ]),
    12: ("Definir un objetivo de rendimiento que si se puede verificar", [
        "Escriba la frase mala: «la app debe ser rapida». Pregunte al grupo como la comprobarian; deje que fallen.",
        "Reescribala en vivo: «el p95 del endpoint de consulta responde en menos de 300 ms con 50 peticiones por segundo».",
        "Explique el p95 con 20 numeros en el tablero: ordene y marque el que deja 95% por debajo.",
        "Cierre pidiendo el bottleneck sospechado: «cual pieza creen que revienta primero, y por que esa».",
    ]),
    13: ("Vertical vs horizontal, y lo que NO escala", [
        "Dibuje una caja «API» y agrandela: eso es vertical (mas CPU/RAM a la misma maquina, con techo fisico).",
        "Borre y dibuje 3 cajas «API» iguales con un balanceador arriba: eso es horizontal.",
        "Agregue la base de datos abajo, conectada a las 3, y encierrela en rojo: «esta no se multiplica igual; aqui esta el limite real».",
        "Escriba el trigger y el limite: «CPU > 70% por 5 min -> +1 instancia, maximo 4» y amarre con el costo de la Clase 10.",
    ]),
    15: ("Modelar una sustentacion de 6 minutos y un Q&A", [
        "Presente usted mismo un CloudLite de ejemplo en 6 minutos cronometrados, con la estructura: problema, decision clave, evidencia, limite conocido.",
        "Hagase una pregunta dificil en voz alta y respondala: «por que no uso microservicios? Porque el proyecto lo sostiene una sola persona y la frontera no se justificaba».",
        "Muestre la rubrica proyectada y senale donde habria perdido puntos su propia demo.",
        "Recuerde la regla de los 60 segundos: quien sustenta debe poder explicar cualquier parte del paquete, y si hubo equipo autorizado, cualquier integrante.",
    ]),
}

ERRORES_ARQ = {
    1: ["Dominio vago tipo «una red social» o «un e-commerce»: sin problema concreto no hay decisiones que tomar. Exija sector, usuario y dolor observable.",
        "Dibujar lo que hay DENTRO del sistema en el nivel Context (base de datos, API). Se corrige recordando que eso es el nivel Containers de la Clase 4.",
        "Confundir capacidad con pantalla: «tener un login» no es capacidad; «autenticar usuarios» si."],
    2: ["Elegir el modelo de servicio por moda y no por trade-off. Pida la frase «aceptamos perder X para ganar Y» escrita en el ADR.",
        "ADR sin alternativas descartadas: un ADR con una sola opcion no documenta una decision, documenta un hecho.",
        "Nombrar productos de marca en vez del modelo conceptual; el modelo aplica a cualquier proveedor."],
    3: ["Decir que el contenedor «es una VM ligera». Insista en la diferencia real: kernel propio vs kernel compartido.",
        "Confundir imagen con contenedor al hablar. Corrija en el momento: la imagen es el molde, el contenedor la instancia corriendo.",
        "Perder el trabajo porque la sesion del lab caduco a la hora. Es el error mas comun del dia: recuerdeles que el Dockerfile se escribe en la carpeta del PI y se PEGA en el lab, nunca al contrario."],
    4: ["Inventar 6 u 8 servicios para verse sofisticados. Pregunte por cada uno: que responsabilidad de negocio propia tiene y quien lo despliega por separado.",
        "Flechas sin etiqueta, o con media etiqueta. No basta «HTTP» ni «SQL»: toda flecha lleva protocolo Y formato de datos («HTTPS/JSON», «TCP/SQL»).",
        "Marcar la base de datos como `Container` y no como `ContainerDb`. Son 2 puntos y se pierden en un solo caracter; revise ese renglon del codigo antes de que envien.",
        "Renombrar las cajas respecto al C4 Context de la pregunta 3 («Pagos» donde antes decia «Pasarela de pagos»). Pidales los dos diagramas lado a lado y compare palabra por palabra.",
        "Responder «un poco de los dos» en la decision de la pregunta 12: vale cero. Devuelvala pidiendo UNA opcion y las dos mitades del trade-off, lo que se gana y lo que se pierde.",
        "Riesgos genericos tipo «los microservicios son mas complejos» o «puede haber latencia». No nombran caja, ni salto, ni dato: exija los tres riesgos concretos que pide el enunciado, y en el primero, que digan tambien que SIGUE funcionando."],
    # El tercer punto decia «3 bien argumentadas»: la pregunta 1 reparte 1.75 pts por
    # amenaza hasta CINCO, asi que el guion le estaba dando al docente un numero que le
    # cuesta 3.5 puntos al estudiante. Y el primero hablaba de la columna «evidencia»,
    # que en la rubrica se llama «donde se ve (caja o flecha)» y no admite un archivo.
    6: ["Entregar una lista generica de buenas practicas en vez de amenaza -> control -> donde se ve. Devuelva la tabla si no tiene las 3 columnas con esos nombres.",
        "Escribir credenciales en el Dockerfile o en el repositorio. Es el error mas costoso y hay que cortarlo el mismo dia.",
        "Cubrir las 6 categorias STRIDE de forma superficial en vez de las CINCO amenazas bien argumentadas que pide el enunciado. Dos amenazas de la misma letra son validas si el camino y el control son distintos.",
        "Amenazas sin sujeto ni camino: «podrian hackear la base de datos». Valen la mitad. La prueba rapida es preguntar si esa frase se podria copiar en el trabajo de otro estudiante sin cambiar nada.",
        "Poner un nombre de archivo en la tercera columna («.dockerignore», «el contrato del endpoint»). No suma: son 2.5 pts que se reparten por senalar una caja o una flecha del diagrama.",
        "Menor privilegio recitado y no aplicado. Pida las dos frases: sobre que componente, y que deja de poder hacer al aplicarlo.",
        "Politica de secretos sin responsable ni frecuencia («deberian rotarse periodicamente»). Son 3 de los 7.5 pts de la pregunta 3, y se pierden por escribir en tercera persona."],
    7: ["Dibujar «la nube» como una caja difusa. Exija las dos zonas, publica y privada, explicitas.",
        "Poner la base de datos en la subred publica «para que sea mas facil probar». Es exactamente lo que la Clase 6 acaba de prohibir.",
        "Renombrar servicios respecto al C4 Containers, con lo que los dos diagramas dejan de ser el mismo sistema."],
    8: ["Un workflow que solo hace `echo ok`: es un pipeline decorativo. Exija que corra algo que pueda fallar de verdad.",
        "Decir que ya tienen CD porque el YAML dice deploy. En este curso el despliegue se simula; que lo digan asi.",
        "Golden signals sin umbral: «medimos latencia» no sirve; falta a partir de que valor se considera un problema."],
    10: ["Pedir precios exactos de un proveedor. No es el objetivo: el analisis es cualitativo y por driver.",
         "Marcar todo como costo «Medio» para no pensar. Fuerce al menos un Alto y un Bajo con justificacion.",
         "Olvidar el trafico de red saliente, que es el driver que mas sorprende en facturas reales."],
    11: ["Traer el paquete de la Clase 1 sin actualizar y presentarlo como avance. Compare contra la version anterior.",
         "Conocer solo la parte que se copio de una plantilla y no el paquete completo. Pregunte al azar por cualquier seccion; si hubo equipo autorizado, pregunte a un integrante distinto del que presenta y si solo uno responde, ese es el hallazgo principal.",
         "Confundir este checkpoint con la sustentacion final o con el Parcial 3. Aclarelo al abrir la sesion."],
    12: ["Objetivo de rendimiento sin numero, sin escenario de carga o sin bottleneck. Falta cualquiera de los tres y no es un analisis.",
         "Usar el promedio en vez del p95 y concluir que todo esta bien. Muestre por que el promedio esconde los casos malos.",
         "Ensayar el pitch leyendo las diapositivas. Cronometre y corte a los 8 minutos."],
    13: ["Prometer autoescalado infinito sin limite maximo ni control de costo (Clase 10).",
         "Escalar horizontalmente un servicio que guarda la sesion en memoria local: al repartir la carga, el usuario pierde su sesion.",
         "No documentar QUE NO escala. La base de datos relacional es casi siempre la respuesta y hay que decirlo."],
    15: ["Describir el diagrama en vez de justificar la decision. Reoriente con «por que asi y no de la otra forma».",
         "En equipos autorizados, que un integrante presente y el resto observe. Distribuya el Q&A a proposito entre todos.",
         "Presentar sin mencionar ningun limite del diseno. Quien no reconoce limites no entendio el trade-off."],
}

PREGUNTAS_ARQ = {
    1: ["Cual es la diferencia entre arquitectura y stack tecnologico?",
        "Que va DENTRO y que va FUERA de la caja en un diagrama C4 Context?",
        "Digan una capacidad de su CloudLite que NO sea una pantalla."],
    2: ["Quien administra el sistema operativo en IaaS, en PaaS y en SaaS?",
        "Que perdieron y que ganaron con el modelo que eligieron?",
        "Por que un ADR necesita las alternativas que descartaron?"],
    3: ["Que comparten los contenedores de una misma maquina que las VM no comparten?",
        "Cual es la diferencia entre imagen y contenedor?",
        "Que pasa con su trabajo cuando caduca la sesion del lab, y donde deberia vivir el Dockerfile?"],
    4: ["Que justifica que dos funciones vivan en servicios separados?",
        "Que cambia cuando una llamada de funcion se vuelve una llamada de red?",
        "Como se llama en su C4 Containers el servicio que expone la API? Se llama igual en su C4 Context de la Clase 1?",
        "Cual de sus cajas es un almacen y como se escribe en el codigo del diagrama?",
        "Cuantos saltos de red tiene la operacion principal de su sistema, contados sobre su propio diagrama?",
        "Si se cae su base de datos, que deja de funcionar y que SIGUE funcionando?"],
    6: ["Que significa la T de STRIDE y una amenaza concreta de su CloudLite?",
        "Donde guardan una API key y por que NO dentro de la imagen?",
        "Sobre que caja o sobre que flecha de SU diagrama cae uno de sus controles? Digalo con el nombre que tiene alli.",
        "Quien rota los secretos de su repositorio y cada cuanto? Un rol y un evento del calendario, no «periodicamente».",
        "Si manana se filtra su cadena de conexion, cual es el PRIMER paso y por que no es borrar el commit?",
        "Sobre que componente aplican menor privilegio, y que deja de poder hacer ese componente al aplicarlo?"],
    7: ["Que va en la subred publica y que en la privada, y por que?",
        "Que hace un balanceador de carga en una frase?",
        "Cuando conviene object storage y cuando la base de datos?"],
    8: ["Que valida CI y que hace CD, y cual de los dos construyeron hoy?",
        "Digan las 4 golden signals y el umbral de una de ellas.",
        "Que pasaria en su pipeline si alguien sube codigo que no compila?"],
    10: ["Cual es el componente mas caro de su CloudLite y cual es su driver?",
         "Que es right-sizing en una frase?",
         "Como se conecta el autoescalado con el costo?"],
    11: ["Que gap concreto identificaron hoy y quien lo cierra?",
         "Su diagrama de despliegue usa los mismos nombres que su C4 Containers?",
         "Que evidencia de la rubrica les falta todavia?"],
    12: ["Que es el p95 y por que no usamos el promedio?",
         "Cual es su bottleneck sospechado y en que se basan?",
         "Diferencia entre stress test y spike test?"],
    13: ["Vertical u horizontal: cual eligieron y por que?",
         "Cual es su trigger y cual su limite maximo?",
         "Que pieza de su sistema NO escala, y que harian al respecto?"],
    15: ["Justifiquen su decision de arquitectura mas importante en 60 segundos.",
         "Cual es el limite conocido de su diseno actual?",
         "Si tuvieran un mes mas, que cambiarian primero y por que?"],
}


def _lista_md(items, bullet="- "):
    return "\n".join(f"{bullet}{x}" for x in items) if items else ""


# Codigo de referencia del diagrama que el docente dibuja en vivo. Existe porque el
# bloque «Demo en vivo» pedia dibujar un diagrama sin dar ninguna referencia de como
# debe quedar: si el proyector o la red fallan, o si el docente prefiere no dibujar a
# mano, pega este codigo en la pregunta de diagrama de ExamLab y lo proyecta
# renderizado en un minuto. Es el mismo contenido que la imagen de Capturas/.
DEMO_MERMAID = {
    1: ("C4 Context de la demo (el mismo de `Capturas/demo-clase01.png`)", """C4Context
    title CloudLite App - nivel Context (demo de clase)
    Person(usuario, "Usuario final", "Consulta y usa el servicio")
    Person(admin, "Administrador", "Configura y opera")
    System(cloudlite, "CloudLite App", "El sistema completo, como caja negra")
    System_Ext(pagos, "Pasarela de pagos", "Servicio de terceros")
    Rel(usuario, cloudlite, "consulta", "HTTPS")
    Rel(admin, cloudlite, "administra", "HTTPS")
    Rel(cloudlite, pagos, "cobra", "API REST sobre HTTPS")"""),
    # El resultado de la demo de la Clase 4 es el mismo Context de arriba con la caja
    # del sistema abierta: los mismos actores y el mismo externo, con los mismos
    # nombres. Es exactamente lo que la pregunta 13 cobra en sus 2 puntos de
    # trazabilidad, y aqui queda al lado para que el docente lo pueda comparar.
    4: ("C4 Container de la demo (el Context de la Clase 1, ya abierto)", """C4Container
    title CloudLite App - nivel Container (demo de clase)
    Person(usuario, "Usuario final", "Consulta y usa el servicio")
    Person(admin, "Administrador", "Configura y opera")
    System_Boundary(cloudlite, "CloudLite App") {
      Container(web, "App web", "React", "Pantallas del usuario y del administrador")
      Container(api, "API CloudLite", "Node.js", "Reglas de negocio y validaciones")
      ContainerDb(db, "Base de datos", "PostgreSQL", "Datos del dominio")
    }
    System_Ext(pagos, "Pasarela de pagos", "Servicio de terceros")
    Rel(usuario, web, "consulta", "HTTPS")
    Rel(admin, web, "administra", "HTTPS")
    Rel(web, api, "pide y envia datos", "HTTPS/JSON")
    Rel(api, db, "lee y escribe", "TCP/SQL")
    Rel(api, pagos, "cobra", "API REST sobre HTTPS")"""),
    # Sobre CloudLite, que es el dominio que se proyecta desde la Clase 1: los nombres son
    # los del C4 Container de la demo de la Clase 4 (`App web`, `API CloudLite`, `Base de
    # datos`, `Pasarela de pagos`) y el puerto de la API es el `EXPOSE 8080` de la Clase 3.
    # Antes esto estaba resuelto sobre AgendaU, que es el dominio del `mermaid_esperado` del
    # kit — el modelo que la solucion declara que el estudiante NO ve. Proyectarlo rompia esa
    # declaracion y ademas invitaba a copiar «API de agenda» en la pregunta 6, que cobra que
    # los nombres sean los del propio C4 Containers.
    7: ("Despliegue en tres zonas de CloudLite (el resultado de la demo)", """flowchart LR
    cliente["Cliente / navegador<br/>Usuario final o Administrador"]
    subgraph publica["Zona publica - internet"]
        edge["Edge / balanceador<br/>443 HTTPS"]
        web["App web<br/>React - estatico<br/>443 HTTPS"]
    end
    subgraph privada["Zona privada - solo alcanzable desde el edge"]
        api["API CloudLite<br/>Node.js<br/>8080 HTTP"]
    end
    subgraph datos["Zona de datos - sin salida a internet"]
        db[("Base de datos<br/>PostgreSQL<br/>5432 TCP")]
    end
    pagos["Pasarela de pagos<br/>externo - 443"]
    cliente -->|"HTTPS 443 - frontera de confianza"| edge
    cliente -->|"HTTPS 443 - descarga el bundle"| web
    edge -->|"HTTP 8080"| api
    api -->|"TCP 5432"| db
    api -->|"HTTPS 443 - frontera de confianza"| pagos"""),
}


def _demo_md(n: int) -> str:
    d = DEMO_ARQ.get(n)
    if not d:
        return "Muestre en vivo la herramienta del dia con un mini-ejemplo de CloudLite.\n"
    titulo, pasos = d
    cuerpo = "\n".join(f"{i + 1}. {p}" for i, p in enumerate(pasos))
    md = f"**Demo que usted debe poder repetir:** {titulo}\n\n{cuerpo}\n"
    ref = DEMO_MERMAID.get(n)
    if ref:
        rotulo, codigo = ref
        md += (
            f"\n**Referencia del resultado:** {rotulo}. Si la red falla o prefiere no "
            "dibujar a mano, pegue este codigo en la pregunta de diagrama de ExamLab y "
            "proyectelo renderizado; tambien sirve para volver a generar la imagen en "
            "cualquier editor que soporte Mermaid.\n\n"
            "```mermaid\n" + codigo + "\n```\n"
        )
    return md


def _guion_parcial_md(c: dict) -> str:
    """Guion del dia de parcial.

    Era un esqueleto de 1.100 caracteres que no cumplia el checklist del repo: no decia
    la modalidad, no decia por que canal se entrega —y el propio enunciado remite a «el
    medio que el docente indique al abrir la sesion», asi que el guion era el unico
    lugar donde podia estar—, no decia que temas cubre el instrumento y no cerraba con
    las dos secciones obligatorias («errores tipicos del docente» y «preguntas
    frecuentes del grupo»). En un bloque de 120 minutos donde 90 son de silencio, lo
    unico que el guion tiene que resolver son los 30 restantes, y eran justo los que
    faltaban.
    """
    # Mismo mecanismo que el fundamento de una clase regular: el guion nombra la
    # diapositiva por su titulo y el numero se resuelve contra el deck real, asi que si
    # el deck del dia de parcial cambia de orden el build falla en vez de mandar al
    # docente a proyectar la diapositiva equivocada.
    return _resolver_slides(_guion_parcial_cuerpo(c), _slide_map(c), c["n"])


def _guion_parcial_cuerpo(c: dict) -> str:
    """Cuerpo del guion, con los tokens `{{slide:…}}` todavia sin resolver."""
    n = c["n"]
    p = PARCIALES_ARQ[n]
    corte = p["corte"]
    m = _parcial_meta(corte)
    archivo = m["archivo"] + ".docx"
    ultima = p.get("ultima")
    prep = p.get("prep_pi")
    nota_prep = (f"- La preparación del pitch del PI fue la **Clase {prep}**; hoy no se "
                 "prepara nada del proyecto." if prep else
                 "- En este corte no hay preparación de pitch: el PI se retoma en la "
                 "siguiente clase regular.")
    # Los tokens se arman fuera del f-string: dentro, `{{` es una llave literal
    # escapada y el token salia como «{slide:evalua hoy}», que ni resuelve ni es
    # detectable como marcador crudo por su forma habitual.
    tok_alcance, tok_entrega = "{{slide:evalua hoy}}", "{{slide:Como se responde}}"
    secciones = "\n".join(f"   - {s}" for s in m["secciones_resumen"])
    lista_temas = "\n".join(f"   - {t}" for t in m["temas"])
    no_resp = " · ".join(f"«{d}»" for d in p["dudas_no"])
    # La clase que el grupo cree que no cuenta es justamente la autonoma: no hubo sesion
    # en vivo, asi que alguien va a preguntar si entra. La portada del instrumento la
    # marca «(sesion autonoma)», de ahi se detecta sin volver a escribir el dato. Cuando
    # la autonoma ES la ultima clase evaluada —el caso del Parcial 2, cuya Clase 10 cayo
    # en la sesion del festivo— la aclaracion va dentro de la misma respuesta, para no
    # preguntar dos veces por la misma clase.
    _auton = [f"{t.split(' · ')[0]} («{t.split(' · ', 2)[-1].replace(' (sesión autónoma)', '')}»)"
              for t in m["temas"] if "autónoma" in t]
    _razon = (" y se evalúa igual, aunque se haya trabajado sin sesión en vivo: es la que "
              "más se olvida al estudiar." if _auton else ".")
    if _auton and _auton[0].startswith(f"Clase {ultima} "):
        faq_entra = (f"**¿Entra lo de la Clase {ultima}?** Sí, es lo más reciente que entra"
                     + _razon)
    else:
        faq_entra = f"**¿Entra lo de la Clase {ultima}?** Sí, y es lo más reciente que entra."
        if _auton:
            faq_entra += ("\n\n**¿La clase autónoma también entra?** Sí: "
                          + " y ".join(_auton) + " se trabajó sin sesión en vivo y se "
                          "evalúa igual. Es la que más se olvida al estudiar.")
    frase_cierre = _cierre_parcial_pi(n, hablada=True)
    return f"""# Guion docente — Clase {n}: {c['tema']}

## Información de la clase
- Asignatura: Arquitectura de Sistemas Computacionales (FI303380)
- Duración del bloque: **120 min** · **virtual síncrona por Google Meet**
- Tipo: **Solo evaluación (Parcial {corte})** — sin tema nuevo y sin taller dirigido
- Proyecto Integrador: CloudLite App (hoy **no** se avanza)
- Instrumento: `Parciales/{archivo}` · la solución es el mismo nombre con «- SOLUCION» y **no se publica en `Clases/`**
- Fecha: {m['fecha']} · tiempo de resolución previsto por el instrumento: **{m['tiempo']}**
- Peso: **{m['valor_corte']}** · 100 puntos · nota = puntos / 20 sobre 5.0

## Objetivo del bloque
Aplicar el **{m['titulo_parcial']}**, que evalúa **solo** estas clases de material:

{lista_temas}

Sus cuatro secciones y lo que vale cada una:

{secciones}

## Antes del bloque (10 min de preparación)
1. Abre `Parciales/{archivo}` y **decide el canal de entrega**, porque el enunciado dice
   «entréguelo por el medio que el docente indique al abrir la sesión» y esa frase te
   deja la decisión a ti. Dos opciones que funcionan en Meet:
   - **Documento editable** (recomendado): compartes el `.docx` por el chat del Meet al
     minuto 0, cada estudiante lo llena y lo devuelve por el mismo canal o por correo.
     Ventaja: las líneas del enunciado ya están donde se responde.
   - **Foto de hoja escrita a mano**: solo si la conexión de alguien no aguanta el
     documento. Entonces exige que la foto se lea y que traiga nombre en cada página.
2. Ten el archivo de **solución a la mano pero cerrado**: hoy no se califica en vivo.
3. Revisa que el enunciado no pregunte nada fuera de las clases listadas arriba. Si algo
   se cuela de otro corte, se anula esa pregunta y se reparten sus puntos entre las demás,
   no se descuenta al estudiante.

## Plan minuto a minuto (120 min)

### 0–10 · Encuadre y canal de entrega
Di, con estas dos frases: «Hoy es **solo parcial**: no hay tema nuevo ni taller del
proyecto.» y «El parcial se entrega **por el canal que voy a nombrar ahora** —el que
decidiste en la preparación— **antes del minuto 110; lo que llegue después no se recibe.**»
Verifica asistencia por lista, no por «los que están conectados».
Proyecta la **{tok_alcance}** mientras lo dices: ahí está el alcance —las clases que
entran— y el reparto de puntos por sección, que es lo primero que van a preguntar.
Pasa a la **{tok_entrega}** para el canal, el minuto de cierre y qué dudas vas a
responder; dejarla en pantalla los primeros minutos ahorra la mitad de los mensajes por
privado.
Comparte el enunciado y **confirma en voz alta que todos lo abrieron** antes de arrancar
el reloj: en virtual, el que no lo pudo descargar pierde 15 minutos en silencio.
Di también qué material está autorizado (por defecto: **nada**) y que las cámaras se
quedan encendidas si el Acuerdo lo permite.

### 10–100 · Desarrollo (90 min de silencio de evaluación)
Tú te quedas en el Meet con el micrófono abierto y la cámara encendida: es la única
supervisión que hay. Anuncia el tiempo a los **50** y a los **80** minutos.
**Resuelves dudas de forma, no de contenido.** La línea es esta: si la respuesta a la
duda es un dato que la pregunta evalúa, no se responde.
- Se responde: «¿esta pregunta pide una tabla o un párrafo?», «¿cuántas líneas?»,
  «¿el punto b) es obligatorio?», «no puedo abrir el archivo».
- No se responde, en este parcial: {no_resp}.
Cuando la duda es de contenido, la respuesta es siempre la misma: «Eso es lo que la
pregunta evalúa; responde con lo que recuerdes de la Clase correspondiente.» Dilo igual
para todos, porque en Meet las preguntas llegan por privado y nadie ve que a otro le
dijiste lo mismo.

### 100–110 · Cierre de envío
Aviso de 10 minutos. Recibe las entregas y **acusa recibo por el chat, uno por uno**:
en virtual el estudiante no tiene forma de saber que su archivo llegó, y el reclamo de
«yo sí lo envié» se previene aquí y no después. Anota quién no entregó.

### 110–120 · Cierre del bloque
Di: «Gracias. {frase_cierre}» No adelantes respuestas ni comentes el parcial: todavía hay
quien está subiendo el archivo, y cualquier comentario tuyo se convierte en la respuesta
correcta.

## Errores típicos del docente que no domina el tema
- **Responder la duda de contenido porque parece inofensiva.** «{p['dudas_no'][0]}» es
  literalmente la respuesta de una pregunta del parcial. La regla es la de arriba y se
  aplica sin excepción: si la respuesta es lo que se evalúa, no se responde.
- **No decidir el canal de entrega antes de empezar.** El enunciado remite a «el medio
  que el docente indique»; si no lo indicaste al minuto 0, lo vas a improvisar al minuto
  105 con la mitad del grupo escribiendo por privado.
- **No acusar recibo.** Es la fuente número uno de reclamos de un parcial virtual, y se
  resuelve escribiendo el nombre de cada uno en el chat a medida que llega el archivo.
- **Calificar con la solución mientras el bloque corre.** Además de que no da el tiempo,
  el archivo de solución abierto en pantalla compartida es un accidente que ya pasó.
- **Descontar por lo que no se enseñó.** Antes de restar puntos por un término que el
  estudiante no usó, busca la diapositiva donde se proyectó. Si no existe, el punto no se
  descuenta: los decks explican con vocabulario concreto donde el parcial usa el término
  de manual, y el estudiante responde con el que se le dio. Vale la respuesta que describe
  el mecanismo correcto aunque no lo nombre.
- **Tratar el día como clase.** Ni tema nuevo, ni avance del PI, ni «aprovechemos que
  terminaron temprano»: el bloque es solo evaluación.

## Preguntas frecuentes del grupo
**¿Puedo usar mis apuntes?** Lo que digas al minuto 0 y nada más. Por defecto: no.
Decídelo antes de abrir la sesión, porque cambiarlo a mitad del parcial invalida el de
quien ya respondió sin ellos.

**¿Se me cayó el internet, qué hago?** Que siga respondiendo el documento sin conexión y
te escriba por correo al reconectarse. El tiempo perdido por una caída comprobable no se
descuenta; el criterio se anuncia al minuto 0 para que nadie lo use como excusa después.

{faq_entra} La portada del enunciado lista las clases evaluadas con su fecha; fuera de esa
lista no hay nada.

**¿Cuánto vale cada sección?** Está en la portada, y es esto:
{secciones}
   El total es 100 puntos y la nota es puntos / 20. Este parcial pesa
   **{m['valor_corte']}**.

**¿Puedo responder con viñetas en vez de párrafos?** Sí, salvo que la pregunta pida una
forma concreta (una tabla, un número de filas). Lo que se califica es el contenido.

**¿Cuánto tiempo tengo?** El instrumento prevé **{m['tiempo']}** y el bloque son 120: hay
holgura, pero el envío cierra en el minuto 110 y eso no se mueve.

**¿Cuándo veo la nota?** En la siguiente sesión; la retroalimentación es escrita sobre el
mismo documento que entregaste.

## Notas
- No mezclar «Tema · Parcial»: hoy no se dicta nada.
{nota_prep}
- No publicar la solución en `Clases/`.
"""


def guion_md(c: dict) -> str:
    n = c["n"]
    if c["tipo"] == "parcial":
        return _guion_parcial_md(c)

    # El desarrollo completo del tema vive en arq_fundamentos.py (modulo de datos,
    # misma convencion que prog2_clases_data / seminario_clases_data). Antes estaba
    # inline y en 4-5 parrafos, insuficiente para la regla de oro del workspace:
    # el guion debe permitir dictar la clase sin consultar otra fuente.
    fund = FUNDAMENTOS.get(
        n, "Teoria al servicio del entregable PI de hoy. Ver diapositivas de la clase."
    )
    fund = _resolver_slides(fund, _slide_map(c), n)

    # Conceptos reales de esta clase (titulos de las slides de teoria), para que el
    # plan diga QUE cubrir en cada tramo en vez de "recorre las slides".
    conceptos = [x[0] for x in c.get("slides_extra", [])]
    minutos_por_concepto = max(5, 30 // max(1, len(conceptos)))

    # Modalidad de trabajo (individual por defecto, equipos si el docente los autoriza).
    # Las dos ramas de abajo usan las MISMAS variables para no volver a divergir.
    arranque_cita = mod(c, "arranque_cita")
    arranque_nota = mod(c, "arranque_nota")
    voluntario_word = mod(c, "voluntario_word")
    taller_modalidad_word = mod(c, "taller_modalidad_word")
    retro_word = mod(c, "retro_word")
    criterio_60s_note = mod(c, "criterio_60s_note")

    # Mapeo real del deck: el guion se lee CON la presentacion proyectada, asi que
    # cada tramo del plan dice a que diapositiva corresponde. Sale de _slide_map(),
    # la misma lista que arma build_pptx, no de numeros escritos a mano.
    mapa = _slide_map(c)
    sl_agenda = _slide_tag(mapa, "Agenda de hoy").strip()
    sl_obj = _slide_tag(mapa, "Objetivos de la clase").strip()
    sl_entreg = _slide_tag(mapa, "entregable de hoy").strip()
    sl_teoria = _slide_tag(mapa, *(conceptos or ["Objetivos"])).strip()
    sl_flujo = _slide_tag(mapa, FLUJO_SLIDE_TITULO).strip()
    sl_taller = _slide_tag(mapa, "paso a paso").strip()
    sl_cierre = f"[Slide {len(mapa)}]"
    mapa_md = "\n".join(f"{i}. {t}" for i, t in enumerate(mapa, 1))

    # Cada concepto de teoria con SU numero de diapositiva. Antes la lista eran titulos
    # sueltos: el docente leia «Cubre estos conceptos» y tenia que buscarlos a ojo en el
    # deck, y no habia forma de notar que una diapositiva de teoria no aparecia en
    # ningun tramo del plan. Si un titulo no esta en el mapa se rotula, no se inventa.
    conceptos_md = _lista_md(
        [f"**{t}** · {_slide_tag(mapa, t).strip() or '(sin diapositiva — avisar)'}"
         for t in conceptos]
    ) or "- Ver diapositivas de la clase."

    # La diapositiva de codigo/plantilla es el artefacto de la demo, asi que la demo la
    # cita. Sin esto, las clases sin diagrama publicaban «Demo en vivo · » — un separador
    # colgando y ninguna referencia — y la diapositiva de codigo no la nombraba ningun
    # tramo del plan.
    _cod = CODIGO_SLIDE.get(n)
    _cod_tag = _slide_tag(mapa, _cod[0]).strip() if _cod else ""
    sl_demo = sl_flujo or _cod_tag
    demo_tag = f" · {sl_demo}" if sl_demo else ""

    # Respaldo si falla la red o el proyector. Antes las 12 clases con taller mandaban al
    # docente a «Capturas/», pero solo 4 tienen imagen: en las otras 8 abria una carpeta
    # con un README dentro. Ahora se cita lo que existe de verdad en cada clase.
    if CAPTURAS_CLASE.get(n):
        respaldo_demo = f"proyecta las capturas de `Kit docente/Clase {n}/Capturas/`"
    elif _cod_tag:
        respaldo_demo = (f"proyecta la {_cod_tag}, que ya trae el resultado de la demo, y "
                         "recórrela rótulo por rótulo")
    else:
        respaldo_demo = (f"proyecta la solución docente de este kit (`Solucion Taller "
                         f"Clase {n} - CloudLite.md`), que trae el resultado esperado")

    # Bloque del flujo de diagramacion: solo si el taller de hoy pide un diagrama.
    if _tiene_diagrama(n):
        _dial = examlab_talleres._dialectos_del_taller(TALLERES_EXAMLAB[n])
        flujo_guion = (
            "\n**Cierra la demo dentro de ExamLab** " + sl_flujo + " — es el paso que el "
            "estudiante no adivina: pasa el boceto a codigo Mermaid con ayuda de una IA, "
            "pegalo en la pregunta de diagrama y muestralo renderizado.\n\n"
            + examlab_talleres.flujo_diagrama_md(
                _dial[0] if len(_dial) == 1 else "el tipo que pide el enunciado")
            + "\n"
        )
    else:
        flujo_guion = ""

    if not c.get("separar_notas_docente"):
        # Plantilla compartida por las clases que no separan «lo que el docente dice»
        # de «las notas para el docente». Cualquier cambio aqui se propaga a esas
        # clases en su proximo build; las frases de modalidad de trabajo estan
        # parametrizadas arriba, no las vuelvas a escribir a mano.
        plan_blocks = f"""### 0–10 · Encuadre PI · {sl_agenda}{sl_obj}{sl_entreg}
Di casi literal: «Hoy avanzamos el PI CloudLite App en: **{c['pi_hoy']}**.
Entregable concreto: {c['entregable']}.
Teoría breve y luego taller; no es un lab suelto.»
Pasa la diapositiva de agenda y la de objetivos. Abre el enunciado PI si alguien aún no lo tiene.
Pregunta de arranque (1 min): «{arranque_cita}» — sirve para detectar estudiantes rezagados antes de avanzar.

### 10–40 · Teoría Core (al servicio del taller) · desde {sl_teoria}
Cubre estos conceptos, en este orden, ~{minutos_por_concepto} min cada uno, con su diapositiva:
{conceptos_md}

**Ninguna se salta**: cada una de esas diapositivas es el mecanismo con que se resuelve
al menos una pregunta de la actividad calificada de hoy.
El desarrollo completo de cada uno está arriba, en «Fundamento teórico para el docente»:
esa sección está escrita para que puedas dictarla sin consultar otra fuente.
Cada 8–10 min amarra al artefacto: «esto es lo que van a dejar hoy en su informe/diagrama/repo».
Pide un {voluntario_word} voluntario y usa SU dominio como ejemplo en vivo (no el de la demo).

### 40–55 · Demo en vivo{demo_tag}
Herramienta del día: **{c['herramienta']}**.
{_demo_md(n)}
Narra los clics en voz alta. Si falla la red, {respaldo_demo}.
Cierra la demo con: «copien la estructura, no el dominio de mi ejemplo.»
{flujo_guion}{_capturas_md(n)}

### 55–100 · Taller guiado PI ({taller_modalidad_word}) · {sl_taller}
Proyecta la lista de pasos del taller del estudiante (está en la sección «Actividad / taller» de este guion).
Circula por mesas/Meet con la lista de errores frecuentes de abajo en la mano: son los que vas a ver hoy.
A los 80 min anuncia: «faltan 20 min. Falta evidencia: PNG/YAML/enlace. Empiecen a subir borrador.»

### 100–115 · Comprobación y evidencias
Haz 3–4 de las preguntas de comprobación oral de abajo, a personas distintas y al azar
(no al que levanta la mano). Es el mecanismo para verificar la regla de los 60 segundos.
Aplica el quiz corto de `Kit docente/Clase {n}/Quiz Clase {n} - {c['slug']}.docx`
(la clave va en archivo aparte y **no se proyecta**).
Mientras responden, verifica que el entregable esté realmente subido.
Retroalimenta 2–3 {retro_word} en voz alta, nombrando el error y la corrección concreta.

### 115–120 · Cierre · {sl_cierre}
Di: «Queda avanzado: {c['pi_hoy']}.
Criterio de éxito: {criterio_60s_note}.
Entrega domingo 23:59 en ExamLab. Siguiente hito del PI según el plan.»
"""
    else:
        # Rama exclusiva para clases con actividad individual (hoy: Clase 1), donde se
        # separa explicitamente lo que el docente DICE (bloques "> ...") de las
        # instrucciones PARA el docente (bloques "**[Nota docente]:**").
        plan_blocks = f"""### 0–10 · Encuadre PI · {sl_agenda}{sl_obj}{sl_entreg}
Di casi literal:
> "Hoy avanzamos el PI CloudLite App en: {c['pi_hoy']}. Entregable concreto: {c['entregable']}. Teoría breve y luego taller; no es un lab suelto."

**[Nota docente]:** pasa la diapositiva de agenda y la de objetivos. Abre el enunciado PI si alguien aún no lo tiene.

**[Nota docente]:** {arranque_nota}
> "{arranque_cita}"

### 10–40 · Teoría Core (al servicio del taller) · desde {sl_teoria}
Cubre estos conceptos, en este orden, ~{minutos_por_concepto} min cada uno, con su diapositiva:
{conceptos_md}

**[Nota docente]: ninguna se salta** — cada una de esas diapositivas es el mecanismo con que se
resuelve al menos una pregunta de la actividad calificada de hoy.
El desarrollo completo de cada uno está arriba, en «Fundamento teórico para el docente», ya dividido
por diapositiva: esa sección está escrita para que puedas dictarla sin consultar otra fuente.

**[Nota docente]:** cada 8–10 min amarra al artefacto («esto es lo que van a dejar hoy en su informe/diagrama/repo»)
y pide un {voluntario_word} voluntario para usar SU dominio como ejemplo en vivo (no el de la demo).

### 40–55 · Demo en vivo{demo_tag}
Herramienta del día: **{c['herramienta']}**.
{_demo_md(n)}

**[Nota docente]:** narra los clics en voz alta. Si falla la red, {respaldo_demo}.
Cierra la demo diciendo:
> "Copien la estructura, no el dominio de mi ejemplo."
{flujo_guion}{_capturas_md(n)}

### 55–100 · Taller guiado PI ({taller_modalidad_word}) · {sl_taller}
**[Nota docente]:** proyecta la lista de pasos del taller del estudiante (está en la sección «Actividad / taller»
de este guion). Circula por mesas/Meet con la lista de errores frecuentes de abajo en la mano: son los que vas
a ver hoy. A los 80 min anuncia:
> "Faltan 20 min. Falta evidencia: PNG/YAML/enlace. Empiecen a subir borrador."

### 100–115 · Comprobación y evidencias
**[Nota docente]:** haz 3–4 de las preguntas de comprobación oral de abajo, a personas distintas y al azar
(no al que levanta la mano). Es el mecanismo para verificar la regla de los 60 segundos.
Aplica el quiz corto de `Kit docente/Clase {n}/Quiz Clase {n} - {c['slug']}.docx`
(la clave va en archivo aparte y **no se proyecta**).
Mientras responden, verifica que el entregable esté realmente subido.
Retroalimenta 2–3 {retro_word} en voz alta, nombrando el error y la corrección concreta.

### 115–120 · Cierre · {sl_cierre}
Di:
> "Queda avanzado: {c['pi_hoy']}. Criterio de éxito: {criterio_60s_note}. Entrega domingo 23:59 en ExamLab. Siguiente hito del PI según el plan."
"""

    if c["tipo"] == "autonoma":
        plan_blocks = f"""### Modalidad autónoma (festivo)
Esta clase cae en festivo: no hay encuentro síncrono obligatorio. El estudiante trabaja solo,
con `Presentacion.pptx` + el taller de la carpeta `Clases/`. Por eso el material publicado
tiene que ser **autosuficiente**: lo que no quede escrito, nadie lo va a explicar en vivo.

### Qué publicar (antes del día de la clase)
1. En ExamLab: las diapositivas, el taller y el recordatorio del hito del PI.
2. La sección «Fundamento teórico para el docente» de este guion, adaptada como **lectura guía**
   del estudiante — es el reemplazo de la explicación en vivo, no un anexo opcional.
3. La **salida esperada** del ejercicio (ver la demo de abajo), para que el estudiante autónomo
   pueda comparar y saber si le quedó bien sin preguntarte.
4. Mensaje sugerido: «Clase {n} autónoma (festivo). Hoy avanzamos el PI en: {c['pi_hoy']}.
   Entregable: {c['entregable']}. Fecha límite: domingo 23:59. Dudas por foro/correo institucional.»

### Cómo debería repartir su tiempo el estudiante (120 min equivalentes)
- **0–15** Leer el encuadre y el objetivo del día; ubicar en qué quedó su CloudLite.
- **15–45** Leer la teoría (lectura guía) y tomar notas directamente en el informe del PI.
- **45–60** Revisar la salida esperada del ejercicio resuelto.
- **60–105** Desarrollar el taller sobre su propio CloudLite.
- **105–120** Empaquetar la evidencia y subirla a ExamLab.

### La demo, en versión asíncrona
{_demo_md(n)}
Publica esto como pasos escritos o como un video corto (3–5 min) grabado con estos mismos pasos.
Sin uno de los dos, el estudiante autónomo no tiene con qué comparar su resultado.
{_capturas_md(n)}

### Seguimiento (lo que sí es tu trabajo esa semana)
1. Revisa las entregas del domingo 23:59 con la lista de errores frecuentes de abajo:
   en modalidad autónoma esos errores aparecen más, porque nadie los corrigió en el momento.
2. Deja feedback breve orientado a la rúbrica del PI, nombrando el error y la corrección.
3. En la siguiente clase regular, dedica los primeros 10 min a los 2 errores más repetidos.
   Es el sustituto de la retroalimentación en vivo que esta clase no tuvo.

### Si ofreces office hours voluntario (opcional, 20–30 min)
Resuelve bloqueos concretos de diagrama/ADR/lab. Usa las preguntas de comprobación de abajo
para detectar quién entendió y quién solo copió la plantilla. No adelantes contenido de Parcial.
"""

    if c["tipo"] == "sustentacion":
        plan_blocks = f"""### Modalidad de la sesión: sustentaciones EN VIVO
Este bloque de 120 min se dedica íntegramente a las sustentaciones del Proyecto Integrador,
con encuentro síncrono. **No es clase autónoma y no es parcial.** No autorices reemplazar la
defensa por un video grabado: la sustentación es el único instrumento con el que verificas
autoría de los otros puntos del PI, y el Q&A en vivo no se puede sustituir por un documento.
El día cae en festivo de calendario, pero la sesión está destinada por decisión docente a
sustentar: anúncialo por escrito una semana antes para que nadie asuma que no hay clase.

### Antes de la sesión (semana previa)
1. Publica el orden y la duración exacta del turno: **6 min de pitch + 2–4 min de Q&A**.
   Con 12 sustentaciones eso es ~110 min; si el grupo es más grande, baja a 5 + 2 y avísalo
   antes, nunca el mismo día.
2. Exige el paquete subido a ExamLab (módulo Proyectos) **antes** del bloque: quien llega a
   subir archivos consume el tiempo de otro. Verifica tú mismo que los enlaces abren.
3. Ten a mano la rúbrica impresa por estudiante y la lista de preguntas de comprobación de
   abajo, para no improvisar el Q&A ni preguntar lo mismo a todos.
4. Si en la clase anterior no alcanzaron a ensayar, modela tú el formato antes de abrir turnos
   (no dentro de este bloque: no hay tiempo para eso y una sustentación menos):

{_demo_md(n)}

### 0–10 · Encuadre y orden de turnos
Di casi literal:
> "Hoy sustentamos. 6 minutos de pitch y hasta 4 de preguntas. Yo corto a los 6 minutos: si no
> llegaron a seguridad, costos y escala, esa parte no se califica. El orden lo sorteo ahora."

**[Nota docente]:** sortea el orden delante del grupo (evita el reclamo de «me tocó primero»),
proyecta el cronómetro y pide que el resto escuche: cerramos el curso entre todos.

### 10–110 · Sustentaciones (turnos consecutivos)
Por cada turno, en este orden:
1. **6 min de pitch.** No interrumpas ni siquiera para corregir un error: se anota y se pregunta
   después. Corta seco a los 6 min.
2. **2–4 min de Q&A.** Haz siempre una pregunta de verificación («muéstrame el .yml del
   workflow»), una de profundización («¿por qué la base de datos no está en el mismo contenedor
   que la API?») y, si queda tiempo, una hipotética («si el tráfico se multiplica por diez el
   lunes, ¿qué pieza se rompe primero?»). En equipo autorizado, dirige cada pregunta a un
   integrante distinto y **no** dejes que responda siempre el mismo.
3. **Cierra el turno con la nota puesta**, no al final del día: la rúbrica se llena en caliente
   mientras recuerdas la respuesta exacta.

**[Nota docente]:** frase de rescate cuando el estudiante se bloquea, para no perder el turno:
> "Déjame la respuesta pendiente y sigue con el siguiente bloque; vuelvo a preguntar al final."

### 110–120 · Cierre del curso
Di casi literal:
> "Lo que entregaron —diagramas, Dockerfile, workflow, informe— es un portafolio real: no lo
> borren al terminar el semestre. Arquitectura no es una lista de logos de proveedores, es un
> conjunto de decisiones documentadas con sus consecuencias."

Recuerda los pesos sin abrir discusión de notas: el PI vale **20% del Corte 3** y el Parcial 3
ya se aplicó en su propia sesión; el proyecto no reemplaza ni compensa el parcial.

### Si un estudiante no se presenta o falla la conexión
Deja constancia escrita en el momento (hora, motivo) y reprograma dentro de la misma semana con
Meet, sustentando igualmente en vivo. Aceptar un video grabado «por esta vez» convierte la
excepción en la regla del semestre siguiente y elimina el Q&A, que es la mitad de lo que evalúas.
"""

    pasos = "\n".join(f"{i+1}. {p}" for i, p in enumerate(_pasos(c)))
    return f"""# Guion docente — Clase {n}: {c['tema']}

## Información de la clase
- Asignatura: Arquitectura de Sistemas Computacionales (FI303380)
- Duración del bloque: **120 min**
- Tipo: {TIPO_LABEL[c["tipo"]]}
- Enfoque: **Proyecto Integrador CloudLite App** (parte práctica)
- Sin fechas de periodo · sin bio · sin mapa completo del curso

## Objetivos de la clase
{chr(10).join("- " + o.replace("**", "") for o in c["objetivos"])}

## Hoy avanzamos el PI en…
**{c["pi_hoy"]}**

**Entregable concreto:** {c["entregable"]}

**Herramienta:** {c["herramienta"]}

## Fundamento teórico para el docente
{fund}

## Referencias a diapositivas
Numeración real del deck `Clases/Clase {n} - {c['slug']}/Presentacion.pptx` (solo tema
de esta clase). Las etiquetas [Slide N] del plan y del fundamento apuntan aquí.

{mapa_md}

## Plan de clase minuto a minuto (120 min)

{plan_blocks}

## Actividad / taller (detalle)
{pasos}

### Criterio de éxito
- Artefacto integrado al paquete PI (no archivo huérfano).
- Evidencia adjunta.
- {mod(c, "criterio_oral_note")}

## Errores frecuentes del estudiante (y cómo corregirlos en el momento)
{_lista_md(ERRORES_ARQ.get(n, [])) or "- Sin errores catalogados para esta clase."}

## Preguntas de comprobación oral (no son del quiz)
{"Úsalas como Q&A al cerrar cada turno de sustentación, variándolas entre estudiantes." if c["tipo"] == "sustentacion" else "Úsalas en el tramo 100–115, a personas distintas y al azar."}
{_lista_md(PREGUNTAS_ARQ.get(n, []), bullet="1. ") or "- Sin preguntas catalogadas para esta clase."}

## Solución del taller (privada)
`Kit docente/Clase {n}/Solucion Taller Clase {n} - CloudLite.docx` — es la referencia con la que
comparas lo que entregan {mod(c, "entregan_word")}. **No proyectarla completa** antes de que trabajen.

## Quiz
`Kit docente/Clase {n}/Quiz Clase {n} - {c['slug']}.docx` (versión estudiante, sin respuestas)
y `Kit docente/Clase {n}/Quiz Clase {n} - CLAVE DOCENTE.docx` (clave, privada).

## Capturas sugeridas
- 📸 La herramienta del día en uso con el artefacto CloudLite [[captura: demo-clase{n:02d}.png | receta: 1) Abre {c['herramienta']} y repite la demo de este guion.  2) Captura solo la ventana útil, no el escritorio completo.  3) Recorta a ~1200 px de ancho.  4) Guárdala como Kit docente/Clase {n}/Capturas/demo-clase{n:02d}.png.  5) Vuelve a generar el guion y la imagen queda embebida aquí sola. Detalle en Capturas/README.txt.]]
- 📸 Evidencia del entregable de un estudiante (diagrama / YAML / lab) [[captura: evidencia-clase{n:02d}.png | receta: 1) Con permiso del estudiante, captura su artefacto de hoy.  2) Recorta nombre y correo antes de guardar.  3) Guárdala como Kit docente/Clase {n}/Capturas/evidencia-clase{n:02d}.png.  4) Es para tu registro del corte; no se proyecta en clase.]]

## Notas operativas
- Plataforma de entrega: ExamLab (https://uniaj.examlab.workers.dev/). No es la plataforma oficial de la UNIAJC; la universidad no tiene campus virtual propio.{chr(10) + "- " + c["entrega_oficial_nota"] if c.get("entrega_oficial_nota") else ""}
- Prohibido pedir cloud con tarjeta: todo el curso corre con free tier o en el navegador.
- Día de parcial = solo evaluación (no aplica a esta clase).
"""


def _escribir_readme_capturas(c: dict, cap: Path) -> None:
    """README de la carpeta Capturas/ con el paso a paso de cada imagen.

    Antes era una linea generica («preferir PNG reales») que no decia que abrir ni
    con que nombre guardar, asi que las carpetas quedaban vacias y el .docx salia
    con la caja de captura en blanco el dia de la clase. Se reescribe en cada build
    para que refleje la herramienta y la demo actuales de la clase.
    """
    n = c["n"]
    demo = DEMO_ARQ.get(n)
    L = [f"Capturas de la Clase {n} — Arquitectura de Sistemas Computacionales",
         "=" * 62, "",
         "El guion embebe automaticamente cualquier PNG que exista aqui con el nombre",
         "esperado. Mientras no exista, el .docx imprime la receta en su lugar.", ""]
    ya = CAPTURAS_CLASE.get(n) or []
    if ya:
        L.append("Ya generadas por `python config/slides/mockups.py` (no hay que tomarlas):")
        for rotulo, fn in ya:
            L.append(f"  - {fn} — {rotulo}")
        L.append("")
    L += [f"Pendiente: demo-clase{n:02d}.png — la herramienta del dia en uso",
          f"  1. Abrir {c['herramienta']}."]
    if demo:
        L.append(f"  2. Repetir la demo: {demo[0]}.")
        for i, paso in enumerate(demo[1], 1):
            L.append(f"     {i}. {paso}")
    else:
        L.append("  2. Reproducir el ejercicio del bloque sobre el dominio CloudLite de ejemplo.")
    L += ["  3. Capturar solo la ventana util, no el escritorio completo.",
          "  4. Recortar a ~1200 px de ancho.",
          f"  5. Guardar aqui como demo-clase{n:02d}.png.",
          "",
          "Pendiente: evidencia del entregable (diagrama / YAML / lab)",
          "  - Con permiso del estudiante, capturar su artefacto de hoy.",
          "  - Recortar nombre y correo antes de guardar. No se proyecta en clase.",
          "",
          "Despues de agregar una imagen, regenerar el guion:",
          f"  SOLO_CLASES={n} python config/slides/build_uniajc_arq_clases_batch.py",
          ""]
    cap.mkdir(parents=True, exist_ok=True)
    (cap / "README.txt").write_text("\n".join(L), encoding="utf-8")


def build_guion(c: dict) -> Path:
    n = c["n"]
    kit = CURSO / "Kit docente" / f"Clase {n}"
    cap = kit / "Capturas"
    kit.mkdir(parents=True, exist_ok=True)
    cap.mkdir(parents=True, exist_ok=True)
    # placeholder readme capturas
    _escribir_readme_capturas(c, cap)
    path = kit / f"Guion Docente Clase {n} - {c['slug']}.md"
    path.write_text(guion_md(c), encoding="utf-8")
    print("OK guion md ->", path)
    return path


def build_solucion(c: dict):
    """Solucion del taller para el Kit docente, pregunta por pregunta.

    El render vive en `solucion_taller` porque BD II usa el mismo formato; aqui
    solo se aportan los datos y el contexto del curso. Las clases sin datos de
    solucion no se tocan: conservan su archivo actual.
    """
    n = c["n"]
    sol = soluciones.SOLUCION.get(n)
    if not sol:
        return None

    md = solucion_taller.render_md(
        n, sol,
        contexto={
            "alineacion": [
                ("Taller del estudiante", f"`Clases/Clase {n} - {c['slug']}/`"),
                ("Configuracion en la plataforma",
                 f"`Kit docente/Clase {n}/Taller en ExamLab - Clase {n} (configuracion).md`"),
                ("Hito del PI", c.get("pi_hoy", "—")),
                ("Entregable", c.get("entregable", "—")),
            ],
            "politica_extra": ("Gratis + navegador; sin cuentas cloud de pago ni tarjeta."),
        },
        opciones=soluciones.opciones,
        mermaid_referencia=soluciones.mermaid_referencia,
        dominio_referencia=soluciones.DOMINIO_REFERENCIA,
    )

    kit = CURSO / "Kit docente" / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    out = kit / f"Solucion Taller Clase {n} - CloudLite.md"
    out.write_text(md, encoding="utf-8")
    print("OK solucion md ->", out)
    return out


def build_parcial_kit_note(c: dict) -> Path:
    """Nota de una pagina para el dia de parcial: el resumen que se lee de pie.

    Los datos salen de `PARCIALES_ARQ` y no de un `mapping` local, porque el guion de la
    misma clase los necesita iguales y tenerlos dos veces ya habia producido una
    divergencia: la nota decia «p. ej. Clase 12 para P3» dentro del archivo de la
    Clase 5, donde el dato correcto es que en el Corte 1 no hay preparacion de pitch.
    """
    n = c["n"]
    kit = CURSO / "Kit docente" / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    path = kit / f"NOTA Docente - Clase {n} Parcial.md"
    p = PARCIALES_ARQ[n]
    m = _parcial_meta(p["corte"])
    prep = p.get("prep_pi")
    linea_pi = (f"- El pitch del PI se preparó en la **Clase {prep}**; hoy no se avanza "
                "en el proyecto." if prep else
                "- En este corte no hubo preparación de pitch; hoy no se avanza en el "
                "proyecto y el PI se retoma en la siguiente clase regular.")
    # Las clases evaluadas se sacan de la portada del instrumento («Clase 7 · 05/10 ·
    # …»), no de una lista escrita aqui: si el parcial cambia de cobertura, la nota
    # cambia con el o queda mandando al docente a estudiar otro corte.
    clases_ev = [t.split(" · ")[0].replace("Clase ", "") for t in m["temas"]]
    lista = (", ".join(clases_ev[:-1]) + " y " + clases_ev[-1]
             if len(clases_ev) > 1 else clases_ev[0])
    return _escribir_nota(path, f"""# Clase {n} — Solo Parcial {p['corte']}

- Bloque 120 min · **virtual síncrona por Google Meet** · **sin taller PI**.
- Enunciado: `Parciales/{m['archivo']}.docx` · la solución es el mismo nombre con
  `- SOLUCION` y **no se publica en `Clases/`**.
- Evalúa **solo** las Clases {lista} · 100 puntos · nota = puntos / 20 ·
  peso **{m['valor_corte']}**.
- **Decide el canal de entrega antes del minuto 0** y anúncialo: el enunciado dice «por el
  medio que el docente indique al abrir la sesión», así que si no lo dices, no existe.
- Durante el parcial se responden dudas de **forma**, nunca de **contenido**.
- Acusa recibo por el chat, uno por uno, a medida que llegan los archivos.
{linea_pi}

El detalle minuto a minuto, los errores típicos y las preguntas frecuentes están en
`Guion Docente Clase {n} - {c['slug']}.md`.
""")


def _escribir_nota(path: Path, texto: str) -> Path:
    path.write_text(texto, encoding="utf-8")
    print("OK nota parcial ->", path)
    return path


def convert_guiones(paths: list[Path]) -> None:
    conv = SLIDES_DIR / "guion_md_a_docx.py"
    for md in paths:
        if not md.exists():
            continue
        os.system(f'python "{conv}" "{md}"')


def build_examlab_guia(c):
    """Guia para armar el taller de esta clase dentro de ExamLab.

    Va en el Kit docente porque la plataforma no importa preguntas desde archivo:
    el docente las crea en la UI y necesita el texto exacto de cada campo.
    """
    taller = TALLERES_EXAMLAB.get(c["n"])
    if not taller:
        return None
    d = CURSO / "Kit docente" / f"Clase {c['n']}"
    d.mkdir(parents=True, exist_ok=True)
    md = examlab_talleres.guia_docente_md(
        c["n"], taller, "Arquitectura de Sistemas Computacionales (FI303380)",
        hito=c.get("pi_hoy"), entregable=c.get("entregable"),
    )
    out = d / f"Taller en ExamLab - Clase {c['n']} (configuracion).md"
    out.write_text(md, encoding="utf-8")
    print("OK examlab ->", out)
    return out


def build_all(solo_clases=None):
    """Regenera el batch completo, o solo un subconjunto de clases si se pasa
    ``solo_clases`` (iterable de numeros de clase) o se define la variable de
    entorno SOLO_CLASES="1,3,7" (coma-separada). Las clases no incluidas no se
    tocan, para poder aislar un cambio (p.ej. reemplazo de una herramienta de
    lab) a las clases afectadas sin regenerar el resto del curso.
    """
    if solo_clases is None:
        env_val = os.environ.get("SOLO_CLASES")
        if env_val:
            solo_clases = {int(x.strip()) for x in env_val.split(",") if x.strip()}
    else:
        solo_clases = set(solo_clases)
    guiones = []
    for c in CLASSES:
        n = c["n"]
        if solo_clases is not None and n not in solo_clases:
            continue
        print(f"\n=== Clase {n} ({c['tipo']}) ===")
        if c["tipo"] == "parcial":
            build_parcial_kit_note(c)
            # pptx mínimo de indicaciones (opcional en Clases para no confundir; se genera carpeta clara)
            build_pptx(c)
            g = build_guion(c)
            guiones.append(g)
            continue
        build_pptx(c)
        build_taller_docx(c)
        build_quiz_docx(c)
        build_examlab_guia(c)
        sol = build_solucion(c)
        if sol:
            guiones.append(sol)   # se convierte a .docx con el mismo conversor
        g = build_guion(c)
        guiones.append(g)
    convert_guiones(guiones)
    print("\nDONE batch Arquitectura PI-first")


if __name__ == "__main__":
    build_all()
