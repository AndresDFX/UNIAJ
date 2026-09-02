# -*- coding: utf-8 -*-
"""Genera material completo BD II 2026-2 centrado en PI VetCare DB.

Salidas:
  Clases/Clase N - <slug>/Presentacion.pptx + Taller PI ....docx
  Kit docente/Clase N/Guion....md|.docx + Quiz + Codigo + Capturas/
  Kit docente/Clase N/Guia aplicacion parcial (dias 5/9/14)
"""
from __future__ import annotations
import os, sys, re, subprocess, unicodedata
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
SLIDES = Path(__file__).resolve().parent
sys.path.insert(0, str(SLIDES))

from uniajc_slides_engine import (
    new_prs, content_slide, table_content, box_note_slide, closing_slide,
    class_cover, herramientas_slide, steps_visual_slide, checklist_slide,
    block_timeline_slide, diagram_boxes_slide, pseudo_code_slide,
    before_after_slide, AMARILLO, NAVY, CIAN,
    RED as PPTX_RED,
)
from uniajc_quiz_helpers import clave_text, pptx_chunks, q_abierta, q_om, q_vf, student_lines
import calendario_2026_2 as cal
from vetcare_contexto import CLIENTE, INTERESADOS, NOMENCLATURA, PROBLEMAS
from bd2_taller_data import HERRAMIENTAS_DIA, TALLER_BLOQUE, SOLUCION
from bd2_fundamentos import FUNDAMENTOS
import bd2_solucion_data as soluciones_bd2
import solucion_taller
from bd2_examlab_data import EXAMLAB as TALLERES_EXAMLAB
import examlab_talleres
from docx import Document
from docx.shared import Pt as DocPt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CURSO = ROOT / "Bases de Datos II"
CLASES_DIR = CURSO / "Clases"
KIT_DIR = CURSO / "Kit docente"
PARCIALES = CURSO / "Parciales"
AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN_D = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"
# Fuente de ancho fijo para las plantillas en blanco del entregable. Con Calibri la
# matriz rol x objeto quedaba desalineada (las columnas de una tabla de pipes solo
# cuadran en monoespaciada) y el estudiante la recibia rota justo en el artefacto
# que se califica por su estructura.
MONO = "Consolas"

# ---------------------------------------------------------------------------
# Tipo de bloque: SE DERIVA DEL CALENDARIO, no se escribe a mano por clase
# ---------------------------------------------------------------------------
# El material se genero con el calendario viejo, donde el festivo del 17/08 caia
# en la Clase 2 y el del 16/11 cerraba el curso. Al acortarse el semestre
# (24/08-22/11, 13 sesiones para 15 temas) el 17/08 quedo FUERA del rango y la
# Sesion 13 (16/11) se destino a las sustentaciones del PI en vivo; como el tipo
# estaba escrito a mano, el material siguio anunciando «clase autonoma» para dos
# clases que si tienen encuentro sincrono. Leerlo del calendario en cada build
# impide que una regeneracion futura reintroduzca el defecto.
CAL_KEY = "bases_datos_ii"


def tipo_de_clase(n_clase):
    """'parcial' | 'autonoma' | 'sustentacion' | 'regular' para una Clase de material.

    'presencial' y 'virtual' colapsan a 'regular': las dos son sincronas y no
    cambian la estructura del bloque (lo que cambia es el aula vs Meet).
    """
    s = cal.sesion_de_clase(CAL_KEY, n_clase)
    if s is None:
        raise SystemExit(
            "Clase %d no aparece en 'clases_material' de ninguna sesion de '%s' en %s. "
            "Corrija el calendario, no este build." % (n_clase, CAL_KEY, cal.JSON_PATH)
        )
    if s.get("parcial"):
        return "parcial"
    t = (s.get("tipo") or "").lower()
    return t if t in ("autonoma", "sustentacion") else "regular"


TIPO_LABEL = {
    "regular": "REGULAR (sincrona)",
    "autonoma": "AUTONOMA (festivo, sin encuentro sincrono)",
    "sustentacion": "SUSTENTACION DEL PI **EN VIVO** (sincrona)",
    "parcial": "PARCIAL (solo evaluacion)",
}

# Sin clave "tipo": se inyecta desde el calendario justo despues de esta lista.
CLASES = [
  dict(n=1, slug="Revision BD I y arranque VetCare",
    titulo="Revision BD I · Arranque VetCare DB",
    subtitulo="Diagnostico · dominio PI · primer modelo",
    herramienta="draw.io + DB Fiddle",
    hito_pi="Arranque PI: dominio, alcance y borrador ER de VetCare DB",
    entregable="Ficha del PI (plantilla) + ER en Mermaid renderizado en ExamLab (PNG para tu carpeta) + 3 reglas Condicion -> Accion",
    teoria=["Nivel conceptual = entidades y relaciones; nivel fisico = tablas con tipos y longitudes. El ER de hoy es el conceptual (Dueno posee Mascota); el CREATE TABLE de la demo es el mismo modelo en fisico (dueno.telefono VARCHAR(30)). Una tabla = conjunto de entidades del mismo tipo; cada fila es una instancia, cada columna un atributo. La clave primaria (PK) identifica sin ambiguedad cada fila: nunca se repite, nunca es nula.",
            "Clave foranea (FK): columna que apunta a la PK de otra tabla y materializa una relacion (1-N o N-N via tabla intermedia). Garantiza integridad referencial: la BD rechaza una Cita con id_mascota que no existe.",
            "Normalizacion 1FN-3FN en una frase cada una: 1FN = nada de listas dentro de una celda (una fila = una mascota, no varias); 2FN = ningun atributo depende solo de una parte de una PK compuesta; 3FN = ningun atributo depende de otro atributo que no sea la PK. Sub-normalizar genera anomalias de insercion/actualizacion/borrado (ej.: cambiar el telefono de un dueno en 5 filas distintas); sobre-normalizar multiplica JOINs sin necesidad real.",
            "Error de docente que no domina el tema: confundir PK con 'el primer campo de la tabla', o asumir que normalizar siempre mejora el rendimiento (en lectura intensiva a veces se denormaliza a proposito, y eso se vera en Clase 6-7).",
            "Dominio VetCare y sus relaciones: Dueno 1-N Mascota, Mascota 1-N Cita, Veterinario 1-N Cita, Consulta 1-1 Cita (una consulta documenta una cita atendida), Factura 1-N DetalleFactura N-1 Insumo.",
            "Reglas de negocio del PI que ya anticipan clases futuras: mascota inactiva no puede tener cita nueva (se validara con un procedimiento en Clase 3), stock de insumo nunca queda negativo (transacciones, Clase 8), cambios sensibles quedan auditados (triggers, Clase 4)."],
    demo="Boceto ER en draw.io (Dueno-Mascota-Cita) + CREATE TABLE minimo en DB Fiddle, y cierre pasando el boceto a Mermaid con IA para pegarlo renderizado en ExamLab.",
    taller=["Registrar el proyecto con el nombre exacto VetCare - [Apellido] (trabajo individual por defecto; equipo de 2-3 solo si el docente lo autoriza).",
            "Llenar la plantilla de la ficha del PI: alcance SI / alcance NO y 3 reglas de negocio propias en formato Condicion -> Accion.",
            "Dibujar el ER borrador en Excalidraw o draw.io, pasarlo a Mermaid (erDiagram) con ayuda de una IA y pegarlo renderizado en ExamLab.",
            "Exportar tambien el PNG del ER a la carpeta del PI y verificar que los nombres coincidan con el DDL (minusculas, singular, id_<entidad>)."],
    quiz=True, sql="01_arranque_vetcare.sql"),
  dict(n=2, slug="Administracion de bases de datos",
    titulo="Administracion de BD · Roles VetCare",
    subtitulo="Privilegios y usuarios del PI",
    herramienta="ExamLab (PostgreSQL) + Google Docs",
    hito_pi="Plan de roles/privilegios de VetCare",
    entregable="Documento Roles_VetCare + script GRANT/REVOKE ejecutado en ExamLab",
    teoria=["Administracion de BD = gestionar QUIEN puede hacer QUE sobre CADA objeto. Tres piezas: usuario (identidad que se conecta), rol (paquete de privilegios con nombre, ej. recepcion), privilegio (permiso atomico: SELECT, INSERT, UPDATE, DELETE, EXECUTE sobre un objeto concreto). En PostgreSQL usuario y rol son la misma cosa: un usuario es un rol con LOGIN.",
            "Principio de minimo privilegio: cada rol recibe solo lo que necesita para su funcion, ni un privilegio mas. No es paranoia, es reduccion de superficie de dano: si roban la sesion de un recepcionista, no debe poder borrar el historial clinico ni ver nomina.",
            "Separacion de funciones (segregation of duties): quien disena/modifica el esquema (DDL: CREATE/ALTER/DROP) no deberia ser la misma cuenta que opera datos del dia a dia (DML: INSERT/UPDATE/DELETE), y quien audita solo deberia leer (SELECT), nunca escribir.",
            "GRANT otorga un privilegio a un rol o usuario; REVOKE lo retira. Un rol se puede asignar a varios usuarios (todos los recepcionistas heredan el rol recepcion) y modificar en un solo lugar en vez de uno por uno.",
            "Un GRANT no es todo-o-nada: se puede recortar la superficie con una vista (CREATE VIEW deja fuera filas y columnas) o con privilegios por columna (GRANT SELECT (id_dueno, nombre) ON dueno TO veterinario_rol). Asi el rol llega al dato que necesita sin ver el resto de la tabla.",
            "Un permiso no vale nada sin evidencia: information_schema.role_table_grants y information_schema.column_privileges son las dos consultas que prueban que la matriz quedo como se decidio. Sin ellas la matriz es una intencion, no un hecho verificable.",
            "Error de docente que no domina el tema: crear un unico usuario 'admin' que todos comparten (rompe la trazabilidad de auditoria) o dar ALL PRIVILEGES a todo el mundo 'para que no falle nada' — exactamente lo opuesto a minimo privilegio."],
    demo="Los 4 roles de VetCare con CREATE ROLE/GRANT/REVOKE en ExamLab, verificados con information_schema.role_table_grants.",
    taller=["Crear los 4 roles (admin_bd, recepcion, veterinario_rol, auditor) con GRANT/REVOKE que corran.",
            "Recortar la superficie: vista v_agenda_recepcion + privilegio por columna sobre dueno.",
            "Matriz rol x objeto x privilegio de los 10 objetos, justificando privilegio minimo.",
            "Redactar 1 pagina: politica de altas/bajas de usuarios, con la prueba negativa (SET ROLE) corrida y su mensaje de error."],
    quiz=True, sql="02_roles_vetcare.sql"),
  dict(n=3, slug="Procedimientos almacenados",
    titulo="Procedimientos almacenados · VetCare",
    subtitulo="Logica de negocio en la BD del PI",
    # El motor es el que califica: las 5 preguntas de ExamLab son PL/pgSQL que
    # corre en PGlite. Oracle Live SQL queda como contraste de sintaxis, no como
    # sitio de trabajo — igual que en la Clase 2.
    herramienta="ExamLab (PostgreSQL) + Google Docs",
    hito_pi=">=1 procedimiento de negocio (agendar cita / registrar consulta)",
    entregable="2 procedimientos en PL/pgSQL corriendo en ExamLab + bateria de pruebas con su tabla resultado_prueba + contrato del proc (6 bloques)",
    teoria=["Un procedimiento almacenado es logica de negocio guardada DENTRO de la base, y se llama con CALL. No es una consulta con nombre: recibe parametros tipados y ejecuta varias sentencias como una sola unidad logica, de modo que la regla vive UNA vez y toda la app la respeta.",
            "El molde de PL/pgSQL es fijo: CREATE PROCEDURE nombre(params) LANGUAGE plpgsql AS $proc$ ... $proc$;. Dentro van DECLARE (variables), BEGIN y END. Los delimitadores $proc$ (dollar-quoting) existen porque el cuerpo lleva punto y coma y el motor necesita saber donde termina. Nada de IS en vez de AS, ni VARCHAR2, ni NUMBER, ni RAISE_APPLICATION_ERROR, ni la barra / final: eso es Oracle y aqui no compila.",
            "Parametros: IN es el defecto y no se escribe; OUT e INOUT existen pero hoy no se usan para reportar errores. Los tipos son los de PostgreSQL: INT, NUMERIC, TEXT, TIMESTAMP, BOOLEAN.",
            "La validacion no devuelve un mensaje: aborta con RAISE EXCEPTION 'ERROR: ... %', variable;. El % se sustituye en orden por las variables que siguen a la coma. Al abortar, todo lo que el procedimiento hubiera hecho se deshace, asi que es imposible que quede una cita a medias. Con un mensaje en un parametro OUT el INSERT seguiria corriendo: la regla no se cumpliria.",
            "Un procedimiento sin prueba no esta terminado: la bateria son bloques DO que capturan el error. Cada caso va en su propio bloque DO $$ BEGIN ... EXCEPTION WHEN OTHERS THEN ... SQLERRM ... END $$;, y el resultado se escribe en una tabla resultado_prueba (caso, esperado, obtenido, paso). Un caso OK y tres casos error, mas el COUNT(*) que demuestra que la tabla cita paso de 10 a 11 filas.",
            "Procedimiento y funcion se diferencian hoy, no en la Clase 4: CALL sp_x(...) frente a SELECT fn_x(...). El procedimiento se ejecuta como una accion y puede manejar transacciones; la funcion retorna un valor y se invoca dentro de una expresion SQL. En PostgreSQL una funcion no puede hacer COMMIT ni ROLLBACK, y eso es lo que decide cual de los dos se usa.",
            "El contrato del proc es lo que consume la futura app: firma, precondiciones, postcondiciones y errores. Son 6 bloques: la firma exacta con tipos, un ejemplo de CALL, las precondiciones, las postcondiciones, la tabla de errores con su mensaje literal, y la decision de diseno que explica por que se aborta en vez de devolver un codigo.",
            "Error de docente que no domina el tema: escribir el proc sin validar nada (solo el INSERT) y llamarlo 'logica de negocio' — un proc sin reglas de validacion es solo una consulta con nombre. El segundo error es dictar el molde de Oracle porque es el que uno recuerda: en ExamLab ese codigo no compila, y el estudiante pierde los 35 puntos de la pregunta 1 por sintaxis, no por no entender el tema."],
    demo="sp_agendar_cita en PL/pgSQL dentro de ExamLab: las 3 validaciones con RAISE EXCEPTION y la bateria de bloques DO que las prueba.",
    taller=["Escribir sp_agendar_cita en PL/pgSQL y ejecutarlo en ExamLab (LANGUAGE plpgsql, dollar-quoting, sin sintaxis de Oracle).",
            "Incluir las 3 validaciones de negocio del PI, cada una con su RAISE EXCEPTION y su mensaje literal.",
            "Correr la bateria de pruebas con bloques DO: 1 caso OK + 3 casos error, escritos en resultado_prueba, mas el COUNT(*) de cita antes y despues.",
            "Escribir sp_registrar_consulta, comprobando con EXISTS antes de chocar contra la restriccion UNIQUE.",
            "Redactar el contrato del proc en sus 6 bloques (plantilla en este documento) y pegarlo en la pregunta 5."],
    quiz=True, sql="03_procs_vetcare.sql"),
  dict(n=4, slug="Funciones disparadores seguridad respaldo",
    titulo="Funciones · Triggers · Seguridad y respaldo",
    subtitulo="Integridad + RAA1 del PI VetCare",
    herramienta="ExamLab (PostgreSQL) + Google Docs",
    hito_pi=">=1 funcion + >=1 trigger + borrador plan de respaldo",
    entregable="fn_precio_consulta + 2 triggers corriendo en ExamLab + Plan_Backup_VetCare con sus 6 secciones (1 pag.)",
    teoria=["Funcion (Clase 3 vio procedimiento): retorna un valor y se usa DENTRO de una expresion SQL, ej. SELECT fn_precio_consulta(especie, urgencia) FROM mascota. Su molde es CREATE FUNCTION nombre(params) RETURNS tipo LANGUAGE plpgsql AS $fn$ ... $fn$;. Si no toca datos se marca IMMUTABLE, que le dice al motor que puede memorizar el resultado. Nada de RETURN NUMBER IS: eso es Oracle.",
            "Trigger (disparador): bloque de codigo que el motor ejecuta AUTOMATICAMENTE cuando ocurre un evento (BEFORE/AFTER INSERT, UPDATE o DELETE) sobre una tabla, sin que nadie lo llame explicitamente. Dos usos tipicos aqui: auditoria (guardar quien/cuando cancelo una cita) y validacion de invariantes (que el stock nunca quede negativo tras un UPDATE).",
            "En PostgreSQL un trigger son SIEMPRE dos objetos, no uno: la funcion y la asociacion. Primero CREATE FUNCTION fn_trg_x() RETURNS TRIGGER, que termina en RETURN NEW (o RETURN OLD si el evento es DELETE); despues CREATE TRIGGER trg_x AFTER UPDATE OF estado ON cita FOR EACH ROW EXECUTE FUNCTION fn_trg_x();. Dentro de la funcion las filas se leen como NEW y OLD, SIN los dos puntos: NEW.estado, no :NEW.estado. Escribir el cuerpo dentro del CREATE TRIGGER es la herencia de Oracle que mas cuesta puntos, porque no compila.",
            "BEFORE o AFTER no es un detalle de estilo: un trigger que VALIDA va BEFORE, porque tiene que abortar antes de que el dato quede escrito; un trigger que AUDITA va AFTER, porque registra un hecho ya consumado. Y la clausula WHEN (OLD.estado IS DISTINCT FROM NEW.estado) evita registrar los UPDATE que no cambiaron nada: es la diferencia entre auditar 2 filas y auditar 3.",
            "Riesgo real de los triggers: son invisibles en el codigo de la app (un desarrollador que solo mira el INSERT no ve que ademas se dispara una auditoria), y pueden encadenarse (un trigger que dispara otro trigger) generando efectos dificiles de rastrear. Se usan para pocas reglas criticas, no para toda la logica de negocio: una regla sobre una sola columna es un CHECK, una regla que compara filas es un trigger, y una regla de interfaz es de la app.",
            "Seguridad y respaldo van juntos: seguridad evita que datos se corrompan o se filtren; respaldo (backup) asume que igual algo saldra mal y prepara la recuperacion. Full backup (copia completa), incremental (solo lo que cambio desde el ultimo backup) y diferencial (todo lo que cambio desde el ultimo FULL) son las tres estrategias base. En PostgreSQL las herramientas son pg_dump (una base), pg_dumpall --globals-only (los roles del cluster, que pg_dump NO respalda) y pg_basebackup con archivado de WAL.",
            "RPO (Recovery Point Objective) = cuantos datos se puede permitir perder, medido en tiempo ('maximo 1 hora de citas perdidas'). RTO (Recovery Time Objective) = cuanto tiempo puede estar caida la BD antes de restaurar. Un backup diario sin probar el restore no cumple ningun RPO/RTO real: un plan de respaldo sin prueba de restauracion es solo una promesa.",
            "Error de docente que no domina el tema: presentar el backup como 'copiar el archivo de vez en cuando' sin frecuencia, retencion (cuantas copias se guardan) ni prueba de restore — eso es lo que el taller de esta clase pide explicitamente que el estudiante defina. El segundo error es dictar el trigger como en Oracle, con el cuerpo dentro del CREATE TRIGGER y :NEW/:OLD: la rubrica lo penaliza expresamente, asi que el docente estaria proyectando el codigo por el que va a descontar."],
    demo="fn_precio_consulta + fn_trg_audit_cita con su CREATE TRIGGER ... EXECUTE FUNCTION, en ExamLab, y el esqueleto del plan de respaldo.",
    taller=["Escribir fn_precio_consulta(especie, urgencia) RETURNS NUMERIC en PL/pgSQL y probarla con SELECT sobre las 3 especies.",
            "Crear la tabla audit_cita y el trigger de auditoria en sus dos objetos: fn_trg_audit_cita() RETURNS TRIGGER + CREATE TRIGGER ... EXECUTE FUNCTION.",
            "Crear el trigger de stock no negativo (BEFORE UPDATE), evidenciando primero que sin el el stock llega a -7.",
            "Decidir donde vive cada validacion: CHECK, trigger o aplicacion (pregunta 4).",
            "Redactar Plan_Backup_VetCare con sus 6 secciones (plantilla en este documento): que se respalda y con que, frecuencia, retencion, RPO/RTO, restore de prueba con quien firma, y que NO cubre el plan."],
    quiz=True, sql="04_func_trigger_backup.sql"),
  dict(n=5, slug=None, titulo="Parcial 1", subtitulo="Solo evaluacion — Corte 1",
    herramienta=None, hito_pi="No avanza PI", entregable=None, teoria=[], demo=None, taller=[],
    quiz=False, sql=None, parcial="Parcial 1 - Administracion procedimientos y seguridad.docx"),
  dict(n=6, slug="Optimizacion de consultas",
    titulo="Optimizacion de consultas · VetCare",
    subtitulo="Antes/despues sobre el DDL del PI",
    # El taller se resuelve y se califica en ExamLab, que corre PostgreSQL (PGlite) y
    # trae la base con VOLUMEN sembrado: 30.010 citas, sin indices y con ANALYZE ya
    # corrido. Decia «DB Fiddle / SQLTest.online», que es donde NO se califica, y
    # ExamLab no aparecia en ninguna parte de la clase.
    herramienta="ExamLab (PostgreSQL) + Google Docs",
    hito_pi="Primera pareja de consultas antes/despues del PI",
    entregable="2 consultas (antes/despues) + justificacion (media pag.)",
    teoria=["Optimizar consultas parte de entender que el motor NO ejecuta el SQL tal cual se escribe: primero lo transforma en un plan de ejecucion (que tablas leer, en que orden, con o sin indice) y ese plan es lo que realmente determina el tiempo de respuesta.",
            "Tres cuellos de botella clasicos: (1) SELECT * trae columnas que nadie usa y aumenta el trafico/memoria; (2) JOIN sin filtro temprano obliga a cruzar tablas completas antes de descartar filas; (3) aplicar una funcion sobre la columna en el WHERE (ej. WHERE UPPER(nombre)='LUNA') impide que el motor use un indice normal sobre esa columna (esto se llama 'no-sargable').",
            "Reescritura tipica: proyectar solo columnas necesarias (SELECT nombre, fecha en vez de SELECT *), aplicar el filtro mas selectivo primero (WHERE fecha >= hoy antes del JOIN si reduce mucho el conjunto), y mover comparaciones a la forma que el motor pueda usar con indice.",
            # El cuarto cuello de botella: es la pregunta 3 del taller (20 pts) y una de
            # las cuatro afirmaciones correctas de la 4. No estaba en ninguna bala ni en
            # ninguna diapositiva, asi que se evaluaban ~22 puntos sin haberlo ensenado.
            "El cuarto cuello de botella, y el mas caro: una subconsulta correlacionada en la lista de columnas se evalua UNA VEZ POR FILA del exterior — el plan lo dice con loops=2006 —, y se elimina reescribiendola como LEFT JOIN + GROUP BY, que hace una sola pasada. Con LEFT JOIN hay que contar la columna, COUNT(c.id_cita), y no COUNT(*): el LEFT JOIN fabrica una fila de NULL por cada dueno sin citas, y COUNT(*) cuenta filas, asi que reportaria 1 donde la respuesta es 0.",
            "Optimizar NO puede cambiar el resultado, y eso se demuestra, no se afirma: un COUNT(*) de cada version que coincida, y para conjuntos completos un EXCEPT en los DOS sentidos que devuelva cero filas (EXCEPT no es simetrico: A EXCEPT B vacio no dice nada sobre filas de mas en B).",
            "EXPLAIN muestra el plan que el motor ESTIMA; EXPLAIN ANALYZE lo ejecuta de verdad y agrega actual rows, loops y Execution Time. En PostgreSQL el recorrido completo de tabla se llama Seq Scan (en Oracle, TABLE ACCESS FULL): verlo sobre una tabla grande donde se esperaba un indice es la senal de que el WHERE o el tipo de dato bloquea el indice.",
            "Conexion con Clase 7: optimizar consultas y crear indices son las dos caras de la misma moneda — una consulta mal escrita no aprovecha ni el mejor indice, y el mejor indice no compensa una consulta que fuerza un escaneo completo. Hoy NO se crea ningun indice: por eso la mejora que se mide es la de filas procesadas y pasadas sobre la tabla, no un cambio de Seq Scan a Index Scan.",
            "Error de docente que no domina el tema: pedir 'la consulta más rápida' sin definir contra que se compara (volumen de datos, indices existentes) — optimizar siempre es relativo a un antes medible, por eso el taller pide guardar la version antes Y despues, no solo la version final."],
    demo="Consulta pesada citas+mascotas+duenos -> version filtrada y proyectada, con EXPLAIN ANALYZE antes y despues, en ExamLab.",
    taller=["Reescribir la agenda del dia corrigiendo sus 4 antipatrones (SELECT *, joins con coma, to_char sobre la fecha, UPPER sobre el estado) y probar con COUNT(*) que las dos versiones devuelven las mismas 91 filas.",
            "Medir con EXPLAIN (ANALYZE, BUFFERS) las dos versiones, y con EXPLAIN ANALYZE una tercera que le anada LIMIT 50, y anotar las tres en comentarios: nodo mas costoso, filas estimadas vs reales y tiempo.",
            "Matar la subconsulta correlacionada del ranking de duenos: LEFT JOIN + GROUP BY + COUNT(c.id_cita), y demostrar la equivalencia con EXCEPT en los dos sentidos.",
            "Responder la de seleccion multiple sobre antipatrones (6 afirmaciones, 4 correctas).",
            "Escribir la justificacion tecnica de media pagina y guardar 06_opt_antes.sql / 06_opt_despues.sql en la carpeta del PI."],
    quiz=True, sql="06_opt_consultas.sql"),
  dict(n=7, slug="Indices y particionamiento",
    titulo="Indices y particionamiento · VetCare",
    subtitulo="Diseno fisico al servicio del PI",
    # Decia «DB Fiddle + draw.io (opcional)», y los 100 puntos del dia se califican en
    # ExamLab, que corre PostgreSQL sobre PGlite. En DB Fiddle no existe la base sembrada
    # de 30.010 citas, asi que quien mida ahi no puede reproducir el cambio de plan que la
    # rubrica exige; y el diagrama opcional se anunciaba en draw.io en esta linea y en
    # Excalidraw en el paso 4 del taller: dos herramientas para el mismo dibujo.
    herramienta="ExamLab (PostgreSQL/PGlite)",
    # El hito decia «>=2 indices» y la actividad exige TRES con nombre exacto (dos de
    # ellos compuestos en la pregunta 2) mas la tabla particionada de la pregunta 3.
    hito_pi="3 indices justificados (uno parcial) + historico particionado por ano",
    entregable="Script CREATE INDEX + cita_hist particionada + tabla justificacion consulta->indice",
    teoria=["Un indice es una estructura auxiliar (tipicamente un arbol B-Tree) que el motor mantiene ordenada por una o mas columnas, para encontrar filas sin recorrer toda la tabla — como el indice de un libro en vez de leer pagina por pagina.",
            "El costo no es gratis: cada INSERT/UPDATE/DELETE sobre una columna indexada obliga al motor a actualizar tambien el indice, asi que mas indices = lecturas mas rapidas pero escrituras mas lentas. Por eso 'indexar todo' es un error, no una optimizacion.",
            "Buen candidato a indice: columna usada muy frecuentemente en WHERE, JOIN u ORDER BY, con alta cardinalidad (muchos valores distintos, ej. id_dueno) — indexar una columna de baja cardinalidad (ej. un booleano activo S/N con solo 2 valores) rara vez ayuda porque el motor igual debe leer una fraccion enorme de la tabla.",
            # Decia «Cita(fecha_hora) ... Mascota(id_dueno) ... DetalleFactura(id_factura)».
            # Las tablas del curso son minusculas y con guion bajo —es la convencion que la
            # Clase 1 declara vinculante para todo el semestre, «nunca camelCase»— y el
            # estudiante tiene que escribir `cita`, `mascota` y `detalle_factura` para que su
            # CREATE INDEX corra. Un `DetalleFactura` proyectado aqui es un error que se
            # copia y no compila.
            "Candidatos reales en VetCare: cita(fecha_hora) para listar la agenda del dia, mascota(id_dueno) porque cada consulta de historial parte de un dueno, detalle_factura(id_factura) para armar el total de una factura sin escanear toda la tabla.",
            # Decia «idea conceptual, no se implementa hoy». La pregunta 3 vale 20 puntos y
            # se implementa entera: PARTITION BY RANGE, dos particiones, migracion, prueba
            # de enrutamiento y poda en el plan. El docente leia que hoy no se implementa y
            # el estudiante lo entregaba ejecutado.
            "Particionamiento (hoy SI se implementa, y es la pregunta 3 del taller): dividir una tabla logica en fragmentos fisicos por rango de fecha, de modo que la consulta de un ano no toque los datos del otro. PostgreSQL lo hace con PARTITION BY RANGE (fecha_hora) y una particion por ano. El indice ORDENA los datos; la particion los SEPARA.",
            "Con 5.010 filas la particion no acelera nada y hay que decirlo: lo que si se comprueba hoy es la poda de particiones en el plan (solo aparece cita_hist_2026) y el archivado, porque tirar un ano completo es un DROP TABLE de la particion en vez de un DELETE masivo.",
            "Error de docente que no domina el tema: crear un indice sobre CADA columna 'por si acaso' sin mirar que consultas realmente lo necesitan — el taller exige justificar cada indice con la consulta concreta que lo aprovecha."],
    # El nombre exacto que califica la pregunta 1 es `idx_cita_fecha_hora`. La demo decia
    # `idx_cita_fecha`, asi que el estudiante que copiaba la demo perdia puntos por el
    # nombre. El script 07 tambien lo decia mal y quedo corregido.
    demo="EXPLAIN ANALYZE con Seq Scan, CREATE INDEX idx_cita_fecha_hora, ANALYZE, y el mismo EXPLAIN mostrando Index Scan.",
    taller=["Medir la linea base con EXPLAIN ANALYZE de las dos consultas frecuentes: hay que ver Seq Scan.",
            "Crear los tres indices con el nombre exacto, incluido el parcial idx_cita_programada_fecha, y correr ANALYZE.",
            "Repetir los EXPLAIN y decir cual indice eligio el planeador y por que.",
            "Construir cita_hist particionada por ano, migrar las citas y demostrar el enrutamiento y la poda.",
            "Llenar la tabla de justificacion consulta->indice (7 columnas) y el veredicto de particionamiento."],
    quiz=True, sql="07_indices_vetcare.sql"),
  dict(n=8, slug="Tuning y transacciones",
    titulo="Tuning · Transacciones · VetCare",
    subtitulo="Atomicidad en facturacion e insumos",
    # Decia «Oracle Live SQL / DB Fiddle». Las tres preguntas de SQL del dia (75 de los
    # 100 puntos) son PL/pgSQL: `CALL`, `GET DIAGNOSTICS ... ROW_COUNT`, `RAISE EXCEPTION`
    # y una funcion que devuelve BOOLEAN. Oracle Live SQL no ejecuta ninguna de las tres.
    # Se anunciaba la herramienta en la que el entregable no compila.
    herramienta="ExamLab (PostgreSQL/PGlite)",
    hito_pi="Transaccion de negocio (factura + stock) + notas de tuning",
    entregable="sp_facturar + fn_descontar_stock + seccion Transacciones y tuning del informe (1 pag.)",
    teoria=["Una transaccion agrupa varias sentencias SQL en una sola unidad de todo-o-nada: si facturar implica INSERT en factura, INSERT en detalle_factura Y UPDATE de stock en insumo, las tres deben aplicarse juntas o ninguna — nunca queda una factura sin descontar stock, ni stock descontado sin factura.",
            "Propiedades ACID en una frase cada una: Atomicidad (todo o nada, ya explicado), Consistencia (la BD pasa de un estado valido a otro, respetando reglas como stock>=0), Aislamiento (transacciones concurrentes no se pisan entre si — se profundiza en Clase 10), Durabilidad (una vez hecho COMMIT, el dato sobrevive aunque el sistema se caiga un segundo despues).",
            # Decia «Sin ROLLBACK explicito ante el error, quedaria una factura registrada
            # sin el descuento real». En PostgreSQL eso es falso para el caso del dia y es
            # justo la opcion INCORRECTA de la pregunta 4 (10 pts): el `CALL` de nivel
            # superior es su propia transaccion, asi que la excepcion que se propaga deshace
            # todo sin que nadie escriba ROLLBACK. Se ensenaba el distractor como respuesta.
            "COMMIT confirma la transaccion de forma permanente; ROLLBACK deshace todo lo hecho desde que se abrio. En PostgreSQL no hace falta escribir ROLLBACK dentro del procedimiento: el CALL de nivel superior es su propia transaccion, y si la excepcion se propaga hasta afuera, el motor deshace todo lo que el procedimiento habia hecho. En Oracle si hay que escribirlo, y ese contraste es la pregunta 4 del taller.",
            "Lo que si hay que escribir es el guardia: UPDATE insumo SET stock = stock - p_cantidades[i] WHERE id_insumo = p_insumos[i] AND stock >= p_cantidades[i], y despues GET DIAGNOSTICS v_filas = ROW_COUNT. Si v_filas es 0 no hubo stock, y ahi se decide: RAISE EXCEPTION si el fallo debe abortar la factura, o devolver FALSE si el 'no hay stock' es una respuesta y no un error.",
            "Dirty read (lectura sucia): una transaccion lee un dato que otra transaccion modifico pero AUN NO ha confirmado con COMMIT; si esa segunda transaccion hace ROLLBACK despues, la primera trabajo con un dato que nunca existio de verdad. Es uno de los problemas que el nivel de aislamiento intenta evitar.",
            "Tuning en este contexto no es magia, son habitos concretos: mantener estadisticas del optimizador actualizadas (para que EXPLAIN elija bien), apoyarse en los indices ya justificados en Clase 7, y mantener las transacciones lo mas CORTAS posible — una transaccion larga retiene bloqueos (locks) sobre filas y puede frenar a otras transacciones que esperan esas mismas filas.",
            "Error de docente que no domina el tema: envolver TODA la sesion de trabajo en una sola transaccion gigante 'para no perder nada' — eso maximiza el tiempo que otros usuarios quedan bloqueados esperando esas filas, exactamente el problema que Clase 10 (concurrencia) va a diagnosticar."],
    # Los dos CALL son los del taller, con sus arreglos: el que funciona y el que falla a
    # mitad. El segundo descuenta el insumo 3 y se estrella en el 2, y el stock del 3 vuelve
    # solo a 40 sin que nadie escriba ROLLBACK: eso es lo que hay que ver proyectado.
    demo="CALL sp_facturar(4, ARRAY[1,6,5], ARRAY[1,2,3]) que factura 27.400, y CALL sp_facturar(4, ARRAY[3,2], ARRAY[2,10]) que falla en la segunda linea: el stock del insumo 3 vuelve a 40 sin ROLLBACK escrito.",
    taller=["Escribir sp_facturar(p_id_consulta, p_insumos INT[], p_cantidades INT[]) en PL/pgSQL: cabecera con total 0, bucle por linea con el guardia stock >= cantidad, y UPDATE del total al final.",
            "Probar el fallo a mitad con ARRAY[3,2] / ARRAY[2,10] y demostrar con foto inicial y final que el stock del insumo 3 volvio a 40.",
            "Encapsular el descuento en fn_descontar_stock, que devuelve BOOLEAN y no lanza excepcion.",
            "Llenar la seccion Transacciones y tuning del informe: inventario de 3 transacciones y checklist de 7 items.",
            "Declarar el gap de concurrencia: PGlite corre una sola sesion, y eso es la Clase 10."],
    quiz=True, sql="08_transacciones_vetcare.sql"),
  dict(n=9, slug=None, titulo="Parcial 2", subtitulo="Solo evaluacion — Corte 2",
    herramienta=None, hito_pi="No avanza PI", entregable=None, teoria=[], demo=None, taller=[],
    quiz=False, sql=None, parcial="Parcial 2 - Optimizacion indices y transacciones.docx"),
  dict(n=10, slug="Control de concurrencia",
    titulo="Control de concurrencia · VetCare",
    subtitulo="Clase autonoma · refuerzo sin parcial",
    herramienta="Google Docs + Live SQL",
    hito_pi="Escenarios de concurrencia del PI documentados",
    entregable="Informe corto: 2 escenarios (cita doble / stock) + mitigacion",
    teoria=["Concurrencia = varias transacciones ejecutandose al mismo tiempo sobre los mismos datos. El problema clasico de VetCare: dos recepcionistas, en dos computadores distintos, intentan agendar la MISMA franja horaria para el MISMO veterinario en el mismo instante; sin control, ambas lecturas ven la franja libre y ambas insertan — doble reserva.",
            "Control pesimista: asumir que el conflicto va a ocurrir, asi que se bloquea la fila (o el recurso) apenas se empieza a leer para modificar, y otras transacciones deben esperar a que termine (SELECT ... FOR UPDATE es el ejemplo tipico). Simple y seguro, pero puede generar esperas largas si hay muchas transacciones compitiendo.",
            "Control optimista: asumir que el conflicto es raro, dejar que todos lean libremente, y verificar SOLO al momento de escribir si alguien mas cambio el dato mientras tanto (comparando una version o timestamp); si hubo cambio, se rechaza y se reintenta. Mejor rendimiento cuando los conflictos son poco frecuentes.",
            "Deadlock (mencion breve): dos transacciones se bloquean mutuamente esperando un recurso que la otra tiene — T1 espera la fila que T2 bloqueo, y T2 espera la fila que T1 bloqueo. El motor detecta esto y aborta una de las dos automaticamente.",
            "Mitigaciones concretas y accesibles para el PI: una restriccion UNIQUE sobre (id_veterinario, fecha_hora) hace que el segundo INSERT falle automaticamente en vez de crear la doble reserva; transacciones cortas reducen la ventana de tiempo en la que puede ocurrir un conflicto; centralizar la logica en un procedimiento (Clase 3) evita que cada pantalla de la app implemente su propia validacion de forma inconsistente.",
            "Error de docente que no domina el tema: creer que 'poner una transaccion' ya resuelve la concurrencia — una transaccion garantiza atomicidad, pero sin un mecanismo de bloqueo o una restriccion UNIQUE, dos transacciones concurrentes pueden seguir generando la doble reserva porque ambas leen 'libre' antes de que la otra confirme."],
    demo="Narrativa paso a paso T1/T2 sobre tabla Cita.",
    taller=["Describir escenario doble reserva con tiempos T1/T2.",
            "Describir escenario doble descuento de stock.",
            "Proponer mitigacion SQL.",
            "Anadir seccion al informe PI."],
    quiz=True, sql="10_concurrencia_vetcare.sql"),
  dict(n=11, slug="Avance del proyecto final",
    titulo="Avance PI · VetCare DB",
    subtitulo="Checklist viva + demo parcial",
    herramienta="Live SQL / DB Fiddle + draw.io + ExamLab",
    hito_pi="Demo parcial + checklist de avance (hito formal PI)",
    entregable="Checklist firmada + enlace/ZIP avance (DDL+procs+ER)",
    teoria=["Hoy no hay tema nuevo: se cierran huecos del PI con rubrica.",
            "Evidencias: ER, DDL, roles, >=2 procs, >=1 fn, >=2 triggers, 1 opt.",
            "Revision cruzada entre estudiantes: 10 min por persona."],
    demo="Recorrido de checklist + ejemplo demo de 3 min.",
    taller=["Completar checklist de avance (si/no/parcial).",
            "Demo 3-5 min: ER + 1 proc + 1 trigger.",
            "Lista de gaps con responsable.",
            "Subir avance intermedio a ExamLab (Talleres) si se pide."],
    quiz=True, sql="11_checklist_seed.sql"),
  dict(n=12, slug="Integracion y preparacion final",
    titulo="Integracion app <-> BD · Prep. presentacion",
    subtitulo="Contrato de operaciones + ensayo PI",
    herramienta="Google Docs + Live SQL + Excalidraw",
    hito_pi="Contrato integracion + preparacion de entrega/sustentacion",
    entregable="Contrato app<->BD + outline de slides de sustentacion (5-8 min)",
    teoria=["Integrar app<->BD significa que la aplicacion NUNCA arma SQL dinamico contra las tablas directamente; llama procedimientos y funciones ya construidos (Clases 3-4). Esto evita SQL injection (nadie concatena texto de usuario dentro de una consulta), centraliza la regla de negocio en un solo lugar, y permite cambiar el esquema interno sin romper la app mientras el contrato del proc se mantenga igual.",
            "Un contrato de integracion documenta, por cada operacion: nombre del proc, parametros de entrada con su tipo, que retorna (valor OUT o codigo de resultado), y que errores puede lanzar y con que significado (ej. 'ERROR: mascota inactiva' vs una excepcion no controlada del motor). Sin este contrato, cualquier desarrollador que use la BD debe adivinar el comportamiento leyendo el codigo SQL directamente.",
            "Manejo de errores en la frontera app-BD: la app no deberia mostrar al usuario final un error crudo de base de datos (ej. 'ORA-00001: unique constraint violated'); el proc devuelve un mensaje o codigo de negocio legible, y la app lo traduce a un mensaje humano ('Ya existe una cita en ese horario').",
            "Autenticacion/autorizacion en este punto es conceptual, no de implementacion: la app se conecta con una cuenta de servicio que respeta los roles definidos en Clase 2 (principio de minimo privilegio) — la app de recepcion no deberia poder ejecutar procs reservados a auditoria o administracion.",
            "Preparar la sustentacion no es 'hacer diapositivas bonitas': es organizar la evidencia tecnica en una narrativa logica -> problema real que resuelve VetCare, modelo de datos (ER + normalizacion), seguridad (roles), automatizacion (procs/triggers), rendimiento (indices/optimizacion), y una demo en vivo que conecte todo eso con una operacion real (agendar una cita, facturar).",
            "Error de docente que no domina el tema: dejar que la 'integracion' quede como una idea abstracta sin contrato escrito — el entregable de hoy exige documentar minimo 3 operaciones con su firma completa, no solo mencionarlas de palabra."],
    demo="Plantilla contrato sp_agendar_cita + storyboard 6 slides.",
    taller=["Redactar contrato de >=3 operaciones.",
            "Diagrama flujo app->BD (Excalidraw) opcional.",
            "Outline presentacion 5-8 min + quien habla que.",
            "Empaquetar borrador entrega final."],
    quiz=True, sql="12_contrato_ops.sql"),
  dict(n=13, slug="Analisis de casos reales",
    titulo="Analisis de casos reales · VetCare",
    subtitulo="Clase autonoma · lecciones para el PI",
    herramienta="Google Docs",
    hito_pi="Informe de caso -> mejoras concretas al PI",
    entregable="Informe 1-2 pag.: caso + 3 mejoras aplicables a VetCare",
    teoria=["Caso 1 — falta de backup real: una organizacion que 'hacia backup' copiando el archivo de datos una vez al mes sin probar nunca el restore. Cuando el disco fallo, el archivo copiado estaba corrupto (nunca se verifico) y perdieron meses de informacion. Leccion para VetCare: un backup que nunca se restauro de prueba no cuenta como backup funcional (conecta con Clase 4: RPO/RTO y prueba de restore).",
            "Caso 2 — indices mal disenados: un sistema con un indice sobre CADA columna 'por si acaso', que volvia cada INSERT/UPDATE mas lento de lo aceptable, sin que nadie hubiera medido si esos indices realmente se usaban en consultas reales. Leccion: indexar sin justificar la consulta que lo aprovecha (conecta con Clase 7) desperdicia recursos y no mejora nada.",
            "Caso 3 — inyeccion SQL: una aplicacion que concatenaba directamente el texto escrito por el usuario dentro de una consulta (ej. \"SELECT * FROM usuarios WHERE nombre='\" + input + \"'\"), permitiendo que alguien escribiera un valor que alterara la consulta completa y expusiera o borrara datos ajenos. Leccion: por eso la app llama procedimientos con parametros tipados (Clase 3 y Clase 12) en vez de armar SQL con texto libre.",
            "Estructura para analizar cualquier caso real: (1) contexto — que sistema era y que se suponia que hacia bien; (2) fallo — que paso exactamente y por que la causa raiz no era 'mala suerte' sino una decision tecnica evitable; (3) leccion — que principio general se puede extraer; (4) cambio concreto — que se ajusta HOY en su propio VetCare, no en abstracto.",
            "Esta clase es autonoma (sin encuentro sincrono) precisamente porque no introduce tecnica nueva: aplica en modo reflexivo/critico todo lo visto en Clases 1-10 sobre un caso real, cerrando el ciclo antes de entrar a integracion y cierre del PI.",
            "Error de docente que no domina el tema: dejar que el informe describa el caso ajeno sin conectar ninguna leccion con una accion verificable en VetCare — el entregable exige 3 mejoras concretas aplicadas al proyecto propio, no un resumen de noticia."],
    demo="Plantilla: contexto -> fallo -> leccion -> cambio en VetCare.",
    taller=["Elegir 1 caso (backup, rendimiento o seguridad).",
            "Resumir en media pagina que paso.",
            "Proponer 3 mejoras concretas a su VetCare.",
            "Actualizar informe PI con lecciones de casos."],
    quiz=True, sql=None),
  dict(n=14, slug=None, titulo="Parcial 3", subtitulo="Solo evaluacion — Corte 3",
    herramienta=None, hito_pi="Prep PI fue Clase 12; cierre en Clase 15", entregable=None,
    teoria=[], demo=None, taller=[], quiz=False, sql=None,
    parcial="Parcial 3 - Integracion casos y cierre de proyecto.docx"),
  dict(n=15, slug="Presentacion del proyecto y cierre",
    titulo="Presentacion PI · Cierre VetCare",
    subtitulo="Sustentacion en vivo del PI · cierre del curso",
    herramienta="ExamLab (Proyectos) + slides propias",
    hito_pi="Sustentacion en vivo y entrega final del PI (20% Corte 3)",
    entregable="ZIP/PDF final subido antes del turno + sustentacion en vivo 5-8 min + Q&A",
    teoria=["Cierre: evidencias completas segun rubrica (100 pts -> 20%).",
            "Sustentacion breve alineada a criterios del enunciado.",
            "Autoevaluacion del propio proceso de trabajo."],
    demo="Checklist final de empaquetado del ZIP.",
    taller=["Subir el paquete final a ExamLab (modulo Proyectos) ANTES de su turno.",
            "Sustentar en vivo 5-8 min con el ER y una ejecucion real en pantalla.",
            "Responder el Q&A en vivo del docente (preguntas al azar sobre su modelo).",
            "Autoevaluacion: que harian distinto.",
            "Cierre del curso."],
    quiz=True, sql=None),
]

# Fuente de verdad del tipo de bloque: el calendario del periodo (ver tipo_de_clase).
for _c in CLASES:
    _c["tipo"] = tipo_de_clase(_c["n"])

SQL_BODIES = {
"01_arranque_vetcare.sql": """-- VetCare DB · Clase 1 · DDL minimo demo (DB Fiddle / PostgreSQL o MySQL)
-- Objetivo PI: dejar entidades base para el ER.

CREATE TABLE dueno (
  id_dueno INT PRIMARY KEY,
  nombre VARCHAR(80) NOT NULL,
  telefono VARCHAR(30),
  email VARCHAR(120)
);

CREATE TABLE mascota (
  id_mascota INT PRIMARY KEY,
  id_dueno INT NOT NULL REFERENCES dueno(id_dueno),
  nombre VARCHAR(60) NOT NULL,
  especie VARCHAR(40),
  activa CHAR(1) DEFAULT 'S' CHECK (activa IN ('S','N'))
);

CREATE TABLE cita (
  id_cita INT PRIMARY KEY,
  id_mascota INT NOT NULL REFERENCES mascota(id_mascota),
  fecha_hora TIMESTAMP NOT NULL,
  estado VARCHAR(20) DEFAULT 'PROGRAMADA'
);

INSERT INTO dueno VALUES (1, 'Ana Perez', '3001112233', 'ana@mail.com');
INSERT INTO mascota VALUES (10, 1, 'Luna', 'Canino', 'S');
INSERT INTO cita VALUES (100, 10, '2026-09-01 09:00:00', 'PROGRAMADA');
SELECT m.nombre, d.nombre AS dueno, c.fecha_hora
FROM cita c JOIN mascota m ON m.id_mascota=c.id_mascota
JOIN dueno d ON d.id_dueno=m.id_dueno;
""",
"02_roles_vetcare.sql": """-- VetCare DB · Clase 2 · Roles y privilegios · PostgreSQL
-- Este es el script de la DEMO: corre tal cual en ExamLab (PostgreSQL en el
-- navegador), sobre el esquema de VetCare ya creado. Los nombres de rol son los
-- mismos que pide el taller: minusculas, y el del veterinario con sufijo `_rol`
-- porque `veterinario` ya es una tabla.
--
-- Se ejecuta de arriba abajo, narrando cada bloque. El bloque 5 es el que convence
-- al grupo: es la unica parte donde se VE que un permiso negado detiene una
-- sentencia — y por eso mismo va SENTENCIA POR SENTENCIA, no de un solo tiro: dos
-- de sus lineas tienen que fallar, y un runner que aborta al primer error se
-- llevaria las siguientes. Correr el script completo una vez antes de la clase.

-- ============ 1) Los cuatro roles ============
-- NOLOGIN: son paquetes de privilegios, no cuentas con las que alguien entra.
-- Una persona recibe el rol despues, con GRANT recepcion TO ana_gomez.
CREATE ROLE admin_bd NOLOGIN;
CREATE ROLE recepcion NOLOGIN;
CREATE ROLE veterinario_rol NOLOGIN;
CREATE ROLE auditor NOLOGIN;

-- ============ 2) Los privilegios, uno por uno ============
-- admin_bd es el unico con privilegios amplios, y sobre las 8 tablas.
GRANT ALL PRIVILEGES ON dueno, mascota, veterinario, cita,
                        consulta, insumo, factura, detalle_factura TO admin_bd;

-- recepcion opera citas y solo LEE los datos con que identifica a quien llama.
GRANT SELECT, INSERT, UPDATE ON cita TO recepcion;
GRANT SELECT ON dueno, mascota, veterinario TO recepcion;

-- veterinario_rol registra la consulta; la cita y la mascota solo las lee.
GRANT SELECT ON cita, mascota TO veterinario_rol;
GRANT SELECT, INSERT, UPDATE ON consulta TO veterinario_rol;

-- auditor verifica: solo lectura, sobre todo lo sensible.
GRANT SELECT ON dueno, mascota, cita, consulta, factura TO auditor;

-- El REVOKE se deja escrito aunque sea redundante (nunca se otorgo DELETE):
-- es la evidencia de una decision de diseno, no una correccion.
REVOKE DELETE ON cita FROM recepcion;

-- ============ 3) La matriz sale del motor, no del documento ============
SELECT grantee, table_name, privilege_type
FROM information_schema.role_table_grants
WHERE grantee IN ('admin_bd', 'recepcion', 'veterinario_rol', 'auditor')
ORDER BY grantee, table_name, privilege_type;

-- ============ 4) Cuando el GRANT es demasiado ============
-- La vista recorta filas (las canceladas) y columnas (el email nunca sale).
-- Se ejecuta con los privilegios de SU PROPIETARIO: por eso recepcion puede
-- consultarla aunque le quitemos el SELECT sobre la tabla dueno.
CREATE VIEW v_agenda_recepcion AS
SELECT c.id_cita, c.fecha_hora, c.estado,
       m.nombre AS mascota, d.nombre AS dueno, d.telefono,
       v.nombre AS veterinario
FROM cita c
JOIN mascota m     ON m.id_mascota = c.id_mascota
JOIN dueno d       ON d.id_dueno = m.id_dueno
JOIN veterinario v ON v.id_veterinario = c.id_veterinario
WHERE c.estado <> 'CANCELADA';

GRANT SELECT ON v_agenda_recepcion TO recepcion;
REVOKE SELECT ON dueno FROM recepcion;   -- ahora solo llega por la vista

-- Privilegio por columna: dos columnas y ninguna otra, sin crear objeto nuevo.
GRANT SELECT (id_dueno, nombre) ON dueno TO veterinario_rol;

-- Evidencia: tiene que devolver EXACTAMENTE dos filas (id_dueno y nombre).
SELECT grantee, table_name, column_name, privilege_type
FROM information_schema.column_privileges
WHERE grantee = 'veterinario_rol' AND table_name = 'dueno'
ORDER BY column_name;

-- ============ 5) La prueba negativa: ver el permiso NEGADO ============
-- No hay una segunda conexion (el entorno tiene un solo usuario con login), pero
-- no hace falta: SET ROLE cambia el rol efectivo DENTRO de la misma sesion, y a
-- partir de ahi los permisos que se revisan son los del rol, no los del dueno.
SET ROLE recepcion;

SELECT id_cita, fecha_hora, dueno FROM v_agenda_recepcion;  -- OK: la vista si
SELECT nombre, email FROM dueno;   -- debe fallar: permission denied for table dueno
DELETE FROM cita WHERE id_cita = 1;  -- debe fallar: permission denied for table cita

RESET ROLE;   -- volver al propietario ANTES de seguir con cualquier otra cosa

-- Si su entorno no permite SET ROLE, no lo esconda: dejelo en pantalla, diga que
-- eso vuelve la prueba negativa una brecha de verificacion del entregable, y
-- muestre cual seria el comando en un servidor real.

-- ============ 6) Ciclo de vida, para la politica ============
-- Alta:   GRANT recepcion TO ana_gomez;
-- Cambio: GRANT veterinario_rol TO ana_gomez; REVOKE recepcion FROM ana_gomez;
--         (los dos, siempre: los permisos NO se acumulan)
-- Baja:   REASSIGN OWNED BY ana_gomez TO admin_bd;  -- antes de borrar el rol
--         DROP ROLE ana_gomez;                      -- falla si todavia posee objetos
""",
"03_procs_vetcare.sql": """-- VetCare DB · Clase 3 · Procedimientos almacenados · PostgreSQL
-- Script de la DEMO: corre tal cual en ExamLab (PostgreSQL/PGlite en el navegador),
-- sobre el esquema de VetCare ya creado y poblado: 8 mascotas (Rocky=3 y Kiara=8
-- estan INACTIVAS), 4 veterinarios, 10 citas, y una cita del veterinario 1 el
-- 2026-09-01 08:00:00.
--
-- NO es Oracle: nada de IS en vez de AS, VARCHAR2, NUMBER, RAISE_APPLICATION_ERROR
-- ni barra / de terminacion. Ese codigo aqui no compila, y es la forma mas facil de
-- perder los puntos de sintaxis de la pregunta 1.
--
-- Se ejecuta de arriba abajo narrando cada bloque. El bloque 3 es el que convence al
-- grupo: es donde se VE que la validacion detiene el INSERT. Correr el script
-- completo una vez antes de la clase.

-- ============ 1) El procedimiento con sus 3 validaciones ============
-- id_cita no se pasa como parametro: es SERIAL y lo genera el motor.
CREATE OR REPLACE PROCEDURE sp_agendar_cita(
  p_id_mascota     INT,
  p_id_veterinario INT,
  p_fecha_hora     TIMESTAMP
)
LANGUAGE plpgsql
AS $proc$
DECLARE
  v_activa CHAR(1);
BEGIN
  -- Validacion 1: la mascota existe. SELECT ... INTO deja FOUND en FALSE cuando no
  -- devolvio ninguna fila, y eso es lo que pregunta IF NOT FOUND.
  SELECT activa INTO v_activa
    FROM mascota
   WHERE id_mascota = p_id_mascota;

  IF NOT FOUND THEN
    RAISE EXCEPTION 'ERROR: la mascota % no existe', p_id_mascota;
  END IF;

  -- Validacion 2: la regla de negocio del PI.
  IF v_activa <> 'S' THEN
    RAISE EXCEPTION 'ERROR: la mascota % esta inactiva; no se agenda cita',
                    p_id_mascota;
  END IF;

  -- Validacion 3: la franja del veterinario esta libre. Una cita CANCELADA libera
  -- la franja, asi que no cuenta.
  IF EXISTS (SELECT 1 FROM cita
              WHERE id_veterinario = p_id_veterinario
                AND fecha_hora     = p_fecha_hora
                AND estado <> 'CANCELADA') THEN
    RAISE EXCEPTION 'ERROR: el veterinario % ya tiene cita en %',
                    p_id_veterinario, p_fecha_hora;
  END IF;

  INSERT INTO cita (id_mascota, id_veterinario, fecha_hora, estado)
  VALUES (p_id_mascota, p_id_veterinario, p_fecha_hora, 'PROGRAMADA');
END;
$proc$;

-- ============ 2) El caso valido ============
CALL sp_agendar_cita(1, 2, TIMESTAMP '2026-09-15 10:00:00');

SELECT id_cita, id_mascota, id_veterinario, fecha_hora, estado
  FROM cita ORDER BY id_cita DESC LIMIT 3;   -- la nueva es la primera fila

-- ============ 3) Los tres errores, SENTENCIA POR SENTENCIA ============
-- Estas tres lineas DEBEN fallar, y por eso no van en un solo tiro: la gracia es
-- leer en pantalla el mensaje exacto que la app va a recibir. Un runner que aborta
-- al primer error se llevaria las siguientes.
CALL sp_agendar_cita(3,  2, TIMESTAMP '2026-09-21 08:00:00');  -- Rocky, INACTIVA
CALL sp_agendar_cita(99, 2, TIMESTAMP '2026-09-22 08:00:00');  -- no existe
CALL sp_agendar_cita(2,  1, TIMESTAMP '2026-09-01 08:00:00');  -- franja ocupada

-- Y la prueba de que no dejaron basura: sigue habiendo 11 citas, no 14.
SELECT COUNT(*) AS citas_totales FROM cita;

-- ============ 4) La bateria: un bloque DO por caso ============
-- Por que un bloque por caso: si los CALL van seguidos, el primero que falla aborta
-- el resto. DO es un bloque anonimo -- se ejecuta una vez y no se guarda -- y su
-- EXCEPTION atrapa el error y deja seguir al caso siguiente.
CREATE TABLE IF NOT EXISTS resultado_prueba (
  id_prueba SERIAL PRIMARY KEY,
  caso      TEXT,
  esperado  TEXT,
  obtenido  TEXT,
  paso      BOOLEAN
);

-- Caso POSITIVO: el exito es que NO haya excepcion.
DO $$
BEGIN
  CALL sp_agendar_cita(1, 2, TIMESTAMP '2026-09-20 08:00:00');
  INSERT INTO resultado_prueba (caso, esperado, obtenido, paso)
  VALUES ('P1 mascota activa', 'OK: cita creada', 'OK: cita creada', TRUE);
EXCEPTION WHEN OTHERS THEN
  INSERT INTO resultado_prueba (caso, esperado, obtenido, paso)
  VALUES ('P1 mascota activa', 'OK: cita creada', SQLERRM, FALSE);
END $$;

-- Caso NEGATIVO: el exito es que SI haya excepcion, y ademas que sea LA esperada.
-- Por eso se verifica el TEXTO con ILIKE y no basta WHEN OTHERS a secas: un typo en
-- el nombre de una columna tambien lanza excepcion, y un WHEN OTHERS pelado lo
-- reportaria como prueba superada.
DO $$
BEGIN
  CALL sp_agendar_cita(3, 2, TIMESTAMP '2026-09-21 08:00:00');
  INSERT INTO resultado_prueba (caso, esperado, obtenido, paso)
  VALUES ('P2 mascota inactiva', 'EXCEPCION: mascota inactiva',
          'NO lanzo excepcion: la cita se creo', FALSE);
EXCEPTION WHEN OTHERS THEN
  INSERT INTO resultado_prueba (caso, esperado, obtenido, paso)
  VALUES ('P2 mascota inactiva', 'EXCEPCION: mascota inactiva',
          SQLERRM, SQLERRM ILIKE '%inactiva%');
END $$;

SELECT caso, esperado, obtenido, paso
  FROM resultado_prueba ORDER BY id_prueba;

-- Nota de lectura: aqui `paso` significa «el resultado coincidio con lo esperado»,
-- asi que las dos filas quedan en t. La otra lectura -- «la operacion se completo»,
-- que deja las negativas en f -- tambien es valida. Lo que hay que hacer es usar UNA
-- de las dos para las cuatro filas y decir cual, porque si no, `paso` no significa
-- nada.

-- ============ 5) El contrato, que es el otro entregable ============
-- Firma        : sp_agendar_cita(p_id_mascota INT, p_id_veterinario INT,
--                               p_fecha_hora TIMESTAMP)
-- Llamada      : CALL sp_agendar_cita(1, 2, TIMESTAMP '2026-09-15 10:00:00');
-- Precondicion : la mascota existe y tiene activa = 'S'; la franja del veterinario
--                esta libre (una cita CANCELADA no la ocupa).
-- Postcondicion: 1 fila nueva en cita con estado 'PROGRAMADA'. Si falla, NINGUNA.
-- Errores      : 'ERROR: la mascota % no existe'
--                'ERROR: la mascota % esta inactiva; no se agenda cita'
--                'ERROR: el veterinario % ya tiene cita en %'
-- Decision     : se aborta con RAISE EXCEPTION en vez de devolver un codigo en un
--                parametro OUT, porque abortar deshace lo hecho; un codigo que
--                nadie revise deja la cita creada igual.
""",
"04_func_trigger_backup.sql": """-- VetCare DB · Clase 4 · Funcion, triggers y respaldo · PostgreSQL
-- Script de la DEMO: corre tal cual en ExamLab (PostgreSQL/PGlite en el navegador).
--
-- En PostgreSQL el trigger son SIEMPRE dos objetos: una funcion que RETURNS TRIGGER
-- y un CREATE TRIGGER que dice cuando dispararla. No existe el trigger con el cuerpo
-- adentro que se escribe en Oracle, ni los dos puntos de :NEW y :OLD, ni
-- RAISE_APPLICATION_ERROR. Eso aqui no compila y la rubrica lo descuenta.

-- ============ 1) La funcion de precio ============
-- IMMUTABLE: para los mismos argumentos siempre devuelve lo mismo y no lee tablas.
-- COALESCE porque la app puede mandar NULL en la casilla de urgencia, y NULL * 1.35
-- es NULL: la factura saldria vacia en vez de salir mal, que es peor.
CREATE OR REPLACE FUNCTION fn_precio_consulta(
  p_especie  TEXT,
  p_urgencia BOOLEAN
)
RETURNS NUMERIC
LANGUAGE plpgsql
IMMUTABLE
AS $fn$
DECLARE
  v_base NUMERIC;
BEGIN
  v_base := CASE UPPER(p_especie)
              WHEN 'CANINO' THEN 45000
              WHEN 'FELINO' THEN 40000
              ELSE 35000
            END;
  IF COALESCE(p_urgencia, FALSE) THEN
    v_base := v_base * 1.35;
  END IF;
  RETURN v_base;
END;
$fn$;

-- Una funcion se llama con SELECT, no con CALL. Es la diferencia con la Clase 3.
SELECT fn_precio_consulta('Canino', FALSE) AS normal,     -- 45000
       fn_precio_consulta('Canino', TRUE)  AS urgencia,   -- 60750
       fn_precio_consulta('canino', TRUE)  AS minusculas, -- 60750, por UPPER()
       fn_precio_consulta('Conejo', FALSE) AS otra_especie, -- 35000
       fn_precio_consulta('Felino', NULL)  AS urgencia_nula; -- 40000, por COALESCE

-- Y donde se usa de verdad: junto a la tabla, como una columna calculada.
SELECT m.nombre, m.especie,
       fn_precio_consulta(m.especie, FALSE) AS precio_normal,
       fn_precio_consulta(m.especie, TRUE)  AS precio_urgencia
  FROM mascota m
 WHERE m.id_mascota IN (1, 4)
 ORDER BY m.id_mascota;

-- ============ 2) Trigger de auditoria: los DOS objetos ============
CREATE TABLE IF NOT EXISTS audit_cita (
  id_audit        SERIAL PRIMARY KEY,
  id_cita         INT  NOT NULL,
  accion          TEXT NOT NULL,
  valor_anterior  TEXT,
  valor_nuevo     TEXT,
  usuario_bd      TEXT      DEFAULT current_user,
  fecha_evento    TIMESTAMP DEFAULT now()
);

-- Objeto 1: la funcion. NEW y OLD sin dos puntos, y RETURN NEW obligatorio.
CREATE OR REPLACE FUNCTION fn_trg_audit_cita()
RETURNS TRIGGER
LANGUAGE plpgsql
AS $fn$
BEGIN
  INSERT INTO audit_cita (id_cita, accion, valor_anterior, valor_nuevo)
  VALUES (NEW.id_cita, 'CAMBIO_ESTADO', OLD.estado, NEW.estado);
  RETURN NEW;   -- quien y cuando los pone los DEFAULT de la tabla
END;
$fn$;

-- Objeto 2: la asociacion. AFTER porque solo se registra lo que ya paso.
DROP TRIGGER IF EXISTS trg_audit_cita ON cita;
CREATE TRIGGER trg_audit_cita
AFTER UPDATE OF estado ON cita
FOR EACH ROW
WHEN (OLD.estado IS DISTINCT FROM NEW.estado)
EXECUTE FUNCTION fn_trg_audit_cita();

-- La prueba: TRES updates que dejan DOS filas de auditoria.
UPDATE cita SET estado = 'CANCELADA'  WHERE id_cita = 1;  -- cambia  -> audita
UPDATE cita SET estado = 'ATENDIDA'   WHERE id_cita = 3;  -- cambia  -> audita
UPDATE cita SET estado = 'PROGRAMADA' WHERE id_cita = 6;  -- ya era  -> NO audita

SELECT id_audit, id_cita, accion, valor_anterior, valor_nuevo, usuario_bd
  FROM audit_cita ORDER BY id_audit;   -- 2 filas: citas 1 y 3

-- El WHEN es lo que hace la diferencia. Sin el, la tercera fila tambien se escribe y
-- la auditoria se llena de eventos donde no cambio nada. Con IS DISTINCT FROM y no
-- con <> porque <> devuelve NULL si un lado es NULL, y un WHEN que da NULL no
-- dispara: un estado que pasa de NULL a 'PROGRAMADA' se quedaria sin auditar.

-- ============ 3) Trigger que IMPIDE: stock negativo ============
-- El CHECK de la tabla se retira a proposito para mostrar el hueco que tapa el
-- trigger. Un CHECK vigila el valor final de UNA fila; el trigger, ademas, puede
-- mirar el valor anterior y decidir con la fila completa.
ALTER TABLE insumo DROP CONSTRAINT IF EXISTS insumo_stock_check;

-- Sin defensa: el insumo 2 (Vacuna triple felina) tiene 3 unidades.
UPDATE insumo SET stock = stock - 10 WHERE id_insumo = 2;
SELECT id_insumo, nombre, stock FROM insumo WHERE id_insumo = 2;   -- stock = -7 (!)
UPDATE insumo SET stock = 3 WHERE id_insumo = 2;                   -- se restaura

CREATE OR REPLACE FUNCTION fn_trg_stock_no_negativo()
RETURNS TRIGGER
LANGUAGE plpgsql
AS $fn$
BEGIN
  IF NEW.stock < 0 THEN
    RAISE EXCEPTION 'ERROR: el stock de % no puede quedar negativo (resultado: %)',
                    OLD.nombre, NEW.stock;
  END IF;
  RETURN NEW;   -- BEFORE: lo que se retorna es lo que se guarda
END;
$fn$;

-- BEFORE, no AFTER: la unica forma de impedir el cambio es correr antes de que se
-- escriba. Un AFTER que lanza excepcion tambien deshace la transaccion, pero para
-- cuando corre el motor ya hizo el trabajo -- y con AFTER no se puede corregir el
-- valor, solo abortar.
DROP TRIGGER IF EXISTS trg_stock_no_negativo ON insumo;
CREATE TRIGGER trg_stock_no_negativo
BEFORE UPDATE OF stock ON insumo
FOR EACH ROW
EXECUTE FUNCTION fn_trg_stock_no_negativo();

-- Con defensa: el mismo UPDATE, ahora rechazado. RAISE NOTICE imprime el mensaje sin
-- abortar el bloque, para que el grupo lea la excepcion en pantalla.
DO $$
BEGIN
  UPDATE insumo SET stock = stock - 10 WHERE id_insumo = 2;
  RAISE NOTICE 'FALLO LA PRUEBA: el UPDATE paso y no debia';
EXCEPTION WHEN OTHERS THEN
  RAISE NOTICE 'RECHAZADO (correcto): %', SQLERRM;
END $$;

-- Y el descuento legitimo sigue funcionando: no se bloqueo la operacion, se bloqueo
-- el resultado invalido.
DO $$
BEGIN
  UPDATE insumo SET stock = stock - 2 WHERE id_insumo = 2;
  RAISE NOTICE 'ACEPTADO (correcto): quedan 1 unidades';
EXCEPTION WHEN OTHERS THEN
  RAISE NOTICE 'FALLO LA PRUEBA: %', SQLERRM;
END $$;

SELECT id_insumo, nombre, stock FROM insumo WHERE id_insumo = 2;   -- stock = 1

-- ============ 4) El respaldo: las herramientas reales ============
-- Estos comandos NO corren dentro de ExamLab -- son de linea de comandos, no SQL --
-- pero son los que hay que nombrar en el plan. Se proyectan como referencia.
--
--   pg_dump -Fc -d vetcare -f vetcare_2026-09-15.dump   respaldo logico de LA base
--   pg_dumpall --globals-only -f roles.sql              roles y privilegios: pg_dump
--                                                       NO los incluye
--   pg_basebackup -D /backup/base -Ft -z                copia fisica del cluster
--   pg_restore -d vetcare_prueba vetcare_2026-09-15.dump   el ensayo de restauracion
--
-- La consulta de validacion despues de restaurar, que es lo que convierte «restaure»
-- en «restaure bien»:
--   SELECT (SELECT COUNT(*) FROM cita)     AS citas,
--          (SELECT COUNT(*) FROM consulta) AS consultas,
--          (SELECT COUNT(*) FROM factura)  AS facturas,
--          (SELECT MAX(fecha_hora) FROM cita) AS ultima_cita;
--
-- Y el esqueleto del plan (1 pagina, en Google Docs): 1) que se respalda y con que
-- herramienta cada cosa · 2) frecuencia y ventana, justificada contra el horario
-- lunes-sabado 7:00-19:00 · 3) retencion, en >=2 ubicaciones · 4) RPO y RTO con su
-- justificacion por impacto · 5) el ensayo de restauracion: cada cuanto, la consulta
-- de validacion y quien firma · 6) que NO cubre este plan y cual es el riesgo
-- residual que se asume.
""",
# El script anterior eran 14 lineas: la pareja antes/despues SIN datos, sin
# EXPLAIN y con fecha 2026-09-01, mientras el taller pide 2026-03-10. Corrido
# sobre la base de 20 filas de la Clase 1 devolvia 3 filas y ninguna medicion, o
# sea exactamente el «celebrar que el tiempo bajo de 0,9 a 0,4 ms» que el propio
# fundamento marca como el error del docente que no domina el tema. Tampoco
# cubria la subconsulta correlacionada (pregunta 3, 20 puntos) ni la prueba de
# equivalencia. Ahora siembra el MISMO volumen que el `setup_sql` de ExamLab, de
# modo que los numeros de la demo son los que el estudiante va a ver.
"06_opt_consultas.sql":"""-- VetCare DB · Clase 6 · Optimizacion de consultas (demo del docente)
-- ============================================================================
-- Motor: PostgreSQL. Se corre en ExamLab (PGlite en el navegador), que es donde
-- se califica el taller. Es AUTOCONTENIDO: crea el esquema, siembra el volumen y
-- deja las estadisticas listas, igual que el `setup_sql` de las preguntas 1, 2 y 3.
-- Volumen resultante: 2.006 duenos · 5.008 mascotas · 16 veterinarios · 30.010 citas.
-- SIN indices adicionales: crearlos es la Clase 7, y por eso hoy la evidencia NO es
-- un cambio de Seq Scan a Index Scan sino menos filas y menos pasadas.
-- ============================================================================

DROP TABLE IF EXISTS cita, mascota, veterinario, dueno;

CREATE TABLE dueno (
  id_dueno SERIAL PRIMARY KEY,
  nombre   TEXT NOT NULL,
  telefono TEXT,
  email    TEXT,
  ciudad   TEXT DEFAULT 'Cali'
);
CREATE TABLE mascota (
  id_mascota SERIAL PRIMARY KEY,
  id_dueno   INT NOT NULL REFERENCES dueno(id_dueno),
  nombre     TEXT NOT NULL,
  especie    TEXT NOT NULL,
  fecha_nac  DATE,
  activa     CHAR(1) NOT NULL DEFAULT 'S' CHECK (activa IN ('S','N'))
);
CREATE TABLE veterinario (
  id_veterinario SERIAL PRIMARY KEY,
  nombre         TEXT NOT NULL,
  especialidad   TEXT,
  activo         CHAR(1) NOT NULL DEFAULT 'S' CHECK (activo IN ('S','N'))
);
CREATE TABLE cita (
  id_cita        SERIAL PRIMARY KEY,
  id_mascota     INT NOT NULL REFERENCES mascota(id_mascota),
  id_veterinario INT NOT NULL REFERENCES veterinario(id_veterinario),
  fecha_hora     TIMESTAMP NOT NULL,
  estado         TEXT NOT NULL DEFAULT 'PROGRAMADA'
                 CHECK (estado IN ('PROGRAMADA','ATENDIDA','CANCELADA'))
);

-- Los 6 duenos, 4 veterinarios, 8 mascotas y 10 citas con nombre propio de VetCare.
INSERT INTO dueno (nombre) VALUES
  ('Ana Gomez'), ('Carlos Ruiz'), ('Marcela Diaz'),
  ('Jorge Pineda'), ('Luisa Cardona'), ('Andres Vallejo');
INSERT INTO veterinario (nombre, especialidad) VALUES
  ('Laura Restrepo','General'), ('Diego Moreno','Cirugia'),
  ('Paula Salazar','Dermatologia'), ('Ivan Ortiz','General');
INSERT INTO mascota (id_dueno, nombre, especie, activa) VALUES
  (1,'Firulais','Canino','S'), (1,'Luna','Felino','S'), (2,'Rocky','Canino','N'),
  (3,'Mishi','Felino','S'),    (3,'Bobby','Canino','S'), (4,'Nube','Felino','S'),
  (5,'Toby','Canino','S'),     (6,'Kiara','Canino','N');
INSERT INTO cita (id_mascota, id_veterinario, fecha_hora, estado) VALUES
  (1,1,TIMESTAMP '2026-09-01 08:00','PROGRAMADA'), (2,1,TIMESTAMP '2026-09-01 09:00','ATENDIDA'),
  (4,2,TIMESTAMP '2026-09-01 10:00','PROGRAMADA'), (5,3,TIMESTAMP '2026-09-02 08:30','CANCELADA'),
  (6,2,TIMESTAMP '2026-09-02 11:00','ATENDIDA'),   (7,4,TIMESTAMP '2026-09-03 07:45','PROGRAMADA'),
  (1,1,TIMESTAMP '2026-09-05 15:00','ATENDIDA'),   (2,3,TIMESTAMP '2026-09-08 16:00','PROGRAMADA'),
  (4,4,TIMESTAMP '2026-09-10 08:00','PROGRAMADA'), (6,1,TIMESTAMP '2026-09-10 09:00','ATENDIDA');

-- Volumen. Sin esto la demo no se puede hacer: con 10 citas todo cabe en una
-- pagina de 8 KB y no hay plan mas barato que leerla, asi que la consulta pesima
-- y la optima miden lo mismo y la diferencia se esconde en el ruido de medicion.
INSERT INTO dueno (nombre, telefono, email)
SELECT 'Dueno '||g, '300'||LPAD(g::text,7,'0'), 'dueno'||g||'@mail.com'
FROM generate_series(1,2000) AS g;                       -- duenos 7..2006

INSERT INTO veterinario (nombre, especialidad)
SELECT 'Veterinario '||g,
       CASE WHEN g%3=0 THEN 'Cirugia' WHEN g%3=1 THEN 'General' ELSE 'Dermatologia' END
FROM generate_series(1,12) AS g;                         -- veterinarios 5..16

INSERT INTO mascota (id_dueno, nombre, especie, activa)
SELECT 1+(g%2000), 'Mascota '||g,
       CASE WHEN g%2=0 THEN 'Canino' ELSE 'Felino' END,
       CASE WHEN g%17=0 THEN 'N' ELSE 'S' END
FROM generate_series(1,5000) AS g;                       -- mascotas 9..5008

INSERT INTO cita (id_mascota, id_veterinario, fecha_hora, estado)
SELECT 1+(g%5000), 1+(g%12),
       TIMESTAMP '2026-01-05 08:00' + ((g%200)*INTERVAL '1 day')
                                    + ((g%9)*INTERVAL '45 minutes'),
       CASE WHEN g%11=0 THEN 'CANCELADA' WHEN g%3=0 THEN 'ATENDIDA' ELSE 'PROGRAMADA' END
FROM generate_series(1,30000) AS g;                      -- citas 11..30010

-- Sin ANALYZE el optimizador trabaja con estimaciones por omision y el «estimado
-- contra real» de la pregunta 2 sale disparatado por una razon que no es el tema.
ANALYZE dueno; ANALYZE mascota; ANALYZE veterinario; ANALYZE cita;

-- Cifras de control que conviene proyectar antes de empezar (200 dias × 150 citas):
--   30.010 citas · 150 el 2026-03-10 · de esas 91 PROGRAMADA, 45 ATENDIDA, 14 CANCELADA.
SELECT COUNT(*) AS total_citas FROM cita;
SELECT estado, COUNT(*) FROM cita
WHERE fecha_hora >= TIMESTAMP '2026-03-10' AND fecha_hora < TIMESTAMP '2026-03-11'
GROUP BY estado ORDER BY estado;

-- ============================================================================
-- BLOQUE 1 · La agenda del dia: los 4 antipatrones (pregunta 1 del taller)
-- ============================================================================

-- ANTES. Cuatro defectos: SELECT * · joins con coma · to_char() sobre la columna
-- · UPPER() sobre el estado. Devuelve 91 filas.
SELECT *
FROM cita c, mascota m, dueno d, veterinario v
WHERE c.id_mascota = m.id_mascota
  AND m.id_dueno = d.id_dueno
  AND c.id_veterinario = v.id_veterinario
  AND to_char(c.fecha_hora,'YYYY-MM-DD') = '2026-03-10'
  AND UPPER(c.estado) = 'PROGRAMADA';

-- DESPUES. Proyeccion de 6 columnas · JOIN ... ON · predicado de RANGO (sargable)
-- · comparacion directa del estado, que el CHECK ya normalizo. Tambien 91 filas.
SELECT c.id_cita, c.fecha_hora, m.nombre AS mascota, d.nombre AS dueno,
       v.nombre AS veterinario, c.estado
FROM cita c
JOIN mascota m     ON m.id_mascota = c.id_mascota
JOIN dueno d       ON d.id_dueno = m.id_dueno
JOIN veterinario v ON v.id_veterinario = c.id_veterinario
WHERE c.fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
  AND c.fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
  AND c.estado = 'PROGRAMADA'
ORDER BY c.fecha_hora;

-- La evidencia (pregunta 2). Se lee: nodo mas costoso, rows= estimadas frente a
-- actual rows=, y Execution Time. Ojo: `actual time` es POR VUELTA y el tiempo de
-- un nodo INCLUYE el de sus hijos.
EXPLAIN (ANALYZE, BUFFERS)
SELECT *
FROM cita c, mascota m, dueno d, veterinario v
WHERE c.id_mascota = m.id_mascota
  AND m.id_dueno = d.id_dueno
  AND c.id_veterinario = v.id_veterinario
  AND to_char(c.fecha_hora,'YYYY-MM-DD') = '2026-03-10'
  AND UPPER(c.estado) = 'PROGRAMADA';

EXPLAIN (ANALYZE, BUFFERS)
SELECT c.id_cita, c.fecha_hora, m.nombre AS mascota, d.nombre AS dueno,
       v.nombre AS veterinario, c.estado
FROM cita c
JOIN mascota m     ON m.id_mascota = c.id_mascota
JOIN dueno d       ON d.id_dueno = m.id_dueno
JOIN veterinario v ON v.id_veterinario = c.id_veterinario
WHERE c.fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
  AND c.fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
  AND c.estado = 'PROGRAMADA'
ORDER BY c.fecha_hora;

-- Lo que la pantalla de agenda realmente necesita. El LIMIT deja de leer en
-- cuanto tiene 50 filas: por eso baja el tiempo aunque el plan sea el mismo.
EXPLAIN ANALYZE
SELECT c.id_cita, c.fecha_hora, m.nombre AS mascota, d.nombre AS dueno,
       v.nombre AS veterinario, c.estado
FROM cita c
JOIN mascota m     ON m.id_mascota = c.id_mascota
JOIN dueno d       ON d.id_dueno = m.id_dueno
JOIN veterinario v ON v.id_veterinario = c.id_veterinario
WHERE c.fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
  AND c.fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
  AND c.estado = 'PROGRAMADA'
ORDER BY c.fecha_hora
LIMIT 50;

-- ============================================================================
-- BLOQUE 2 · La subconsulta correlacionada (pregunta 3 del taller)
-- ============================================================================

-- ANTES. La subconsulta esta en la LISTA DE COLUMNAS y menciona d.id_dueno, del
-- exterior: no se puede calcular una vez y reusar. El plan lo delata con un nodo
-- SubPlan y loops=2006 — un dueno, una ejecucion.
EXPLAIN ANALYZE
SELECT d.id_dueno, d.nombre,
       (SELECT COUNT(*) FROM cita c JOIN mascota m ON m.id_mascota = c.id_mascota
         WHERE m.id_dueno = d.id_dueno) AS total_citas
FROM dueno d
ORDER BY total_citas DESC;

-- DESPUES. Una sola pasada: el SubPlan desaparece y queda un HashAggregate.
-- COUNT(c.id_cita) y NO COUNT(*): el LEFT JOIN fabrica una fila de NULL por cada
-- dueno sin citas, y COUNT(*) cuenta filas, asi que reportaria 1 donde va 0.
-- Y LEFT y no INNER: el INNER es mas rapido y borra del ranking a los 6 duenos
-- sin mascotas (2001..2006). Mas rapido devolviendo otra cosa no es optimizar.
EXPLAIN ANALYZE
SELECT d.id_dueno, d.nombre, COUNT(c.id_cita) AS total_citas
FROM dueno d
LEFT JOIN mascota m ON m.id_dueno = d.id_dueno
LEFT JOIN cita c    ON c.id_mascota = m.id_mascota
GROUP BY d.id_dueno, d.nombre
ORDER BY total_citas DESC, d.id_dueno
LIMIT 20;

-- ============================================================================
-- BLOQUE 3 · Optimizar no cambio el resultado: la prueba
-- ============================================================================

-- Prueba 1 · los dos COUNT(*) de la agenda, en la misma corrida. Las dos columnas
-- tienen que decir 91.
SELECT (SELECT COUNT(*) FROM cita c, mascota m, dueno d, veterinario v
         WHERE c.id_mascota = m.id_mascota AND m.id_dueno = d.id_dueno
           AND c.id_veterinario = v.id_veterinario
           AND to_char(c.fecha_hora,'YYYY-MM-DD') = '2026-03-10'
           AND UPPER(c.estado) = 'PROGRAMADA')                       AS filas_antes,
       (SELECT COUNT(*) FROM cita c
          JOIN mascota m ON m.id_mascota = c.id_mascota
          JOIN dueno d ON d.id_dueno = m.id_dueno
          JOIN veterinario v ON v.id_veterinario = c.id_veterinario
         WHERE c.fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
           AND c.fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
           AND c.estado = 'PROGRAMADA')                              AS filas_despues;

-- Prueba 2 · EXCEPT en los DOS sentidos, sin LIMIT. A EXCEPT B vacio NO prueba la
-- igualdad: B puede traer filas de mas. Tiene que devolver CERO filas.
WITH antes AS (
  SELECT d.id_dueno,
         (SELECT COUNT(*) FROM cita c JOIN mascota m ON m.id_mascota = c.id_mascota
           WHERE m.id_dueno = d.id_dueno) AS total_citas
  FROM dueno d
), despues AS (
  SELECT d.id_dueno, COUNT(c.id_cita) AS total_citas
  FROM dueno d
  LEFT JOIN mascota m ON m.id_dueno = d.id_dueno
  LEFT JOIN cita c    ON c.id_mascota = m.id_mascota
  GROUP BY d.id_dueno
)
SELECT 'sobra en ANTES' AS lado, * FROM (SELECT * FROM antes EXCEPT SELECT * FROM despues) a
UNION ALL
SELECT 'sobra en DESPUES', * FROM (SELECT * FROM despues EXCEPT SELECT * FROM antes) b;

-- El contraejemplo que vale la pena proyectar 30 segundos: con COUNT(*) en vez de
-- COUNT(c.id_cita), estas 6 filas dicen 1 y la respuesta correcta es 0.
SELECT d.id_dueno, COUNT(c.id_cita) AS bien, COUNT(*) AS mal
FROM dueno d
LEFT JOIN mascota m ON m.id_dueno = d.id_dueno
LEFT JOIN cita c    ON c.id_mascota = m.id_mascota
WHERE d.id_dueno BETWEEN 2001 AND 2006
GROUP BY d.id_dueno ORDER BY d.id_dueno;

-- Lo que NO se puede medir aqui, y hay que decirlo (es la seccion 5 de la
-- pregunta 5): tiempos con la memoria intermedia vacia —vaciarla exige
-- privilegios de administrador—, concurrencia (eso es la Clase 10) y cualquier
-- comparacion por encima de unos cientos de miles de filas.
""",
"07_indices_vetcare.sql": """-- VetCare DB · Clase 7 · Indices y particionamiento
-- Ejecutable en PostgreSQL, incluido PGlite (la consola de ExamLab). Corre completo y
-- EN ORDEN: el valor de la clase esta en el antes/despues, no en el CREATE INDEX.
--
-- Los CINCO nombres de indice de aqui son los EXACTOS que califica la actividad. No los
-- cambie: el plan de ejecucion imprime "Index Scan using <nombre>" y la tabla de
-- justificacion de la pregunta 5 se llena con estos nombres.
--
-- ATENCION: el BLOQUE 0 recrea las tablas desde cero. Correlo en una base vacia o en la
-- consola de ExamLab, no sobre una VetCare DB con datos que quiera conservar. Si ya tiene
-- las 30.010 citas cargadas, salte al BLOQUE 1.

-- =====================================================================
-- BLOQUE 0 · Volumen. Con 50 filas el planeador prefiere Seq Scan por
-- muchos indices que existan: sin volumen esta clase no se puede medir.
-- Reproduce la siembra sintetica de la actividad: 30.000 citas del
-- 2026-01-05 al 2026-07-23, 5.000 mascotas, 2.000 duenos, 12 veterinarios.
-- (En ExamLab hay 10 citas mas puestas a mano en septiembre: 30.010.)
-- =====================================================================
DROP TABLE IF EXISTS cita_hist;
DROP TABLE IF EXISTS cita;
DROP TABLE IF EXISTS mascota;
DROP TABLE IF EXISTS veterinario;
DROP TABLE IF EXISTS dueno;

CREATE TABLE dueno (
  id_dueno SERIAL PRIMARY KEY,
  nombre   TEXT NOT NULL,
  ciudad   TEXT DEFAULT 'Cali'
);
CREATE TABLE veterinario (
  id_veterinario SERIAL PRIMARY KEY,
  nombre         TEXT NOT NULL,
  especialidad   TEXT
);
CREATE TABLE mascota (
  id_mascota SERIAL PRIMARY KEY,
  id_dueno   INT NOT NULL REFERENCES dueno(id_dueno),
  nombre     TEXT NOT NULL,
  especie    TEXT NOT NULL
);
CREATE TABLE cita (
  id_cita        SERIAL PRIMARY KEY,
  id_mascota     INT NOT NULL REFERENCES mascota(id_mascota),
  id_veterinario INT NOT NULL REFERENCES veterinario(id_veterinario),
  fecha_hora     TIMESTAMP NOT NULL,
  estado         TEXT NOT NULL DEFAULT 'PROGRAMADA'
    CHECK (estado IN ('PROGRAMADA','ATENDIDA','CANCELADA'))
);

INSERT INTO dueno (nombre) SELECT 'Dueno ' || g FROM generate_series(1, 2000) AS g;
INSERT INTO veterinario (nombre, especialidad)
SELECT 'Veterinario ' || g,
       CASE WHEN g % 3 = 0 THEN 'Cirugia'
            WHEN g % 3 = 1 THEN 'General'
            ELSE 'Dermatologia' END
FROM generate_series(1, 12) AS g;
INSERT INTO mascota (id_dueno, nombre, especie)
SELECT 1 + (g % 2000), 'Mascota ' || g,
       CASE WHEN g % 2 = 0 THEN 'Canino' ELSE 'Felino' END
FROM generate_series(1, 5000) AS g;
INSERT INTO cita (id_mascota, id_veterinario, fecha_hora, estado)
SELECT 1 + (g % 5000),
       1 + (g % 12),
       TIMESTAMP '2026-01-05 08:00:00'
         + ((g % 200) * INTERVAL '1 day')
         + ((g % 9) * INTERVAL '45 minutes'),
       CASE WHEN g % 11 = 0 THEN 'CANCELADA'
            WHEN g % 3  = 0 THEN 'ATENDIDA'
            ELSE 'PROGRAMADA' END
FROM generate_series(1, 30000) AS g;

ANALYZE dueno;  ANALYZE veterinario;  ANALYZE mascota;  ANALYZE cita;

-- Control: 30.000 | 18.182 PROGRAMADA | 9.091 ATENDIDA | 2.727 CANCELADA. En la base de
-- ExamLab hay 10 citas mas sembradas a mano, y ahi el reparto es 30.010 / 18.187 / 9.095 /
-- 2.728. Si su corrida da otros numeros, el resto del script no cuadra.
SELECT estado, COUNT(*) FROM cita GROUP BY estado ORDER BY estado;

-- =====================================================================
-- BLOQUE 1 · LINEA BASE. Sin este paso no hay clase: el "despues" solo
-- significa algo contra un "antes" medido. Tiene que salir Seq Scan.
-- =====================================================================
EXPLAIN ANALYZE   -- C1 · agenda del dia (rango de fecha + estado)
SELECT id_cita, fecha_hora, estado
  FROM cita
 WHERE fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
   AND fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
   AND estado = 'PROGRAMADA';
-- Esperado: Seq Scan on cita, filas devueltas = 91 (de 150 citas ese dia).

EXPLAIN ANALYZE   -- C2 · mascotas de un dueno
SELECT id_mascota, nombre, especie FROM mascota WHERE id_dueno = 1234;
-- Esperado: Seq Scan on mascota, 2 filas devueltas (id_dueno = 1 + (g % 2000) hace que solo
-- las mascotas g=1233 y g=3233 caigan en el dueno 1234). La FK NO crea indice sola en PostgreSQL.

-- =====================================================================
-- BLOQUE 2 · LOS TRES INDICES DE LA PREGUNTA 1
-- =====================================================================
-- (a) Simple: sirve a cualquier consulta por rango de fecha, con o sin estado.
CREATE INDEX idx_cita_fecha_hora ON cita (fecha_hora);

-- (b) Sobre la FK: "las mascotas de un dueno", y ademas abarata el borrado de un dueno.
CREATE INDEX idx_mascota_dueno ON mascota (id_dueno);

-- (c) PARCIAL: el WHERE es parte de la DEFINICION del indice, no de la consulta. Indexa
--     18.182 de las 30.000 de este script (18.187 de 30.010 en ExamLab) porque la pantalla
--     de agenda nunca pregunta por atendidas ni por canceladas.
CREATE INDEX idx_cita_programada_fecha ON cita (fecha_hora) WHERE estado = 'PROGRAMADA';

-- El paso que se salta la mitad del salon. Crear el indice NO actualiza estadisticas.
ANALYZE cita;
ANALYZE mascota;

-- Las MISMAS dos consultas, sin cambiar una coma.
EXPLAIN ANALYZE
SELECT id_cita, fecha_hora, estado
  FROM cita
 WHERE fecha_hora >= TIMESTAMP '2026-03-10 00:00:00'
   AND fecha_hora <  TIMESTAMP '2026-03-11 00:00:00'
   AND estado = 'PROGRAMADA';
-- Esperado: Index Scan using idx_cita_programada_fecha (gana el PARCIAL: recorre 91
-- entradas y ya sabe que todas cumplen el estado; el completo recorreria 150 y tendria
-- que descartar 59 despues de leer la tabla). Reporte el que VEA, no el que diga esto.

EXPLAIN ANALYZE
SELECT id_mascota, nombre, especie FROM mascota WHERE id_dueno = 1234;
-- Esperado: Index Scan (o Bitmap Index Scan) using idx_mascota_dueno.

-- Evidencia de que existen. indexdef devuelve el CREATE INDEX completo, asi que aqui se
-- ve tambien el WHERE del parcial.
SELECT indexname, tablename, indexdef
  FROM pg_indexes
 WHERE tablename IN ('cita','mascota')
 ORDER BY tablename, indexname;

-- =====================================================================
-- BLOQUE 3 · ORDEN DE COLUMNAS (pregunta 2). Los dos indices llevan las
-- MISMAS dos columnas en orden inverso, y existen para demostrar que el
-- orden decide. Regla: igualdad primero, rango al final.
-- =====================================================================
CREATE INDEX idx_cita_estado_fecha ON cita (estado, fecha_hora);
CREATE INDEX idx_cita_fecha_estado ON cita (fecha_hora, estado);
ANALYZE cita;

EXPLAIN ANALYZE   -- Q1 · estado (igualdad) + fecha (rango) -> favorece (estado, fecha_hora)
SELECT id_cita, fecha_hora FROM cita
 WHERE estado = 'PROGRAMADA'
   AND fecha_hora >= TIMESTAMP '2026-03-01' AND fecha_hora < TIMESTAMP '2026-04-01';

EXPLAIN ANALYZE   -- Q2 · solo rango de fecha -> favorece (fecha_hora, estado)
SELECT id_cita, estado FROM cita
 WHERE fecha_hora >= TIMESTAMP '2026-03-01' AND fecha_hora < TIMESTAMP '2026-04-01';

EXPLAIN ANALYZE   -- Q3 · solo estado, sin fecha -> columna lider ausente en el de fecha
SELECT COUNT(*) FROM cita WHERE estado = 'CANCELADA';

-- Fuerce el experimento: quite el que Q2 estaba usando y vuelva a medir.
DROP INDEX idx_cita_fecha_estado;
ANALYZE cita;
EXPLAIN ANALYZE
SELECT id_cita, estado FROM cita
 WHERE fecha_hora >= TIMESTAMP '2026-03-01' AND fecha_hora < TIMESTAMP '2026-04-01';
-- Esperado: cae en idx_cita_fecha_hora o vuelve a Seq Scan, pero NO usa
-- idx_cita_estado_fecha: su columna lider (estado) no aparece en el WHERE.

-- =====================================================================
-- BLOQUE 4 · PARTICIONAMIENTO (pregunta 3). HOY SE IMPLEMENTA.
-- =====================================================================
-- La trampa: en una tabla particionada la PK DEBE incluir la columna de particion.
-- PRIMARY KEY (id_cita) a secas no compila, y el mensaje del motor no lo dice asi.
CREATE TABLE cita_hist (
  id_cita        INT,
  id_mascota     INT,
  id_veterinario INT,
  fecha_hora     TIMESTAMP NOT NULL,
  estado         TEXT,
  PRIMARY KEY (id_cita, fecha_hora)
) PARTITION BY RANGE (fecha_hora);

-- Rango cerrado por abajo, abierto por arriba: el TO de una es el FROM de la siguiente.
CREATE TABLE cita_hist_2025 PARTITION OF cita_hist
  FOR VALUES FROM (TIMESTAMP '2025-01-01') TO (TIMESTAMP '2026-01-01');
CREATE TABLE cita_hist_2026 PARTITION OF cita_hist
  FOR VALUES FROM (TIMESTAMP '2026-01-01') TO (TIMESTAMP '2027-01-01');

INSERT INTO cita_hist
SELECT id_cita, id_mascota, id_veterinario, fecha_hora, estado FROM cita;

-- Prueba del enrutamiento. tableoid es la columna de sistema que dice en que tabla FISICA
-- vive cada fila; ::regclass la traduce a nombre. Sin esto no hay evidencia: solo un
-- INSERT que no dio error.
SELECT tableoid::regclass AS particion, COUNT(*), MIN(fecha_hora), MAX(fecha_hora)
  FROM cita_hist GROUP BY 1 ORDER BY 1;
-- Con la siembra del BLOQUE 0 (todas las citas son de 2026) cae TODO en cita_hist_2026 y
-- cita_hist_2025 queda vacia: eso ya demuestra el enrutamiento. La base de la pregunta 3
-- en ExamLab reparte 5.010 citas entre 2025 y 2026 y ahi se ven las dos particiones.

-- Poda de particiones: lo unico que mejora hoy de verdad.
EXPLAIN ANALYZE
SELECT COUNT(*) FROM cita_hist
 WHERE fecha_hora >= TIMESTAMP '2026-01-01' AND fecha_hora < TIMESTAMP '2027-01-01';
-- Esperado: en el plan aparece SOLO cita_hist_2026. El tiempo no baja de forma apreciable
-- con este volumen, y hay que decirlo: lo que se demuestra es que el motor descarta
-- particiones enteras ANTES de leer.

-- El beneficio real es de mantenimiento: archivar un ano es DROP TABLE de su particion
-- --una operacion de metadatos-- en vez de un DELETE masivo que toca millones de filas,
-- infla el registro de transacciones y sostiene bloqueos largos. Eso es la Clase 8.
-- DROP TABLE cita_hist_2025;
""",
"08_transacciones_vetcare.sql": """-- VetCare DB · Clase 8 · Transaccion de facturacion + descuento de stock
-- Ejecutable en PostgreSQL, incluido PGlite (la consola de ExamLab). Corre completo y
-- EN ORDEN: el bloque 3 solo tiene sentido si antes se tomo la foto del bloque 2.
--
-- ESTO ES PL/pgSQL, NO PL/SQL DE ORACLE. Aqui no existen NUMBER, SQL%ROWCOUNT ni
-- RAISE_APPLICATION_ERROR, y NO se escribe COMMIT ni ROLLBACK dentro del procedimiento:
-- el CALL de nivel superior ya es su propia transaccion.

-- =====================================================================
-- BLOQUE 0 · Esquema minimo y datos. Los stocks son los de la actividad.
-- =====================================================================
-- Los DROP van primero para que el script se pueda correr dos veces sin limpiar a mano.
DROP PROCEDURE IF EXISTS sp_facturar(INT, INT[], INT[]);
DROP FUNCTION  IF EXISTS fn_descontar_stock(INT, INT);
DROP TABLE     IF EXISTS detalle_factura;
DROP TABLE     IF EXISTS factura;
DROP TABLE     IF EXISTS insumo;

CREATE TABLE insumo (
  id_insumo   SERIAL PRIMARY KEY,
  nombre      TEXT NOT NULL,
  stock       INT NOT NULL CHECK (stock >= 0),
  precio_unit NUMERIC(12,2) NOT NULL
);
CREATE TABLE factura (
  id_factura  SERIAL PRIMARY KEY,
  id_consulta INT NOT NULL,
  total       NUMERIC(12,2) NOT NULL DEFAULT 0
);
CREATE TABLE detalle_factura (
  id_detalle  SERIAL PRIMARY KEY,
  id_factura  INT NOT NULL REFERENCES factura(id_factura),
  id_insumo   INT NOT NULL REFERENCES insumo(id_insumo),
  cantidad    INT NOT NULL CHECK (cantidad > 0),
  precio_unit NUMERIC(12,2) NOT NULL
);

INSERT INTO insumo (nombre, stock, precio_unit) VALUES
  ('Vacuna antirrabica',   12, 22000),   -- 1
  ('Vacuna triple felina',  3, 31000),   -- 2  <- el que se va a quedar corto
  ('Antiparasitario oral', 40,  9500),   -- 3
  ('Suero fisiologico',    25,  7000),   -- 4
  ('Gasa esteril',          8,  1200),   -- 5
  ('Jeringa 5ml',          60,   900);   -- 6

-- =====================================================================
-- BLOQUE 1 · EL PROCEDIMIENTO. Una factura tiene VARIAS lineas, asi que
-- la firma recibe dos arreglos paralelos, no un insumo suelto.
-- =====================================================================
CREATE PROCEDURE sp_facturar(
  p_id_consulta INT,
  p_insumos     INT[],
  p_cantidades  INT[]
)
LANGUAGE plpgsql
AS $proc$
DECLARE
  v_id_factura INT;
  v_total   NUMERIC(12,2) := 0;
  v_precio  NUMERIC(12,2);
  v_filas   INT;
  i         INT;
BEGIN
  -- El llamador se equivoco: se rechaza antes de tocar la base.
  IF array_length(p_insumos, 1) IS DISTINCT FROM array_length(p_cantidades, 1) THEN
    RAISE EXCEPTION 'ERROR: insumos y cantidades deben tener la misma longitud';
  END IF;

  -- Total en 0: todavia no se sabe. RETURNING ... INTO evita otro SELECT.
  INSERT INTO factura (id_consulta, total) VALUES (p_id_consulta, 0)
  RETURNING id_factura INTO v_id_factura;

  FOR i IN 1 .. array_length(p_insumos, 1) LOOP
    SELECT precio_unit INTO v_precio FROM insumo WHERE id_insumo = p_insumos[i];
    IF NOT FOUND THEN
      RAISE EXCEPTION 'ERROR: el insumo % no existe', p_insumos[i];
    END IF;

    -- EL GUARDIA. La comprobacion viaja DENTRO del WHERE: comprobar y escribir son
    -- una sola sentencia atomica y nadie puede colarse entre las dos.
    UPDATE insumo
       SET stock = stock - p_cantidades[i]
     WHERE id_insumo = p_insumos[i]
       AND stock >= p_cantidades[i];
    GET DIAGNOSTICS v_filas = ROW_COUNT;   -- 1 alcanzo, 0 no habia stock
    IF v_filas = 0 THEN
      RAISE EXCEPTION 'ERROR: stock insuficiente del insumo % (se pidieron %)',
        p_insumos[i], p_cantidades[i];
    END IF;

    INSERT INTO detalle_factura (id_factura, id_insumo, cantidad, precio_unit)
    VALUES (v_id_factura, p_insumos[i], p_cantidades[i], v_precio);

    v_total := v_total + (v_precio * p_cantidades[i]);
  END LOOP;

  UPDATE factura SET total = v_total WHERE id_factura = v_id_factura;
  RAISE NOTICE 'Factura % creada por %', v_id_factura, v_total;
END;
$proc$;

-- =====================================================================
-- BLOQUE 2 · CASO EXITOSO
-- =====================================================================
CALL sp_facturar(4, ARRAY[1, 6, 5], ARRAY[1, 2, 3]);
-- Esperado: 22000*1 + 900*2 + 1200*3 = 27.400, y los stocks 1, 6 y 5 bajan a 11, 58 y 5.
SELECT id_factura, id_consulta, total FROM factura ORDER BY id_factura;
SELECT id_insumo, nombre, stock FROM insumo ORDER BY id_insumo;

-- =====================================================================
-- BLOQUE 3 · ATOMICIDAD. Aqui esta la clase entera.
-- =====================================================================
-- Foto inicial: estos numeros son el punto de comparacion.
SELECT (SELECT COUNT(*) FROM factura)         AS facturas,
       (SELECT COUNT(*) FROM detalle_factura) AS lineas,
       (SELECT stock FROM insumo WHERE id_insumo = 3) AS stock_3,
       (SELECT stock FROM insumo WHERE id_insumo = 2) AS stock_2;
-- Esperado tras el bloque 2: 1 | 3 | 40 | 3

-- Intento que falla A MITAD: la primera linea (2 del insumo 3, que tiene 40) SI alcanza;
-- la segunda (10 del insumo 2, que solo tiene 3) NO. El DO ... EXCEPTION es para que el
-- script no se detenga; el que decide sigue siendo el procedimiento.
DO $$
BEGIN
  CALL sp_facturar(4, ARRAY[3, 2], ARRAY[2, 10]);
  RAISE NOTICE 'No deberia llegar aqui';
EXCEPTION WHEN OTHERS THEN
  RAISE NOTICE 'Fallo esperado: %', SQLERRM;
END $$;

-- Foto final: EXACTAMENTE la misma consulta.
SELECT (SELECT COUNT(*) FROM factura)         AS facturas,
       (SELECT COUNT(*) FROM detalle_factura) AS lineas,
       (SELECT stock FROM insumo WHERE id_insumo = 3) AS stock_3,
       (SELECT stock FROM insumo WHERE id_insumo = 2) AS stock_2;
-- Esperado: 1 | 3 | 40 | 3, identico a la foto inicial.
--   * no quedo una factura huerfana,
--   * no quedo ninguna linea de detalle,
--   * y sobre todo el stock del insumo 3 VOLVIO A 40: el descuento que si habia
--     alcanzado se deshizo. Nadie escribio ROLLBACK.

-- Y ahora la misma factura con una cantidad viable del insumo 2.
CALL sp_facturar(4, ARRAY[3, 2], ARRAY[2, 3]);
SELECT id_factura, id_consulta, total FROM factura ORDER BY id_factura;
SELECT id_insumo, nombre, stock FROM insumo ORDER BY id_insumo;
-- Esperado: factura 2 por 9500*2 + 31000*3 = 112.000; insumo 3 en 38 e insumo 2 en 0.

-- =====================================================================
-- BLOQUE 4 · EL MISMO PATRON COMO FUNCION REUTILIZABLE
-- Aqui "no hay stock" es una RESPUESTA, no un error: la funcion informa y
-- el llamador decide. El procedimiento del bloque 1 abortaba.
-- =====================================================================
CREATE FUNCTION fn_descontar_stock(p_id_insumo INT, p_cantidad INT)
RETURNS BOOLEAN
LANGUAGE plpgsql
AS $fn$
DECLARE
  v_filas INT;
BEGIN
  -- Una cantidad no positiva no es "no hay stock", es una llamada mal hecha.
  IF p_cantidad <= 0 THEN
    RAISE EXCEPTION 'ERROR: la cantidad debe ser positiva (llego %)', p_cantidad;
  END IF;

  UPDATE insumo
     SET stock = stock - p_cantidad
   WHERE id_insumo = p_id_insumo
     AND stock >= p_cantidad;
  GET DIAGNOSTICS v_filas = ROW_COUNT;

  RETURN v_filas = 1;
END;
$fn$;

-- Reiniciar los stocks para que la prueba de abajo de los valores esperados.
UPDATE insumo SET stock = 8 WHERE id_insumo = 5;
UPDATE insumo SET stock = 3 WHERE id_insumo = 2;

SELECT fn_descontar_stock(5, 3)  AS caso_ok,
       fn_descontar_stock(2, 10) AS caso_sin_stock,
       fn_descontar_stock(2, 3)  AS caso_limite;
-- Esperado: true | false | true. El tercero es el interesante: pide EXACTAMENTE el stock
-- que queda, y con >= en el guardia tiene que pasar.

SELECT id_insumo, nombre, stock FROM insumo ORDER BY id_insumo;
-- Esperado: insumo 5 en 5, insumo 2 en 0, y NINGUN stock negativo.

-- La diferencia con leer primero y decidir despues:
--   SELECT stock ... ; IF stock >= cantidad THEN UPDATE ...
-- deja una VENTANA entre la lectura y la escritura. Con dos recepcionistas facturando el
-- mismo insumo, las dos leen 3, las dos deciden que alcanza, y el stock termina en -2 (o
-- el CHECK revienta). El UPDATE con la condicion en el WHERE no tiene ventana.
-- Aqui no se puede demostrar: PGlite corre UNA SOLA sesion. Ese es el gap que se declara
-- en la pregunta 5 y lo que abre la Clase 10.
""",
"10_concurrencia_vetcare.sql": """-- VetCare DB · Clase 10 · Demo ejecutable: doble reserva y su mitigacion
-- Ejecutar EN ORDEN: primero se ve el problema, despues la solucion.

-- Paso 1: tabla de demo SIN restriccion (asi llegaria si nadie penso en concurrencia)
CREATE TABLE cita_demo (
  id_cita INT PRIMARY KEY,
  id_mascota INT NOT NULL,
  id_veterinario INT NOT NULL,
  fecha_hora TIMESTAMP NOT NULL,
  estado VARCHAR(20) DEFAULT 'PROGRAMADA'
);

-- Paso 2: T1 (Recepcion A) agenda la franja - OK
INSERT INTO cita_demo VALUES (1, 10, 5, TIMESTAMP '2026-10-12 09:00:00', 'PROGRAMADA');

-- Paso 3: T2 (Recepcion B) agenda OTRA mascota, MISMO veterinario, MISMA franja.
-- Sin restriccion esto se inserta SIN ERROR -> aqui esta la doble reserva.
INSERT INTO cita_demo VALUES (2, 22, 5, TIMESTAMP '2026-10-12 09:00:00', 'PROGRAMADA');

-- Evidencia del problema: dos citas para el mismo veterinario en la misma franja
SELECT id_veterinario, fecha_hora, COUNT(*) AS citas_en_la_misma_franja
FROM cita_demo
GROUP BY id_veterinario, fecha_hora
HAVING COUNT(*) > 1;

-- Paso 4: la mitigacion real - la restriccion que debio existir desde el diseño
ALTER TABLE cita_demo
  ADD CONSTRAINT uq_cita_demo_vet_fecha UNIQUE (id_veterinario, fecha_hora);

-- Paso 5: repetir el intento de doble reserva - AHORA debe fallar
INSERT INTO cita_demo VALUES (3, 35, 5, TIMESTAMP '2026-10-12 09:00:00', 'PROGRAMADA');
-- Esperado: error de restriccion unica (ORA-00001 en Oracle) -> la BD rechaza la doble reserva.
""",
"11_checklist_seed.sql": """-- VetCare DB · Clase 11 · Seed ejecutable para la demo de checklist
-- Autocontenido: cree estas tablas minimas si aun no las tiene, o
-- adapte los nombres a su propio DDL (Clases 1-8) antes de correr los INSERT.

CREATE TABLE dueno_demo (id_dueno INT PRIMARY KEY, nombre VARCHAR(80));
CREATE TABLE mascota_demo (
  id_mascota INT PRIMARY KEY, id_dueno INT REFERENCES dueno_demo(id_dueno),
  nombre VARCHAR(60), activa CHAR(1) DEFAULT 'S'
);
CREATE TABLE cita_demo11 (
  id_cita INT PRIMARY KEY, id_mascota INT REFERENCES mascota_demo(id_mascota),
  fecha_hora TIMESTAMP, estado VARCHAR(20)
);
CREATE TABLE insumo_demo (id_insumo INT PRIMARY KEY, nombre VARCHAR(60), stock INT);

-- Datos que permiten mostrar EN VIVO cada punto del checklist:
INSERT INTO dueno_demo VALUES (1, 'Ana Perez');
INSERT INTO dueno_demo VALUES (2, 'Carlos Ruiz');
INSERT INTO mascota_demo VALUES (10, 1, 'Luna', 'S');   -- mascota activa: SI puede agendar
INSERT INTO mascota_demo VALUES (11, 2, 'Rocky', 'N');  -- mascota inactiva: NO debe poder agendar
INSERT INTO cita_demo11 VALUES (100, 10, TIMESTAMP '2026-10-19 09:00:00', 'PROGRAMADA');
INSERT INTO insumo_demo VALUES (50, 'Vacuna antirrabica', 3);  -- stock bajo a proposito

-- Punto del checklist "regla de negocio se cumple": intente agendar la mascota
-- inactiva (id 11) con su sp_agendar_cita y confirme que el proc la rechaza.
-- Punto "stock nunca negativo": intente facturar 5 unidades del insumo 50
-- (solo hay 3) y confirme que su transaccion de Clase 8 hace ROLLBACK.
SELECT m.nombre, m.activa, d.nombre AS dueno FROM mascota_demo m JOIN dueno_demo d ON d.id_dueno = m.id_dueno;
""",
"12_contrato_ops.sql": """-- VetCare DB · Clase 12 · Contrato app<->BD (Oracle PL/SQL, ejecutable)
-- Regla: la app NUNCA hace INSERT directo a cita/consulta/factura; solo llama estos procs.

CREATE OR REPLACE PROCEDURE sp_agendar_cita (
  p_id_cita IN NUMBER, p_id_mascota IN NUMBER, p_fecha IN TIMESTAMP, p_msg OUT VARCHAR2
) AS
BEGIN
  INSERT INTO cita(id_cita, id_mascota, fecha_hora, estado) VALUES (p_id_cita, p_id_mascota, p_fecha, 'PROGRAMADA');
  p_msg := 'OK: cita agendada'; COMMIT;
EXCEPTION WHEN OTHERS THEN p_msg := 'ERROR: ' || SQLERRM; ROLLBACK;
END;
/

CREATE OR REPLACE PROCEDURE sp_registrar_consulta (
  p_id_consulta IN NUMBER, p_id_cita IN NUMBER, p_notas IN VARCHAR2, p_precio IN NUMBER, p_msg OUT VARCHAR2
) AS
BEGIN
  INSERT INTO consulta(id_consulta, id_cita, notas, precio) VALUES (p_id_consulta, p_id_cita, p_notas, p_precio);
  p_msg := 'OK: consulta registrada'; COMMIT;
EXCEPTION WHEN OTHERS THEN p_msg := 'ERROR: ' || SQLERRM; ROLLBACK;
END;
/

-- Contrato para la sustentacion (documentar tal cual en el informe):
-- sp_agendar_cita(id_cita, id_mascota, fecha)      -> p_msg: 'OK: ...' | 'ERROR: ...'
-- sp_registrar_consulta(id_consulta, id_cita, notas, precio) -> p_msg idem
-- sp_facturar(id_factura, id_consulta, lineas...)  -> ver Clase 8 (transaccion factura+stock)
""",
}

QUIZ = {
1: [
    q_om("En el modelo VetCare, ¿qué diferencia PK de FK?",
         ["A) No hay diferencia", "B) PK identifica la fila; FK referencia otra PK",
          "C) FK solo sirve para índices", "D) PK es solo texto"], "B"),
    q_om("¿Cuáles son entidades mínimas típicas del PI VetCare?",
         ["A) Solo Factura", "B) Dueño, Mascota, Cita (y relacionadas)",
          "C) Solo logs de servidor", "D) Solo usuarios LMS"], "B"),
    q_vf("El ER del PI se entrega en draw.io (PNG/SVG) como evidencia visual, no solo como párrafo.", "V"),
    q_vf("Normalización 1–3FN se aplica en el PI solo donde aporta al modelo VetCare (no teoría aislada).", "V"),
    q_om("Regla de negocio típica VetCare:",
         ["A) Mascota inactiva puede agendar sin control", "B) Stock puede ser negativo sin auditar",
          "C) Mascota inactiva no agenda; stock ≥ 0; auditoría sensible", "D) Sin reglas"], "C"),
    q_vf("El diagnóstico de Clase 1 evalúa la plataforma de entrega, no saberes previos de BD I.", "F"),
    q_abierta("Nombre 3 entidades de su ER VetCare y una FK entre dos de ellas.",
              "Ej. Mascota.id_dueno → Dueño.id; Cita.id_mascota → Mascota.id."),
    q_abierta("Escriba en una frase el alcance SI / NO de su PI esta semana.",
              "Respuesta alineada a su ficha (qué sí modela / qué queda fuera)."),
],
2: [
    q_om("Un rol con privilegios mínimos en VetCare sirve para:",
         ["A) Dar SUPER a todos", "B) Least privilege por perfil (app/lectura/admin)",
          "C) Evitar backups", "D) Eliminar PK"], "B"),
    q_om("¿Qué privilegio NO debería tener el rol de solo lectura de agenda?",
         ["A) SELECT sobre cita", "B) DROP TABLE", "C) SELECT sobre mascota", "D) Consultar veterinario"], "B"),
    q_vf("GRANT y REVOKE forman parte de la administración de privilegios del PI.", "V"),
    q_vf("Usar el usuario root/admin de playground para la app en producción es buena práctica.", "F"),
    q_om("Evidencia típica del taller de roles:",
         ["A) Solo un screenshot del logo", "B) Script de roles + prueba de permiso denegado",
          "C) Factura AWS", "D) Diagrama sin usuarios"], "B"),
    q_vf("Separar rol de aplicación y rol de DBA reduce el impacto de una filtración de credenciales.", "V"),
    q_abierta("Defina 2 roles VetCare (nombre + privilegio clave de cada uno).",
              "Ej. app_vetcare: EXECUTE procs; lector_agenda: SELECT cita/mascota."),
    q_abierta("¿Qué operación debería fallar con el rol de lectura y cómo lo evidencia?",
              "Ej. DELETE/UPDATE stock → error de privilegio; captura del playground."),
],
3: [
    q_om("¿Para qué sirve un procedimiento almacenado en VetCare?",
         ["A) Solo formatear texto", "B) Centralizar regla de negocio invocable desde la app",
          "C) Reemplazar el ER", "D) Dibujar el diagrama"], "B"),
    q_om("Validación mínima de sp_agendar_cita:",
         ["A) Ninguna", "B) Que la mascota exista y esté activa",
          "C) Solo validar el color del logo", "D) Borrar stock"], "B"),
    q_vf("IN entra al procedimiento; OUT/devuelve resultado al llamador.", "V"),
    q_vf("Es correcto duplicar la validación de negocio solo en la UI y nunca en la BD del PI.", "F"),
    q_om("Ventaja de invocar un proc desde la app VetCare:",
         ["A) Oculta reglas y evita SQL suelto inconsistente", "B) Impide transacciones",
          "C) Elimina índices", "D) Obliga a SELECT *"], "A"),
    q_vf("Un procedimiento puede validar reglas y devolver mensaje/código de error controlado.", "V"),
    q_abierta("Escriba la firma (nombre + 2 parámetros) de su sp_agendar_cita.",
              "Ej. sp_agendar_cita(IN id_mascota, IN fecha_hora, OUT msg)."),
    q_abierta("¿Qué pasa si intentan agendar una mascota inactiva? (comportamiento esperado)",
              "Proc rechaza / no inserta cita / mensaje de error; sin fila inválida."),
],
4: [
    q_om("Función vs procedimiento en el PI:",
         ["A) Son idénticos", "B) Función retorna valor usable en SQL; proc orquesta acciones",
          "C) Proc no puede validar", "D) Función borra tablas siempre"], "B"),
    q_om("Un trigger de cancelación/auditoría sirve para:",
         ["A) Decoración", "B) Auditar cambios sensibles sin depender solo de la app",
          "C) Eliminar backups", "D) Quitar FK"], "B"),
    q_vf("RPO = cantidad máxima de dato que se acepta perder (punto de recuperación).", "V"),
    q_vf("RTO mide el tiempo objetivo para volver a operación tras un incidente.", "V"),
    q_om("Evidencia de respaldo mínima razonable en el PI:",
         ["A) «Ya lo hacemos» sin archivo", "B) Plan RPO/RTO + script/export de prueba",
          "C) Solo meme", "D) Apagar el playground"], "B"),
    q_vf("La seguridad del PI termina al crear un usuario; no hace falta auditar cambios sensibles.", "F"),
    q_abierta("Proponga un trigger VetCare (evento + tabla + qué registra).",
              "Ej. AFTER UPDATE cita → auditar cancelación (quién/cuándo/estado)."),
    q_abierta("Indique RPO y RTO objetivo cualitativos para su VetCare.",
              "Ej. RPO ≤ 24h (export diario); RTO ≤ 4h (restaurar playground/script)."),
],
6: [
    q_om("Anti-patrón de consulta en VetCare:",
         ["A) Filtrar por fecha", "B) SELECT * con producto cartesiano / sin filtro",
          "C) Proyectar columnas", "D) Usar índice selectivo"], "B"),
    q_om("Evidencia de optimización pedida en el PI:",
         ["A) Ninguna", "B) ≥2 consultas antes/después con justificación",
          "C) Solo cambiar el color del ER", "D) Borrar la BD"], "B"),
    q_vf("Proyectar solo columnas necesarias reduce I/O y aclara el plan.", "V"),
    q_vf("EXPLAIN/plan de ejecución no aporta nada al entregable de optimización.", "F"),
    q_om("Consulta típica a optimizar en agenda VetCare:",
         ["A) Citas del día con joins necesarios y filtro de fecha", "B) SELECT * FROM todas las tablas",
          "C) Cross join sin WHERE", "D) Actualizar sin WHERE a propósito"], "A"),
    q_vf("«Antes/después» debe mostrar qué cambió (predicado, proyección, índice) y por qué mejora.", "V"),
    q_abierta("Escriba una consulta VetCare «mala» (anti-patrón) en una línea.",
              "Ej. SELECT * FROM cita, mascota; (sin join/filtro)."),
    q_abierta("¿Qué cambio concreto haría en esa consulta y qué evidencia guardaría?",
              "JOIN explícito + filtro fecha + columnas; captura plan/tiempo playground."),
],
7: [
    q_om("¿Cuándo NO conviene crear un índice?",
         ["A) Filtros frecuentes selectivos", "B) Baja selectividad o escritura intensa sin lectura asociada",
          "C) Agenda por fecha_hora", "D) FK muy consultada"], "B"),
    q_om("Índice típico para agenda del día:",
         ["A) Sobre cita.fecha_hora (y opcionalmente estado)", "B) Sobre un comentario libre nunca filtrado",
          "C) Sobre todas las columnas a la vez sin criterio", "D) Ninguno nunca"], "A"),
    q_vf("Particionamiento aquí se trabaja como concepto (p. ej. rango de fecha); no exige SGBD local de pago.", "V"),
    q_vf("Más índices siempre aceleran inserts/updates de stock y facturas.", "F"),
    q_om("Selectividad alta significa:",
         ["A) Casi todas las filas comparten el valor", "B) El predicado discrimina bien (pocas filas)",
          "C) Que no hay PK", "D) Que hay SELECT *"], "B"),
    q_vf("Documentar por qué se creó (o no) un índice es parte de la evidencia del PI.", "V"),
    q_abierta("Proponga un índice VetCare (tabla + columna(s) + consulta que lo usa).",
              "Ej. INDEX cita(fecha_hora) para agenda del día."),
    q_abierta("¿Particionaría citas por mes? Justifique en 1 frase (sí/no + motivo).",
              "Sí si volumen histórico alto por fecha; no si dataset demo pequeño."),
],
8: [
    q_om("COMMIT en la factura VetCare garantiza:",
         ["A) Solo un SELECT", "B) Atomicidad: factura + detalle + stock juntos",
          "C) Borrar auditorías", "D) Ignorar stock"], "B"),
    q_om("Si falta stock, ROLLBACK debe:",
         ["A) Dejar la factura a medias", "B) Deshacer cambios parciales y conservar consistencia",
          "C) Negar el stock a -100", "D) Borrar dueños"], "B"),
    q_vf("Transacciones cortas reducen bloqueos y problemas de concurrencia.", "V"),
    q_vf("Es aceptable hacer la factura en la app con tres commits independientes sin control.", "F"),
    q_om("Propiedad ACID más visible al facturar + descontar insumo:",
         ["A) Atomicidad", "B) Tipografía", "C) Color del ER", "D) Nombre del repo"], "A"),
    q_vf("Un script demo con BEGIN/COMMIT/ROLLBACK es evidencia válida del hito de transacciones.", "V"),
    q_abierta("Liste los 3 pasos atómicos de su caso de facturación VetCare.",
              "Insert factura + detalles + update stock (o equivalente) en una TX."),
    q_abierta("Describa la prueba de fallo (falta stock) y el estado final esperado.",
              "Fuerza stock insuficiente → ROLLBACK → sin factura huérfana ni stock negativo."),
],
10: [
    q_om("Control de concurrencia busca evitar principalmente:",
         ["A) Diagramas bonitos", "B) Anomalías por lecturas/escrituras concurrentes",
          "C) Tener PK", "D) Usar draw.io"], "B"),
    q_om("Escenario VetCare sensible a concurrencia:",
         ["A) Dos cajeros descontando el mismo insumo", "B) Leer el microcurrículo",
          "C) Cambiar el color del logo", "D) Exportar PNG del ER"], "A"),
    q_vf("Aislar transacciones (niveles/locks) es tema de esta clase autónoma.", "V"),
    q_vf("Si «funciona en mi prueba serial», ya está probada la concurrencia.", "F"),
    q_om("Estrategia conceptual frente a doble descuento de stock:",
         ["A) Ignorar", "B) TX + verificación de stock / bloqueo o control optimista",
          "C) SELECT * sin WHERE", "D) Borrar la tabla"], "B"),
    q_vf("Documentar el escenario concurrente (pasos A/B) es parte del entregable autónomo.", "V"),
    q_abierta("Diseñe un caso A/B: dos sesiones tocando el mismo insumo VetCare.",
              "Sesión A y B leen stock=1; ambas intentan vender → una OK, otra falla/rollback."),
    q_abierta("¿Qué evidencia capturaría en el playground para demostrar el control?",
              "Capturas de ambas sesiones + resultado final de stock/factura."),
],
11: [
    q_om("El checkpoint de avance PI VetCare:",
         ["A) Reemplaza el Parcial 3", "B) Es revisión de evidencias, no la sustentación final",
          "C) Cierra el semestre", "D) Elimina la rúbrica"], "B"),
    q_om("Evidencia mínima esperable en el avance:",
         ["A) Solo narrativa", "B) ER + scripts clave + demo parcial + checklist rúbrica",
          "C) Solo el logo", "D) Factura cloud de pago"], "B"),
    q_vf("Si falta el contrato app↔BD, conviene marcarlo como pendiente en el checklist.", "V"),
    q_vf("Seed de datos de prueba no aporta a la demo del PI.", "F"),
    q_om("Durante la viva/demo parcial, lo más útil es:",
         ["A) Improvisar sin script", "B) Mostrar flujo real (agendar/facturar) con datos seed",
          "C) Leer el microcurrículo completo", "D) Ocultar errores siempre"], "B"),
    q_vf("Anti-patrón: SQL suelto en la app que bypasea procs y auditoría.", "V"),
    q_abierta("Liste 3 ítems del checklist PI que YA tienen evidencia en su proyecto.",
              "Respuesta concreta: ER PNG, sp_*, trigger, TX, índices, etc."),
    q_abierta("Nombre el mayor hueco actual y la clase/hito donde lo cerrarán.",
              "Ej. contrato ops → Clase 12; concurrencia → ya/ajustes."),
],
12: [
    q_om("El contrato app↔BD es:",
         ["A) Un meme", "B) Lista de operaciones (procs), parámetros, errores y ejemplos",
          "C) Solo el ER sin operaciones", "D) La nota del parcial"], "B"),
    q_om("¿Por qué la app no debe hacer SQL suelto de negocio?",
         ["A) Es más bonito", "B) Duplica reglas y rompe auditoría/seguridad del PI",
          "C) Acelera siempre", "D) Obliga a DROP"], "B"),
    q_vf("La sustentación del PI dura típicamente 5 a 8 minutos.", "V"),
    q_vf("El ensayo de pitch de hoy sustituye el Parcial 3.", "F"),
    q_om("Entrada típica del contrato para agendar:",
         ["A) sp_agendar_cita(params) + errores + ejemplo CALL", "B) Solo «usar la BD»",
          "C) SELECT * sin firma", "D) Token de AWS"], "A"),
    q_vf("El contrato debe alinearse con los procedimientos ya implementados en el PI.", "V"),
    q_abierta("Escriba 2 operaciones de su contrato (nombre + 1 parámetro clave cada una).",
              "Ej. sp_agendar_cita(id_mascota); sp_facturar(id_consulta)."),
    q_abierta("Indique un error de negocio que la app debe manejar (código/mensaje).",
              "Ej. MASCOTA_INACTIVA / STOCK_INSUFICIENTE."),
],
13: [
    q_om("Analizar un caso real sirve en el PI para:",
         ["A) Copiar el logo", "B) Extraer lecciones de modelo, seguridad o operaciones a VetCare",
          "C) Evitar evidencias", "D) Eliminar el ER"], "B"),
    q_om("Lección típica transferable a VetCare:",
         ["A) Ignorar backups", "B) Auditoría, privilegios y consistencia en operaciones críticas",
          "C) SELECT * siempre", "D) Sin transacciones"], "B"),
    q_vf("Esta clase autónoma pide relacionar el caso con decisiones del propio PI.", "V"),
    q_vf("Un caso real «bonito» sin amarre a VetCare cumple el entregable.", "F"),
    q_om("Salida esperada del análisis:",
         ["A) Solo resumen de noticias", "B) 3 lecciones + impacto concreto en su VetCare",
          "C) Solo emojis", "D) Borrar scripts"], "B"),
    q_vf("Incidentes de datos (fugas, inconsistencias) motivan controles que ya vieron en el curso.", "V"),
    q_abierta("Nombre el caso analizado y 1 riesgo que también tiene VetCare.",
              "Respuesta propia: caso + riesgo (privilegios, backup, TX…)."),
    q_abierta("¿Qué cambio harán en su PI por esa lección (artefacto concreto)?",
              "Ej. añadir trigger auditoría / endurecer rol / documentar RPO."),
],
15: [
    q_om("En la sustentación VetCare, lo mínimo a demostrar es:",
         ["A) Solo hablar sin demo", "B) Modelo + operaciones clave + evidencias de rúbrica",
          "C) Solo un repositorio de archivos", "D) AWS de pago"], "B"),
    q_om("El PI (20% Corte 3):",
         ["A) Reemplaza el Parcial 3", "B) No sustituye el Parcial 3",
          "C) Elimina asistencia", "D) Es sin rúbrica"], "B"),
    q_vf("Mostrar secretos/credenciales en pantalla durante la demo es aceptable.", "F"),
    q_vf("La demo debe usar su propio dominio VetCare (no otro caso improvisado).", "V"),
    q_om("Orden razonable de pitch:",
         ["A) Problema → modelo → ops/demo → decisiones → cierre", "B) Solo memes",
          "C) Leer 40 diapositivas de teoría", "D) Empezar por la factura AWS"], "A"),
    q_vf("Un fallo controlado explicado (rollback/validación) puede reforzar el diseño.", "V"),
    q_abierta("Liste las evidencias que proyectarán (archivos/URLs) en la sustentación.",
              "ER, scripts procs/TX, contrato, capturas playground, checklist."),
    q_abierta("Trade-off principal que defenderán (1–2 frases).",
              "Respuesta alineada a sus propias decisiones (procs vs app, índices, etc.)."),
],
}


def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def run(r, *, size=11, bold=False, color=GRIS, font=None):
    f = font or FONT
    r.font.name = f
    r._element.rPr.rFonts.set(qn('w:eastAsia'), f)
    r.font.size = DocPt(size); r.bold = bold; r.font.color.rgb = color

def _sin_code_spans(text) -> str:
    """`code span` de Markdown -> «code span». El .docx no es Markdown: los acentos
    graves salian impresos tal cual en lo que lee el estudiante. El fuente puede
    seguir escribiendose en Markdown, que es lo que necesita el guion .md."""
    return re.sub(r"`([^`\n]+)`", r"«\1»", str(text))

def para(doc, text, *, size=11, bold=False, color=GRIS, space_after=6, shade_fill=None,
         font=None):
    """`font` solo se pasa para las plantillas: ver MONO."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = DocPt(space_after)
    if shade_fill: shade(p, shade_fill)
    r = p.add_run(_sin_code_spans(text)); run(r, size=size, bold=bold, color=color, font=font); return p

def banda(doc, text):
    return para(doc, "  "+text, size=13, bold=True, color=BLANCO, shade_fill="095292", space_after=8)

def add_inline_docx(p, text, *, size=11, color=GRIS):
    """Igual que uniajc_slides_engine._rich pero para runs de docx: soporta @@negrita@@.

    Y como alli, los `code spans` de Markdown pasan a «comillas angulares»: el docx del
    estudiante los imprimia con los acentos graves a la vista («los `GRANT` de hoy»).
    """
    text = _sin_code_spans(text)
    for part in re.split(r'(@@.*?@@)', text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith('@@') and part.endswith('@@'):
            r.text = part[2:-2]
            run(r, size=size, bold=True, color=color)
        else:
            r.text = part
            run(r, size=size, color=color)

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = DocPt(2)
        add_inline_docx(p, it)

DIAGRAMAS_BD2 = {
    1: {
        "titulo": "ER minimo VetCare (con cardinalidad)",
        "sub": "La FK siempre vive en la tabla del lado “N” de la relacion",
        "boxes": [
            {"id": "dueno", "label": "Dueño\nid_dueno (PK)\nnombre, telefono", "x": 0.6, "y": 2.2, "w": 2.7, "h": 1.6, "color": NAVY},
            {"id": "mascota", "label": "Mascota\nid_mascota (PK)\nid_dueno (FK)\nactiva", "x": 4.1, "y": 2.2, "w": 2.7, "h": 1.6, "color": CIAN},
            {"id": "cita", "label": "Cita\nid_cita (PK)\nid_mascota (FK)\nfecha_hora", "x": 7.6, "y": 2.2, "w": 2.7, "h": 1.6, "color": NAVY},
            {"id": "vet", "label": "Veterinario\nid_veterinario (PK)", "x": 7.6, "y": 4.4, "w": 2.7, "h": 1.1, "color": CIAN},
        ],
        "arrows": [
            {"src": "dueno", "dst": "mascota", "label": "1 : N"},
            {"src": "mascota", "dst": "cita", "label": "1 : N"},
            {"src": "vet", "dst": "cita", "label": "1 : N"},
        ],
        "note": "“Un dueño tiene N mascotas”, “una mascota tiene N citas”. La PK identifica cada fila; la FK materializa la relación y la BD la hace cumplir.",
    },
    10: {
        "titulo": "Doble reserva sin control de concurrencia",
        "sub": "Dos transacciones concurrentes leen “libre” antes de que ninguna confirme",
        "boxes": [
            {"id": "t1a", "label": "T1 (Recepción A)\nlee franja: LIBRE", "x": 0.8, "y": 2.2, "w": 3.1, "h": 1.0, "color": NAVY},
            {"id": "t1b", "label": "T1: INSERT cita\nCOMMIT → OK", "x": 4.6, "y": 2.2, "w": 3.1, "h": 1.0, "color": NAVY},
            {"id": "t2a", "label": "T2 (Recepción B)\nlee franja: LIBRE", "x": 0.8, "y": 3.7, "w": 3.1, "h": 1.0, "color": CIAN},
            {"id": "t2b", "label": "T2: INSERT cita\nCOMMIT → ¡DOBLE RESERVA!", "x": 4.6, "y": 3.7, "w": 3.1, "h": 1.0, "color": PPTX_RED},
        ],
        "arrows": [
            {"src": "t1a", "dst": "t1b", "label": "tiempo →"},
            {"src": "t2a", "dst": "t2b", "label": "tiempo →"},
        ],
        "note": "Mitigación: UNIQUE(id_veterinario, fecha_hora) hace que el segundo INSERT falle automáticamente en vez de crear la doble reserva.",
    },
}


# Codigo PROYECTABLE por clase: el fragmento minimo que el estudiante debe ver en
# pantalla mientras el docente explica (no el script completo — ese va en Codigo/).
CODIGO_SLIDE = {
    1: ("El DDL minimo que sostiene el ER", [
        "CREATE TABLE mascota (",
        "  id_mascota INT PRIMARY KEY,               -- PK: identifica la fila",
        "  id_dueno   INT NOT NULL",
        "             REFERENCES dueno(id_dueno),    -- FK: materializa la relacion",
        "  nombre     VARCHAR(60) NOT NULL,",
        "  activa     CHAR(1) DEFAULT 'S'",
        "             CHECK (activa IN ('S','N'))    -- regla de negocio en la BD",
        ");",
    ], "La FK vive en el lado 'N' de la relacion. El CHECK ya es una regla de negocio."),
    2: ("Reducir la superficie: vista y privilegio por columna", [
        "-- 1) La vista deja fuera el email y las citas canceladas",
        "CREATE VIEW v_agenda_recepcion AS",
        "SELECT c.id_cita, c.fecha_hora, d.nombre AS dueno, d.telefono",
        "  FROM cita c JOIN mascota m ON m.id_mascota = c.id_mascota",
        "  JOIN dueno d ON d.id_dueno = m.id_dueno",
        " WHERE c.estado <> 'CANCELADA';",
        "GRANT SELECT ON v_agenda_recepcion TO recepcion;",
        "REVOKE SELECT ON dueno FROM recepcion;   -- solo llega por la vista",
        "",
        "-- 2) O se recorta la tabla columna por columna",
        "GRANT SELECT (id_dueno, nombre) ON dueno TO veterinario_rol;",
    ], "Un GRANT no es todo-o-nada: la vista recorta filas y columnas, y el privilegio "
       "por columna recorta columnas sin crear objeto nuevo."),
    # El molde completo, no solo el IF: la pregunta 1 vale 35 puntos y 6 de ellos son
    # de sintaxis PL/pgSQL. Antes esta diapositiva proyectaba un parametro OUT con el
    # mensaje de error, que es exactamente lo que la rubrica NO acepta.
    3: ("El molde de PL/pgSQL y la validacion que aborta", [
        "CREATE OR REPLACE PROCEDURE sp_agendar_cita(",
        "  p_id_mascota INT, p_id_veterinario INT,",
        "  p_fecha_hora TIMESTAMP)          -- id_cita es SERIAL: no se pasa",
        "LANGUAGE plpgsql AS $proc$         -- ni IS, ni VARCHAR2, ni / final",
        "DECLARE v_activa CHAR(1);",
        "BEGIN",
        "  SELECT activa INTO v_activa FROM mascota",
        "   WHERE id_mascota = p_id_mascota;",
        "  IF NOT FOUND THEN",
        "    RAISE EXCEPTION 'ERROR: la mascota % no existe', p_id_mascota;",
        "  END IF;",
        "  IF v_activa <> 'S' THEN",
        "    RAISE EXCEPTION 'ERROR: la mascota % esta inactiva', p_id_mascota;",
        "  END IF;                          -- aborta: no inserta NADA",
        "  INSERT INTO cita(id_mascota, id_veterinario, fecha_hora, estado)",
        "  VALUES (p_id_mascota, p_id_veterinario, p_fecha_hora, 'PROGRAMADA');",
        "END; $proc$;",
    ], "`RAISE EXCEPTION` aborta el CALL completo y deshace lo hecho. Con el mensaje "
       "en un parametro OUT, el INSERT seguiria corriendo."),
    # En PostgreSQL el trigger son DOS objetos. La rubrica de la pregunta 2 descuenta
    # por `:NEW`/`:OLD` y por omitir `EXECUTE FUNCTION`, asi que la diapositiva tiene
    # que proyectar la forma correcta: antes proyectaba la de Oracle.
    4: ("Un trigger son DOS objetos: la funcion y la asociacion", [
        "-- 1) La funcion: NEW y OLD SIN dos puntos",
        "CREATE OR REPLACE FUNCTION fn_trg_audit_cita()",
        "RETURNS TRIGGER LANGUAGE plpgsql AS $fn$",
        "BEGIN",
        "  INSERT INTO audit_cita(id_cita, accion,",
        "                         valor_anterior, valor_nuevo)",
        "  VALUES (NEW.id_cita, 'CAMBIO_ESTADO',",
        "          OLD.estado, NEW.estado);   -- usuario y fecha: DEFAULT",
        "  RETURN NEW;                        -- obligatorio",
        "END; $fn$;",
        "",
        "-- 2) La asociacion: cuando se dispara y a quien llama",
        "CREATE TRIGGER trg_audit_cita",
        "AFTER UPDATE OF estado ON cita FOR EACH ROW",
        "WHEN (OLD.estado IS DISTINCT FROM NEW.estado)",
        "EXECUTE FUNCTION fn_trg_audit_cita();",
    ], "El `WHEN` es lo que hace que 3 UPDATE dejen 2 filas de auditoria. Se dispara "
       "solo: el riesgo es que sea invisible para quien solo lee el codigo de la app."),
    # El nombre del indice se proyectaba como `idx_cita_fecha` y la pregunta 1 califica
    # `idx_cita_fecha_hora`: esta es la diapositiva de la que el estudiante copia, asi que
    # era la que le costaba los puntos. Se agrega el parcial, que tambien se califica.
    7: ("Un indice se justifica con la consulta que lo usa", [
        "-- Consulta frecuente: la agenda del dia (siempre PROGRAMADA)",
        "SELECT ... FROM cita",
        " WHERE fecha_hora >= :hoy AND estado = 'PROGRAMADA';",
        "",
        "CREATE INDEX idx_cita_fecha_hora ON cita (fecha_hora);",
        "CREATE INDEX idx_cita_programada_fecha ON cita (fecha_hora)",
        "  WHERE estado = 'PROGRAMADA';        -- indice PARCIAL",
        "",
        "-- Mal candidato: baja cardinalidad (solo 'S' o 'N')",
        "-- CREATE INDEX idx_mascota_activa ON mascota(activa);",
    ], "Cada indice acelera lecturas y encarece INSERT/UPDATE/DELETE. El nombre se califica letra por letra."),
    # Decia «ROLLBACK; -- de lo contrario: COMMIT;», que es literalmente la opcion
    # INCORRECTA de la pregunta 4 (10 pts) proyectada como respuesta. En PostgreSQL el
    # procedimiento no lleva control de transaccion: la excepcion que se propaga fuera del
    # `CALL` deshace todo, y quien copiaba esta diapositiva perdia los 10 puntos.
    8: ("Todo o nada: la transaccion de facturacion", [
        "CREATE PROCEDURE sp_facturar(p_id_consulta INT,",
        "        p_insumos INT[], p_cantidades INT[]) ...",
        "  INSERT INTO factura(...) RETURNING id_factura INTO v_id_factura;",
        "  FOR i IN 1 .. array_length(p_insumos, 1) LOOP",
        "    UPDATE insumo SET stock = stock - p_cantidades[i]",
        "     WHERE id_insumo = p_insumos[i]",
        "       AND stock >= p_cantidades[i];   -- 0 filas si no alcanza",
        "    GET DIAGNOSTICS v_filas = ROW_COUNT;",
        "    IF v_filas = 0 THEN RAISE EXCEPTION '...'; END IF;",
        "  END LOOP;   -- sin COMMIT y sin ROLLBACK: los pone el CALL",
    ], "La condicion stock >= cantidad evita el stock negativo; la excepcion que sale del CALL deshace todo sola."),
    10: ("La restriccion que hace imposible la doble reserva", [
        "ALTER TABLE cita",
        "  ADD CONSTRAINT uq_cita_vet_franja",
        "  UNIQUE (id_veterinario, fecha_hora);",
        "",
        "-- T2 intenta la misma franja que T1:",
        "-- ORA-00001: unique constraint violated",
    ], "Poner solo BEGIN/COMMIT no basta: ambas transacciones leen 'libre' antes de confirmar."),
    12: ("El contrato que la app consume (no SQL suelto)", [
        "sp_agendar_cita(",
        "   p_id_cita     IN  NUMBER,",
        "   p_id_mascota  IN  NUMBER,",
        "   p_fecha       IN  TIMESTAMP,",
        "   p_msg         OUT VARCHAR2   -- 'OK: ...' | 'ERROR: ...'",
        ")",
    ], "La app llama el proc con parametros tipados: por eso no hay inyeccion SQL."),
}

# Comparaciones antes/despues (lo que mejor se entiende visualmente).
ANTES_DESPUES = {
    # Decia «Plan: TABLE ACCESS FULL (~120.000 filas)» -> «INDEX RANGE SCAN (~340
    # filas)». Dos cosas falsas en una diapositiva que ve el estudiante. Una: esos son
    # nombres de nodo de ORACLE, y el taller corre sobre PostgreSQL (Seq Scan). Dos, y
    # peor: hoy NO se crea ningun indice, asi que ninguna version puede dar un Index
    # Scan — la solucion docente lo dice expresamente («sin indice las dos versiones
    # siguen leyendo las 30.010 filas»). El estudiante que buscara el Index Scan
    # prometido concluiria que su respuesta correcta esta mal, o lo inventaria.
    6: {
        "titulo": "Optimizar es un ANTES medible, no una opinion",
        "b_t": "Antes",
        "b": ["`SELECT *` arrastra las columnas de 4 tablas por el join",
              "`to_char(fecha_hora,…)` y `UPPER(estado)`: **funcion sobre la columna**",
              "Joins con coma: si falta una condicion, **producto cartesiano**",
              "Subconsulta por fila: el plan dice **loops=2006**"],
        "a_t": "Despues",
        "a": ["Las 6 columnas que la pantalla de agenda usa",
              "`>= TIMESTAMP '2026-03-10' … <` y `estado = '…'`: **sargable**",
              # La opcion 5 de la pregunta 4 de ExamLab es «cambiar la coma por JOIN…ON por
              # si solo hace la consulta mas rapida», y es FALSA: mismo plan. Si la
              # diapositiva lo pusiera en la columna «Despues» sin decirlo, estaria
              # ensenando justo la opcion que descuenta.
              "`JOIN … ON`: **no acelera**, impide el cartesiano",
              "`LEFT JOIN` + `GROUP BY`: de 2.006 pasadas a **una**"],
        "sub": "Mismas **91 filas** en las dos versiones. Y sin indices no hay Index Scan "
               "—eso es la Clase 7—: hoy se miden filas procesadas y pasadas sobre la tabla",
    },
    2: {
        "titulo": "Minimo privilegio, en concreto",
        "b_t": "Lo que suele pasar",
        "b": ["Un solo usuario `admin` que todos comparten",
              "GRANT ALL 'para que no falle nada'",
              "Nadie sabe quien borro que (sin trazabilidad)",
              "Se van del cargo y la cuenta sigue viva"],
        "a_t": "Lo que pide el taller",
        # Los nombres van tal cual se escriben en ExamLab: minusculas, y el del
        # veterinario con sufijo `_rol` porque `veterinario` ya es una tabla.
        "a": ["4 roles: admin_bd · recepcion · veterinario_rol · auditor",
              "Privilegio por objeto y operacion (matriz)",
              "recepcion nunca hace DELETE: cancelar es UPDATE",
              "Politica de baja: revocacion el mismo dia"],
    },
    3: {
        "titulo": "Por que un procedimiento y no SQL en cada pantalla",
        "b_t": "Sin procedimiento",
        "b": ["Cada pantalla reescribe la regla a su manera",
              "Alguien la escribe distinta... o la olvida",
              "La app arma SQL con texto -> inyeccion SQL",
              "Cambiar la regla = buscarla en N lugares"],
        "a_t": "Con procedimiento",
        "a": ["La regla vive UNA vez, dentro de la BD",
              "Toda la app la respeta sin excepcion",
              "Parametros tipados: sin concatenar texto",
              "Cambiar la regla = un solo lugar"],
    },
}


# Diapositivas de teoria que no caben en ninguno de los tres moldes anteriores
# (diagrama, antes/despues, codigo) y que cierran la corrida teorica: van despues
# de CODIGO_SLIDE y antes de la demo.
#
# Por que existe: el taller cobra puntos por entregables que ninguna diapositiva
# enseñaba. El caso que la creo es la politica de altas y bajas de la Clase 2, que
# vale 15 de 100 puntos y cuyo unico ancla era «Para el PI esta semana», una
# diapositiva que trae el hito y la fecha limite pero ni una palabra de la
# politica. Antes de inventar un molde nuevo, mire si el contenido es una lista de
# secciones: si lo es, va aqui.
#
#   {n: [(titulo, [vinetas], subtitulo_o_None), ...]}
TEORIA_EXTRA = {
    2: [(
        "Ciclo de vida de una cuenta: alta, cambio, baja, revision",
        [
            "@@1. Alta.@@ Quien solicita y quien aprueba; la cuenta nace con @@un solo rol@@ "
            "y con credencial temporal que caduca en el primer ingreso.",
            "@@2. Cambio de rol.@@ `GRANT` del nuevo y @@`REVOKE` del anterior@@: los permisos "
            "@@no se acumulan@@. Quien pasa de recepcion a auditoria y conserva los dos "
            "termina auditando lo que el mismo modifica.",
            "@@3. Baja: el mismo dia.@@ Revocar los roles, dejar la cuenta sin login y "
            "@@reasignar los objetos antes de borrar el rol@@ (`REASSIGN OWNED BY ... TO "
            "admin_bd`): PostgreSQL no deja hacer `DROP ROLE` de un rol que todavia posee "
            "objetos. Lo que esto evita es la @@cuenta huerfana@@.",
            "@@4. Revision periodica.@@ Cada 3 a 6 meses se audita la matriz con "
            "`information_schema.role_table_grants`, y @@alguien firma@@ la evidencia. "
            "Sin cuentas compartidas: si tres personas entran como `recepcion1`, la "
            "auditoria no puede decir quien cancelo la cita.",
            "@@5. Probar que el permiso NO esta.@@ No hace falta otra conexion (aqui hay un "
            "solo usuario con login): `SET ROLE recepcion;` cambia el rol efectivo, y ahi "
            "`DELETE FROM cita ...` debe responder @@permission denied@@. Se cierra con "
            "`RESET ROLE;`. Si el entorno no deja cambiar de rol, se documenta el intento: "
            "eso es una @@brecha de verificacion@@ del entregable.",
        ],
        "Cada flecha del ciclo es un GRANT o un REVOKE. La politica es lo que le pone "
        "responsable y plazo a cada uno: son las 5 secciones del entregable",
    )],
    # Clase 3: la bateria de pruebas vale 25 puntos y el contrato 15 — 40 de los 100.
    # Ninguna diapositiva las mencionaba: la teoria core llegaba hasta «hay que probar
    # el proc» y el molde del bloque DO solo existia dentro de ExamLab.
    3: [(
        "La bateria de pruebas: un bloque DO por caso",
        [
            "@@Por que un bloque por caso.@@ Si los cuatro `CALL` van seguidos, el primero que "
            "falla @@aborta el resto@@ y la evidencia queda a medias. `DO $$ ... $$;` es un "
            "bloque anonimo —se ejecuta una vez y no se guarda— y su `EXCEPTION` atrapa el "
            "error y deja seguir al siguiente caso.",
            "@@El molde de un caso error.@@ `DO $$ BEGIN CALL sp_agendar_cita(...); INSERT INTO "
            "resultado_prueba VALUES (..., 'FALLO: no lanzo error', FALSE); EXCEPTION WHEN "
            "OTHERS THEN INSERT INTO resultado_prueba VALUES (..., SQLERRM, TRUE); END $$;` — "
            "`SQLERRM` es el texto del error que se acaba de capturar.",
            "@@Los 4 casos, y la prueba de que si inserto.@@ 1 caso OK + mascota inexistente + "
            "mascota inactiva + franja ocupada. Y un `COUNT(*)` sobre `cita` @@antes y "
            "despues@@: tiene que pasar de @@10 a 11@@ filas. Sin ese conteo nada demuestra que "
            "el caso OK hizo algo.",
        ],
        "Un procedimiento sin bateria no esta terminado: son 25 de los 100 puntos",
    ), (
        # Segunda mitad de la bateria. Se separo de la diapositiva anterior porque con el
        # cuerpo a 20 pt las cinco vinetas ya no caben en una sola, y porque el guion
        # docente venia tratando esto como una explicacion aparte: la seccion «Que
        # significa la columna paso, y la trampa del WHEN OTHERS» apuntaba a la misma
        # diapositiva que «por que un bloque DO por caso». Ahora cada seccion tiene la
        # suya y el reparto de minutos del plan las cuenta como dos.
        "La columna paso y la trampa de WHEN OTHERS",
        [
            "@@Ojo con la logica invertida, y con `WHEN OTHERS` a secas.@@ En un caso error, "
            "@@llegar al final sin excepcion es el fallo@@. Y no basta con que falle: hay que "
            "verificar @@que falle por lo esperado@@ (`SQLERRM ILIKE '%inactiva%'`), porque un "
            "typo en un nombre de columna tambien lanza excepcion y se reportaria como prueba "
            "superada.",
            "@@La tabla de resultados.@@ `resultado_prueba(id_prueba SERIAL, caso TEXT, esperado "
            "TEXT, obtenido TEXT, paso BOOLEAN)`, y un `SELECT * ... ORDER BY id_prueba;` con las "
            "4 filas. `paso` admite dos lecturas —«coincidio con lo esperado» o «la operacion se "
            "completo»— y las dos valen; lo que se exige es @@usar la misma para las 4 y decir "
            "cual@@.",
        ],
        "Los 25 puntos se califican por esta tabla, no por el CALL suelto",
    ), (
        "El contrato del procedimiento: los 6 bloques que consume la app",
        [
            "@@1. Firma exacta.@@ `sp_agendar_cita(p_id_mascota INT, p_id_veterinario INT, "
            "p_fecha_hora TIMESTAMP)`: nombre, orden y @@tipo@@ de cada parametro. Sin los tipos "
            "no es un contrato, es una descripcion. @@2. Como se llama:@@ un `CALL` de ejemplo "
            "con valores reales, copiable tal cual.",
            "@@3. Precondiciones@@ — lo que garantiza quien llama: la mascota existe y esta "
            "activa, la franja del veterinario esta libre.",
            "@@4. Postcondiciones@@ — lo que queda cierto si termina bien: una fila nueva en "
            "`cita` con estado `PROGRAMADA`, y @@ninguna@@ si termina mal.",
            "@@5. Tabla de errores.@@ Una fila por cada `RAISE EXCEPTION`, con el @@mensaje "
            "literal@@ y que debe hacer la app al recibirlo. Es lo que vuelve programable el "
            "contrato: sin el texto exacto, la app no puede distinguir un error de otro.",
            "@@6. Decision de diseno.@@ Por que se aborta en vez de devolver un codigo en un "
            "`OUT`: porque abortar @@deshace@@ lo hecho, y un codigo que nadie revise deja la "
            "cita creada igual.",
        ],
        "Es el entregable de la pregunta 5, vale 15 puntos, y la plantilla en blanco de los "
        "6 bloques esta en el Taller del PI",
    ), (
        # La pregunta 3 vale 10 puntos y contrasta PROCEDURE con FUNCTION. En las 18
        # diapositivas no aparecia ni una vez FUNCTION, RETURNS ni «funcion»: el deck
        # solo enseñaba procedimientos, asi que el contraste se deducia de las opciones
        # del quiz. Cada vineta refuta uno de los cinco distractores.
        "PROCEDURE o FUNCTION: cual se puede usar dentro de un SELECT",
        [
            "@@`FUNCTION`: devuelve un valor y se invoca DENTRO de una consulta.@@ "
            "`CREATE FUNCTION fn_precio(p_especie TEXT) RETURNS NUMERIC LANGUAGE plpgsql "
            "AS $fn$ ... RETURN v_precio; ... $fn$;` y se usa `SELECT nombre, "
            "fn_precio(especie) FROM mascota;`. Es la @@unica de las dos@@ que cabe en la "
            "lista de columnas de un `SELECT`, en un `WHERE` o en un `ORDER BY`.",
            "@@`PROCEDURE`: ejecuta pasos y se invoca como sentencia suelta.@@ "
            "`CALL sp_agendar_cita(3, 7, '2026-09-01 10:00');`. @@No@@ se puede poner en "
            "un `SELECT`. Si lo intentas, el motor responde `sp_agendar_cita(...) is not a "
            "function`.",
            "@@Las dos pueden ser `LANGUAGE plpgsql`.@@ El lenguaje no decide nada: es "
            "falso que una funcion necesite `LANGUAGE sql` para poder retornar un valor. "
            "`sp_agendar_cita` y `fn_precio` son las dos `plpgsql`.",
            "@@Un `PROCEDURE` si admite parametros `OUT`@@ —y aun asi @@sigue sin poder "
            "usarse dentro de un `SELECT`@@. Devolver por `OUT` no es lo mismo que ser "
            "invocable en una consulta, y no es «la unica forma de retornar un valor».",
            "@@Lo que solo puede el procedimiento: manejar la transaccion.@@ Un "
            "`PROCEDURE` puede hacer `COMMIT` o `ROLLBACK` adentro; una `FUNCTION` corre "
            "dentro de la transaccion de quien la llamo. Por eso agendar es procedimiento "
            "y @@la tarifa de la Clase 4 sera funcion@@.",
        ],
        "PROCEDURE y FUNCTION no son sinonimos: si el resultado tiene que entrar en un "
        "SELECT, es FUNCTION. Son 10 de los 100 puntos y es la bisagra con la Clase 4",
    )],
    # Clase 4: la pregunta 4 (donde vive cada validacion) vale 15 y el plan de respaldo
    # 25 — 40 de los 100. La teoria core nombraba RPO/RTO pero ninguna diapositiva decia
    # las 6 secciones que se califican ni con que herramienta de PostgreSQL.
    # Y la pregunta 1 vale 20: pedia escribir fn_precio_consulta con CASE, COALESCE e
    # IMMUTABLE, y la unica sintaxis de funcion proyectada era el RETURNS TRIGGER de la
    # diapositiva de codigo. La funcion normal solo estaba NOMBRADA en la demo y en los
    # criterios de exito, que es enunciar el entregable, no enseñar el mecanismo.
    4: [(
        "La funcion de tarifas: RETURNS NUMERIC, CASE, COALESCE e IMMUTABLE",
        [
            "@@La firma decide todo lo demas.@@ `CREATE FUNCTION fn_precio_consulta("
            "p_especie TEXT, p_urgencia BOOLEAN) RETURNS NUMERIC LANGUAGE plpgsql "
            "IMMUTABLE AS $fn$ ... $fn$;`. Es `RETURNS NUMERIC`, @@no `RETURNS TRIGGER`@@: "
            "esta funcion no se asocia a ninguna tabla, se llama desde una consulta.",
            "@@La tabla de tarifas es un `CASE`.@@ `CASE UPPER(p_especie) WHEN 'CANINO' "
            "THEN 45000 WHEN 'FELINO' THEN 40000 ELSE 35000 END`. El `UPPER()` es lo que "
            "hace la comparacion @@insensible a mayusculas@@: sin el, «canino» en "
            "minuscula cae en el `ELSE` y cobra 35000.",
            "@@El `NULL` de la urgencia: hazlo explicito con `COALESCE`.@@ "
            "`IF COALESCE(p_urgencia, FALSE) THEN v_base := v_base * 1.35; END IF;`. Un "
            "`IF` a secas ya trata `NULL` como falso, pero en cuanto el booleano entra en "
            "una @@cuenta@@ el `NULL` se propaga y la funcion devuelve `NULL`.",
            "@@`IMMUTABLE` es una promesa que el motor cobra:@@ @@misma entrada, misma "
            "salida, sin leer tablas@@. Por eso el filtro `WHERE m.activa = 'S'` va en la "
            "consulta y @@no@@ dentro de la funcion. Una funcion que consulta tablas seria "
            "`STABLE`, nunca `IMMUTABLE`.",
            "@@Y se usa dentro del `SELECT`, que es lo que la hace funcion@@ (el contraste "
            "de la Clase 3): `SELECT nombre, especie, fn_precio_consulta(especie, FALSE) "
            "AS tarifa_normal, fn_precio_consulta(especie, TRUE) AS tarifa_urgencia FROM "
            "mascota ORDER BY id_mascota;` — un canino en urgencia da @@60750@@ "
            "(45000 × 1.35).",
        ],
        "Son 20 de los 100 puntos, y se pierden casi siempre por tres cosas: un RETURNS "
        "TRIGGER copiado del trigger, el UPPER() olvidado y un IMMUTABLE que lee tablas",
    ), (
        "Donde vive cada validacion: CHECK, trigger o aplicacion",
        [
            "@@`CHECK`: una sola fila, una sola columna.@@ `stock >= 0`, `precio > 0`, `estado "
            "IN (...)`. Es lo mas barato y lo mas dificil de saltarse, porque lo aplica el "
            "motor sin una linea de codigo. @@Si la regla cabe en un CHECK, no se hace "
            "trigger.@@",
            "@@`UNIQUE` y `FK`: la relacion entre filas o entre tablas.@@ «un veterinario no "
            "puede tener dos citas en la misma franja» es `UNIQUE (id_veterinario, "
            "fecha_hora)`, no un trigger.",
            "@@Trigger: cuando la regla necesita OTRA fila u OTRA tabla.@@ Comparar `OLD` con "
            "`NEW`, mirar el stock de `insumo` mientras se actualiza otra tabla, o @@escribir "
            "en una segunda tabla@@ (auditoria). Nada de eso cabe en un `CHECK`.",
            "@@Aplicacion: solo lo que la base no puede saber.@@ El formato del correo, el "
            "permiso de la pantalla, la traduccion del mensaje. Y el criterio para decidir: "
            "@@una validacion que solo vive en la app se salta conectandose por `psql`@@.",
            "@@BEFORE valida, AFTER audita.@@ Un trigger que rechaza corre `BEFORE`, porque "
            "despues el dato ya esta escrito; uno que registra un hecho consumado corre "
            "`AFTER`. Poner `AFTER` en el de stock es el error mas comun del dia.",
        ],
        "La pregunta 4 son 15 puntos y no pide codigo: pide ubicar cada validacion en su "
        "capa y justificar por que ahi",
    ), (
        "Plan de respaldo: 6 secciones y herramientas reales de PostgreSQL",
        [
            "@@1. Que se respalda y con que.@@ `pg_dump -Fc -d vetcare -f "
            "vetcare_AAAAMMDD.dump` para los datos, y `pg_dumpall --globals-only` para los "
            "@@roles de la Clase 2@@: `pg_dump` es de UNA base y los roles son del cluster, asi "
            "que no los respalda. Si falta, se restaura la base y ningun rol tiene permisos.",
            "@@2. Frecuencia, con su razon.@@ No «diario»: «`pg_dump -Fc` a las 20:30 @@porque@@ "
            "la facturacion de Huellitas cierra a las 20:00». La hora se justifica con un hecho "
            "del negocio.",
            "@@3. Retencion@@ —cuantas copias y donde: 7 diarias, 4 semanales, 12 mensuales, y "
            "@@al menos una fuera del mismo servidor@@— y @@4. RPO y RTO en numeros@@: cuantos "
            "datos se acepta perder («maximo 1 hora de citas») y cuanto puede estar caida la "
            "base. Un dump diario implica un RPO de 24 h: si no alcanza, hace falta "
            "@@archivado de WAL@@.",
            "@@5. Restore de prueba@@ con `pg_restore` sobre una base vacia: fecha, duracion "
            "medida, la consulta que confirma el conteo, cada cuanto se ensaya y quien firma. "
            "@@Un respaldo que nunca se restauro no es un respaldo:@@ es un archivo.",
            "@@6. Que NO cubre el plan@@ y cual es el riesgo residual que se asume: el borrado "
            "por error que se descubre tres dias despues, la hora de citas que cabe en el RPO. "
            "@@Es la seccion que separa un plan de una lista de comandos@@, y la que mas se olvida.",
        ],
        "Son 25 de los 100 puntos de la clase y se califican seccion por seccion. La "
        "plantilla en blanco esta en el Taller del PI",
    )],
    # Clase 6: tres huecos del criterio rector, y los tres valian puntos.
    #
    # 1. LEER EL PLAN es la pregunta 2 completa — 20 de 100 — mas la afirmacion 6 de la
    #    pregunta 4: pide el «nodo mas costoso», «filas estimadas vs reales» y el «tiempo
    #    total» de tres EXPLAIN. Ninguna diapositiva mostraba un plan ni nombraba un solo
    #    campo. El fundamento si lo explicaba, pero la unica bala de teoria que menciona
    #    EXPLAIN es la sexta, y `_slide_summary` corta a 5: no se proyectaba.
    # 2. LA SUBCONSULTA CORRELACIONADA es la pregunta 3 y una de las cuatro afirmaciones
    #    correctas de la 4 — unos 22 de 100 —, y no aparecia en ninguna diapositiva ni en
    #    ninguna bala de teoria. La palabra «correlacionad-» si estaba en el fundamento,
    #    pero referida a PREDICADOS correlacionados, que es otro tema (estadisticas), asi
    #    que el docente podia buscarla, encontrarla y creer que estaba cubierta.
    # 3. LA PRUEBA DE EQUIVALENCIA (los dos COUNT(*) de la pregunta 1 y el EXCEPT de la 3)
    #    suma otros 6 puntos y tampoco estaba.
    6: [(
        "Leer un plan: es un arbol y se lee de adentro hacia afuera",
        [
            "@@El orden de lectura.@@ El plan @@no@@ es una lista de pasos de arriba hacia "
            "abajo: es un arbol. Los nodos @@mas indentados son las hojas@@ y se ejecutan "
            "primero; cada nodo consume las filas de sus hijos. La @@primera linea impresa es "
            "la ULTIMA operacion@@.",
            "@@Los cuatro campos.@@ `cost=270.00..4821.50` son arranque y total en una unidad "
            "@@relativa@@ (1.0 = leer una pagina de 8 KB), @@no@@ milisegundos; `rows=` son las "
            "filas que el motor @@estima@@; `width=` el ancho de la fila en bytes; `loops=` "
            "cuantas @@veces se repitio@@ ese nodo.",
            "@@`EXPLAIN` estima, `EXPLAIN ANALYZE` ejecuta.@@ El segundo agrega `actual "
            "time=`, `actual rows=` y el `Execution Time` final. Y ejecuta @@de verdad@@: "
            "sobre un `UPDATE` o `DELETE` hay que envolverlo en `BEGIN` … `ROLLBACK` (eso es "
            "la Clase 8). Con `SELECT` no hay riesgo.",
            "@@Estimado contra real: la senal.@@ Una divergencia de 2 veces es normal; de @@10 "
            "veces o mas@@ delata estadisticas viejas o predicados que el motor cree "
            "independientes y no lo son. Es lo que se anota en la tabla de la pregunta 2.",
            "@@El nodo mas costoso, sin trampa.@@ El tiempo de un nodo @@incluye el de sus "
            "hijos@@, y `actual time` es @@por vuelta@@: hay que multiplicarlo por `loops`. Un "
            "nodo de 0,5 ms con `loops=2006` cuesta un segundo — mas que el `Seq Scan` de "
            "arriba que parece el culpable.",
        ],
        "Es la pregunta 2 completa — 20 de los 100 puntos — y se responde leyendo estos "
        "campos, no pegando el plano tal cual",
    ), (
        "La subconsulta correlacionada: 2.006 pasadas o una sola",
        [
            "@@Que la hace correlacionada.@@ La subconsulta esta en la @@lista de columnas@@ y "
            "menciona una columna del exterior (`WHERE m.id_dueno = d.id_dueno`). No se puede "
            "calcular una vez y reusar: depende de la fila que se este mirando, asi que el motor "
            "@@la ejecuta una vez por fila@@ del exterior.",
            "@@El numero que lo delata en el plan.@@ `EXPLAIN ANALYZE` muestra un nodo `SubPlan` "
            "con @@`loops=2006`@@ — un dueno, una ejecucion. Y cada ejecucion recorre las 30.010 "
            "citas. `loops` es el campo que hay que buscar: es el unico lugar donde el plan dice "
            "«esto se repitio».",
            "@@La reescritura.@@ `dueno LEFT JOIN mascota LEFT JOIN cita` + `GROUP BY d.id_dueno, "
            "d.nombre` + `COUNT(...)`. El `SubPlan` desaparece y en su lugar queda un solo "
            "`HashAggregate`: @@una pasada@@. Es la unica mejora del dia que es de ordenes de "
            "magnitud, y no necesita ningun indice — lo que se elimino no fue un escaneo, fueron "
            "2.005 escaneos.",
            "@@`COUNT(c.id_cita)`, nunca `COUNT(*)`.@@ El `LEFT JOIN` fabrica @@una fila llena de "
            "`NULL`@@ por cada dueno sin citas. `COUNT(*)` cuenta filas y reporta @@1@@; `COUNT` de "
            "una columna ignora los `NULL` y reporta @@0@@. El sintoma es exacto: los duenos sin "
            "citas dicen 1 en vez de 0.",
            "@@Y `LEFT`, no `INNER`.@@ Un `INNER JOIN` es mas rapido y esta @@mal@@: borra del "
            "reporte a los duenos sin mascotas, y el ranking deja de cuadrar con el total de "
            "clientes de la clinica. Mas rapido devolviendo otra cosa no es optimizar.",
        ],
        "Es la pregunta 3 del taller — 20 de los 100 puntos — y la afirmacion de la "
        "pregunta 4 que mas se falla",
    ), (
        "Optimizar no cambia el resultado: como se prueba",
        [
            "@@La regla.@@ Correccion y tiempo son ejes @@independientes@@. Si la version DESPUES "
            "devuelve algo distinto, no se optimizo nada: se rompio la consulta, y se rompio @@sin "
            "avisar@@, porque ningun motor va a lanzar un error por eso.",
            "@@Prueba 1, para una consulta con filtro: los dos `COUNT(*)`.@@ Se envuelve cada "
            "version y se comparan los conteos en la @@misma corrida@@. En la agenda del "
            "2026-03-10 las dos tienen que decir @@91@@. Un conteo distinto es la respuesta "
            "equivocada, no una version mas rapida.",
            "@@Prueba 2, para conjuntos completos: `EXCEPT` en los DOS sentidos.@@ `A EXCEPT B` "
            "devuelve lo que esta en A y no en B. Que salga vacio @@no@@ prueba la igualdad: B "
            "puede tener filas de mas. Se corren las dos direcciones unidas con `UNION ALL` y se "
            "exige @@cero filas@@.",
            "@@Sin `LIMIT` en la prueba.@@ Comparar solo las primeras 20 filas deja fuera "
            "justamente las que fallan — las de los duenos con cero citas, que es donde "
            "`COUNT(*)` miente. El `LIMIT` va en la consulta que se entrega, no en la que verifica.",
            "@@Que NO sirve como prueba.@@ «Se ve igual», «trae mas o menos lo mismo» o mirar la "
            "primera pantalla de resultados. La equivalencia se afirma con una consulta cuyo "
            "resultado @@se conoce de antemano@@: un numero que coincide o un conjunto vacio.",
        ],
        "Vale puntos dos veces: los dos `COUNT(*)` de la pregunta 1 y el `EXCEPT` de la "
        "pregunta 3",
    )],
    # Clase 7: el deck llegaba hasta «un indice se justifica con la consulta que lo usa» y
    # la actividad cobraba 70 de los 100 puntos por tres mecanismos que ninguna diapositiva
    # proyectaba. Uno: los nombres son EXACTOS y la demo ensenaba `idx_cita_fecha` cuando la
    # pregunta 1 califica `idx_cita_fecha_hora`. Dos: el indice parcial, que la pregunta 1
    # exige por nombre —y su rubrica descuenta si falta— y que la pregunta 4 vuelve a cobrar,
    # no aparecia en ningun lado. Tres: el particionamiento, que la teoria core declaraba
    # «idea conceptual, no se implementa hoy» mientras la pregunta 3 lo implementaba entero
    # por 20 puntos. La regla del prefijo izquierdo si estaba en el fundamento; lo que
    # faltaba era el experimento con el que la pregunta 2 la hace demostrar.
    7: [(
        "Los cinco indices de hoy, con su nombre exacto",
        [
            "@@El nombre se califica.@@ La actividad dice «nombres exactos» y los compara letra "
            "por letra: `idx_cita_fecha_hora`, `idx_mascota_dueno`, `idx_cita_programada_fecha` "
            "en la pregunta 1, y `idx_cita_estado_fecha` + `idx_cita_fecha_estado` en la 2. "
            "@@`idx_cita_fecha` no es ninguno de los cinco@@: el sufijo es `_fecha_hora`, como la "
            "columna.",
            "@@Simple o compuesto.@@ Los tres primeros son de una columna; los dos ultimos llevan "
            "las mismas dos columnas @@en orden inverso@@, y existen precisamente para que se vea "
            "que el orden cambia el resultado.",
            "@@La secuencia que hay que respetar.@@ `EXPLAIN ANALYZE` @@antes@@ (tiene que salir "
            "`Seq Scan`) -> `CREATE INDEX` -> @@`ANALYZE cita;`@@ -> el mismo `EXPLAIN` otra vez. "
            "Sin el `ANALYZE` intermedio el planeador sigue con las estadisticas viejas y puede "
            "ignorar el indice que acabas de crear: la mitad de los «no me sirvio» del taller "
            "son este paso saltado.",
            "@@La evidencia de que existen.@@ `SELECT indexname, tablename, indexdef FROM "
            "pg_indexes WHERE tablename IN ('cita','mascota') ORDER BY tablename, indexname;` — "
            "`indexdef` devuelve el `CREATE INDEX` completo, asi que ahi se ve tambien el `WHERE` "
            "del parcial.",
            "@@Y el experimento del orden.@@ Con los dos compuestos creados se miden tres "
            "consultas: estado + rango de fecha, solo rango, solo estado. Despues se hace `DROP "
            "INDEX idx_cita_fecha_estado` y se vuelve a medir la del rango solo. @@La columna de "
            "igualdad va primero y la de rango al final@@; un indice cuya columna lider no "
            "aparece en el `WHERE` normalmente no se usa.",
        ],
        "Preguntas 1 y 2 del taller — 50 de los 100 puntos — y los nombres son los que se "
        "escriben en la plantilla del entregable",
    ), (
        "El indice parcial: el mismo beneficio, una fraccion del tamano",
        [
            "@@Que es.@@ Un indice que solo contiene las filas que cumplen una condicion: "
            "`CREATE INDEX idx_cita_programada_fecha ON cita (fecha_hora) WHERE estado = "
            "'PROGRAMADA';`. El `WHERE` no es el de la consulta: es @@parte de la definicion del "
            "indice@@ y decide que filas entran en el arbol.",
            "@@Cuanto se ahorra, con los datos de hoy.@@ De las 30.010 citas sembradas, "
            "@@18.187@@ estan `PROGRAMADA` (el 61 %): el indice completo indexa 30.010 entradas y "
            "el parcial 18.187, cuatro de cada diez menos. Menos entradas es menos disco, menos "
            "cache ocupada y menos trabajo en cada escritura de una cita @@que no este "
            "programada@@.",
            "@@La condicion para que sirva.@@ El planeador solo lo usa si puede demostrar que la "
            "consulta @@trae la misma condicion@@ del indice. `WHERE estado = 'PROGRAMADA' AND "
            "fecha_hora >= ...` si lo aprovecha; la misma consulta sin el filtro de estado @@no@@, "
            "porque el indice no contiene las filas atendidas ni las canceladas.",
            "@@El caso de VetCare que lo justifica.@@ La agenda del dia de la recepcion @@siempre@@ "
            "filtra por `PROGRAMADA`: nadie abre la pantalla para ver las citas que ya se "
            "atendieron. Cuando el filtro es parte del caso de uso y no del capricho de una "
            "consulta, el parcial es la respuesta correcta.",
            "@@Cual de los dos elige.@@ Con `idx_cita_fecha_hora` y `idx_cita_programada_fecha` "
            "compitiendo por la misma consulta, el plan nombra al ganador: hay que leer el "
            "`Index Scan using ...` y @@escribir cual salio@@, porque la rubrica descuenta si no "
            "se comenta.",
        ],
        "La pregunta 1 lo exige por nombre y su rubrica descuenta si falta; la pregunta 4 lo "
        "vuelve a cobrar como afirmacion correcta",
    ), (
        "Particionar hoy de verdad: rango por ano, poda y archivado",
        [
            "@@El DDL, en tres sentencias.@@ `CREATE TABLE cita_hist (... ) PARTITION BY RANGE "
            "(fecha_hora);` y despues una particion por ano: `CREATE TABLE cita_hist_2025 "
            "PARTITION OF cita_hist FOR VALUES FROM (TIMESTAMP '2025-01-01') TO (TIMESTAMP "
            "'2026-01-01');`. El rango es @@cerrado por abajo y abierto por arriba@@, asi que el "
            "`TO` de una particion es el `FROM` de la siguiente y nunca se solapan.",
            "@@La trampa que cuesta la pregunta.@@ En una tabla particionada la clave primaria "
            "@@debe incluir la columna de particion@@: `PRIMARY KEY (id_cita, fecha_hora)`. Un "
            "`PRIMARY KEY (id_cita)` a secas no compila, y el error del motor no dice «te falta "
            "la columna de particion» con esas palabras.",
            "@@La prueba de que el reparto ocurrio.@@ `SELECT tableoid::regclass AS particion, "
            "COUNT(*), MIN(fecha_hora), MAX(fecha_hora) FROM cita_hist GROUP BY 1 ORDER BY 1;` — "
            "`tableoid` es la columna de sistema que dice @@en que tabla fisica vive cada fila@@, "
            "y `::regclass` la traduce a nombre. Sin esta consulta no hay evidencia del "
            "enrutamiento, solo un `INSERT` que no dio error.",
            "@@La poda, que es lo unico que mejora hoy.@@ En `EXPLAIN ANALYZE` de una consulta "
            "acotada a 2026 el plan tiene que mencionar @@solo `cita_hist_2026`@@. Con 5.010 filas "
            "el tiempo no baja de forma apreciable y hay que decirlo: lo que se demuestra es que "
            "el motor @@descarta particiones enteras@@ antes de leer.",
            "@@Y el beneficio real, que es de mantenimiento.@@ Archivar 2025 es `DROP TABLE "
            "cita_hist_2025`, una operacion de @@metadatos@@ que tarda un instante. El `DELETE "
            "FROM cita WHERE fecha_hora < ...` equivalente toca millones de filas, infla el "
            "registro de transacciones y sostiene bloqueos largos — el contraste conecta directo "
            "con la Clase 8.",
        ],
        "Es la pregunta 3 completa —20 de los 100 puntos— y el segundo parrafo de la "
        "pregunta 5: hoy se implementa, no se cuenta",
    )],
    # Clase 8: el fundamento ensenaba el procedimiento canonico en PL/SQL de Oracle
    # (`IN NUMBER`, `SQL%ROWCOUNT`, `RAISE_APPLICATION_ERROR`, `EXCEPTION WHEN OTHERS THEN
    # ROLLBACK`) y la actividad califica PL/pgSQL: 75 de los 100 puntos no compilan con lo
    # que se proyectaba. Peor todavia, la pregunta 4 —10 puntos— tiene como opcion
    # INCORRECTA «porque el procedimiento incluia un ROLLBACK explicito en su bloque
    # EXCEPTION, igual que en Oracle», que es exactamente la forma que el guion presentaba
    # como canonica: el estudiante que estudiaba bien marcaba mal.
    8: [(
        "sp_facturar en PL/pgSQL: el molde que se califica",
        [
            "@@La firma exacta, que se factura por lineas.@@ `CREATE PROCEDURE "
            "sp_facturar(p_id_consulta INT, p_insumos INT[], p_cantidades INT[]) LANGUAGE "
            "plpgsql AS $proc$ ... $proc$;` — @@dos arreglos paralelos@@, no un insumo suelto: "
            "una factura tiene varias lineas. Se llama `CALL sp_facturar(4, ARRAY[1,6,5], "
            "ARRAY[1,2,3]);`. Los tipos son `INT` y `NUMERIC`, no `NUMBER`. Primera linea del "
            "cuerpo: `IF array_length(p_insumos, 1) IS DISTINCT FROM array_length(p_cantidades, 1) "
            "THEN RAISE EXCEPTION ...` — dos arreglos de longitud distinta se rechazan @@antes de "
            "tocar la base@@.",
            "@@La cabecera y el id que se acaba de generar.@@ `INSERT INTO factura (id_consulta, "
            "total) VALUES (p_id_consulta, 0) RETURNING id_factura INTO v_id_factura;` — el total "
            "entra en @@0@@ y se corrige al final, porque todavia no se sabe. `RETURNING ... INTO` "
            "evita ir a buscarlo con otro `SELECT`.",
            "@@El bucle, y el guardia que es el corazon del dia.@@ `FOR i IN 1 .. "
            "array_length(p_insumos, 1) LOOP` y dentro: `UPDATE insumo SET stock = stock - "
            "p_cantidades[i] WHERE id_insumo = p_insumos[i] AND @@stock >= p_cantidades[i]@@;` — "
            "la comprobacion viaja @@dentro del `WHERE`@@: comprobar y escribir son @@una sola "
            "sentencia@@ y nadie puede colarse entre las dos.",
            "@@Como se sabe si alcanzo, y como se aborta.@@ `GET DIAGNOSTICS v_filas = ROW_COUNT;` "
            "guarda cuantas filas toco el `UPDATE`: `1` alcanzo, `0` no habia stock. Entonces `IF "
            "v_filas = 0 THEN RAISE EXCEPTION 'ERROR: stock insuficiente del insumo %', "
            "p_insumos[i]; END IF;`. @@No es `SQL%ROWCOUNT` ni `RAISE_APPLICATION_ERROR`@@: eso es "
            "Oracle y aqui no existe.",
            "@@Y el cierre.@@ Cada iteracion inserta su linea en `detalle_factura` con el "
            "`precio_unit` @@vigente@@ y acumula `v_total := v_total + v_precio * "
            "p_cantidades[i];`. Al salir del bucle, `UPDATE factura SET total = v_total WHERE "
            "id_factura = v_id_factura;`. El caso de prueba da @@27.400@@ y deja los stocks en 11, "
            "58 y 5.",
        ],
        "Es la pregunta 1 del taller — 35 de los 100 puntos — y el molde que la pregunta 2 "
        "hace fallar a proposito",
    ), (
        "Por que el procedimiento no lleva COMMIT ni ROLLBACK",
        [
            "@@La regla de PostgreSQL.@@ Un `CALL` escrito por fuera de cualquier `BEGIN` es "
            "@@su propia transaccion@@. Si la excepcion se propaga hasta afuera del "
            "procedimiento, el motor deshace @@todo@@ lo que ese `CALL` habia hecho: la cabecera "
            "de la factura, el detalle y el descuento del primer insumo. @@Nadie escribio "
            "ROLLBACK@@ y la base queda intacta.",
            "@@Y el savepoint que nadie declara.@@ Un bloque `BEGIN ... EXCEPTION WHEN ... END` "
            "en PL/pgSQL crea un @@savepoint implicito@@ al entrar. Por eso, si tu capturas el "
            "error, se revierte solo lo hecho @@dentro de ese bloque@@ y el resto sigue vivo. "
            "Capturar no es lo mismo que dejar propagar.",
            "@@El contraste con Oracle, que es lo que se pregunta.@@ En Oracle el procedimiento "
            "es parte de la transaccion del llamador y ahi si se escribe `EXCEPTION WHEN OTHERS "
            "THEN ROLLBACK; RAISE;`. En PostgreSQL ese `ROLLBACK` dentro de un procedimiento "
            "invocado por un `CALL` de nivel superior @@ni siquiera esta permitido@@.",
            "@@La consecuencia practica.@@ Quien decide el `COMMIT` es @@uno solo@@: el llamador. "
            "Un procedimiento que confirma por su cuenta le quita al llamador la posibilidad de "
            "deshacer, y es la fuente numero uno de facturas a medias cuando la Clase 12 conecte "
            "la aplicacion.",
            "@@El error que no se perdona.@@ Capturar la excepcion y no volver a lanzarla. "
            "`EXCEPTION WHEN OTHERS THEN NULL;` @@convierte el fallo en silencio@@: la factura "
            "queda registrada, el stock no se descuenta y nadie se entera hasta el inventario.",
        ],
        "Es la pregunta 4 del taller —10 puntos, seleccion unica— y la explicacion de por "
        "que la pregunta 2 sale bien sin escribir un ROLLBACK",
    ), (
        "fn_descontar_stock: cuando «no hay stock» es una respuesta, no un error",
        [
            "@@La diferencia.@@ El procedimiento @@aborta@@ la factura completa; la funcion "
            "@@informa@@ y deja que el llamador decida. La misma regla de negocio, dos contratos "
            "distintos, y hay que saber cual se pide.",
            "@@La firma, y su palabra clave.@@ `CREATE OR REPLACE FUNCTION "
            "fn_descontar_stock(p_id_insumo INT, p_cantidad INT) RETURNS BOOLEAN LANGUAGE "
            "plpgsql AS $$ ...` — @@`RETURNS BOOLEAN`@@, no `PROCEDURE`. Devuelve `TRUE` si "
            "desconto y `FALSE` si no habia suficiente, @@sin lanzar excepcion@@ en ese segundo "
            "caso.",
            "@@Lo que si es un error.@@ Una cantidad negativa o cero no es «no hay stock», es una "
            "llamada mal hecha: eso @@si@@ va con `RAISE EXCEPTION`. Distinguir el dato invalido "
            "del resultado negativo es la mitad de la pregunta.",
            "@@Se prueba en una sola consulta.@@ `SELECT fn_descontar_stock(5,3) AS caso_ok, "
            "fn_descontar_stock(2,10) AS caso_sin_stock, fn_descontar_stock(2,3) AS caso_limite;` "
            "-> @@`true`, `false`, `true`@@. El tercer caso es el interesante: pide @@exactamente@@ "
            "el stock que queda, y con `>=` en el guardia tiene que pasar.",
            "@@Por que no leer primero.@@ `SELECT stock ...; IF stock >= cantidad THEN UPDATE ...` "
            "deja una @@ventana@@ entre la lectura y la escritura, y con dos recepcionistas "
            "facturando el mismo insumo las dos leen 3, las dos deciden que alcanza y el stock "
            "termina en @@-2@@. El `UPDATE` con la condicion en el `WHERE` no tiene ventana. Aqui "
            "no se puede demostrar —PGlite corre @@una sola sesion@@— y ese es el gap que se "
            "declara y que abre la Clase 10.",
        ],
        "Es la pregunta 3 del taller —15 puntos— y la decision documentada que pide la "
        "pregunta 5",
    )],
}


# Titulo de la diapositiva del flujo de diagramacion (Excalidraw -> IA -> Mermaid
# -> ExamLab). Vive aqui y no en examlab_talleres porque es el rotulo del deck del
# curso; el CONTENIDO de los 4 pasos si es compartido entre los cuatro cursos.
FLUJO_SLIDE_TITULO = "Del boceto a ExamLab (diagrama)"

# Diapositiva de contexto del cliente. Solo en la Clase 1: es donde el estudiante
# empieza a modelar y necesita saber para quien. Antes conocia a la clinica
# «Huellitas» por primera vez dentro de ExamLab, el dia que se le calificaba.
CLIENTE_SLIDE_TITULO = "El cliente · " + CLIENTE


def _fundamento_md(c):
    """Desarrollo en prosa del tema, SOLO para el guion docente.

    Va aparte de `teoria` a proposito: `teoria` son vinetas que alimentan tambien la
    diapositiva del estudiante (via _slide_summary, que toma la primera frase de cada
    una), asi que no se puede engordar sin romper la slide. La regla de oro del
    workspace pide que el docente pueda dictar sin consultar otra fuente, y para eso
    hacen falta parrafos desarrollados: eso vive aqui.
    """
    fund = (c.get("fundamento") or FUNDAMENTOS.get(c["n"]) or "").strip()
    if not fund:
        return ""
    fund = _resolver_slides(fund, _slide_map(c), c["n"])
    return "\n\n### Desarrollo del tema (para dictar sin consultar otra fuente)\n\n" + fund + "\n"


def _slide_summary(bullets_, max_chars=110, max_items=5):
    """Resume cada viñeta de teoria (pensada para el guion, muy detallada) a su
    idea central para que la diapositiva de estudiante no quede sobrecargada de texto.
    El guion docente conserva el texto completo; solo la slide usa esta versión corta.

    Dos filtros que la slide del ESTUDIANTE necesita y el guion no:
    - Se descarta la viñeta «Error de docente que no domina el tema...»: es material
      de preparación del docente y no tiene sentido proyectado al grupo.
    - Se corta a `max_items` viñetas (regla del workspace: máximo 5 por diapositiva)."""
    out = []
    for b in bullets_:
        if re.match(r"\s*Error\s+(tipico|típico|de)\s+d(el|e)\s+docente", b, re.I):
            continue
        first = re.split(r"(?<=[a-záéíóúü0-9\)])\.\s", b, maxsplit=1)[0].strip()
        if len(first) > max_chars or len(first) < 12:
            # colon-only label (ej. "Control pesimista:") o frase larga: usa un corte por longitud
            base = b if len(first) < 12 else first
            first = base[:max_chars].rsplit(" ", 1)[0].rstrip(":,;") + "…"
        out.append(first)
    return out[:max_items]


def _tiene_diagrama(n):
    """True si el taller de ExamLab de esta clase tiene una pregunta tipo `diagrama`.

    Es lo que decide si la clase necesita la diapositiva del flujo Excalidraw ->
    IA -> Mermaid -> ExamLab: sin pregunta de diagrama no aporta nada.
    """
    taller = TALLERES_EXAMLAB.get(n) or {}
    return any(p.get("tipo") == "diagrama" for p in taller.get("preguntas", []))


def _slide_map(c):
    """Titulos de las diapositivas de esta clase, EN ORDEN.

    Por que existe: el guion docente tenia una lista de «Referencias a diapositivas»
    escrita a mano que no coincidia con el deck real (anunciaba la demo en la slide 5
    cuando estaba en la 7), asi que el docente no podia leer el guion con la
    presentacion proyectada. Esta funcion refleja los mismos condicionales que
    `build_pptx`, y `build_pptx` verifica al final que las dos listas tengan la misma
    longitud: si alguien agrega una diapositiva y olvida este mapa, el build falla en
    vez de publicar un guion desincronizado.
    """
    if c['tipo'] == 'parcial':
        return [f"Portada · Clase {c['n']} · {c['titulo']}",
                "Que se evalua hoy",
                "Como se responde y como se entrega",
                f"{c['titulo']} · Clase {c['n']}"]
    m = [f"Portada · Clase {c['n']} · {c['titulo']}",
         "Encuadre de hoy · Objetivo PI",
         "Mapa del bloque de hoy (120 min)"]
    if c['n'] == 1:
        m.append(CLIENTE_SLIDE_TITULO)
    m.append("Teoria Core (breve)")
    dg = DIAGRAMAS_BD2.get(c['n'])
    if dg:
        m.append(dg["titulo"])
    ad = ANTES_DESPUES.get(c['n'])
    if ad:
        m.append(ad["titulo"])
    cs = CODIGO_SLIDE.get(c['n'])
    if cs:
        m.append(cs[0])
    for extra in TEORIA_EXTRA.get(c['n'], []):
        m.append(extra[0])
    m.append("Como se ordena la sesion de hoy" if c['tipo'] == 'sustentacion' else "Demo del dia")
    if HERRAMIENTAS_DIA.get(c["n"]):
        m.append("Herramientas de hoy")
    if _tiene_diagrama(c['n']):
        m.append(FLUJO_SLIDE_TITULO)
    tb = TALLER_BLOQUE.get(c["n"], {})
    label = {"autonoma": "Actividad autonoma",
             "sustentacion": "Sustentacion del PI"}.get(c["tipo"], "Taller PI VetCare")
    if tb.get("contexto"):
        m.append(f"{label} — contexto / por que importa")
    m.append(f"{label} — objetivo y criterios")
    if tb.get("escenario"):
        m.append(f"{label} — escenario / datos de partida")
    m.append(f"{label} — pasos guiados")
    if tb.get("pistas"):
        m.append(f"{label} — pistas (checklist vacio)")
    m.append("Criterios de exito / entregable")
    m.append("Cierre del PI" if c['tipo'] == 'sustentacion' else "Para el PI esta semana")
    m.append(f"Cierre · Clase {c['n']}")
    return m


def _plano(s):
    """Minusculas y sin tildes, para comparar fragmentos con titulos de diapositiva.

    Los modulos de datos se escriben SIN tildes por convencion, pero algunos titulos
    del deck si las llevan. Sin plegar los acentos habia que buscar a mano un trozo
    del titulo sin ninguna tilde, y eso es lo que hacia los fragmentos fragiles.
    """
    return "".join(
        ch for ch in unicodedata.normalize("NFD", (s or "").lower())
        if unicodedata.category(ch) != "Mn"
    )


def _slide_no(mapa, *fragmentos):
    """Numero (1-based) de la primera diapositiva cuyo titulo contiene el fragmento.

    Sirve para etiquetar el plan minuto a minuto con `[Slide N]` reales en vez de
    numeros escritos a mano.

    La portada se excluye a proposito: su titulo es «Portada · Clase N · <tema>» y el
    tema repite justo las palabras que uno querria usar como fragmento, asi que se
    quedaba con matches destinados a una diapositiva de teoria. Ningun bloque del
    fundamento se ancla nunca a la portada.
    """
    for frag in fragmentos:
        f = _plano(frag)
        for i, t in enumerate(mapa, 1):
            if _plano(t).startswith("portada"):
                continue
            if f in _plano(t):
                return i
    return None


def _slide_tag(mapa, *fragmentos):
    """`[Slide 7]` listo para pegar en el guion, o cadena vacia si no aplica."""
    n = _slide_no(mapa, *fragmentos)
    return f"[Slide {n}] " if n else ""


def _slides_teoria(mapa):
    """Corrida de teoria del deck: de «Teoria Core» hasta la anterior a la demo.

    Devuelve `[(numero, titulo), ...]`.

    Por que existe: el plan minuto a minuto anclaba la teoria a UNA diapositiva
    («Teoria Core») y luego saltaba a la demo, asi que las de diagrama,
    antes/despues, codigo y cierre teorico quedaban en el deck pero fuera del
    plan. El docente que lee el guion pasaba de la 4 a la 7 y el mecanismo que
    despues cobra el taller no se proyectaba nunca: en la Clase 2 eran 20 puntos
    (vista y privilegio por columna, diapositiva 6) proyectados por nadie.

    La corrida se deriva del mismo `_slide_map()` que arma el deck, de modo que
    agregar una diapositiva de teoria la mete en el plan sin tocar esta funcion.
    """
    ini = _slide_no(mapa, "Teoria Core")
    if ini is None:
        return []
    fin = _slide_no(mapa, "Demo del dia", "Como se ordena") or (len(mapa) + 1)
    return [(i, mapa[i - 1]) for i in range(ini, fin)]


# --- Referencias a diapositivas dentro del fundamento teorico ---------------
# En la prosa se escribe «{{slide:Teoria Core}}» y aqui se convierte en
# «diapositiva 4» usando el mapa real del deck. Asi el numero no se escribe a mano
# en ningun sitio y no puede quedar corrido.
_SLIDE_TOKEN = re.compile(r"\{\{\s*slide:\s*([^}]+?)\s*\}\}")


def _resolver_slides(texto, mapa, n_clase):
    def _rep(m):
        frag = m.group(1)
        i = _slide_no(mapa, frag)
        if i is None:
            raise SystemExit(
                f"Clase {n_clase}: el fundamento referencia la diapositiva "
                f"«{frag}», que ya no existe en el deck. Corrige el texto o "
                "_slide_map()."
            )
        return f"diapositiva {i}"
    return _SLIDE_TOKEN.sub(_rep, texto)


def cover_pptx(prs, c):
    """Portada limpia: marca + título + subtítulo. Meta PI/agenda → 2ª slide."""
    class_cover(prs, c['titulo'], subtitulo=c['subtitulo'], clase_n=c['n'], idx=1)

def build_pptx(c):
    if c['tipo'] == 'parcial':
        # Este deck es lo UNICO que el estudiante recibe el dia del parcial y no traia
        # ninguno de los datos que se preguntan al minuto 0: ni el alcance —que clases
        # entran—, ni el reparto de puntos por seccion, ni el peso en el corte, ni el canal
        # de entrega, ni la hora de cierre, ni que el SQL se responde escrito. Ademas
        # estaba sin acentos, y esto no es fuente de fundamento: es lo que se proyecta. Y
        # el cierre afirmaba «El PI VetCare DB continua la proxima clase» tambien en el
        # Parcial 3, cuando lo que sigue ahi es la sustentacion.
        p = PARCIALES_BD2[c['n']]
        m = _parcial_meta_bd2(p['corte'])
        prs = new_prs()
        class_cover(prs, c['titulo'], subtitulo="Solo evaluación · sin tema ni taller",
                    clase_n=c['n'], idx=1)
        content_slide(prs, "Qué se evalúa hoy",
                      [f"**{s.split(' — ')[0]}** — {s.split(' — ')[1]}"
                       for s in m['secciones_resumen']]
                      + [f"Total **100 puntos** · nota = puntos / 20 · "
                         f"este parcial pesa {m['valor_corte']}."],
                      sub="**Solo** " + " · ".join(t.split(' · ')[0] for t in m['temas'])
                          + " — fuera de esa lista no hay nada",
                      idx=2)
        content_slide(prs, "Cómo se responde y cómo se entrega", [
            "El envío **cierra en el minuto 110**: lo que llegue después no se recibe.",
            "Canal de entrega: el que se anuncia ahora. Confirmo cada recibido por el chat.",
            "El SQL se escribe **como texto**: no se pide captura ni se abre ExamLab.",
            "Pregunta de **forma** sí (cuántas líneas, si pide tabla). De **contenido** no.",
            "Si se te cae el internet: sigue respondiendo y avisa por correo al volver.",
        ], sub=f"Tiempo previsto **{m['tiempo']}** dentro del bloque de 120", idx=3)
        closing_slide(prs, f"{c['titulo']} · Clase {c['n']}",
                      ["Hoy solo se evalúa el corte", _cierre_pi_parcial(c['n'])],
                      accent="Solo evaluación")
        _verificar_mapa(c, prs)
        out_dir = CLASES_DIR / f"Clase {c['n']} - {c['titulo']}"
        out_dir.mkdir(parents=True, exist_ok=True)
        out = out_dir / "Presentacion.pptx"
        prs.save(str(out)); print("PPTX", out)
        return out
    prs = new_prs(); cover_pptx(prs, c); idx = 2
    tipo_lbl = {"autonoma": "autonoma (festivo)",
                "sustentacion": "sustentacion en vivo"}.get(c['tipo'], "regular")
    # 2ª slide: encuadre / objetivos (contenido que salió de la portada)
    if c['tipo'] == 'sustentacion':
        content_slide(prs, "Encuadre de hoy · Objetivo PI", [
            f"**Hoy cerramos el PI:** {c['hito_pi']}",
            "Sesión **síncrona** de sustentaciones · Bloque **120 min** · turnos consecutivos.",
            f"**Entregable de hoy:** {c['entregable']}",
            "Sube el paquete a ExamLab **antes** de tu turno: presentando no se sube nada.",
            "La sustentación es **en vivo** y con preguntas: no se reemplaza por video grabado.",
        ], idx=idx); idx += 1
    else:
        content_slide(prs, "Encuadre de hoy · Objetivo PI", [
            f"**Hoy avanzamos el PI en:** {c['hito_pi']}",
            f"Herramienta: **{c['herramienta']}** · Bloque **120 min** · Teoría breve · Taller PI",
            f"**Entregable de hoy:** {c['entregable']}",
            "Gratis + navegador · free tier · sin software de pago obligatorio.",
            "La teoría no es un tema aislado: alimenta evidencias de la rúbrica del PI.",
            "Al salir: avance concreto en su paquete VetCare.",
        ], idx=idx); idx += 1
    if c['tipo'] == 'sustentacion':
        # El bloque no se reparte en teoría + taller: son turnos de sustentación.
        block_timeline_slide(prs, "Mapa del bloque de hoy (120 min)", [
            ("0-10", "Encuadre · sorteo del orden de turnos"),
            ("10-110", "Sustentaciones: 5-8 min de pitch + Q&A por turno"),
            ("110-120", "Cierre del curso · autoevaluacion"),
        ], idx=idx); idx += 1
    else:
        block_timeline_slide(prs, "Mapa del bloque de hoy (120 min)", [
            ("0-10", f"Encuadre · clase {tipo_lbl} · VetCare"),
            ("10-35", "Teoría Core breve (al servicio del PI)"),
            ("35-55", "Demo con la herramienta del día"),
            ("55-105", "Taller guiado = tarea del PI"),
            ("105-120", "Criterios de exito · cierre · dudas del PI"),
        ], idx=idx); idx += 1
    if c['n'] == 1:
        content_slide(prs, CLIENTE_SLIDE_TITULO, [
            f"@@{CLIENTE}@@ (Cali): unas @@150 citas al día@@, @@16 veterinarios@@, "
            "~5.000 mascotas de ~2.000 dueños — y todo en @@carpetas de papel@@.",
            "Duele en tres puntos: se extravían fichas, buscar un historial genera filas "
            "en la sala de espera, y no hay métricas.",
            "Tres personas van a usar el sistema, y esperan cosas distintas: "
            "@@el dueño@@ métricas, @@la recepcionista@@ agendar rápido, "
            "@@el veterinario@@ el historial a la mano.",
            "Sus intereses @@entran en conflicto@@: más datos dan mejores métricas, pero "
            "hacen más lento el agendamiento. Ahí están las decisiones de diseño del semestre.",
            "@@Caso completo:@@ anexo «Caso de estudio Clínica Huellitas» en "
            "Clases/Proyecto Integrador — 8 entidades, 3 reglas y el elenco de nombres.",
        ], sub=NOMENCLATURA, idx=idx); idx += 1
    content_slide(prs, "Teoria Core (breve)", _slide_summary(c['teoria']), idx=idx,
                  ); idx += 1
    dg = DIAGRAMAS_BD2.get(c['n'])
    if dg:
        diagram_boxes_slide(
            prs, dg["titulo"], dg["boxes"], arrows=dg.get("arrows"),
            sub=dg.get("sub"), note=dg.get("note"), idx=idx,
        )
        idx += 1
    ad = ANTES_DESPUES.get(c['n'])
    if ad:
        # `sub` es opcional: lo usa la Clase 6 para decir en la propia diapositiva que
        # hoy no hay indices, y que por eso la evidencia no es un cambio de nodo.
        before_after_slide(prs, ad["titulo"], ad["b_t"], ad["b"], ad["a_t"], ad["a"],
                           sub=ad.get("sub"), idx=idx)
        idx += 1
    cs = CODIGO_SLIDE.get(c['n'])
    if cs:
        pseudo_code_slide(prs, cs[0], cs[1], caption=cs[2], idx=idx)
        idx += 1
    for extra in TEORIA_EXTRA.get(c['n'], []):
        content_slide(prs, extra[0], extra[1],
                      sub=(extra[2] if len(extra) > 2 else None), idx=idx,
                      )
        idx += 1
    if c['tipo'] == 'sustentacion':
        # Hoy no hay demo del docente: el que demuestra es el estudiante, en su turno.
        content_slide(prs, "Como se ordena la sesion de hoy", [
            "Sustentación **en vivo**, en este bloque: no se reemplaza por video grabado.",
            "Turnos de **5–8 min de pitch + Q&A**; el orden se sortea al empezar.",
            f"Ten listo: **{c['entregable']}**",
            "En tu turno muestra una **ejecución real** (procedimiento OK + caso rechazado), no solo capturas.",
            "Mientras otros presentan, escuchas: el cierre del curso se hace con todo el grupo.",
        ], idx=idx); idx += 1
    else:
        content_slide(prs, "Demo del dia", [
            f"**Herramienta:** {c['herramienta']}",
            f"**Demo:** {c['demo']}",
            "Sigan el mismo dominio VetCare (no inventen otro caso).",
            "Al final de la demo: dejar enlace/script compartible al grupo.",
        ], idx=idx); idx += 1
    tools = HERRAMIENTAS_DIA.get(c["n"])
    if tools:
        herramientas_slide(prs, tools, title="Herramientas de hoy",
                           sub="Gratis · navegador o free tier", idx=idx)
        idx += 1
    # Del boceto al codigo: solo en las clases cuyo taller tiene pregunta de
    # diagrama. El estudiante disena en Excalidraw/draw.io y entrega Mermaid; sin
    # esta diapositiva llegaba con un PNG a una caja que espera texto.
    if _tiene_diagrama(c['n']):
        dialectos = examlab_talleres._dialectos_del_taller(TALLERES_EXAMLAB[c['n']])
        steps_visual_slide(
            prs, FLUJO_SLIDE_TITULO,
            examlab_talleres.flujo_diagrama_pasos(
                dialectos[0] if len(dialectos) == 1 else "el tipo que pide el enunciado"),
            sub="El diagrama se entrega como codigo Mermaid dentro de ExamLab, no como imagen",
            idx=idx)
        idx += 1
    tb = TALLER_BLOQUE.get(c["n"], {})
    label = {"autonoma": "Actividad autonoma",
             "sustentacion": "Sustentacion del PI"}.get(c["tipo"], "Taller PI VetCare")
    if tb.get("contexto"):
        content_slide(prs, f"{label} — contexto / por que importa", tb["contexto"], idx=idx,
                      )
        idx += 1
    obj = tb.get("objetivo") or c["hito_pi"]
    crit = [f"@@Exito:@@ {x}" for x in tb.get("criterios", [])] or [
        f"@@Entregable:@@ {c['entregable']}",
        # No hay «playground»: la evidencia vive dentro de ExamLab, que guarda la
        # consulta y lo que devolvio la base en cada pregunta. Este texto solo se
        # usa si la clase no tiene `criterios` en TALLER_BLOQUE.
        "Evidencia: lo que ejecutaste y su salida quedan guardados en la pregunta de ExamLab.",
    ]
    content_slide(prs, f"{label} — objetivo y criterios", [f"@@Objetivo:@@ {obj}", *crit], idx=idx,
                  )
    idx += 1
    if tb.get("escenario"):
        content_slide(prs, f"{label} — escenario / datos de partida", tb["escenario"], idx=idx,
                      )
        idx += 1
    steps_visual_slide(prs, f"{label} — pasos guiados", [(t, "") for t in c["taller"]], idx=idx)
    idx += 1
    if tb.get("pistas"):
        checklist_slide(prs, f"{label} — pistas (checklist vacio)", tb["pistas"], idx=idx)
        idx += 1
    if c['tipo'] == 'sustentacion':
        content_slide(prs, "Criterios de exito / entregable", [
            f"**Entregable:** {c['entregable']}",
            "Evidencia ejecutable en pantalla durante tu turno (no solo capturas).",
            "Puedes explicar **cualquier** parte de tu modelo en 60 segundos.",
            "@@Entrega en ExamLab@@ (https://uniaj.examlab.workers.dev/ · módulo Proyectos) — **antes** de tu turno.",
        ], idx=idx); idx += 1
        box_note_slide(prs, "Cierre del PI", [
            ("info", f"Hito: {c['hito_pi']}"),
            ("aclaracion", "Enunciado completo y rubrica: Clases/Proyecto Integrador/ (VetCare DB)."),
            ("advertencia", "El PI vale 20% del Corte 3 y NO reemplaza el Parcial 3, que ya se aplico en su propia sesion."),
        ], idx=idx); idx += 1
    else:
        content_slide(prs, "Criterios de exito / entregable", [
            f"**Entregable:** {c['entregable']}",
            # Decia «Evidencia en playground (enlace)»: no hay playground ni enlace que
            # compartir — la entrega es dentro de ExamLab, que guarda la consulta y la
            # salida en cada pregunta. Es una diapositiva compartida por las 15 clases.
            "Evidencia: el SQL y su salida quedan guardados en cada pregunta de ExamLab.",
            "Conserva copia en tu carpeta del PI (los .sql que pide el entregable).",
            "Actualizar checklist PI (que criterio de rubrica avanzo).",
            "@@Entrega en ExamLab@@ (https://uniaj.examlab.workers.dev/) — domingo 23:59 cuando aplique taller.",
        ], idx=idx); idx += 1
        box_note_slide(prs, "Para el PI esta semana", [
            ("info", f"Hito: {c['hito_pi']}"),
            ("aclaracion", "Enunciado completo: Clases/Proyecto Integrador/ (VetCare DB)."),
            ("advertencia", "Taller de la semana en ExamLab (https://uniaj.examlab.workers.dev/): domingo 23:59 (regla del Acuerdo) cuando aplique."),
        ], idx=idx); idx += 1
    # El QUIZ no va en el material del estudiante: ni proyectado ni anunciado.
    # Vive solo en Kit docente/Clase N/ (enunciados + CLAVE DOCENTE aparte), que
    # el docente aplica por el canal que decida. Anticiparlo en la diapositiva le
    # quita sentido como comprobacion.
    if c['tipo'] == 'sustentacion':
        closing_slide(prs, f"Clase {c['n']} · cierre de VetCare DB", [
            c['hito_pi'],
            f"Entregable: {c['entregable']}",
            "Conserva el paquete (ER, DDL, roles, procs, triggers, optimizacion) como portafolio",
        ], accent="Sustentar es justificar decisiones, no describir tablas")
    else:
        closing_slide(prs, f"Clase {c['n']} · VetCare avanza", [
            c['hito_pi'],
            f"Entregable: {c['entregable']}",
            "Siguiente clase: continuar el hilo del PI segun plan",
        ], accent="Teoria breve · practica = PI")
    _verificar_mapa(c, prs)
    out_dir = CLASES_DIR / f"Clase {c['n']} - {c['slug']}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "Presentacion.pptx"
    prs.save(str(out)); print("PPTX", out)
    return out


def _verificar_mapa(c, prs):
    """Aborta si `_slide_map` y el deck real dejaron de coincidir.

    Sin esta guarda el guion vuelve a anunciar numeros de diapositiva equivocados en
    cuanto alguien agregue o quite una slide, que es exactamente el defecto que se
    esta corrigiendo.
    """
    esperado, real = len(_slide_map(c)), len(prs.slides)
    if esperado != real:
        raise SystemExit(
            f"Clase {c['n']}: el deck tiene {real} diapositivas y _slide_map() "
            f"declara {esperado}. Actualiza _slide_map() en build_uniajc_bd2_all.py "
            "para que el guion siga apuntando a las diapositivas correctas."
        )


# Plantillas de ficha por clase: el esqueleto literal que el estudiante copia y
# llena. Refleja campo por campo lo que califica la pregunta abierta de ExamLab
# (que es la fuente de verdad de la nota); el documento solo evita que cada quien
# invente su propia estructura.
PLANTILLA_FICHA = {
    1: [
        "Nombre del proyecto: VetCare - [Apellido]",
        "Autor: [nombre completo]        Codigo: [__________]",
        "Integrantes: [solo si el docente autorizo equipo; si trabajas solo, escribe «individual»]",
        "Descripcion en una frase: [que es VetCare DB para la clinica Huellitas]",
        "",
        "1) QUE SI HARA EL PI  (5 a 8 lineas)",
        "   - [proceso de la clinica que la base va a soportar]",
        "   - [ ... ]",
        "",
        "2) QUE NO HARA EL PI  (3 a 5 lineas)",
        "   - [limite explicito: p. ej. no habra pasarela de pagos]",
        "   - [ ... ]",
        "",
        "3) TRES REGLAS DE NEGOCIO PROPIAS  —  formato Condicion -> Accion",
        "   R1. Si [condicion verificable] -> entonces [accion que hace la base].",
        "       Se implementa con: [ ] CHECK  [ ] UNIQUE  [ ] FK  [ ] procedimiento  [ ] trigger",
        "   R2. Si [ ... ] -> entonces [ ... ].",
        "       Se implementa con: [ ] CHECK  [ ] UNIQUE  [ ] FK  [ ] procedimiento  [ ] trigger",
        "   R3. Si [ ... ] -> entonces [ ... ].",
        "       Se implementa con: [ ] CHECK  [ ] UNIQUE  [ ] FK  [ ] procedimiento  [ ] trigger",
        "   (deben ser TUYAS: distintas de las tres del enunciado general)",
        "",
        "4) RIESGO PRINCIPAL PARA TERMINAR EL PI Y COMO LO MITIGAS  (2 lineas)",
        "   Riesgo: [ ... ]        Mitigacion: [ ... ]",
    ],
    # Clase 2: la matriz (pregunta 4) y la politica (pregunta 5) son 40 de los 100
    # puntos y las dos se califican POR SU ESTRUCTURA — la matriz por cubrir los 10
    # objetos x 4 roles sin celdas vacias, la politica por tener las 5 secciones con
    # responsable y plazo. Los nombres de rol van en minuscula y con el sufijo `_rol`
    # del veterinario, exactamente como los pide ExamLab y como los verifica la
    # solucion: escribirlos de otra forma en la matriz es una inconsistencia con el
    # script de la pregunta 1.
    2: [
        "A) MATRIZ ROL x OBJETO x PRIVILEGIO  (pregunta 4)",
        "   Leyenda de las celdas:  S=SELECT  I=INSERT  U=UPDATE  D=DELETE  E=EXECUTE  -=ninguno",
        "   Regla: ninguna celda queda vacia, y la matriz tiene que decir lo mismo que",
        "   los GRANT que ejecutaste en la pregunta 1.",
        "",
        "   | Objeto           | admin_bd | recepcion | veterinario_rol | auditor |",
        "   |------------------|----------|-----------|-----------------|---------|",
        "   | dueno            |          |           |                 |         |",
        "   | mascota          |          |           |                 |         |",
        "   | veterinario      |          |           |                 |         |",
        "   | cita             |          |           |                 |         |",
        "   | consulta         |          |           |                 |         |",
        "   | insumo           |          |           |                 |         |",
        "   | factura          |          |           |                 |         |",
        "   | detalle_factura  |          |           |                 |         |",
        "   | sp_agendar_cita  |          |           |                 |         |",
        "   | sp_facturar      |          |           |                 |         |",
        "",
        "   Justificacion (4 a 6 lineas, tres decisiones concretas de minimo privilegio):",
        "   1. [por que ningun rol operativo tiene D] ...",
        "   2. [por que auditor es de solo lectura, incluso sobre la tabla de auditoria] ...",
        "   3. [por que los dos sp_ van con E y no con I] ...",
        "",
        "B) POLITICA DE ALTAS Y BAJAS  —  maximo una pagina  (pregunta 5)",
        "   Cada seccion se califica por tener responsable y plazo concretos, no genericos.",
        "",
        "   1) ALTA",
        "      Quien la solicita: [cargo]        Quien la aprueba: [cargo]",
        "      Rol con el que nace la cuenta: [uno solo]",
        "      Como se entrega la credencial inicial: [ ... ]   Caduca en: [ ... ]",
        "",
        "   2) CAMBIO DE ROL  (recepcionista -> auxiliar veterinaria)",
        "      Se otorga: GRANT [ ... ] TO [ ... ];",
        "      Se REVOCA: REVOKE [ ... ] FROM [ ... ];   <- sin esto los permisos se acumulan",
        "",
        "   3) BAJA  (el mismo dia del retiro)",
        "      Paso 1 - revocar roles: [ ... ]",
        "      Paso 2 - dejar la cuenta sin login: [ ... ]",
        "      Paso 3 - objetos de los que era dueno: REASSIGN OWNED BY [ ... ] TO [ ... ];",
        "      Traza de auditoria: se conserva [cuanto tiempo] porque [ ... ]",
        "",
        "   4) REVISION PERIODICA",
        "      Cada [3 / 6] meses.        Quien firma la evidencia: [cargo]",
        "      Consulta con la que se saca la evidencia:",
        "      SELECT ______________ FROM information_schema.______________ ;",
        "",
        "   5) PRUEBA NEGATIVA Y LIMITE DEL ENTORNO",
        "      Comando con el que compruebo que el privilegio NO esta:",
        "        SET ROLE ____________;  ->  [ sentencia que debe fallar ]  ->  RESET ROLE;",
        "      Error que esperaba: [ ... ]        Lo que realmente obtuve: [ ... ]",
        "      Diferencia entre SET ROLE y conectarse como otro usuario: [ ... ]",
        "      Lo que este entorno NO permite (un solo usuario con login): [ ... ]",
        "      Si no pude correr la prueba: por que eso es una brecha de verificacion: [ ... ]",
    ],
    # Clase 3: el contrato (pregunta 5) vale 15 puntos y se califica por tener los
    # 6 bloques para CADA procedimiento y por que las firmas coincidan con el codigo
    # de las preguntas 1 y 4. Sin plantilla, la entrega tipica trae los dos
    # procedimientos mezclados y sin tabla de errores, que es la mitad del puntaje.
    # Las preguntas 1 a 4 son SQL: no llevan plantilla, se entregan como codigo.
    3: [
        "CONTRATO DE LOS PROCEDIMIENTOS  —  pregunta 5",
        "Los 6 bloques se llenan DOS veces, uno por procedimiento. La firma tiene que decir",
        "exactamente lo que dice tu codigo de las preguntas 1 y 4: mismo nombre, mismo orden,",
        "mismos tipos. Si al escribirla descubres que no coinciden, corrige el codigo o la",
        "firma, no dejes las dos versiones.",
        "",
        "=========================  A) sp_agendar_cita  =========================",
        "",
        "1) FIRMA EXACTA  (nombre, parametros con tipo PostgreSQL y orden)",
        "   sp_agendar_cita( ______________ ______ , ______________ ______ ,",
        "                    ______________ ______ )",
        "",
        "2) COMO SE INVOCA  (CALL de ejemplo con valores reales, no con descripciones)",
        "   CALL sp_agendar_cita( ____ , ____ , TIMESTAMP '____-__-__ __:__:__' );",
        "",
        "3) PRECONDICIONES  (que debe ser verdadero ANTES de llamarlo)",
        "   - [ ... ]",
        "   - [ ... ]",
        "   - [ ... ]",
        "",
        "4) POSTCONDICIONES  (que cambia en la base si la llamada tiene exito)",
        "   Filas insertadas: [cuantas, en que tabla, con que estado]",
        "   Filas actualizadas: [cuales, o «ninguna»]",
        "   Si la llamada falla: [que queda cambiado]",
        "",
        "5) TABLA DE ERRORES  ->  las 3 primeras filas de la tabla del final",
        "",
        "6) UNA DECISION DE DISENO JUSTIFICADA  (2 o 3 lineas: por que la validacion vive",
        "   en la base de datos y no solo en la aplicacion)",
        "   [ ... ]",
        "",
        "=====================  B) sp_registrar_consulta  =====================",
        "",
        "1) FIRMA EXACTA",
        "   sp_registrar_consulta( ______________ ______ , ______________ ______ ,",
        "                          ______________ ______ )",
        "",
        "2) COMO SE INVOCA",
        "   CALL sp_registrar_consulta( ____ , '________________________' , ______ );",
        "",
        "3) PRECONDICIONES",
        "   - [ ... ]                                        (son cuatro)",
        "",
        "4) POSTCONDICIONES",
        "   Fila insertada en: [ ... ]      Fila actualizada en: [ ... ] a estado [ ... ]",
        "   Las dos cosas o ninguna: [por que no puede quedar solo una]",
        "",
        "5) TABLA DE ERRORES  ->  las 4 ultimas filas de la tabla del final",
        "",
        "6) UNA DECISION DE DISENO JUSTIFICADA",
        "   [ ... ]",
        "",
        "==============  TABLA DE ERRORES  (bloque 5 de los dos contratos)  ==============",
        "Una fila por cada excepcion que lanzaste: 3 del primer procedimiento y 4 del",
        "segundo. El texto del mensaje se copia del codigo, no se parafrasea.",
        "",
        "| Procedimiento         | Excepcion | Texto del mensaje | Que debe hacer la aplicacion |",
        "|-----------------------|-----------|-------------------|------------------------------|",
        "| sp_agendar_cita       |           |                   |                              |",
        "| sp_agendar_cita       |           |                   |                              |",
        "| sp_agendar_cita       |           |                   |                              |",
        "| sp_registrar_consulta |           |                   |                              |",
        "| sp_registrar_consulta |           |                   |                              |",
        "| sp_registrar_consulta |           |                   |                              |",
        "| sp_registrar_consulta |           |                   |                              |",
        "",
        "REGLA DEL PI  —  frase de cierre, obligatoria",
        "   La aplicacion de Huellitas nunca hara ______________ directo sobre ____________",
        "   ni sobre ____________. Su unico acceso de escritura a esas dos tablas es",
        "   ____________ sobre ______________________ y ______________________.",
    ],
    # Clase 4: el Plan_Backup_VetCare (pregunta 5) vale 25 puntos, el puntaje mas alto
    # del taller, y se califica por tener las 6 secciones con numeros y herramientas
    # reales de PostgreSQL. Las secciones 1, 2 y 3 se responden en una sola tabla —esa
    # es la forma que espera la solucion—, y la 6 es la que el estudiante omite si
    # nadie le deja el renglon abierto. Las preguntas 1 a 3 son SQL y la 4 es de
    # seleccion: no llevan plantilla.
    4: [
        "PLAN_BACKUP_VETCARE  —  pregunta 5   ·   maximo una pagina",
        "Clinica Huellitas, PostgreSQL, atencion de lunes a sabado de 7:00 a 19:00.",
        "Todo numero va justificado contra ese horario. «Diario» sin hora no se califica;",
        "«diario a las 20:30 porque la facturacion cierra a las 19:45» si.",
        "",
        "SECCIONES 1, 2 y 3  —  QUE SE RESPALDA, CON QUE, CUANDO Y CUANTO SE GUARDA",
        "Minimo 4 filas: el enunciado nombra esquema (DDL), datos, rutinas",
        "(procedimientos / funciones / triggers) y scripts de migracion. Agrega las filas",
        "que tu plan necesite. La retencion exige AL MENOS DOS UBICACIONES distintas.",
        "",
        "| Que se respalda | Herramienta | Frecuencia y ventana | Retencion y ubicacion |",
        "|-----------------|-------------|----------------------|-----------------------|",
        "|                 |             |                      |                       |",
        "|                 |             |                      |                       |",
        "|                 |             |                      |                       |",
        "|                 |             |                      |                       |",
        "|                 |             |                      |                       |",
        "",
        "4) RPO Y RTO OBJETIVO",
        "   RPO objetivo: ______   (cuanta informacion aceptas perder, en tiempo)",
        "     Que se pierde exactamente en esa ventana, en datos de la clinica: [ ... ]",
        "     Por que la clinica lo tolera (o que hace falta para bajarlo): [ ... ]",
        "   RTO objetivo: ______   (en cuanto tiempo debes estar operando de nuevo)",
        "     Como se reparte ese tiempo:  detectar y decidir ____ + restaurar ____ +",
        "     validar ____ + margen ____",
        "     Impacto para la clinica si se excede: [ ... ]",
        "",
        "5) RESTORE DE PRUEBA",
        "   Pasos (nunca sobre la base de produccion):",
        "     1. [ ... ]",
        "     2. [ ... ]",
        "     3. [ ... ]",
        "     4. [ ... ]",
        "   Consulta de validacion post-restauracion, con el valor que debe salir:",
        "     SELECT (SELECT COUNT(*) FROM ________)  AS ____________ ,  -- esperado: ____",
        "            (SELECT COUNT(*) FROM ________)  AS ____________ ,  -- esperado: ____",
        "            (SELECT COUNT(*) FROM ________)  AS ____________ ,  -- esperado: ____",
        "            (SELECT ______(________) FROM ________) AS __________ ; -- esperado: ____",
        "   Cada cuanto se ensaya: ____________     Quien firma la evidencia: ____________",
        "   Donde queda la evidencia: ____________________________________",
        "",
        "6) QUE NO CUBRE ESTE PLAN, Y EL RIESGO RESIDUAL QUE SE ASUME",
        "   No cubre: [ ... ]",
        "   No cubre: [ ... ]",
        "   No cubre: [ ... ]",
        "   Riesgo residual asumido, escrito para que sea una decision y no un olvido:",
        "   [que puede pasar]  ·  [por que se acepta]  ·  [quien lo acepta]",
        "",
        "CIERRE  —  CHECKLIST DEL PI  (seguridad y respaldo)",
        "   Listo:        [ ... ]",
        "   En progreso:  [ ... ]",
        "   Falta:        [ ... ]",
    ],
    # Clase 6: dos entregables se califican POR SU ESTRUCTURA y suman 40 de los 100
    # puntos. La mini tabla de la pregunta 2 (9 de sus 20 pts) tiene tres filas y tres
    # columnas fijas, y la justificacion de la pregunta 5 (20 pts) se califica con «las
    # 5 secciones estan presentes»: sin el esqueleto, la entrega tipica trae tres
    # parrafos seguidos y pierde las secciones 4 y 5, que son las que nadie escribe si
    # no ve el renglon. Las preguntas 1 y 3 son SQL y la 4 es de seleccion: no llevan
    # plantilla. Los nombres de fila y de seccion son los de ExamLab, letra por letra.
    6: [
        "A) MINI TABLA DE LECTURA DEL PLAN  —  pregunta 2",
        "   Va como comentarios SQL (lineas con `--`) DESPUES de los tres EXPLAIN, en el",
        "   mismo cuadro de respuesta. Los numeros se leen del plan: se comparan `rows=`",
        "   con `actual rows=`, no se estiman.",
        "",
        "-- VERSION       | nodo mas costoso | filas estimadas vs reales | tiempo total (ms)",
        "-- ANTES         |                  |                           |",
        "-- DESPUES       |                  |                           |",
        "-- DESPUES+LIM50 |                  |                           |",
        "-- CONCLUSION: factor de mejora aproximado ____x  ( ____ ms -> ____ ms )",
        "",
        "   La fila DESPUES+LIM50 es la que mas ensena: antes de correrla, escribe aqui que",
        "   esperas que pase con el `LIMIT 50` y luego comprueba si acertaste.",
        "   Esperaba: [ ... ]        Paso: [ ... ]",
        "",
        "B) JUSTIFICACION TECNICA DEL ANTES/DESPUES  —  pregunta 5   ·   media pagina",
        "   Las 5 secciones se califican por estar presentes. Van con estos titulos.",
        "",
        "   1) CONSULTA ELEGIDA Y PARA QUE SIRVE EN HUELLITAS  (una frase)",
        "      Pantalla o reporte que la usa: [ ... ]     Frecuencia: [veces por dia]",
        "",
        "   2) TRES CAMBIOS CONCRETOS",
        "      Cambio 1 - que cambiaste: [ ... ]",
        "         Por que mejora: [sargabilidad / proyeccion / cardinalidad / numero de pasadas]",
        "         Evidencia del EXPLAIN ANALYZE: [nodo que desaparecio, tiempo que bajo,",
        "                                         filas que dejaron de leerse]",
        "      Cambio 2 - que cambiaste: [ ... ]",
        "         Por que mejora: [ ... ]",
        "         Evidencia del EXPLAIN ANALYZE: [ ... ]",
        "      Cambio 3 - que cambiaste: [ ... ]",
        "         Por que mejora: [ ... ]",
        "         Evidencia del EXPLAIN ANALYZE: [ ... ]",
        "",
        "   3) QUE NO CAMBIO  (la equivalencia)",
        "      Afirmacion: las dos versiones devuelven [ ____ ] filas, las mismas.",
        "      Como lo verifique:  COUNT(*) ANTES = ____   ·   COUNT(*) DESPUES = ____",
        "                          EXCEPT en los dos sentidos: ____ filas y ____ filas",
        "",
        "   4) QUE SIGUE  (el indice de la Clase 7)",
        "      CREATE INDEX ______________ ON ________ ( ________________ );",
        "      Por que ayudaria, dicho con tu propio plan: [ ... ]",
        "",
        "   5) LIMITES DE LA MEDICION",
        "      Medi sobre PostgreSQL en el navegador, con 30.010 citas y sin concurrencia.",
        "      Que cambiaria con millones de citas y varios usuarios: [ ... ]",
        "      Si no pude usar la opcion BUFFERS, aqui es donde se dice: [ ... ]",
    ],
    # Clase 7: la pregunta 5 vale 20 puntos y su rubrica dice «la tabla cubre al menos 3
    # indices con las 7 columnas». Siete columnas es justo el formato que nadie reproduce
    # de memoria: la entrega tipica trae tres o cuatro y pierde cardinalidad, costo y
    # veredicto, que son las que distinguen una justificacion de una lista de indices. El
    # veredicto de particionamiento del segundo parrafo tiene tres exigencias separadas
    # (numeros propios, la ganancia que NO es medible, y que si se comprobo), asi que va
    # con sus renglones. Las preguntas 1, 2 y 3 son SQL y la 4 es de seleccion: sin
    # plantilla. Los encabezados son los de ExamLab, letra por letra.
    7: [
        "A) TABLA DE JUSTIFICACION CONSULTA -> INDICE  —  pregunta 5",
        "   Una fila por indice, minimo 3 (los que creaste en las preguntas 1 y 2). Las 7",
        "   columnas se califican; una fila con 4 columnas llenas es una fila incompleta.",
        "   Cabe en horizontal, pero si no te cabe, usa el bloque de abajo por indice.",
        "",
        "   | Indice | Tabla y columnas | Consulta del PI que lo usa | Cardinalidad estimada"
        " de la columna lider | Evidencia en EXPLAIN | Costo de mantenimiento | Veredicto |",
        "   |--------|------------------|----------------------------|----------------------"
        "----------------|----------------------|------------------------|-----------|",
        "   |        |                  |                            |                      "
        "                |                      |                        |           |",
        "",
        "   Un bloque por indice, si prefieres escribirlo asi (repite los 7 renglones):",
        "",
        "   INDICE 1",
        "     Indice ....................: ______________________________  (nombre exacto)",
        "     Tabla y columnas ..........: ________ ( ____________________ )",
        "     Consulta del PI que lo usa : [la pantalla o el reporte, no «varias consultas»]",
        "     Cardinalidad de la lider ..: ______ valores distintos  ->  alta / baja",
        "        y por que eso lo hace util o inutil: [ ... ]",
        "     Evidencia en EXPLAIN ......: nodo: ____________________________________",
        "        (`Index Scan using ...`, `Bitmap Heap Scan`)   tiempo: ____ ms -> ____ ms",
        "     Costo de mantenimiento ....: que escritura de VetCare lo paga: [ ... ]",
        "     Veredicto .................: se queda / se cambia por parcial / se cambia por",
        "                                  compuesto / se descarta   ->  ______________",
        "",
        "B) REGLA DE SOBRE-INDEXACION QUE ADOPTAS  —  pregunta 5, parrafo 1",
        "   Tiene que ser verificable: alguien debe poder mirar tu proyecto y decir si la",
        "   cumpliste o no. «Voy a indexar con cuidado» no es una regla.",
        "   [ ... ]",
        "",
        "C) PARTICIONAMIENTO: VEREDICTO PARA VETCARE  —  pregunta 5, parrafo 2",
        "   Los tres renglones se califican por separado.",
        "",
        "   1) Volumen que espera Huellitas, con TUS numeros:",
        "      ______ citas por dia  x  ______ dias de operacion al ano  =  ______ citas/ano",
        "      A ______ anos de historia:  ______ citas en total",
        "   2) Veredicto:   particionar cita  SI  /  NO   ->  ______",
        "      Por que, atado al numero de arriba: [ ... ]",
        "   3) Lo que demostraste hoy, y lo que no:",
        "      Con 5.010 filas la ganancia de RENDIMIENTO no es apreciable  <-- decirlo suma",
        "      Lo que si quedo comprobado: la poda de particiones en el plan (solo aparecio",
        "      ______________________ ) y la facilidad de archivado ( DROP TABLE de la",
        "      particion en vez de DELETE masivo ).",
    ],
    # Clase 8: la pregunta 5 vale 15 puntos y se califica por cuatro bloques, uno de ellos
    # un checklist de 7 items donde la rubrica exige «estado y evidencia concreta (nombre
    # de indice, archivo, consulta), no solo casillas marcadas». Sin los renglones de
    # evidencia la entrega llega con siete casillas marcadas y cero evidencias, que es
    # exactamente lo que la rubrica no acepta. El inventario pide tres campos por
    # transaccion y el gap de concurrencia es un punto aparte. Las preguntas 1, 2 y 3 son
    # SQL y la 4 es de seleccion unica: sin plantilla.
    8: [
        "SECCION «TRANSACCIONES Y TUNING» DEL INFORME DEL PI  —  pregunta 5 · 1 pagina",
        "",
        "A) INVENTARIO DE TRANSACCIONES DE NEGOCIO DE VETCARE",
        "   Minimo TRES. Cada una con sus tres datos: sin el punto de fallo no es un",
        "   inventario de transacciones, es una lista de operaciones.",
        "",
        "   1. Operacion .........: ____________________________________________________",
        "      Tablas que toca ...: ____________________________________________________",
        "      Paso que puede fallar: ______________________________________________",
        "      Que debe pasar si falla: ____________________________________________",
        "   2. Operacion .........: ____________________________________________________",
        "      Tablas que toca ...: ____________________________________________________",
        "      Paso que puede fallar: ______________________________________________",
        "      Que debe pasar si falla: ____________________________________________",
        "   3. Operacion .........: ____________________________________________________",
        "      Tablas que toca ...: ____________________________________________________",
        "      Paso que puede fallar: ______________________________________________",
        "      Que debe pasar si falla: ____________________________________________",
        "",
        "B) CHECKLIST DE TUNING   ·   los 7 items, con estado Y evidencia",
        "   Estado: listo / parcial / pendiente.  Evidencia: un nombre, un archivo o una",
        "   consulta. Una casilla marcada sin evidencia no puntua.",
        "",
        "   | # | Item                                                    | Estado | Evidencia |",
        "   |---|---------------------------------------------------------|--------|-----------|",
        "   | 1 | indices sobre las columnas de filtro y join frecuentes  |        |           |",
        "   | 2 | consultas sin SELECT * en los reportes del PI           |        |           |",
        "   | 3 | predicados sargables (sin funciones sobre la columna)    |        |           |",
        "   | 4 | transacciones cortas: nada de esperar al usuario         |        |           |",
        "   | 5 | validaciones criticas en la base (CHECK, trigger, proc)  |        |           |",
        "   | 6 | ANALYZE / estadisticas al dia tras cargas masivas        |        |           |",
        "   | 7 | plan de respaldo con restore probado (viene de Clase 4)  |        |           |",
        "",
        "C) DECISION DOCUMENTADA   ·   una frase que puedas defender en la sustentacion",
        "   Por que el descuento de stock se hace con",
        "   UPDATE ... SET stock = stock - :cant WHERE id_insumo = :id AND stock >= :cant",
        "   y NO leyendo primero el stock y decidiendo despues:",
        "   [ ... ]",
        "",
        "D) GAP HONESTO   ·   lo que no se pudo comprobar y como se aborda",
        "   No pude comprobar: ____________________________________________________",
        "     (PostgreSQL en el navegador corre UNA SOLA sesion: dos recepcionistas",
        "      facturando el mismo insumo al mismo tiempo no se puede montar aqui)",
        "   Como lo abordare en la Clase 10: ______________________________________",
    ],
}


def build_taller_docx(c):
    if c['tipo']=='parcial': return None
    tb = TALLER_BLOQUE.get(c['n'], {})
    doc = Document()
    _titulo_banda = ("Guia de sustentacion PI · Clase %d · Bases de Datos II" % c['n']
                     if c['tipo'] == 'sustentacion'
                     else "Taller PI · Clase %d · Bases de Datos II" % c['n'])
    banda(doc, _titulo_banda)
    para(doc, c['titulo'], size=14, bold=True, color=AZUL)
    para(doc, "Hilo conductor: Proyecto Integrador VetCare DB (no es un ejercicio desconectado).", size=11, bold=True)
    para(doc, f"Herramienta: {c['herramienta']}")
    para(doc, f"Hoy avanzamos el PI en: {c['hito_pi']}", shade_fill="FFF8D6")
    if c['n'] == 1:
        # El cliente se presentaba solo en la diapositiva. Aqui queda en el
        # documento que el estudiante conserva, con el puntero al anexo completo.
        para(doc, "0. El cliente: " + CLIENTE, size=12, bold=True, color=AZUL)
        _p = doc.add_paragraph(); _p.paragraph_format.space_after = DocPt(6)
        add_inline_docx(_p, "Atiende unas @@150 citas al día@@ con @@16 veterinarios@@, sobre "
                            "unas 5.000 mascotas de unos 2.000 dueños, y hoy lleva todo en "
                            "@@carpetas de papel@@. La administración reporta tres problemas:")
        bullets(doc, [p.replace("@@", "") for p in PROBLEMAS])
        _p = doc.add_paragraph(); _p.paragraph_format.space_after = DocPt(6)
        add_inline_docx(_p, "Usted construye la @@capa de datos@@ de VetCare: el modelo, la "
                            "integridad, la seguridad y el rendimiento. La aplicación no se "
                            "pide en esta asignatura.")
        _p = doc.add_paragraph(); _p.paragraph_format.space_after = DocPt(6)
        shade(_p, "E8F4FA")
        add_inline_docx(_p, "@@Caso completo (téngalo a mano todo el semestre):@@ "
                            "Clases/Proyecto Integrador/Anexo - Caso de estudio Clinica "
                            "Huellitas — las 8 entidades, las 3 reglas de negocio, el elenco de "
                            "nombres y cómo crece la base hasta las 30.010 citas que se "
                            "optimizan en las Clases 6 y 7.")
    para(doc, "1. Contexto / por que importa al PI", size=12, bold=True, color=AZUL)
    bullets(doc, tb.get('contexto') or ["Trabaje sobre su propio dominio VetCare."])
    para(doc, "2. Objetivo", size=12, bold=True, color=AZUL)
    para(doc, tb.get('objetivo', c['hito_pi']))
    para(doc, "3. Escenario / datos de partida", size=12, bold=True, color=AZUL)
    bullets(doc, tb.get('escenario') or ["Usar su propio DDL/ER de VetCare."])
    para(doc, "4. Actividades (pasos guiados)", size=12, bold=True, color=AZUL)
    bullets(doc, c['taller'])
    _n_sec = 5
    if PLANTILLA_FICHA.get(c['n']):
        # La plantilla vive DENTRO del taller a proposito: cuando estaba solo en el
        # enunciado del PI, el estudiante improvisaba la estructura y cada entrega
        # llegaba con secciones distintas, imposibles de comparar contra la rubrica.
        para(doc, f"{_n_sec}. Plantilla del entregable (copia esto y llenalo)",
             size=12, bold=True, color=AZUL)
        para(doc, "Estos son exactamente los campos que califica ExamLab. Copia el bloque "
                  "tal cual en tu documento, llenalo, y pega cada parte en la pregunta que "
                  "corresponda.", size=10)
        for linea in PLANTILLA_FICHA[c['n']]:
            para(doc, linea or " ", size=9, shade_fill="F2F2F3", space_after=0, font=MONO)
        para(doc, " ", size=6, space_after=0)
        _n_sec += 1
    para(doc, f"{_n_sec}. Entregable", size=12, bold=True, color=AZUL)
    para(doc, c['entregable'], shade_fill="E8F4FA")
    para(doc, f"{_n_sec + 1}. Criterios de exito", size=12, bold=True, color=AZUL)
    bullets(doc, tb.get('criterios') or [
        "Avance real de su propio VetCare.",
        "Evidencia ejecutable o diagrama exportado.",
        "Criterio de rubrica del PI movido hoy.",
    ])
    para(doc, f"{_n_sec + 2}. Pistas (checklist vacio — sin solucion)", size=12, bold=True, color=AZUL)
    bullets(doc, tb.get('pistas') or ["Revisar evidencia antes de subir."])
    para(doc, f"{_n_sec + 3}. Entrega", size=12, bold=True, color=AZUL)
    _p_entrega = doc.add_paragraph(); _p_entrega.paragraph_format.space_after = DocPt(6)
    if c['tipo'] == 'sustentacion':
        # No es un taller con plazo del domingo: la sesion es la sustentacion en vivo,
        # asi que el paquete tiene que estar arriba ANTES del bloque.
        add_inline_docx(_p_entrega, "@@Sube el paquete final en ExamLab@@ (https://uniaj.examlab.workers.dev/ · módulo Proyectos) "
                                    "ANTES de tu turno: presentando no se sube nada. La sustentación es EN VIVO en la sesión "
                                    "de clase (5-8 min + Q&A del docente); no se reemplaza por un video grabado.")
    else:
        add_inline_docx(_p_entrega, "@@Sube tu taller en ExamLab@@ (https://uniaj.examlab.workers.dev/ · módulo Talleres) — domingo 23:59 cuando aplique.")
    # 9. Que encuentra en la plataforma. Antes el taller terminaba en «sube esto a
    # ExamLab» sin decir en que forma se responde cada cosa; con esto el estudiante
    # sabe de antemano que hay editor de SQL con PostgreSQL real, cuadro de texto, etc.
    _taller_el = TALLERES_EXAMLAB.get(c["n"])
    if _taller_el:
        examlab_talleres.render_estudiante(
            doc, _taller_el, para=para, bullets=bullets,
            add_inline=add_inline_docx, color_titulo=AZUL,
            titulo=f"{_n_sec + 4}. Que vas a resolver en ExamLab",
        )
    out_dir = CLASES_DIR / f"Clase {c['n']} - {c['slug']}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"Taller PI - Clase {c['n']} - VetCare.docx"
    doc.save(str(out)); print("TALLER", out); return out


def _build_solucion_completa(c):
    """Solucion pregunta por pregunta, con el render compartido con Arquitectura."""
    n = c['n']
    md = solucion_taller.render_md(
        n, soluciones_bd2.SOLUCION[n],
        contexto={
            "alineacion": [
                ("Taller del estudiante",
                 f"`Clases/Clase {n} - {c['slug']}/Taller PI - Clase {n} - VetCare.docx`"),
                ("Configuracion en la plataforma",
                 f"`Kit docente/Clase {n}/Taller en ExamLab - Clase {n} (configuracion).md`"),
                ("Caso de estudio",
                 "`Clases/Proyecto Integrador/Anexo - Caso de estudio Clinica Huellitas - "
                 "Bases de Datos II.docx`"),
                ("Hito del PI", c.get('hito_pi', '—')),
                ("Entregable", c.get('entregable', '—')),
            ],
            "politica_extra": ("El motor de la plataforma es PostgreSQL (PGlite en el "
                               "navegador), no Oracle."),
        },
        opciones=soluciones_bd2.opciones,
        # Sin `mermaid_referencia` a proposito: en BD II el dominio es fijo (VetCare),
        # asi que el modelo de la solucion Y el que ve el estudiante son el mismo y
        # mostrarlo dos veces solo alarga el documento. En Arquitectura si tiene
        # sentido, porque alli la solucion usa un dominio distinto del proyectado.
    )
    kit = KIT_DIR / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    out_md = kit / f"Solucion Taller Clase {n} - VetCare.md"
    out_md.write_text(md, encoding='utf-8')
    print("SOLUCION md", out_md)
    convert_guion(out_md)          # mismo conversor que los guiones: tablas y bloques
    return out_md


def build_solucion_docx(c):
    if c['tipo']=='parcial': return None
    # Formato nuevo (pregunta por pregunta, con criterio de calificacion) si esta
    # migrada; si no, el formato corto de siempre. Asi la migracion va clase por
    # clase sin romper el build de las que faltan.
    if c['n'] in soluciones_bd2.SOLUCION:
        return _build_solucion_completa(c)
    sol = SOLUCION.get(c['n'])
    if not sol: return None
    kit = KIT_DIR / f"Clase {c['n']}"
    kit.mkdir(parents=True, exist_ok=True)
    stem = f"Solucion Taller Clase {c['n']} - VetCare"
    lines = [f"# {sol['titulo']}", "",
             "> DOCUMENTO DOCENTE — PRIVADO. No publicar en Clases/.", "",
             f"**Resumen:** {sol['resumen']}", "",
             "## Alineacion",
             f"- Taller: `Clases/Clase {c['n']} - {c['slug']}/Taller PI - Clase {c['n']} - VetCare.docx`",
             f"- Hito: {c['hito_pi']}", f"- Entregable: {c['entregable']}", "",
             "## Solucion paso a paso"]
    for i, s in enumerate(sol.get('pasos', []), 1):
        lines.append(f"{i}. {s}")
    lines += ["", "## Ejemplo / SQL / artefactos"]
    for e in sol.get('ejemplo', []):
        lines.append(f"- {e}")
    if c.get('sql'):
        lines.append(f"- Script demo: `Kit docente/Clase {c['n']}/Codigo/{c['sql']}`")
    lines += ["", "## Rubrica corta"]
    for r in sol.get('rubrica', []):
        lines.append(f"- [ ] {r}")
    lines += ["", "## Errores frecuentes"]
    for e in sol.get('errores', []):
        lines.append(f"- {e}")
    lines += ["", "Entrega en ExamLab.", ""]
    (kit / f"{stem}.md").write_text("\n".join(lines), encoding="utf-8")
    doc = Document(); banda(doc, sol['titulo'])
    para(doc, "DOCUMENTO DOCENTE — PRIVADO (no va en Clases/)", bold=True, color=RGBColor(0xA0,0x20,0x30), shade_fill="FBE4E4")
    para(doc, sol['resumen'], shade_fill="E8F4FA")
    para(doc, "Alineacion al enunciado", size=12, bold=True, color=AZUL)
    bullets(doc, [
        f"Taller: Clases/Clase {c['n']} - {c['slug']}/Taller PI - Clase {c['n']} - VetCare.docx",
        f"Hito: {c['hito_pi']}",
        f"Entregable: {c['entregable']}",
    ])
    para(doc, "Solucion paso a paso", size=12, bold=True, color=AZUL)
    bullets(doc, sol.get('pasos', []))
    para(doc, "Ejemplo / SQL / artefactos", size=12, bold=True, color=AZUL)
    bullets(doc, sol.get('ejemplo', []) + ([f"Script demo: Codigo/{c['sql']}"] if c.get('sql') else []))
    para(doc, "Rubrica corta", size=12, bold=True, color=AZUL)
    bullets(doc, ["[ ] "+r for r in sol.get('rubrica', [])])
    para(doc, "Errores frecuentes", size=12, bold=True, color=AZUL)
    bullets(doc, sol.get('errores', []))
    out = kit / f"{stem}.docx"
    doc.save(str(out)); print("SOLUCION", out); return out


# Imagenes de SALIDA ESPERADA por clase (generadas por config/slides/mockups.py).
# El guion las embebe con el token [[captura: archivo.png]]; si el archivo no
# existe, guion_md_a_docx.py deja la caja "inserta aqui la captura" y no rompe.
CAPTURAS_CLASE = {
    1: [("Resultado del JOIN de verificacion del ER (lo que debe salir tras los INSERT)",
         "salida-join-vetcare.png")],
    3: [("Bateria de pruebas de sp_agendar_cita: P1 OK y P2 rechazado por mascota inactiva",
         "salida-proc-ok-y-error.png")],
    4: [("trg_audit_cita: los 3 UPDATE dejan 2 filas de auditoria (el WHEN filtra el tercero)",
         "cap01_demo.png")],
    # Decia «(FULL SCAN -> INDEX RANGE SCAN)»: nombres de nodo de Oracle, y ademas
    # una promesa que la clase no puede cumplir, porque hoy no se crea ningun indice
    # (eso es la Clase 7). El docente que intentara producir esa captura no podia.
    6: [("EXPLAIN ANALYZE ANTES vs DESPUES: el nodo no cambia, las pasadas si (loops 2006 -> 1)",
         "salida-explain-antes-despues.png")],
    7: [("El plan de C1 antes y despues: Seq Scan -> Index Scan using idx_cita_programada_fecha",
         "salida-indice-antes-despues.png")],
    8: [("CALL sp_facturar que falla a mitad: foto inicial y foto final identicas, sin ROLLBACK escrito",
         "salida-rollback-stock.png")],
    10: [("Evidencia del problema: dos citas en la misma franja (sin restriccion)",
          "salida-doble-reserva.png"),
         ("El MISMO INSERT ya con UNIQUE: la BD lo rechaza sola",
          "salida-unique-rechaza.png")],
}


#: Receta por defecto para la captura de la demo cuando todavia no hay PNG. El
#: marcador `[CAP: ...]` que habia antes llegaba literal al .docx y no le decia al
#: docente ni que abrir ni con que nombre guardar, asi que la caja seguia vacia.
def _receta_demo(c):
    return (
        f"1) Abra {c['herramienta']} y repita la demo de este bloque sobre el dominio "
        "VetCare (no otro ejemplo).  "
        "2) Capture la ventana en el momento en que se ve el resultado, no el "
        "escritorio completo.  "
        "3) Recorte a ~1200 px de ancho.  "
        f"4) Guardela como Kit docente/Clase {c['n']}/Capturas/cap01_demo.png.  "
        "5) Vuelva a generar el guion: la imagen queda embebida aqui sola."
    )


def _capturas_md(n, c=None):
    """Lineas 📸 del guion para la clase n (con token explicito de archivo).

    Cuando la clase no tiene ilustracion de «salida esperada» generada por
    `mockups.py`, en vez de un marcador crudo se emite el token con la receta para
    producir la captura: es la regla transversal de recursos no generables.
    """
    items = CAPTURAS_CLASE.get(n)
    if not items:
        receta = _receta_demo(c) if c else (
            f"1) Abra la herramienta de la demo y repita los pasos sobre VetCare.  "
            f"2) Capture solo la ventana util.  3) Recorte a ~1200 px de ancho.  "
            f"4) Guardela como Kit docente/Clase {n}/Capturas/cap01_demo.png.  "
            "5) Vuelva a generar el guion."
        )
        return (f"📸 Salida esperada de la demo de la Clase {n} "
                f"[[captura: cap01_demo.png | receta: {receta}]]\n")
    return "".join(f"📸 {cap} [[captura: {fn}]]\n" for cap, fn in items)


#: Lo del dia de parcial que NO esta en el instrumento. El resto —archivo, temas,
#: secciones con sus puntos, peso, fecha— se lee de `contenido_parciales_2026_2`.
#:   corte    · numero del corte que cierra; indexa BD2_P1/P2/P3
#:   ultima   · ultima clase dictada antes del parcial, la que el grupo tiene fresca
#:   dudas_no · tres dudas REALES de contenido de ESE corte, para que el ejemplo de «esto
#:              no se responde» sea del parcial que se esta aplicando
PARCIALES_BD2 = {
    5: {"corte": 1, "ultima": 4,
        "dudas_no": ["¿la diferencia entre funcion y procedimiento es que la funcion retorna valor?",
                     "¿GRANT y REVOKE tienen que ver con el control de acceso?",
                     "¿esta opcion es la correcta?"]},
    9: {"corte": 2, "ultima": 10,
        "dudas_no": ["¿que nivel de aislamiento evita la lectura fantasma?",
                     "¿que es un deadlock?",
                     "¿esta opcion es la correcta?"]},
    14: {"corte": 3, "ultima": 13,
         "dudas_no": ["¿embeber la contraseña en el codigo pasa si el repo es privado?",
                      "¿hay que validar las entradas de la API?",
                      "¿esta opcion es la correcta?"]},
}


def _parcial_meta_bd2(corte):
    """Portada del instrumento del corte `corte`, leida de la fuente del .docx.

    Se lee y no se copia: la guia decia «Entregar enunciado (impreso / PDF)» y nada mas,
    asi que ni el alcance ni el reparto de puntos estaban en ninguna parte del kit, y el
    unico dato operativo que daba —«impreso»— es imposible en una sesion virtual.
    """
    _dir = str(ROOT / "config" / "parciales")
    if _dir not in sys.path:
        sys.path.insert(0, _dir)
    import contenido_parciales_2026_2 as cp  # noqa: PLC0415
    return getattr(cp, f"BD2_P{corte}")["meta"]


def _cierre_pi_parcial(n, hablada=False):
    """Que pasa con el PI despues de este parcial.

    El deck y la guia decian los dos «el PI continua en la siguiente clase», que es falso
    en el Parcial 3: lo que sigue es la sustentacion, y es justo el aviso que ese dia hay
    que dar. Se calcula de `CLASES` para no depender de un numero escrito a mano.
    """
    sig = next((x['n'] for x in CLASES if x['n'] > n), None)
    if sig and sig == CLASES[-1]['n']:
        return (f"En la Clase {sig} es la **sustentación del Proyecto Integrador**: se "
                "presenta el VetCare DB que ya construyeron, no algo nuevo."
                if hablada else
                f"Clase {sig}: sustentación del PI VetCare DB, no hay tema nuevo")
    return ("El PI VetCare continúa en la siguiente clase; hoy no hay tarea nueva."
            if hablada else "El PI VetCare DB continúa en la siguiente clase")


def _guia_parcial_md(c):
    """Guia del dia de parcial, con los `{{slide:…}}` resueltos contra el deck real.

    Mismo mecanismo que el fundamento de una clase regular: la guia nombra la diapositiva
    por su titulo y el numero se resuelve contra el deck, asi que si el deck del dia de
    parcial cambia de orden el build falla en vez de mandar al docente a proyectar la
    diapositiva equivocada.
    """
    return _resolver_slides(_guia_parcial_cuerpo(c), _slide_map(c), c['n'])


def _guia_parcial_cuerpo(c):
    """Guia de aplicacion del dia de parcial.

    Eran 20 lineas que no cumplian el checklist del repo: no decian la modalidad, la
    unica instruccion de reparto era «impreso / PDF» —imposible en una sesion virtual
    sincrona por Meet—, no decian por que canal se recibe el parcial, no listaban los
    temas ni las secciones del instrumento, y no traian ni «errores tipicos del docente»
    ni las preguntas frecuentes del grupo. Con 90 minutos de silencio en el bloque, los
    30 que quedan son todo lo que la guia tiene que resolver.
    """
    n = c['n']
    p = PARCIALES_BD2[n]
    m = _parcial_meta_bd2(p['corte'])
    clases_ev = [t.split(' · ')[0].replace('Clase ', '') for t in m['temas']]
    lista = (', '.join(clases_ev[:-1]) + ' y ' + clases_ev[-1]
             if len(clases_ev) > 1 else clases_ev[0])
    temas_md = '\n'.join(f'- {t}' for t in m['temas'])
    secciones_md = '\n'.join(f'- {s}' for s in m['secciones_resumen'])
    no_resp = ' · '.join(f'«{d}»' for d in p['dudas_no'])
    sig = next((x['n'] for x in CLASES if x['n'] > n), None)
    cierre_pi = _cierre_pi_parcial(n, hablada=True)
    # Los tokens se arman fuera del f-string: dentro, `{{` es una llave literal escapada y
    # el token saldria como «{slide:evalua hoy}», que ni resuelve ni se detecta como
    # marcador crudo por su forma habitual.
    tok_alcance, tok_entrega = '{{slide:evalua hoy}}', '{{slide:Como se responde}}'
    # La clase que el grupo cree que no cuenta es la autonoma: no hubo sesion en vivo. La
    # portada del instrumento la marca «(sesion autonoma)». Cuando la autonoma ES la ultima
    # clase evaluada —el Parcial 2, cuya Clase 10 cayo en la sesion del festivo— la
    # aclaracion va dentro de la misma respuesta, para no preguntar dos veces por la misma
    # clase.
    auton = [t.split(' · ')[0] for t in m['temas'] if 'autónoma' in t]
    razon_auton = (" y se evalua igual, aunque se haya trabajado sin sesion en vivo: es la "
                   "que mas se olvida al estudiar")
    if auton and auton[0] == f"Clase {p['ultima']}":
        faq_entra = (f"**¿Entra lo de la Clase {p['ultima']}?** Si, es lo mas reciente que "
                     f"entra{razon_auton}.")
        faq_auton = ""
    else:
        faq_entra = (f"**¿Entra lo de la Clase {p['ultima']}?** Si, y es lo mas reciente que "
                     "entra.")
        faq_auton = (f"\n**¿La clase autonoma tambien entra?** Si: la {auton[0]} se trabajo sin "
                     "sesion en vivo y se evalua igual. Es la que mas se olvida al estudiar.\n"
                     if auton else "")
    # Las tres clases del PI se leen de `CLASES`: la nota decia «Prep PI / sustentacion:
    # Clases 11-12» en los tres parciales, que ademas de estar escrito a mano ponia la
    # sustentacion en la 11-12 cuando es la ultima clase del curso.
    prep = next((x['n'] for x in CLASES if 'Prep' in x['titulo']), None)
    sust = next((x['n'] for x in CLASES if x.get('tipo') == 'sustentacion'), None)
    if prep and prep < n:
        linea_pi = (f"- La preparacion de la presentacion fue la **Clase {prep}**; la "
                    f"sustentacion del PI es la **Clase {sust}**. Hoy no se avanza en el "
                    "proyecto.")
    else:
        # Detras del Parcial 2 la siguiente clase es la autonoma del festivo: si la nota
        # dice «se retoma en la Clase 10» a secas, el docente la lee como sesion en vivo y
        # programa un avance de PI que ese dia no tiene con quien hacerse.
        auton_sig = next((x['n'] for x in CLASES
                          if x['n'] == sig and x.get('tipo') == 'autonoma'), None)
        etiqueta = f"**Clase {sig}**" + (" (sesion autonoma)" if auton_sig else "")
        linea_pi = (f"- Hoy no se avanza en el PI: se retoma en la {etiqueta}. La "
                    f"preparacion de la presentacion es la **Clase {prep}** y la sustentacion "
                    f"la **Clase {sust}**.")
    return f"""# Guia docente · Clase {n} · {c['titulo']} (solo evaluacion)

> Dia de **parcial = solo evaluacion**. No hay tema nuevo ni avance de PI en clase.
> Bloque **120 min** · **virtual sincrona por Google Meet**.
> Enunciado: `Parciales/{c.get('parcial','')}` · la solucion es el mismo nombre con
> «- SOLUCION» y **no se publica** en `Clases/`.

## Que evalua el instrumento

Solo estas clases de material (asi las lista la portada del enunciado):

{temas_md}

Sus cuatro secciones y lo que vale cada una:

{secciones_md}

Total **100 puntos** · nota = puntos / 20 sobre 5.0 · peso **{m['valor_corte']}** ·
fecha **{m['fecha']}** · tiempo de resolucion previsto **{m['tiempo']}**.

## Antes de abrir la sesion (10 min)

1. Abre `Parciales/{c.get('parcial','')}` y **decide el canal de entrega**. El enunciado
   remite a «el medio que el docente indique al abrir la sesion», asi que si no lo
   decides tu, no existe. Lo que funciona en Meet:
   - **Documento editable** (recomendado): compartes el .docx por el chat al minuto 0,
     cada estudiante lo llena y lo devuelve por el mismo canal o por correo. El SQL se
     escribe como texto; **no** se pide captura de ejecucion en el parcial.
   - **Foto de hoja escrita a mano**: solo como plan B si a alguien no le abre el
     documento. Exige que se lea y que traiga nombre en cada pagina.
2. Ten la solucion **a la mano pero cerrada**: hoy no se califica en vivo, y menos con
   la pantalla compartida.
3. Revisa que el enunciado no pregunte nada fuera de las Clases {lista}. Si algo se cuela
   de otro corte se anula esa pregunta y se reparten sus puntos, no se descuenta.

## Checklist 120 min

| Min | Accion |
|---|---|
| 0-10 | Asistencia por lista. Proyecta la **{tok_alcance}** (alcance y reparto de puntos, que es lo primero que preguntan) y luego la **{tok_entrega}**. Anuncia: canal de entrega, cierre en el minuto 110, que material esta autorizado (por defecto **nada**) y que las dudas de contenido no se responden. |
| 10-15 | Comparte el enunciado y **confirma en voz alta que todos lo abrieron** antes de arrancar el reloj. Deja la **{tok_entrega}** en pantalla: ahorra la mitad de los mensajes por privado. |
| 15-100 | Desarrollo (silencio de evaluacion). Camara y microfono abiertos: es la unica supervision que hay. Avisa el tiempo a los 50 y a los 80 minutos. |
| 100-110 | Aviso de 10 min. Recibe las entregas y **acusa recibo por el chat, uno por uno**. Anota quien no entrego. |
| 110-120 | Cierre. «{cierre_pi}» Sin comentarios sobre el parcial: todavia hay quien esta subiendo el archivo. |

## Que se responde y que no durante el parcial

La linea es una sola: **si la respuesta a la duda es un dato que la pregunta evalua, no se responde.**

- Se responde: «¿esto pide una consulta o una explicacion?», «¿cuantas lineas?», «¿el
  punto b) es obligatorio?», «no puedo abrir el archivo».
- No se responde, en este parcial: {no_resp}.

Cuando la duda es de contenido, la respuesta es siempre la misma: «Eso es lo que la
pregunta evalua; responde con lo que recuerdes de la clase.» Dila igual para todos: en
Meet las preguntas llegan por privado y nadie ve que a otro le dijiste lo mismo.

## Errores tipicos del docente que no domina el tema

- **Responder la duda de contenido porque parece inofensiva.** «{p['dudas_no'][0]}» es
  literalmente la respuesta de una pregunta de este parcial.
- **No decidir el canal de entrega antes de empezar.** Si no lo anunciaste al minuto 0,
  lo vas a improvisar al minuto 105 con medio grupo escribiendo por privado.
- **No acusar recibo.** Es la fuente numero uno de reclamos de un parcial virtual y se
  resuelve escribiendo el nombre de cada uno en el chat cuando llega su archivo.
- **Exigir salida de ejecucion.** El parcial se responde con SQL escrito; ExamLab no
  interviene y nadie tiene que abrir un motor. Si un estudiante escribe una consulta
  correcta con un nombre de tabla que no existe en el enunciado, se descuenta por el
  nombre, no por la consulta.
- **Descontar por el termino de manual.** Antes de restar puntos por una palabra que el
  estudiante no uso, busca la diapositiva donde se proyecto. Vale la respuesta que
  describe el mecanismo correcto aunque no lo nombre.
- **Sintaxis de otro motor.** Las clases se dictan sobre **PostgreSQL**; una respuesta
  con sintaxis de Oracle que expresa bien la idea no pierde los puntos del concepto.
- **Tratar el dia como clase.** Ni tema nuevo, ni avance del PI, ni «aprovechemos que
  terminaron temprano».

## Preguntas frecuentes del grupo

**¿Puedo usar mis apuntes?** Lo que digas al minuto 0 y nada mas. Por defecto: no.
Decidelo antes de abrir la sesion: cambiarlo a mitad del parcial invalida el de quien ya
respondio sin ellos.

**¿Se me cayo el internet?** Que siga respondiendo el documento sin conexion y te escriba
por correo al reconectarse. El tiempo perdido por una caida comprobable no se descuenta, y
el criterio se anuncia al minuto 0 para que nadie lo use como excusa despues.

{faq_entra} La portada del
enunciado lista las clases evaluadas con su fecha; fuera de esa lista no hay nada.
{faq_auton}
**¿Tengo que ejecutar el SQL?** No. Se responde escrito y se califica la logica de la
consulta: los alias, el JOIN, el WHERE, el orden. No se pide captura.

**¿Cuanto vale cada seccion?** Esta en la portada: {' · '.join(m['secciones_resumen'])}.
Total 100 puntos, nota = puntos / 20.

**¿Cuanto tiempo tengo?** El instrumento preve {m['tiempo']} y el bloque son 120: hay
holgura, pero el envio cierra en el minuto 110 y eso no se mueve.

**¿Cuando veo la nota?** En la siguiente sesion, con la retroalimentacion escrita sobre el
mismo documento que entregaste.

## Notas

- No mezclar «Tema · Parcial»: hoy no se dicta nada.
{linea_pi}
- Solucion privada: archivo «- SOLUCION.docx» en `Parciales/`, nunca en `Clases/`.
"""


def build_guion_md(c):
    kit = KIT_DIR / f"Clase {c['n']}"
    kit.mkdir(parents=True, exist_ok=True)
    (kit/"Capturas").mkdir(exist_ok=True)
    (kit/"Codigo").mkdir(exist_ok=True)
    # Las clases sin script (parcial, autonoma, sustentacion) dejaban Codigo/ vacia,
    # y una carpeta vacia no existe en git: al clonar el repo desaparecia y quedaba
    # la duda de si el .sql se perdio. El README dice que no hay ninguno y por que.
    if not c.get('sql'):
        _titulo_cod = f"Codigo de la Clase {c['n']} — Bases de Datos II"
        (kit/"Codigo"/"README.txt").write_text(
            f"{_titulo_cod}\n{'=' * len(_titulo_cod)}\n\n"
            f"Esta clase es de tipo «{c['tipo']}» y no trae script ejecutable: no falta\n"
            "ningun archivo. Los .sql del curso viven en las clases regulares y en la\n"
            "Clase 10; el catalogo esta en SQL_BODIES, dentro de\n"
            "config/slides/build_uniajc_bd2_all.py.\n",
            encoding='utf-8')
    if c['tipo']=='parcial':
        md = _guia_parcial_md(c)
        path = kit / f"Guia aplicacion {c['titulo']} - Clase {c['n']}.md"
        path.write_text(md, encoding='utf-8')
        # placeholder capturas
        (kit/"Capturas"/".gitkeep").write_text("", encoding='utf-8')
        return path

    # Referencias a diapositivas: se derivan del mismo mapa que arma el deck, para
    # que el docente pueda leer el guion con la presentacion proyectada sin perder
    # el hilo. Antes era una lista escrita a mano y estaba corrida.
    mapa = _slide_map(c)
    sl_teoria = _slide_tag(mapa, "Teoria Core")
    # La teoria son VARIAS diapositivas, no solo «Teoria Core»: el plan las enumera
    # todas con su numero real y reparte los 25 min entre ellas, para que ninguna
    # quede proyectada por nadie.
    teoria_slides = _slides_teoria(mapa)
    teoria_slides_md = "\n".join(
        f"{k}. **[Slide {i}] {t}**" for k, (i, t) in enumerate(teoria_slides, 1))
    # El suelo era de 4 min y no se comprobaba contra el bloque: con 7 diapositivas de
    # teoria el plan anunciaba «~4 min cada una», o sea 28 min dentro del tramo 10-35
    # que el propio guion declara de 25. Un docente que sigue el reparto al pie de la
    # letra llega tarde a la demo. Con el suelo en 2 el reparto siempre cabe (hasta 12
    # diapositivas) y sigue sin proponer medios minutos.
    min_por_slide = max(2, 25 // max(1, len(teoria_slides)))
    sl_demo = _slide_tag(mapa, "Demo del dia", "Como se ordena")
    sl_flujo = _slide_tag(mapa, FLUJO_SLIDE_TITULO)
    sl_taller = _slide_tag(mapa, "pasos guiados")
    sl_crit = _slide_tag(mapa, "Criterios de exito")
    sl_cierre = f"[Slide {len(mapa)}] "  # el cierre es siempre la ultima
    tipo = TIPO_LABEL[c['tipo']]
    # Si el taller pide un diagrama, la demo tiene que terminar en ExamLab y no en
    # el PNG: el docente demuestra tambien la conversion a Mermaid.
    if _tiene_diagrama(c['n']):
        _dial = examlab_talleres._dialectos_del_taller(TALLERES_EXAMLAB[c['n']])
        flujo_guion = (
            "\n**Cierre la demo dentro de ExamLab** " + sl_flujo + "— es la parte que el "
            "estudiante no adivina: pase el boceto a codigo Mermaid con ayuda de una IA, "
            "peguelo en la pregunta de diagrama y muestrelo renderizado.\n\n"
            + examlab_talleres.flujo_diagrama_md(
                _dial[0] if len(_dial) == 1 else "el tipo que pide el enunciado")
            + "\n"
        )
    else:
        flujo_guion = ""
    bloques = ""
    if c['tipo'] == 'sustentacion':
        plan = """## Plan minuto a minuto (120 min) — sesion de SUSTENTACIONES EN VIVO

> Este bloque es sincrono y se dedica completo a las sustentaciones del PI VetCare DB.
> **No es clase autonoma y no es parcial.** No autorice reemplazar la defensa por un video
> grabado: el Q&A dirigido al azar es el unico instrumento con el que verifica que el modelo,
> los procedimientos y la optimizacion son de quien los presenta. El dia cae en festivo de
> calendario, pero la sesion esta destinada por decision docente a sustentar: anunciela por
> escrito la semana anterior para que nadie asuma que no hay clase.

### Antes de la sesion (semana previa)
1. Publique el orden y la duracion del turno: **5-8 min de pitch + 2-4 min de Q&A**. Con 12
   sustentaciones son ~110 min; si el grupo es mas grande, baje a 5 + 2 y avisele antes.
2. Exija el paquete subido a ExamLab (modulo Proyectos) **antes** del bloque, y abra usted
   mismo dos o tres ZIP en un playground limpio: quien llega a subir archivos consume su turno.
3. Tenga la rubrica impresa por estudiante y las preguntas de Q&A ya escogidas por tipo
   (verificacion, profundizacion, hipotetica), para no preguntar lo mismo a todos.

### 0-10 · Encuadre y orden de turnos
**Decir:** «Hoy sustentamos. De 5 a 8 minutos de pitch y hasta 4 de preguntas. Corto a los 8
minutos: si no llegaron a optimizacion, esa parte no se califica. El orden lo sorteo ahora.»
Sortee el orden delante del grupo, proyecte el cronometro y pida que el resto escuche.

### 10-110 · Sustentaciones (turnos consecutivos)
Por cada turno:
1. **5-8 min de pitch.** No interrumpa ni para corregir: anote y pregunte despues. Exija que se
   vea al menos **una ejecucion real** (procedimiento con caso valido e invalido, o el plan de
   ejecucion antes/despues), no solo capturas fijas.
2. **2-4 min de Q&A.** Una pregunta de verificacion («muestreme el DDL de esa tabla»), una de
   profundizacion («por que esa regla esta en un disparador y no en la aplicacion») y, si hay
   tiempo, una hipotetica («si manana entran cien mil citas, que consulta se cae primero»). Si
   autorizo equipo, dirija cada pregunta a un integrante distinto.
3. **Cierre el turno con la nota puesta**, no al final del dia.

### 110-120 · Cierre del curso
**Decir:** «Lo que entregaron —ER justificado, DDL con restricciones, matriz de privilegios,
procedimientos con manejo de errores, disparadores y analisis de plan— es el contenido real de
las tareas de un desarrollador de base de datos junior. Conserven el repositorio.»
Recuerde los pesos sin abrir discusion de notas: el PI vale **20% del Corte 3** y el Parcial 3
ya se aplico en su propia sesion; el proyecto no lo reemplaza ni lo compensa.

### Si alguien no se presenta o falla la conexion
Deje constancia escrita en el momento (hora, motivo) y reprograme dentro de la misma semana por
Meet, sustentando igualmente en vivo. Aceptar un video «por esta vez» elimina el Q&A, que es la
mitad de lo que se evalua, y vuelve regla la excepcion el semestre siguiente.
"""
    elif c['tipo']=='autonoma':
        plan = """## Plan minuto a minuto (120 min equivalentes — trabajo autonomo)

> El estudiante trabaja sin encuentro sincrono. Usted publica este guion resumido + taller en ExamLab.

### Bloque A (0-20) · Encuadre PI
**Decir/publicar:** «Hoy avanzamos el PI en: {hito}. No es un taller suelto.»
Referencia slides: Agenda + Objetivo PI.

### Bloque B (20-45) · Teoria minima
Leer Teoria Core. Tomar notas en el informe del PI.

### Bloque C (45-100) · Practica = entregable PI
Seguir el taller estudiante. Herramienta: {herr}.
Salida esperada de la practica (publiquela junto al enunciado para que el
estudiante autonomo sepa si le quedo bien):
{caps}
### Bloque D (100-120) · Empaquetado y cierre
Subir entregable a ExamLab. Actualizar el checklist PI del proyecto.
""".format(hito=c['hito_pi'], herr=c['herramienta'], caps=_capturas_md(c['n'], c))
    else:
        plan = f"""## Plan minuto a minuto (120 min) — texto casi literal

### 0-10 · Encuadre · {_slide_tag(mapa, "Encuadre de hoy").strip()}{_slide_tag(mapa, "Mapa del bloque").strip()}
**Decir:** «Buenas. Hoy el hilo es VetCare DB. Avanzamos el PI en: {c['hito_pi']}.
La teoria sera corta; el peso esta en el taller del proyecto.»
Proyectar {_slide_tag(mapa, "Encuadre de hoy") or "la slide de"}«Encuadre de hoy · Objetivo PI» y {_slide_tag(mapa, "Mapa del bloque")}«Mapa del bloque de hoy».
Pasar asistencia. Recordar herramientas gratis+nube.

### 10-35 · Teoria Core (breve) · desde {sl_teoria.strip()}
**Decir:** «Solo lo necesario para el entregable de hoy.»
Proyecte estas diapositivas, en este orden, ~{min_por_slide} min cada una. Son la teoria
completa del dia: **ninguna se salta**, porque el taller cobra puntos por lo que se
proyecta en todas ellas.
{teoria_slides_md}

El desarrollo completo de cada una esta arriba, en «Fundamento teorico», dividido por
diapositiva: esa seccion esta escrita para dictarla sin consultar otra fuente.
Ideas que tienen que quedar dichas:
""" + "\n".join(f"- {t}" for t in c['teoria']) + f"""
Pregunta al aire (2 min): ¿como se conecta esto con su VetCare?

### 35-55 · Demo paso a paso · {sl_demo.strip()}{sl_flujo.strip()}
**Decir:** «Miren mi pantalla. Dominio VetCare — no otro ejemplo.»
Demo: {c['demo']}
Herramienta: {c['herramienta']}
{flujo_guion}""" + _capturas_md(c['n'], c) + f"""Dejar script/enlace en el chat o en ExamLab.

### 55-105 · Taller guiado = tarea del PI · {sl_taller.strip()}
**Decir:** «Abran su carpeta VetCare. Esto suma a la rubrica del PI. Al final suben el taller en ExamLab.»
Usar bloque Taller ampliado (contexto->pistas). Solucion en Kit docente/Solucion Taller... (no proyectar completa).
Actividades:
""" + "\n".join(f"{i+1}. {t}" for i,t in enumerate(c['taller'])) + f"""
Circular por estudiantes (o salas). Empujar evidencia, no perfectionismo.
Entregable: {c['entregable']}
📸 Evidencia de avance de un estudiante (para su registro del corte) [[captura: cap02_taller.png | receta: 1) Con permiso del estudiante, capture SU pantalla con el artefacto de hoy a medio construir.  2) Recorte datos personales (nombre, correo) antes de guardar.  3) Guardela como Kit docente/Clase {c['n']}/Capturas/cap02_taller.png.  4) Sirve de referencia del nivel esperado en el proximo semestre; no se proyecta.]]

### 105-115 · Criterios de exito + quiz corto · {sl_crit.strip()}
Repasar checklist del dia con {sl_crit}«Criterios de exito / entregable».
""" + (
        f"Pasar quiz 8–10 min **en ExamLab** (preguntas de esta clase; ver Guia Docente - Parte Practica). "
        f"Version impresa/proyectable de respaldo: `Quiz Clase {c['n']} - VetCare.docx`. "
        f"Clave para usted: `Quiz Clase {c['n']} - CLAVE DOCENTE.docx` (**no proyectar**)."
        if c['quiz'] else
        "Sin quiz formal: 2 preguntas orales de cierre."
    ) + f"""

### 115-120 · Cierre · {sl_cierre.strip()}
**Decir:** «Queda avanzado: {c['hito_pi']}. Suban el taller a ExamLab hoy domingo 23:59 si aplica. Enunciado PI en Clases/Proyecto Integrador.»
Proyectar {sl_cierre}slide de cierre. Dudas finales.
"""

    md = f"""# Guion docente · Clase {c['n']} · {c['titulo']}

- **Curso:** Bases de Datos II (FI303215) · 120 min
- **Tipo:** {tipo}
- **Hilo:** Proyecto Integrador **VetCare DB**
- **{"Hoy cerramos el PI" if c['tipo'] == 'sustentacion' else "Hoy avanzamos el PI en"}:** {c['hito_pi']}
- **Entregable de hoy:** {c['entregable']}
- **Herramienta:** {c['herramienta']}
- **Slides:** Clases/Clase {c['n']} - {c['slug']}/Presentacion.pptx
- **Caso de estudio (anexo del estudiante):** `Clases/Proyecto Integrador/Anexo - Caso de estudio Clinica Huellitas - Bases de Datos II.docx`
  — perfil de la clinica, las 8 entidades, las 3 reglas, el elenco de nombres y la escala por clase.
  Remita a este anexo cada vez que alguien pregunte «que datos guarda» o «de que tamano es esto».

> Sin mapa completo del curso, sin bio del docente, sin fechas de periodo.
> Presentacion del Curso / Acuerdo cubren logistica global.

## Fundamento teorico para el docente (al servicio del PI)

El objetivo de la clase no es «cubrir un capitulo» aislado, sino producir evidencia
del PI VetCare. La teoria se limita a desbloquear el taller.

""" + "\n".join(f"- {t}" for t in c['teoria']) + _fundamento_md(c) + f"""

**Demo que usted debe poder repetir:** {c['demo']}

## Referencias a diapositivas
Numeracion real del deck `Clases/Clase {c['n']} - {c['slug']}/Presentacion.pptx`.
Las etiquetas [Slide N] del plan y del fundamento apuntan aqui.

""" + "\n".join(f"{i+1}. {t}" for i, t in enumerate(mapa)) + f"""

> Privado, no se proyecta: `Kit docente/Clase {c['n']}/Solucion Taller Clase {c['n']} - VetCare.docx`
""" + "\n" + plan + f"""

## Codigo / scripts
Carpeta Codigo/ — archivo {c['sql'] or 'N/A'}.

## Capturas
Carpeta `Kit docente/Clase {c['n']}/Capturas/`. Cada linea de pantallazo de arriba trae
el nombre exacto del archivo y, si todavia no existe, el paso a paso para producirlo:
tomelo, guardelo con ese nombre y vuelva a generar el guion — la imagen se embebe sola.
Detalle por captura en `Capturas/README_capturas.txt`.

## Criterios de exito del dia
- Cada estudiante tiene el entregable o sus gaps escritos.
- Queda claro el vinculo con la rubrica del PI (modelo, seguridad, procs, opt, integracion).
"""
    path = kit / f"Guion Docente Clase {c['n']} - {c['slug']}.md"
    path.write_text(md, encoding='utf-8')
    # capturas placeholder readme
    (kit/"Capturas"/"README_capturas.txt").write_text(
        f"Capturas de la Clase {c['n']} — Bases de Datos II\n"
        f"{'=' * 46}\n\n"
        "El guion embebe automaticamente cualquier PNG que exista en esta carpeta con\n"
        "el nombre esperado. Mientras no exista, el .docx imprime la receta en su lugar.\n\n"
        "1) cap01_demo.png — salida de la demo del docente\n"
        f"   - Abrir {c['herramienta']} y repetir la demo: {c['demo']}\n"
        "   - Capturar solo la ventana con el resultado (no el escritorio completo).\n"
        "   - Recortar a ~1200 px de ancho y guardar aqui como cap01_demo.png.\n\n"
        "2) cap02_taller.png — evidencia de avance de un estudiante\n"
        "   - Con permiso del estudiante, capturar su artefacto a medio construir.\n"
        "   - Recortar nombre y correo antes de guardar. No se proyecta en clase.\n\n"
        "Despues de agregar una imagen, regenerar el guion:\n"
        f'   SOLO_CLASES={c["n"]} python config/slides/build_uniajc_bd2_all.py\n',
        encoding='utf-8')
    if c['sql'] and c['sql'] in SQL_BODIES:
        (kit/"Codigo"/c['sql']).write_text(SQL_BODIES[c['sql']], encoding='utf-8')
    return path

def build_quiz(c):
    """Quiz estudiante (sin claves) + CLAVE DOCENTE aparte. Nunca claves en Clases/."""
    if not c.get('quiz') or c['n'] not in QUIZ:
        return None
    kit = KIT_DIR / f"Clase {c['n']}"
    kit.mkdir(parents=True, exist_ok=True)
    quiz = QUIZ[c['n']]

    out = kit / f"Quiz Clase {c['n']} - VetCare.docx"
    doc = Document()
    banda(doc, f"Quiz · Clase {c['n']} · VetCare PI")
    para(doc, "Versión para aplicar en clase (sin claves). Corrección: CLAVE DOCENTE.", size=10, bold=True)
    para(doc, "8–10 min · individual · VetCare + tema del día. OM + V/F + abiertas cortas.", size=10)
    for i, item in enumerate(quiz, 1):
        for line in student_lines(item, i):
            text = line.replace("**", "")
            para(doc, text.lstrip(), bold=text.startswith(f"{i}."), size=11)
    doc.save(str(out))
    print("QUIZ", out)

    clave_out = kit / f"Quiz Clase {c['n']} - CLAVE DOCENTE.docx"
    doc_k = Document()
    banda(doc_k, f"CLAVE DOCENTE · Quiz Clase {c['n']} · VetCare")
    para(doc_k, "PRIVADO DOCENTE — no compartir ni proyectar", size=11, bold=True,
         color=RGBColor(0xA0, 0x20, 0x30), shade_fill="FBE4E4")
    para(doc_k, "Usar tras el quiz para retro. No proyectar en Presentacion.pptx.", size=10)
    for i, item in enumerate(quiz, 1):
        para(doc_k, clave_text(item, i), size=10, shade_fill="E8F4FA")
    doc_k.save(str(clave_out))
    print("CLAVE", clave_out)
    return out

def convert_guion(md_path: Path):
    conv = SLIDES / "guion_md_a_docx.py"
    if not conv.exists(): return None
    try:
        subprocess.run([sys.executable, str(conv), str(md_path)], check=False)
    except Exception as e:
        print("WARN convert", e)

def build_readme():
    text = """# Kit docente — Bases de Datos II (2026-2)

Material **privado** del docente. Los estudiantes solo ven Clases/.

## Enfoque
Todo el material de clase esta orientado a avanzar el **Proyecto Integrador VetCare DB**.
Teoria breve; talleres = entregables del PI.

## Por clase
- Guion Docente Clase N - ….md + .docx (via guion_md_a_docx.py)
- Quiz Clase N - VetCare.docx (sin claves) + Quiz Clase N - CLAVE DOCENTE.docx (privado; no proyectar)
- Codigo/*.sql demos VetCare
- Capturas/ con README de paso a paso; si falta el PNG, el guion imprime la receta
- Dias 5/9/14: Guia aplicacion Parcial N (solo evaluacion)

## Builds
`ash
python .config/slides/build_uniajc_bd2_all.py
python .config/slides/build_uniajc_bd2_curso.py
`

## PI
- Estudiante: Clases/Proyecto Integrador/
- Docente: Kit docente/Proyecto Integrador/
"""
    (KIT_DIR / "README.md").write_text(text, encoding='utf-8')

EXAMLAB_URL = "https://uniaj.examlab.workers.dev/"

# Boceto sugerido para la Pizarra (whiteboard) de cada clase — agnostico de
# herramienta: sirve igual en el whiteboard de ExamLab, en una pizarra fisica
# o en draw.io/Excalidraw si el docente prefiere otra.
PIZARRA = {
    1: "ER minimo: Dueño —1:N→ Mascota —1:N→ Cita. Marcar PK subrayada y FK con flecha.",
    2: "Tabla simple 3 columnas: Rol | Objeto | Privilegio (llenar en vivo con los 4 roles del taller).",
    3: "Flujo: App → llama sp_agendar_cita → valida mascota.activa → INSERT o mensaje de error.",
    4: "Mismo ER de Clase 1 + una nota junto a Cita: 'AQUI dispara el trigger de auditoria' y junto a Mascota: 'AQUI vive la fn_precio_base'.",
    6: "Dos columnas: 'Antes' (consulta con SELECT * y JOIN sin filtro) vs 'Despues' (columnas puntuales + filtro temprano) sobre el mismo dibujo de tablas.",
    7: "Tabla caliente (Cita) con una flecha grande hacia un rectangulo 'INDICE idx_cita_fecha_hora' y la palabra 'acelera lectura / cuesta escritura'.",
    8: "Linea de tiempo horizontal: BEGIN → INSERT factura → INSERT detalle → UPDATE stock → COMMIT/ROLLBACK con una bifurcacion visual en el ROLLBACK.",
    10: "La MISMA linea de tiempo T1/T2 de la diapositiva de Clase 10, pero redibujada en vivo con los IDs reales que use el script de demo.",
    11: "Checklist en 2 columnas: Evidencia (ER, DDL, roles, procs, fn, triggers, opt) | Si/No/Parcial — llenar en vivo con el curso.",
    12: "Caja 'App' — flecha rotulada con el nombre del proc (ej. sp_agendar_cita) — caja 'Base de datos'. Sin flecha directa App→tablas.",
    13: "Tabla 4 columnas: Contexto | Fallo | Leccion | Cambio en VetCare — una fila por caso discutido.",
    15: "Checklist final de empaquetado: ER, DDL, roles, procs, triggers, optimizacion, transacciones, concurrencia, contrato, informe — marcar completos.",
}

# Prompt generico reutilizable: se llena con el tema puntual de la clase.
_PROMPT_TEMPLATE = (
    "Actua como docente de Bases de Datos II. Usando el dominio VetCare (Dueño, "
    "Mascota, Cita, Veterinario, Insumo, Factura), dame un ejemplo minimo en SQL "
    "(Oracle/PostgreSQL) sobre «{tema}»: (1) el DDL de las tablas que necesito, "
    "(2) datos de ejemplo realistas de una clinica veterinaria (INSERT), (3) el "
    "codigo que ilustra «{tema}» paso a paso, (4) en 3 lineas, que debe notar el "
    "estudiante cuando lo vea ejecutar."
)


def build_guia_practica():
    """Guia docente de la PARTE PRACTICA de cada clase: que actividad hacer,
    objetivo, boceto de pizarra + prompt de apoyo + script SQL completo para la
    demo, pasos guiados, entregable y criterios de exito. Agnostica de
    plataforma — ExamLab se menciona solo como el lugar de entrega/quiz, sin
    detalle de configuracion. Reorganiza el mismo contenido de CLASES/
    TALLER_BLOQUE/SOLUCION/SQL_BODIES que ya usan el Guion Docente y el Taller."""
    lines = [
        "# Guia Docente — Parte Practica por Clase (Bases de Datos II, 2026-2)",
        "",
        "> Cada clase = una practica con objetivo propio. La demo se apoya en un",
        "> boceto de pizarra + un script SQL completo (con datos) para que usted lo",
        "> ejecute en vivo **en la consola de ExamLab**, que corre PostgreSQL (PGlite)",
        "> en el navegador: es el mismo motor donde se califica el taller, y varios de",
        "> estos scripts son PL/pgSQL que Oracle Live SQL no compila. El taller y el",
        f"> quiz se entregan/presentan en ExamLab (`{EXAMLAB_URL}`) — no es la",
        "> plataforma oficial de la UNIAJC, pero es la que usamos para eso en este curso.",
        "",
    ]
    for c in CLASES:
        if c["tipo"] == "parcial":
            continue
        tb = TALLER_BLOQUE.get(c["n"], {})
        sol = SOLUCION.get(c["n"], {})
        lines += [
            f"## Clase {c['n']} — {c['titulo']}",
            "",
            f"**Objetivo practico:** {c['hito_pi']}",
        ]
        contexto = (tb.get("contexto") or [None])[0]
        if contexto:
            contexto = re.sub(r"@@Por qu[eé] importa[^@]*@@\s*", "", contexto).replace("@@", "")
            lines.append(f"**Por que importa:** {contexto}")
        lines += [
            "",
            "**Demo en vivo:**",
            f"- Pizarra: {PIZARRA.get(c['n'], 'Boceto libre del ejemplo de hoy sobre el ER de VetCare.')}",
            f"- Prompt de apoyo (IA, opcional si le falta tiempo de preparar): \"{_PROMPT_TEMPLATE.format(tema=c['titulo'].split('·')[0].strip())}\"",
        ]
        sql_key = c.get("sql")
        if sql_key and sql_key in SQL_BODIES:
            lines.append(f"- Script SQL completo para correr en vivo (con datos de ejemplo):")
            lines.append("```sql")
            lines.append(SQL_BODIES[sql_key].strip())
            lines.append("```")
        else:
            lines.append("- Esta clase no requiere script SQL nuevo: es analisis/discusion sobre lo ya construido.")
        lines += [
            "",
            "**Pasos guiados del taller:**",
        ]
        for i, t in enumerate(c["taller"], 1):
            lines.append(f"{i}. {t}")
        lines += [
            "",
            f"**Entregable:** {c['entregable']}",
            "**Criterios de exito:**",
        ]
        for cr in (tb.get("criterios") or sol.get("rubrica") or ["Avance verificable del PI VetCare."]):
            lines.append(f"- {cr}")
        n_preg = len(QUIZ.get(c["n"], []))
        lines += [
            "",
            f"**Quiz de cierre:** {n_preg} preguntas (banco completo en `Kit docente/Clase {c['n']}/Quiz Clase {c['n']} - VetCare.docx`).",
            f"**Entrega:** taller y quiz en ExamLab · domingo 23:59 cuando aplique el taller.",
            "",
            "---",
            "",
        ]
    lines += [
        "## Parciales (Clase 5 / 9 / 14)",
        "",
        "Solo evaluacion — enunciado + solucion en `Parciales/` (dominio VetCare).",
        "Se presentan en ExamLab con proctoring activado (es evaluacion formal",
        "virtual sincrona por Meet); duracion 90-100 min dentro del bloque de 120.",
        "",
        "## Proyecto Integrador VetCare DB",
        "",
        "Hilo conductor de todas las clases regulares/autonomas. Avance formal en",
        "Clase 11 (checkpoint) y entrega/sustentacion en Clase 15. Se sube a ExamLab",
        "como Proyecto (individual por defecto; equipo de 2-3 solo si el docente lo autoriza); pesa 20% del Corte 3.",
        "",
    ]
    KIT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = KIT_DIR / "Guia Docente - Parte Practica por Clase.md"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    print("OK guia practica (md) ->", md_path)
    convert_guion(md_path)
    return md_path


def build_examlab_guia(c):
    """Guia para armar el taller de esta clase dentro de ExamLab.

    Va en el Kit docente porque la plataforma no importa preguntas desde archivo:
    el docente las crea en la UI y necesita el texto exacto de cada campo (incluido
    el SQL de partida, que aqui es PostgreSQL y no Oracle).
    """
    taller = TALLERES_EXAMLAB.get(c["n"])
    if not taller:
        return None
    d = KIT_DIR / f"Clase {c['n']}"
    d.mkdir(parents=True, exist_ok=True)
    md = examlab_talleres.guia_docente_md(
        c["n"], taller, "Bases de Datos II (FI303215)",
        hito=c.get("hito_pi"), entregable=c.get("entregable"),
    )
    out = d / f"Taller en ExamLab - Clase {c['n']} (configuracion).md"
    out.write_text(md, encoding="utf-8")
    print("EXAMLAB", out)
    return out


def main(solo_clases=None):
    """Regenera todo el curso, o solo un subconjunto de clases.

    ``solo_clases`` (iterable de numeros) o la variable de entorno
    SOLO_CLASES="2,15" limitan la regeneracion: las clases no incluidas no se
    tocan, para poder aislar una correccion sin reescribir el curso completo.
    Los archivos globales (README, Guia Docente - Parte Practica) se regeneran
    solo en la corrida completa, porque agregan las 15 clases.
    """
    if solo_clases is None:
        env_val = os.environ.get("SOLO_CLASES")
        if env_val:
            solo_clases = {int(x.strip()) for x in env_val.split(",") if x.strip()}
    else:
        solo_clases = set(solo_clases)
    KIT_DIR.mkdir(parents=True, exist_ok=True)
    CLASES_DIR.mkdir(parents=True, exist_ok=True)
    if solo_clases is None:
        build_readme()
    for c in CLASES:
        if solo_clases is not None and c['n'] not in solo_clases:
            continue
        print("=== Clase", c['n'], c['tipo'], "===")
        build_pptx(c)
        build_taller_docx(c)
        build_solucion_docx(c)
        md = build_guion_md(c)
        if md: convert_guion(md)
        build_quiz(c)
        build_examlab_guia(c)
    build_guia_practica()
    print("DONE")

if __name__ == '__main__':
    main()