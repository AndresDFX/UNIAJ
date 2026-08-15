# -*- coding: utf-8 -*-
"""Part 3: builders PPTX/DOCX/guion for Arquitectura PI-first."""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pptx.util import Pt as EPt

from part1 import CURSO, PARCIAL_FILE, RAW, SLIDES_DIR
from part2 import EXTRA, FUND, PASOS, QUIZ

sys.path.insert(0, str(SLIDES_DIR))
from uniajc_slides_engine import (  # noqa: E402
    CONTENT_W, CIAN, GRAY, MARGIN, MSO_ANCHOR, NAVY, PP_ALIGN, SW, WHITE,
    _rich, _run, add_logo, bg_white, blank, box_note_slide, closing_slide,
    content_slide, footer_num, new_prs, rect, textbox,
)

AZUL = RGBColor(0x09, 0x52, 0x92)
CIAN_D = RGBColor(0x26, 0x9C, 0xCB)
GRIS = RGBColor(0x2B, 0x2B, 0x2B)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xA0, 0x20, 0x30)
FONT = "Calibri"


def classes():
    out = []
    for n, tipo, slug, tema, sub, pi, ent, tool in RAW:
        out.append(dict(
            n=n, tipo=tipo, slug=slug, tema=tema, sub=sub, pi_hoy=pi, entregable=ent,
            herramienta=tool,
            objetivos=[
                f"Avanzar el PI CloudLite: {pi}.",
                "Aplicar el concepto del dia solo en funcion del entregable.",
                "Dejar evidencia integrable a tu informe/repo del PI.",
            ],
            slides_extra=EXTRA.get(n, []),
            taller_pasos=PASOS.get(n, []),
            quiz=QUIZ.get(n, []),
            taller_titulo=(
                f"{'Actividad autonoma' if tipo == 'autonoma' else 'Taller'} Clase {n} - CloudLite"
            ),
        ))
    return out


def _shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _rd(run, size=11, bold=False, color=GRIS):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def para(doc, text, size=11, bold=False, color=GRIS, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=6, shade=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if shade:
        _shade(p, shade)
    r = p.add_run(text)
    _rd(r, size=size, bold=bold, color=color)
    return p


def banda(doc, text):
    return para(doc, f"  {text}", size=13, bold=True, color=BLANCO, shade="095292")


def h2(doc, text):
    return para(doc, text, size=12, bold=True, color=AZUL, space_after=6)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        _rd(r)


def margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.85)


def cover(prs, n, tema, sub, pi):
    s = blank(prs)
    bg_white(s)
    rect(s, 0, 0, SW, 3.0, NAVY)
    rect(s, 0, 3.0, SW, 0.08, CIAN)
    add_logo(s, width=2.0, corner="left-top", mt=0.3, mr=0.5, variant="white")
    tn = textbox(s, SW - 2.2, 0.35, 1.8, 0.4)
    pn = tn.paragraphs[0]
    pn.alignment = PP_ALIGN.RIGHT
    _run(pn.add_run(), f"Clase {n}", 12, CIAN, bold=True)
    tf = textbox(s, MARGIN, 1.0, CONTENT_W, 1.5, anchor=MSO_ANCHOR.MIDDLE)
    p = tf.paragraphs[0]
    p.alignment = PP_ALIGN.CENTER
    _run(p.add_run(), tema, 28, WHITE, bold=True)
    ps = tf.add_paragraph()
    ps.alignment = PP_ALIGN.CENTER
    ps.space_before = EPt(8)
    _run(ps.add_run(), sub, 15, CIAN)
    tm = textbox(s, MARGIN, 3.5, CONTENT_W, 2.8)
    for i, ln in enumerate([
        f"**Hoy avanzamos el PI en:** {pi}",
        "Bloque **120 min** · Teoria breve · Taller PI · Quiz/cierre.",
        "Gratis + navegador · sin AWS/GCP/Oracle Cloud.",
    ]):
        pp = tm.paragraphs[0] if i == 0 else tm.add_paragraph()
        pp.space_after = EPt(8)
        _rich(pp, ln, 15, GRAY)
    footer_num(s, 1)


def build_pptx(c):
    n = c["n"]
    folder = CURSO / "Clases" / f"Clase {n} - {c['slug']}"
    folder.mkdir(parents=True, exist_ok=True)
    out = folder / "Presentacion.pptx"
    prs = new_prs()
    if c["tipo"] == "parcial":
        cover(prs, n, c["tema"], "Solo evaluacion · sin trabajo dirigido", c["pi_hoy"])
        content_slide(prs, "Indicaciones", [
            "Hoy es **solo Parcial** (presencial sincrono).",
            "No hay taller ni avance dirigido del PI.",
            "Prep del PI/pitch fue en la clase regular anterior.",
        ], idx=2)
        closing_slide(prs, f"Parcial · Clase {n}", [
            "Evaluacion del corte", "PI CloudLite continua despues",
        ], accent="Solo evaluacion")
    else:
        cover(prs, n, c["tema"], c["sub"], c["pi_hoy"])
        idx = 2
        content_slide(prs, "Agenda (120 min)", [
            "**0-10** Encuadre PI + entregable",
            "**10-40** Teoria breve al servicio del taller",
            "**40-100** Taller guiado PI",
            "**100-115** Quiz / evidencias",
            "**115-120** Cierre domingo 23:59",
        ], idx=idx)
        idx += 1
        content_slide(prs, "Objetivos", c["objetivos"], idx=idx)
        idx += 1
        content_slide(prs, "PI CloudLite — entregable de hoy", [
            f"@@Entregable:@@ {c['entregable']}",
            f"Herramienta: **{c['herramienta']}**",
            "Entra al informe/repo del PI (no lab suelto).",
        ], idx=idx)
        idx += 1
        for t, b in c["slides_extra"]:
            content_slide(prs, t, b, idx=idx)
            idx += 1
        content_slide(
            prs, "Taller PI",
            [f"**{i+1}.** {p}" for i, p in enumerate(c["taller_pasos"])],
            idx=idx,
        )
        idx += 1
        box_note_slide(prs, "Para continuar (PI)", [
            ("info", f"Entregable: {c['entregable']}"),
            ("aclaracion", "ExamLab domingo 23:59."),
            ("advertencia", "Sin cloud de pago ni Docker Desktop obligatorio."),
        ], idx=idx)
        closing_slide(prs, f"Clase {n} · PI en movimiento", [
            c["pi_hoy"], c["entregable"], "Siguiente hito CloudLite",
        ], accent="Teoria al servicio del proyecto")
    prs.save(str(out))
    print("OK pptx", out)


def build_taller(c):
    if c["tipo"] == "parcial":
        return
    n = c["n"]
    folder = CURSO / "Clases" / f"Clase {n} - {c['slug']}"
    path = folder / f"{c['taller_titulo']}.docx"
    doc = Document()
    margins(doc)
    banda(doc, c["taller_titulo"])
    para(doc, "Arquitectura · CloudLite App (PI)", bold=True, color=AZUL,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Documento estudiante — avance del Proyecto Integrador", color=CIAN_D,
         align=WD_ALIGN_PARAGRAPH.CENTER, shade="E8F4FA")
    h2(doc, "1. Hoy avanzamos el PI en…")
    para(doc, c["pi_hoy"])
    h2(doc, "2. Entregable")
    para(doc, c["entregable"], shade="E8F4FA")
    h2(doc, "3. Herramientas")
    para(doc, c["herramienta"])
    para(doc, "Prohibido: AWS/GCP/Oracle/Azure con tarjeta; hipervisores/Docker Desktop obligatorio.",
         shade="FBE4E4")
    h2(doc, "4. Pasos")
    bullets(doc, c["taller_pasos"])
    h2(doc, "5. Criterio de exito")
    bullets(doc, [
        "Artefacto en paquete PI.",
        "Explicacion oral de 60 s (si hay equipo, cualquier integrante).",
        "Evidencia adjunta.",
    ])
    h2(doc, "6. Entrega")
    para(doc, "Entrega en ExamLab · domingo 23:59 · envio individual.")
    doc.save(str(path))
    print("OK taller", path)


def build_quiz(c):
    if c["tipo"] == "parcial" or not c["quiz"]:
        return
    n = c["n"]
    kit = CURSO / "Kit docente" / f"Clase {n}"
    kit.mkdir(parents=True, exist_ok=True)
    path = kit / f"Quiz Clase {n} - {c['slug']}.docx"
    doc = Document()
    margins(doc)
    banda(doc, f"Quiz Clase {n} — {c['tema']}")
    para(doc, "DOCUMENTO DOCENTE — con respuestas", bold=True, color=ROJO,
         align=WD_ALIGN_PARAGRAPH.CENTER, shade="FBE4E4")
    for i, (q, a) in enumerate(c["quiz"], 1):
        h2(doc, f"Pregunta {i}")
        para(doc, q)
        para(doc, f"Respuesta: {a}", shade="E8F4FA")
    doc.save(str(path))
    print("OK quiz", path)


def guion_md(c):
    n = c["n"]
    if c["tipo"] == "parcial":
        return (
            f"# Guion docente — Clase {n}: {c['tema']}\n\n"
            f"## Info\n- FI303380 · **120 min** · Solo evaluacion (Parcial)\n"
            f"- PI CloudLite: sin taller dirigido hoy\n\n"
            f"## Plan\n### 0-10 Organizacion\n"
            f"Di: «Hoy es solo Parcial. No hay taller PI.»\n\n"
            f"### 10-100 Aplicacion\n"
            f"Enunciado en `Parciales/{PARCIAL_FILE.get(n, '')}`. Dudas solo de enunciado.\n\n"
            f"### 100-120 Recoleccion y cierre\n"
            f"«PI CloudLite continua en la siguiente clase regular/autonoma.»\n"
        )
    pasos = "\n".join(f"{i+1}. {p}" for i, p in enumerate(c["taller_pasos"]))
    if c["tipo"] == "autonoma":
        plan = (
            f"### Autonoma (festivo)\n"
            f"Publica Presentacion.pptx + Taller.docx en ExamLab.\n"
            f"Mensaje: «Clase {n} autonoma. Hoy avanzamos el PI en: {c['pi_hoy']}. "
            f"Entregable: {c['entregable']}.»\n"
            f"Revisa domingo 23:59 con feedback a rubrica PI.\n"
        )
    else:
        plan = (
            f"### 0-10 Encuadre PI\n"
            f"Di: «Hoy avanzamos el PI CloudLite en: **{c['pi_hoy']}**. "
            f"Entregable: {c['entregable']}.»\n\n"
            f"### 10-40 Teoria Core\n"
            f"Slides al servicio del taller. Cada 7-8 min amarra al artefacto PI.\n\n"
            f"### 40-55 Demo\n"
            f"Herramienta: **{c['herramienta']}**. "
            f"Capturas en `Kit docente/Clase {n}/Capturas/` si falla la red.\n\n"
            f"### 55-100 Taller PI\n"
            f"Pasos del taller estudiante. Bloquea dominios vagos. A los 80 min pedir evidencia.\n\n"
            f"### 100-115 Quiz\nQuiz del Kit + retro a 2-3 equipos.\n\n"
            f"### 115-120 Cierre\n"
            f"«Explicar artefacto en 60 s. Entrega domingo 23:59.»\n"
        )
    objs = "\n".join("- " + o for o in c["objetivos"])
    return (
        f"# Guion docente — Clase {n}: {c['tema']}\n\n"
        f"## Info\n- Arquitectura de Sistemas Computacionales (FI303380) · **120 min**\n"
        f"- Enfoque: **PI CloudLite App**\n"
        f"- Sin fechas de periodo · sin bio · sin mapa del curso\n\n"
        f"## Objetivos\n{objs}\n\n"
        f"## Hoy avanzamos el PI en…\n**{c['pi_hoy']}**\n\n"
        f"**Entregable:** {c['entregable']}\n\n"
        f"**Herramienta:** {c['herramienta']}\n\n"
        f"## Fundamento para el docente\n{FUND.get(n, 'Teoria al servicio del entregable PI.')}\n\n"
        f"Slides: `Clases/Clase {n} - {c['slug']}/Presentacion.pptx`\n\n"
        f"## Plan (120 min)\n\n{plan}\n"
        f"## Taller\n{pasos}\n\n"
        f"## Criterio de exito\nArtefacto en paquete PI · evidencia · explicacion 60 s.\n\n"
        f"## Quiz\n`Kit docente/Clase {n}/Quiz Clase {n} - {c['slug']}.docx`\n\n"
        f"## Capturas\n- Pantallazo herramienta del dia con artefacto CloudLite\n"
        f"- Pantallazo evidencia del entregable\n\n"
        f"Entrega en ExamLab. Sin cloud con tarjeta.\n"
    )


def build_guion(c):
    n = c["n"]
    kit = CURSO / "Kit docente" / f"Clase {n}"
    cap = kit / "Capturas"
    kit.mkdir(parents=True, exist_ok=True)
    cap.mkdir(exist_ok=True)
    (cap / "README.txt").write_text(f"Capturas Clase {n}\n", encoding="utf-8")
    if c["tipo"] == "parcial":
        (kit / f"NOTA Docente - Clase {n} Parcial.md").write_text(
            f"# Clase {n} Solo Parcial\n\n- Sin taller PI.\n"
            f"- Enunciado: `Parciales/{PARCIAL_FILE.get(n, '')}` (+ SOLUCION).\n"
            f"- Prep PI en clase regular anterior.\n",
            encoding="utf-8",
        )
    path = kit / f"Guion Docente Clase {n} - {c['slug']}.md"
    path.write_text(guion_md(c), encoding="utf-8")
    print("OK guion", path)
    return path


def build_all():
    guiones = []
    for c in classes():
        print(f"\n=== Clase {c['n']} ({c['tipo']}) ===")
        build_pptx(c)
        if c["tipo"] != "parcial":
            build_taller(c)
            build_quiz(c)
        guiones.append(build_guion(c))
    conv = SLIDES_DIR / "guion_md_a_docx.py"
    for md in guiones:
        subprocess.run([sys.executable, str(conv), str(md)], check=False)
    print("\nDONE")


if __name__ == "__main__":
    build_all()
