# -*- coding: utf-8 -*-
"""Genera calendario 2026-2 + Acuerdos pedagógicos prellenados (sin listado de estudiantes).

FUENTE DE VERDAD: `config/calendario/semestre_2026_2.json` (este script lo LEE, no lo
recalcula ni lo reescribe). El JSON trae, por curso, el array `clases` con las 13
sesiones del semestre acortado (inicio 2026-08-24 · fin fijo 2026-11-22) y, en cada
sesión, `clases_material` = los "Clase N" del material ya construido que se dictan ahí
(las carpetas `Clases/Clase N - …` y `Kit docente/Clase N/` NO se renumeran).

Salidas por curso:
  <Curso>/Plan curso/2026-2/CALENDARIO_2026-2.md
  <Curso>/Plan curso/2026-2/calendario_eventos_2026-2.csv   (14 columnas · UTF-8 BOM)
  <Curso>/Plan curso/2026-2/Cronograma 2026-2.md            (documento del estudiante)
  <Curso>/Plan curso/2026-2/PLAN_DE_CURSO_2026-2.md         (cabecera regenerada · secciones propias preservadas)
  <Curso>/Entregas docente/2026-2/ACUERDO PEDAGOGICO - <Curso> - 2026-2.docx
  <Curso>/Plan curso/2026-2/CORREO_BIENVENIDA - <Curso> - 2026-2.md (fechas y bloques gestionados)
  config/calendario/eventos_<curso>_*.csv + eventos_todos_cursos_2026-2.csv
"""
from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
JSON_PATH = Path(__file__).with_name("semestre_2026_2.json")
TEMPLATE = (
    ROOT
    / "Programacion II"
    / "Entregas docente"
    / "ACUERDO PEDAGOGICO ACTUALIZADO.docx"
)

DATA = json.loads(JSON_PATH.read_text(encoding="utf-8"))

START = DATA["inicio"]          # 2026-08-24 (semestre acortado)
END = DATA["fin"]               # 2026-11-22 (NO se mueve)
FESTIVOS = DATA["festivos_en_rango"]
CORTES = DATA["cortes_teoricos"]
DOCENTE = DATA["docente"]["nombre_completo"]
CORREO = DATA["docente"]["correo"]
N_TEMAS = 15

LOGICA_EVALUACION = DATA["logica_evaluacion"]

# ---------------------------------------------------------------- utilidades


def dmy(iso: str) -> str:
    y, m, d = iso.split("-")
    return f"{d}/{m}/{y}"


def dm(iso: str) -> str:
    y, m, d = iso.split("-")
    return f"{d}/{m}"


def festivos_en_rango() -> dict[str, str]:
    return {k: v for k, v in FESTIVOS.items() if START <= k <= END}


# Los `tema` del JSON vienen sin tildes (ASCII). Estos documentos se comparten con
# estudiantes, así que se restituye la ortografía palabra por palabra (sin tocar el JSON).
ACENTOS = {
    "Presentacion": "Presentación",
    "presentacion": "presentación",
    "Introduccion": "Introducción",
    "introduccion": "introducción",
    "Diagnostico": "Diagnóstico",
    "diagnostico": "diagnóstico",
    "dinamicas": "dinámicas",
    "graficas": "gráficas",
    "Evaluacion": "Evaluación",
    "evaluacion": "evaluación",
    "Documentacion": "Documentación",
    "documentacion": "documentación",
    "Refactorizacion": "Refactorización",
    "Revision": "Revisión",
    "revision": "revisión",
    "Integracion": "Integración",
    "integracion": "integración",
    "codigo": "código",
    "modulos": "módulos",
    "Diseno": "Diseño",
    "diseno": "diseño",
    "Sustentacion": "Sustentación",
    "sustentacion": "sustentación",
    "sustentaciones": "sustentaciones",
    "Metodologias": "Metodologías",
    "agiles": "ágiles",
    "Analisis": "Análisis",
    "Optimizacion": "Optimización",
    "optimizacion": "optimización",
    "Administracion": "Administración",
    "Indices": "Índices",
    "gestion": "gestión",
    "automatica": "automática",
    "Virtualizacion": "Virtualización",
    "Preparacion": "Preparación",
    "preparacion": "preparación",
    "Escalabilidad": "Escalabilidad",
}
_ACENTOS_RE = re.compile(r"\b(" + "|".join(sorted(ACENTOS, key=len, reverse=True)) + r")\b")


def tema_txt(cl: dict) -> str:
    # El prefijo "DOBLE: " del JSON es redundante: las sesiones dobles ya se marcan
    # aparte (columna de material y etiqueta «Sesión doble»).
    tema = re.sub(r"^DOBLE:\s*", "", cl["tema"])
    return _ACENTOS_RE.sub(lambda m: ACENTOS[m.group(1)], tema)


TIPO_LABEL = {
    "presencial": "Presencial (síncrona)",
    "virtual": "Virtual (síncrona)",
    "autonoma": "Autónoma (festivo)",
    "sustentacion": "Sustentación PI (festivo)",
}


def corte_rangos() -> list[tuple[int, int, int]]:
    """[(sesion_ini, sesion_fin, n_corte)] leído del JSON (no hardcodeado a 15)."""
    out: list[tuple[int, int, int]] = []
    for i, key in enumerate(sorted(CORTES), start=1):
        a, b = CORTES[key]["clases"].split("-")
        out.append((int(a), int(b), i))
    return out


def corte_de(n: int) -> int:
    for a, b, pn in corte_rangos():
        if a <= n <= b:
            return pn
    return corte_rangos()[-1][2]


def parcial_de_corte(meta: dict, pn: int) -> dict | None:
    return meta.get("parciales", {}).get(f"parcial_{pn}")


def material(cl: dict) -> str:
    ms = cl.get("clases_material") or []
    if not ms:
        return "—"
    if len(ms) == 1:
        return f"Clase {ms[0]}"
    return " + ".join(f"Clase {m}" for m in ms)


def etiqueta_sesion(cl: dict) -> str:
    base = f"Sesión {cl['n']}"
    if cl.get("parcial"):
        base += f" · Parcial {cl['parcial_n']}"
    elif cl.get("tipo") == "sustentacion":
        base += " · Sustentación PI"
    base += f" · {material(cl)}"
    if cl.get("sesion_doble"):
        base += " (doble)"
    return base


def nota_sesion(cl: dict) -> str:
    notas: list[str] = []
    if cl.get("festivo"):
        notas.append(f"festivo: {cl['festivo']}")
    if cl.get("tipo") == "autonoma":
        notas.append("clase autónoma (trabajo independiente guiado)")
    if cl.get("tipo") == "sustentacion":
        notas.append("sesión de sustentaciones del Proyecto Integrador (no es parcial)")
    if cl.get("parcial"):
        notas.append("parcial presencial síncrono")
    if cl.get("sesion_doble"):
        notas.append(f"sesión doble: cubre {material(cl)} del material en un bloque de 120 min")
    return "; ".join(notas)


def resumen_compresion() -> str:
    return (
        f"Semestre 2026-2 acortado: inicio **{dmy(START)}** (fin fijo **{dmy(END)}**) = "
        f"**13 sesiones**. Se conservan los **{N_TEMAS} temas** del microcurrículo: "
        "**2 sesiones son dobles** (dos temas afines en el mismo bloque de 120 min). "
        "El material existente (`Clases/Clase N - …`, `Kit docente/Clase N/`) **no se renumera**: "
        "cambia solo el mapeo Sesión → Clase(s) de material."
    )


def logica_curso(meta: dict) -> str:
    """Lógica de evaluación de ESTE curso.

    Se prefiere sobre DATA["logica_evaluacion"], que describe los 4 cursos a la vez
    y no corresponde a un documento de un solo curso.
    """
    clases = meta["clases"]
    parciales = [cl for cl in clases if cl.get("parcial")]
    nums = "/".join(str(cl["n"]) for cl in parciales)
    partes = [
        f'**{meta["nombre"]}** ({meta["dia"]} {meta["horario"]}) · '
        f'Modalidad: {meta["modalidad"]}.',
        f"Sesión 1 presencial (encuadre) · parciales en las sesiones **{nums}**, "
        "presencial síncrono · resto de sesiones regulares virtual síncrona.",
    ]
    autonomas = [cl for cl in clases if cl["tipo"] == "autonoma"]
    if autonomas:
        det = " · ".join(f'Sesión {cl["n"]} ({dmy(cl["fecha"])}, {cl["festivo"]})' for cl in autonomas)
        partes.append(f"Festivos = clase autónoma, no se omiten: {det}.")
    else:
        partes.append(f'No hay festivos en {meta["dia"].lower()}: todas las sesiones son regulares.')
    sust = [cl for cl in clases if cl.get("sustentacion_pi")]
    if sust:
        cl = sust[0]
        partes.append(
            f'Sesión {cl["n"]} ({dmy(cl["fecha"])}) se dedica a las **sustentaciones del '
            "Proyecto Integrador** (no es parcial)."
        )
    partes.append(
        "Día de parcial = solo evaluación. Los parciales NUNCA se programan en festivo ni "
        "en clase autónoma. Sesión 0 = Presentación del Curso (no es sesión temática)."
    )
    return " ".join(partes)


# ---------------------------------------------------------------- metadatos de acuerdo
# Campos que NO viven en el JSON (objetivos/RAA del microcurrículo, semestre, etc.).
ACUERDO_EXTRA = {
    "programacion_ii": {
        "grupo_acuerdo": "341-C",
        "semestre": "4",
        "programa": "Ingeniería de Sistemas",
        "objetivos": (
            "Comprender y aplicar los pilares de la Programación Orientada a Objetos (POO) en Java.\n"
            "Implementar y manipular Estructuras de Datos dinámicas en memoria.\n"
            "Desarrollar Interfaces Gráficas de Usuario (GUI) interactivas.\n"
            "Aplicar patrones de diseño y refactorización con apoyo de IA.\n"
            "Construir persistencia básica integrando lectura y escritura de archivos."
        ),
        "metodologia_base": (
            "Estructura de clase: Teoría Core · Taller Guiado calificable "
            "(entrega máx. domingo 23:59) · Quiz corto.\n"
            "Enfoque: aprendizaje activo / ABPr con Proyecto Integrador."
        ),
    },
    "seminario": {
        "grupo_acuerdo": "341-C",
        "semestre": "4",
        "programa": "Ingeniería de Sistemas",
        "objetivos": (
            "Levantar, analizar y documentar requerimientos funcionales y no funcionales "
            "de sistemas de información.\n"
            "Diseñar la arquitectura de software aplicando Patrones de Diseño "
            "(Singleton, Factory) y Diagramas UML.\n"
            'Emplear el paradigma "Docs as Code" (Mermaid, Markdown) para la '
            "documentación técnica y manuales de usuario.\n"
            "Desarrollar habilidades de comunicación asertiva para el Storytelling "
            "y sustentación de proyectos tecnológicos."
        ),
        "metodologia_base": (
            "Metodología orientada a Role-Playing (Arquitectos / Analistas QA), "
            "talleres con Draw.io/Mermaid, peer review."
        ),
    },
    "bases_datos_ii": {
        "grupo_acuerdo": "641A-2",
        "semestre": "Sexto Semestre",
        "programa": "Ingeniería de Sistemas",
        "objetivos": (
            "Objetivo de aprendizaje: Diseñar, administrar y optimizar bases de datos "
            "relacionales avanzadas, garantizando seguridad, integridad y eficiencia en el "
            "manejo de grandes volúmenes de información.\n\n"
            "Resultados de Aprendizaje de la Asignatura (RAA):\n"
            "RAA1: Administra bases de datos aplicando estrategias de seguridad y respaldo.\n"
            "RAA2: Implementa procedimientos almacenados y disparadores para la automatización "
            "de procesos.\n"
            "RAA3: Optimiza consultas y estructuras de bases de datos para mejorar el "
            "rendimiento del sistema.\n\n"
            "Fuente: Microcurrículo FI303215 — Bases de Datos II (Plan curso/)."
        ),
        "metodologia_base": (
            "Estructura sugerida (ajustar al Acuerdo/Plan): "
            "Teoría Core · Taller Guiado · Quiz/comprobación."
        ),
    },
    "arquitectura": {
        "grupo_acuerdo": "6303C",
        "semestre": "Sexto Semestre",
        "programa": "Ingeniería de Sistemas",
        "objetivos": (
            "Objetivo de aprendizaje: Diseñar e implementar arquitecturas de sistemas "
            "computacionales aplicando principios de computación en la nube, virtualización y "
            "escalabilidad, asegurando eficiencia y sostenibilidad.\n\n"
            "Resultados de Aprendizaje de la Asignatura (RAA):\n"
            "RAA1: Comprende y aplica modelos de servicio cloud (IaaS, PaaS, SaaS).\n"
            "RAA2: Configura entornos virtualizados y despliega sistemas distribuidos.\n"
            "RAA3: Evalúa la seguridad, rendimiento y sostenibilidad de arquitecturas en la nube.\n\n"
            "Fuente: Microcurrículo FI303380 — Arquitectura de Sistemas Computacionales "
            "(Enfoque Cloud) (Plan curso/)."
        ),
        "metodologia_base": (
            "Estructura sugerida: Teoría Core · Taller Guiado sobre el Proyecto Integrador "
            "CloudLite · Quiz/comprobación."
        ),
    },
}

# Nombre del CSV consolidado por curso en config/calendario/. Todos llevan el periodo
# vigente: antes programacion_ii y seminario se llamaban *_2026-1.csv aunque su contenido
# ya era 2026-2, y eso choca con que 2026-1 SÍ es un periodo real de esos dos cursos
# (ver <Curso>/Plan curso/2026-1/).
CSV_CONFIG_NAME = {
    "programacion_ii": "eventos_programacion_ii_2026-2.csv",
    "seminario": "eventos_seminario_2026-2.csv",
    "bases_datos_ii": "eventos_bases_datos_ii_2026-2.csv",
    "arquitectura": "eventos_arquitectura_2026-2.csv",
}

CSV_HEADER = [
    "curso",
    "codigo_fi",
    "grupo",
    "clase_n",
    "fecha",
    "dia",
    "hora_inicio",
    "hora_fin",
    "tipo_clase",
    "es_parcial",
    "parcial_n",
    "sesion_etiqueta",
    "tema",
    "notas",
]


# ---------------------------------------------------------------- textos del acuerdo


def metodologia_text(key: str, meta: dict) -> str:
    extra = ACUERDO_EXTRA[key]
    clases = meta["clases"]
    fest = [cl for cl in clases if cl.get("festivo")]
    if fest:
        fest_line = "Festivos = clase autónoma:\n" + ", ".join(
            f"{dm(cl['fecha'])} ({cl['festivo']})" for cl in fest
        ) + "."
    else:
        fest_line = f"No hay festivos en {meta['dia'].lower()}; todas las sesiones son regulares."
    dobles = [cl["n"] for cl in clases if cl.get("sesion_doble")]
    sust = meta.get("sustentacion_pi")
    lines = [
        "Acuerdo sobre los aspectos metodológicos",
        f"Periodo 2026-2 · Grupo {meta['grupo']} · {meta['dia']} {meta['horario']} "
        f"({meta['duracion_min']} min).",
        "Modalidad: Presencialidad asistida (Sesión 1 y parciales presencial síncrono · "
        "resto de sesiones regulares virtual síncrona · festivos = clase autónoma).",
        extra["metodologia_base"],
        f"Calendario: 13 sesiones ({N_TEMAS} temas del microcurrículo) "
        f"({dmy(clases[0]['fecha'])}–{dmy(clases[-1]['fecha'])}). "
        f"Semestre acortado: inicio {dmy(START)}, fin {dmy(END)}; se conservan los "
        f"{N_TEMAS} temas porque las sesiones "
        + " y ".join(str(n) for n in dobles)
        + " son dobles (dos temas afines en un bloque de 120 min).",
        fest_line,
    ]
    if sust:
        lines.append(
            f"Sesión {sust['clase']} ({dmy(sust['fecha'])}): sustentaciones del proyecto final."
        )
    lines.append("[Detalle en Plan curso/2026-2/CALENDARIO_2026-2.md del curso.]")
    return "\n".join(lines)


def eval_text(meta: dict) -> str:
    clases = meta["clases"]
    by_n = {cl["n"]: cl for cl in clases}
    ordinal = {1: "Primer", 2: "Segundo", 3: "Tercer"}
    detalle = {
        1: "10% Talleres o Quiz | 10% Asistencia",
        2: "10% Talleres o Quiz | 10% Asistencia",
        3: "20% Proyecto Integrador | 5% Asistencia",
    }
    peso = {1: "10%", 2: "10%", 3: "15%"}
    out = [
        "Acuerdo sobre los aspectos de evaluación",
        "(Cálculo teórico 2026-2 · 30% / 30% / 40% — validar en socialización con el grupo)",
        "Parciales: presenciales y síncronos; NUNCA en festivo ni en clase autónoma. "
        "Criterio: última sesión regular del corte.",
        "",
    ]
    for a, b, pn in corte_rangos():
        c = CORTES[f"corte_{pn}"]
        p = parcial_de_corte(meta, pn)
        out.append(
            f"{ordinal[pn]} corte ({c['pct']}) — [{dmy(c['inicio'])} al {dmy(c['fin'])}] · "
            f"Sesiones {a}–{b}:"
        )
        if p:
            out.append(
                f"* {peso[pn]} Parcial {pn} (cierre de corte, Sesión {p['clase']} — "
                f"{dmy(p['fecha'])}, presencial) | {detalle[pn]}"
            )
        out.append("")
    sust = meta.get("sustentacion_pi")
    nota = (
        f"Nota: semestre 2026-2 acortado (inicio {dmy(START)} · fin {dmy(END)}) = 13 sesiones que "
        f"cubren los {N_TEMAS} temas del microcurrículo (2 sesiones dobles). Los porcentajes "
        "30/30/40 y su desglose NO cambian. Los parciales nunca caen en festivo ni en clase "
        "autónoma."
    )
    if sust:
        s = by_n.get(sust["clase"], {})
        nota += (
            f" La Sesión {sust['clase']} ({dmy(sust['fecha'])}"
            + (f", {s.get('festivo')}" if s.get("festivo") else "")
            + ") se dedica a las sustentaciones del Proyecto Integrador (no es parcial)."
        )
    out.append(nota)
    return "\n".join(out)


APROBACION = (
    "[PRELLENADO 2026-2 — pendiente socialización con el grupo]\n"
    f"Periodo académico: {dmy(START)} al {dmy(END)}.\n"
    "Pendiente: aprobación/ajustes con estudiantes, listado oficial, vocero y firmas.\n"
    "No se inventan nombres ni códigos de estudiantes."
)


# ---------------------------------------------------------------- docx helpers


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = True
    for p in paras:
        if first:
            if p.runs:
                p.runs[0].text = text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(text)
            first = False
        else:
            for r in p.runs:
                r.text = ""


def set_merged_row_value(row, start_col: int, text: str, end_col: int | None = None) -> None:
    end = end_col if end_col is not None else len(row.cells)
    for i in range(start_col, end):
        set_cell_text(row.cells[i], text)


def fill_acuerdo(key: str, meta: dict) -> Path:
    extra = ACUERDO_EXTRA[key]
    doc = Document(str(TEMPLATE))
    t0 = doc.tables[0]
    set_merged_row_value(t0.rows[0], 1, extra["programa"])
    set_merged_row_value(t0.rows[1], 1, meta["nombre"])
    set_cell_text(t0.rows[2].cells[1], extra["grupo_acuerdo"])
    set_cell_text(t0.rows[2].cells[3], extra["semestre"])
    set_cell_text(t0.rows[3].cells[1], "2026-2")
    set_cell_text(t0.rows[3].cells[3], "[PENDIENTE — fecha socialización]")
    set_merged_row_value(t0.rows[4], 1, DOCENTE)
    set_cell_text(doc.tables[1].rows[1].cells[0], extra["objetivos"])
    set_cell_text(doc.tables[2].rows[1].cells[0], APROBACION)
    set_cell_text(doc.tables[3].rows[0].cells[0], metodologia_text(key, meta))
    set_cell_text(doc.tables[3].rows[1].cells[0], eval_text(meta))
    t4 = doc.tables[4]
    for c in t4.rows[0].cells[3:]:
        set_cell_text(c, "[PENDIENTE — listado]")
    for c in t4.rows[1].cells[2:]:
        set_cell_text(c, "[PENDIENTE — vocero]")
    set_cell_text(t4.rows[2].cells[1], "[PENDIENTE]")
    set_cell_text(t4.rows[2].cells[2], "[PENDIENTE]")
    set_cell_text(t4.rows[2].cells[3], "[PENDIENTE]")
    set_cell_text(t4.rows[2].cells[5], "[PENDIENTE — email vocero]")
    set_cell_text(t4.rows[3].cells[1], "Cali")
    set_cell_text(t4.rows[3].cells[2], "Cali")
    set_cell_text(t4.rows[3].cells[3], "Cali")
    set_cell_text(t4.rows[3].cells[5], "[PENDIENTE — fecha]")
    t6 = doc.tables[6]
    for row in t6.rows[1:]:
        if len(row.cells) >= 3:
            set_cell_text(row.cells[1], "")
            set_cell_text(row.cells[2], "")
    note = doc.add_paragraph()
    note.add_run(
        f"Docente: {DOCENTE} · Correo: {CORREO} · "
        f'Horario: {meta["dia"]} {meta["horario"]} ({meta["duracion_min"]} min) · '
        f'Modalidad: {meta["modalidad"]} (Sesión 1 y parciales presencial síncrono · '
        "resto virtual síncrona · festivos = clase autónoma) · "
        f'Código: {meta["codigo"]} · '
        f"Periodo: {dmy(START)}–{dmy(END)} · 13 sesiones / {N_TEMAS} temas · "
        "PRELLENADO 2026-2 — campos de estudiantes pendientes."
    )
    out = (
        ROOT
        / meta["folder"]
        / "Entregas docente"
        / "2026-2"
        / f'ACUERDO PEDAGOGICO - {meta["nombre"]} - 2026-2.docx'
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------- markdown


def cortes_table(meta: dict) -> list[str]:
    lines = [
        "| Corte | % | Ventana | Sesiones | Parcial de cierre | Desglose teórico |",
        "|---|---|---|---|---|---|",
    ]
    for a, b, pn in corte_rangos():
        c = CORTES[f"corte_{pn}"]
        p = parcial_de_corte(meta, pn)
        parcial = (
            f"Parcial {pn} · Sesión {p['clase']} ({dmy(p['fecha'])}) · Presencial (síncrona)"
            if p
            else "—"
        )
        lines.append(
            f"| Corte {pn} | {c['pct']} | {dmy(c['inicio'])} → {dmy(c['fin'])} | {a}-{b} | "
            f"{parcial} | {c['desglose']} |"
        )
    return lines


def calendario_md(meta: dict) -> str:
    clases = meta["clases"]
    fest = festivos_en_rango()
    dobles = [cl for cl in clases if cl.get("sesion_doble")]
    lines = [
        f'# Calendario 2026-2 — {meta["nombre"]}',
        "",
        f'- **Código:** {meta["codigo"]}',
        f'- **Grupo:** {meta["grupo"]}',
        f"- **Periodo:** 2026-2 · **{dmy(START)} – {dmy(END)}**",
        f'- **Horario:** {meta["dia"]} **{meta["horario"]}** ({meta["duracion_min"]} min)',
        f'- **Modalidad:** **{meta["modalidad"]}** (Sesión 1 y parciales presencial síncrono · '
        "resto virtual síncrona · festivos = clase autónoma)",
        f"- **Docente:** {DOCENTE} · `{CORREO}`",
        f"- **Total sesiones:** {len(clases)} · **temas del microcurrículo:** {N_TEMAS} "
        f"({len(dobles)} sesiones dobles) — festivos = **clase autónoma**, no se omiten",
        "",
        f"> {resumen_compresion()}",
        "",
        "## Cortes teóricos (30% / 30% / 40%)",
        "",
        logica_curso(meta),
        "",
        *cortes_table(meta),
        "",
        "> **Día de parcial = solo evaluación** (sin tema de trabajo dirigido nuevo). "
        "Detalle temático en PLAN_DE_CURSO_2026-2.md.",
        "",
        "## Sesiones (mapeo Sesión → Clase de material)",
        "",
        "> La columna **Clase(s) de material** indica qué carpeta `Clases/Clase N - …` y "
        "`Kit docente/Clase N/` se usa en esa sesión. **No se renumeró nada**: las sesiones "
        "marcadas *(doble)* dictan dos clases de material en el mismo bloque de 120 min.",
        "",
        "| Sesión | Fecha | Tipo | Clase(s) de material | Tema | Nota |",
        "|---|---|---|---|---|---|",
    ]
    for cl in clases:
        mat = material(cl) + (" **(doble)**" if cl.get("sesion_doble") else "")
        nota = nota_sesion(cl) or "—"
        lines.append(
            f'| {cl["n"]} | {dmy(cl["fecha"])} | {TIPO_LABEL.get(cl["tipo"], cl["tipo"])} | '
            f'{mat} | {tema_txt(cl)} | {nota} |'
        )
    lines += ["", "## Sesiones dobles", ""]
    for cl in dobles:
        lines.append(f'- **Sesión {cl["n"]}** ({dmy(cl["fecha"])}) — {material(cl)}: {tema_txt(cl)}')
    lines += [
        "",
        "## Festivos Colombia 2026 (rango del periodo)",
        "",
    ]
    for iso, nombre in sorted(fest.items()):
        cl = next((c for c in clases if c["fecha"] == iso), None)
        extra = ""
        if cl:
            extra = (
                f" — Sesión {cl['n']}: "
                + ("sustentaciones del PI" if cl["tipo"] == "sustentacion" else "clase autónoma")
            )
        lines.append(f"- {dmy(iso)} — {nombre}{extra}")
    lines += [
        "",
        f"> Fuera de rango tras el nuevo inicio ({dmy(START)}): 17/08/2026 (Asunción de la Virgen) "
        "ya no cae en el periodo.",
        "",
        "Fuente: `config/calendario/semestre_2026_2.json` (generado por "
        "`config/calendario/generar_semestre_2026_2.py`).",
        "",
    ]
    return "\n".join(lines)


def cronograma_md(meta: dict) -> str:
    clases = meta["clases"]
    lines = [
        f'# Cronograma 2026-2 — {meta["nombre"]}',
        "",
        "Documento para estudiantes (carpeta compartida `Clases/`).",
        "",
        f'- **Código:** {meta["codigo"]} · **Grupo:** **{meta["grupo"]}**',
        f'- **Horario:** **{meta["dia"]} {meta["horario"]}** ({meta["duracion_min"]} min)',
        f"- **Periodo:** 2026-2 · **{dmy(START)} – {dmy(END)}**",
        f'- **Modalidad:** **{meta["modalidad"]}** (Sesión 1 y parciales presencial síncrono · '
        "resto virtual síncrona · festivos = clase autónoma)",
        f"- **Sesiones:** **{len(clases)}** (cubren los **{N_TEMAS} temas** del curso; "
        "2 sesiones son dobles)",
        "",
        "> El **día 1** incluye la **Sesión 0** (Presentación del curso: acuerdo, logística, "
        "Padlet, evaluación, socialización del Proyecto Integrador) **y** el arranque temático.",
        "",
        "| Sesión | Fecha | Tipo | Tema | Material | Parcial |",
        "|---|---|---|---|---|---|",
    ]
    for cl in clases:
        tema = tema_txt(cl)
        if cl.get("sesion_doble"):
            tema = f"**Sesión doble** · {tema}"
        mat = material(cl)
        lines.append(
            f'| {cl["n"]} | {dmy(cl["fecha"])} | {TIPO_LABEL.get(cl["tipo"], cl["tipo"])} | '
            f'{tema} | {mat} | {"sí" if cl.get("parcial") else "no"} |'
        )
    lines += [
        "",
        "## Evaluación por cortes",
        "",
        "| Corte | % | Ventana | Detalle |",
        "|---|---|---|---|",
    ]
    for a, b, pn in corte_rangos():
        c = CORTES[f"corte_{pn}"]
        p = parcial_de_corte(meta, pn)
        det = (
            f"Parcial {pn} en Sesión {p['clase']} ({dmy(p['fecha'])}) · {c['desglose']}"
            if p
            else c["desglose"]
        )
        lines.append(f"| {pn} | {c['pct']} | {dmy(c['inicio'])} – {dmy(c['fin'])} | {det} |")
    sust = meta.get("sustentacion_pi")
    if sust:
        lines += [
            "",
            f"> **Sustentación del Proyecto Integrador:** Sesión {sust['clase']} — "
            f"{dmy(sust['fecha'])}.",
        ]
    lines += [
        "",
        "Fuente derivada del Plan de curso 2026-2. Detalle docente interno no se comparte aquí.",
        "",
    ]
    return "\n".join(lines)


def plan_head(meta: dict) -> str:
    clases = meta["clases"]
    parc = meta.get("parciales", {})
    nums = [str(parc[f"parcial_{i}"]["clase"]) for i in (1, 2, 3) if f"parcial_{i}" in parc]
    fechas = [dmy(parc[f"parcial_{i}"]["fecha"]) for i in (1, 2, 3) if f"parcial_{i}" in parc]
    dobles = [cl for cl in clases if cl.get("sesion_doble")]
    lines = [
        f'# Plan de curso — {meta["nombre"]} · 2026-2',
        "> **CSV eventos:** `Plan curso/2026-2/calendario_eventos_2026-2.csv` (UTF-8 BOM). "
        "Importar en hoja/calendario cuando exista el listado de estudiantes "
        "(una fila = una sesión; filtrar `es_parcial=si` para parciales síncronos).",
        "",
        f'- **Código:** {meta["codigo"]} · **Grupo:** **{meta["grupo"]}**',
        f"- **Periodo:** **2026-2** · **{dmy(START)} – {dmy(END)}**",
        f'- **Horario:** **{meta["dia"]} {meta["horario"]}** ({meta["duracion_min"]} min)',
        f'- **Modalidad:** **{meta["modalidad"]}** (Sesión 1 y parciales presencial síncrono · '
        "resto virtual síncrona · festivos = clase autónoma)",
        f"- **Docente:** {DOCENTE} · `{CORREO}`",
        "- **Calendario:** `Plan curso/2026-2/CALENDARIO_2026-2.md` · "
        "`config/calendario/semestre_2026_2.json`",
        "",
        f"## Ajuste 2026-2: 13 sesiones para {N_TEMAS} temas",
        "",
        resumen_compresion(),
        "",
        "Sesiones dobles de este curso: "
        + " · ".join(
            f'**Sesión {cl["n"]}** ({dmy(cl["fecha"])}) = {material(cl)}' for cl in dobles
        )
        + ".",
        "",
        logica_curso(meta),
        "",
        f"Parciales de este curso: Sesiones **{' / '.join(nums)}** "
        f"({', '.join(fechas)}) — presencial síncrono.",
        "",
        "> **Día de parcial = solo evaluación:** sin tema de trabajo dirigido nuevo.",
        "",
        "> **Sesión 0 (no es sesión temática):** `Clases/Presentacion del Curso - ….pptx` "
        "(logística, acuerdo, Padlet, evaluación, CONTENIDO, socialización del Proyecto "
        "Integrador). En el **día 1** va Sesión 0 + Sesión 1 en el bloque de 120 min.",
        "",
    ]
    sust = meta.get("sustentacion_pi")
    if sust:
        lines += [
            f"> **Sesión {sust['clase']} ({dmy(sust['fecha'])}) = sustentaciones del Proyecto "
            f"Integrador** ({sust['nota']})",
            "",
        ]
    lines += [
        "## Tabla Sesión · Fecha · Tipo · Clase(s) de material · Tema",
        "",
        "| Sesión | Fecha | Tipo | Clase(s) de material | Tema (Trabajo dirigido) |",
        "|---|---|---|---|---|",
    ]
    for cl in clases:
        mat = material(cl) + (" **(doble)**" if cl.get("sesion_doble") else "")
        tema = tema_txt(cl)
        if cl.get("parcial"):
            tema = f"**Parcial {cl['parcial_n']}** (solo evaluación)"
        lines.append(
            f'| {cl["n"]} | {dmy(cl["fecha"])} | {TIPO_LABEL.get(cl["tipo"], cl["tipo"])} | '
            f"{mat} | {tema} |"
        )
    lines += [
        "",
        "## Evaluación teórica (Acuerdo 2026-2)",
        "",
        "| Corte | % | Ventana | Parcial de cierre |",
        "|---|---|---|---|",
    ]
    for a, b, pn in corte_rangos():
        c = CORTES[f"corte_{pn}"]
        p = parcial_de_corte(meta, pn)
        det = (
            f"Parcial {pn} en Sesión {p['clase']} ({dmy(p['fecha'])}) · {c['desglose']}"
            if p
            else c["desglose"]
        )
        lines.append(f"| {pn} | {c['pct']} | {dmy(c['inicio'])} – {dmy(c['fin'])} | {det} |")
    lines.append("")
    return "\n".join(lines)


PRESERVE_FROM = "## Herramientas del curso"


def plan_de_curso_md(meta: dict, existing: str | None) -> str:
    """Cabecera + tabla regeneradas; secciones propias del curso preservadas."""
    tail_parts: list[str] = []
    if existing and PRESERVE_FROM in existing:
        tail = existing[existing.index(PRESERVE_FROM):]
        for chunk in re.split(r"\n(?=## )", tail):
            title = chunk.lstrip("# ").splitlines()[0].strip() if chunk.strip() else ""
            if title.lower().startswith("evaluación teórica"):
                continue  # se regenera arriba
            if title.lower().startswith("ajuste"):
                continue
            tail_parts.append(chunk.rstrip())
    head = plan_head(meta)
    if tail_parts:
        return head + "\n" + "\n\n".join(tail_parts) + "\n"
    return head


# ---------------------------------------------------------------- csv


def csv_rows(meta: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for cl in meta["clases"]:
        hora_i, hora_f = [h.strip() for h in meta["horario"].replace("–", "-").split("-")]
        rows.append(
            [
                meta["nombre"],
                meta["codigo"],
                meta["grupo"],
                str(cl["n"]),
                cl["fecha"],
                meta["dia"],
                hora_i,
                hora_f,
                cl["tipo"],
                "si" if cl.get("parcial") else "no",
                str(cl["parcial_n"]) if cl.get("parcial") else "",
                etiqueta_sesion(cl),
                tema_txt(cl),
                nota_sesion(cl),
            ]
        )
    return rows


def write_csv(path: Path, rows: list[list[str]]) -> None:
    buf = io.StringIO(newline="")
    w = csv.writer(buf, lineterminator="\n")
    w.writerow(CSV_HEADER)
    w.writerows(rows)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(buf.getvalue(), encoding="utf-8-sig")


# ---------------------------------------------------------------- correo bienvenida


def patch_correo(meta: dict, path: Path) -> bool:
    if not path.exists():
        return False
    txt = path.read_text(encoding="utf-8")
    orig = txt
    txt = txt.replace("10/08/2026 – 22/11/2026", f"{dmy(START)} – {dmy(END)}")
    parc = meta.get("parciales", {})
    nums = [str(parc[f"parcial_{i}"]["clase"]) for i in (1, 2, 3) if f"parcial_{i}" in parc]
    dobles = [cl["n"] for cl in meta["clases"] if cl.get("sesion_doble")]
    dia = meta["dia"].lower()
    bullets = [
        f"- **Modalidad por sesión:** **Sesión 1 presencial**; resto de sesiones regulares "
        f"**virtual síncrona**; **parciales presencial** síncrono "
        f"(Sesiones {', '.join(nums[:-1])} y {nums[-1]}). Festivo → **clase autónoma**.",
        f"- **Calendario:** **13 sesiones de {dia}** ({dmy(START)} – {dmy(END)}) que cubren los "
        f"{N_TEMAS} temas del curso; las Sesiones "
        + " y ".join(str(n) for n in dobles)
        + " son **dobles** (dos temas en el mismo bloque de 120 min).",
    ]
    sust = meta.get("sustentacion_pi")
    if sust:
        bullets.append(
            f"- **Sustentación del proyecto final:** Sesión {sust['clase']} — "
            f"{dmy(sust['fecha'])}."
        )
    # elimina bullets previos de modalidad/calendario/encuentros y reinserta los nuevos
    lines = [
        ln
        for ln in txt.splitlines()
        if not re.match(
            r"^- \*\*(Encuentros de|Modalidad por sesión:|Calendario:|Sustentación del proyecto)",
            ln,
        )
    ]
    out: list[str] = []
    inserted = False
    for ln in lines:
        out.append(ln)
        if not inserted and ln.startswith("- **Modalidad:**"):
            out.extend(bullets)
            inserted = True
    if not inserted:
        return False
    txt = "\n".join(out)
    txt = _bloques_gestionados(meta, txt)
    if txt != orig:
        path.write_text(txt, encoding="utf-8")
    return True


# Bloques que este script gestiona dentro del correo. Van entre marcadores para poder
# regenerarlos sin duplicar y sin tocar lo que el docente escribió a mano (por ejemplo
# el link de la carpeta compartida).
MARCA_FECHAS = ("<!-- fechas-clave: generado -->", "<!-- /fechas-clave -->")
MARCA_VOCERO = ("<!-- vocero: generado -->", "<!-- /vocero -->")
MARCA_EXAMLAB = ("<!-- examlab: generado -->", "<!-- /examlab -->")
MARCA_CARPETAS = ("<!-- carpetas: generado -->", "<!-- /carpetas -->")

EXAMLAB_AUTH = "https://examlab.lovable.app/auth"
EXAMLAB_MANUAL = (
    "https://uxxpzfsfcnqiwwdxoelm.supabase.co/storage/v1/object/public/"
    "help-docs/manual-estudiante.pdf"
)
EXAMLAB_VIDEO = (
    "https://uxxpzfsfcnqiwwdxoelm.supabase.co/storage/v1/object/public/"
    "help-videos/serie-estudiante.mp4"
)


def _quitar_bloque(txt: str, marcas: tuple[str, str]) -> str:
    ini, fin = marcas
    while ini in txt and fin in txt:
        a = txt.index(ini)
        b = txt.index(fin) + len(fin)
        txt = (txt[:a].rstrip("\n") + "\n\n" + txt[b:].lstrip("\n")).rstrip() + "\n"
    return txt


def fechas_clave_md(meta: dict) -> str:
    """Tabla de fechas clave del curso, incluida la de la PRIMERA CLASE.

    La primera clase no siempre coincide con el inicio del periodo: el periodo abre el
    24/08 (lunes) pero Programación II arranca el miércoles 26 y Seminario el jueves 27.
    """
    clases = meta["clases"]
    primera, ultima = clases[0], clases[-1]
    dia = meta["dia"]
    filas = [
        ("Inicio del periodo académico", dmy(START), "—"),
        (f"**Primera clase** ({dia})", f"**{dmy(primera['fecha'])}**",
         "Sesión 0 (presentación del curso) + Clase 1 (diagnóstico) en el mismo bloque"),
    ]
    parc = meta.get("parciales", {})
    for i in (1, 2, 3):
        p = parc.get(f"parcial_{i}")
        if p:
            filas.append((f"Parcial {i}", dmy(p["fecha"]),
                          f"Sesión {p['clase']} · presencial síncrono · solo evaluación"))
    sust = meta.get("sustentacion_pi")
    if sust:
        filas.append(("Sustentación del Proyecto Integrador", dmy(sust["fecha"]),
                      f"Sesión {sust['clase']} · en vivo"))
    tipo_ult = "sustentaciones del PI" if ultima.get("sustentacion_pi") else (
        f"Parcial {ultima['parcial_n']}" if ultima.get("parcial") else "cierre del curso")
    filas.append((f"Última clase ({dia})", dmy(ultima["fecha"]),
                  f"Sesión {ultima['n']} · {tipo_ult}"))
    filas.append(("Cierre del periodo académico", dmy(END), "—"))

    out = [MARCA_FECHAS[0], "", "### Fechas clave", "",
           "| Hito | Fecha | Detalle |", "|---|---|---|"]
    out += [f"| {a} | {b} | {c} |" for a, b, c in filas]
    out += ["",
            f"> El curso son **{len(clases)} sesiones de {meta['dia'].lower()}**, "
            f"una por semana, de {dmy(primera['fecha'])} a {dmy(ultima['fecha'])}.",
            ""]

    # Los eventos del calendario llevan el tipo de encuentro al principio del título.
    autonomas = [cl for cl in clases if cl["tipo"] == "autonoma"]
    out += ["Les voy a compartir los eventos del curso por calendario. Cada uno empieza "
            "con el tipo de encuentro, para que sepan de un vistazo si tienen que "
            "conectarse a esa hora:", "",
            "- **`[SINCRONICO]`** — hay encuentro en el horario del curso: presencial, "
            "virtual en vivo, parcial o sustentación. **Deben asistir.**",
            "- **`[AUTONOMO]`** — **no hay encuentro**. Es trabajo independiente guiado: "
            "les dejo el material y la actividad, y ustedes la resuelven por su cuenta "
            "antes de la fecha de cierre."]
    if autonomas:
        det = " y ".join(f"{dmy(cl['fecha'])} ({cl['festivo']})" for cl in autonomas)
        out += ["",
                f"En este curso hay **{len(autonomas)} sesiones autónomas**, porque caen "
                f"en festivo: {det}. **No se pierden**: la clase existe, con material y "
                "entrega, solo que sin encuentro en vivo."]
    else:
        out += ["",
                f"En este curso **todas las sesiones son sincrónicas**: ningún festivo "
                f"cae en {meta['dia'].lower()}."]
    out += ["", MARCA_FECHAS[1]]
    return "\n".join(out)


def examlab_md() -> str:
    """Bloque de ExamLab: qué es, qué se hace ahí y verificación de acceso.

    Se pide verificar el acceso ANTES de la primera clase: si alguien no entra, se
    resuelve fuera del bloque de clase y no se pierde tiempo de la sesión 1.
    """
    return "\n".join([
        MARCA_EXAMLAB[0],
        "",
        "### Plataforma del curso — ExamLab",
        "",
        f"Trabajaremos en **ExamLab**: {EXAMLAB_AUTH}",
        "",
        "**No es una plataforma oficial de la UNIAJC**, pero es donde se desarrolla "
        "todo lo evaluable del curso:",
        "",
        "- **Asistencia**",
        "- **Talleres** (se resuelven y se entregan dentro de la plataforma)",
        "- **Quices y parciales**",
        "- **Entrega del proyecto integrador**",
        "",
        "**Por favor verifiquen que pueden entrar ANTES de la primera clase**, con su "
        "**correo institucional** y esta contraseña temporal:",
        "",
        "> **Contraseña temporal:** ________________________",
        "",
        "Al ingresar por primera vez, cámbienla. Si no logran entrar, escríbanme "
        "respondiendo este correo **antes de la primera sesión**: resolverlo en clase nos "
        "quita tiempo de clase.",
        "",
        "Material de apoyo para usar la plataforma:",
        "",
        f"- **Manual del estudiante (PDF):** {EXAMLAB_MANUAL}",
        f"- **Todas las funcionalidades (video):** {EXAMLAB_VIDEO}",
        "",
        MARCA_EXAMLAB[1],
    ])


def carpetas_md(meta: dict) -> str:
    """Carpetas de Drive del curso: material y grabaciones."""
    c = meta.get("carpetas_drive") or {}
    if not c:
        return ""
    out = [MARCA_CARPETAS[0], "", "### Carpetas del curso en Drive", ""]
    if c.get("clases"):
        out.append(f'- **Clases** (Presentación del Curso, diapositivas y talleres): '
                   f'{c["clases"]["url"]}')
    if c.get("grabadas"):
        out.append(f'- **Clases grabadas** (queda la grabación de cada sesión sincrónica): '
                   f'{c["grabadas"]["url"]}')
    out += ["",
            "Las grabaciones se suben después de cada sesión. Sirven para repasar, y sobre "
            "todo si faltaron: **no reemplazan la asistencia**, que se toma en la sesión.",
            "", MARCA_CARPETAS[1]]
    return "\n".join(out)


def vocero_md() -> str:
    return "\n".join([
        MARCA_VOCERO[0],
        "",
        "**Una cosa que necesito de ustedes:** que el **vocero del grupo** me **responda "
        "este correo con su número de WhatsApp**. Lo uso solo para avisos urgentes del "
        "curso (un cambio de sala, una caída de la plataforma el día de un parcial) y para "
        "tener un canal directo con el grupo. Si todavía no han elegido vocero, lo "
        "definimos en la primera clase y me escribe después.",
        "",
        MARCA_VOCERO[1],
    ])


def _bloques_gestionados(meta: dict, txt: str) -> str:
    txt = _quitar_bloque(txt, MARCA_FECHAS)
    txt = _quitar_bloque(txt, MARCA_VOCERO)
    txt = _quitar_bloque(txt, MARCA_EXAMLAB)
    txt = _quitar_bloque(txt, MARCA_CARPETAS)

    # El bloque de carpetas reemplaza al parrafo heredado con el placeholder a mano.
    txt = "\n".join(
        ln for ln in txt.splitlines()
        if not ln.startswith("**Contenido de las clases**")
        and "[PEGAR AQUÍ LINK DE LA CARPETA CLASES]" not in ln
    )

    # El bullet suelto de ExamLab queda redundante (y traía la URL vieja /app):
    # lo cubre el bloque de plataforma, que además da la URL de acceso correcta.
    txt = "\n".join(ln for ln in txt.splitlines()
                    if not ln.startswith("- **Entrega de talleres y evaluaciones:**"))

    # Fechas clave: después de la lista de bullets, antes del bloque de contenido.
    lineas = txt.splitlines()
    corte = next((i for i, ln in enumerate(lineas)
                  if ln.startswith("**Contenido de las clases**")), None)
    if corte is None:
        corte = next((i for i, ln in enumerate(lineas)
                      if ln.startswith("Por favor **revisen")), len(lineas))
    lineas[corte:corte] = ["", fechas_clave_md(meta), "", carpetas_md(meta), "",
                           examlab_md(), ""]

    # Vocero: justo antes del cierre.
    cierre = next((i for i, ln in enumerate(lineas) if ln.startswith("Nos vemos pronto")),
                  len(lineas))
    lineas[cierre:cierre] = [vocero_md(), ""]

    return re.sub(r"\n{3,}", "\n\n", "\n".join(lineas)).rstrip() + "\n"


# ---------------------------------------------------------------- main


def main() -> None:
    todos: list[list[str]] = []
    for key, meta in DATA["cursos"].items():
        folder = ROOT / meta["folder"]
        periodo_dir = folder / "Plan curso" / "2026-2"
        for sub in [
            "Plan curso/2026-2",
            "Entregas docente/2026-2",
            "Kit docente",
            "Clases",
            "Parciales",
            "Clases grabadas",
        ]:
            (folder / sub).mkdir(parents=True, exist_ok=True)

        cal_path = periodo_dir / "CALENDARIO_2026-2.md"
        cal_path.write_text(calendario_md(meta), encoding="utf-8")

        crono_path = periodo_dir / "Cronograma 2026-2.md"
        crono_path.write_text(cronograma_md(meta), encoding="utf-8")

        plan_path = periodo_dir / "PLAN_DE_CURSO_2026-2.md"
        prev = plan_path.read_text(encoding="utf-8") if plan_path.exists() else None
        plan_path.write_text(plan_de_curso_md(meta, prev), encoding="utf-8")

        rows = csv_rows(meta)
        write_csv(periodo_dir / "calendario_eventos_2026-2.csv", rows)
        write_csv(Path(__file__).with_name(CSV_CONFIG_NAME[key]), rows)
        todos.extend(rows)

        acuerdo = fill_acuerdo(key, meta)
        # El correo de bienvenida es material de PLANEACION del curso, no una entrega a
        # la universidad: vive en Plan curso/<periodo>/. En "Entregas docente/" va solo lo
        # que el docente le entrega a la institucion (acuerdo, diagnostico).
        correo = periodo_dir / f'CORREO_BIENVENIDA - {meta["folder"]} - 2026-2.md'
        if not correo.exists():
            cand = list(periodo_dir.glob("CORREO_BIENVENIDA*.md"))
            if not cand:  # compatibilidad: si quedo en la ubicacion vieja, se migra
                cand = list((folder / "Entregas docente" / "2026-2").glob("CORREO_BIENVENIDA*.md"))
                if cand:
                    destino = periodo_dir / cand[0].name
                    cand[0].replace(destino)
                    print(f"  - correo movido a Plan curso/2026-2/: {destino.name}")
                    cand = [destino]
            correo = cand[0] if cand else correo
        ok_correo = patch_correo(meta, correo)

        auton = sum(1 for c in meta["clases"] if c["tipo"] == "autonoma")
        dobles = sum(1 for c in meta["clases"] if c.get("sesion_doble"))
        print(
            f"OK {meta['nombre']}: {len(meta['clases'])} sesiones "
            f"({auton} autónomas · {dobles} dobles · {N_TEMAS} temas)"
        )
        for p in (cal_path, crono_path, plan_path, acuerdo):
            print(f"  - {p.relative_to(ROOT)}")
        print(f"  - {'correo actualizado' if ok_correo else 'CORREO NO ENCONTRADO'}: {correo.name}")

    write_csv(Path(__file__).with_name("eventos_todos_cursos_2026-2.csv"), todos)
    print(f"OK consolidado: {len(todos)} filas (4 cursos × 13 sesiones)")
    print("NOTA: semestre_2026_2.json es ENTRADA de este script; no se sobreescribe.")


if __name__ == "__main__":
    main()
