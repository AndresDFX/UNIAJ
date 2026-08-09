# -*- coding: utf-8 -*-
"""Aplica modalidad 2026-2 + limpia placeholders Presentación del Curso."""
from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]

REGLA = (
    "Criterio de modalidad por sesión (fijo 2026-2): modalidad del curso = Presencialidad asistida. "
    "Clase 1 = presencial; demás clases regulares = virtual (síncrona); "
    "parciales = siempre presenciales y síncronos (aunque la sesión de otro modo sería virtual); "
    "festivos = clase autónoma (sin parcial). Los parciales NUNCA se programan en día festivo ni en clase autónoma. "
    "Si el cierre teórico del corte cae en festivo/autónoma, el parcial se mueve a la última clase regular "
    "anterior del mismo corte; la clase autónoma de cierre queda como refuerzo sin parcial."
)

MODALIDAD_LABEL = {
    "presencial": "Presencial",
    "virtual": "Virtual (síncrona)",
    "autonoma": "Autónoma (festivo)",
}

CSV_KEY = {
    "programacion_ii": "eventos_programacion_ii_2026-2.csv",
    "seminario": "eventos_seminario_2026-2.csv",
    "bases_datos_ii": "eventos_bases_datos_ii_2026-2.csv",
    "arquitectura": "eventos_arquitectura_2026-2.csv",
}


def tipo_sesion(n: int, festivo, parcial: bool) -> str:
    if festivo:
        return "autonoma"
    if parcial or n == 1:
        return "presencial"
    return "virtual"


def dmy(iso: str) -> str:
    y, m, d = iso.split("-")
    return f"{d}/{m}/{y}"


def dmy_short(iso: str) -> str:
    y, m, d = iso.split("-")
    return f"{d}/{m}"


def read_text(path: Path) -> str:
    raw = path.read_bytes()
    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16")
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    return raw.decode("utf-8")


def write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8", newline="\n")


def replace_regla(text: str) -> str:
    old = (
        "Criterio de parciales (fijo 2026-2): modalidad del curso = Presencialidad asistida. "
        "Las clases pueden ser presencial, virtual (síncrona) o autónoma (p. ej. festivo). "
        "Los parciales son siempre síncronos (presencial o virtual síncrona según la franja del curso) "
        "y NUNCA se programan en día festivo ni en clase autónoma. "
        "Si el cierre teórico del corte cae en festivo/autónoma, el parcial se mueve a la última clase regular "
        "(presencial/virtual síncrona) anterior del mismo corte; la clase autónoma de cierre queda como refuerzo sin parcial."
    )
    if old in text:
        return text.replace(old, REGLA)
    # fallback: replace any prior criterio block starting the same way
    pattern = r"Criterio de (parciales|modalidad por sesión) \(fijo 2026-2\):.*?(?=\n\n|\nParciales|\n##|\n\|)"
    return re.sub(pattern, REGLA + "\n", text, count=1, flags=re.S)


def update_json() -> dict:
    jpath = ROOT / ".config/calendario/semestre_2026_2.json"
    data = json.loads(jpath.read_text(encoding="utf-8"))
    data["logica_evaluacion"] = (
        "Tres cortes 30%/30%/40%. Parcial al finalizar cada corte en la última clase regular del corte "
        "(no en autónoma). Corte 3 incluye Proyecto Integrador. " + REGLA
    )
    data["regla_parciales"] = REGLA
    data["regla_modalidad_sesion"] = (
        "Clase 1 presencial; demás regulares virtual (síncrona); parciales presencial; festivos autónoma."
    )
    data["modalidad_cursos"] = "Presencialidad asistida (todos los cursos)"
    for curso in data["cursos"].values():
        for cl in curso["clases"]:
            cl["tipo"] = tipo_sesion(cl["n"], cl.get("festivo"), bool(cl.get("parcial")))
        for pv in curso.get("parciales", {}).values():
            if isinstance(pv, dict) and "clase" in pv:
                pv["tipo"] = "presencial"
    jpath.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return data


def update_calendarios(data: dict) -> None:
    cortes_info = [
        ("Corte 1", "30%", "2026-08-10 → 2026-09-13", "1-5", 1,
         "10% Parcial 1 · 10% Talleres y Quiz · 10% Asistencia"),
        ("Corte 2", "30%", "2026-09-14 → 2026-10-18", "6-10", 2,
         "10% Parcial 2 · 10% Talleres y Quiz · 10% Asistencia"),
        ("Corte 3", "40%", "2026-10-19 → 2026-11-22", "11-15", 3,
         "15% Parcial 3 · 20% Proyecto Integrador · 5% Asistencia"),
    ]
    for curso in data["cursos"].values():
        folder = ROOT / curso["folder"]
        lines = [
            f"# Calendario 2026-2 — {curso['nombre']}",
            "",
            f"- **Código:** {curso['codigo']}",
            f"- **Grupo:** {curso['grupo']}",
            "- **Periodo:** 2026-2 · **10/08/2026 – 22/11/2026**",
            f"- **Horario:** {curso['dia']} **{curso['horario']}** ({curso['duracion_min']} min)",
            "- **Modalidad:** Presencialidad asistida "
            "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)",
            "- **Docente:** Julian Andres Castaño Espinosa · `julianacastano@profesores.uniajc.edu.co`",
            f"- **Total clases:** {curso['n_clases']} (festivos = **clase autónoma**, no se omiten)",
            "",
            "## Cortes teóricos (30% / 30% / 40%)",
            "",
            REGLA,
            "",
            "| Corte | % | Ventana | Clases | Parcial de cierre | Desglose teórico |",
            "|---|---|---|---|---|---|",
        ]
        pmap = curso.get("parciales", {})
        for title, pct, ventana, clases, pn, desg in cortes_info:
            p = pmap.get(f"parcial_{pn}", {})
            lines.append(
                f"| {title} | {pct} | {ventana} | {clases} | "
                f"Parcial {pn} · Clase {p.get('clase')} ({dmy(p.get('fecha', '?'))}) · Presencial | {desg} |"
            )
        lines += [
            "",
            "> Validar en Acuerdo pedagógico / socialización. Ventanas de corte = bloques temáticos; "
            "la fecha del parcial = última clase regular del corte.",
            "",
            "## Clases",
            "",
            "| Clase | Fecha | Tipo | Nota |",
            "|---|---|---|---|",
        ]
        for cl in curso["clases"]:
            notas = []
            if cl.get("festivo"):
                notas.append(cl["festivo"])
                if not cl.get("parcial"):
                    notas.append("refuerzo sin parcial")
            if cl.get("parcial"):
                pn = cl.get("parcial_n")
                notas.append(f"Parcial {pn} (cierre Corte {pn})")
            nota = " · ".join(notas) if notas else "—"
            lines.append(
                f"| {cl['n']} | {dmy(cl['fecha'])} | {MODALIDAD_LABEL[cl['tipo']]} | {nota} |"
            )
        lines += [
            "",
            "## Festivos Colombia 2026 (rango del periodo)",
            "",
            "- 17/08/2026 — Asunción de la Virgen",
            "- 12/10/2026 — Día de la Diversidad Étnica y Cultural",
            "- 02/11/2026 — Todos los Santos",
            "- 16/11/2026 — Independencia de Cartagena",
            "",
        ]
        write_text(folder / "CALENDARIO_2026-2.md", "\n".join(lines))


def patch_plan_tipo_column(text: str, clases: list[dict]) -> str:
    """Rewrite Tipo column in PLAN markdown tables based on clases list."""
    by_n = {c["n"]: c for c in clases}
    out_lines = []
    for line in text.splitlines():
        m = re.match(r"^\|\s*(\d+)\s*\|\s*([^|]+)\|\s*([^|]+)\|\s*(.*)$", line)
        if not m:
            out_lines.append(line)
            continue
        n = int(m.group(1))
        if n not in by_n:
            out_lines.append(line)
            continue
        fecha = m.group(2).strip()
        tema = m.group(4).rstrip()
        tipo = MODALIDAD_LABEL[by_n[n]["tipo"]]
        out_lines.append(f"| {n} | {fecha} | {tipo} | {tema}")
    return "\n".join(out_lines)


def update_plans(data: dict) -> None:
    for curso in data["cursos"].values():
        path = ROOT / curso["folder"] / "Plan curso" / "PLAN_DE_CURSO_2026-2.md"
        if not path.exists():
            continue
        text = read_text(path)
        text = replace_regla(text)
        # BD II: remove 'franja virtual síncrona' as sole modality hint on horario
        text = text.replace(" (120 min) · franja virtual síncrona", " (120 min)")
        text = text.replace(" · franja virtual síncrona", "")
        # modality bullet
        text = re.sub(
            r"- \*\*Modalidad:\*\* \*\*Presencialidad asistida\*\*.*",
            "- **Modalidad:** **Presencialidad asistida** "
            "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)",
            text,
        )
        text = patch_plan_tipo_column(text, curso["clases"])
        # Prog2/Seminario used "Regular" in tipo — already patched via table rewrite when n matches
        write_text(path, text if text.endswith("\n") else text + "\n")


def update_csv_file(path: Path, clases_by_n: dict[int, dict]) -> None:
    text = path.read_text(encoding="utf-8-sig")
    rows = list(csv.DictReader(text.splitlines()))
    if not rows:
        return
    fieldnames = list(rows[0].keys())
    for row in rows:
        n = int(row["clase_n"])
        cl = clases_by_n[n]
        row["tipo_clase"] = cl["tipo"]
        notes = row.get("notas") or ""
        # normalize parcial notes
        if row.get("es_parcial") == "si":
            notes = re.sub(
                r"(parcial[^;]*;?\s*)+",
                "parcial presencial sincrono; ",
                notes,
                flags=re.I,
            )
            if "parcial presencial sincrono" not in notes.lower():
                notes = ("parcial presencial sincrono; " + notes).strip()
            notes = notes.replace("curso con franja virtual; parcial sincrono presencial; ", "")
            notes = notes.replace("curso con franja virtual; ", "")
        row["notas"] = notes
    out = ROOT / "._tmp_csv.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(rows)


def update_csvs(data: dict) -> None:
    all_rows = []
    fieldnames = None
    for ckey, curso in data["cursos"].items():
        by_n = {c["n"]: c for c in curso["clases"]}
        course_csv = ROOT / curso["folder"] / "Entregas docente" / "calendario_eventos_2026-2.csv"
        cfg_csv = ROOT / ".config/calendario" / CSV_KEY[ckey]
        for path in [course_csv, cfg_csv]:
            if path.exists():
                update_csv_file(path, by_n)
        # collect for todos
        if course_csv.exists():
            with course_csv.open(encoding="utf-8-sig", newline="") as f:
                r = csv.DictReader(f)
                fieldnames = r.fieldnames
                all_rows.extend(list(r))
    todos = ROOT / ".config/calendario/eventos_todos_cursos_2026-2.csv"
    if fieldnames and all_rows:
        with todos.open("w", encoding="utf-8-sig", newline="") as f:
            w = csv.DictWriter(f, fieldnames=fieldnames)
            w.writeheader()
            w.writerows(all_rows)


def update_notas(data: dict) -> None:
    for key in ("bases_datos_ii", "arquitectura"):
        curso = data["cursos"][key]
        path = ROOT / curso["folder"] / "Entregas docente" / "NOTA - Acuerdo y listado 2026-2.md"
        if not path.exists():
            continue
        text = read_text(path)
        text = replace_regla(text)
        text = text.replace(
            "Modalidad: **Presencialidad asistida** (franja virtual síncrona lun 18:00–20:00)",
            "Modalidad: **Presencialidad asistida** "
            "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)",
        )
        text = text.replace(
            "Modalidad: **Presencialidad asistida** · Grupo:",
            "Modalidad: **Presencialidad asistida** "
            "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos) · Grupo:",
        )
        write_text(path, text if text.endswith("\n") else text + "\n")


def patch_generator() -> None:
    path = ROOT / ".config/calendario/generar_semestre_2026_2.py"
    text = read_text(path)
    text = text.replace(
        'LOGICA_EVALUACION = (\n'
        '    "Tres cortes 30%/30%/40%. Parciales siempre sincronos (presencial/virtual sincrona). "\n'
        '    "NUNCA en festivo ni clase autonoma. Criterio: ultima clase regular del corte "\n'
        '    "(mie/jue: 5/10/15; lun Arq/BD II: 5/9/14). Autonomas de cierre = refuerzo sin parcial."\n'
        ')',
        'LOGICA_EVALUACION = (\n'
        '    "' + REGLA.replace('"', '\\"') + '"\n'
        ')',
    )
    # class_dates: keep autonoma for festivo; tipo refined later — update to presencial/virtual
    old_fn = '''def class_dates(weekday: int) -> list[dict]:
    out: list[dict] = []
    d = START
    n = 0
    while d <= END:
        if d.weekday() == weekday:
            n += 1
            iso = d.isoformat()
            out.append(
                {
                    "n": n,
                    "fecha": iso,
                    "tipo": "autonoma" if iso in FESTIVOS else "regular",
                    "festivo": FESTIVOS.get(iso),
                }
            )
        d += timedelta(days=1)
    return out'''
    new_fn = '''def class_dates(weekday: int) -> list[dict]:
    out: list[dict] = []
    d = START
    n = 0
    while d <= END:
        if d.weekday() == weekday:
            n += 1
            iso = d.isoformat()
            festivo = FESTIVOS.get(iso)
            if festivo:
                tipo = "autonoma"
            elif n == 1:
                tipo = "presencial"
            else:
                tipo = "virtual"
            out.append(
                {
                    "n": n,
                    "fecha": iso,
                    "tipo": tipo,
                    "festivo": festivo,
                    "parcial": False,
                }
            )
        d += timedelta(days=1)
    return out


def apply_parciales(clases: list[dict]) -> list[dict]:
    """Marca parciales en última regular del corte; esos días pasan a presencial."""
    ranges = [(1, 5, 1), (6, 10, 2), (11, 15, 3)]
    for a, b, pn in ranges:
        regs = [cl for cl in clases if a <= cl["n"] <= b and cl["tipo"] != "autonoma"]
        if not regs:
            continue
        target = regs[-1]
        target["parcial"] = True
        target["parcial_n"] = pn
        target["tipo"] = "presencial"
    return clases'''
    if old_fn in text:
        text = text.replace(old_fn, new_fn)
    # modalidad strings in CURSOS
    text = text.replace(
        '"modalidad": (\n'
        '            "Presencialidad asistida "\n'
        '            "(confirmar calendario virtual del mes en Campus Virtual)"\n'
        '        )',
        '"modalidad": "Presencialidad asistida '
        '(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)"',
    )
    text = text.replace(
        "Modalidad: Presencialidad asistida (confirmar semanas virtuales en Campus Virtual).\n",
        "Modalidad: Presencialidad asistida "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).\n",
    )
    text = text.replace(
        "Periodo 2026-2 · Grupo 641A-2 · Lunes 18:00–20:00 (120 min) · Modalidad: Virtual.\n",
        "Periodo 2026-2 · Grupo 641A-2 · Lunes 18:00–20:00 (120 min) · "
        "Modalidad: Presencialidad asistida "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).\n",
    )
    text = text.replace(
        "Grupo: [PENDIENTE]. Modalidad: [PENDIENTE].\n",
        "Grupo: [PENDIENTE]. Modalidad: Presencialidad asistida "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).\n",
    )
    # calendario_md tipo labels + apply parciales in main
    text = text.replace(
        'clases = class_dates(meta["weekday"])\n',
        'clases = apply_parciales(class_dates(meta["weekday"]))\n',
    )
    text = text.replace(
        'tipo = "Autónoma (festivo)" if cl["tipo"] == "autonoma" else "Regular"',
        'tipo = {\n'
        '            "autonoma": "Autónoma (festivo)",\n'
        '            "presencial": "Presencial",\n'
        '            "virtual": "Virtual (síncrona)",\n'
        '        }.get(cl["tipo"], cl["tipo"])',
    )
    write_text(path, text)


def patch_build(path: Path, data: dict, ckey: str) -> None:
    text = read_text(path)
    curso = data["cursos"][ckey]

    # Remove Campus Virtual / listado / Meet pendiente lines in string lists
    lines = text.splitlines()
    new_lines = []
    for line in lines:
        low = line.lower()
        drop = False
        if "campus virtual:" in low and "pendiente" in low:
            drop = True
        if "listado de estudiantes" in low and "pendiente" in low:
            drop = True
        if "meet" in low and "pendiente" in low:
            drop = True
        if "campus virtual uniajc:" in low and "pendiente" in low:
            drop = True
        if drop:
            continue
        # strip listado mentions from notes/advertencias
        line = re.sub(r"\s*Listado:\s*\[PENDIENTE listado\]\.?", "", line)
        line = re.sub(r"\s*Listado:\s*\[PENDIENTE listado\]", "", line)
        line = line.replace(" Listado: [PENDIENTE listado].", "")
        line = line.replace(" Listado: [PENDIENTE listado]", "")
        new_lines.append(line)
    text = "\n".join(new_lines)

    # modality wording
    text = text.replace(
        "Modalidad: **presencialidad asistida** (dos semanas al mes la clase puede ser virtual).",
        "Modalidad: **Presencialidad asistida** "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).",
    )
    text = text.replace(
        "Modalidad: **presencialidad asistida** (confirmar virtuales del mes en Campus Virtual).",
        "Modalidad: **Presencialidad asistida** "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).",
    )
    text = text.replace(
        "Modalidad: **Presencialidad asistida** (virtual sincrona)",
        "Modalidad: **Presencialidad asistida** "
        "(Clase 1 presencial · resto virtual · parciales presencial)",
    )
    text = text.replace(
        "Modalidad: **Presencialidad asistida** · franja virtual sincrona (grupo 641A-2).",
        "Modalidad: **Presencialidad asistida** "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos).",
    )
    text = text.replace(
        "Presencialidad asistida (virtual sincrona)",
        "Presencialidad asistida (Clase 1 presencial · resto virtual · parciales presencial)",
    )
    text = text.replace(
        "'Modalidad: **Presencialidad asistida**',",
        "'Modalidad: **Presencialidad asistida** "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)',",
    )

    # Rewrite cronograma table rows (#, fecha short, tipo, tema) from JSON when possible
    # Match lines like: ["1", "12/08", "Regular", "..."],
    def repl_row(m: re.Match) -> str:
        n = int(m.group(1))
        cl = next((c for c in curso["clases"] if c["n"] == n), None)
        if not cl:
            return m.group(0)
        tipo = {
            "presencial": "Presencial",
            "virtual": "Virtual",
            "autonoma": "Autónoma",
        }[cl["tipo"]]
        tema = m.group(4)
        quote = m.group(2)  # ' or "
        return (
            f"{m.group(0)[:m.start(1)-m.start()]}{n}{m.group(0)[m.end(1)-m.start():]}"
        )  # unused fallback

    pattern = re.compile(
        r"""([\[\(]\s*)(['"])(\d+)\2\s*,\s*(['"])([^'"]+)\4\s*,\s*(['"])([^'"]+)\6\s*,\s*(['"])([^'"]*)\8"""
    )

    def repl(m: re.Match) -> str:
        n = int(m.group(3))
        cl = next((c for c in curso["clases"] if c["n"] == n), None)
        if not cl:
            return m.group(0)
        tipo = {
            "presencial": "Presencial",
            "virtual": "Virtual",
            "autonoma": "Autónoma",
        }[cl["tipo"]]
        fecha = dmy_short(cl["fecha"])
        q1, q2, q3, q4 = m.group(2), m.group(4), m.group(6), m.group(8)
        tema = m.group(9)
        return (
            f"{m.group(1)}{q1}{n}{q1}, {q2}{fecha}{q2}, {q3}{tipo}{q3}, {q4}{tema}{q4}"
        )

    text = pattern.sub(repl, text)
    write_text(path, text if text.endswith("\n") else text + "\n")


def update_uniajc_json() -> None:
    path = ROOT / ".config/universidades/uniajc.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    data["_version"] = "1.5"
    data["_actualizado"] = "2026-08-07"
    data["_cursos_workspace"] = [
        "Programacion II (FI303204 · grupo 341C · 2026-2 · mié 18:00-20:00 · 120 min · Presencialidad asistida · Clase 1 presencial / resto virtual / parciales presencial)",
        "Seminario de Sistemas (FI303301 · grupo 341C · 2026-2 · jue 18:00-20:00 · 120 min · Presencialidad asistida · Clase 1 presencial / resto virtual / parciales presencial)",
        "Bases de Datos II (FI303215 · grupo 641A-2 · 2026-2 · lun 18:00-20:00 · 120 min · Presencialidad asistida · Clase 1 presencial / resto virtual / parciales presencial)",
        "Arquitectura de Sistemas Computacionales (FI303380 · 2026-2 · lun 10:00-13:00 · 180 min · Presencialidad asistida · Clase 1 presencial / resto virtual / parciales presencial · grupo pendiente)",
        "0. Base (referencia Ingeniería de Software 1 / plantillas institucionales)",
    ]
    # keep URL Campus as known pendiente for ops, but not for slides
    data["_pendientes"] = [
        "URL Campus Virtual / plataforma de entregas por curso (NO incluir placeholder en Presentación del Curso)",
        "URL Meet (si aplica) por curso",
        "Grupo de Arquitectura",
        "Listados de estudiantes / voceros / firmas de Acuerdos 2026-2 (NO incluir placeholder de listado en Presentación del Curso)",
        "Confirmar tipografía oficial del Manual de Identidad (hoy Calibri como fallback web/docente)",
        "Migrar gradualmente .gslides/.gdoc existentes a .pptx/.docx locales sin borrar los atajos de Google",
    ]
    data["periodo_vigente"]["regla_parciales"] = REGLA
    data["periodo_vigente"]["regla_modalidad_sesion"] = (
        "Clase 1 presencial; demás regulares virtual (síncrona); parciales presencial; festivos autónoma."
    )
    data["periodo_vigente"]["modalidad_cursos"] = "Presencialidad asistida"
    est = data["estandar_material"]["presentacion_del_curso"]
    est["incluye"] = [
        "docente",
        "grupo/periodo en negrita",
        "horario",
        "evaluación Acuerdo",
        "cronograma Plan de curso",
        "Padlet + QR institucional (mismo para todos los cursos)",
        "cierre día/hora semanal",
    ]
    est["prohibido"] = [
        "placeholder de URL Campus Virtual",
        "Listado de estudiantes / [PENDIENTE listado]",
        "placeholder de Meet (salvo URL real confirmada)",
    ]
    est["negrita_oferta"] = "grupo, periodo, horario (sin placeholders de URL Campus/listado)"
    data["pedagogia"]["_nota_duracion"] = (
        "Periodo 2026-2 (10/08/2026–22/11/2026). Todos los cursos: Presencialidad asistida. "
        "Modalidad por sesión: Clase 1 presencial; demás regulares virtual (síncrona); "
        "parciales presencial; festivos autónoma. "
        "Duraciones: Programación II 341C mié 18:00–20:00 = 120 min; Seminario 341C jue 18:00–20:00 = 120 min; "
        "Bases de Datos II 641A-2 lun 18:00–20:00 = 120 min; Arquitectura lun 10:00–13:00 = 180 min."
    )
    data["pedagogia"]["metodologia_observada_prog2"]["modalidad"] = (
        "Presencialidad asistida (Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)"
    )
    data["evaluacion"]["_nota"] = (
        "Copia SIEMPRE del Acuerdo pedagógico del grupo. Teórico 2026-2: 30/30/40. " + REGLA
    )
    data["evaluacion"]["teorico_2026_2"]["_logica"] = (
        "Tres cortes 30%/30%/40%. Parcial al finalizar cada corte en la última clase regular del corte "
        "(no en autónoma). Corte 3 incluye Proyecto Integrador. " + REGLA
    )
    data["evaluacion"]["teorico_2026_2"]["criterio_parciales"] = REGLA
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def update_rules_and_agents() -> None:
    rule_path = ROOT / ".cursor/rules/uniajc-docente.mdc"
    text = read_text(rule_path)
    text = text.replace(
        "- **Modalidad (todos los cursos):** **Presencialidad asistida** (clases presencial / virtual síncrona / autónoma).\n"
        "- **Regla de festivos:** no omitir clases; marcarlas como **clase autónoma**.\n"
        "- **Regla de parciales:** siempre síncronos; **NUNCA** en festivo ni clase autónoma. "
        "Si el cierre del corte es autónoma → parcial en la **última clase regular anterior** del corte.",
        "- **Modalidad (todos los cursos):** **Presencialidad asistida**.\n"
        "- **Modalidad por sesión:** Clase 1 = **presencial**; demás regulares = **virtual**; "
        "**parciales = presencial** (síncronos); festivos = **autónoma**.\n"
        "- **Regla de festivos:** no omitir clases; marcarlas como **clase autónoma**.\n"
        "- **Regla de parciales:** siempre **presenciales** y síncronos; **NUNCA** en festivo ni clase autónoma. "
        "Si el cierre del corte es autónoma → parcial en la **última clase regular anterior** del corte.",
    )
    text = text.replace(
        "3. **Bases de Datos II** (FI303215 · grupo **641A-2** · lun **18:00–20:00** · **120 min** · **Presencialidad asistida** · franja virtual síncrona)",
        "3. **Bases de Datos II** (FI303215 · grupo **641A-2** · lun **18:00–20:00** · **120 min** · **Presencialidad asistida**)",
    )
    # Presentación del curso bullets
    text = text.replace(
        "4. Presentación del curso (para **estudiantes**): grupo(s), docente, recursos, evaluación, cronograma, Padlet rompe-hielo, cierre. **Sin** checklist/rutina operativa del docente ni Clear posts.",
        "4. Presentación del curso (para **estudiantes**): grupo(s), docente, recursos, evaluación, cronograma, Padlet rompe-hielo, cierre. "
        "**Sin** checklist/rutina operativa del docente ni Clear posts. "
        "**Sin** placeholder de Campus Virtual ni de listado de estudiantes.",
    )
    text = text.replace(
        "5. Datos de oferta del periodo en **negrita** en Presentación del Curso (grupo, periodo, horario, placeholders de URL). **No** forzar negrita en docente/correo (fijos), temario ni textos pedagógicos.",
        "5. Datos de oferta del periodo en **negrita** en Presentación del Curso (grupo, periodo, horario). "
        "**No** forzar negrita en docente/correo (fijos), temario ni textos pedagógicos. "
        "**No** incluir placeholders de URL Campus Virtual ni listado.",
    )
    text = text.replace(
        "- URL Campus Virtual por curso — placeholder: `[URL Campus Virtual UNIAJC — pendiente]`\n"
        "- URL Meet (si aplica)\n"
        "- Códigos oficiales BD II y Arquitectura; grupo y modalidad de Arquitectura\n"
        "- Tipografía oficial del Manual de Identidad (hoy Calibri)\n"
        "- Listados de estudiantes / voceros / firmas de Acuerdos 2026-2",
        "- URL Campus Virtual / Meet (ops docente; **no** poner placeholders en Presentación del Curso)\n"
        "- Grupo de Arquitectura\n"
        "- Tipografía oficial del Manual de Identidad (hoy Calibri)\n"
        "- Listados / voceros / firmas de Acuerdos 2026-2 (**no** en Presentación del Curso)",
    )
    # where bio section
    text = text.replace(
        "   - **SÍ** docente + correo + logística (horario, **grupo(s)**, periodo, Campus Virtual, evaluación, fechas de oferta en **negrita** si aplican).",
        "   - **SÍ** docente + correo + logística (horario, **grupo(s)**, periodo, evaluación, fechas de oferta en **negrita** si aplican). "
        "**NO** placeholder Campus Virtual ni listado de estudiantes.",
    )
    write_text(rule_path, text if text.endswith("\n") else text + "\n")

    agent_snippet_old = (
        "8. Evaluación teórica: 30/30/40. Parciales **siempre síncronos**; **NUNCA** en festivo/autónoma. "
        "Criterio: última clase regular del corte (mié/jue: 5/10/15; lun Arq/BD II: **5/9/14**). "
        "Festivos = clase autónoma de refuerzo. Modalidad de todos los cursos: **Presencialidad asistida**."
    )
    agent_snippet_new = (
        "8. Evaluación teórica: 30/30/40. Parciales **siempre presenciales y síncronos**; **NUNCA** en festivo/autónoma. "
        "Criterio: última clase regular del corte (mié/jue: 5/10/15; lun Arq/BD II: **5/9/14**). "
        "Festivos = clase autónoma de refuerzo. Modalidad del curso: **Presencialidad asistida**. "
        "Por sesión: Clase 1 presencial · resto virtual · parciales presencial · festivos autónoma."
    )
    for ap in [
        ROOT / ".claude/agents/disenador-curricular-uniajc.md",
        ROOT / ".cursor/agents/disenador-curricular-uniajc.md",
    ]:
        if not ap.exists():
            continue
        t = read_text(ap)
        t = t.replace(agent_snippet_old, agent_snippet_new)
        t = t.replace(
            "No inventar listado de estudiantes: dejar [PENDIENTE listado] en Presentación del Curso / Acuerdo.",
            "No inventar listado de estudiantes. En Presentación del Curso **NO** incluir placeholder de listado "
            "ni de URL Campus Virtual. En Acuerdo sí puede quedar [PENDIENTE] en campos de nómina/firmas.",
        )
        t = t.replace(
            "1. Portada (asignatura, código FI…, grupo, periodo, programa, horario en **negrita**; listado [PENDIENTE listado])",
            "1. Portada (asignatura, código FI…, grupo, periodo, programa, horario en **negrita**; "
            "**sin** listado ni placeholder Campus Virtual)",
        )
        t = t.replace(
            "6. Metodología (ABPr, presencialidad asistida / virtual según curso, Teoría + Taller + Quiz)",
            "6. Metodología (ABPr, Presencialidad asistida: Clase 1 presencial · resto virtual · parciales presencial, "
            "Teoría + Taller + Quiz)",
        )
        t = t.replace(
            "10. Recursos (Campus Virtual placeholder, bibliografía del microcurrículo, IDE)",
            "10. Recursos (bibliografía del microcurrículo, IDE/herramientas; **sin** placeholder Campus Virtual)",
        )
        t = t.replace(
            "9. Pendientes conocidos: URL Campus Virtual, Meet, grupo Arquitectura, listados de estudiantes. Padlet URL ya fija.",
            "9. Pendientes ops (no van en PPTX curso): URL Campus/Meet, grupo Arquitectura, listados/firmas Acuerdo. Padlet URL ya fija.",
        )
        t = t.replace(
            "*v1.3 — Variante UNIAJC · 2026-2 · Presencialidad asistida · parciales síncronos (nunca en autónoma) · última regular del corte · Motor `uniajc_slides_engine.py`.*",
            "*v1.4 — Variante UNIAJC · 2026-2 · Presencialidad asistida · Clase 1 presencial / resto virtual / parciales presencial · "
            "festivos autónoma · sin placeholders Campus/listado en Presentación del Curso · Motor `uniajc_slides_engine.py`.*",
        )
        # also system prompt copy inside agent if duplicated in ENTREGABLE
        write_text(ap, t if t.endswith("\n") else t + "\n")

    dudas = ROOT / ".claude/agents/uniajc-dudas-material.md"
    if dudas.exists():
        t = read_text(dudas)
        t = t.replace(
            "| Bases de Datos II | 641A-2 · lun 18:00–20:00 virtual (120 min) | Skeleton + Acuerdo/Calendario 2026-2 |",
            "| Bases de Datos II | 641A-2 · lun 18:00–20:00 · Presencialidad asistida (120 min) | Skeleton + Acuerdo/Calendario 2026-2 |",
        )
        write_text(dudas, t if t.endswith("\n") else t + "\n")


def update_acuerdos_docx() -> None:
    """Patch methodology cell text in Acuerdos 2026-2 if present."""
    try:
        from docx import Document
    except ImportError:
        print("WARN: python-docx no disponible; Acuerdos no tocados")
        return
    modality = (
        "Modalidad: Presencialidad asistida "
        "(Clase 1 presencial · resto virtual · parciales presencial · festivos autónomos)."
    )
    paths = list((ROOT).glob("**/ACUERDO PEDAGOGICO*2026-2.docx"))
    for path in paths:
        doc = Document(str(path))
        changed = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text
                    if not txt:
                        continue
                    new = txt
                    new = new.replace(
                        "Modalidad: Presencialidad asistida (confirmar semanas virtuales en Campus Virtual).",
                        modality,
                    )
                    new = new.replace(
                        "Modalidad: Virtual.",
                        modality,
                    )
                    new = new.replace(
                        "Modalidad: [PENDIENTE].",
                        modality,
                    )
                    if "Modalidad:" in new and "Clase 1 presencial" not in new and "Presencialidad asistida" in new:
                        new = re.sub(
                            r"Modalidad:\s*Presencialidad asistida[^\n]*",
                            modality.rstrip("."),
                            new,
                        )
                        if not new.endswith("\n") and modality in new:
                            pass
                    if new != txt:
                        # set cell text via paragraphs
                        paras = cell.paragraphs
                        if paras:
                            runs = paras[0].runs
                            if runs:
                                runs[0].text = new
                                for r in runs[1:]:
                                    r.text = ""
                            else:
                                paras[0].add_run(new)
                            for p in paras[1:]:
                                for r in p.runs:
                                    r.text = ""
                        changed = True
        if changed:
            doc.save(str(path))
            print("OK ACUERDO", path.relative_to(ROOT))


def rebuild_pptx() -> None:
    builds = [
        "build_uniajc_prog2_curso.py",
        "build_uniajc_seminario_curso.py",
        "build_uniajc_bd2_curso.py",
        "build_uniajc_arq_curso.py",
    ]
    for name in builds:
        script = ROOT / ".config/slides" / name
        print("BUILD", name)
        r = subprocess.run([sys.executable, str(script)], cwd=str(ROOT), capture_output=True, text=True)
        print(r.stdout)
        if r.returncode != 0:
            print(r.stderr)
            raise SystemExit(f"Build falló: {name}")


def print_summary(data: dict) -> None:
    print("\n=== RESUMEN MODALIDAD ===")
    for ckey, curso in data["cursos"].items():
        c1 = next(c for c in curso["clases"] if c["n"] == 1)
        parciales = [c for c in curso["clases"] if c.get("parcial")]
        print(f"\n{curso['nombre']}:")
        print(f"  Clase 1: {c1['fecha']} → {c1['tipo']}")
        for p in parciales:
            print(f"  Parcial {p.get('parcial_n')}: Clase {p['n']} {p['fecha']} → {p['tipo']}")
        counts = {}
        for c in curso["clases"]:
            counts[c["tipo"]] = counts.get(c["tipo"], 0) + 1
        print("  Conteos:", counts)


def main() -> None:
    data = update_json()
    update_calendarios(data)
    update_plans(data)
    update_csvs(data)
    update_notas(data)
    patch_generator()
    update_uniajc_json()
    update_rules_and_agents()
    for ckey, script in [
        ("programacion_ii", "build_uniajc_prog2_curso.py"),
        ("seminario", "build_uniajc_seminario_curso.py"),
        ("bases_datos_ii", "build_uniajc_bd2_curso.py"),
        ("arquitectura", "build_uniajc_arq_curso.py"),
    ]:
        patch_build(ROOT / ".config/slides" / script, data, ckey)
    update_acuerdos_docx()
    rebuild_pptx()
    print_summary(data)
    print("\nDONE")


if __name__ == "__main__":
    main()
