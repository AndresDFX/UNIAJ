# -*- coding: utf-8 -*-
"""One-shot: Arquitectura 10:00-13:00/180 -> 10:00-12:00/120."""
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
touched: list[str] = []


def replace_in_file(path: Path, pairs: list[tuple[str, str]], encoding: str = "utf-8") -> bool:
    text = path.read_text(encoding=encoding)
    orig = text
    for a, b in pairs:
        text = text.replace(a, b)
    if text != orig:
        path.write_text(text, encoding=encoding)
        touched.append(str(path.relative_to(ROOT)).replace("\\", "/"))
        return True
    return False


def patch_csvs() -> None:
    csv_paths = [
        ROOT / "Arquitectura de Sistemas Computacionales/Plan curso/calendario_eventos_2026-2.csv",
        ROOT / ".config/calendario/eventos_arquitectura_2026-2.csv",
        ROOT / ".config/calendario/eventos_todos_cursos_2026-2.csv",
    ]
    for p in csv_paths:
        text = p.read_text(encoding="utf-8-sig")
        lines = text.splitlines(keepends=True)
        new_lines = []
        changed = False
        for line in lines:
            if line.startswith("Arquitectura de Sistemas Computacionales,") and ",10:00,13:00," in line:
                line = line.replace(",10:00,13:00,", ",10:00,12:00,", 1)
                changed = True
            new_lines.append(line)
        if changed:
            p.write_text("".join(new_lines), encoding="utf-8-sig")
            touched.append(str(p.relative_to(ROOT)).replace("\\", "/"))


def patch_acuerdo() -> None:
    acuerdo = (
        ROOT
        / "Arquitectura de Sistemas Computacionales"
        / "Entregas docente"
        / "ACUERDO PEDAGOGICO - Arquitectura de Sistemas Computacionales - 2026-2.docx"
    )
    doc = Document(str(acuerdo))
    replacements = [
        ("10:00–13:00 (180 min)", "10:00–12:00 (120 min)"),
        ("10:00 – 13:00 (180 min)", "10:00 – 12:00 (120 min)"),
        ("10:00-13:00 (180 min)", "10:00-12:00 (120 min)"),
        ("10:00\u201313:00 (180 min)", "10:00\u201312:00 (120 min)"),
        ("10:00 \u2013 13:00 (180 min)", "10:00 \u2013 12:00 (120 min)"),
    ]

    def patch_paragraph(p) -> bool:
        full = p.text
        new = full
        for a, b in replacements:
            new = new.replace(a, b)
        if new != full:
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new)
            return True
        return False

    changed_doc = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if patch_paragraph(p):
                        changed_doc = True
    for p in doc.paragraphs:
        if patch_paragraph(p):
            changed_doc = True

    if changed_doc:
        doc.save(str(acuerdo))
        touched.append(str(acuerdo.relative_to(ROOT)).replace("\\", "/"))
    else:
        print("ACUERDO: no change")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "10:00" in cell.text or "180" in cell.text:
                        print(repr(cell.text[:250]))


def main() -> None:
    patch_csvs()

    file_pairs: dict[Path, list[tuple[str, str]]] = {
        ROOT / "Arquitectura de Sistemas Computacionales/Plan curso/PLAN_DE_CURSO_2026-2.md": [
            ("**Lunes 10:00 – 13:00** (180 min)", "**Lunes 10:00 – 12:00** (120 min)"),
            ("**Lunes 10:00–13:00** (180 min)", "**Lunes 10:00–12:00** (120 min)"),
            ("**Lunes 10:00 \u2013 13:00** (180 min)", "**Lunes 10:00 \u2013 12:00** (120 min)"),
        ],
        ROOT / "Arquitectura de Sistemas Computacionales/Plan curso/CALENDARIO_2026-2.md": [
            ("Lunes **10:00 – 13:00** (180 min)", "Lunes **10:00 – 12:00** (120 min)"),
            ("Lunes **10:00–13:00** (180 min)", "Lunes **10:00–12:00** (120 min)"),
            ("Lunes **10:00 \u2013 13:00** (180 min)", "Lunes **10:00 \u2013 12:00** (120 min)"),
        ],
        ROOT / ".config/calendario/semestre_2026_2.json": [
            ('"horario": "10:00 – 13:00"', '"horario": "10:00 – 12:00"'),
            ('"horario": "10:00 \u2013 13:00"', '"horario": "10:00 \u2013 12:00"'),
            ('"duracion_min": 180', '"duracion_min": 120'),
        ],
        ROOT / ".config/calendario/generar_semestre_2026_2.py": [
            ('"horario": "10:00 – 13:00"', '"horario": "10:00 – 12:00"'),
            ('"horario": "10:00 \u2013 13:00"', '"horario": "10:00 \u2013 12:00"'),
            ('"duracion_min": 180', '"duracion_min": 120'),
            ("Periodo 2026-2 · Lunes 10:00–13:00 (180 min).\n", "Periodo 2026-2 · Lunes 10:00–12:00 (120 min).\n"),
            (
                "Periodo 2026-2 · Lunes 10:00\u201313:00 (180 min).\n",
                "Periodo 2026-2 · Lunes 10:00\u201312:00 (120 min).\n",
            ),
        ],
        ROOT / ".config/universidades/uniajc.json": [
            ("lun 10:00-13:00 · 180 min", "lun 10:00-12:00 · 120 min"),
            ("Arquitectura lun 10:00–13:00 = 180 min.", "Arquitectura lun 10:00–12:00 = 120 min."),
            (
                "Arquitectura lun 10:00\u201313:00 = 180 min.",
                "Arquitectura lun 10:00\u201312:00 = 120 min.",
            ),
            ('"Arquitectura de Sistemas Computacionales": 180', '"Arquitectura de Sistemas Computacionales": 120'),
        ],
        ROOT / ".cursor/rules/uniajc-docente.mdc": [
            ("lun **10:00–13:00** · **180 min**", "lun **10:00–12:00** · **120 min**"),
            ("lun **10:00\u201313:00** · **180 min**", "lun **10:00\u201312:00** · **120 min**"),
            (
                "usar duración real del curso: 120 o 180 min",
                "usar duración real del curso: **120 min**",
            ),
            (
                "Bloque según curso (Prog. II / Seminario / BD II = **120 min**; Arquitectura = **180 min**).",
                "Bloque según curso (todos los cursos activos 2026-2 = **120 min**, incl. Arquitectura).",
            ),
        ],
        ROOT / ".cursor/agents/disenador-curricular-uniajc.md": [
            (
                "5. Tiempo real del bloque (120 min Prog. II / Seminario / BD II; 180 min Arquitectura).",
                "5. Tiempo real del bloque (**120 min** en todos los cursos activos 2026-2, incl. Arquitectura lun 10:00–12:00).",
            ),
        ],
        ROOT / ".claude/agents/disenador-curricular-uniajc.md": [
            (
                "5. Tiempo real del bloque (120 min Prog. II / Seminario / BD II; 180 min Arquitectura).",
                "5. Tiempo real del bloque (**120 min** en todos los cursos activos 2026-2, incl. Arquitectura lun 10:00–12:00).",
            ),
        ],
        ROOT / ".cursor/agents/uniajc-dudas-material.md": [
            (
                "| Arquitectura de Sist. Comp. | lun 10:00–13:00 (180 min) |",
                "| Arquitectura de Sist. Comp. | lun 10:00–12:00 (120 min) |",
            ),
            (
                "| Arquitectura de Sist. Comp. | lun 10:00\u201313:00 (180 min) |",
                "| Arquitectura de Sist. Comp. | lun 10:00\u201312:00 (120 min) |",
            ),
        ],
        ROOT / ".claude/agents/uniajc-dudas-material.md": [
            (
                "| Arquitectura de Sist. Comp. | lun 10:00–13:00 (180 min) · Presencialidad asistida |",
                "| Arquitectura de Sist. Comp. | lun 10:00–12:00 (120 min) · Presencialidad asistida |",
            ),
            (
                "| Arquitectura de Sist. Comp. | lun 10:00\u201313:00 (180 min) · Presencialidad asistida |",
                "| Arquitectura de Sist. Comp. | lun 10:00\u201312:00 (120 min) · Presencialidad asistida |",
            ),
        ],
        ROOT / "LEEME - Mapa de cursos y agentes.md": [
            ("lun **10:00–13:00** (180 min)", "lun **10:00–12:00** (120 min)"),
            ("lun **10:00\u201313:00** (180 min)", "lun **10:00\u201312:00** (120 min)"),
        ],
        ROOT / ".config/slides/build_uniajc_arq_curso.py": [
            (
                "'Horario: **Lunes 10:00 – 13:00** (180 min)'",
                "'Horario: **Lunes 10:00 – 12:00** (120 min)'",
            ),
            (
                "'Horario: **Lunes 10:00 \u2013 13:00** (180 min)'",
                "'Horario: **Lunes 10:00 \u2013 12:00** (120 min)'",
            ),
            (
                "'Bloque de **180 min** (lunes 10:00–13:00).'",
                "'Bloque de **120 min** (lunes 10:00–12:00).'",
            ),
            (
                "'Bloque de **180 min** (lunes 10:00\u201313:00).'",
                "'Bloque de **120 min** (lunes 10:00\u201312:00).'",
            ),
            (
                "('info', 'Lunes 10:00-13:00 (180 min). Modalidad: Presencialidad asistida (Clase 1 presencial / resto virtual / parciales presencial). Grupo: [PENDIENTE].')",
                "('info', 'Lunes 10:00-12:00 (120 min). Modalidad: Presencialidad asistida (Clase 1 presencial / resto virtual / parciales presencial). Grupo: [PENDIENTE].')",
            ),
            ("'Lunes **10:00 – 13:00**'", "'Lunes **10:00 – 12:00**'"),
            ("'Lunes **10:00 \u2013 13:00**'", "'Lunes **10:00 \u2013 12:00**'"),
        ],
    }

    for path, pairs in file_pairs.items():
        ok = replace_in_file(path, pairs)
        if not ok:
            print("NO CHANGE:", path.relative_to(ROOT))
            t = path.read_text(encoding="utf-8")
            for i, line in enumerate(t.splitlines(), 1):
                if ("13:00" in line or "180" in line) and (
                    "10:00" in line or "Arquitectura" in line or "duracion" in line or "180 min" in line
                ):
                    print(f"  L{i}: {line[:140]}")

    patch_acuerdo()

    print("TOUCHED", len(touched))
    for t in touched:
        print(" -", t)


if __name__ == "__main__":
    main()
