# -*- coding: utf-8 -*-
"""Corrección crítica: parciales nunca en festivo/autónoma (2026-2)."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]

CRITERIO = (
    "Criterio de parciales (fijo 2026-2): modalidad del curso = Presencialidad asistida. "
    "Las clases pueden ser presencial, virtual (síncrona) o autónoma (p. ej. festivo). "
    "Los parciales son siempre síncronos (presencial o virtual síncrona según la franja del curso) "
    "y NUNCA se programan en día festivo ni en clase autónoma. "
    "Si el cierre teórico del corte cae en festivo/autónoma, el parcial se mueve a la "
    "última clase regular (presencial/virtual síncrona) anterior del mismo corte; "
    "la clase autónoma de cierre queda como refuerzo sin parcial."
)

LOGICA_CORTES = (
    "Tres cortes 30%/30%/40%. Parcial al finalizar cada corte en la última clase regular "
    "del corte (no en autónoma). Corte 3 incluye Proyecto Integrador."
)

PARCIALES = {
    "programacion_ii": {1: 5, 2: 10, 3: 15},
    "seminario": {1: 5, 2: 10, 3: 15},
    "bases_datos_ii": {1: 5, 2: 9, 3: 14},
    "arquitectura": {1: 5, 2: 9, 3: 14},
}

FECHA_ES = {
    "2026-08-10": "10/08/2026",
    "2026-08-12": "12/08/2026",
    "2026-08-13": "13/08/2026",
    "2026-08-17": "17/08/2026",
    "2026-08-19": "19/08/2026",
    "2026-08-20": "20/08/2026",
    "2026-08-24": "24/08/2026",
    "2026-08-26": "26/08/2026",
    "2026-08-27": "27/08/2026",
    "2026-08-31": "31/08/2026",
    "2026-09-02": "02/09/2026",
    "2026-09-03": "03/09/2026",
    "2026-09-07": "07/09/2026",
    "2026-09-09": "09/09/2026",
    "2026-09-10": "10/09/2026",
    "2026-09-14": "14/09/2026",
    "2026-09-16": "16/09/2026",
    "2026-09-17": "17/09/2026",
    "2026-09-21": "21/09/2026",
    "2026-09-23": "23/09/2026",
    "2026-09-24": "24/09/2026",
    "2026-09-28": "28/09/2026",
    "2026-09-30": "30/09/2026",
    "2026-10-01": "01/10/2026",
    "2026-10-05": "05/10/2026",
    "2026-10-07": "07/10/2026",
    "2026-10-08": "08/10/2026",
    "2026-10-12": "12/10/2026",
    "2026-10-14": "14/10/2026",
    "2026-10-15": "15/10/2026",
    "2026-10-19": "19/10/2026",
    "2026-10-21": "21/10/2026",
    "2026-10-22": "22/10/2026",
    "2026-10-26": "26/10/2026",
    "2026-10-28": "28/10/2026",
    "2026-10-29": "29/10/2026",
    "2026-11-02": "02/11/2026",
    "2026-11-04": "04/11/2026",
    "2026-11-05": "05/11/2026",
    "2026-11-09": "09/11/2026",
    "2026-11-11": "11/11/2026",
    "2026-11-12": "12/11/2026",
    "2026-11-16": "16/11/2026",
    "2026-11-18": "18/11/2026",
    "2026-11-19": "19/11/2026",
}


def fecha_corta(iso: str) -> str:
    # dd/mm
    y, m, d = iso.split("-")
    return f"{d}/{m}"


def fecha_larga(iso: str) -> str:
    return FECHA_ES.get(iso, iso)


def tipo_label(tipo: str, festivo: str | None) -> str:
    if tipo == "autonoma":
        return "Autónoma (festivo)" if festivo else "Autónoma"
    if tipo == "virtual":
        return "Virtual (síncrona)"
    return "Presencial"


def update_calendario_json() -> dict:
    path = ROOT / ".config" / "calendario" / "semestre_2026_2.json"
    cal = json.loads(path.read_text(encoding="utf-8"))
    cal["_actualizado"] = "2026-08-07"
    cal["modalidad_cursos"] = "Presencialidad asistida (todos los cursos)"
    cal["tipos_clase"] = ["presencial", "virtual", "autonoma"]
    cal["regla_festivos"] = (
        "No omitir. Marcar como clase autónoma. Los parciales NUNCA van en festivo/autónoma."
    )
    cal["regla_parciales"] = CRITERIO
    cal["logica_evaluacion"] = f"{LOGICA_CORTES} {CRITERIO}"
    cal["cortes_teoricos"] = {
        "corte_1": {
            "pct": "30%",
            "inicio": "2026-08-10",
            "fin": "2026-09-13",
            "clases": "1-5",
            "parcial_cierre_regla": "Última clase regular del corte (típicamente Clase 5 si es regular)",
            "desglose": "10% Parcial 1 (cierre de corte) · 10% Talleres y Quiz · 10% Asistencia",
        },
        "corte_2": {
            "pct": "30%",
            "inicio": "2026-09-14",
            "fin": "2026-10-18",
            "clases": "6-10",
            "parcial_cierre_regla": "Última clase regular del corte (si Clase 10 es autónoma → Clase 9)",
            "desglose": "10% Parcial 2 (cierre de corte) · 10% Talleres y Quiz · 10% Asistencia",
        },
        "corte_3": {
            "pct": "40%",
            "inicio": "2026-10-19",
            "fin": "2026-11-22",
            "clases": "11-15",
            "parcial_cierre_regla": "Última clase regular del corte (si Clase 15 es autónoma → Clase 14)",
            "desglose": "15% Parcial 3 (cierre de corte) · 20% Proyecto Integrador · 5% Asistencia",
        },
    }

    for key, curso in cal["cursos"].items():
        curso["modalidad"] = "Presencialidad asistida"
        pmap = PARCIALES[key]
        for c in curso["clases"]:
            if c.get("festivo"):
                c["tipo"] = "autonoma"
            else:
                c["tipo"] = "virtual" if key == "bases_datos_ii" else "presencial"
            c["parcial"] = False
            c.pop("parcial_n", None)
            for pn, cn in pmap.items():
                if c["n"] == cn:
                    c["parcial"] = True
                    c["parcial_n"] = pn
        by_n = {c["n"]: c for c in curso["clases"]}
        curso["parciales"] = {
            "criterio": "Última clase regular del corte; nunca festivo/autónoma",
        }
        for pn, cn in pmap.items():
            curso["parciales"][f"parcial_{pn}"] = {
                "clase": cn,
                "fecha": by_n[cn]["fecha"],
                "tipo": by_n[cn]["tipo"],
            }

    path.write_text(json.dumps(cal, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"OK {path.relative_to(ROOT)}")
    return cal


def write_calendario_md(curso: dict) -> None:
    folder = ROOT / curso["folder"]
    path = folder / "CALENDARIO_2026-2.md"
    pmap = curso["parciales"]
    rows = []
    for c in curso["clases"]:
        nota = c.get("festivo") or "—"
        if c.get("parcial"):
            pn = c["parcial_n"]
            tag = f"Parcial {pn} (cierre Corte {pn})"
            nota = tag if nota == "—" else f"{nota} · {tag}"
        elif c["tipo"] == "autonoma" and c.get("festivo"):
            # autónoma de cierre teórico sin parcial
            if c["n"] in (10, 15) and curso["folder"] in (
                "Arquitectura de Sistemas Computacionales",
                "Bases de Datos II",
            ):
                nota = f"{c['festivo']} · refuerzo sin parcial"
        rows.append(
            f"| {c['n']} | {fecha_larga(c['fecha'])} | {tipo_label(c['tipo'], c.get('festivo'))} | {nota} |"
        )

    p1, p2, p3 = pmap["parcial_1"], pmap["parcial_2"], pmap["parcial_3"]
    grupo = curso.get("grupo", "")
    md = f"""# Calendario 2026-2 — {curso['nombre']}

- **Código:** {curso['codigo']}
- **Grupo:** {grupo}
- **Periodo:** 2026-2 · **10/08/2026 – 22/11/2026**
- **Horario:** {curso['dia']} **{curso['horario']}** ({curso['duracion_min']} min)
- **Modalidad:** Presencialidad asistida (clases presencial / virtual síncrona / autónoma)
- **Docente:** Julian Andres Castaño Espinosa · `julianacastano@profesores.uniajc.edu.co`
- **Total clases:** 15 (festivos = **clase autónoma**, no se omiten)

## Cortes teóricos (30% / 30% / 40%)

{CRITERIO}

| Corte | % | Ventana | Clases | Parcial de cierre | Desglose teórico |
|---|---|---|---|---|---|
| Corte 1 | 30% | 2026-08-10 → 2026-09-13 | 1-5 | Parcial 1 · Clase {p1['clase']} ({fecha_larga(p1['fecha'])}) · {tipo_label(p1['tipo'], None)} | 10% Parcial 1 · 10% Talleres y Quiz · 10% Asistencia |
| Corte 2 | 30% | 2026-09-14 → 2026-10-18 | 6-10 | Parcial 2 · Clase {p2['clase']} ({fecha_larga(p2['fecha'])}) · {tipo_label(p2['tipo'], None)} | 10% Parcial 2 · 10% Talleres y Quiz · 10% Asistencia |
| Corte 3 | 40% | 2026-10-19 → 2026-11-22 | 11-15 | Parcial 3 · Clase {p3['clase']} ({fecha_larga(p3['fecha'])}) · {tipo_label(p3['tipo'], None)} | 15% Parcial 3 · 20% Proyecto Integrador · 5% Asistencia |

> Validar en Acuerdo pedagógico / socialización. Ventanas de corte = bloques temáticos; la fecha del parcial = última clase regular del corte.

## Clases

| Clase | Fecha | Tipo | Nota |
|---|---|---|---|
{chr(10).join(rows)}

## Festivos Colombia 2026 (rango del periodo)

- 17/08/2026 — Asunción de la Virgen
- 12/10/2026 — Día de la Diversidad Étnica y Cultural
- 02/11/2026 — Todos los Santos
- 16/11/2026 — Independencia de Cartagena
"""
    path.write_text(md, encoding="utf-8")
    print(f"OK {path.relative_to(ROOT)}")


def update_plan_arq_bd(curso: dict, temas: dict[int, str]) -> None:
    folder = ROOT / curso["folder"]
    path = folder / "Plan curso" / "PLAN_DE_CURSO_2026-2.md"
    pmap = curso["parciales"]
    p1, p2, p3 = pmap["parcial_1"], pmap["parcial_2"], pmap["parcial_3"]

    rows = []
    for c in curso["clases"]:
        tema = temas[c["n"]]
        if c.get("parcial"):
            tema = f"{tema} · **Parcial {c['parcial_n']} (cierre Corte {c['parcial_n']})**"
        elif c["tipo"] == "autonoma" and c["n"] in (10, 15):
            tema = f"{tema} · refuerzo sin parcial"
        rows.append(
            f"| {c['n']} | {fecha_larga(c['fecha'])} | {tipo_label(c['tipo'], c.get('festivo'))} | {tema} |"
        )

    header_extra = ""
    if curso["folder"] == "Arquitectura de Sistemas Computacionales":
        header_extra = f"""- **Código:** FI303380
- **Grupo:** **[PENDIENTE — grupo]**
- **Periodo:** **2026-2** · 10/08/2026 – 22/11/2026
- **Horario:** **Lunes 10:00 – 13:00** (180 min)
- **Modalidad:** **Presencialidad asistida**
- **Docente:** Julian Andres Castaño Espinosa · `julianacastano@profesores.uniajc.edu.co`
- **Fuente oficial:** Microcurrículo FI303380 (Enfoque Cloud) + Plan_de_curso FI303380
- **Calendario:** `../CALENDARIO_2026-2.md` · `.config/calendario/semestre_2026_2.json`
- **Listado de estudiantes:** `[PENDIENTE listado]`

## Ajuste 16 → 15 clases

El Plan oficial trae 16 sesiones. En 2026-2 hay **15 clases**: el cierre de la sesión 16 se integra en la **Clase 15**.
"""
    else:
        header_extra = f"""- **Código:** FI303215
- **Grupo:** **641A-2**
- **Periodo:** **2026-2** · 10/08/2026 – 22/11/2026
- **Horario:** **Lunes 18:00 – 20:00** (120 min) · franja virtual síncrona
- **Modalidad:** **Presencialidad asistida**
- **Docente:** Julian Andres Castaño Espinosa · `julianacastano@profesores.uniajc.edu.co`
- **Fuente oficial:** Microcurrículo FI303215 + Plan_de_curso FI303215
- **Calendario:** `../CALENDARIO_2026-2.md` · `.config/calendario/semestre_2026_2.json`
- **Listado de estudiantes:** `[PENDIENTE listado]`

## Ajuste 16 → 15 clases

El Plan oficial trae 16 sesiones. En 2026-2 hay **15 clases**: el cierre de la sesión 16 se integra en la **Clase 15**.
"""

    md = f"""# Plan de curso propuesto — {curso['nombre']} · 2026-2

{header_extra}
{CRITERIO}

Parciales de este curso: Clases **{p1['clase']} / {p2['clase']} / {p3['clase']}** ({fecha_larga(p1['fecha'])}, {fecha_larga(p2['fecha'])}, {fecha_larga(p3['fecha'])}).

## Tabla Clase · Fecha · Tipo · Tema

| Clase | Fecha | Tipo | Tema (Trabajo dirigido) |
|---|---|---|---|
{chr(10).join(rows)}

## Evaluación teórica (Acuerdo 2026-2)

| Corte | % | Ventana | Parcial de cierre |
|---|---|---|---|
| 1 | 30% | 10/08 – 13/09/2026 | Parcial 1 en Clase {p1['clase']} ({fecha_larga(p1['fecha'])}) (10%) + Talleres/Quiz 10% + Asistencia 10% |
| 2 | 30% | 14/09 – 18/10/2026 | Parcial 2 en Clase {p2['clase']} ({fecha_larga(p2['fecha'])}) (10%) + Talleres/Quiz 10% + Asistencia 10% |
| 3 | 40% | 19/10 – 22/11/2026 | Parcial 3 en Clase {p3['clase']} ({fecha_larga(p3['fecha'])}) (15%) + Proyecto Integrador 20% + Asistencia 5% |
"""
    # preserve objetivo/RAA if present in old file
    old = path.read_text(encoding="utf-8") if path.exists() else ""
    if "## Objeto" in old or "## Objeto / objetivo" in old:
        # keep from first ## Objeto
        idx = old.find("## Objeto")
        if idx == -1:
            idx = old.find("## Objeto / objetivo")
        if idx != -1:
            md = md.rstrip() + "\n\n" + old[idx:]
    elif "# Objeto" in old:
        pass
    else:
        # append RAA sections from known stubs if file ends without them
        if curso["folder"].startswith("Arquitectura"):
            md += """
## Objeto / objetivo / RAA (microcurrículo)

- **Objeto:** Arquitecturas de sistemas computacionales en entornos cloud.
- **Objetivo:** Diseñar e implementar arquitecturas con cloud, virtualización y escalabilidad.
- **RAA1** IaaS/PaaS/SaaS · **RAA2** Virtualización y distribuidos · **RAA3** Seguridad, rendimiento y sostenibilidad.
"""
        else:
            md += """
## Objeto / objetivo / RAA (microcurrículo)

- **Objeto:** Gestión avanzada y optimización de bases de datos relacionales.
- **Objetivo:** Diseñar, administrar y optimizar bases de datos relacionales avanzadas.
- **RAA1** Seguridad y respaldo · **RAA2** Procedimientos y disparadores · **RAA3** Optimización.
"""

    path.write_text(md, encoding="utf-8")
    print(f"OK {path.relative_to(ROOT)}")


def update_plan_mie_jue(curso: dict) -> None:
    """Solo añade nota de criterio; parciales 5/10/15 ya OK."""
    path = ROOT / curso["folder"] / "Plan curso" / "PLAN_DE_CURSO_2026-2.md"
    text = path.read_text(encoding="utf-8")
    nota = (
        f"\n\n> {CRITERIO} En este curso (sin festivos en {curso['dia'].lower()}) "
        f"los parciales quedan en Clases **5 / 10 / 15** (todas regulares).\n"
    )
    if "Criterio de parciales (fijo 2026-2)" in text:
        # replace existing blockquote if any
        import re

        text = re.sub(
            r"\n> Criterio de parciales \(fijo 2026-2\):.*?\n",
            nota,
            text,
            count=1,
            flags=re.S,
        )
        if "Criterio de parciales (fijo 2026-2)" not in text:
            # insert after header block
            lines = text.splitlines()
            insert_at = 1
            for i, line in enumerate(lines[:20]):
                if line.startswith("- **"):
                    insert_at = i + 1
            lines.insert(insert_at, nota.strip("\n"))
            text = "\n".join(lines) + ("\n" if not text.endswith("\n") else "")
    else:
        lines = text.splitlines()
        insert_at = 0
        for i, line in enumerate(lines[:25]):
            if line.startswith("- **Listado") or line.startswith("- **Docente"):
                insert_at = i + 1
        if insert_at == 0:
            insert_at = min(8, len(lines))
        lines.insert(insert_at, nota.strip("\n"))
        text = "\n".join(lines)
        if not text.endswith("\n"):
            text += "\n"
        # ensure modalidad line
        if "Modalidad" not in text[:800]:
            for i, line in enumerate(lines):
                if line.startswith("- **Horario"):
                    lines.insert(i + 1, "- **Modalidad:** **Presencialidad asistida**")
                    text = "\n".join(lines)
                    if not text.endswith("\n"):
                        text += "\n"
                    break
    path.write_text(text, encoding="utf-8")
    print(f"OK {path.relative_to(ROOT)}")


def update_uniajc_json(cal: dict) -> None:
    path = ROOT / ".config" / "universidades" / "uniajc.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    data["_version"] = "1.4"
    data["_actualizado"] = "2026-08-07"
    data["_cursos_workspace"] = [
        "Programacion II (FI303204 · grupo 341C · 2026-2 · mié 18:00-20:00 · 120 min · Presencialidad asistida)",
        "Seminario de Sistemas (FI303301 · grupo 341C · 2026-2 · jue 18:00-20:00 · 120 min · Presencialidad asistida)",
        "Bases de Datos II (FI303215 · grupo 641A-2 · 2026-2 · lun 18:00-20:00 virtual síncrona · 120 min · Presencialidad asistida)",
        "Arquitectura de Sistemas Computacionales (FI303380 · 2026-2 · lun 10:00-13:00 · 180 min · Presencialidad asistida · grupo pendiente)",
        "0. Base (referencia Ingeniería de Software 1 / plantillas institucionales)",
    ]
    data["periodo_vigente"]["regla_festivos"] = (
        "No omitir. Marcar como clase autónoma. Parciales NUNCA en festivo/autónoma."
    )
    data["periodo_vigente"]["regla_parciales"] = CRITERIO
    data["periodo_vigente"]["modalidad_cursos"] = "Presencialidad asistida"
    data["pedagogia"]["_nota_duracion"] = (
        "Periodo 2026-2 (10/08/2026–22/11/2026). Todos los cursos: Presencialidad asistida "
        "(clases presencial / virtual síncrona / autónoma). Duraciones: Programación II 341C mié "
        "18:00–20:00 = 120 min; Seminario 341C jue 18:00–20:00 = 120 min; Bases de Datos II 641A-2 "
        "lun 18:00–20:00 = 120 min; Arquitectura lun 10:00–13:00 = 180 min. "
        "Festivos = clase autónoma; parciales en última regular del corte "
        "(config/calendario/semestre_2026_2.json)."
    )
    data["pedagogia"]["metodologia_observada_prog2"]["modalidad"] = (
        "Presencialidad asistida (clases presencial / virtual síncrona / autónoma)"
    )
    arq = cal["cursos"]["arquitectura"]["parciales"]
    bd = cal["cursos"]["bases_datos_ii"]["parciales"]
    prog = cal["cursos"]["programacion_ii"]["parciales"]
    data["evaluacion"]["_nota"] = (
        "Copia SIEMPRE del Acuerdo pedagógico del grupo. Teórico 2026-2: 30/30/40. " + CRITERIO
    )
    data["evaluacion"]["teorico_2026_2"] = {
        "_logica": LOGICA_CORTES + " " + CRITERIO,
        "criterio_parciales": CRITERIO,
        "corte_1_30": {
            "ventana": "10/08/2026 al 13/09/2026",
            "clases": "1-5",
            "parcial_cierre": "Parcial 1 en la última clase regular del corte",
            "desglose": {
                "parcial_1": "10%",
                "talleres_y_quiz": "10%",
                "asistencia": "10%",
            },
        },
        "corte_2_30": {
            "ventana": "14/09/2026 al 18/10/2026",
            "clases": "6-10",
            "parcial_cierre": "Parcial 2 en la última clase regular del corte (lun: Clase 9 si Clase 10 es autónoma)",
            "desglose": {
                "parcial_2": "10%",
                "talleres_y_quiz": "10%",
                "asistencia": "10%",
            },
        },
        "corte_3_40": {
            "ventana": "19/10/2026 al 22/11/2026",
            "clases": "11-15",
            "parcial_cierre": "Parcial 3 en la última clase regular del corte (lun: Clase 14 si Clase 15 es autónoma)",
            "desglose": {
                "parcial_3": "15%",
                "proyecto_integrador": "20%",
                "asistencia": "5%",
            },
        },
        "por_curso": {
            "programacion_ii": {
                "parcial_1": f"Clase {prog['parcial_1']['clase']}",
                "parcial_2": f"Clase {prog['parcial_2']['clase']}",
                "parcial_3": f"Clase {prog['parcial_3']['clase']}",
            },
            "seminario": {
                "parcial_1": "Clase 5",
                "parcial_2": "Clase 10",
                "parcial_3": "Clase 15",
            },
            "bases_datos_ii": {
                "parcial_1": f"Clase {bd['parcial_1']['clase']} ({bd['parcial_1']['fecha']})",
                "parcial_2": f"Clase {bd['parcial_2']['clase']} ({bd['parcial_2']['fecha']})",
                "parcial_3": f"Clase {bd['parcial_3']['clase']} ({bd['parcial_3']['fecha']})",
            },
            "arquitectura": {
                "parcial_1": f"Clase {arq['parcial_1']['clase']} ({arq['parcial_1']['fecha']})",
                "parcial_2": f"Clase {arq['parcial_2']['clase']} ({arq['parcial_2']['fecha']})",
                "parcial_3": f"Clase {arq['parcial_3']['clase']} ({arq['parcial_3']['fecha']})",
            },
        },
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"OK {path.relative_to(ROOT)}")


def patch_rule_and_agents() -> None:
    replacements = [
        (
            ROOT / ".cursor" / "rules" / "uniajc-docente.mdc",
            [
                (
                    "- **Regla de festivos:** no omitir clases; marcarlas como **clase autónoma**.\n"
                    "- Festivos en rango: **17/08**, **12/10**, **02/11**, **16/11** de 2026.\n"
                    "- Cortes teóricos (validar en Acuerdo): **30% / 30% / 40%** — misma lógica Prog. II / Seminario\n"
                    "  - Corte 1: 10/08–13/09 · clases 1–5 · **Parcial 1 al cierre (Clase 5)** · 10% Parcial + 10% Talleres/Quiz + 10% Asistencia\n"
                    "  - Corte 2: 14/09–18/10 · clases 6–10 · **Parcial 2 al cierre (Clase 10)** · 10% Parcial + 10% Talleres/Quiz + 10% Asistencia\n"
                    "  - Corte 3: 19/10–22/11 · clases 11–15 · **Parcial 3 al cierre (Clase 15)** · 15% Parcial + 20% Proyecto Integrador + 5% Asistencia\n"
                    "  - Si la clase de cierre cae en festivo → parcial en esa **clase autónoma** / Campus Virtual\n",
                    "- **Modalidad (todos los cursos):** **Presencialidad asistida** (clases presencial / virtual síncrona / autónoma).\n"
                    "- **Regla de festivos:** no omitir clases; marcarlas como **clase autónoma**.\n"
                    "- **Regla de parciales:** siempre síncronos; **NUNCA** en festivo ni clase autónoma. Si el cierre del corte es autónoma → parcial en la **última clase regular anterior** del corte.\n"
                    "- Festivos en rango: **17/08**, **12/10**, **02/11**, **16/11** de 2026.\n"
                    "- Cortes teóricos (validar en Acuerdo): **30% / 30% / 40%**\n"
                    "  - Corte 1: 10/08–13/09 · clases 1–5 · Parcial 1 en última regular del corte (mié/jue: Clase 5; lun: Clase 5)\n"
                    "  - Corte 2: 14/09–18/10 · clases 6–10 · Parcial 2 en última regular (mié/jue: Clase 10; lun: **Clase 9**, no 10 festivo)\n"
                    "  - Corte 3: 19/10–22/11 · clases 11–15 · Parcial 3 en última regular (mié/jue: Clase 15; lun: **Clase 14**, no 15 festivo)\n"
                    "  - Desglose: C1/C2 = 10% Parcial + 10% Talleres/Quiz + 10% Asistencia; C3 = 15% Parcial + 20% PI + 5% Asistencia\n",
                ),
                (
                    "3. **Bases de Datos II** (FI303215 · grupo **641A-2** · lun **18:00–20:00** · **120 min** · **virtual**)\n"
                    "4. **Arquitectura de Sistemas Computacionales** (FI303380 · lun **10:00–13:00** · **180 min** · grupo/modalidad pendientes)\n",
                    "3. **Bases de Datos II** (FI303215 · grupo **641A-2** · lun **18:00–20:00** · **120 min** · **Presencialidad asistida** · franja virtual síncrona)\n"
                    "4. **Arquitectura de Sistemas Computacionales** (FI303380 · lun **10:00–13:00** · **180 min** · **Presencialidad asistida** · grupo pendiente)\n",
                ),
            ],
        ),
    ]

    agent_old = (
        "8. Evaluación teórica: 30/30/40 con **Parcial al cierre de cada corte** (Clases 5, 10, 15). "
        "Festivos = clase autónoma."
    )
    agent_new = (
        "8. Evaluación teórica: 30/30/40. Parciales **siempre síncronos**; **NUNCA** en festivo/autónoma. "
        "Criterio: última clase regular del corte (mié/jue: 5/10/15; lun Arq/BD II: **5/9/14**). "
        "Festivos = clase autónoma de refuerzo. Modalidad de todos los cursos: **Presencialidad asistida**."
    )
    agent_paso0_old = (
        "4. Lee también `.config/calendario/semestre_2026_2.json` (clases, festivos=autónoma, "
        "cortes 30/30/40 con Parcial 1/2/3 al cierre)."
    )
    agent_paso0_new = (
        "4. Lee también `.config/calendario/semestre_2026_2.json` (clases presencial/virtual/autónoma, "
        "cortes 30/30/40; parciales en última regular del corte — nunca en festivo)."
    )
    footer_old = (
        "*v1.2 — Variante UNIAJC · 2026-2 · correo profesores.uniajc · 15 clases · "
        "Parcial al cierre de cada corte · Festivos=autónoma · Motor `uniajc_slides_engine.py`.*"
    )
    footer_new = (
        "*v1.3 — Variante UNIAJC · 2026-2 · Presencialidad asistida · parciales síncronos "
        "(nunca en autónoma) · última regular del corte · Motor `uniajc_slides_engine.py`.*"
    )

    for path, pairs in replacements:
        text = path.read_text(encoding="utf-8")
        for old, new in pairs:
            if old not in text:
                raise SystemExit(f"No match in {path}: {old[:60]}...")
            text = text.replace(old, new)
        # also fix evaluation bullet later in file if present
        text = text.replace(
            "- Evaluación teórica: 30/30/40 con **Parcial al cierre de cada corte** (Clases 5, 10, 15). Festivos = clase autónoma.",
            "- Evaluación teórica: 30/30/40; parciales síncronos en última regular del corte (nunca festivo/autónoma).",
        )
        path.write_text(text, encoding="utf-8")
        print(f"OK {path.relative_to(ROOT)}")

    for agent in (
        ROOT / ".cursor" / "agents" / "disenador-curricular-uniajc.md",
        ROOT / ".claude" / "agents" / "disenador-curricular-uniajc.md",
    ):
        text = agent.read_text(encoding="utf-8")
        text = text.replace(agent_paso0_old, agent_paso0_new)
        text = text.replace(agent_old, agent_new)
        text = text.replace(footer_old, footer_new)
        # modalidad note in pendientes
        text = text.replace(
            "grupo/modalidad Arquitectura",
            "grupo Arquitectura (modalidad = Presencialidad asistida)",
        )
        agent.write_text(text, encoding="utf-8")
        print(f"OK {agent.relative_to(ROOT)}")


def update_notas() -> None:
    notes = {
        ROOT
        / "Arquitectura de Sistemas Computacionales"
        / "Entregas docente"
        / "NOTA - Acuerdo y listado 2026-2.md": (
            "# Pendientes Acuerdo pedagógico — Arquitectura de Sistemas Computacionales · 2026-2\n\n"
            "- Prefill del Acuerdo: generar desde plantilla `Programacion II/Entregas docente/ACUERDO PEDAGOGICO ACTUALIZADO.docx` cuando se socialice.\n"
            "- Listado de estudiantes / vocero / firmas: **[PENDIENTE listado]**\n"
            "- Modalidad: **Presencialidad asistida** · Grupo: **[PENDIENTE]**\n"
            "- Evaluación teórica: 30/30/40 · Parciales en Clases **5 / 9 / 14** (07/09, 05/10, 09/11) — nunca en autónoma.\n"
            f"- {CRITERIO}\n"
            "- Docente: Julian Andres Castaño Espinosa · julianacastano@profesores.uniajc.edu.co\n"
        ),
        ROOT
        / "Bases de Datos II"
        / "Entregas docente"
        / "NOTA - Acuerdo y listado 2026-2.md": (
            "# Pendientes Acuerdo pedagógico — Bases de Datos II · 2026-2\n\n"
            "- Prefill del Acuerdo: generar desde plantilla `Programacion II/Entregas docente/ACUERDO PEDAGOGICO ACTUALIZADO.docx` cuando se socialice.\n"
            "- Listado de estudiantes / vocero / firmas: **[PENDIENTE listado]**\n"
            "- Modalidad: **Presencialidad asistida** (franja virtual síncrona lun 18:00–20:00) · Grupo **641A-2**\n"
            "- Evaluación teórica: 30/30/40 · Parciales en Clases **5 / 9 / 14** (07/09, 05/10, 09/11) — nunca en autónoma.\n"
            f"- {CRITERIO}\n"
            "- Docente: Julian Andres Castaño Espinosa · julianacastano@profesores.uniajc.edu.co\n"
        ),
    }
    for path, content in notes.items():
        path.write_text(content, encoding="utf-8")
        print(f"OK {path.relative_to(ROOT)}")


def set_cell_text(cell, text: str) -> None:
    # Keep first paragraph, clear rest
    if not cell.paragraphs:
        cell.text = text
        return
    cell.paragraphs[0].text = text
    for p in cell.paragraphs[1:]:
        p.text = ""


def update_acuerdos(cal: dict) -> None:
    nota_eval = (
        "Acuerdo sobre los aspectos de evaluación\n"
        "(Cálculo teórico 2026-2 — 30% / 30% / 40% — validar en socialización con el grupo)\n"
        "Parciales: siempre síncronos; NUNCA en festivo ni clase autónoma. "
        "Criterio: última clase regular (presencial/virtual síncrona) del corte.\n\n"
    )

    specs = {
        "arquitectura": {
            "glob": "Arquitectura de Sistemas Computacionales/Entregas docente/ACUERDO PEDAGOGICO - Arquitectura de Sistemas Computacionales - 2026-2.docx",
            "p15": (
                "Docente: Julian Andres Castaño Espinosa — Correo: julianacastano@profesores.uniajc.edu.co — "
                "Horario: Lunes 10:00 – 13:00 (180 min) — Modalidad: Presencialidad asistida — "
                "Código: FI303380 — PRELLENADO 2026-2 — campos de estudiantes pendientes."
            ),
            "metod": (
                "Acuerdo sobre los aspectos metodológicos\n"
                "Periodo 2026-2 — Lunes 10:00–13:00 (180 min).\n"
                "Grupo: [PENDIENTE]. Modalidad: Presencialidad asistida "
                "(clases presencial / virtual síncrona / autónoma).\n"
                "Calendario: 15 clases (10/08/2026–16/11/2026). Festivos = clase autónoma:\n"
                "17/08 (Asunción), 12/10 (Diversidad étnica), 02/11 (Todos los Santos), 16/11 (Independencia de Cartagena).\n"
                "[Detalle en CALENDARIO_2026-2.md del curso.]\n"
            ),
            "eval_body": (
                "Primer corte (30%) — [10/08/2026 al 13/09/2026] — Clases 1–5:\n"
                "* 10% Parcial 1 (Clase 5 · 07/09/2026, presencial) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Segundo corte (30%) — [14/09/2026 al 18/10/2026] — Clases 6–10:\n"
                "* 10% Parcial 2 (Clase 9 · 05/10/2026, presencial; Clase 10 festivo = refuerzo sin parcial) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Tercer corte (40%) — [19/10/2026 al 22/11/2026] — Clases 11–15:\n"
                "* 15% Parcial 3 (Clase 14 · 09/11/2026, presencial; Clase 15 festivo = refuerzo sin parcial) | 20% Proyecto Integrador | 5% Asistencia\n\n"
                f"Nota: {CRITERIO}\n"
            ),
        },
        "bases_datos_ii": {
            "glob": "Bases de Datos II/Entregas docente/ACUERDO PEDAGOGICO - Bases de Datos II - 2026-2.docx",
            "p15": (
                "Docente: Julian Andres Castaño Espinosa — Correo: julianacastano@profesores.uniajc.edu.co — "
                "Horario: Lunes 18:00 – 20:00 (120 min) — Modalidad: Presencialidad asistida (franja virtual síncrona) — "
                "Código: FI303215 — PRELLENADO 2026-2 — campos de estudiantes pendientes."
            ),
            "metod": (
                "Acuerdo sobre los aspectos metodológicos\n"
                "Periodo 2026-2 — Grupo 641A-2 — Lunes 18:00–20:00 (120 min) — "
                "Modalidad: Presencialidad asistida (franja virtual síncrona; autónomas en festivo).\n"
                "Estructura sugerida (ajustar al Acuerdo/Plan): Teoría Core — Taller Guiado — Quiz/comprobación.\n"
                "Calendario: 15 clases (10/08/2026–16/11/2026). Festivos = clase autónoma:\n"
                "17/08 (Asunción), 12/10 (Diversidad étnica), 02/11 (Todos los Santos), 16/11 (Independencia de Cartagena).\n"
                "[Detalle en CALENDARIO_2026-2.md del curso.]\n"
            ),
            "eval_body": (
                "Primer corte (30%) — [10/08/2026 al 13/09/2026] — Clases 1–5:\n"
                "* 10% Parcial 1 (Clase 5 · 07/09/2026, virtual síncrona) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Segundo corte (30%) — [14/09/2026 al 18/10/2026] — Clases 6–10:\n"
                "* 10% Parcial 2 (Clase 9 · 05/10/2026, virtual síncrona; Clase 10 festivo = refuerzo sin parcial) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Tercer corte (40%) — [19/10/2026 al 22/11/2026] — Clases 11–15:\n"
                "* 15% Parcial 3 (Clase 14 · 09/11/2026, virtual síncrona; Clase 15 festivo = refuerzo sin parcial) | 20% Proyecto Integrador | 5% Asistencia\n\n"
                f"Nota: {CRITERIO}\n"
            ),
        },
        "seminario": {
            "glob": "Seminario de Sistemas/Entregas docente/ACUERDO PEDAGOGICO - Seminario de Sistemas - 2026-2.docx",
            "p15": None,
            "metod": None,
            "eval_body": (
                "Primer corte (30%) — [10/08/2026 al 13/09/2026] — Clases 1–5:\n"
                "* 10% Parcial 1 (cierre de corte, Clase 5 · 10/09/2026) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Segundo corte (30%) — [14/09/2026 al 18/10/2026] — Clases 6–10:\n"
                "* 10% Parcial 2 (cierre de corte, Clase 10 · 15/10/2026) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Tercer corte (40%) — [19/10/2026 al 22/11/2026] — Clases 11–15:\n"
                "* 15% Parcial 3 (cierre de corte, Clase 15 · 19/11/2026) | 20% Proyecto Integrador | 5% Asistencia\n\n"
                f"Nota: {CRITERIO} En este curso no hay festivos en jueves; parciales en Clases 5/10/15.\n"
            ),
        },
        "programacion_ii": {
            "glob": None,  # resolve by glob
            "p15": None,
            "metod": None,
            "eval_body": (
                "Primer corte (30%) — [10/08/2026 al 13/09/2026] — Clases 1–5:\n"
                "* 10% Parcial 1 (cierre de corte, Clase 5 · 09/09/2026) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Segundo corte (30%) — [14/09/2026 al 18/10/2026] — Clases 6–10:\n"
                "* 10% Parcial 2 (cierre de corte, Clase 10 · 14/10/2026) | 10% Talleres y Quiz | 10% Asistencia\n\n"
                "Tercer corte (40%) — [19/10/2026 al 22/11/2026] — Clases 11–15:\n"
                "* 15% Parcial 3 (cierre de corte, Clase 15 · 18/11/2026) | 20% Proyecto Integrador | 5% Asistencia\n\n"
                f"Nota: {CRITERIO} En este curso no hay festivos en miércoles; parciales en Clases 5/10/15.\n"
            ),
        },
    }

    # resolve prog2 path
    prog_docs = list((ROOT / "Programacion II" / "Entregas docente").glob("ACUERDO PEDAGOGICO*2026-2.docx"))
    if not prog_docs:
        raise SystemExit("No se encontró Acuerdo Prog II 2026-2")
    specs["programacion_ii"]["path"] = prog_docs[0]

    for key, spec in specs.items():
        path = spec.get("path") or (ROOT / spec["glob"])
        doc = Document(str(path))
        if spec["p15"] is not None:
            doc.paragraphs[15].text = spec["p15"]
        if spec["metod"] is not None:
            set_cell_text(doc.tables[3].rows[0].cells[0], spec["metod"])
        set_cell_text(doc.tables[3].rows[1].cells[0], nota_eval + spec["eval_body"])
        doc.save(str(path))
        print(f"OK {path.relative_to(ROOT)}")


def main() -> None:
    cal = update_calendario_json()
    for curso in cal["cursos"].values():
        write_calendario_md(curso)

    temas_arq = {
        1: "Presentación del curso · Introducción a arquitecturas cloud",
        2: "Modelos de servicio: IaaS, PaaS, SaaS (Asunción de la Virgen)",
        3: "Virtualización y contenedores",
        4: "Microservicios",
        5: "Arquitecturas distribuidas",
        6: "Seguridad en la nube",
        7: "Redes y almacenamiento cloud",
        8: "Monitoreo y optimización",
        9: "Integración continua y despliegue (CI/CD)",
        10: "Costos y sostenibilidad cloud",
        11: "Avance del proyecto final",
        12: "Pruebas de rendimiento",
        13: "Escalabilidad automática (Todos los Santos)",
        14: "Preparación de presentación final",
        15: "Presentación del proyecto + cierre",
    }
    temas_bd = {
        1: "Presentación del curso · Revisión de Bases de Datos I",
        2: "Administración de bases de datos (Asunción de la Virgen)",
        3: "Procedimientos almacenados",
        4: "Funciones y disparadores",
        5: "Seguridad y respaldo",
        6: "Optimización de consultas",
        7: "Índices y particionamiento",
        8: "Tuning de bases de datos",
        9: "Gestión de transacciones",
        10: "Control de concurrencia",
        11: "Avance del proyecto final",
        12: "Integración de aplicaciones externas",
        13: "Análisis de casos reales (Todos los Santos)",
        14: "Preparación de presentación final",
        15: "Presentación del proyecto + cierre",
    }
    update_plan_arq_bd(cal["cursos"]["arquitectura"], temas_arq)
    update_plan_arq_bd(cal["cursos"]["bases_datos_ii"], temas_bd)
    update_plan_mie_jue(cal["cursos"]["programacion_ii"])
    update_plan_mie_jue(cal["cursos"]["seminario"])
    update_uniajc_json(cal)
    patch_rule_and_agents()
    update_notas()
    update_acuerdos(cal)
    print("DONE base updates")


if __name__ == "__main__":
    main()
